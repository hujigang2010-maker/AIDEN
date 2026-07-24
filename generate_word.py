# -*- coding: utf-8 -*-
"""生成《赴宁波港及宁波经济技术开发区考察交流策划案》Word 文档。"""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import content as C

OUT_DIR = os.path.join(os.path.dirname(__file__), "output")
CN_FONT = "微软雅黑"
ACCENT = RGBColor(0x0B, 0x3D, 0x5C)
TEAL = RGBColor(0x1A, 0x6B, 0x6B)
DARK = RGBColor(0x22, 0x2A, 0x33)
GREY = RGBColor(0x5A, 0x63, 0x6E)


def set_cn_font(run, size=None, bold=None, color=None, font=CN_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def para(doc, text, size=12, bold=False, color=None, align=None,
         space_after=6, first_line_indent=None, space_before=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=bold, color=color)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(first_line_indent)
    return p


def heading(doc, text, level=1):
    sizes = {1: 16, 2: 13.5, 3: 12}
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, size=sizes[level], bold=True, color=ACCENT)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    return p


def bullet(doc, text, size=11.5):
    p = doc.add_paragraph(style="List Bullet")
    # 清空默认 run，自行写入以控制字体
    if p.runs:
        p.runs[0].text = ""
    r = p.add_run(text)
    set_cn_font(r, size=size, color=DARK)
    p.paragraph_format.space_after = Pt(3)
    return p


def style_table(table, header_row=True):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_cn_font(r, size=10.5, bold=(header_row and i == 0),
                                color=RGBColor(0xFF, 0xFF, 0xFF) if (header_row and i == 0) else DARK)
            if header_row and i == 0:
                shading = cell._tc.get_or_add_tcPr()
                el = shading.makeelement(qn("w:shd"), {
                    qn("w:val"): "clear", qn("w:fill"): "0B3D5C"})
                shading.append(el)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    for j, t in enumerate(headers):
        table.rows[0].cells[j].text = t
    for row_data in rows:
        cells = table.add_row().cells
        for j, v in enumerate(row_data):
            cells[j].text = str(v)
    style_table(table)
    return table


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.6)

    # ---------------- 封面 ----------------
    for _ in range(4):
        doc.add_paragraph()
    para(doc, "考察交流策划案", size=14, bold=True, color=TEAL,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    para(doc, C.PROJECT_TITLE, size=22, bold=True, color=ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    para(doc, C.PROJECT_SUBTITLE, size=13, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28)
    para(doc, C.PROJECT_TAG + "  ·  " + C.VERSION, size=11, color=TEAL,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

    for _ in range(3):
        doc.add_paragraph()

    para(doc, C.ORGANIZER_LINE, size=12.5, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    para(doc, "提交对象：" + C.RECIPIENT, size=12.5, bold=True, color=ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    para(doc, C.PARTICIPANT_LINE, size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    para(doc, C.DOC_DATE, size=12, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    doc.add_page_break()

    # ---------------- 一、背景与目的 ----------------
    heading(doc, "一、项目背景与目的", 1)
    heading(doc, "1.1 背景", 2)
    for b in C.BACKGROUND:
        para(doc, b, size=11.5, first_line_indent=22, space_after=8)

    heading(doc, "1.2 考察目的", 2)
    for i, p in enumerate(C.PURPOSE, 1):
        bullet(doc, f"{p}")

    # ---------------- 二、组织与接待 ----------------
    heading(doc, "二、组织方与接待单位", 1)
    para(doc,
         "本次考察由复旦大学住房政策研究中心与上海市杨浦区科技企业联合会联合组织，"
         "恳请宁波经济技术开发区投资促进局作为接待单位予以统筹支持。",
         size=11.5, first_line_indent=22, space_after=10)

    for party in C.PARTIES:
        para(doc, f"{party['role']}", size=11, bold=True, color=TEAL, space_after=2)
        para(doc, party["name"], size=12, bold=True, color=ACCENT, space_after=4)
        for d in party["duties"]:
            bullet(doc, d, size=11)

    # ---------------- 三、人员构成 ----------------
    heading(doc, "三、考察人员构成", 1)
    para(doc, C.SCALE, size=11.5, bold=True, space_after=6)
    para(doc, C.TRAVEL_NOTE, size=11, color=GREY, space_after=10)

    add_table(doc,
              ["类别", "说明"],
              [(a, b) for a, b in C.PARTICIPANT_PROFILE])
    para(doc, "", size=6)

    # ---------------- 四、合作方向 ----------------
    heading(doc, "四、考察主题与长期合作方向", 1)
    para(doc,
         "本次交流不仅聚焦单次参访，更着眼于围绕以下事项建立长期合作与投洽通道：",
         size=11.5, first_line_indent=22, space_after=8)

    for i, theme in enumerate(C.COOP_THEMES, 1):
        heading(doc, f"4.{i} {theme['name']}", 2)
        for pt in theme["points"]:
            bullet(doc, pt, size=11)

    # ---------------- 四点五、宁波港资源筛选 ----------------
    doc.add_page_break()
    heading(doc, "五、宁波港及周边参访资源筛选", 1)
    para(doc, C.SCREEN_INTRO, size=11.5, first_line_indent=22, space_after=8)
    para(doc, C.VALUE_CHAIN, size=11.5, bold=True, color=TEAL, space_after=10)

    heading(doc, "5.1 筛选原则", 2)
    add_table(doc,
              ["原则", "说明"],
              list(C.SCREEN_CRITERIA))
    para(doc, "", size=6)

    heading(doc, "5.2 港口线精选点位", 2)
    for site in C.PORT_SITES:
        para(doc, f"【{site['tier']}】{site['name']}", size=12, bold=True,
             color=ACCENT, space_after=4)
        bullet(doc, f"为何值得去：{site['why']}", size=11)
        bullet(doc, f"看什么：{site['see']}", size=11)
        bullet(doc, f"现场互动：{site['interact']}", size=11)
        bullet(doc, f"与沪方链接：{site['link']}", size=11)

    heading(doc, "5.3 产业与孵化线精选点位", 2)
    for site in C.INDUSTRY_SITES:
        para(doc, f"【{site['tier']}】{site['name']}", size=12, bold=True,
             color=ACCENT, space_after=4)
        bullet(doc, f"为何值得去：{site['why']}", size=11)
        bullet(doc, f"看什么：{site['see']}", size=11)
        bullet(doc, f"现场互动：{site['interact']}", size=11)
        bullet(doc, f"与沪方链接：{site['link']}", size=11)

    heading(doc, "5.4 互动链接设计（把参访变成沟通与转化）", 2)
    for link in C.INTERACTION_LINKS:
        para(doc, link["name"], size=12, bold=True, color=TEAL, space_after=2)
        para(doc, f"嵌入节点：{link['anchor']}　｜　形式：{link['form']}",
             size=11, color=GREY, space_after=4)
        for a in link["agenda"]:
            bullet(doc, a, size=11)

    # ---------------- 六、时间安排 ----------------
    doc.add_page_break()
    heading(doc, "六、时间安排建议", 1)
    para(doc, C.TIME_WINDOW, size=12, bold=True, color=ACCENT, space_after=8)
    para(doc, C.DATE_CONFIRM, size=11, color=GREY, space_after=10)

    heading(doc, "6.1 行程深度选项", 2)
    add_table(doc,
              ["方案", "时长", "适用情形"],
              list(C.TIME_OPTIONS))
    para(doc, "", size=6)

    heading(doc, "6.2 建议窗口", 2)
    for m in C.PREFERRED_MONTHS:
        bullet(doc, m, size=11)

    # ---------------- 七、行程方案 ----------------
    heading(doc, "七、建议行程方案（已嵌入精选点位与互动链接）", 1)
    para(doc,
         "以下行程已嵌入第五节筛选点位与四大互动链接，具体开放范围可由双方按接待条件微调。"
         "推荐采用「方案 B · 两日深度行」。",
         size=11.5, first_line_indent=22, space_after=10)

    heading(doc, "7.1 第一日（精华行程，一日行至此结束）", 2)
    add_table(doc,
              ["时间", "事项", "说明"],
              list(C.DAY1_ITINERARY))
    para(doc, "", size=6)

    heading(doc, "7.2 第二日（两日行加深对接，推荐）", 2)
    add_table(doc,
              ["时间", "事项", "说明"],
              list(C.DAY2_ITINERARY))
    para(doc, "", size=6)

    heading(doc, "7.3 参访资源总表（请贵局协助勾选协调）", 2)
    add_table(doc,
              ["建议点位", "定位与重点"],
              list(C.SUGGESTED_SITES))
    para(doc, "", size=6)

    # ---------------- 八、预期成果 ----------------
    heading(doc, "八、预期成果与长效合作", 1)
    heading(doc, "8.1 本次预期成果", 2)
    for o in C.OUTCOMES:
        bullet(doc, o, size=11)

    heading(doc, "8.2 长效合作设想", 2)
    for name, desc in C.LONG_TERM:
        bullet(doc, f"{name}：{desc}", size=11)

    # ---------------- 九、恳请支持 ----------------
    heading(doc, "九、恳请接待单位支持事项", 1)
    para(doc,
         "为保障考察交流顺利、高效、务实，恳请贵局在以下方面给予支持与指导：",
         size=11.5, first_line_indent=22, space_after=8)
    add_table(doc,
              ["支持事项", "具体说明"],
              list(C.SUPPORT_REQUESTS))
    para(doc, "", size=6)

    # ---------------- 十、费用原则 ----------------
    heading(doc, "十、费用与分工原则", 1)
    for c in C.COST_PRINCIPLES:
        bullet(doc, c, size=11)

    # ---------------- 十一、下一步 ----------------
    heading(doc, "十一、下一步工作安排", 1)
    add_table(doc,
              ["步骤", "工作内容"],
              list(C.NEXT_STEPS))
    para(doc, "", size=6)
    para(doc, C.CONTACT_NOTE, size=11, color=GREY, space_after=14)

    para(doc, "结语", size=13, bold=True, color=ACCENT, space_after=8)
    para(doc, C.CLOSING, size=11.5, first_line_indent=22, space_after=16)

    para(doc, "此致", size=12, space_after=4)
    para(doc, "敬礼", size=12, space_after=20)
    para(doc, "复旦大学住房政策研究中心", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)
    para(doc, "上海市杨浦区科技企业联合会", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)
    para(doc, C.DOC_DATE, size=11, color=GREY,
         align=WD_ALIGN_PARAGRAPH.RIGHT)

    os.makedirs(OUT_DIR, exist_ok=True)
    out = os.path.join(
        OUT_DIR,
        "赴宁波港及宁波经济技术开发区考察交流策划案.docx",
    )
    doc.save(out)
    print("saved:", out)


if __name__ == "__main__":
    main()
