# -*- coding: utf-8 -*-
"""
生成《沪牌额度使用及车辆管理协议》两版 Word 文件：
  1) 沪牌额度使用及车辆管理协议（乙方提议版）.docx
     —— 完全按乙方曹老师 2026-06 提出的两条修改原话写入，便于胡先生看清楚原话的法律影响。
  2) 沪牌额度使用及车辆管理协议（甲方建议版）.docx
     —— 站在甲方（胡先生）视角，对乙方两条修改做合理限缩 / 反提，并将
        附件 PDF 中的填空全部补齐、乙方电话更新为 18734434007。

执行：python3 generate_agreement.py
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


CN_FONT = "宋体"
EN_FONT = "Times New Roman"
HEADING_FONT = "黑体"


# ============================================================
# Word 样式工具
# ============================================================
def set_run_font(run, size_pt=12, bold=False, cn_font=CN_FONT, en_font=EN_FONT):
    run.font.name = en_font
    run.font.size = Pt(size_pt)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), cn_font)
    rfonts.set(qn("w:ascii"), en_font)
    rfonts.set(qn("w:hAnsi"), en_font)


def set_paragraph_format(p, first_line_indent=False, line_spacing=1.5, space_after=4, align=None):
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    if first_line_indent:
        pf.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align


def add_title(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5, space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_run_font(run, size_pt=18, bold=True, cn_font=HEADING_FONT)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.3, space_after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_run_font(run, size_pt=11, bold=False, cn_font=HEADING_FONT)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    set_run_font(run, size_pt=13, bold=True, cn_font=HEADING_FONT)


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    run = p.add_run(text)
    set_run_font(run, size_pt=12, bold=bold)
    return p


def add_info_line(doc, label, value, label_bold=True):
    """姓名 / 身份证号 / 联系电话等信息行。"""
    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=2)
    r1 = p.add_run(label)
    set_run_font(r1, size_pt=12, bold=label_bold)
    r2 = p.add_run(value if value else "_" * 28)
    set_run_font(r2, size_pt=12)


def add_signature_block(doc, jia_id, yi_id, sign_date):
    table = doc.add_table(rows=3, cols=2)
    table.autofit = True
    cells = [
        ("甲方（签字 / 电子签）：________________", "乙方（签字 / 电子签）：________________"),
        (f"身份证号:{jia_id}", f"身份证号:{yi_id}"),
        (f"日期：{sign_date}", f"日期：{sign_date}"),
    ]
    for row_idx, (left, right) in enumerate(cells):
        for col_idx, txt in enumerate((left, right)):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_format(p, space_after=6)
            run = p.add_run(txt)
            set_run_font(run, size_pt=12)


# ============================================================
# 协议主体（共用条款）
# ============================================================
INFO = {
    "jia_name":  "胡继刚",
    "jia_id":    "220283198704250614",
    "jia_phone": "13262607888",
    "jia_addr":  "上海市杨浦区爱国路 389 号",

    "yi_name":   "曹成林",
    "yi_id":     "142421197106030012",
    "yi_phone":  "18734434007",  # 按本次要求由 18635485488 改为 18734434007
    "yi_addr":   "山西省榆社县东升东街南一巷 1 号",

    "car_model": "别克牌 SGM6521UBA4",
    "car_plate": "沪 EAV853",
    "car_vin":   "LSGUA83L8NG060792",
    "car_engine":"222692021",

    "start_date":"2026 年 07 月 16 日",
    "end_date":  "2029 年 07 月 15 日",

    "bank_user": "胡继刚",
    "bank_name": "上海市杨浦区创智天地支行",
    "bank_acct": "6214 8551 9096 1892",
    "bank_wx":   "18678408669",

    "sign_date": "2026 年 06 月 10 日",
}


def add_parties_and_intro(doc):
    add_section_heading(doc, "甲方（沪牌额度登记权利人）")
    add_info_line(doc, "姓    名：", INFO["jia_name"])
    add_info_line(doc, "身份证号：", INFO["jia_id"])
    add_info_line(doc, "联系电话：", INFO["jia_phone"])
    add_info_line(doc, "联系地址：", INFO["jia_addr"])

    add_section_heading(doc, "乙方（车辆实际出资及使用人）")
    add_info_line(doc, "姓    名：", INFO["yi_name"])
    add_info_line(doc, "身份证号：", INFO["yi_id"])
    add_info_line(doc, "联系电话：", INFO["yi_phone"])
    add_info_line(doc, "联系地址：", INFO["yi_addr"])

    add_body(
        doc,
        "鉴于甲方系沪牌额度及车辆登记相关权利人，乙方为该车辆的实际出资人、实际使用人，"
        "双方前期已合作三年，合作期间乙方用车情况良好，未发生交通违章，且每年均为车辆投保充足商业险种。"
        "现双方经友好协商，本着平等、自愿、诚实信用的原则，就沪牌额度继续使用及车辆管理事宜，"
        "达成如下协议，以资共同遵守：",
    )


def add_clause_1(doc):
    add_section_heading(doc, "第一条  车辆及沪牌基本信息")
    for t in [
        f"1. 车辆品牌 / 型号：{INFO['car_model']}",
        f"2. 车牌号码:{INFO['car_plate']}",
        f"3. 车辆识别代码（VIN）:{INFO['car_vin']}",
        f"4. 发动机号:{INFO['car_engine']}",
        "5. 车辆登记车主：以《机动车登记证书》《行驶证》记载为准。",
        "6. 双方确认：上述车辆由乙方全额出资购买，并由乙方实际占有、使用、管理；"
        "车辆使用过程中产生的费用、风险与责任由乙方承担。本协议系双方关于沪牌额度使用与车辆代为登记的"
        "内部约定，不构成《上海市非营业性客车额度证明》的买卖或对外转让。",
    ]:
        add_body(doc, t)


def add_clause_2(doc):
    add_section_heading(doc, "第二条  使用期限")
    for t in [
        f"1. 本次续约期限为叁年，自 {INFO['start_date']} 起至 {INFO['end_date']} 止。",
        "2. 叁年期满后，如双方均同意继续合作，可再续签贰年。续签贰年的使用费暂按每年人民币 7,000 元计算，"
        "即两年合计人民币 14,000 元；最终费用与条件以届时双方另行签署的书面或电子续约协议为准。",
        "3. 乙方应在本协议期满前 30 日内与甲方书面或微信确认是否续约；未达成续约协议的，"
        "乙方应按本协议约定配合办理退牌、过户、车辆处置或其他必要手续。",
    ]:
        add_body(doc, t)


def add_clause_3(doc):
    add_section_heading(doc, "第三条  使用费及支付方式")
    add_body(doc, "1. 本协议项下叁年期使用费合计为人民币 21,000 元（大写：贰万壹仟元整）。")
    add_body(doc, "2. 乙方应于本协议签署之日起 3 日内，将上述款项一次性支付至甲方指定账户。")
    add_body(doc, "3. 甲方收款信息：")
    add_body(doc, f"    户  名:{INFO['bank_user']}")
    add_body(doc, f"    开户行:{INFO['bank_name']}")
    add_body(doc, f"    账  号:{INFO['bank_acct']}")
    add_body(doc, f"    微信 / 支付宝:{INFO['bank_wx']}")
    add_body(doc, "4. 甲方收到款项后应及时回复确认。双方通过微信、短信或电子签平台留存的转账记录、收款回执，"
                  "均可作为付款及履约凭证。")


def add_clause_4(doc):
    add_section_heading(doc, "第四条  违章备用金")
    for t in [
        "1. 鉴于双方前三年合作期间车辆零违章，本次续约甲方暂不向乙方收取车辆违章备用金。",
        "2. 协议期内如发生下列任一情形，甲方有权书面（含微信）通知乙方，在 3 日内补缴人民币 5,000 元"
        "作为违章及风险备用金：",
        "    （1）乙方发生交通违章后，未在甲方通知或系统提示后 15 日内处理完毕；",
        "    （2）因乙方原因导致车辆年检、保险续保或违章处理受到影响；",
        "    （3）车辆发生重大交通事故、行政处罚，或因乙方原因被司法机关查封、扣押；",
        "    （4）乙方未按本协议约定足额、按期投保；",
        "    （5）出现其他可能导致甲方承担费用、责任或个人征信受损的情形。",
        "3. 备用金不计利息。协议期满且车辆相关违章、事故、罚款、欠费及其他责任全部结清后，"
        "甲方应在 15 日内将剩余备用金无息退还乙方。",
    ]:
        add_body(doc, t)


def add_clause_5(doc):
    add_section_heading(doc, "第五条  保险约定")
    for t in [
        "1. 协议期内车辆保险由乙方负责购买，保险费用由乙方全额承担，且应保证保险连续有效，不得脱保。",
        "2. 每年续保时，乙方应购买交强险及商业险，其中第三者责任险保额不得低于人民币 300 万元。",
        "3. 建议乙方同时购买车辆损失险、车上人员责任险（含驾乘人员）、医保外用药责任险等必要险种，"
        "维持与现有投保水平相当或更高的保障。",
        "4. 乙方应在每年保险到期前完成续保，并在新保单生效后 7 日内将电子保单或扫描件发送甲方备案。",
        "5. 因乙方未及时投保、保额不足、脱保或拒绝配合理赔造成的全部损失，由乙方承担；"
        "如因此导致甲方被追偿或承担连带责任，乙方应足额赔偿。",
    ]:
        add_body(doc, t)


def add_clause_6(doc):
    add_section_heading(doc, "第六条  车辆使用及管理责任")
    for t in [
        "1. 乙方应合法、安全使用车辆，并保证车辆实际驾驶人持有合法、有效的机动车驾驶证。",
        "2. 协议期内，车辆使用过程中产生的油费 / 电费、停车费、通行费、保养费、维修费、年检费、"
        "保险费、违章罚款、扣分处理、事故赔偿、律师费、诉讼费、执行费等一切费用，均由乙方承担。",
        "3. 乙方不得将车辆用于违法犯罪活动、非法营运、网约车 / 出租等营运用途，"
        "亦不得将车辆用于抵押、质押、担保、转租、转借给无证人员驾驶或其他可能损害甲方权益的行为。",
        "4. 未经甲方书面同意，乙方不得擅自办理车辆转让、抵押、变更登记、补领牌证、注销、报废等手续；"
        "同样，未经乙方书面同意，甲方亦不得以该车辆设定任何抵押、担保或将车辆变卖、过户。",
        "5. 机动车登记证书、沪牌额度相关凭证由甲方保管；行驶证、保险单等日常用车所需资料"
        "可由乙方保管或随车携带。",
    ]:
        add_body(doc, t)


def add_clause_7(doc):
    add_section_heading(doc, "第七条  违章、事故及年检处理")
    for t in [
        "1. 乙方应主动查询并及时处理车辆违章。凡因乙方使用车辆产生的违章、罚款、扣分、滞纳金等，"
        "由乙方承担，乙方应在甲方通知或系统提示后 15 日内处理完毕，最晚不得迟于车辆年检之日。",
        "2. 如违章处理、事故理赔、年检或其他事项需要甲方配合，乙方应提前通知甲方，"
        "并承担因此产生的合理交通、误工等费用。",
        "3. 如因乙方未及时处理违章、事故或年检事项，导致甲方被行政处罚、个人征信受损、被诉讼、"
        "被强制执行、车辆被扣押或产生其他损失，乙方应承担全部赔偿责任。",
        "4. 发生交通事故后，乙方应立即依法报警、报险，并及时通知甲方，"
        "不得隐瞒、拖延或私自作出可能损害甲方权益的处理。",
    ]:
        add_body(doc, t)


# ============================================================
# 第八条 / 第十二条：两个版本不同
# ============================================================
def add_clause_8(doc, version):
    add_section_heading(doc, "第八条  甲方义务")
    add_body(doc, "1. 甲方应在协议期内按约定配合乙方正常使用车辆，及配合办理必要的保险、年检、违章处理等手续。")
    add_body(doc, "2. 在乙方无违约情形下，甲方不得无故提前收回沪牌额度或要求乙方停止使用车辆。")
    add_body(doc,
        "3. 如因甲方个人原因需提前终止合作的，应提前 30 日书面通知乙方，并按未使用期间比例无息退还"
        "相应使用费；因甲方违约给乙方造成损失的，应另行赔偿。但因法律法规调整、行政机关或司法机关要求、"
        "或乙方违约导致无法继续合作的，不属于本款约定的甲方违约情形。")

    if version == "yi":
        add_body(doc,
            "4. 租赁期间如因甲方原因被司法机关扣押查封时，由甲方承担一切责任（包括给乙方造成的一切全部损失）。")
    else:
        add_body(doc,
            "4. 协议期内，若完全因甲方自身原因（包括但不限于甲方个人债务纠纷、对外担保、"
            "民事 / 行政 / 刑事案件、个人征信问题、甲方擅自将车辆设定抵押 / 担保等）导致案涉车辆被司法机关"
            "查封、扣押或限制处分的，由甲方负责自费协调解除，并应在被通知后 30 日内消除该等查封 / 扣押。")
        add_body(doc,
            "5. 因前款情形给乙方造成的直接损失，由甲方承担赔偿责任。赔偿范围限于：")
        add_body(doc, "    （1）查封 / 扣押期间合理的同等档次替代用车费用（凭票据据实结算，最长不超过 30 日）；")
        add_body(doc, "    （2）已支付但因查封 / 扣押无法继续使用部分的使用费（按未使用月份比例无息退还）；")
        add_body(doc, "    （3）当年已缴保险费按未使用月份的合理摊销部分。")
        add_body(doc,
            "上述赔偿不包含乙方的间接损失、可得利益损失、商业机会损失、惩罚性赔偿及精神损害赔偿。")
        add_body(doc,
            "6. 如查封 / 扣押系因乙方使用车辆产生的违章、事故、营运、违法犯罪、抵押 / 处置等原因，"
            "或系因法律法规调整、行政机关 / 司法机关一般性政策要求、不可抗力所致，"
            "不属于本协议第八条第 4、5 款约定的甲方责任范围。")
        add_body(doc,
            "7. 甲方未在前述 30 日内消除查封 / 扣押的，乙方有权书面通知甲方解除本协议，"
            "甲方应在收到通知后 15 日内将剩余未使用月份的使用费无息退还乙方。")


def add_clause_9(doc):
    add_section_heading(doc, "第九条  提前解除及期满处理")
    for t in [
        "1. 乙方有下列情形之一的，甲方有权提前解除本协议：",
        "    （1）逾期支付使用费超过 7 日；",
        "    （2）车辆脱保、未按约定足额投保，或拒绝按本协议第四条补缴备用金；",
        "    （3）逾期处理违章、事故或年检事项；",
        "    （4）擅自转租、转借、抵押、出售或处置车辆；",
        "    （5）车辆因乙方原因被查封、扣押，或涉及重大事故、违法犯罪行为；",
        "    （6）其他严重损害甲方权益的行为。",
        "2. 因乙方违约导致协议解除的，已支付费用不予退还；如甲方另有损失，乙方应继续赔偿。",
        "3. 协议期满或协议解除后，乙方应在 15 日内配合甲方办理退牌、过户、车辆转移、注销、处置或其他"
        "必要手续，使甲方恢复对沪牌额度的完整控制和使用。",
        "4. 乙方逾期不配合的，每逾期一日，应向甲方支付人民币 200 元违约金；"
        "上述违约金不足以弥补甲方实际损失的，乙方应继续就差额部分予以赔偿。",
    ]:
        add_body(doc, t)


def add_clause_10(doc):
    add_section_heading(doc, "第十条  政策变化与不可抗力")
    add_body(doc,
        "如因法律法规、上海市沪牌管理政策、车辆登记政策、行政机关要求或不可抗力，导致本协议无法继续履行的，"
        "双方应及时协商处理。已支付费用按实际使用期间结算；任何一方另有过错的，由过错方承担相应责任。")


def add_clause_11(doc):
    add_section_heading(doc, "第十一条  通知及电子签署")
    for t in [
        "1. 双方确认，本协议载明的电话、微信、短信、电子邮箱等均可作为有效联系方式与通知送达方式。",
        "2. 双方通过微信确认、短信确认、电子签平台签署、扫描件签署或纸质签署的本协议及补充文件，均具有同等效力。",
        "3. 如采用电子签（如腾讯电子签等微信小程序），双方应确保签署人为本人，"
        "并妥善保存完整签署记录、身份核验信息、付款凭证及关键聊天记录，以备争议时使用。",
    ]:
        add_body(doc, t)


def add_clause_12(doc, version):
    add_section_heading(doc, "第十二条  争议解决")
    if version == "yi":
        add_body(doc,
            "本协议在履行过程中如发生争议，双方应先友好协商解决；协商不成的，"
            "任一方有权向原告常驻地有管辖权的人民法院提起诉讼。")
    else:
        add_body(doc,
            "本协议在履行过程中如发生争议，双方应先友好协商解决；协商不成的，"
            "任一方有权向车辆登记地或被告住所地有管辖权的人民法院提起诉讼。")


def add_clause_13(doc):
    add_section_heading(doc, "第十三条  其他")
    for t in [
        "1. 本协议未尽事宜，双方可另行签署补充协议，补充协议与本协议具有同等法律效力。",
        "2. 本协议一式两份，甲乙双方各执一份；电子签署版本与纸质版本具有同等效力。",
        "3. 本协议自双方签字（或电子签章）且甲方收到全额使用费之日起生效。",
    ]:
        add_body(doc, t)


# ============================================================
# 装配两个版本
# ============================================================
def build_document(version):
    """version: 'yi' = 乙方提议版；'jia' = 甲方建议版"""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    add_title(doc, "沪牌额度使用及车辆管理协议")

    if version == "yi":
        add_subtitle(doc, "（乙方曹老师提议版 —— 按乙方原话补充第八条第 4 款 / 修改第十二条管辖）")
        out_name = "沪牌额度使用及车辆管理协议（乙方提议版）.docx"
    else:
        add_subtitle(doc, "（甲方建议版 —— 站在甲方视角对乙方两条诉求做合理限缩与平衡）")
        out_name = "沪牌额度使用及车辆管理协议（甲方建议版）.docx"

    add_parties_and_intro(doc)
    add_clause_1(doc)
    add_clause_2(doc)
    add_clause_3(doc)
    add_clause_4(doc)
    add_clause_5(doc)
    add_clause_6(doc)
    add_clause_7(doc)
    add_clause_8(doc, version)
    add_clause_9(doc)
    add_clause_10(doc)
    add_clause_11(doc)
    add_clause_12(doc, version)
    add_clause_13(doc)

    add_section_heading(doc, "（以下无正文，为签署页）")
    add_signature_block(doc, INFO["jia_id"], INFO["yi_id"], INFO["sign_date"])

    doc.save(out_name)
    print(f"已生成：{out_name}")


if __name__ == "__main__":
    build_document("yi")
    build_document("jia")
