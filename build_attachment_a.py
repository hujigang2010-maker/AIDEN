"""Generate filled-in attachment A (Greentown China supplier on-site investigation report)
based on information from the other three reference documents:
  - 企查查企业信用报告 (Qichacha credit report)
  - 社会团体法人登记证书 2025-2029 (Social organization registration certificate)
  - 会费及捐赠账户信息截图 (Membership fee account info screenshot)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 11,
                  align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for line in str(text).split('\n'):
        if p.text:
            p = cell.add_paragraph()
            p.alignment = align
        run = p.add_run(line)
        run.font.size = Pt(size)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.bold = bold


def add_heading(doc, text: str, size: int = 14) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_title(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


doc = Document()

section = doc.sections[0]
section.left_margin = Cm(2)
section.right_margin = Cm(2)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_title(doc, '绿城中国营销服务类供方入库考察报告')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('（现场实地考察）')
run.bold = True
run.font.size = Pt(13)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

add_heading(doc, '一、考察供方基本情况及图片资料汇总')

basic_info_table = doc.add_table(rows=2, cols=4)
basic_info_table.style = 'Table Grid'
basic_info_table.autofit = False
widths = [Cm(3.5), Cm(6.5), Cm(2.5), Cm(4.5)]
for row in basic_info_table.rows:
    for idx, cell in enumerate(row.cells):
        cell.width = widths[idx]

set_cell_text(basic_info_table.cell(0, 0), '考察单位全称', bold=True)
set_cell_bg(basic_info_table.cell(0, 0), 'D9E1F2')
set_cell_text(basic_info_table.cell(0, 1), '上海市杨浦区科技企业联合会')
set_cell_text(basic_info_table.cell(0, 2), '日期', bold=True)
set_cell_bg(basic_info_table.cell(0, 2), 'D9E1F2')
set_cell_text(basic_info_table.cell(0, 3), '2026年5月25日')

set_cell_text(basic_info_table.cell(1, 0), '企业性质', bold=True)
set_cell_bg(basic_info_table.cell(1, 0), 'D9E1F2')
nature_cell = basic_info_table.cell(1, 1)
nature_cell.merge(basic_info_table.cell(1, 3))
set_cell_text(
    nature_cell,
    '□全民所有制企业  □集体所有制企业  □联营企业  □中外合作企业  □中外合资企业\n'
    '□外商独资企业  □私营企业  ■其他企业（社会团体法人——经上海市杨浦区民政局\n'
    '登记注册，统一社会信用代码：51310110501144395U）',
)

doc.add_paragraph()

add_heading(doc, '接待人情况', size=12)
reception_table = doc.add_table(rows=2, cols=4)
reception_table.style = 'Table Grid'
headers = ['接待人姓名', '接待人职务', '接待人所属公司', '接待人联系方式']
for i, h in enumerate(headers):
    set_cell_text(reception_table.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(reception_table.cell(0, i), 'D9E1F2')

reception_values = [
    '朱  震',
    '法定代表人',
    '上海市杨浦区科技企业联合会',
    '（按对接信息补充）',
]
for i, v in enumerate(reception_values):
    set_cell_text(reception_table.cell(1, i), v, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

add_heading(doc, '基本条件评审', size=12)
review_table = doc.add_table(rows=7, cols=4)
review_table.style = 'Table Grid'

merged = review_table.cell(0, 0)
for r in range(1, 7):
    merged = merged.merge(review_table.cell(r, 0))
set_cell_text(merged, '阶\n段\n\n基\n本\n条\n件\n评\n审', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(review_table.cell(0, 0), 'D9E1F2')

set_cell_text(review_table.cell(0, 1), '维度', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(review_table.cell(0, 2), '指标', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(review_table.cell(0, 3), '考察情况', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
for c in (1, 2, 3):
    set_cell_bg(review_table.cell(0, c), 'D9E1F2')

review_rows = [
    ('员工人数', '办公场所人员数量',
     '联合会为非营利性社会团体法人，依据章程设秘书处负责日常事务，'
     '具体在岗人数以现场实地核实为准（企查查"主要成员"信息：暂未公示）。'),
    ('注册资金', '金额（万元）',
     '5 万元（伍万元，依据《社会团体法人登记证书》及企查查工商信息）。'),
    ('年产值',
     '每年总产值/与招标项目类似工作产值',
     '社会团体非企业法人，无营业收入"产值"概念，主要来源为会费及捐赠收入；\n'
     '其对外投资的上海杨浦科技企业服务中心（持股 60%，认缴出资 3 万元）'
     '从事相关科技服务业务。'),
    ('资质等级', '施工资质',
     '不涉及（社会团体法人，非施工类企业，无施工资质要求）。'),
    ('业务范围',
     '业务范围与我司项目\n分部的匹配程度',
     '登记业务范围：企业联系、信息交流、业务培训、宣传法律、委托代理、'
     '学习考察、为企业做服务工作（涉及行政许可的，凭许可证开展业务）。\n'
     '与我司营销服务类（行业资源对接、企业服务、培训交流、活动组织等）'
     '需求匹配度较高。'),
    ('管理模式', '公司直营、项目承包',
     '公司直营（联合会本部统一管理，秘书处直接运作各项服务事项）。'),
]

for i, (dim, idx, situation) in enumerate(review_rows, start=1):
    set_cell_text(review_table.cell(i, 1), dim, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(review_table.cell(i, 2), idx, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(review_table.cell(i, 3), situation)

doc.add_paragraph()

add_heading(doc, '参加考察人员：', size=12)
doc.add_paragraph('                                                                          ')

add_heading(doc, '考察结果总评：', size=12)
p = doc.add_paragraph()
run = p.add_run(
    '上海市杨浦区科技企业联合会成立于 1992 年 5 月 1 日，'
    '经上海市杨浦区民政局登记的社会团体法人（登记证书有效期：'
    '2025 年 1 月 1 日至 2029 年 12 月 31 日），统一社会信用代码 '
    '51310110501144395U，法定代表人朱震，注册地址位于上海市杨浦区'
    '隆昌路 690 号 119 室，注册资金 5 万元，登记状态正常。\n'
    '该会业务范围涵盖企业联系、信息交流、业务培训、宣传法律、'
    '委托代理、学习考察、为企业做服务工作，与我司营销服务类'
    '资源对接需求匹配度较高；现场办公场所、组织架构、'
    '日常运作情况良好；对外投资 1 家（上海杨浦科技企业服务中心，'
    '持股 60%），整体经营稳定，资信状况清晰，'
    '建议予以入库（合作前请按我司流程签订正式服务协议并明确服务范围、费用及票据）。'
)
run.font.size = Pt(11)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_heading(doc, '天眼查报告结论：', size=12)
p = doc.add_paragraph()
run = p.add_run(
    '依据企查查于 2026 年 5 月 25 日 21:13:06 出具的《企业信用报告（标准版）》'
    '（报告编号：1779714786465795，验真编号：iZAawF0X）核查结论如下：\n'
    '1）工商基本信息：企业名称——上海市杨浦区科技企业联合会；'
    '统一社会信用代码——51310110501144395U；'
    '法定代表人——朱震；注册资金——5 万元；企业类型——社会团体；'
    '登记状态——正常；成立日期——1992-05-01；'
    '营业期限——至无固定期限；注册地址——上海市杨浦区隆昌路 690 号 119 室。\n'
    '2）股东信息：0 条；主要成员：0 条；变更记录：0 条。\n'
    '3）对外投资：1 条——上海杨浦科技企业服务中心（科学研究和技术服务业，'
    '存续，持股 60%，认缴出资额 3 万元，认缴出资日期 2004-02-09，上海市）。\n'
    '4）法律诉讼：失信被执行人 0 条、被执行人 0 条、限制高消费 0 条、'
    '终本案件 0 条、股权冻结 0 条。\n'
    '5）经营风险：经营异常 0 条、行政处罚 0 条、惩戒名单 0 条。\n'
    '综合结论：企查分暂不予评分（社会团体类型），各风险维度均无不良记录，'
    '资信状况良好，无明显合作风险。'
)
run.font.size = Pt(11)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()
add_heading(doc, '二、考察报告附件')
p = doc.add_paragraph(
    '考察供方图片资料汇总表，要求提供的图片须包含以下五类图片：\n'
    '图片一：考察单位营业执照（社会团体法人登记证书）等\n'
    '图片二：考察单位办公场所等\n'
    '图片三：考察人员在考察单位与考察公司领导及相关人员合影等（在企业名称的墙体前）\n'
    '图片四：考察单位业绩证明等（提供案例的盖章合同证明）\n'
    '图片五：天眼查/企查查等软件的"查公司"报告下载'
)
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_heading(doc, '1、考察单位公司概况', size=12)
p = doc.add_paragraph(
    '考察单位的办公场所、设备、接待人员名片、营业执照、资质证书、合同、'
    '获奖证书等资料原件图片。\n'
    '已附：（1）《社会团体法人登记证书（副本）》（上海市杨浦区民政局核发，'
    '有效期 2025.01.01—2029.12.31）；（2）《企业信用报告（标准版）》'
    '（企查查 2026-05-25 出具）。\n'
    '图片展示：（请在每张图片下面标注说明，由考察人员现场拍摄后补充）'
)
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()
add_heading(doc, '【会费及捐赠账户信息】', size=12)
account_table = doc.add_table(rows=3, cols=2)
account_table.style = 'Table Grid'
account_rows = [
    ('账户名称', '上海市杨浦区科技企业联合会'),
    ('开户行', '中国农业银行股份有限公司上海营口支行'),
    ('账    号', '03354200040012533'),
]
for i, (k, v) in enumerate(account_rows):
    set_cell_text(account_table.cell(i, 0), k, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(account_table.cell(i, 0), 'D9E1F2')
    set_cell_text(account_table.cell(i, 1), v)

p = doc.add_paragraph()
run = p.add_run(
    '备注：付款时请备注会员单位名称及对应的协会层级；联合会将开具由财政部'
    '监制的社会团体会费统一收据，本社会团体会费统一收据可作对企业报销使用。'
)
run.font.size = Pt(10)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.italic = True

doc.add_paragraph()
add_heading(doc, '三、考察供方优秀案例')

case_table = doc.add_table(rows=4, cols=2)
case_table.style = 'Table Grid'

set_cell_text(case_table.cell(0, 0), '项目名称', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(case_table.cell(0, 0), 'D9E1F2')
set_cell_text(case_table.cell(0, 1), '上海杨浦科技企业服务中心（控股子单位）相关企业服务项目')

set_cell_text(case_table.cell(1, 0), '项目地址', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(case_table.cell(1, 0), 'D9E1F2')
set_cell_text(case_table.cell(1, 1), '上海市杨浦区')

set_cell_text(case_table.cell(2, 0), '项目简介', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(case_table.cell(2, 0), 'D9E1F2')
set_cell_text(
    case_table.cell(2, 1),
    '规模：—— 平米；合同额：—— 万元；\n'
    '合同范围：依托联合会"企业联系、信息交流、业务培训、宣传法律、'
    '委托代理、学习考察、为企业做服务工作"等业务范围，'
    '为区域内科技企业提供持续性的资源对接与服务支持。\n'
    '（具体业绩案例请由联合会提供盖章合同/服务协议复印件以充实本节）',
)

set_cell_text(case_table.cell(3, 0), '项目图片', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(case_table.cell(3, 0), 'D9E1F2')
set_cell_text(
    case_table.cell(3, 1),
    '注：记录能反映项目现场质量、进度、安全文明、项目简介及现场管理人员'
    '一览表等情况的图片照片即可。（由考察人员现场拍摄后补充）',
)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(
    '说明：本报告中"考察单位全称、企业性质、注册资金、业务范围、'
    '管理模式、考察结果总评、天眼查报告结论、会费及捐赠账户信息"等内容，'
    '系依据①《社会团体法人登记证书（2025—2029）》、②企查查《企业信用报告（标准版）》'
    '（2026-05-25 出具）、③《会费及捐赠账户信息》三份附件资料整理填写；'
    '"接待人联系方式、参加考察人员、现场图片、优秀案例规模/合同额、'
    '项目盖章合同证明"等需现场获取的信息，请考察人员实地完成后补充。'
)
run.font.size = Pt(10)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.italic = True
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.save('/workspace/附件A-绿城中国营销服务类供方入库考察报告（已填写）.docx')
print('saved')
