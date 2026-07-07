# -*- coding: utf-8 -*-
"""Word 商务大气样式辅助(合集文档共用)。"""
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = "微软雅黑"
NAVY = RGBColor(0x0F, 0x2A, 0x4A)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
GOLD = RGBColor(0xA8, 0x80, 0x2A)
GRAY = RGBColor(0x59, 0x59, 0x59)
NAVY_HEX = "0F2A4A"
BLUE_HEX = "1F4E79"
GOLD_HEX = "C79A3B"
LIGHT_HEX = "EEF2F8"
LIGHTGOLD_HEX = "F6EEDA"


def cn(run, size=None, bold=None, color=None, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(el, hex_color):
    tcPr = el.get_or_add_tcPr() if el.tag.endswith("}tc") else el.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def para(doc, text, size=11.5, bold=False, color=None, align=None,
         space_after=6, space_before=0, indent=None, line=1.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    cn(r, size=size, bold=bold, color=color)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line
    if indent:
        p.paragraph_format.first_line_indent = Pt(indent)
    return p


def bullet(doc, text, size=11.5, color=None):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    cn(r, size=size, color=color)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    return p


def h1(doc, text, no=None):
    """一级标题:金色色块编号 + 深蓝标题 + 下方金线。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    if no:
        rn = p.add_run(f"  {no}  ")
        cn(rn, size=15, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(p._p, GOLD_HEX)
    rt = p.add_run(("  " if no else "") + text)
    cn(rt, size=16, bold=True, color=NAVY)
    _bottom_border(p, GOLD_HEX, sz=14)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    cn(r, size=13, bold=True, color=BLUE)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p


def _bottom_border(p, hex_color, sz=12):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), hex_color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def banner(doc, lines, fill_hex=NAVY_HEX, height_cm=None):
    """整宽色块横幅(单元格表格),lines: list of (text, size, color, bold)。"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell._tc, fill_hex)
    cell.paragraphs[0].clear() if hasattr(cell.paragraphs[0], "clear") else None
    first = True
    for (text, size, color, bold) in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(text)
        cn(r, size=size, bold=bold, color=color)
    _set_cell_margins(cell, top=240, bottom=240, left=160, right=160)
    _remove_table_borders(table)
    return table


def styled_table(doc, headers, rows, col_widths_cm=None, header_hex=NAVY_HEX,
                 zebra_hex=LIGHT_HEX, header_size=10.5, body_size=10,
                 center_cols=None, last_row_bold=False):
    center_cols = center_cols or set()
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        shade(cell._tc, header_hex)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        cn(r, size=header_size, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    # body
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cell = cells[j]
            if zebra_hex and i % 2 == 1:
                shade(cell._tc, zebra_hex)
            p = cell.paragraphs[0]
            if j in center_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            cn(r, size=body_size, color=RGBColor(0x22, 0x2B, 0x35))
            p.paragraph_format.line_spacing = 1.15
    if last_row_bold:
        for cell in table.rows[-1].cells:
            shade(cell._tc, LIGHTGOLD_HEX)
            for p in cell.paragraphs:
                for r in p.runs:
                    cn(r, bold=True, color=NAVY)
    if col_widths_cm:
        for j, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    return table


def _set_cell_margins(cell, top=100, bottom=100, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left),
                     ("end", right), ("left", left), ("right", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)
