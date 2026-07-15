# -*- coding: utf-8 -*-
"""生成 Word《策划服务建议书(报价版)》——含分工、工作分解与16万报价。"""
import os

from docx import Document
from docx.shared import Cm, RGBColor

import content_framework as C
from style_docx import (BLUE, GOLD, GRAY, NAVY, WD_ALIGN_PARAGRAPH,
                        banner, bullet, h1, h2, para, styled_table)

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "deliverables")
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHTGOLD = RGBColor(0xE8, 0xC9, 0x77)


def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    banner(doc, [
        (C.PROJECT_NAME, 24, WHITE, True),
        ("GLOBAL FREE TRADE ZONE · HUANGPU JIUFO TOD", 10, LIGHTGOLD, True),
    ])
    para(doc, "", space_after=2)
    banner(doc, [
        (C.SUBTITLE, 24, NAVY, True),
        ("（策划服务框架 · 分工与报价）", 12, GRAY, False),
    ], fill_hex="F6EEDA")
    for _ in range(6):
        doc.add_paragraph()
    para(doc, f"提供方:{C.PROVIDER_LINE}", size=13, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(doc, f"第四部分承接方:{C.PARTNER}", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(doc, f"提交对象:{C.CLIENT}", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(doc, C.DOC_DATE, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def main():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.4); sec.bottom_margin = Cm(2.4)
        sec.left_margin = Cm(2.6); sec.right_margin = Cm(2.6)
    cover(doc)

    # 一、项目与服务框架概述
    h1(doc, "项目与服务框架概述", no="一")
    para(doc, C.PREFACE, size=11.5, indent=24)
    h2(doc, "分工说明")
    para(doc, f"第一至第三部分:由{C.OUR_SIDE}负责。", size=11.5, indent=24)
    para(doc, f"第四部分(场景营造与空间引导):由{C.PARTNER}负责,我方不承担该部分工作。",
         size=11.5, indent=24)

    # 二、服务内容与工作分解
    h1(doc, "服务内容与工作分解", no="二")
    for (no, name, owner, fee, desc, subs) in C.PARTS:
        title = f"（{no}）{name}　【负责方:{owner}】"
        para(doc, title, size=12.5, bold=True, color=BLUE, space_before=8, space_after=3)
        para(doc, desc, size=11, indent=24, space_after=4)
        for s in subs:
            bullet(doc, s, size=11)

    # 三、成果交付(我方负责范围)
    h1(doc, "成果交付（我方负责:第一至第三部分）", no="三")
    for d in C.DELIVERABLES:
        bullet(doc, d)
    para(doc, "注:第四部分“场景营造与空间引导”的成果由内里集交付。", size=10.5,
         color=GRAY, indent=24, space_before=4)

    # 四、服务报价
    h1(doc, "服务报价", no="四")
    rows = []
    for (no, name, owner, fee, desc, subs) in C.PARTS:
        fee_txt = f"{fee}" if fee is not None else "—（内里集负责）"
        rows.append((f"第{no}部分  {name}", owner, fee_txt))
    rows.append(("我方合计（第一至第三部分）", "我方", str(C.OUR_TOTAL)))
    styled_table(doc, ["服务部分", "负责方", "报价(万元)"], rows,
                 col_widths_cm=[8.6, 4.6, 2.6], center_cols={1, 2},
                 body_size=10.5, header_size=11, last_row_bold=True)
    para(doc, "", size=4)
    for n in C.QUOTATION_NOTES:
        bullet(doc, n, size=10.5)

    # 五、服务团队与说明
    h1(doc, "服务团队与说明", no="五")
    para(doc,
         "本项目第一至第三部分由复旦大学住房政策研究中心与上海市杨浦区科技企业联合会"
         "联合组建服务团队:研究中心负责政策研究、市场洞察与战略定位;科技企业联合会负责"
         "产业资源组织、招商与出海路径设计。第四部分由内里集负责,双方按框架分工协同推进。",
         size=11.5, indent=24)
    para(doc, "联系方式:另行提供(以正式合同联络人为准)。", size=11.5, indent=24)

    os.makedirs(OUT_DIR, exist_ok=True)
    out = os.path.join(OUT_DIR, "广州黄埔九佛TOD全球自贸区_策划服务建议书_报价版.docx")
    doc.save(out)
    print("saved:", out)


if __name__ == "__main__":
    main()
