#!/usr/bin/env python3
"""生成《接手失败项目的通用框架》操作手册 Word。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


TEAL = RGBColor(0x0F, 0x76, 0x6E)
INK = RGBColor(0x0F, 0x17, 0x2A)
SLATE = RGBColor(0x47, 0x55, 0x69)


def set_run_font(run, *, size=12, bold=False, color=None, name="微软雅黑"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading_cn(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=18 if level == 1 else 14, bold=True, color=TEAL if level == 1 else INK)
    return p


def add_para(doc, text, *, size=11, bold=False, color=None, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_bullets(doc, items, *, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=size)
        p.paragraph_format.space_after = Pt(4)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=INK)
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = ""
            p = cells[c_i].paragraphs[0]
            run = p.add_run(val)
            set_run_font(run, size=10)
    doc.add_paragraph()


def build_docx(output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("接手失败项目的通用框架")
    set_run_font(run, size=22, bold=True, color=INK)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("全量摸底 → 按自己的逻辑重构 → 完成后及时转向")
    set_run_font(run, size=12, color=TEAL)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("适用：AI 项目管理 · 产品接手 · 技术债抢救 · 失败业务盘活")
    set_run_font(run, size=10, color=SLATE)

    add_heading_cn(doc, "一、为什么需要这套框架", 1)
    add_para(
        doc,
        "失败项目最危险的不是「已经坏了」，而是接手者用错误节奏继续坏下去："
        "一上来救火、默认继承前任逻辑、清理没有终点。个人英雄主义或许能扛过一两次，"
        "但无法迁移到 AI 项目管理、产品接手等重复场景。",
    )
    add_para(doc, "本框架把接手压缩为三步闭环，目标是：建立真相 → 夺回主权 → 按时离开清理态。", bold=True)

    add_heading_cn(doc, "二、框架总览", 1)
    add_table(
        doc,
        ["步骤", "核心问题", "关键输出", "禁止事项"],
        [
            ["01 全量摸底", "现在到底是什么状态？", "真相图、风险清单、停做清单", "边摸边大改"],
            ["02 逻辑重构", "按谁的逻辑前进？", "新成功定义、杀/留/重写决策", "双轨目标并行"],
            ["03 及时转向", "什么时候算接手完成？", "完成标准、Owner、playbook", "无限收尾"],
        ],
    )

    add_heading_cn(doc, "三、步骤一：全量摸底", 1)
    add_para(doc, "摸底不是拖延，是防止在错误地图上加速。建议设 3–5 天时间盒，对外宣布「只读期」。")
    add_heading_cn(doc, "3.1 摸什么", 2)
    add_bullets(
        doc,
        [
            "目标与承诺：当初卖给客户/老板什么，现在还剩什么。",
            "资产：代码、数据、文档、账号权限、合同与供应商。",
            "债务：技术债、合规债、信誉债、人情债。",
            "干系人：谁拍板、谁被影响、谁在暗中等待结果。",
            "约束：时间、预算、法务、品牌、不可碰的红线。",
        ],
    )
    add_heading_cn(doc, "3.2 怎么摸", 2)
    add_bullets(
        doc,
        [
            "只读优先：先读后改；任何「顺手修一下」都要登记。",
            "三色标注：绿可用 / 黄存疑 / 红不可用。",
            "事实与观点分离：先记证据，再写假设。",
            "输出一页「现状真相图」：可复用 / 需改造 / 冻结 / 待查。",
        ],
    )
    add_heading_cn(doc, "3.3 摸底完成标准", 2)
    add_bullets(
        doc,
        [
            "能提出「为什么失败」的主因假设，并标明证据强弱。",
            "能指出约 20% 可 salvage 的高价值资产。",
            "能列出必须立刻停掉的危险动作。",
            "关键干系人对现状陈述基本无重大异议。",
        ],
    )

    add_heading_cn(doc, "四、步骤二：按自己的逻辑重构", 1)
    add_para(
        doc,
        "接手不是继承前任世界观。若继续沿用失败路径上的目标、架构与叙事，"
        "所谓「重构」只是给旧系统续命。",
    )
    add_heading_cn(doc, "4.1 六条原则", 2)
    add_bullets(
        doc,
        [
            "重定义成功：在当前约束下重写目标与非目标。",
            "先杀后留：默认不信任旧路线；能删则删。",
            "主权在己：优先级与验收标准由接手方定义。",
            "小闭环验证：最短路径证明新逻辑，再扩大半径。",
            "叙事同步改：对内对外统一失败归因与新策略。",
            "工具服从逻辑：先定操作系统，再选 AI/流程/看板。",
        ],
    )
    add_heading_cn(doc, "4.2 杀 / 留 / 重写", 2)
    add_table(
        doc,
        ["动作", "适用信号", "典型动作", "做错的代价"],
        [
            ["杀", "无主责、无证据、无用户", "下线、冻结、归档、切断入口", "继续烧资源养空壳"],
            ["留", "稳定、可证伪、低维护成本", "最小改动纳入新边界，写清契约", "误删核心资产"],
            ["重写", "主路径但旧设计不可救", "按新逻辑重做最小可行切片", "大爆炸重写失控"],
        ],
    )

    add_heading_cn(doc, "五、步骤三：完成后及时转向", 1)
    add_para(
        doc,
        "没有「接手完成」定义，就会永远停在修补循环。提前写下可验收标准，"
        "并用 WIP 上限与公开倒计时防止拖延。",
        bold=True,
    )
    add_heading_cn(doc, "5.1 完成定义示例", 2)
    add_bullets(
        doc,
        [
            "主路径可演示且可复现。",
            "关键风险已封堵或明确 Owner。",
            "干系人接受新目标与边界。",
            "文档与职责完成书面交接。",
        ],
    )
    add_heading_cn(doc, "5.2 转向去向", 2)
    add_bullets(
        doc,
        [
            "进入增长/交付创造期。",
            "交给稳态 Owner 运营。",
            "正式关停并沉淀复盘。",
            "拆成新项目重新立项。",
        ],
    )

    add_heading_cn(doc, "六、场景应用", 1)
    add_heading_cn(doc, "6.1 AI 项目管理", 2)
    add_para(doc, "AI 项目常败在：演示代替上线、Agent 链路失控、评测缺失、成本账单无人看。")
    add_bullets(
        doc,
        [
            "摸底：目标指标、数据质量、模型/Agent 边界、评测集、成本。",
            "重构：业务问题 → 任务分解 → 人机分工；杀无效链路，建验收与回退。",
            "转向：主链路达标即冻结探索实验；探索预算单列；交稳态团队。",
        ],
    )
    add_heading_cn(doc, "6.2 产品接手", 2)
    add_para(doc, "建议用 2–4 周节奏：第 1 周摸底，第 2–3 周重构，第 4 周转向常态迭代。")
    add_bullets(
        doc,
        [
            "第 1 周：用户路径走查 + 收入/留存/成本三张表 + 干系人访谈。",
            "第 2–3 周：重写北极星与非目标，砍到一条主路径，统一叙事。",
            "第 4 周：演示可验收切片，明确 Owner/SLA，关闭接手专项看板。",
        ],
    )

    add_heading_cn(doc, "七、反模式清单", 1)
    add_table(
        doc,
        ["反模式", "为什么危险", "替代动作"],
        [
            ["边摸边大改", "无法归因新旧问题", "宣布只读期，改动一律登记"],
            ["讨好式继承", "双目标并行撕碎资源", "书面重定目标与非目标"],
            ["完美主义清理", "错过窗口期", "设完成定义与 WIP 上限"],
            ["工具崇拜", "无逻辑的效率幻觉", "先定操作系统再选工具"],
            ["口头交接", "问题回流到接手人", "书面完成标准 + Owner"],
            ["英雄主义", "经验不可迁移", "沉淀 playbook 并交班"],
        ],
    )

    add_heading_cn(doc, "八、一页作战手册", 1)
    add_bullets(
        doc,
        [
            "设时间盒，宣布只读摸底期——禁止大改。",
            "产出真相图：可复用 / 需改造 / 冻结 / 待查。",
            "重写成功定义与非目标，获关键干系人书面确认。",
            "用杀/留/重写矩阵压缩到一条主路径。",
            "跑通最短可验证闭环，再扩大重构半径。",
            "写下接手完成标准；到期转向，清理任务设上限。",
            "沉淀 playbook，交给稳态 Owner，自己进入下一创造。",
        ],
    )

    add_heading_cn(doc, "九、使用建议", 1)
    add_para(
        doc,
        "把本手册与配套 PPT、Excel 清单一起使用：PPT 用于对齐认知，"
        "Excel 用于落地执行，Word 用于归档与培训。每次接手结束后，"
        "用半小时复盘更新清单——框架才会越用越强。",
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— 完 —")
    set_run_font(run, size=10, color=SLATE)

    doc.save(str(output_path))
    print(f"已生成: {output_path}")


if __name__ == "__main__":
    out = Path(__file__).resolve().parents[1] / "deliverables" / "接手失败项目的通用框架_操作手册.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    build_docx(out)
