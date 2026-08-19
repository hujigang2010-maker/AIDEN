#!/usr/bin/env python3
"""生成给律师阅卷用的完整经过说明（Markdown + Word + PDF）。不是律师函。"""

from pathlib import Path
import html

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.colors import HexColor, white
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from comm_plan import CN_FONT_PATH
from content import LAWYER_BRIEF, LAWYER_SECTIONS

CHINESE_FONT = "宋体"
HEADING_FONT = "黑体"
ACCENT = RGBColor(0x0B, 0x2F, 0x5B)
MUTED = RGBColor(0x5C, 0x6B, 0x7A)
RED = RGBColor(0xA6, 0x3D, 0x2F)

MD_NAME = "给律师的事故完整经过说明_2026-08-19.md"
DOCX_NAME = "青岛抚顺路和哈尔滨路路口交通事故_给律师的完整经过说明_20260819.docx"
PDF_NAME = "青岛抚顺路和哈尔滨路路口交通事故_给律师的完整经过说明_20260819.pdf"


def _set_run_font(run, font_name: str, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _add_paragraph(
    doc: Document,
    text: str = "",
    *,
    font: str = CHINESE_FONT,
    size: float = 12,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: float | None = None,
    space_after: float = 6,
    space_before: float = 0,
    line_spacing: float = 1.5,
    color: RGBColor | None = None,
):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if text:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold, color=color)
    return p


def render_markdown() -> str:
    b = LAWYER_BRIEF
    lines = [
        f"# {b['title']}",
        "",
        b["subtitle"],
        "",
        f"> {b['file_note']}",
        "",
        "## 请律师先做这几件事",
        "",
    ]
    for i, item in enumerate(b["ask"], 1):
        lines.append(f"{i}. {item}")
    lines.append("")
    for sec in LAWYER_SECTIONS:
        lines.append(f"## {sec['h']}")
        lines.append("")
        for p in sec["paras"]:
            lines.append(p)
            lines.append("")
    lines.append("## 十、现有材料与仍缺文件")
    lines.append("")
    lines.append("已经可以交给律师的：")
    lines.append("")
    for x in b["have"]:
        lines.append(f"- {x}")
    lines.append("")
    lines.append("请指导家属尽快补齐的：")
    lines.append("")
    for x in b["need"]:
        lines.append(f"- {x}")
    lines.append("")
    lines.append("重新生成：`python3 scripts/build_all.py`。3D 示范动画留作待用，不是证据。")
    lines.append("")
    return "\n".join(lines)


def build_markdown(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(render_markdown(), encoding="utf-8")


def build_document(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.6)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = CHINESE_FONT
    normal_style.font.size = Pt(12)
    rpr = normal_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CHINESE_FONT)

    b = LAWYER_BRIEF
    _add_paragraph(
        doc,
        b["title"],
        font=HEADING_FONT,
        size=18,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
        line_spacing=1.3,
        color=ACCENT,
    )
    _add_paragraph(
        doc,
        b["subtitle"],
        font=HEADING_FONT,
        size=11,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
        line_spacing=1.3,
        color=MUTED,
    )
    _add_paragraph(doc, b["file_note"], bold=True, color=RED, space_after=10)

    _add_paragraph(doc, "请律师先做这几件事", font=HEADING_FONT, size=14, bold=True, space_before=8, space_after=8, color=ACCENT, line_spacing=1.3)
    for item in b["ask"]:
        _add_paragraph(doc, f"• {item}", first_line_indent=0.37, space_after=3, line_spacing=1.4)

    for sec in LAWYER_SECTIONS:
        _add_paragraph(doc, sec["h"], font=HEADING_FONT, size=14, bold=True, space_before=14, space_after=8, color=ACCENT, line_spacing=1.3)
        for p in sec["paras"]:
            _add_paragraph(doc, p, first_line_indent=0.74, space_after=6)

    _add_paragraph(doc, "十、现有材料与仍缺文件", font=HEADING_FONT, size=14, bold=True, space_before=14, space_after=8, color=ACCENT, line_spacing=1.3)
    _add_paragraph(doc, "已经可以交给律师的：", bold=True, space_after=4)
    for x in b["have"]:
        _add_paragraph(doc, f"• {x}", first_line_indent=0.37, space_after=3, line_spacing=1.4)
    _add_paragraph(doc, "请指导家属尽快补齐的：", bold=True, space_before=8, space_after=4)
    for x in b["need"]:
        _add_paragraph(doc, f"• {x}", first_line_indent=0.37, space_after=3, line_spacing=1.4)

    _add_paragraph(
        doc,
        "本文根据 2026 年 8 月 14–19 日材料整理。重新生成：python3 scripts/build_all.py。3D 示范动画留作待用，不是证据。",
        size=10.5,
        color=MUTED,
        space_before=12,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _pdf_font() -> str:
    kwargs = {"subfontIndex": 0} if CN_FONT_PATH.endswith(".ttc") else {}
    pdfmetrics.registerFont(TTFont("CN", CN_FONT_PATH, **kwargs))
    return "CN"


def _pdf_styles(font: str) -> dict[str, ParagraphStyle]:
    navy = HexColor("#0B2F5B")
    muted = HexColor("#5C6B7A")
    red = HexColor("#A63D2F")
    dark = HexColor("#1F2A37")
    return {
        "title": ParagraphStyle("t", fontName=font, fontSize=16, leading=24, textColor=navy, alignment=TA_CENTER, spaceAfter=6),
        "sub": ParagraphStyle("s", fontName=font, fontSize=10, leading=16, textColor=muted, alignment=TA_CENTER, spaceAfter=8),
        "warn": ParagraphStyle("w", fontName=font, fontSize=10, leading=16, textColor=red, alignment=TA_LEFT, spaceAfter=10),
        "h1": ParagraphStyle("h", fontName=font, fontSize=13, leading=20, textColor=navy, spaceBefore=12, spaceAfter=6),
        "body": ParagraphStyle("b", fontName=font, fontSize=10.5, leading=17, textColor=dark, alignment=TA_JUSTIFY, firstLineIndent=22, spaceAfter=6),
        "bullet": ParagraphStyle("u", fontName=font, fontSize=10.5, leading=16, textColor=dark, leftIndent=14, spaceAfter=3),
        "label": ParagraphStyle("l", fontName=font, fontSize=10.5, leading=16, textColor=dark, spaceBefore=6, spaceAfter=4),
        "foot": ParagraphStyle("f", fontName=font, fontSize=8, leading=12, textColor=muted, spaceBefore=10),
    }


def _p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(html.escape(str(text)).replace("\n", "<br/>"), style)


def _header_footer(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(HexColor("#0B2F5B"))
    canvas.rect(0, A4[1] - 12 * mm, A4[0], 12 * mm, fill=1, stroke=0)
    canvas.setFillColor(white)
    canvas.setFont("CN", 8)
    canvas.drawString(18 * mm, A4[1] - 8 * mm, "青岛抚顺路和哈尔滨路路口交通事故 · 给律师的完整经过说明 · 内部阅卷")
    canvas.setFillColor(HexColor("#C4A35A"))
    canvas.rect(0, 0, A4[0], 10 * mm, fill=1, stroke=0)
    canvas.setFillColor(HexColor("#0B2F5B"))
    canvas.setFont("CN", 8)
    canvas.drawString(18 * mm, 4 * mm, "不是律师函 · 不要发给骑手或刘孝春 · 现阶段不谈总赔偿数额")
    canvas.drawRightString(A4[0] - 18 * mm, 4 * mm, f"第 {doc.page} 页")
    canvas.restoreState()


def build_pdf(output_path: Path) -> None:
    font = _pdf_font()
    styles = _pdf_styles(font)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=16 * mm,
        title=LAWYER_BRIEF["title"],
        author="内部整理",
    )
    story = [
        _p(LAWYER_BRIEF["title"], styles["title"]),
        _p(LAWYER_BRIEF["subtitle"], styles["sub"]),
        _p(LAWYER_BRIEF["file_note"], styles["warn"]),
        _p("请律师先做这几件事", styles["h1"]),
    ]
    for item in LAWYER_BRIEF["ask"]:
        story.append(_p("• " + item, styles["bullet"]))
    for sec in LAWYER_SECTIONS:
        story.append(_p(sec["h"], styles["h1"]))
        for para in sec["paras"]:
            story.append(_p(para, styles["body"]))
    story.append(_p("十、现有材料与仍缺文件", styles["h1"]))
    story.append(_p("已经可以交给律师的：", styles["label"]))
    for x in LAWYER_BRIEF["have"]:
        story.append(_p("• " + x, styles["bullet"]))
    story.append(_p("请指导家属尽快补齐的：", styles["label"]))
    for x in LAWYER_BRIEF["need"]:
        story.append(_p("• " + x, styles["bullet"]))
    story.append(Spacer(1, 8))
    story.append(_p("本文根据 2026 年 8 月 14–19 日材料整理。重新生成：python3 scripts/build_all.py。3D 示范动画留作待用，不是证据。", styles["foot"]))
    doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "deliverables"
    build_markdown(out / MD_NAME)
    build_document(out / DOCX_NAME)
    build_pdf(out / PDF_NAME)
    print(f"Wrote {out / MD_NAME}")
    print(f"Wrote {out / DOCX_NAME}")
    print(f"Wrote {out / PDF_NAME}")


if __name__ == "__main__":
    main()
