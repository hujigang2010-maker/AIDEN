"""Generate the enriched Tencent Cloud Strategic Sponsorship Agreement Word doc.

参照绿城附件四的结构、深度与法务严密度，对腾讯云 30 万元首席战略协办协议
进行扩写，作为评标参考版本。保留原协议「一揽子综合对价整体交付、不作分项
标价」的核心商业逻辑，并通过权益清单的丰富度体现 30 万元对应的真实价值。"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


GREEN = RGBColor(0x00, 0x6B, 0x3F)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)
TENCENT_BLUE = RGBColor(0x00, 0x52, 0xD9)
DARK_BLUE = RGBColor(0x00, 0x3A, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _force_font(run, font="宋体"):
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rpr.append(rfonts)


def set_default_font(doc, name="宋体"):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def heading(doc, text, level=1, color=None, center=False, font="黑体"):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(20)
    elif level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    if color is not None:
        run.font.color.rgb = color
    _force_font(run, font=font)
    return p


def para(doc, text, bold=False, size=11, color=DARK, align=None,
         first_indent=True, font="宋体", line_spacing=1.5):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = line_spacing
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    _force_font(run, font=font)
    return p


def clause(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(0)
    run_num = p.add_run(num + "  ")
    run_num.bold = True
    run_num.font.size = Pt(11)
    _force_font(run_num)
    run = p.add_run(text)
    run.font.size = Pt(11)
    _force_font(run)
    return p


def blank(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=DARK, size=10.5,
                  align_center=False, font="宋体"):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    _force_font(run, font=font)


def add_kv_table(doc, rows, col_widths=(4.0, 12.5)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(col_widths[0])
        row.cells[1].width = Cm(col_widths[1])
    for i, (k, v) in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], k, bold=True, size=10.5)
        shade_cell(table.rows[i].cells[0], "E6EFFF")
        set_cell_text(table.rows[i].cells[1], v, size=10.5)
    return table


def add_table(doc, headers, rows, col_widths=None, header_fill="0052D9",
              header_color=RGBColor(0xFF, 0xFF, 0xFF)):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True,
                      color=header_color, size=10.5, align_center=True)
        shade_cell(table.rows[0].cells[i], header_fill)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            set_cell_text(table.rows[r + 1].cells[c], str(val), size=10)
    return table


def build():
    doc = Document()
    set_default_font(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

    # ====== 封面页 ======
    blank(doc, 4)
    heading(doc, "2026 人工智能商业化落地与硬核投资破局峰会",
            level=2, color=GREY, center=True)
    blank(doc)
    heading(doc, "《战略协办赞助合作协议》",
            level=0, color=DARK_BLUE, center=True)
    blank(doc)
    heading(doc,
            "—— 首席战略合作伙伴 · 算力生态独家伙伴 ——",
            level=2, color=GREY, center=True)
    blank(doc, 3)

    cover = doc.add_table(rows=4, cols=2)
    cover.style = "Table Grid"
    cover.autofit = False
    for row in cover.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(11.0)
    set_cell_text(cover.rows[0].cells[0], "甲方（主办方 / 组委会）",
                  bold=True, size=11.5)
    shade_cell(cover.rows[0].cells[0], "E6EFFF")
    set_cell_text(cover.rows[0].cells[1],
                  "人工智能商业化落地峰会组委会\n复旦大学住房政策研究中心\n上海市杨浦区科技企业联合会（联合主办）",
                  size=11)
    set_cell_text(cover.rows[1].cells[0], "乙方（首席战略合作伙伴）",
                  bold=True, size=11.5)
    shade_cell(cover.rows[1].cells[0], "E6EFFF")
    set_cell_text(cover.rows[1].cells[1],
                  "腾讯云计算（北京）有限责任公司",
                  size=11)
    set_cell_text(cover.rows[2].cells[0], "合作级别",
                  bold=True, size=11.5)
    shade_cell(cover.rows[2].cells[0], "E6EFFF")
    set_cell_text(cover.rows[2].cells[1],
                  "首席战略合作伙伴 + 算力生态独家伙伴（同级别唯一席位）",
                  size=11)
    set_cell_text(cover.rows[3].cells[0], "赞助金额",
                  bold=True, size=11.5)
    shade_cell(cover.rows[3].cells[0], "E6EFFF")
    set_cell_text(cover.rows[3].cells[1],
                  "人民币叁拾万元整（￥300,000.00，含税）",
                  bold=True, color=TENCENT_BLUE, size=12)
    blank(doc, 3)

    para(doc,
         "依据本协议第 8.4 条「独立比对声明」，甲方有权在对外招商、生态展示及向同级别合作方进行横向对位时，合法引用本协议第二条项下的合作总金额与合作级别作为优秀案例说明。",
         align="center", first_indent=False, size=10, color=GREY)
    blank(doc, 2)
    para(doc, "签 署 日 期：二〇二六 年 五 月 十二 日",
         align="center", first_indent=False, size=11, color=DARK, bold=True)
    para(doc, "签 署 地 点：上海市",
         align="center", first_indent=False, size=11, color=DARK)
    para(doc, "合 同 编 号：AIBIZ-2026-SP-013",
         align="center", first_indent=False, size=11, color=DARK)

    doc.add_page_break()

    # ====== 正文抬头 ======
    heading(doc, "2026 人工智能商业化落地与硬核投资破局峰会",
            level=2, color=GREY, center=True)
    heading(doc, "《战略协办赞助合作协议》",
            level=0, color=DARK_BLUE, center=True)
    para(doc,
         "Strategic Co-Organizer Sponsorship Agreement",
         align="center", size=10, color=GREY, first_indent=False)
    blank(doc)

    para(doc, "合同编号：AIBIZ-2026-SP-013      签署地点：上海市      签署日期：2026 年 5 月 12 日",
         first_indent=False, size=10.5, align="center")
    blank(doc)

    # ====== 鉴于 ======
    para(doc, "鉴于：", bold=True, first_indent=False, size=12)
    para(doc,
         "1. 甲方拟于 2026 年 5 月在上海·北外滩核心地标（北外滩 · 一滴水）举办「重构与突围 — 2026 人工智能商业化落地与硬核投资破局峰会」（以下简称「本次峰会」），峰会聚焦人工智能在新质生产力、大模型商业化、产业数字化转型、算力底层基础设施及一二级市场资本配置等领域的落地实践，预计现场规模 500+ 位高净值嘉宾，包括上市公司董事长、独角兽高管、一二级市场基金合伙人、AI/算力/大模型创业团队创始人、北大复旦双校核心校友及政府、产业园区、媒体合作伙伴；")
    para(doc,
         "2. 乙方系国内领先的云计算与人工智能基础设施服务商，在算力调度、大模型 API 与底层架构层面具备产业级标杆能力，希望通过本次峰会向核心政企客户与产业生态展示其在算力与大模型领域的商业化能力，并以「首席战略协办伙伴」及「算力生态独家伙伴」身份深度联动；")
    para(doc,
         "3. 双方已就本次合作的核心商业要素达成原则一致，并就乙方战略协办本次峰会的赞助金额、权益范围、物料植入、执行节点、付款发票、知识产权、保密合规、违约责任、争议解决等事项进行了多轮协商；")
    para(doc,
         "4. 本协议作为双方就上述合作事项的正式书面文件，与其后双方书面确认的附件、补充协议、往来函件具有同等法律效力。")
    blank(doc)

    para(doc,
         "双方本着平等、自愿、诚实信用、互利共赢的原则，依据《中华人民共和国民法典》《中华人民共和国广告法》《中华人民共和国反不正当竞争法》《中华人民共和国商标法》《中华人民共和国著作权法》《中华人民共和国网络安全法》《中华人民共和国数据安全法》《中华人民共和国个人信息保护法》及其他相关法律法规之规定，就乙方以「首席战略协办伙伴」及「算力生态独家伙伴」身份赞助本次峰会事宜达成如下协议，以资共同遵守。")
    blank(doc)

    # ====== 一、双方主体信息 ======
    heading(doc, "一、双方主体信息", level=1, color=DARK_BLUE)

    para(doc, "甲方（主办方 / 组委会）：", bold=True, first_indent=False)
    add_kv_table(doc, rows=[
        ["联合主办", "人工智能商业化落地峰会组委会 / 复旦大学住房政策研究中心 / 上海市杨浦区科技企业联合会"],
        ["统一社会信用代码（牵头方）", "12100000425005665L（复旦大学住房政策研究中心 · 示意填写，最终以盖章方为准）"],
        ["通讯地址", "上海市杨浦区国权路 600 号 复旦大学经济学院"],
        ["授权代表 / 职务", "____________________ / 组委会执行秘书长"],
        ["项目联系人", "王诗潼      联系电话：132 6260 7888"],
        ["电子邮箱", "office@ai-biz-summit.cn"],
        ["收款单位", "______________________________（以盖章方银行预留名称为准）"],
        ["开户银行", "______________________________"],
        ["银行账号", "______________________________"],
        ["税号", "______________________________"],
    ])
    blank(doc)

    para(doc, "乙方（首席战略合作伙伴 / 算力生态独家伙伴）：",
         bold=True, first_indent=False)
    add_kv_table(doc, rows=[
        ["公司名称", "腾讯云计算（北京）有限责任公司"],
        ["统一社会信用代码", "91110108576649585L"],
        ["注册地址", "北京市海淀区中关村东路 1 号院 9 号楼 7 层 705"],
        ["授权代表 / 职务", "____________________ / 腾讯云副总裁（华东渠道生态总经理）"],
        ["项目联系人", "____________________      联系电话：____________________"],
        ["电子邮箱", "____________________@tencent.com"],
    ])
    blank(doc)

    # ====== 二、合作级别与赞助标的 ======
    heading(doc, "二、合作级别与赞助标的", level=1, color=DARK_BLUE)
    clause(doc, "2.1",
           "合作级别：乙方为本次峰会唯一的「首席战略协办伙伴」及「算力生态独家伙伴」，享有同一级别下不可被并列、不可被超越的最高排序权益。在峰会全场景品牌露出（含主背景板、官方议程、邀请函、白皮书、官方门户、官方公众号、官方视频号、媒体通稿、回顾视频等）中，乙方 logo、企业全称及合作身份标识均位列首位。")
    clause(doc, "2.2",
           "赞助标的额：乙方为本次峰会提供专项赞助资金，总计金额为人民币（大写）叁拾万元整（小写：￥300,000.00 元），含税。")
    clause(doc, "2.3",
           "全包制商业对价：双方明确确认，本协议第三条项下乙方享有的全部权益、物料、席位、宣发、长效圈层、算力延伸等综合内容，作为「一揽子综合对价整体交付」（All-in Bundled Deliverables），不再就单项权益进行分项估值与单独定价；任何一方不得就单项权益要求拆分估价、单独退款或单独补差。该等约定为本次合作的核心商业前提，亦构成行业标杆类合作的通用商业惯例之一。")
    clause(doc, "2.4",
           "资金定向用途：上述人民币 300,000 元赞助款项仅用于且全部用于本次峰会当期的专项支出，包括但不限于：")
    para(doc, "（1）北外滩核心地标场地租赁与场地保障费用；")
    para(doc, "（2）主会场高规格舞美设计、视觉搭建与现场技术支持费用；")
    para(doc, "（3）顶尖政企大咖、智库专家及核心嘉宾的接待、出行与安保运营费用；")
    para(doc, "（4）峰会当期主视觉物料制作、官方传播、白皮书印制、回顾视频制作与现场运营所产生的其他直接费用。")
    clause(doc, "2.5",
           "排他承诺：在本次峰会筹备、执行及后续传播全过程中，甲方不得引入与乙方在云计算、大模型、智能算力调度、AI 基础设施领域形成直接竞争关系的同级别合作方（含但不限于其他头部公有云、智算云、大模型基础设施服务商等）；如需引入相关领域非同级合作方（如黄金、铂金、基础曝光等），应事先以书面形式告知乙方。")
    blank(doc)

    # ====== 三、甲方权益交付明细 ======
    heading(doc, "三、甲方权益交付明细（全包制 · 不作分项标价）", level=1, color=DARK_BLUE)
    para(doc,
         "为对应乙方上述赞助投入，甲方在本次峰会及其后续生态互动中，向乙方综合交付以下顶级权益。双方再次确认：下列权益作为一揽子综合对价整体交付，不再就单项权益进行分项估值与单独定价。如双方对个别细项落地执行需另作调整，应经书面确认（含电子邮件、企业微信、钉钉等电子形式），但不影响第二条 2.2 之总赞助金额。")
    blank(doc)

    heading(doc, "（一）核心现场曝光与顶级排序", level=2, color=DARK)
    add_table(doc,
        headers=["序号", "权益内容", "数量 / 规格", "责任方"],
        rows=[
            ["1", "大会主背景板「首席战略协办伙伴」顶级 logo 位（并排首位）", "1 处（设计稿以甲方主背景板模板为准）", "甲方"],
            ["2", "官方议程手册、邀请函、白皮书、回顾视频、官方门户、官方公众号统称权与 logo 露出", "全套（峰会前后周期）", "甲方"],
            ["3", "现场签到背板、立体导视、嘉宾胸卡、主舞台 LED 顶级位 logo 露出", "全套", "甲方 / 活动公司"],
            ["4", "媒体通稿标题级或副标题级冠名「腾讯云联合呈现 / 算力生态独家伙伴」", "全部官方对外通稿", "甲方媒体组"],
            ["5", "现场签到台联名定制（含「腾讯云 × 2026 AI 商业化峰会」共同体标识）", "1 套", "甲方 / 活动公司"],
        ],
        col_widths=[1.2, 8.5, 5.0, 3.0],
    )
    blank(doc)

    heading(doc, "（二）议题主导与高管发声", level=2, color=DARK)
    add_table(doc,
        headers=["序号", "权益内容", "数量 / 规格", "责任方"],
        rows=[
            ["6", "主会场核心高管独立主旨演讲席位", "1 个（≥ 15 分钟，议题与高管以双方协商确定）", "甲方提供位置 / 乙方主讲"],
            ["7", "「AI 硬核圆桌：技术突围与新质生产力落地」或「云端生态与 AI 商业化天花板」巅峰对话核心席位", "1 个圆桌席位（由乙方核心高管出席）", "甲方提供位置 / 乙方主讲"],
            ["8", "担任「第三届 2026 人工智能商业化落地颁奖典礼」战略支持单位 + 参与颁奖", "1 次（含联合颁奖致辞 + 媒体合影）", "甲方"],
            ["9", "白皮书章节联合发布署名", "1 章（章节主题与内容方向由双方协商，乙方提供章节素材）", "甲方编委会 / 乙方供稿"],
        ],
        col_widths=[1.2, 8.5, 5.0, 3.0],
    )
    blank(doc)

    heading(doc, "（三）顶级嘉宾接待与晚宴主桌", level=2, color=DARK)
    add_table(doc,
        headers=["序号", "权益内容", "数量 / 规格", "责任方"],
        rows=[
            ["10", "VIP 闭门晚宴主桌核心席位 + 联合主办方致祝酒辞", "主桌 6 人（最终名单提前 5 个工作日互相确认）", "甲方"],
            ["11", "重量级嘉宾 1V1 闭门沟通定向引荐", "≥ 5 位（涵盖头部算力需求方、高端汽车流通集团、头部地产与城市运营集团、金融及产业资本机构）", "甲方"],
            ["12", "晚宴期间 LED 大屏品牌片轮播", "1 段（≤ 60 秒循环）", "甲方上屏 / 乙方供片"],
            ["13", "晚宴桌卡、菜单、席卡的「首席战略协办伙伴」字样植入", "全场所有桌卡 / 菜单 / 席卡", "活动公司印刷"],
        ],
        col_widths=[1.2, 8.5, 5.0, 3.0],
    )
    blank(doc)

    heading(doc, "（四）现场展示与品牌渗透", level=2, color=DARK)
    add_table(doc,
        headers=["序号", "权益内容", "数量 / 规格", "responsible"],
        rows=[
            ["14", "主会场 1 号位独家品牌展位（大型 / 顶级位）", "1 处（不少于 24㎡，含品牌包装 / 桌台 / 互动设备位）", "甲方场地 / 活动公司展位包装"],
            ["15", "现场算力 / 大模型 demo 体验区（如适用）", "1 处（位置由双方协商，乙方可派驻技术人员现场演示）", "甲方场地 / 乙方运营"],
            ["16", "500 份现场手拎袋顶级合作伙伴位（logo 印刷 + 物料夹页）", "500 套（袋型与印刷规格以甲方模板为准）", "印刷单位 / 活动公司装袋"],
            ["17", "主舞台中场休息时段品牌片轮播", "≤ 60 秒，2 次循环", "甲方上屏"],
            ["18", "议程手册扉页整版广告 + 内页章节首页 logo 露出", "1 处扉页整版 + 4 处章节首页", "甲方设计 / 印刷单位"],
        ],
        col_widths=[1.2, 8.5, 5.0, 3.0],
    )
    blank(doc)

    heading(doc, "（五）媒体宣发与回顾沉淀", level=2, color=DARK)
    add_table(doc,
        headers=["序号", "权益内容", "数量 / 规格", "责任方"],
        rows=[
            ["19", "大会回顾视频片头独家鸣谢 + 片尾 logo 墙重点位", "1 条 1920×1080 mp4", "甲方"],
            ["20", "朋友圈九宫格冠名图（至少 3 张含乙方 logo / KV）", "9 张 1080×1080 PNG", "甲方媒体组"],
            ["21", "官方公众号、视频号峰会专题文章 / 短视频联合署名", "≥ 3 篇 / 条", "甲方媒体组"],
            ["22", "现场摄影 / 摄像精修图 + 短视频素材交付乙方品牌部使用", "≥ 100 张精修图 + ≥ 5 条短视频", "甲方"],
            ["23", "大会结束后 7 个自然日内出具《赞助权益执行回执》（含现场图片、媒体链接、嘉宾合影、媒体监测）", "1 份", "甲方"],
        ],
        col_widths=[1.2, 8.5, 5.0, 3.0],
    )
    blank(doc)

    heading(doc, "（六）算力生态自带延伸权益（不计入第二条赞助标的额）", level=2, color=DARK)
    add_table(doc,
        headers=["序号", "权益内容", "数量 / 规格", "责任方"],
        rows=[
            ["24", "向甲方生态内符合资质的科创企业与「超级个体」定向发放腾讯云专属算力包及大模型 API 调用支持", "由乙方按其内部资源池与生态准入标准自主审定", "乙方"],
            ["25", "甲方通过自有渠道（公众号、视频号、白皮书、生态社群）配合发布与宣传", "≥ 2 轮联合宣发", "甲方"],
        ],
        col_widths=[1.2, 8.5, 5.0, 3.0],
    )
    para(doc,
         "说明：本（六）项作为乙方首席算力生态伙伴身份的自带延伸权益，由乙方依其自有商业政策与资源池独立执行；不构成乙方对甲方或任何第三方的固定服务承诺，亦不计入本协议第二条项下的赞助标的额。",
         size=10, color=GREY)
    blank(doc)

    heading(doc, "（七）智库长效背书（不计入第二条赞助标的额）", level=2, color=DARK)
    add_table(doc,
        headers=["序号", "权益内容", "数量 / 规格", "责任方"],
        rows=[
            ["26", "「见微知海新质商业生态」首批年度战略理事单位身份", "永久（首批入册）", "甲方"],
            ["27", "复旦大学住房政策研究中心宏观经济与产业研究成果共享", "年度研究简报 + 重大议题定向通报", "甲方"],
            ["28", "后续长三角系列闭门局、专题沙龙、行业研究发布会的优先参与权", "1 年（自本次峰会结束之日起）", "甲方"],
            ["29", "双校长三角校友产业联盟战略合作伙伴永久入册", "永久", "甲方"],
        ],
        col_widths=[1.2, 8.5, 5.0, 3.0],
    )
    para(doc,
         "说明：本（七）项作为乙方战略级合作身份的自带长效权益，不计入本协议第二条项下的赞助标的额；如乙方希望进一步深化年度战略理事单位合作，可由双方另行签订年度合作框架协议。",
         size=10, color=GREY)
    blank(doc)

    # ====== 四、物料交付与执行时间节点 ======
    heading(doc, "四、物料交付与执行时间节点", level=1, color=DARK_BLUE)
    para(doc,
         "双方共同确认以下关键时间节点。如任一方因不可抗力以外的原因延迟，应提前 24 小时书面通知对方并协商解决：")
    add_table(doc,
        headers=["日期", "节点", "乙方动作", "甲方动作"],
        rows=[
            ["T+0（签约日）", "本协议签署生效", "盖章后协议回传", "出具盖章合同 + 银行账户"],
            ["T+3 个工作日", "物料源文件互交", "矢量 logo（AI/EPS/SVG）+ 高管简介 + 演讲主题方向 + 品牌片素材 + 算力体验 demo 准备清单", "晚宴 KV / 主背景板 / 议程手册 / 邀请函 / 白皮书设计稿初稿"],
            ["T+5 个工作日", "首席演讲与圆桌确认", "确认高管出席与议题方向", "确认时长与位置；提供舞台技术参数"],
            ["T+7 个工作日", "设计稿确认", "确认设计稿一轮", "完成印刷物料终稿"],
            ["T+10 个工作日", "印刷下单 & 物流", "—", "议程手册、白皮书、签到背板、晚宴桌卡/菜单/席卡、展位 KV 等全部下印"],
            ["大会前 3 日", "现场布展 & 彩排", "派驻品牌与技术对接人员到场", "活动公司完成主背景板、签到背板、展位、LED 调试"],
            ["大会当日", "峰会执行", "高管演讲 + 圆桌 + 颁奖 + 晚宴 + 展位现场对接", "议程执行 + 主持口播 + 宣传片轮播 + 媒体协调"],
            ["大会后 7 个自然日内", "宣发 & 执行回执", "同步乙方自有渠道二次宣发", "媒体通稿 / 九宫格发布；出具《赞助权益执行回执》"],
        ],
        col_widths=[3.0, 3.5, 5.5, 5.0],
    )
    blank(doc)

    # ====== 五、款项支付与发票 ======
    heading(doc, "五、款项支付与发票", level=1, color=DARK_BLUE)
    clause(doc, "5.1",
           "乙方应于本协议签署并经双方盖章生效后 15 个工作日内，将上述款项一次性汇入甲方指定的对公账户（账户信息详见本协议第一条「甲方主体信息」项下）。")
    clause(doc, "5.2",
           "甲方在款项到账后 10 个工作日内向乙方开具等额合法有效的增值税普通发票或专用发票（发票内容：会议服务费 / 赞助费）。")
    clause(doc, "5.3",
           "本协议项下任何因银行手续费、税费产生的费用，由各自一方承担。")
    blank(doc)

    # ====== 六、甲方义务 ======
    heading(doc, "六、甲方义务", level=1, color=DARK_BLUE)
    clause(doc, "6.1",
           "严格按本协议第三条约定，向乙方完整、及时交付各项权益，并保障峰会按计划如期举办，维护峰会的行业影响力与规格档次。")
    clause(doc, "6.2",
           "在峰会筹备、执行、后续传播全过程中，确保乙方作为「首席战略协办伙伴」及「算力生态独家伙伴」的排他性与唯一性，不引入与乙方在云计算及算力领域形成直接竞争关系的同级别合作方。")
    clause(doc, "6.3",
           "甲方应在本协议签订后 3 个工作日内向乙方提供：（1）物料规格清单（含 logo / KV / 宣传片 / PPT / 折页等技术参数与提交规范）；（2）主背景板及晚宴 LED 屏 / 投影准确像素比及刷新率参数；（3）专人对接物料交付与上线（含微信群、邮件双通道）。")
    clause(doc, "6.4",
           "妥善管理乙方支付的赞助资金，确保专款专用于本次峰会当期支出；并妥善保管乙方提供的物料、商标及其他商业资料，未经乙方书面同意，不得用于本次峰会之外的任何用途。")
    clause(doc, "6.5",
           "甲方对其所发布的媒体通稿、朋友圈九宫格、回顾视频等传播物料的合法合规性负责，发布前应将涉及乙方品牌名称、商标、核心宣传口径的内容提交乙方品牌部书面确认。")
    clause(doc, "6.6",
           "甲方应根据《中华人民共和国广告法》《中华人民共和国反不正当竞争法》之规定，在任何宣传物料中真实陈述本次合作的级别及合作身份，不得作出禁止性、误导性表述。")
    blank(doc)

    # ====== 七、乙方义务 ======
    heading(doc, "七、乙方义务", level=1, color=DARK_BLUE)
    clause(doc, "7.1",
           "按本协议约定按时足额支付赞助款项。")
    clause(doc, "7.2",
           "配合甲方完成 logo、高管简介、演讲主题、品牌片、白皮书章节素材等权益落地所需的素材提供与确认工作；并保证其提供的全部物料、商标、文字、图片、视频、PPT 等不侵犯任何第三方知识产权、商业秘密、肖像权及其他合法权益。")
    clause(doc, "7.3",
           "如因乙方提供物料导致甲方或第三方主张权利，由乙方独立承担全部法律责任并赔偿甲方因此遭受的全部直接损失（含但不限于诉讼费、律师费、合理调查费等）。")
    clause(doc, "7.4",
           "不得利用本次峰会平台从事违反国家法律法规及公序良俗的活动；遵守峰会现场及晚宴的相关秩序与流程安排，配合甲方完成议程互动。")
    clause(doc, "7.5",
           "乙方在乙方自有渠道（含官方微信、视频号、行业活动）二次宣发本次合作内容时，使用甲方提供的统一宣传口径，并确保使用的甲方名称、商标、嘉宾形象等元素已取得甲方书面同意。")
    blank(doc)

    # ====== 八、知识产权、保密及合规条款 ======
    heading(doc, "八、知识产权、保密及合规条款", level=1, color=DARK_BLUE)
    clause(doc, "8.1",
           "知识产权与品牌使用：乙方授权甲方在本次峰会的宣传、报道、白皮书、官网、大屏、议程手册、媒体通稿、回顾视频等场景中使用乙方「腾讯云」商标与企业简介；授权范围严格限于本次峰会及其衍生宣传内容，授权期限为本协议签订之日起至本次峰会结束后 12 个月。甲方拥有本次峰会名称、视觉系统、白皮书、回顾视频等衍生品的完整知识产权；乙方在使用大会名称、Logo、嘉宾形象等元素进行二次宣传前，应事先取得甲方书面同意（含微信书面同意）。")
    clause(doc, "8.2",
           "成果物归属：本协议项下双方就合作过程中形成的现场摄影 / 摄像、回顾视频、白皮书署名作品等成果物，知识产权归甲方所有；乙方享有非独家、不可转授权的免费使用权，使用范围限于乙方品牌内部宣传与生态推广。")
    clause(doc, "8.3",
           "商业机密保护：鉴于本协议体现了甲乙双方最核心的战略合作对价，未经甲乙双方书面同意，任何一方不得向任何第三方（包括但不限于其他赞助商、合作方、媒体及外部审计机构）披露本协议项下的资金打款凭证、财务流水、内部核算细则、单项权益估值（如有）等敏感信息。保密期限自本协议签订之日起 3 年。")
    clause(doc, "8.4",
           "独立比对声明：甲方有权在对外招商、生态展示及向同级别合作方进行横向拉齐时，合法引用本协议第二条项下的合作总金额（叁拾万元）及合作级别（首席战略协办伙伴 + 算力生态独家伙伴）作为优秀案例说明，以证明本次峰会的市场估值与顶尖行业含金量。该等引用不视为对本条第 8.3 款保密义务的违反，但该引用不得包含本协议项下的资金打款凭证、财务流水、内部核算细则等敏感信息。")
    clause(doc, "8.5",
           "个人信息合规：双方依照《中华人民共和国个人信息保护法》《中华人民共和国数据安全法》之规定处理本次合作中涉及的个人信息；未经嘉宾事先同意，不得将嘉宾联系方式用于本次峰会以外的任何商业推广。乙方对其通过展位、demo 体验区收集到的意向客户个人信息，应单独取得当事人书面或电子告知同意，并不得提供给本协议双方以外的第三方。")
    clause(doc, "8.6",
           "合规承诺：双方承诺，本协议项下的资金往来与权益交付均合法合规，不存在任何商业贿赂、利益输送及违反《中华人民共和国反不正当竞争法》《关于禁止商业贿赂行为的暂行规定》等相关规定的情形；不得直接或间接向对方工作人员（含其近亲属、关联方）给予任何形式的回扣、佣金、礼品、招待等不正当利益。")
    blank(doc)

    # ====== 九、违约责任 ======
    heading(doc, "九、违约责任", level=1, color=DARK_BLUE)
    clause(doc, "9.1",
           "任何一方未按本协议约定履行义务，给对方造成损失的，应承担相应的违约责任并赔偿对方因此遭受的直接经济损失。")
    clause(doc, "9.2",
           "乙方未按约定时间支付赞助款项的，每逾期一日按未付款项的 0.5% 向甲方支付违约金；逾期超过 15 日的，甲方有权解除本协议并不退还任何已收款项（如有）。")
    clause(doc, "9.3",
           "甲方未按本协议第三条 / 第四条提供权益且无正当理由的，应按未履行权益对应价值（双方协商认定）的 100% 向乙方退款，或在双方协商一致后以同等价值的下一届同类活动权益补偿。鉴于本协议为「全包制不分项标价」，如双方对未履行权益的对应价值存在争议，应参照《附件一》各类别整体占比进行合理认定。")
    clause(doc, "9.4",
           "任何一方违反本协议第八条（知识产权、保密及合规）的，应向守约方支付不低于人民币 100,000 元的违约金；造成实际损失超过违约金的，超过部分应另行赔偿。")
    blank(doc)

    # ====== 十、不可抗力 ======
    heading(doc, "十、不可抗力", level=1, color=DARK_BLUE)
    clause(doc, "10.1",
           "若因不可抗力（包括但不限于地震、洪水、台风等自然灾害；战争、暴乱、罢工；政府管制、重大公共卫生事件、电力中断、网络中断等）导致本次峰会延期或无法举办，双方互不承担违约责任。")
    clause(doc, "10.2",
           "双方应友好协商，按已发生的实际支出比例处理已支付的赞助款项；剩余款项可顺延用于甲方下一届同级别峰会，或经双方书面同意后退还乙方。")
    clause(doc, "10.3",
           "遭受不可抗力的一方应在事件发生后 5 个自然日内书面通知对方，并提供有效证明文件。")
    blank(doc)

    # ====== 十一、争议解决 ======
    heading(doc, "十一、争议解决与法律适用", level=1, color=DARK_BLUE)
    clause(doc, "11.1",
           "本协议的签订、履行、解释及争议解决均适用中华人民共和国法律。")
    clause(doc, "11.2",
           "因本协议产生或与本协议有关的任何争议，双方应首先通过友好协商解决；协商不成的，任何一方均有权向甲方所在地有管辖权的人民法院提起诉讼。")
    blank(doc)

    # ====== 十二、签署与生效 ======
    heading(doc, "十二、签署与生效", level=1, color=DARK_BLUE)
    clause(doc, "12.1",
           "生效条件：本协议自甲乙双方法定代表人或授权代表签字并加盖公章（或合同专用章）之日起生效。")
    clause(doc, "12.2",
           "落款与盖章：甲方加盖组委会公章或复旦大学住房政策研究中心公章；乙方加盖公章或合同专用章。")
    clause(doc, "12.3",
           "物理要求：本协议正本须加盖骑缝章，一式贰份，甲乙双方各执壹份，具有同等法律效力。")
    clause(doc, "12.4",
           "附件效力：本协议如有附件、补充协议或经双方书面确认的往来函件（含电子邮件、企业微信、钉钉等电子形式），均为本协议不可分割的组成部分，与本协议具有同等法律效力。本协议附件包括：《附件一·首席战略协办伙伴顶级权益交付执行清单》《附件二·物料规格与提交规范》《附件三·算力生态延伸权益执行细则》。")
    clause(doc, "12.5",
           "电子签章：双方通过电子邮件、企业微信、钉钉等电子形式签署的扫描件 / 加盖电子签章的文件，与纸质原件具有同等法律效力。")
    blank(doc)

    # ====== 签署页 ======
    doc.add_page_break()
    heading(doc, "签 署 页（Signature Page）", level=0, color=DARK_BLUE, center=True)
    para(doc,
         "（以下无正文，本页为《首席战略协办伙伴合作协议》之签署页）",
         align="center", first_indent=False, size=10, color=GREY)
    blank(doc, 2)

    sig = doc.add_table(rows=6, cols=2)
    sig.style = "Table Grid"
    sig.autofit = False
    for row in sig.rows:
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(8.0)
    set_cell_text(sig.rows[0].cells[0], "甲方（盖章）：",
                  bold=True, size=11)
    set_cell_text(sig.rows[0].cells[1], "乙方（盖章）：",
                  bold=True, size=11)
    set_cell_text(sig.rows[1].cells[0],
                  "人工智能商业化落地峰会组委会 /\n复旦大学住房政策研究中心 /\n上海市杨浦区科技企业联合会",
                  size=10.5)
    set_cell_text(sig.rows[1].cells[1],
                  "腾讯云计算（北京）有限责任公司",
                  size=10.5)
    set_cell_text(sig.rows[2].cells[0],
                  "\n\n\n（公章 / 合同章位置）\n\n",
                  size=10.5, align_center=True)
    set_cell_text(sig.rows[2].cells[1],
                  "\n\n\n（公章 / 合同章位置）\n\n",
                  size=10.5, align_center=True)
    set_cell_text(sig.rows[3].cells[0],
                  "授权代表（签字）：______________________",
                  size=11)
    set_cell_text(sig.rows[3].cells[1],
                  "授权代表（签字）：______________________",
                  size=11)
    set_cell_text(sig.rows[4].cells[0], "职务：__________________", size=11)
    set_cell_text(sig.rows[4].cells[1], "职务：__________________", size=11)
    set_cell_text(sig.rows[5].cells[0],
                  "签字日期：______年______月______日",
                  size=11)
    set_cell_text(sig.rows[5].cells[1],
                  "签字日期：______年______月______日",
                  size=11)

    # ====== 附件一 ======
    doc.add_page_break()
    heading(doc, "附件一  首席战略协办伙伴顶级权益交付执行清单",
            level=0, color=DARK_BLUE, center=True)
    para(doc,
         "Annex 1 — Top Sponsor Rights Delivery Checklist",
         align="center", size=10, color=GREY, first_indent=False)
    blank(doc)
    para(doc,
         "本附件列示本协议第三条项下乙方享有的全部顶级权益的交付明细，作为甲方履约及乙方验收的依据。**本附件不构成对单项权益的分项估值，亦不作为单项退款依据；具体退款或补偿按本协议第 9.3 条处理。**")
    blank(doc)

    add_table(doc,
        headers=["类别", "权益项目", "数量 / 规格", "责任方", "双方确认"],
        rows=[
            ["一、核心现场曝光", "主背景板顶级 logo 并排首位露出", "1 处（首位）", "甲方", "□ 已确认"],
            ["", "议程手册 / 邀请函 / 白皮书 / 回顾视频 / 官方门户 / 官方公众号统称权与 logo 露出", "全套", "甲方", "□ 已确认"],
            ["", "现场签到背板 / 立体导视 / 嘉宾胸卡 / 主舞台 LED 顶级位 logo 露出", "全套", "甲方 / 活动公司", "□ 已确认"],
            ["", "媒体通稿标题或副标题级冠名", "全部官方对外通稿", "甲方媒体组", "□ 已确认"],
            ["", "现场签到台联名定制", "1 套", "甲方 / 活动公司", "□ 已确认"],
            ["二、议题主导", "主会场核心高管独立主旨演讲席位", "1 个（≥ 15 min）", "甲方位置 / 乙方主讲", "□ 已确认"],
            ["", "圆桌或巅峰对话核心席位", "1 个", "甲方位置 / 乙方主讲", "□ 已确认"],
            ["", "颁奖典礼战略支持单位 + 联合颁奖", "1 次", "甲方", "□ 已确认"],
            ["", "白皮书章节联合署名", "1 章", "甲方编委会 / 乙方供稿", "□ 已确认"],
            ["三、晚宴主桌", "VIP 闭门晚宴主桌 + 联合祝酒辞", "主桌 6 人", "甲方", "□ 已确认"],
            ["", "重量级嘉宾 1V1 闭门沟通定向引荐", "≥ 5 位", "甲方", "□ 已确认"],
            ["", "晚宴 LED 大屏品牌片轮播", "≤ 60s 循环", "甲方 / 乙方供片", "□ 已确认"],
            ["", "晚宴桌卡 / 菜单 / 席卡身份字样植入", "全场", "活动公司印刷", "□ 已确认"],
            ["四、现场展示", "主会场 1 号位独家品牌展位", "1 处 ≥ 24㎡", "甲方场地 / 活动公司", "□ 已确认"],
            ["", "算力 / 大模型 demo 体验区", "1 处", "甲方场地 / 乙方运营", "□ 已确认"],
            ["", "500 份手拎袋顶级合作伙伴位 + 夹页", "500 套", "印刷单位 / 活动公司", "□ 已确认"],
            ["", "主舞台中场休息品牌片轮播", "≤ 60s × 2 次", "甲方", "□ 已确认"],
            ["", "议程手册扉页整版广告 + 章节首页 logo", "扉页 1 + 章首 4", "甲方设计 / 印刷单位", "□ 已确认"],
            ["五、媒体宣发", "大会回顾视频片头独家鸣谢", "1 条", "甲方", "□ 已确认"],
            ["", "朋友圈九宫格冠名图", "9 张 1080×1080", "甲方媒体组", "□ 已确认"],
            ["", "官方公众号 / 视频号专题文章 / 短视频联合署名", "≥ 3 篇 / 条", "甲方媒体组", "□ 已确认"],
            ["", "现场摄影精修图 + 短视频素材交付", "≥ 100 张 + ≥ 5 条", "甲方", "□ 已确认"],
            ["", "《赞助权益执行回执》", "7 个自然日内 1 份", "甲方", "□ 已确认"],
            ["六、算力延伸（自带）", "腾讯云专属算力包 + 大模型 API 调用支持", "依乙方资源池", "乙方", "□ 已确认"],
            ["", "甲方自有渠道联合宣发", "≥ 2 轮", "甲方", "□ 已确认"],
            ["七、智库长效（自带）", "「见微知海新质商业生态」首批年度战略理事单位", "永久（首批）", "甲方", "□ 已确认"],
            ["", "复旦大学住房政策研究中心研究成果共享", "年度", "甲方", "□ 已确认"],
            ["", "长三角系列闭门局 / 沙龙优先参与权", "1 年", "甲方", "□ 已确认"],
            ["", "双校长三角校友产业联盟战略合作伙伴永久入册", "永久", "甲方", "□ 已确认"],
        ],
        col_widths=[3.0, 6.0, 3.5, 3.0, 2.5],
    )
    blank(doc)
    para(doc,
         "本附件经双方签字盖章后与本协议正文具有同等法律效力。",
         size=10, color=GREY)
    blank(doc)
    sig2 = doc.add_table(rows=1, cols=2)
    sig2.style = "Table Grid"
    sig2.autofit = False
    for row in sig2.rows:
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(8.0)
    set_cell_text(sig2.rows[0].cells[0],
                  "甲方授权代表（签字 / 盖章）：\n\n\n日期：______年______月______日",
                  bold=True, size=11)
    set_cell_text(sig2.rows[0].cells[1],
                  "乙方授权代表（签字 / 盖章）：\n\n\n日期：______年______月______日",
                  bold=True, size=11)

    # ====== 附件二 ======
    doc.add_page_break()
    heading(doc, "附件二  物料规格与提交规范",
            level=0, color=DARK_BLUE, center=True)
    para(doc, "Annex 2 — Material Specifications",
         align="center", size=10, color=GREY, first_indent=False)
    blank(doc)
    add_table(doc,
        headers=["序号", "物料 / 素材", "比例", "推荐尺寸 / 分辨率", "时长 / 帧率 / dpi", "备注"],
        rows=[
            ["1", "主背景板顶级 logo 位", "—", "矢量 AI / EPS / SVG（无白底）", "—", "另出 PNG 透明底 4 套尺寸"],
            ["2", "主会场宣传片 / 品牌片", "16:9", "1920×1080（4K 备份 3840×2160）", "≤ 60s，25/30 fps", "H.264 .mp4 ≤ 200MB"],
            ["3", "晚宴 LED 大屏品牌片", "16:9（条屏另出）", "1920×1080 或 6144×1080", "≤ 60s 循环", "请场地方提供精确像素比"],
            ["4", "高管主旨演讲 PPT", "16:9", "1920×1080（40×22.5cm）", "≤ 25 张，控制 15 分钟内", "封面/封底/产品截图预留"],
            ["5", "议程手册扉页整版广告", "—", "整版 A4 210×285mm，出血 3mm", "CMYK，300 dpi", "印刷单位统一排版"],
            ["6", "议程手册章节首页 logo", "—", "矢量", "—", "印刷品；4 处章节首页"],
            ["7", "白皮书章节联合署名稿", "—", "Word / PPT 源文件", "—", "由乙方提供章节内容，甲方编委会统一排版"],
            ["8", "现场展位 KV 主视觉", "—", "1920×1080 主稿，另按展位实际尺寸延展", "—", "RGB；提供 AI 源 + PNG"],
            ["9", "现场展位易拉宝", "—", "800×2000mm，出血 5mm", "CMYK 150 dpi", "推荐 4–6 张组合"],
            ["10", "现场展位户型 / 产品 KT 板", "—", "依展位规划尺寸", "CMYK 150 dpi", "由活动公司统一打印"],
            ["11", "晚宴 KV / 桌卡 / 菜单 / 席卡 logo", "—", "矢量 logo + 静帧 PNG", "—", "活动公司统一植入"],
            ["12", "朋友圈九宫格", "1:1", "1080×1080 PNG × 9", "—", "媒体组发布；乙方品牌部同步"],
            ["13", "回顾视频片头鸣谢", "16:9", "1920×1080 mp4", "≤ 5min", "片头 3s 独家鸣谢 + 片尾 logo 墙"],
            ["14", "嘉宾胸卡 / 现场签到背板", "—", "印刷品", "300 dpi", "印刷单位统一制作"],
            ["15", "500 份手拎袋", "—", "印刷品", "300 dpi", "印刷单位 + 活动公司装袋"],
        ],
        col_widths=[1.0, 4.0, 2.0, 4.0, 2.8, 3.0],
    )
    blank(doc)

    # ====== 附件三 ======
    doc.add_page_break()
    heading(doc, "附件三  算力生态延伸权益执行细则",
            level=0, color=DARK_BLUE, center=True)
    para(doc, "Annex 3 — Tencent Cloud Compute Ecosystem Extension Rules",
         align="center", size=10, color=GREY, first_indent=False)
    blank(doc)

    para(doc,
         "本附件用于约定本协议第三条（六）项下乙方自带的「算力生态延伸权益」的具体执行机制。本附件项下乙方所提供的算力资源与大模型 API 调用支持，均为乙方依其内部商业政策与生态准入标准自主审定与发放，不构成本协议第二条赞助标的额之外的额外付费义务，亦不构成乙方对甲方或任何第三方的固定服务承诺。")
    blank(doc)

    clause(doc, "Ⅰ", "发放范围：")
    para(doc,
         "甲方生态内符合以下任一资质的科创企业或「超级个体」，可由甲方推荐、乙方独立审定后纳入本附件项下的资源发放范围：")
    para(doc, "（1）国家高新技术企业 / 专精特新企业；")
    para(doc, "（2）已完成天使轮以上融资的 AI / 大模型 / 算力领域创业公司；")
    para(doc, "（3）入选「见微知海新质商业生态」首批年度评估的优秀「超级个体」；")
    para(doc, "（4）双校长三角校友产业联盟内部认定的会员企业。")
    blank(doc)

    clause(doc, "Ⅱ", "资源类型与额度（指导值，最终以乙方资源池实时供给为准）：")
    add_table(doc,
        headers=["资源类型", "建议发放额度 / 企业", "使用周期", "适用方向"],
        rows=[
            ["腾讯云通用计算资源券", "￥3,000 – ￥10,000 / 企业", "12 个月", "弹性云主机、对象存储等"],
            ["腾讯云智算 / GPU 算力券", "￥5,000 – ￥30,000 / 企业", "12 个月", "AI 训练 / 推理"],
            ["大模型 API 调用券", "1–10 万 tokens / 企业", "12 个月", "大模型应用开发与商业化验证"],
            ["开发者技术支持", "1V1 解决方案咨询 1–2 次", "活动后 60 天内", "云架构设计、降本增效咨询"],
        ],
        col_widths=[4.0, 4.0, 3.0, 5.0],
    )
    blank(doc)

    clause(doc, "Ⅲ", "执行流程：")
    para(doc, "（1）甲方在峰会结束后 14 个自然日内，向乙方推荐意向企业名单（含企业基本信息、联系人、需求方向）；")
    para(doc, "（2）乙方在收到名单后 30 个自然日内，按其内部生态准入标准独立审定，并直接与入选企业对接资源发放与服务条款；")
    para(doc, "（3）甲方在乙方完成发放后 7 个自然日内，通过自有渠道（公众号、视频号、白皮书、生态社群）配合发布联合宣发文章不少于 2 篇；")
    para(doc, "（4）乙方有权根据其自身商业政策调整或终止本附件项下任一资源的发放，但应提前 10 个自然日书面通知甲方。")
    blank(doc)

    clause(doc, "Ⅳ", "数据与个人信息保护：")
    para(doc, "甲方推荐入选企业的相关信息（含企业基本信息、联系人）须经入选企业同意后方可提供给乙方；乙方对该等信息承担保密义务，仅用于本附件项下资源发放与合作沟通，不得用于其他商业用途。")
    blank(doc)

    clause(doc, "Ⅴ", "效力：")
    para(doc, "本附件与本协议正文具有同等法律效力；本附件项下事项的争议解决适用本协议第十一条之约定。")

    blank(doc, 2)
    sig3 = doc.add_table(rows=1, cols=2)
    sig3.style = "Table Grid"
    sig3.autofit = False
    for row in sig3.rows:
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(8.0)
    set_cell_text(sig3.rows[0].cells[0],
                  "甲方授权代表（签字 / 盖章）：\n\n\n日期：______年______月______日",
                  bold=True, size=11)
    set_cell_text(sig3.rows[0].cells[1],
                  "乙方授权代表（签字 / 盖章）：\n\n\n日期：______年______月______日",
                  bold=True, size=11)

    # ====== 附件四：赞助权益执行回执模板 ======
    doc.add_page_break()
    heading(doc, "附件四  赞助权益执行回执模板",
            level=0, color=DARK_BLUE, center=True)
    para(doc, "Annex 4 — Sponsorship Rights Delivery Receipt Template",
         align="center", size=10, color=GREY, first_indent=False)
    blank(doc)
    para(doc,
         "本附件用于规范甲方在本次峰会结束后 7 个自然日内向乙方出具的《赞助权益执行回执》的内容、格式与证据材料标准，作为甲方履约验收的统一依据。")
    blank(doc)

    heading(doc, "一、回执基本信息", level=2, color=DARK)
    add_kv_table(doc, rows=[
        ["回执编号", "AIBIZ-2026-RC-013"],
        ["对应合同编号", "AIBIZ-2026-SP-013"],
        ["乙方名称", "腾讯云计算（北京）有限责任公司"],
        ["合作身份", "首席战略合作伙伴 + 算力生态独家伙伴"],
        ["赞助金额", "￥300,000.00（人民币叁拾万元整）"],
        ["回执出具日期", "______年______月______日"],
        ["回执签发人", "____________________（甲方组委会执行秘书长签字 + 公章）"],
    ])
    blank(doc)

    heading(doc, "二、权益交付确认表（与附件一对位）", level=2, color=DARK)
    add_table(doc,
        headers=["类别", "权益项目", "执行状态", "证据材料类型", "证据材料编号"],
        rows=[
            ["一、核心现场曝光", "主背景板顶级 logo 并排首位", "□ 已交付 / □ 未交付", "现场背板照片", "P-001"],
            ["", "议程 / 邀请函 / 白皮书 / 回顾视频统称权", "□ 已交付 / □ 未交付", "成品扫描件", "P-002"],
            ["", "签到背板 / 导视 / 胸卡 / 主舞台 LED", "□ 已交付 / □ 未交付", "现场照片 + 视频截图", "P-003"],
            ["", "媒体通稿冠名露出", "□ 已交付 / □ 未交付", "媒体链接", "M-001 至 M-00X"],
            ["", "签到台联名定制", "□ 已交付 / □ 未交付", "现场照片", "P-004"],
            ["二、议题主导", "主旨演讲席位", "□ 已交付 / □ 未交付", "现场照片 + 视频", "P-005 + V-001"],
            ["", "圆桌核心席位", "□ 已交付 / □ 未交付", "现场照片 + 视频", "P-006 + V-002"],
            ["", "颁奖典礼战略支持 + 联合颁奖", "□ 已交付 / □ 未交付", "现场照片", "P-007"],
            ["", "白皮书章节署名", "□ 已交付 / □ 未交付", "白皮书扫描件", "P-008"],
            ["三、晚宴主桌", "主桌 + 联合祝酒辞", "□ 已交付 / □ 未交付", "现场照片", "P-009"],
            ["", "1V1 闭门引荐 ≥ 5 位", "□ 已交付 / □ 未交付", "对接清单（脱敏）", "L-001"],
            ["", "晚宴 LED 品牌片轮播", "□ 已交付 / □ 未交付", "屏幕截图 / 视频", "V-003"],
            ["", "晚宴桌卡/菜单/席卡植入", "□ 已交付 / □ 未交付", "成品照片", "P-010"],
            ["四、现场展示", "1 号位品牌展位", "□ 已交付 / □ 未交付", "全景 + 细节照片", "P-011"],
            ["", "算力 / demo 体验区", "□ 已交付 / □ 未交付", "现场照片", "P-012"],
            ["", "500 份手拎袋顶级位 + 夹页", "□ 已交付 / □ 未交付", "成品照片", "P-013"],
            ["", "中场休息品牌片轮播", "□ 已交付 / □ 未交付", "屏幕截图", "V-004"],
            ["", "议程手册扉页 + 章节首页", "□ 已交付 / □ 未交付", "印刷成品扫描件", "P-014"],
            ["五、媒体宣发", "回顾视频片头独家鸣谢", "□ 已交付 / □ 未交付", "回顾视频链接", "V-005"],
            ["", "朋友圈九宫格冠名图", "□ 已交付 / □ 未交付", "九宫格截图", "P-015"],
            ["", "公众号/视频号专题文章 ≥3", "□ 已交付 / □ 未交付", "文章链接", "M-00X"],
            ["", "摄影精修图 + 短视频素材交付", "□ 已交付 / □ 未交付", "云盘链接", "C-001"],
            ["六、算力延伸（自带）", "腾讯云算力包 + API 调用", "□ 已发放 / □ 待发放 / □ N/A", "发放清单（脱敏）", "L-002"],
            ["", "甲方联合宣发 ≥ 2 轮", "□ 已交付 / □ 未交付", "宣发链接", "M-00X"],
            ["七、智库长效（自带）", "「见微知海」首批战略理事", "□ 已交付 / □ 未交付", "聘书扫描件", "P-016"],
            ["", "双校长三角校友联盟入册", "□ 已交付 / □ 未交付", "通讯录截屏（脱敏）", "L-003"],
        ],
        col_widths=[3.0, 6.0, 3.0, 3.5, 2.5],
    )
    blank(doc)

    heading(doc, "三、证据材料目录（与上表编号对应）", level=2, color=DARK)
    para(doc,
         "甲方在本回执出具时，应同时随附下列证据材料文件夹至乙方品牌部及法务对接人："
    )
    para(doc, "（1）P-001 ~ P-016：现场及成品照片（建议 ≥ 100 张精修图）；",
         first_indent=False)
    para(doc, "（2）V-001 ~ V-005：现场视频/截图（建议 ≥ 5 条短视频 + 屏幕截图若干）；",
         first_indent=False)
    para(doc, "（3）M-001 至 M-00X：媒体通稿链接清单（建议主流财经/科技/产业媒体 ≥ 5 家）；",
         first_indent=False)
    para(doc, "（4）L-001 ~ L-003：对接清单 / 发放清单 / 入册名单（脱敏处理后提供）；",
         first_indent=False)
    para(doc, "（5）C-001：云盘下载链接（含原始素材文件夹，30 天有效期内可下载）。",
         first_indent=False)
    blank(doc)

    heading(doc, "四、验收结论", level=2, color=DARK)
    para(doc,
         "□ 完整履约：经乙方品牌部及法务对接人确认，甲方已按本协议第三条及附件一约定全面交付各项权益，证据材料齐全，乙方同意签收本回执并视为本次合作权益交付完成。",
         first_indent=False)
    para(doc,
         "□ 部分履约：经乙方确认，甲方就以下条款未完整履行：____________________________；双方按本协议第 9.3 条另行协商处理。",
         first_indent=False)
    blank(doc, 2)

    sig4 = doc.add_table(rows=1, cols=2)
    sig4.style = "Table Grid"
    sig4.autofit = False
    for row in sig4.rows:
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(8.0)
    set_cell_text(sig4.rows[0].cells[0],
                  "甲方（盖章）：\n\n\n签发人：______________________\n日期：______年______月______日",
                  bold=True, size=11)
    set_cell_text(sig4.rows[0].cells[1],
                  "乙方（签收 / 盖章）：\n\n\n签收人：______________________\n日期：______年______月______日",
                  bold=True, size=11)

    blank(doc)
    para(doc,
         "本回执模板与本协议正文具有同等法律效力；正式回执出具后，构成本协议项下甲方履约的完成证明。",
         align="center", first_indent=False, size=10, color=GREY)

    out = "/workspace/deliverables/腾讯云-首席战略协办伙伴合作协议.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
