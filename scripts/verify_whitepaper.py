# -*- coding: utf-8 -*-
"""校验白皮书 DOCX 是否生成完整。"""

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "dist" / "河南大学住房政策研究中心_住房即服务_医疗养老与提效服务2030展望白皮书.docx"

REQUIRED_HEADINGS = [
    "住房即服务",
    "河南大学住房政策研究中心",
    "摘要：2026 年的五个判断与 2030 年的一条主线",
    "一、导论：住房政策为何必须同时讨论医疗、养老与提效",
    "二、当下形势：三重压力在居住空间里交汇",
    "三、科技与制度的交汇：2025—2026 年的关键进展",
    "四、医疗服务展望 2030：从医院能力到家庭可达",
    "五、养老服务展望 2030：从床位思维到支付—空间—机器人闭环",
    "六、生产效率服务展望 2030：工厂、社区与住房如何共同提效",
    "七、住房作为三类服务的空间操作系统",
    "八、四个情景与河南含义",
    "九、政策建议",
    "附录二  主要公开资料来源",
]

REQUIRED_KEYWORDS = [
    "2030",
    "人工智能+",
    "长护险",
    "医养结合",
    "完整社区",
    "具身智能",
    "WAIC",
    "众擎",
    "好房子",
    "家庭养老床位",
    "9785",
    "智能体",
    "WS/T 876",
    "IEC 63310",
]


def main():
    assert DOCX.exists(), f"未找到文件：{DOCX}"
    size_kb = DOCX.stat().st_size / 1024
    assert size_kb > 80, f"文件过小：{size_kb:.1f} KB"

    doc = Document(str(DOCX))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    blob = "\n".join(texts)
    char_count = len(blob.replace(" ", "").replace("\n", ""))

    missing_h = [h for h in REQUIRED_HEADINGS if h not in blob]
    missing_k = [k for k in REQUIRED_KEYWORDS if k not in blob]

    table_count = len(doc.tables)
    print(f"文件：{DOCX.name}")
    print(f"大小：{size_kb:.1f} KB")
    print(f"段落数：{len(doc.paragraphs)}")
    print(f"非空段落：{len(texts)}")
    print(f"正文字符数（约）：{char_count}")
    print(f"表格数：{table_count}")

    assert not missing_h, f"缺少标题：{missing_h}"
    assert not missing_k, f"缺少关键词：{missing_k}"
    assert char_count >= 14000, f"正文过短：{char_count}"
    assert table_count >= 4, f"表格过少：{table_count}"
    print("校验通过。")


if __name__ == "__main__":
    main()
