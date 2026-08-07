#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成更简洁清晰的百科第二版操作稿：改动处标黄。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "胡继刚_百度百科第二版操作稿_改动标黄.docx"
YELLOW = "FFFF00"


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def highlight_run(run, color=YELLOW):
    from docx.oxml import OxmlElement

    rPr = run._element.get_or_add_rPr()
    # remove existing highlight
    for child in list(rPr):
        if child.tag == qn("w:highlight"):
            rPr.remove(child)
    hl = OxmlElement("w:highlight")
    hl.set(qn("w:val"), "yellow")
    rPr.append(hl)


def add_para(doc, text, *, size=12, bold=False, space_before=0, space_after=6, color=None, align=None, first_line=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_runs(doc, segments, *, size=12, space_before=0, space_after=6, first_line=False):
    """segments: list of (text, bold, highlight)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        pf.first_line_indent = Cm(0.74)
    for text, bold, hl in segments:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
        if hl:
            highlight_run(run)
    return p


def add_h(doc, text, level=1):
    sizes = {1: 16, 2: 13}
    return add_para(
        doc,
        text,
        size=sizes.get(level, 12),
        bold=True,
        space_before=12 if level == 1 else 8,
        space_after=6,
        color=RGBColor(0x1A, 0x3A, 0x5C),
    )


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)

    add_para(doc, "胡继刚｜百度百科第二版操作稿", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, color=RGBColor(0x1A, 0x3A, 0x5C))
    add_runs(
        doc,
        [
            ("用法：", True, False),
            ("白色=原封不动保留；", False, False),
            ("黄色高亮=需要改/新加，只粘贴黄字部分或整段替换黄段。", False, True),
        ],
        size=10.5,
        space_after=4,
    )
    add_para(
        doc,
        "本版只动 3 处：①2025年经历加半句；②2026年经历/执行会长加脚注〔9〕；③上传照片。其余一律不改。",
        size=10.5,
        color=RGBColor(0x55, 0x55, 0x55),
        space_after=10,
    )

    # 操作总览
    add_h(doc, "一、你要做什么（30秒看懂）", 1)
    rows = [
        ("位置", "动作", "粘贴什么"),
        ("人物经历 → 2025年峰会那一段", "整段替换", "下方【改动1】黄段全文"),
        ("人物经历 → 2026年AI峰会那一段末尾", "只加脚注", "新增参考资料〔9〕，挂在该段后"),
        ("社会职务 → 执行会长那一行末尾", "只加脚注", "同样挂〔9〕"),
        ("词条头图 / 图片", "新上传", "见【改动3】"),
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            set_run_font(run, size=10, bold=(i == 0))
            if i > 0 and j == 1:
                highlight_run(run)
    add_para(doc, "", space_after=4)

    # 不动清单
    add_h(doc, "二、这些原封不动（不要动）", 1)
    for t in [
        "概述全文",
        "基本信息栏（中文名 / 主要职务 / 社会职务条目文字）",
        "人物经历：2024年6月锦天城段",
        "人物经历：2025年5月28日工商联段",
        "人物经历：2026年3月北欧会客厅段",
        "社会职务三条的文字本身（只给执行会长加脚注编号）",
        "已有参考资料〔1〕—〔8〕（不要删、不要改链接）",
    ]:
        add_para(doc, f"✓ {t}", size=11, space_after=2)

    # 改动1
    add_h(doc, "三、改动1｜人物经历·2025年峰会（整段替换）", 1)
    add_para(doc, "粘贴位置：百科编辑页 →「人物经历」→ 找到2025年全球新经济增长引擎峰会那一段 → 整段删掉后粘贴下面黄段。", size=10.5, color=RGBColor(0x55, 0x55, 0x55), space_after=4)
    add_para(doc, "【现有原文·不要留着，用下面替换】", bold=True, size=10.5, space_after=2)
    add_para(
        doc,
        "2025年5月，胡继刚以复旦大学住房政策研究中心秘书长身份参加“2025全球新经济增长引擎峰会”，在圆桌对话环节与证券、经济及资产管理领域嘉宾围绕房地产市场、资本布局与产业创新等议题展开交流。〔2〕〔3〕",
        size=11,
        first_line=True,
        space_after=6,
        color=RGBColor(0x66, 0x66, 0x66),
    )
    add_para(doc, "【替换为·整段标黄·复制粘贴】", bold=True, size=10.5, space_after=2)
    add_runs(
        doc,
        [
            (
                "2025年5月，胡继刚以复旦大学住房政策研究中心秘书长身份参加“2025全球新经济增长引擎峰会”，在圆桌对话环节与证券、经济及资产管理领域嘉宾围绕房地产市场、资本布局与产业创新等议题展开交流，并以主要组织方代表身份就峰会主题及住房政策研究中心后续工作方向发表观点。〔2〕〔3〕",
                False,
                True,
            )
        ],
        size=11,
        first_line=True,
        space_after=4,
    )
    add_runs(
        doc,
        [
            ("黄字新增部分是：", False, False),
            ("，并以主要组织方代表身份就峰会主题及住房政策研究中心后续工作方向发表观点", False, True),
            ("。脚注仍挂〔2〕人民网、〔3〕中新网，不用新开栏目。", False, False),
        ],
        size=10.5,
        space_after=8,
    )

    # 改动2
    add_h(doc, "四、改动2｜加脚注〔9〕（正文句子可不动）", 1)
    add_para(
        doc,
        "粘贴位置A：百科编辑页 →「人物经历」→ 2026年人工智能峰会那一段 → 在段末已有脚注后，再插入一条新参考资料〔9〕。",
        size=10.5,
        color=RGBColor(0x55, 0x55, 0x55),
        space_after=2,
    )
    add_para(
        doc,
        "粘贴位置B：百科编辑页 →「社会职务」→「上海市杨浦区科技企业联合会执行会长」那一行末尾 → 同样挂〔9〕。",
        size=10.5,
        color=RGBColor(0x55, 0x55, 0x55),
        space_after=4,
    )

    add_para(doc, "【2026年经历段·正文原封不动，仅末尾脚注变黄】", bold=True, size=10.5, space_after=2)
    add_runs(
        doc,
        [
            (
                "2026年5月，胡继刚以复旦大学住房政策研究中心秘书长、上海市杨浦区科技企业联合会执行会长身份参加“2026人工智能商业化落地与硬核投资破局峰会”，主持“从算力引擎到新质资产（AI全产业链的商业化实战）”圆桌对话，并以主办方代表身份发表观点。〔1〕〔6〕〔7〕",
                False,
                False,
            ),
            ("〔9〕", False, True),
        ],
        size=11,
        first_line=True,
        space_after=6,
    )

    add_para(doc, "【执行会长·文字不动，仅末尾脚注变黄】", bold=True, size=10.5, space_after=2)
    add_runs(
        doc,
        [
            ("上海市杨浦区科技企业联合会执行会长。〔1〕〔4〕〔7〕", False, False),
            ("〔9〕", False, True),
        ],
        size=11,
        first_line=True,
        space_after=6,
    )

    add_para(doc, "【新增参考资料〔9〕·按百科“添加参考资料”填写】", bold=True, size=10.5, space_after=2)
    add_runs(doc, [("标题：", True, False), ("AI新质风口到来 共议超级个体成长路径", False, True)], size=11, space_after=2)
    add_runs(doc, [("来源：", True, False), ("上海市杨浦区人民政府门户网站", False, True)], size=11, space_after=2)
    add_runs(doc, [("日期：", True, False), ("2026年5月29日", False, True)], size=11, space_after=2)
    add_runs(
        doc,
        [("链接：", True, False), ("https://www.shyp.gov.cn/shypq/xwzx-bmdt/20260529/506469.html", False, True)],
        size=11,
        space_after=4,
    )
    add_para(
        doc,
        "原文依据：第二场圆桌论坛由复旦大学住房政策研究中心秘书长、区科技企业联合会执行会长胡继刚主持……",
        size=10.5,
        color=RGBColor(0x66, 0x66, 0x66),
        space_after=8,
    )

    # 改动3
    add_h(doc, "五、改动3｜上传照片", 1)
    add_para(doc, "粘贴位置：百科编辑页 → 插入图片 / 词条头图。", size=10.5, color=RGBColor(0x55, 0x55, 0x55), space_after=4)
    add_runs(doc, [("头图（优先）：", True, False), ("1张正面半身正式照；免冠、清晰、无水印、无二维码。", False, True)], size=11, space_after=2)
    add_runs(doc, [("图注：", True, False), ("胡继刚", False, True)], size=11, space_after=2)
    add_runs(doc, [("来源：", True, False), ("本人提供", False, True)], size=11, space_after=4)
    add_runs(
        doc,
        [
            ("可选活动图：", True, False),
            ("1张能辨认本人的峰会现场照；图注写「胡继刚在2026人工智能商业化落地与硬核投资破局峰会现场」。", False, True),
        ],
        size=11,
        space_after=8,
    )

    # 提交
    add_h(doc, "六、提交顺序（照做即可）", 1)
    for i, t in enumerate(
        [
            "打开已通过的「胡继刚」词条 → 编辑。",
            "只替换2025年峰会那一段（改动1黄段）。",
            "添加参考资料〔9〕，挂到2026年AI段末尾 + 执行会长行末尾。",
            "上传头图（改动3）。",
            "预览：确认黄改处已挂脚注，白字段落未被改乱。",
            "提交理由可填：补充政府官网交叉来源，完善2025年公开活动表述，并添加人物照片；不新增未经公开来源支持的履历内容。",
        ],
        1,
    ):
        add_para(doc, f"{i}. {t}", size=11, space_after=3)

    add_h(doc, "七、一键复制区（仅黄段）", 1)
    add_para(doc, "复制A｜替换2025年经历段：", bold=True, size=11, space_after=2)
    add_runs(
        doc,
        [
            (
                "2025年5月，胡继刚以复旦大学住房政策研究中心秘书长身份参加“2025全球新经济增长引擎峰会”，在圆桌对话环节与证券、经济及资产管理领域嘉宾围绕房地产市场、资本布局与产业创新等议题展开交流，并以主要组织方代表身份就峰会主题及住房政策研究中心后续工作方向发表观点。",
                False,
                True,
            )
        ],
        size=11,
        space_after=6,
    )
    add_para(doc, "复制B｜〔9〕链接：", bold=True, size=11, space_after=2)
    add_runs(
        doc,
        [("https://www.shyp.gov.cn/shypq/xwzx-bmdt/20260529/506469.html", False, True)],
        size=11,
        space_after=6,
    )
    add_para(doc, "复制C｜提交理由：", bold=True, size=11, space_after=2)
    add_runs(
        doc,
        [
            (
                "补充政府官网交叉来源，完善2025年公开活动表述，并添加人物照片；不新增未经公开来源支持的履历内容。",
                False,
                True,
            )
        ],
        size=11,
        space_after=8,
    )

    add_para(
        doc,
        "图例：白色=保持不变　黄色=改动/新增　不要改白字段落。",
        size=10.5,
        color=RGBColor(0x55, 0x55, 0x55),
        space_before=4,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"已生成：{OUT}")
    return OUT


if __name__ == "__main__":
    build()
