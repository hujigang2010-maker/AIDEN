#!/usr/bin/env python3
"""生成给胡继刚口袋的三方对接清单：交警、平安、美团。家属自留。"""

from pathlib import Path
import html

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.colors import HexColor, white
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from comm_plan import CN_FONT_PATH

CHINESE_FONT = "宋体"
HEADING_FONT = "黑体"
ACCENT = RGBColor(0x0B, 0x2F, 0x5B)
MUTED = RGBColor(0x5C, 0x6B, 0x7A)
RED = RGBColor(0xA6, 0x3D, 0x2F)

MD_NAME = "三方对接清单_交警平安美团_2026-08-29.md"
DOCX_NAME = "青岛抚顺路和哈尔滨路路口交通事故_三方对接清单_20260829.docx"
PDF_NAME = "青岛抚顺路和哈尔滨路路口交通事故_三方对接清单_20260829.pdf"

TITLE = "现在要做什么：交警、平安、美团怎么对接"
SUB = "对象：胡继刚口袋　2026-08-29　家属自留，不要转给骑手、刘孝春或保险"
LEAD = (
    "现在不谈一共赔多少。三路并行、互不等待："
    "交警要书面认定书和监控回执；平安走理赔材料和进度；"
    "美团只把骑手和站点留在群里，不找骑手个人拿钱。"
)
NOW = [
    "平安陈老师：如果首次六问还没发，今天就发；已有病历、发票、身份证、受伤照片一起发。微信置顶，只走文字。",
    "市北区交警陈师傅：微信问四件事并请他文字回——事故编号、监控是否已调取保存到何时、家属/律师何时可看能否复制、认定书预计哪天出具。",
    "美团骑手：建一个群，拉你、骑手、站点负责人、平安陈老师。保单号让保险在群里写，不要再逼骑手个人交材料。",
    "家里同步：爸爸签授权委托书；妈妈找刘孝春写 3 万收据五句话 + 欠薪另写一张；收齐护工合同、发票、护理记录、转账凭证。",
]
POLICE = [
    "找谁：市北区交警陈师傅（承办本案）。交警大队不管赔多少钱，只管事实、监控和书面认定。",
    "怎么说话：微信为主，当面也要把结论再发一条微信留痕。地点一律写：抚顺路和哈尔滨路路口（抚顺路批发市场）。",
    "本周只要这四句回执：事故编号；监控已调取、保存到何时；可否查看/复制；认定书预计出具日。",
    "认定书来了：先拍照发给岳父看，再决定签不签、发不发平安。不要当场情绪化拒签，也不要逼交警写赔多少。",
    "领认定书：带爸爸身份证；你去领就出示授权委托书和你的身份证。核对路名、双方身份、过错描述、责任比例。不服有复核期。",
    "简易还是一般：岳父看完监控再定，你不自己拍板。不要主动把驾驶证交出去扩查。",
    "骑手不拉群：请陈师傅转达「请骑手把站点和保险拉进群」，不要升级成律师函。",
]
PINGAN = [
    "找谁：理赔专员陈老师 18929587336（美团众包骑手三者险）。超 5 个工作日不回，可提主管秦文政 18665057471。",
    "怎么对接：只走这个微信，置顶，文字留痕。不要只打电话。首次回复全文在《给平安理赔专员陈老师的回复_2026-08-29.md》，复制「请直接复制发出」那一段。",
    "今天问清六件事：报案号/保单号/三项限额；人民医院二甲认不认；认定书前能否预付；护理 300 元/天认不认及青岛上限；发票原件和分割单怎么走；后续走微信还是理赔通道。",
    "每周固定问一次进度（改报案号和本周补充即可）。材料缺什么请对方一次列全，不要自己猜着补。",
    "承保只有四项：医疗费（医保范围内）、误工费、护理费、物损（免赔 300，要认定书）。营养、康复、交通、精神抚慰金他们不赔——不赔不等于对方不用赔，记台账以后主张。",
    "物损是刘孝春的三轮，让他自己对接陈老师。我们不代办、不代弃。",
    "认定书先给岳父，签字前不发给陈老师。对平安只说「社保已走」，不披露泰康，不报总数、不报残级。",
]
MEITUAN = [
    "找谁：女骑手微信 + 她的站点负责人。美团赔付走的是骑手侧平安险，不是找她个人钱包。",
    "现在保险已经进场，对骑手只办一件事：把站点和陈老师拉进同一个群。不要再连发术语清单。",
    "跟进微信全文在《给美团骑手的跟进微信_2026-08-22.md》。她说没钱、等认定书、不知道保单号：一律回到「拉群、让保险说话」。",
    "不要对她说：全责、一共要多少、70%/30%、岳父是律师、刘孝春怎么分担。私下给一点了结：不签。",
    "她已读不回：当天请陈师傅转达拉群。现在不发律师函、不请诉讼律师当面谈。",
]
NEVER = [
    "总赔偿数额、总包、一次性了结、残级。有人问一共要多少：费用还在发生，等认定书和治疗稳定后再按票算。",
    "内部 70%/30%、刘孝春分担安排、泰康细节、「岳父是律师」。",
    "签任何「一次了结、以后不再主张」，包括微信口头两清。",
    "去人社局报工伤。无劳动合同，工伤排除。",
    "造误工流水、假单位证明。平安这边误工费大概率拿不到，以后用派活微信向侵权方主张。",
]


def render_markdown() -> str:
    lines = [
        f"# {TITLE}",
        "",
        SUB,
        "",
        f"> {LEAD}",
        "",
        "## 这几天你要做的（四路并行，互不等）",
        "",
    ]
    for i, item in enumerate(NOW, 1):
        lines.append(f"{i}. {item}")
    lines.extend(["", "## 一、交警大队（市北区陈师傅）", ""])
    for item in POLICE:
        lines.append(f"- {item}")
    lines.extend(["", "## 二、保险公司（平安陈老师）", ""])
    for item in PINGAN:
        lines.append(f"- {item}")
    lines.extend(["", "## 三、美团（骑手 + 站点）", ""])
    for item in MEITUAN:
        lines.append(f"- {item}")
    lines.extend(["", "## 对外一律不说", ""])
    for item in NEVER:
        lines.append(f"- {item}")
    lines.extend(
        [
            "",
            "完整回复稿和骑手跟进稿仍用原文件复制，本文只当口袋清单。重新生成：`python3 scripts/build_all.py`。不是律师函。",
            "",
        ]
    )
    return "\n".join(lines)


def build_markdown(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(render_markdown(), encoding="utf-8")


def _set_run_font(run, font_name: str, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _add_paragraph(
    doc: Document,
    text: str = "",
    *,
    font: str = CHINESE_FONT,
    size: float = 11,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    space_after: float = 3,
    space_before: float = 0,
    line_spacing: float = 1.25,
    color: RGBColor | None = None,
):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold, color=color)
    return p


def build_document(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    normal_style = doc.styles["Normal"]
    normal_style.font.name = CHINESE_FONT
    rpr = normal_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CHINESE_FONT)

    _add_paragraph(doc, TITLE, font=HEADING_FONT, size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, color=ACCENT)
    _add_paragraph(doc, SUB, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, color=MUTED)
    _add_paragraph(doc, LEAD, bold=True, size=11, color=RED, space_after=8)

    _add_paragraph(doc, "这几天你要做的（四路并行，互不等）", font=HEADING_FONT, size=13, bold=True, space_after=4, color=ACCENT)
    for i, item in enumerate(NOW, 1):
        _add_paragraph(doc, f"{i}. {item}", size=11, space_after=3)

    _add_paragraph(doc, "一、交警大队（市北区陈师傅）", font=HEADING_FONT, size=13, bold=True, space_before=8, space_after=4, color=ACCENT)
    for item in POLICE:
        _add_paragraph(doc, "• " + item, size=11, space_after=3)

    _add_paragraph(doc, "二、保险公司（平安陈老师）", font=HEADING_FONT, size=13, bold=True, space_before=8, space_after=4, color=ACCENT)
    for item in PINGAN:
        _add_paragraph(doc, "• " + item, size=11, space_after=3)

    _add_paragraph(doc, "三、美团（骑手 + 站点）", font=HEADING_FONT, size=13, bold=True, space_before=8, space_after=4, color=ACCENT)
    for item in MEITUAN:
        _add_paragraph(doc, "• " + item, size=11, space_after=3)

    _add_paragraph(doc, "对外一律不说", font=HEADING_FONT, size=13, bold=True, space_before=8, space_after=4, color=RED)
    for item in NEVER:
        _add_paragraph(doc, "• " + item, size=11, space_after=3, color=RED)

    _add_paragraph(
        doc,
        "完整回复稿和骑手跟进稿仍用原文件复制。本文家属自留，不是律师函。",
        size=9.5,
        color=MUTED,
        space_before=8,
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _pdf_font() -> str:
    kwargs = {"subfontIndex": 0} if CN_FONT_PATH.endswith(".ttc") else {}
    pdfmetrics.registerFont(TTFont("CN", CN_FONT_PATH, **kwargs))
    return "CN"


def _p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(html.escape(str(text)).replace("\n", "<br/>"), style)


def _header_footer(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(HexColor("#0B2F5B"))
    canvas.rect(0, A4[1] - 11 * mm, A4[0], 11 * mm, fill=1, stroke=0)
    canvas.setFillColor(white)
    canvas.setFont("CN", 8)
    canvas.drawString(14 * mm, A4[1] - 7 * mm, "三方对接清单 · 交警 / 平安 / 美团 · 家属自留")
    canvas.setFillColor(HexColor("#C4A35A"))
    canvas.rect(0, 0, A4[0], 9 * mm, fill=1, stroke=0)
    canvas.setFillColor(HexColor("#0B2F5B"))
    canvas.setFont("CN", 8)
    canvas.drawString(14 * mm, 3.2 * mm, "不谈总包 · 不报残级 · 不发律师函 · 地点写抚顺路和哈尔滨路路口")
    canvas.drawRightString(A4[0] - 14 * mm, 3.2 * mm, f"第 {doc.page} 页")
    canvas.restoreState()


def build_pdf(output_path: Path) -> None:
    font = _pdf_font()
    navy = HexColor("#0B2F5B")
    muted = HexColor("#5C6B7A")
    red = HexColor("#A63D2F")
    dark = HexColor("#1F2A37")
    styles = {
        "title": ParagraphStyle("t", fontName=font, fontSize=14, leading=20, textColor=navy, alignment=TA_CENTER, spaceAfter=3),
        "sub": ParagraphStyle("s", fontName=font, fontSize=9, leading=13, textColor=muted, alignment=TA_CENTER, spaceAfter=5),
        "lead": ParagraphStyle("l", fontName=font, fontSize=10, leading=15, textColor=red, alignment=TA_LEFT, spaceAfter=6),
        "h1": ParagraphStyle("h", fontName=font, fontSize=11.5, leading=16, textColor=navy, spaceBefore=6, spaceAfter=3),
        "h1red": ParagraphStyle("hr", fontName=font, fontSize=11.5, leading=16, textColor=red, spaceBefore=6, spaceAfter=3),
        "bullet": ParagraphStyle("u", fontName=font, fontSize=9.5, leading=14, textColor=dark, leftIndent=2, spaceAfter=2),
        "redb": ParagraphStyle("r", fontName=font, fontSize=9.5, leading=14, textColor=red, leftIndent=2, spaceAfter=2),
        "foot": ParagraphStyle("f", fontName=font, fontSize=8, leading=12, textColor=muted, spaceBefore=6),
    }
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=13 * mm,
        rightMargin=13 * mm,
        topMargin=14 * mm,
        bottomMargin=12 * mm,
        title=TITLE,
        author="内部整理",
    )
    story = [
        _p(TITLE, styles["title"]),
        _p(SUB, styles["sub"]),
        _p(LEAD, styles["lead"]),
        _p("这几天你要做的（四路并行，互不等）", styles["h1"]),
    ]
    for i, item in enumerate(NOW, 1):
        story.append(_p(f"{i}. {item}", styles["bullet"]))
    story.append(_p("一、交警大队（市北区陈师傅）", styles["h1"]))
    for item in POLICE:
        story.append(_p("• " + item, styles["bullet"]))
    story.append(_p("二、保险公司（平安陈老师）", styles["h1"]))
    for item in PINGAN:
        story.append(_p("• " + item, styles["bullet"]))
    story.append(_p("三、美团（骑手 + 站点）", styles["h1"]))
    for item in MEITUAN:
        story.append(_p("• " + item, styles["bullet"]))
    story.append(_p("对外一律不说", styles["h1red"]))
    for item in NEVER:
        story.append(_p("• " + item, styles["redb"]))
    story.append(Spacer(1, 4))
    story.append(_p("完整回复稿和骑手跟进稿仍用原文件复制。本文家属自留，不是律师函。", styles["foot"]))
    doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "deliverables"
    build_markdown(out / MD_NAME)
    build_document(out / DOCX_NAME)
    build_pdf(out / PDF_NAME)
    print(f"Wrote {out / MD_NAME}")
    print(f"Wrote {out / DOCX_NAME}")
    print(f"Wrote {out / PDF_NAME}")


if __name__ == "__main__":
    main()
