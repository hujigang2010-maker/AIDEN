"""Generate the Guansong Auto Group sponsorship agreement as a .docx file."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


CHINESE_FONT = "宋体"
HEADING_FONT = "黑体"


def _set_run_font(run, font_name: str, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _add_paragraph(
    doc: Document,
    text: str = "",
    *,
    font: str = CHINESE_FONT,
    size: float = 11,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: float | None = None,
    space_after: float = 6,
    space_before: float = 0,
    line_spacing: float = 1.5,
    color: RGBColor | None = None,
):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if text:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold, color=color)
    return p


def _add_rich_paragraph(
    doc: Document,
    segments: list[tuple[str, bool]],
    *,
    font: str = CHINESE_FONT,
    size: float = 11,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: float | None = None,
    space_after: float = 6,
    line_spacing: float = 1.5,
):
    """segments: list of (text, bold)."""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    for text, bold in segments:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold)
    return p


def build_document(output_path: Path) -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = CHINESE_FONT
    normal_style.font.size = Pt(11)
    rpr = normal_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CHINESE_FONT)

    _add_paragraph(
        doc,
        "2026人工智能商业化落地与硬核投资破局峰会",
        font=HEADING_FONT,
        size=20,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
        line_spacing=1.4,
    )
    _add_paragraph(
        doc,
        "官方智慧出行伙伴合作协议",
        font=HEADING_FONT,
        size=18,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=14,
        line_spacing=1.4,
    )

    _add_rich_paragraph(
        doc,
        [("甲方（主办方）：", True), ("人工智能商业化落地峰会组委会 / 复旦大学住房政策研究中心", False)],
        space_after=4,
    )
    _add_rich_paragraph(
        doc,
        [("乙方（赞助方）：", True), ("上海冠松汽车股份有限公司（华为鸿蒙智行授权合作伙伴）", False)],
        space_after=10,
    )

    _add_paragraph(
        doc,
        (
            "鉴于甲方拥有顶尖的政企智库资源及长三角前沿科技产融生态，乙方系国内领先的汽车经销商集团及华为鸿蒙智行"
            "授权合作伙伴。双方本着优势互补、资源共享的原则，就乙方作为本次峰会“官方指定智慧出行独家伙伴”事宜，"
            "达成如下协议："
        ),
        first_line_indent=0.74,
        space_after=10,
    )

    def add_article_heading(text: str) -> None:
        _add_paragraph(
            doc,
            text,
            font=HEADING_FONT,
            size=13,
            bold=True,
            space_before=8,
            space_after=6,
            line_spacing=1.4,
        )

    add_article_heading("第一条　合作级别与赞助标的说明")

    _add_rich_paragraph(
        doc,
        [
            ("1. 合作身份：", True),
            (
                "乙方作为本次峰会及高净值私局的“钻石赞助商”及“官方指定智慧出行独家伙伴”。",
                False,
            ),
        ],
    )
    _add_rich_paragraph(
        doc,
        [
            ("2. 综合赞助标的额：", True),
            (
                "乙方为本次活动提供专项赞助（含现金及高端接驳车队等重资产实物调用对价），总计金额为",
                False,
            ),
            ("人民币 200,000 元（大写：贰拾万元整）", True),
            ("。", False),
        ],
    )
    _add_rich_paragraph(
        doc,
        [
            ("3. 资金与资源定向消耗：", True),
            ("双方在此明确约定，", False),
            (
                "上述 20 万元赞助对价仅限用于且全部消耗于“本次峰会当期”的专属落地执行。",
                True,
            ),
            ("该笔专项对价的消耗范围涵盖但不限于：", False),
        ],
    )
    bullets_art1 = [
        "北外滩核心地标的专属江景车辆特装展位费及动线规划费；",
        "峰会当期华为鸿蒙智行高端车队的独家调度运营（含车辆折旧成本、全险、专职司机劳务及高端接待标准成本）；",
        "峰会现场高净值VIP专属试乘试驾动线的安保及综合运营保障支出。",
    ]
    for item in bullets_art1:
        _add_paragraph(
            doc,
            f"• {item}",
            first_line_indent=0.74,
            space_after=4,
        )

    add_article_heading("第二条　甲方当期核心交付权益")

    _add_rich_paragraph(
        doc,
        [
            ("甲方在", False),
            ("本次峰会现场", True),
            ("，向乙方交付以下顶级品牌与转化权益：", False),
        ],
    )

    art2_items: list[list[tuple[str, bool]]] = [
        [
            ("1. 核心场景曝光：", True),
            (
                "乙方享有本次峰会主视觉背景板、签到处及官方宣传中的“官方指定智慧出行伙伴”专属顶级Logo露出。",
                False,
            ),
        ],
        [
            ("2. 高端实车特展：", True),
            (
                "甲方在北外滩会场核心区域为乙方划定专属江景实车展位，用于全方位展示乙方华为鸿蒙智行系列（如尊界、问界M9等）的智能座舱与高阶智驾体验。",
                False,
            ),
        ],
        [
            ("3. 高净值圈层现场导流：", True),
            (
                "甲方利用自身政商顶级圈层，在峰会期间为乙方定向引荐具备大宗换车及企业集中采购需求的高净值人群（重点覆盖头部房企高管、算力大厂代表、科创企业创始人及顶尖智库学者），并由乙方团队在现场闭门对接。",
                False,
            ),
        ],
    ]
    for segs in art2_items:
        _add_rich_paragraph(doc, segs)

    add_article_heading("第三条　长效生态延伸权益（战略附赠）")

    _add_paragraph(
        doc,
        (
            "鉴于乙方在本次峰会中展现的顶级赞助诚意及双方在智慧出行领域的共识，甲方同意在本次活动之外，"
            "基于长三角新质产融生态的整体战略框架，向乙方额外开放以下延伸赋能权益（本条款下所有权益均为组委会战略附赠，"
            "不计入第一条之当期财务核算体系，亦不作单项计价）："
        ),
    )

    art3_items: list[list[tuple[str, bool]]] = [
        [
            ("1. 常态化圈层入场：", True),
            (
                "乙方自动成为“见微知海新质商业生态”年度战略理事单位，享有未来半年内组委会举办的其他系列闭门私局的优先参与权。",
                False,
            ),
        ],
        [
            ("2. 大宗采购优先匹配：", True),
            (
                "针对后续组委会生态内成功跑通的重资产企业、科创园区及超级个体，如有高端商务用车集中采购或高管换车需求，甲方将优先向乙方及华为鸿蒙智行体系进行线索导入与商业撮合。",
                False,
            ),
        ],
        [
            ("3. 前沿课题共建：", True),
            (
                "甲方智库团队在进行“智慧城市与新质人居”相关学术及商业课题调研时，优先将乙方及华为生态作为研究样板与联合发布方。",
                False,
            ),
        ],
    ]
    for segs in art3_items:
        _add_rich_paragraph(doc, segs)

    add_article_heading("第四条　保密与对标豁免")

    _add_rich_paragraph(
        doc,
        [
            ("1. 商业机密保护：", True),
            (
                "本协议第一条体现了甲乙双方最核心的高端重资产调度折算成本，未经双方书面同意，任何一方不得向第三方披露本协议相关的实物折算明细清单及底层财务核算逻辑。",
                False,
            ),
        ],
    )
    _add_rich_paragraph(
        doc,
        [
            ("2. 独立横向比对豁免：", True),
            ("为促成更广泛的产融生态建设，", False),
            (
                "甲方有权在对外招商及面向同级别头部企业（如头部商业地产、大型科技巨头等）进行横向比对与资质说明时，合法出示本协议及合作总金额",
                True,
            ),
            ("，以证明本组委会系列活动的顶级市场公允估值与合作门槛。", False),
        ],
    )

    add_article_heading("第五条　签署与生效")

    _add_paragraph(
        doc,
        (
            "本协议一式两份，甲乙双方各执一份。自双方法定代表人或授权代表签字并加盖公章（或合同专用章、组委会印章）"
            "之日起生效，具有同等法律效力。"
        ),
        first_line_indent=0.74,
    )

    _add_paragraph(
        doc,
        "（以下无正文）",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=10,
        space_after=18,
    )

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell.width = Cm(7.5)

    def fill_signature_cell(cell, party_label: str, party_name: str) -> None:
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(8)
        r1 = p1.add_run(party_label)
        _set_run_font(r1, CHINESE_FONT, 11, bold=True)
        r2 = p1.add_run(party_name)
        _set_run_font(r2, CHINESE_FONT, 11)

        for label in ("授权代表签字：", "", "日期：2026 年    月    日"):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(14)
            run = p.add_run(label)
            _set_run_font(run, CHINESE_FONT, 11)

    fill_signature_cell(
        table.rows[0].cells[0],
        "甲方：",
        "人工智能商业化落地峰会组委会 /\n上海市杨浦区科技企业联合会（盖章）",
    )
    fill_signature_cell(
        table.rows[0].cells[1],
        "乙方：",
        "上海冠松汽车股份有限公司（盖章）",
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    output = Path(__file__).resolve().parent.parent / "agreements" / "冠松汽车-智慧出行伙伴合作协议.docx"
    build_document(output)
    print(f"Generated: {output}")


if __name__ == "__main__":
    main()
