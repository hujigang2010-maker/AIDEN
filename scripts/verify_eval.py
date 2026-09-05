#!/usr/bin/env python3
"""核对评估稿是否覆盖三份原文的关键事实。"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]


def main() -> int:
    eval_md = (ROOT / "docs/CGC领事会客厅三案评估意见.md").read_text(encoding="utf-8")
    eval_doc = Document(ROOT / "output/CGC领事会客厅三案评估意见.docx")
    eval_text = eval_md + "\n" + "\n".join(p.text for p in eval_doc.paragraphs)
    for table in eval_doc.tables:
        for row in table.rows:
            eval_text += "\n" + " ".join(c.text for c in row.cells)

    semir = Document(ROOT / "source/森马领事会客厅.docx")
    semir_text = "\n".join(p.text for p in semir.paragraphs)
    chuang = Document(ROOT / "source/创智汇领事会客厅.docx")
    chuang_text = "\n".join(p.text for p in chuang.paragraphs)

    checks: list[tuple[str, bool, str]] = []

    def ok(name: str, cond: bool, detail: str = "") -> None:
        checks.append((name, bool(cond), detail))

    ok("森马2000㎡", "2000㎡" in semir_text and "2000㎡" in eval_text)
    ok("森马价值章1000㎡笔误存在于原文", "1000㎡专属空间" in semir_text.replace(" ", ""))
    ok("评估点出面积笔误", "笔误" in eval_text and "1000㎡" in eval_text)
    ok("创智汇1000㎡", "1000㎡" in chuang_text and "创智汇" in eval_text)
    ok("10万流向相反", "流向相反" in eval_text)
    ok("一滴水挂牌费方向", "餐厅支付" in eval_text or "餐厅 → CGC" in eval_text)
    ok("禁止官方授权", "官方授权" in eval_text)
    ok("国资风险", "国有" in eval_text and "国资" in eval_text)
    ok("中建四局冲突", "中建四局" in chuang_text and "中建四局" in eval_text)
    ok("可执行结论三分法", "小改" in eval_text and "大改" in eval_text and "中改" in eval_text)
    ok("90天路径", "90 天" in eval_text or "90天" in eval_text)
    ok("总判断", "可以谈，不能原样发出" in eval_text)

    s1 = [p.text.strip() for p in semir.paragraphs if p.text.strip()]
    c1 = [p.text.strip() for p in chuang.paragraphs if p.text.strip()]
    shared = sum(1 for a, b in zip(s1[3:20], c1[2:19]) if a == b)
    ok("两园区前章高度复用", shared >= 10, f"shared={shared}")

    html = ROOT / "output/CGC领事会客厅三案评估意见.html"
    ok("HTML 已生成", html.exists() and html.stat().st_size > 1000)

    fail = 0
    print("VERIFICATION")
    for name, passed, detail in checks:
        mark = "PASS" if passed else "FAIL"
        if not passed:
            fail += 1
        print(f"  {mark}  {name}  {detail}".rstrip())
    print(f"\n{len(checks) - fail}/{len(checks)} passed")
    return fail


if __name__ == "__main__":
    sys.exit(main())
