#!/usr/bin/env python3
"""生成《CGC领事会客厅三案评估意见》。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output"
DOCS_DIR = ROOT / "docs"
FONT_CN = "WenQuanYi Micro Hei"
FONT_EN = "Calibri"
NAVY = RGBColor(0x1B, 0x3A, 0x5F)
RED = RGBColor(0xB4, 0x23, 0x18)
ORANGE = RGBColor(0xB8, 0x5C, 0x00)
GREEN = RGBColor(0x1F, 0x6B, 0x3A)
GRAY = RGBColor(0x4A, 0x4A, 0x4A)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_run_font(run, size=11, bold=False, color=BLACK, font=FONT_CN):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = FONT_EN
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:cs"), FONT_EN)


def shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell, color="D0D5DD") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, *, size=9, bold=False, color=BLACK, align="left", fill=None):
    if fill:
        shade_cell(cell, fill)
    set_cell_border(cell)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_paragraph(doc, text, *, size=11, bold=False, color=BLACK, space_after=8, first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.35
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading_cn(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0)
    size = 16 if level == 1 else 13
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=NAVY)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "1B3A5F")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_bullet(doc, text, *, bold_prefix=None, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run("• ")
    set_run_font(run, size=11, bold=True, color=NAVY)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True, color=color)
        r2 = p.add_run(text)
        set_run_font(r2, size=11, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=color)
    return p


def add_callout(doc, title, body, fill="FDECEC", border="E4B4B4", title_color=RED):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, border)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(title)
    set_run_font(r1, size=12, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.line_spacing = 1.3
    r2 = p2.add_run(body)
    set_run_font(r2, size=11, color=BLACK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_table(doc, headers, rows, col_widths=None, header_fill="1B3A5F"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        set_cell_text(
            table.cell(0, i),
            h,
            size=9,
            bold=True,
            color=WHITE,
            align="center",
            fill=header_fill,
        )
    for r_idx, row in enumerate(rows, start=1):
        fill = "F7F9FC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            align = "center" if c_idx > 0 else "left"
            set_cell_text(table.cell(r_idx, c_idx), val, size=9, align=align, fill=fill)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, size=9, color=GRAY)
    return p


# ---------------------------------------------------------------------------
# Markdown 同步输出
# ---------------------------------------------------------------------------

MD_LINES: list[str] = []


def md(text: str = "") -> None:
    MD_LINES.append(text)


def md_table(headers, rows) -> None:
    md("| " + " | ".join(headers) + " |")
    md("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in rows:
        md("| " + " | ".join(str(x).replace("\n", "<br>") for x in row) + " |")
    md()


def build_markdown() -> Path:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    path = DOCS_DIR / "CGC领事会客厅三案评估意见.md"
    path.write_text("\n".join(MD_LINES).rstrip() + "\n", encoding="utf-8")
    return path


def build() -> tuple[Path, Path]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("内部评估意见  ·  仅供决策使用  ·  2026-09-05")
    set_run_font(hr, size=8, color=GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("CGC 领事会客厅三案评估  ·  不作为对外发送稿")
    set_run_font(fr, size=8, color=GRAY)

    # 封面标题
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(8)
    t.paragraph_format.space_after = Pt(4)
    r = t.add_run("CGC 领事会客厅三案评估意见")
    set_run_font(r, size=22, bold=True, color=NAVY)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(4)
    sr = st.add_run("森马上海国际运营中心  ·  创智汇  ·  尚9·一滴水")
    set_run_font(sr, size=12, color=GRAY)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(12)
    mr = meta.add_run("评估对象为三份「合作建议」原文  ·  结论：方向成立，文本未达可签约级")
    set_run_font(mr, size=10, bold=True, color=ORANGE)

    md("# CGC 领事会客厅三案评估意见")
    md()
    md("**森马上海国际运营中心 · 创智汇 · 尚9·一滴水**")
    md()
    md("> 内部评估意见，仅供决策使用。评估对象为三份「合作建议」原文。结论：方向成立，文本未达可签约级。")
    md()

    # ------------------------------------------------------------------
    add_heading_cn(doc, "一、一页纸结论")
    md("## 一、一页纸结论")
    md()

    add_callout(
        doc,
        "总判断：可以谈，不能原样发出",
        "三份文本是同一套 CGC「领事会客厅」骨架的场地换皮，不是三套独立策划。"
        "「民间交流平台 + 产业主题空间 + 公益沙龙/市场化深度服务分层」这条产品逻辑成立，"
        "也比空喊「对接各国中央政府」更接近可落地的生意。"
        "但目前对园区是「要 1000–2000㎡ 免租 + 全额装修 + 每年 10 万现金 + 不定额营销」，"
        "对己方几乎没有可核验的场次、领事出席、线索转化和驻场运营承诺。"
        "按这个对价，森马和创智汇的法务、招商、国资部门大概率会直接退回。"
        "一滴水是轻资产挂牌，商业结构相对健康，但仍缺保底场次，且出现「CGC 官方授权」与「非官方民间定位」的口径冲突。",
        fill="FFF6E8",
        border="E8C48A",
        title_color=ORANGE,
    )
    md("**总判断：可以谈，不能原样发出。**")
    md()
    md("三份文本是同一套 CGC「领事会客厅」骨架的场地换皮，不是三套独立策划。"
       "产品逻辑成立，但对园区要价过重、己方承诺不可核验。一滴水结构相对健康，仍缺保底场次。")
    md()

    add_paragraph(
        doc,
        "建议动作很明确：一滴水可在补齐保底场次和对价表后进入商务谈判；"
        "创智汇先改成「活动合作 + 小型展陈」，不要一上来要 1000㎡ 国有物业无偿使用；"
        "森马先改面积笔误，把 2000㎡ 降到可解释的展厅规模，并与创智汇错开 AI 主题。"
        "三案不要同一周发出，否则对方一对口径就会发现是批量定制。",
        first_line=True,
    )
    md("建议动作：一滴水补保底场次后可谈；创智汇改为活动合作+小型展陈；森马先改面积笔误并错开 AI 主题。三案不要同一周发出。")
    md()

    add_caption(doc, "表1  三案可执行评分（10 分制，内部裁量，用于排序而不是对外报价）")
    md("**表1  三案可执行评分（10 分制，内部裁量）**")
    md()
    headers = ["维度", "森马", "创智汇", "一滴水", "说明"]
    rows = [
        ["战略匹配", "7.0", "8.0", "7.5", "创智汇有高校和国企开放任务；森马有潮玩产业故事；一滴水有中阿场景"],
        ["商业对等", "3.5", "3.5", "7.0", "两园区是对方出空间和现金；一滴水是对方付挂牌费、CGC 导流"],
        ["落地可执行", "4.0", "4.0", "6.0", "均缺日历、驻场、KPI；一滴水不新占物业，执行门槛更低"],
        ["合规过会", "5.0", "3.5", "6.5", "创智汇涉及国有物业无偿使用，过会难度最高"],
        ["主题差异", "6.0", "5.0", "8.0", "森马与创智汇都打 AI，会互相稀释；中阿主题最干净"],
        ["建议处置", "改后再发", "大改", "小改可谈", "均不建议以当前文本作为合同附件"],
    ]
    add_table(doc, headers, rows, col_widths=[2.4, 1.8, 1.8, 1.8, 8.8])
    md_table(headers, rows)

    add_paragraph(
        doc,
        "综合分大约 5.5/10。作为「内部讨论稿」可用，作为「给董事长/园区总经理的合作建议」不够。"
        "真正值钱的不是 1000㎡ 挂牌，而是：一年能稳定请来多少领事和商协会、能沉淀多少可跟进线索、"
        "以及会客厅在非活动日到底谁来值守。这三件事，三份文本都没有写清楚。",
        first_line=True,
    )
    md("综合分大约 5.5/10。真正值钱的是领事出席的稳定性、线索沉淀和驻场值守，这三件事文本都没写清。")
    md()

    # ------------------------------------------------------------------
    add_heading_cn(doc, "二、三案分别是什么")
    md("## 二、三案分别是什么")
    md()
    add_paragraph(
        doc,
        "三份文件标题不同，骨架相同：先用同一套 CGC 简介和五项能力（会客厅、企业出海、来华考察、海外中小企业扶持、青年交流），"
        "再贴上「民间交流、不承担官方外交职能」的合规声明，最后才进入场地条款和合作价值。"
        "一至三章几乎可以整段复用，差异集中在第四章的要价和第五章的价值话术。"
        "这不是三案并行的网络规划，而是一份母版对三个业主的销售定制。",
        first_line=True,
    )
    md("三份文件骨架相同：CGC 简介和五项能力复用，差异集中在场地条款和价值话术。这是一份母版对三个业主的销售定制。")
    md()

    add_heading_cn(doc, "（一）森马：重资产入驻时尚产业园", 2)
    md("### （一）森马：重资产入驻时尚产业园")
    md()
    add_paragraph(
        doc,
        "对象是闵行吴泾镇莲花南路 2689 号森马上海国际运营中心，自称 22 万㎡、「国际总部 / 时尚总部 / 创业总部」，"
        "正在往潮玩 IP、AIGC、时尚 AI 方向导流。方案要求园区提供合计 2000㎡ 免租金专属场地，挂牌「亚洲潮玩与AI国际领事会客厅」；"
        "森马全额装修、承担水电物业网络；每年再给 CGC 10 万元挂牌经费，另设不定额营销预算。"
        "CGC 承诺国别沙龙、海外 IP 来华对接、园区企业出海线索台账、每年不少于 3 场青年沙龙，以及对外物料双品牌露出。"
        "合作期建议 3 年，按「年度绩效评估」决定续约，但评估指标未定义。",
        first_line=True,
    )
    md("对象为闵行吴泾森马上海国际运营中心。要 2000㎡ 免租专属场地、全额装修、水电物业、每年 10 万挂牌经费和不定额营销。"
       "CGC 承诺沙龙、IP 对接、线索台账和每年不少于 3 场青年沙龙。3 年期，绩效指标未定义。")
    md()

    add_heading_cn(doc, "（二）创智汇：重资产入驻国企科创园区", 2)
    md("### （二）创智汇：重资产入驻国企科创园区")
    md()
    add_paragraph(
        doc,
        "对象是杨浦大创智核心载体创智汇，资产归属杨浦科技创新集团，由全资子公司上海北岛科技发展有限公司运营。"
        "周边复旦、同济、财大，园区已有中建四局「AI+IP 产业创新中心」。"
        "方案要 1000㎡ 免租专属场地，挂牌「亚洲数字AI产业会客厅（领事会客厅）」；园区全额装修且资产归国有；"
        "水电物业网络由园区承担；同样每年 10 万元挂牌经费 + 不定额营销补贴。"
        "CGC 责任与森马高度同构，只是把潮玩换成 AI / 数字文创 / 智能制造，并强调高校青年交流和国际招商推介。",
        first_line=True,
    )
    md("对象为杨浦创智汇（杨浦科创集团/北岛科技）。要 1000㎡ 免租、国资全额装修、每年 10 万和营销补贴。责任条款与森马同构。")
    md()

    add_heading_cn(doc, "（三）一滴水：轻资产挂牌滨江中餐厅", 2)
    md("### （三）一滴水：轻资产挂牌滨江中餐厅")
    md()
    add_paragraph(
        doc,
        "对象是尚9·一滴水江景中餐厅，文本写「黄浦区外滩及陆家嘴滨江核心区域」，实际控制人长期深耕阿联酋。"
        "这是完全不同的交易结构：不新要专属物业，餐厅在入口、VIP 区和宴会厅挂「CGC阿拉伯领事会客厅·尚9一滴水交流基地」；"
        "宴会厅、包间、10–30 人沙龙区按档期预约给 CGC；餐饮按优惠价另结。"
        "CGC 收取每年 10 万元挂牌服务费（不抵扣餐费），承诺把领事晚宴、中阿闭门沙龙、国际商会宴请、企业家私享会优先放到该餐厅。"
        "文本明确 CGC 不介入餐饮经营、不承担房租人力食材；沙龙和路演免费提供场地，宴请另计。",
        first_line=True,
    )
    md("对象为尚9·一滴水江景中餐厅。不新占物业，餐厅付 CGC 每年 10 万挂牌费，CGC 承诺优先导入领事晚宴和中阿沙龙。")
    md()

    add_paragraph(
        doc,
        "和既有「金茂 86 层国家会客厅」战略相比，这三案是一次明显的产品下沉：从高客单价国别冠名（原方案量级在每年数十万至数百万元），"
        "变成向产业园区要空间、向餐厅收 10 万元挂牌费。"
        "如果金茂旗舰还没做出可参观、可复购的样板，先在上海铺三个卫星点，品牌会被摊薄，价格锚也会被自己打穿。",
        first_line=True,
    )
    md("相对金茂 86 层国家会客厅高客单价冠名，三案是产品下沉。旗舰未成样板就铺卫星点，会摊薄品牌并打穿价格锚。")
    md()

    # ------------------------------------------------------------------
    add_heading_cn(doc, "三、三案对照")
    md("## 三、三案对照")
    md()
    add_caption(doc, "表2  交易结构对照")
    md("**表2  交易结构对照**")
    md()
    headers = ["条款", "森马", "创智汇", "一滴水"]
    rows = [
        ["业主性质", "民营时尚产业园", "区属国企园区", "民营餐饮"],
        ["挂牌名", "亚洲潮玩与AI国际领事会客厅", "亚洲数字AI产业会客厅", "阿拉伯领事会客厅交流基地"],
        ["占用方式", "2000㎡ 专属免租", "1000㎡ 专属免租", "档期预约，不新占面积"],
        ["装修", "森马全额，资产归森马", "园区全额，资产归国有", "无"],
        ["现金流向", "园区 → CGC 10万/年", "园区 → CGC 10万/年", "餐厅 → CGC 10万/年"],
        ["营销费", "园区另设，额度未写", "园区另设，额度未写", "无；餐费另结"],
        ["CGC 可量化义务", "青年沙龙 ≥3 场/年", "无场次 KPI", "无场次 KPI"],
        ["退出", "3 年 + 空泛绩效评估", "同左", "终止后拆除标识"],
        ["转租限制", "禁止转租分租", "禁止转租分租", "禁止转授权品牌"],
    ]
    add_table(doc, headers, rows, col_widths=[3.0, 4.6, 4.6, 4.4])
    md_table(headers, rows)

    add_paragraph(
        doc,
        "一眼能看出两个结构问题。第一，10 万元在三案里数字相同、方向相反：园区是付钱请 CGC 来占地方，餐厅是付钱买挂牌。"
        "对内要解释清楚这不是「统一报价」，而是两种生意，否则财务和商务会把它做成错误的标准套餐。"
        "第二，森马正文要 2000㎡，合作价值章节又写「1000㎡ 专属空间」——这是从创智汇稿粘贴后没改干净的硬伤。"
        "面积差一倍，投资匡算差一倍，不能带着笔误去见森马。",
        first_line=True,
    )
    md("两个结构问题：10 万元在园区和餐厅流向相反；森马正文 2000㎡、价值章节写成 1000㎡，是创智汇稿粘贴残留。")
    md()

    add_caption(doc, "表3  对方三年投入数量级匡算（示意，按保守租金和中档装修，不是估价报告）")
    md("**表3  对方三年投入数量级匡算（示意）**")
    md()
    headers = ["成本项", "森马（按 2000㎡）", "创智汇（按 1000㎡）", "一滴水"]
    rows = [
        ["租金机会成本", "约 800–1200 万/3年\n（按 4–5.5 元/㎡/天）", "约 550–800 万/3年\n（按 5–7 元/㎡/天）", "包间/宴会厅档期机会成本，取决于场次"],
        ["专属装修", "约 600–1200 万\n（3000–6000 元/㎡）", "约 300–600 万", "基本为零"],
        ["现金挂牌", "30 万", "30 万", "餐厅支付 30 万给 CGC"],
        ["营销与水电", "未定价，但文本承诺由对方承担", "同左", "沙龙免费场地；宴请按优惠价"],
        ["三年总成本量级", "约 1500–2500 万", "约 900–1500 万", "净支出取决于 CGC 能否带来宴会"],
        ["CGC 锁定承诺", "青年沙龙 9 场/3年 + 若干未定量活动", "未定量", "优先导流，未定量"],
    ]
    add_table(doc, headers, rows, col_widths=[3.0, 4.6, 4.6, 4.4])
    md_table(headers, rows)

    add_paragraph(
        doc,
        "园区两案的投入产出比，用招商口径也很难圆：要证明 1000–2000㎡ 免租能换来对等的外资项目或头部 IP 落户，"
        "必须拿出可验证的历史案例、领事确认函和首年招商线索目标。现在这些都没有。"
        "一滴水则是「10 万买一块牌 + 赌 CGC 能带宴请」。如果首年能落地 4 场高客单晚宴，餐厅就回本；如果一场都没有，这块牌就成了纯费用。"
        "所以一滴水谈判的关键不是再要空间，而是把「优先落地」写成最低场次和档期规则。",
        first_line=True,
    )
    md("园区两案很难用招商口径证明对等回报。一滴水的关键是把「优先落地」写成最低场次。")
    md()

    # ------------------------------------------------------------------
    add_heading_cn(doc, "四、分案点评")
    md("## 四、分案点评")
    md()

    add_heading_cn(doc, "（一）森马稿：产业故事对，要价和文本都过满", 2)
    md("### （一）森马稿：产业故事对，要价和文本都过满")
    md()
    add_paragraph(
        doc,
        "潮玩、IP、AIGC 和领事网络之间，确实有一条能讲圆的链条：国产 IP 要出海授权，海外 IP 要来华找渠道和制造，"
        "园区缺的是涉外窗口而不是又一栋办公楼。把会客厅嵌进「孵化—IP 开发—跨境对接」也符合吴泾科技时尚小镇的叙事。"
        "问题是：2000㎡ 专属已经接近小型美术馆或独立会所，不是「会客厅」。"
        "CGC 现有节奏如果仍以午餐会、沙龙、年度慈善晚宴为主，这个面积在非活动日会空转。"
        "空转的空间既耗物业，也损害「永不落幕」的观感——参观的人看到的是装修过的闲置，不是国际枢纽。",
        first_line=True,
    )
    md("潮玩/IP 与领事网络的故事能讲圆，但 2000㎡ 专属接近小型会所。非活动日空转会损害「永不落幕」观感。")
    md()
    add_bullet(doc, "正文 2000㎡、价值章节 1000㎡，必须在发出前改掉，并统一功能分区面积表。", bold_prefix="笔误：")
    add_bullet(doc, "与创智汇同时主打 AI，森马应把主词收缩为「潮玩 IP / 消费时尚」，AI 只作工具而不作会客厅名。", bold_prefix="主题：")
    add_bullet(doc, "森马是服饰上市公司体系，园区企业未必已形成可出海的 IP 供给；要先写清首批 20–30 家目标企业，而不是「潮玩 AI 集群」四个字。", bold_prefix="企业盘：")
    add_bullet(doc, "「海外 IP 引进、跨境版权、授权法规」容易被理解成法律服务，需改为信息分享，并避开无资质咨询承诺。", bold_prefix="过界：")
    md("- **笔误：** 正文 2000㎡、价值章节 1000㎡，发出前必须改掉。")
    md("- **主题：** 与创智汇同时主打 AI，森马应收缩为潮玩 IP / 消费时尚。")
    md("- **企业盘：** 要有首批 20–30 家目标企业，不能只写「潮玩 AI 集群」。")
    md("- **过界：** 版权/授权法规表述需降为信息分享，避免无资质咨询承诺。")
    md()

    add_heading_cn(doc, "（二）创智汇稿：战略价值最高，过会概率最低", 2)
    md("### （二）创智汇稿：战略价值最高，过会概率最低")
    md()
    add_paragraph(
        doc,
        "如果只看区位，创智汇是三案里最值得做的：杨浦三区融合、高校密度、大创智国家级文化产业示范区、区属平台要出对外开放业绩。"
        "会客厅放在这里，和政府外事、投促「错位协同」的说法也最容易被区里听进去。"
        "但文本把最难的事写成了最轻的事：1000㎡ 国有存量物业免租三年、全额装修、现金补助，给一家民间机构专属使用，且允许该机构在同一空间向企业收取市场化服务费。"
        "对国企而言，这不是「合作建议」，这是国有资产使用、采购或合作项目，通常要进决策清单、评估或招投标路径、审计可追溯。"
        "现在的写法过不了北岛科技和杨浦科创集团的法务。",
        first_line=True,
    )
    md("创智汇区位最好，但 1000㎡ 国有物业免租+全额装修+允许在场内收商业服务费，过不了国企法务。")
    md()
    add_bullet(doc, "园区已有中建四局 AI+IP 中心。必须写清会客厅与它是「国际对接层」还是重复建设，否则内部会认为叠床架屋。", bold_prefix="存量冲突：")
    add_bullet(doc, "企业清单写成 AI、数字文创、智能制造、生物医药、跨境贸易，主题却是数字 AI 会客厅，产业口径在撒网。", bold_prefix="产业口过宽：")
    add_bullet(doc, "高校是创智汇独有优势，但「每年举办」没有场次，也没有和复旦、同济哪个学院对接的接口人。", bold_prefix="高校虚写：")
    add_bullet(doc, "「协助对外国际招商推介」会让区投促部门问：线索归谁、算谁的招商口径、失败了谁负责。", bold_prefix="招商权属：")
    md("- **存量冲突：** 须写清与中建四局 AI+IP 中心的分工。")
    md("- **产业口过宽：** 企业清单含生物医药等，与数字 AI 会客厅不完全匹配。")
    md("- **高校虚写：** 无场次、无学院接口人。")
    md("- **招商权属：** 线索归谁、算谁的口径，文本没写。")
    md()

    add_heading_cn(doc, "（三）一滴水稿：结构最健康，交付最容易空转", 2)
    md("### （三）一滴水稿：结构最健康，交付最容易空转")
    md()
    add_paragraph(
        doc,
        "这是三案里唯一「对方付钱、我方导流」的正向现金流结构，也是唯一有场景差异的主题：中阿、滨江宴会、包间闭门。"
        "餐厅要的是高客单包场，CGC 要的是领事宴请和中阿小范围会谈场地，供需匹配比「给科创园塞一个会客厅」真实。"
        "实际控制人的阿联酋地产背景，既可以成为中阿主题的私人背书，也可能把合作目标从餐厅经营悄悄换成个人圈层赋能——"
        "文本已经写了「反哺企业主在阿联酋的海外地产业务」。对 CGC 来说，这是资源；对外口径上，不宜把私人地产利益写进合作价值的主句。",
        first_line=True,
    )
    md("一滴水是唯一正向现金流结构，中阿宴会场景匹配真实。不宜把企业主私人地产利益写成合作主句。")
    md()
    add_bullet(doc, "价值章节出现「CGC 官方授权」。这和全文「民间、非官方外交」直接打架，发出即留下口实。", bold_prefix="口径：")
    add_bullet(doc, "「外滩及陆家嘴滨江」跨越黄浦、浦东，选址表述不严谨，对方和媒体都会追问具体地址。", bold_prefix="地理：")
    add_bullet(doc, "CGC 承诺「优先保障档期」「导入全品类活动」，但没有最低场次、领事出席人数、取消通知时限。", bold_prefix="对赌：")
    add_bullet(doc, "沙龙免费、宴请优惠，未附价目表；「最优质的服务和价格」无法签约。", bold_prefix="价格：")
    add_bullet(doc, "慈善晚宴进商业餐厅，须单列善款账户、义拍合规和发票路径，不能和挂牌费、餐费混在一笔往来里。", bold_prefix="慈善：")
    md("- **口径：** 「CGC 官方授权」与民间定位打架。")
    md("- **地理：** 「外滩及陆家嘴滨江」表述不严谨。")
    md("- **对赌：** 无最低场次和领事出席承诺。")
    md("- **价格：** 无餐费/场地价目表。")
    md("- **慈善：** 善款、挂牌费、餐费必须分账。")
    md()

    # ------------------------------------------------------------------
    add_heading_cn(doc, "五、三案共同的硬伤")
    md("## 五、三案共同的硬伤")
    md()

    add_heading_cn(doc, "1. 套模板痕迹过重，会被对方一眼看穿", 2)
    md("### 1. 套模板痕迹过重，会被对方一眼看穿")
    md()
    add_paragraph(
        doc,
        "前三章逐字级复用，价值章节也是同一套「补齐国际化短板 / 差异化招商 IP / 海外线索导流 / 盘活空间」四段论。"
        "这对内部提效应付可以，拿去见森马集团或杨浦科创集团不够。"
        "产业园总经理会问：你们在几家园区同时铺？我们是不是其中之一？线索真的「优先」给我吗？"
        "现在的文本无法回答，因为根本没有网络治理条款：排他、分轨、冲突活动如何分配。",
        first_line=True,
    )
    md("前三章复用、价值章节四段论同构。没有网络治理：排他、分轨、活动如何分配。")
    md()

    add_heading_cn(doc, "2. 把「领事网络」写成了几乎无限的政府能力", 2)
    md("### 2. 把「领事网络」写成了几乎无限的政府能力")
    md()
    add_paragraph(
        doc,
        "文本反复出现「对接各国中央及地方政府产业部门」「提前解读各国法律法规」「降低投资经营风险」。"
        "CGC 作为民间俱乐部，能做的是组织领事或商务官员出席交流、转介商协会、提供非正式介绍。"
        "不能替代律师、不能出具法律意见、不能承诺中央部委对接、更不能把营商风险「降低」写进对园区的责任条款。"
        "这类句子在招商材料里好看，在合同和事后追责里会反噬。"
        "建议统一降级为：邀请驻沪领馆商务官员或商协会代表做信息分享；深度尽调、法务、投资落地另案、另资质、另收费。",
        first_line=True,
    )
    md("「对接各国中央政府」「解读法律法规」「降低投资风险」超出民间俱乐部能力，须降级为信息分享。")
    md()

    add_heading_cn(doc, "3. 公益与收费共用同一空间，防火墙只有一句话", 2)
    md("### 3. 公益与收费共用同一空间，防火墙只有一句话")
    md()
    add_paragraph(
        doc,
        "分层服务本身是对的：公开沙龙免费，一对一出海、尽调、约访收费，园区不垫付商业服务成本。"
        "但空间、品牌、人员是同一套。园区出钱装修的专属场地，CGC 在里面做收费项目，对方一定会问分成、审计和「公益场是不是获客场」。"
        "现在只写了「不在补贴范围内」，没有写：公益场与商业场的时间隔离、名单隔离、合同主体、发票、是否允许在会客厅内推销。"
        "不写清楚，要么被指责用公共补贴做私活，要么 CGC 自己的销售动作被园区管死。",
        first_line=True,
    )
    md("公益沙龙和收费项目共用空间与品牌，缺少时间/名单/合同/发票隔离规则。")
    md()

    add_heading_cn(doc, "4. 「年度绩效评估」是空壳，3 年锁定期对 CGC 并不有利", 2)
    md("### 4. 「年度绩效评估」是空壳，3 年锁定期对 CGC 并不有利")
    md()
    add_paragraph(
        doc,
        "文本想用 3 年换装修和免租，却把考核权交给对方且不定义指标。结果是：空间和装修沉没在对方资产负债表上，"
        "CGC 只有使用权；对方随时可以用「绩效不佳」在年审时卡续约或要求增加场次。"
        "对 CGC 更稳的结构是：先 6–12 个月活动合作，达到场次和出席门槛再谈专属空间；专属空间先 150–300㎡ 展陈 + 借用公共大厅，而不是 1000㎡ 起。"
        "装修若由对方出，应约定合作期内的使用权保障、提前终止时的补偿，而不是单方面「固定装修不可拆除、CGC 走人」。",
        first_line=True,
    )
    md("3 年锁定期但考核权在对方且无指标。更稳的是先 6–12 个月活动合作，专属空间先 150–300㎡。")
    md()

    add_heading_cn(doc, "5. 没有运营方案：谁开门、谁值守、周历是什么", 2)
    md("### 5. 没有运营方案：谁开门、谁值守、周历是什么")
    md()
    add_paragraph(
        doc,
        "会客厅一旦挂牌，领事、企业、媒体按「点位」来访。文本没有开馆时间、驻场人数、中英文接待标准、安保、礼宾流程、突发舆情预案。"
        "也没有第一年 12 个月的活动日历，甚至没有「每月至少 1 场闭门沙龙」这种最低开工率。"
        "森马稿里唯一的数字是青年沙龙每年 3 场，创智汇和一滴水连这个数字都没有。"
        "对内要先做一份最小运营编制：驻场 2 人 + 活动项目经理 1 人 + 领事关系接口 1 人，并写清工资出自挂牌费、市场化收入还是 CGC 自筹。"
        "每年 10 万挂牌费不够养这个班子，这正好说明 10 万不是运营费，只是品牌符号费——文本却把它写成了「专项运营补助」。名实不符。",
        first_line=True,
    )
    md("无开馆时间、驻场编制、年历。10 万不够养运营班子，却写成「专项运营补助」，名实不符。")
    md()

    # ------------------------------------------------------------------
    add_heading_cn(doc, "六、对方会怎么砍")
    md("## 六、对方会怎么砍")
    md()
    add_paragraph(
        doc,
        "评估策划案不能只看自己想讲什么，还要预演对方会议室里的第一轮反应。下面按「最可能发生」而不是「最客气」来写。",
        first_line=True,
    )
    md("按对方会议室里最可能发生的反应预演，而不是按最客气的反应来写。")
    md()

    add_heading_cn(doc, "森马侧", 2)
    md("### 森马侧")
    md()
    add_paragraph(
        doc,
        "招商和运营会感兴趣「国际化窗口」这个词，财务和法务会卡住免租面积。"
        "可预期的还价：专属面积降到 150–300㎡ 展厅；大型活动用中庭和森马客厅按次审批；取消每年 10 万现金或改成按有效场次报销；"
        "要求 CGC 列出可到场的国别清单和近 12 个月活动业绩；潮玩企业名单由园区提供、CGC 对出席率负责。"
        "如果 CGC 不能接受「先活动、后装修」，这次大概率停在宣传合作而不是物业入驻。",
        first_line=True,
    )
    md("森马可能把面积砍到 150–300㎡，取消或对赌 10 万现金，并要国别清单和历史业绩。")
    md()

    add_heading_cn(doc, "创智汇 / 北岛科技 / 杨浦科创集团侧", 2)
    md("### 创智汇 / 北岛科技 / 杨浦科创集团侧")
    md()
    add_paragraph(
        doc,
        "业务部门可能愿意听，决策链条会拉到国资监管。"
        "可预期的还价：不签物业无偿使用，改「场地支持协议」+ 按次免费或优惠；装修走园区自身公共空间改造，不做成 CGC 专属；"
        "10 万现金改成活动委托服务采购，要发票和成果清单；国际招商线索必须进区里台账；市场化收费须报备或禁止使用「创智汇领事会客厅」品牌销售。"
        "若坚持 1000㎡ 专属，准备被要求资产评估、集体决策和上级报备，周期会以季度计，不是「建议 3 年、马上装修」。",
        first_line=True,
    )
    md("创智汇更可能改成按次场地支持和活动采购，而不是 1000㎡ 无偿专属。现金改委托服务、要发票。")
    md()

    add_heading_cn(doc, "一滴水侧", 2)
    md("### 一滴水侧")
    md()
    add_paragraph(
        doc,
        "老板若看重圈层和阿联酋叙事，10 万挂牌费过关的概率高于两园区。"
        "餐厅经理则会要：每月保底场次、周五周六晚高峰是否占用、取消补偿、领事级别嘉宾的最低人数、菜单和酒水结算周期。"
        "CGC 拥有全部物料终审权，餐厅品牌被置于从属地位，对方公关也会要求对等审核。"
        "这一案最接近能签，前提是 CGC 真的拿得出第一季度的两场可售晚宴，而不是只交一块铜牌。",
        first_line=True,
    )
    md("一滴水过关概率最高，但会要保底场次、高峰档期规则和对等审核。关键是第一季能否交出两场晚宴。")
    md()

    # ------------------------------------------------------------------
    add_heading_cn(doc, "七、合规与品牌风险")
    md("## 七、合规与品牌风险")
    md()
    add_bullet(doc, "「领事会客厅」易被听成领馆延伸。对外物料必须固定「民间平台、不代表领馆立场、不承担外交职能」，禁止「官方授权」「官方外交机构」。", bold_prefix="外事口径：")
    add_bullet(doc, "邀请现任总领事出席商业或招商活动，应走外办报备习惯路径，不要在文本里承诺「各国领事常态化坐镇」。", bold_prefix="领事出席：")
    add_bullet(doc, "创智汇免租 + 装修 + 现金，属于国有资产和可能的利益输送审查面，即使动机正当也要走得出书面程序。", bold_prefix="国资：")
    add_bullet(doc, "在补贴场地上向企业收「出海尽调、投融资撮合」费，涉及咨询与中介资质，不能用俱乐部名义一笔代收。", bold_prefix="经营资质：")
    add_bullet(doc, "慈善晚宴、义拍、自闭症和先心病救助，必须与挂牌费、餐费、市场化服务分账，善款不得冲抵合作成本。", bold_prefix="公益资金：")
    add_bullet(doc, "多点挂牌若质量不齐，金茂旗舰和「总领事文化慈善晚宴」会被连带贬值。宁可少挂，不要满城铜牌。", bold_prefix="品牌稀释：")
    md("- **外事口径：** 禁止「官方授权」；固定民间定位。")
    md("- **领事出席：** 不承诺领事常态化坐镇；现任总领事出席应报备。")
    md("- **国资：** 创智汇免租+装修+现金必须走得出书面程序。")
    md("- **经营资质：** 尽调/撮合不能用俱乐部名义一笔代收。")
    md("- **公益资金：** 善款与挂牌费、餐费分账。")
    md("- **品牌稀释：** 多点挂牌质量不齐会伤及金茂旗舰。")
    md()

    add_callout(
        doc,
        "特别提醒：一滴水稿已出现「CGC 官方授权」",
        "这不是文风问题，是合规问题。只要这一句出现在对方转述、海报或公众号里，"
        "民间定位就从内部原则变成对外争议。发出前必须全局替换为「品牌挂牌合作 / 民间交流基地授权使用」。",
    )
    md("**特别提醒：** 一滴水稿「CGC 官方授权」必须全局替换为「品牌挂牌合作 / 民间交流基地授权使用」。")
    md()

    # ------------------------------------------------------------------
    add_heading_cn(doc, "八、建议的网络策略：不要三店齐开")
    md("## 八、建议的网络策略：不要三店齐开")
    md()
    add_paragraph(
        doc,
        "会客厅适合做成「一个旗舰 + 有限卫星」，不适合做成加盟铜牌。建议内部先定网络规则，再对外谈判：",
        first_line=True,
    )
    md("会客厅应是「一个旗舰 + 有限卫星」，不是加盟铜牌。先定网络规则再谈。")
    md()
    add_bullet(doc, "金茂或现有俱乐部空间继续承担国别接待和会员会籍，不把高客单价冠名降到 10 万去和卫星点比价。", bold_prefix="旗舰：")
    add_bullet(doc, "一滴水做中阿宴会和闭门会谈，定位「场景」不是「园区」。这是最该先打样的卫星点。", bold_prefix="卫星 A：")
    add_bullet(doc, "创智汇做数字经济与高校青年，定位「科创对接」。先活动入驻，专属面积可谈 100–200㎡ 展陈。", bold_prefix="卫星 B：")
    add_bullet(doc, "森马做潮玩 IP 与消费出海，定位「产业」。AI 二字从会客厅名称拿掉，避免和创智汇抢同一句话。", bold_prefix="卫星 C：")
    add_bullet(doc, "同一国别、同一周，原则上只在一个点举办对公活动；海外线索按产业标签分流，并在月报里向各业主披露去向。", bold_prefix="分轨：")
    md("- **旗舰：** 金茂继续承担国别接待和会籍，不把冠名降到 10 万。")
    md("- **卫星 A：** 一滴水做中阿宴会场景，最先打样。")
    md("- **卫星 B：** 创智汇做科创对接，先活动、小展陈。")
    md("- **卫星 C：** 森马做潮玩 IP，名称去掉 AI。")
    md("- **分轨：** 同国别同周不对多点举办对公活动；线索按产业标签分流并月报。")
    md()

    add_caption(doc, "表4  建议谈判底线（对内，不直接贴进对方合同）")
    md("**表4  建议谈判底线（对内）**")
    md()
    headers = ["议题", "不建议再要", "可以接受的替代", "CGC 必须拿到的对价"]
    rows = [
        ["专属面积", "森马 2000㎡、创智汇 1000㎡ 一步到位", "150–300㎡ 展陈 + 公共大厅按次", "书面档期优先权和形象点位"],
        ["装修", "对方全额按会所标准一次到位", "基础展陈由对方做，活动物料 CGC 自担", "合作期内使用权和提前终止补偿"],
        ["10 万现金", "无场次的「挂牌补助」", "按有效场次报销或活动委托采购", "一滴水可保留年费，但绑定保底场次"],
        ["期限", "无 KPI 的 3 年", "12 个月试点 + 达标续约 2 年", "试点期达标标准写进协议"],
        ["收费项目", "在对方补贴空间内自由销售", "商业项目另签、另地或分成", "名单和合同不与公益场混淆"],
    ]
    add_table(doc, headers, rows, col_widths=[2.4, 4.8, 5.0, 4.4])
    md_table(headers, rows)

    # ------------------------------------------------------------------
    add_heading_cn(doc, "九、文本改写清单（发出前必改）")
    md("## 九、文本改写清单（发出前必改）")
    md()
    add_paragraph(doc, "下列各项不改，不建议把 Word 发给对方。改完后每案应能单独回答「你们和其他点有何不同」。", first_line=True)
    md("下列各项不改，不建议发给对方。改完后每案应能单独回答「你们和其他点有何不同」。")
    md()

    add_caption(doc, "表5  三案共性修改")
    md("**表5  三案共性修改**")
    md()
    headers = ["序号", "原文问题", "改法"]
    rows = [
        ["1", "前三章完全复用", "各保留半页 CGC 简介，其余改成该场地专属痛点和首年日历"],
        ["2", "无 KPI", "写进：年度对公活动场次、领事/商协会出席场次、企业参会家次、线索条数、月报"],
        ["3", "「对接中央政府」「降低法律风险」", "改为信息分享与转介，深度服务单列资质和合同"],
        ["4", "营销预算无数字", "要么给年度上限，要么改成按次核定"],
        ["5", "无驻场运营", "补充开馆时间、编制、费用承担、非活动日用途"],
        ["6", "公益/商业无隔离", "补时间、名单、合同主体、发票、场内推销规则"],
        ["7", "无线索权属", "「优先导流」改成书面分流规则和并列抄送"],
        ["8", "无样板业绩", "附 12 个月内 3–5 个可核验活动（时间、主题、嘉宾层级、成果）"],
    ]
    add_table(doc, headers, rows, col_widths=[1.4, 6.0, 9.2])
    md_table(headers, rows)

    add_caption(doc, "表6  分案必改")
    md("**表6  分案必改**")
    md()
    headers = ["案", "必改"]
    rows = [
        ["森马", "统一面积数字；会客厅名称去掉或降级「AI」；专属面积降到可解释规模；补潮玩企业样本名单；青年沙龙 3 场升格为整体年历中的一项而不是唯一 KPI"],
        ["创智汇", "删除或重写无偿 1000㎡ 条款；增加国资使用路径；写明与中建四局 AI+IP 中心的分工；收窄产业口径；高校对接落到具体学院"],
        ["一滴水", "删除「官方授权」；写清具体地址；保底场次（建议试点年不少于 6 场对公活动，其中晚宴不少于 2 场）；附餐费和场地优惠表；私人地产业务移出主价值表述"],
    ]
    add_table(doc, headers, rows, col_widths=[2.4, 14.2])
    md_table(headers, rows)

    # ------------------------------------------------------------------
    add_heading_cn(doc, "十、90 天验证路径")
    md("## 十、90 天验证路径")
    md()
    add_paragraph(
        doc,
        "在重装修和大面积免租之前，用 90 天证明「领事会客厅」不是铜牌。三案可以共用一条验证链，只换场地和主题。",
        first_line=True,
    )
    md("在重装修和大面积免租之前，用 90 天证明会客厅不是铜牌。")
    md()
    add_bullet(doc, "各点签「活动合作备忘录」，不签物业无偿使用。明确试点期场次、取消规则、物料口径。", bold_prefix="第 1–2 周：")
    add_bullet(doc, "一滴水落地 1 场中阿或领事主题晚宴；创智汇或森马落地 1 场 20–30 人闭门沙龙。嘉宾名单须可核验，不做「拟邀请」。", bold_prefix="第 3–6 周：")
    add_bullet(doc, "输出线索台账（不少于 15 条有企业主体的跟进项）和一场复盘，向对方总经理级汇报。同时给出专属点位的最小面积方案。", bold_prefix="第 7–10 周：")
    add_bullet(doc, "达标则谈 12 个月空间试点；不达标则停止装修谈判，保留活动合作。禁止用「先装修再请人」倒逼自己。", bold_prefix="第 11–13 周：")
    md("- **第 1–2 周：** 只签活动合作备忘录，不签物业无偿使用。")
    md("- **第 3–6 周：** 一滴水 1 场晚宴，园区 1 场闭门沙龙，嘉宾可核验。")
    md("- **第 7–10 周：** 线索台账不少于 15 条，并给出最小面积方案。")
    md("- **第 11–13 周：** 达标再谈 12 个月空间试点；禁止先装修再请人。")
    md()

    add_paragraph(
        doc,
        "验证期的成功标准建议写死四条，避免事后各说各话：至少 1 名现任驻沪领事官员或领馆商务负责人出席；"
        "单场企业嘉宾不少于 15 家真实主体；活动后 10 个工作日内发出纪要；对方书面确认愿意进入空间条款讨论。"
        "四条里缺两条，就还不是会客厅，只是一场论坛执行。",
        first_line=True,
    )
    md("成功标准建议写死：现任领事或商务负责人出席；企业嘉宾不少于 15 家；10 个工作日内发纪要；对方书面愿意谈空间条款。缺两条就还不是会客厅。")
    md()

    # ------------------------------------------------------------------
    add_heading_cn(doc, "十一、评估结论")
    md("## 十一、评估结论")
    md()
    add_paragraph(
        doc,
        "这三份策划案证明 CGC 已经从「金茂会所活动」转向「把领事关系产品化到城市空间」。这个转向是对的。"
        "错在用同一套话术同时向民营园区、国有园区和餐厅要三种量级完全不同的资源，却没有匹配三种完全不同的交付、治理和合规方案。"
        "会客厅可以成为长三角民间涉外的物理入口，前提是先做出一个让人愿意再来的现场，而不是先铺三块牌。",
        first_line=True,
    )
    md("转向「把领事关系产品化到城市空间」是对的。错在用同一套话术向三类业主要不同量级资源，却没有匹配不同的交付和合规方案。")
    md()

    add_callout(
        doc,
        "可执行结论",
        "一滴水：小改后可以进入商务谈判，绑定保底场次。"
        "创智汇：大改，先活动后空间，按国资程序重写。"
        "森马：中改，先改笔误和面积，主题与创智汇错开。"
        "三案均不建议以当前文本作为合同附件或董事会汇报终稿。",
        fill="EAF6EE",
        border="A9CDB6",
        title_color=GREEN,
    )
    md("**可执行结论：** 一滴水小改可谈；创智汇大改、先活动后空间；森马中改面积和主题。均不建议以当前文本作为合同附件。")
    md()

    add_heading_cn(doc, "附件：评估范围与方法")
    md("## 附件：评估范围与方法")
    md()
    add_paragraph(
        doc,
        "评估依据为三份原文：《CGC领事会客厅入驻森马上海国际运营中心合作建议》《CGC领事会客厅入驻创智汇合作建议》"
        "《CGC领事会客厅入驻尚9·一滴水合作建议》。未做现场踏勘，未核实 CGC 历史活动出席名单，未向三家业主取证。"
        "租金和装修数字为上海办公及产业园区常见量级的内部匡算，仅用于判断对价是否失衡，不构成估价。"
        "评分是排序工具。若后续补充可核验的领事业绩、企业名单和首季日历，园区两案的「落地可执行」分可明显上修。",
        first_line=True,
    )
    md("评估依据为三份原文。未做现场踏勘，未核实历史出席名单。租金装修为数量级匡算。若补充可核验业绩和年历，园区两案评分可上修。")
    md()

    docx_path = OUTPUT_DIR / "CGC领事会客厅三案评估意见.docx"
    doc.save(docx_path)
    md_path = build_markdown()
    return docx_path, md_path


if __name__ == "__main__":
    docx_path, md_path = build()
    print(f"docx: {docx_path}")
    print(f"md:   {md_path}")
