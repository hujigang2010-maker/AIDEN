#!/usr/bin/env python3
"""生成给护工贴床头的无糖清淡食谱。大字、一页能看完。不是医嘱。"""

from pathlib import Path
import html

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.colors import HexColor, white
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from comm_plan import CN_FONT_PATH

CHINESE_FONT = "宋体"
HEADING_FONT = "黑体"
ACCENT = RGBColor(0x0B, 0x2F, 0x5B)
MUTED = RGBColor(0x5C, 0x6B, 0x7A)
RED = RGBColor(0xA6, 0x3D, 0x2F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

MD_NAME = "给护工的无糖床头食谱_2026-08-29.md"
DOCX_NAME = "青岛抚顺路和哈尔滨路路口交通事故_给护工的无糖床头食谱_20260829.docx"
PDF_NAME = "青岛抚顺路和哈尔滨路路口交通事故_给护工的无糖床头食谱_20260829.pdf"

TITLE = "胡志远床头饮食表（无糖清淡）"
SUB = "给护工贴床头用　患者：胡志远　男　64 岁　糖尿病 + 高血压 + 右踝术后开放伤"
NOTE = "本表不是医院营养科医嘱。和当天医生、护士说法冲突时，听当天医嘱。药按护士发的吃，本表不代替吃药。"

MEALS = [
    ["早餐", "7:00–8:00", "无糖豆浆 200ml（或纯牛奶半杯）+ 水煮蛋 1 个 + 全麦馒头半个或米饭小半碗 + 凉拌黄瓜。豆浆、牛奶都不要加糖、蜂蜜。"],
    ["上午加餐", "10:00", "无糖酸奶半盒，或黄瓜一根。口渴喝白开水，不喝饮料。"],
    ["午餐", "11:30–12:30", "清蒸鱼或鸡胸或瘦肉（一掌心大小）+ 青菜一碟（少油少盐）+ 北豆腐一块 + 米饭半碗到大半碗（按原来降糖量，不要自己加量）。"],
    ["下午加餐", "15:30", "苹果小半个，或猕猴桃 1 个。不要一整盘水果，不要果汁。"],
    ["晚餐", "17:30–18:30", "清炖去油瘦肉汤或豆腐汤 + 青菜 + 虾仁或鱼 + 米饭比午餐再少一点。晚上不喝甜粥、不吃油条包子。"],
    ["睡前", "21:00 前", "原则上不宵夜。若确实饿，半杯纯牛奶。21:00 后少喝水，减少夜尿。"],
]

SWAPS = [
    "蛋白质每餐都要有一样：鸡蛋、鱼、虾、鸡胸、瘦猪肉、北豆腐、无糖豆浆。伤口靠这个长，不是靠骨头汤。",
    "青菜可换：菠菜、油菜、西兰花、冬瓜、黄瓜、番茄、菌菇。炒菜淡口，不另加咸菜、酱豆腐、火腿肠。",
    "主食可换：米饭、馒头、清水煮燕麦。不换成糖醋拌面、甜粥、油条、带馅甜包。",
    "全天白开水约 2000ml，分多次喝，不要一次灌一大壶。可在杯上划线，早中晚各完成一段。",
    "做法：蒸、煮、炖。少油、少盐、少酱油。不许勾芡加糖。骨头汤只喝清汤，不当主食。",
]

BAN = [
    "白糖、红糖、冰糖、蜂蜜、红枣桂圆汤、冰糖梨、阿胶",
    "饮料、果汁、奶茶、甜酸奶、糕点、饼干、面包、罐头水果",
    "荔枝、龙眼、大把香蕉、一整盘西瓜",
    "红烧、糖醋、油炸、肥肉、动物内脏",
    "人参、西洋参、三七粉、药酒、白酒、接骨粉、壮骨粉、氨糖套装",
    "外面送来的补品先放着问家属，护工不要自己给吃",
]

HYPO = (
    "出汗、手抖、心慌、说话含糊、突然说特别饿：立刻按铃叫护士。"
    "若医嘱备有葡萄糖片，按护士交代给 3–4 片，或半杯温糖水。"
    "日常无糖，低血糖时必须给糖。处理完马上告诉家属。"
)


def render_markdown() -> str:
    lines = [
        f"# {TITLE}",
        "",
        SUB,
        "",
        f"> {NOTE}",
        "",
        "## 一天怎么吃（按点做，不要两套叠着吃）",
        "",
        "| 餐次 | 时间 | 就按这个做 |",
        "| --- | --- | --- |",
    ]
    for meal, when, how in MEALS:
        lines.append(f"| {meal} | {when} | {how} |")
    lines.extend(["", "## 可以换、不能省", ""])
    for item in SWAPS:
        lines.append(f"- {item}")
    lines.extend(["", "## 禁止（贴在最显眼处）", ""])
    for item in BAN:
        lines.append(f"- {item}")
    lines.extend(
        [
            "",
            "## 低血糖（比无糖更优先）",
            "",
            HYPO,
            "",
            "降压药、降糖药、欧开、利伐沙班按护士发药时间吃。不要用保健品替换药。",
            "",
            "打印两份：一份贴床头，一份护工口袋。重新生成：`python3 scripts/build_all.py`。",
            "",
        ]
    )
    return "\n".join(lines)


def build_markdown(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(render_markdown(), encoding="utf-8")


def _set_run_font(run, font_name: str, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _add_paragraph(
    doc: Document,
    text: str = "",
    *,
    font: str = CHINESE_FONT,
    size: float = 12,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    space_after: float = 4,
    space_before: float = 0,
    line_spacing: float = 1.25,
    color: RGBColor | None = None,
):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold, color=color)
    return p


def _set_cell_text(cell, text, *, bold=False, size=11, color=None, font=CHINESE_FONT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    _set_run_font(run, font, size, bold=bold, color=color)


def _shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = tcPr.makeelement(qn("w:shd"), {})
        tcPr.append(shd)
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")


def build_document(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = CHINESE_FONT
    normal_style.font.size = Pt(12)
    rpr = normal_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CHINESE_FONT)

    _add_paragraph(doc, TITLE, font=HEADING_FONT, size=20, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, color=ACCENT)
    _add_paragraph(doc, SUB, font=HEADING_FONT, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, color=MUTED)
    _add_paragraph(doc, NOTE, bold=True, size=11, color=RED, space_after=8)

    _add_paragraph(doc, "一天怎么吃（按点做，不要两套叠着吃）", font=HEADING_FONT, size=14, bold=True, space_before=2, space_after=4, color=ACCENT)
    table = doc.add_table(rows=1 + len(MEALS), cols=3)
    table.style = "Table Grid"
    headers = ["餐次", "时间", "就按这个做"]
    widths = [2.8, 3.2, 11.4]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, h, bold=True, size=12, color=WHITE, font=HEADING_FONT)
        _shade(cell, "0B2F5B")
        cell.width = Cm(widths[i])
    for r_i, row in enumerate(MEALS):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            _set_cell_text(cell, val, size=11, bold=(c_i == 0))
            if r_i % 2 == 1:
                _shade(cell, "F3F6FA")
            cell.width = Cm(widths[c_i])

    _add_paragraph(doc, "可以换、不能省", font=HEADING_FONT, size=14, bold=True, space_before=10, space_after=4, color=ACCENT)
    for item in SWAPS:
        _add_paragraph(doc, "• " + item, size=11.5, space_after=3)

    _add_paragraph(doc, "禁止（不要买、不要熬、不要让人送来就吃）", font=HEADING_FONT, size=14, bold=True, space_before=8, space_after=4, color=RED)
    for item in BAN:
        _add_paragraph(doc, "• " + item, size=11.5, space_after=2, color=RED)

    _add_paragraph(doc, "低血糖（比无糖更优先）", font=HEADING_FONT, size=14, bold=True, space_before=8, space_after=4, color=RED)
    _add_paragraph(doc, HYPO, size=12, bold=True, space_after=6, color=RED)
    _add_paragraph(
        doc,
        "降压药、降糖药、欧开、利伐沙班按护士发药时间吃。不要用保健品替换药。打印两份：一份贴床头，一份护工口袋。",
        size=10.5,
        color=MUTED,
        space_before=4,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _pdf_font() -> str:
    kwargs = {"subfontIndex": 0} if CN_FONT_PATH.endswith(".ttc") else {}
    pdfmetrics.registerFont(TTFont("CN", CN_FONT_PATH, **kwargs))
    return "CN"


def _p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(html.escape(str(text)).replace("\n", "<br/>"), style)


def _header_footer(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(HexColor("#0B2F5B"))
    canvas.rect(0, A4[1] - 11 * mm, A4[0], 11 * mm, fill=1, stroke=0)
    canvas.setFillColor(white)
    canvas.setFont("CN", 9)
    canvas.drawString(14 * mm, A4[1] - 7 * mm, "胡志远床头饮食表 · 给护工贴床头 · 无糖清淡")
    canvas.setFillColor(HexColor("#A63D2F"))
    canvas.rect(0, 0, A4[0], 9 * mm, fill=1, stroke=0)
    canvas.setFillColor(white)
    canvas.setFont("CN", 8)
    canvas.drawString(14 * mm, 3.2 * mm, "与医嘱冲突听医嘱 · 低血糖先叫护士给糖 · 不是营养科处方")
    canvas.drawRightString(A4[0] - 14 * mm, 3.2 * mm, f"第 {doc.page} 页")
    canvas.restoreState()


def build_pdf(output_path: Path) -> None:
    font = _pdf_font()
    navy = HexColor("#0B2F5B")
    muted = HexColor("#5C6B7A")
    red = HexColor("#A63D2F")
    dark = HexColor("#1F2A37")
    styles = {
        "title": ParagraphStyle("t", fontName=font, fontSize=16, leading=22, textColor=navy, alignment=TA_CENTER, spaceAfter=3),
        "sub": ParagraphStyle("s", fontName=font, fontSize=10, leading=14, textColor=muted, alignment=TA_CENTER, spaceAfter=5),
        "warn": ParagraphStyle("w", fontName=font, fontSize=10, leading=14, textColor=red, alignment=TA_LEFT, spaceAfter=6),
        "h1": ParagraphStyle("h", fontName=font, fontSize=12.5, leading=17, textColor=navy, spaceBefore=6, spaceAfter=3),
        "h1red": ParagraphStyle("hr", fontName=font, fontSize=12.5, leading=17, textColor=red, spaceBefore=6, spaceAfter=3),
        "bullet": ParagraphStyle("u", fontName=font, fontSize=10.5, leading=15, textColor=dark, leftIndent=2, spaceAfter=2),
        "ban": ParagraphStyle("b", fontName=font, fontSize=10.5, leading=15, textColor=red, leftIndent=2, spaceAfter=2),
        "hypo": ParagraphStyle("y", fontName=font, fontSize=11, leading=16, textColor=red, alignment=TA_LEFT, spaceAfter=4),
        "cell": ParagraphStyle("c", fontName=font, fontSize=9.5, leading=13.5, textColor=dark, alignment=TA_LEFT),
        "cellb": ParagraphStyle("cb", fontName=font, fontSize=10, leading=13.5, textColor=navy, alignment=TA_CENTER),
        "head": ParagraphStyle("hd", fontName=font, fontSize=10, leading=13, textColor=white, alignment=TA_CENTER),
        "foot": ParagraphStyle("f", fontName=font, fontSize=8.5, leading=12, textColor=muted, spaceBefore=6),
    }
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=14 * mm,
        bottomMargin=12 * mm,
        title=TITLE,
        author="内部整理",
    )
    head_row = [
        _p("餐次", styles["head"]),
        _p("时间", styles["head"]),
        _p("就按这个做", styles["head"]),
    ]
    data = [head_row]
    for meal, when, how in MEALS:
        data.append(
            [
                _p(meal, styles["cellb"]),
                _p(when, styles["cellb"]),
                _p(how, styles["cell"]),
            ]
        )
    table = Table(data, colWidths=[28 * mm, 32 * mm, 118 * mm])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), navy),
                ("TEXTCOLOR", (0, 0), (-1, 0), white),
                ("BACKGROUND", (0, 1), (-1, 1), HexColor("#F3F6FA")),
                ("BACKGROUND", (0, 3), (-1, 3), HexColor("#F3F6FA")),
                ("BACKGROUND", (0, 5), (-1, 5), HexColor("#F3F6FA")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("GRID", (0, 0), (-1, -1), 0.4, HexColor("#C5CED6")),
            ]
        )
    )
    story = [
        _p(TITLE, styles["title"]),
        _p(SUB, styles["sub"]),
        _p(NOTE, styles["warn"]),
        _p("一天怎么吃（按点做，不要两套叠着吃）", styles["h1"]),
        table,
        Spacer(1, 4),
        _p("可以换、不能省", styles["h1"]),
    ]
    for item in SWAPS:
        story.append(_p("• " + item, styles["bullet"]))
    story.append(_p("禁止（不要买、不要熬、不要让人送来就吃）", styles["h1red"]))
    for item in BAN:
        story.append(_p("• " + item, styles["ban"]))
    story.append(_p("低血糖（比无糖更优先）", styles["h1red"]))
    story.append(_p(HYPO, styles["hypo"]))
    story.append(
        _p(
            "降压药、降糖药、欧开、利伐沙班按护士发药时间吃。不要用保健品替换药。打印两份：一份贴床头，一份护工口袋。",
            styles["foot"],
        )
    )
    doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "deliverables"
    build_markdown(out / MD_NAME)
    build_document(out / DOCX_NAME)
    build_pdf(out / PDF_NAME)
    print(f"Wrote {out / MD_NAME}")
    print(f"Wrote {out / DOCX_NAME}")
    print(f"Wrote {out / PDF_NAME}")


if __name__ == "__main__":
    main()
