#!/usr/bin/env python3
"""核对发给对方的修改建议是否覆盖关键修改点。"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]


def main() -> int:
    md = (ROOT / "docs/CGC领事会客厅合作方案修改建议.md").read_text(encoding="utf-8")
    doc = Document(ROOT / "output/CGC领事会客厅合作方案修改建议.docx")
    text = md + "\n" + "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + " ".join(c.text for c in row.cells)

    present = [
        ("签约拨款装修边界", "暂不作为签约、拨款或启动装修"),
        ("面积笔误", "2000㎡"),
        ("在建工程进度", "52.19%"),
        ("一滴水地址", "东大名路500号"),
        ("品牌授权", "CGC品牌授权"),
        ("公开募捐", "公开募捐"),
        ("三年成本口径", "三年总成本"),
        ("森马试运营面积", "300—500"),
        ("一滴水场次", "3—5 场"),
        ("删除地产表述", "地产"),
        ("创智汇先借会议室", "借用现有会议"),
        ("KPI表", "有效线索"),
    ]
    absent = [
        ("无内部打分4/10", "4/10"),
        ("无内部打分3/10", "3/10"),
    ]

    fail = 0
    print("FEEDBACK VERIFICATION")
    for name, needle in present:
        ok = needle in text
        print(("PASS" if ok else "FAIL"), name)
        fail += 0 if ok else 1
    for name, needle in absent:
        ok = needle not in text
        print(("PASS" if ok else "FAIL"), name)
        fail += 0 if ok else 1
    html = ROOT / "output/CGC领事会客厅合作方案修改建议.html"
    ok = html.exists() and html.stat().st_size > 2000
    print(("PASS" if ok else "FAIL"), "HTML")
    fail += 0 if ok else 1
    total = len(present) + len(absent) + 1
    print(f"\n{total - fail}/{total} passed")
    return fail


if __name__ == "__main__":
    sys.exit(main())
