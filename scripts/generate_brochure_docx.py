# -*- coding: utf-8 -*-
"""生成《复旦杨浦科创生态共建计划》宣传册（DOCX）。

视觉：复旦蓝 + 金色点缀；面向网上传播与线下打印。
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "deliverables" / "复旦杨浦科创生态共建计划_宣传册.docx"

NAVY = RGBColor(0x0F, 0x2E, 0x5C)
NAVY_DEEP = RGBColor(0x08, 0x1C, 0x3A)
GOLD = RGBColor(0xC8, 0x96, 0x3E)
INK = RGBColor(0x1A, 0x1D, 0x21)
GREY = RGBColor(0x5C, 0x65, 0x70)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SOFT = RGBColor(0xF5, 0xF7, 0xFA)

FONT = "微软雅黑"


def set_run(run, *, size=11, bold=False, color=INK, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_para(
    doc,
    text,
    *,
    size=11,
    bold=False,
    color=INK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    space_before=0,
    space_after=8,
    line=1.55,
):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return p


def shade_paragraph(paragraph, hex_color: str):
    """给段落加底色（用于色块标题）。"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def set_cell_shading(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, *, size=10.5, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)


def add_section_label(doc, num: str, title: str):
    add_para(doc, f"{num}　{title}", size=10, bold=True, color=GOLD, space_before=18, space_after=4)


def add_h2(doc, text: str):
    add_para(doc, text, size=16, bold=True, color=NAVY, space_before=2, space_after=10)


def add_divider_note(doc, text: str):
    p = add_para(doc, text, size=11, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
    shade_paragraph(p, "0F2E5C")
    return p


def build():
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    # ===== 封面 =====
    cover = add_para(
        doc,
        "联合发起　·　面向杨浦　·　链接复旦",
        size=10,
        bold=True,
        color=GOLD,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=36,
        space_after=18,
    )
    shade_paragraph(cover, "081C3A")

    add_para(
        doc,
        "复旦杨浦",
        size=28,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=28,
        space_after=0,
    )
    add_para(
        doc,
        "科创生态共建计划",
        size=28,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0,
        space_after=14,
    )
    add_para(
        doc,
        "以住房政策研究为纽带，连接科技企业、产业资源与高端客户群",
        size=12,
        color=GREY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    add_para(
        doc,
        "让资源相遇，让合作发生",
        size=13,
        bold=True,
        color=GOLD,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=18,
    )

    org = doc.add_table(rows=3, cols=1)
    org.autofit = True
    set_cell_shading(org.cell(0, 0), "0F2E5C")
    set_cell_text(org.cell(0, 0), "联合组织单位", size=10, bold=True, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(org.cell(1, 0), "0F2E5C")
    set_cell_text(
        org.cell(1, 0),
        "复旦大学住房政策研究中心　×　杨浦区科技企业联合会",
        size=12,
        bold=True,
        color=WHITE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_shading(org.cell(2, 0), "081C3A")
    set_cell_text(
        org.cell(2, 0),
        "秘书长 / 会长　联合发起",
        size=10,
        color=RGBColor(0xC8, 0xD0, 0xDA),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_para(doc, "", space_after=6)
    kw = add_para(
        doc,
        "促招商　｜　强链条　｜　聚人群",
        size=12,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=10,
        space_after=4,
    )
    shade_paragraph(kw, "F5F7FA")
    add_para(
        doc,
        "线上宣传册｜公众号「复联会」配套传播",
        size=9,
        color=GREY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=18,
    )

    # ===== 01 =====
    add_section_label(doc, "01", "为什么是现在")
    add_h2(doc, "把「资源相遇」，变成「生态共建」")
    add_para(
        doc,
        "一个区域的科创活力，从来不是靠企业数量堆出来的。真正决定高度的，是连接的密度。",
        size=12,
        bold=True,
        color=INK,
        space_after=10,
    )
    add_para(
        doc,
        "科技企业的成长，离不开空间、人才、资本、场景与产业协同。杨浦拥有深厚的科创基因与丰富的产业腹地；复旦汇聚政策研究能力、学术智力与高层次人才网络。",
        color=GREY,
    )
    add_para(
        doc,
        "双方联合发起本计划，是为了搭建一个更高质量、更可持续的连接平台——不是办一场活动，而是运转一条生态链。",
        color=GREY,
    )

    add_divider_note(doc, "核心判断｜招商不是一次性「引进」，而是围绕企业全生命周期，持续链接空间、政策、资本、技术、人才与客户。")

    add_para(doc, "我们希望促成的三类结果", size=12, bold=True, color=NAVY, space_before=14, space_after=8)
    results = [
        ("落地连接", "企业与杨浦产业空间、政策服务、应用场景建立明确联系。"),
        ("合作线索", "企业与高校院所、投资机构、上下游伙伴形成可跟进的合作线索。"),
        ("长期圈层", "参与者从「来参加一次活动」，走向「进入一个长期有价值的圈层」。"),
    ]
    t = doc.add_table(rows=3, cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(results):
        set_cell_shading(t.cell(i, 0), "0F2E5C")
        set_cell_text(t.cell(i, 0), k, size=11, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(t.cell(i, 1), v, size=10.5, color=GREY)

    # ===== 02 =====
    add_section_label(doc, "02", "三大目标")
    add_h2(doc, "我们只做三件事")
    add_para(
        doc,
        "三件事，一条逻辑：让有价值的连接，反复发生。",
        size=11,
        color=GREY,
        space_after=10,
    )

    goals = [
        ("01", "促进招商", "引资源", "让优质科技企业更了解杨浦、走进杨浦、落地杨浦。不做单向招引，而是把空间、政策、场景、客户一次讲清楚——企业自己会算这笔账。"),
        ("02", "促进生态链达成", "强协同", "让企业与产业伙伴、科研机构、资本方及应用场景高效对接，形成可跟进、可落地的合作线索。上下游在同一个房间里，很多事情不需要三个月。"),
        ("03", "促进高端客户群形成", "聚人群", "以高质量活动聚合高质量人群，沉淀长期互信与合作关系。圈层不是被组织出来的，是被反复见面沉淀出来的。"),
    ]
    for idx, title, tag, desc in goals:
        p = add_para(doc, f"{idx}　{title}　·　{tag}", size=12, bold=True, color=NAVY, space_before=10, space_after=4)
        shade_paragraph(p, "F5F7FA")
        add_para(doc, desc, size=10.5, color=GREY, space_after=6)

    # 招商四要素
    add_para(doc, "目标一深化｜把「引进来」做成「留得住」", size=12, bold=True, color=NAVY, space_before=12, space_after=8)
    four = doc.add_table(rows=2, cols=2)
    four.style = "Table Grid"
    items = [
        ("空间", "产业载体、办公与中试场地匹配"),
        ("政策", "专项扶持、人才政策与申报辅导"),
        ("场景", "区域应用场景与首试首用机会"),
        ("客户", "本地及长三角高质量客户资源"),
    ]
    for i, (k, v) in enumerate(items):
        r, c = divmod(i, 2)
        set_cell_text(four.cell(r, c), f"{k}\n{v}", size=10.5, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(four.cell(r, c), "F5F7FA")

    # ===== 03 =====
    add_section_label(doc, "03", "活动怎么做")
    add_h2(doc, "从一次相聚，到一条可持续的生态链")
    add_para(
        doc,
        "活动不止于主题分享，更强调「认识 → 对话 → 匹配 → 跟进」的完整闭环。以小而精的组织方式，提升每一次见面的有效性。人少一点，有效性高一点——这是刻意的选择。",
        color=GREY,
    )

    steps = [
        ("1　主题引导", "围绕住房政策、科技产业、企业发展与区域机会，输出有价值的趋势判断。"),
        ("2　企业呈现", "精选科技企业与产业项目，讲清技术、产品、场景与合作诉求。"),
        ("3　资源对接", "组织企业、投资、科研、园区、服务机构与高端客户进行定向交流。"),
        ("4　持续跟进", "沉淀需求清单、合作线索与后续拜访机制，让活动成果可追踪。"),
    ]
    for title, desc in steps:
        add_para(doc, title, size=11, bold=True, color=NAVY, space_before=8, space_after=2)
        add_para(doc, desc, size=10.5, color=GREY, space_after=2)

    add_para(doc, "建议活动单元", size=12, bold=True, color=NAVY, space_before=14, space_after=8)
    units = doc.add_table(rows=4, cols=2)
    units.style = "Table Grid"
    unit_rows = [
        ("主旨分享", "复旦智力与杨浦实践的交汇　｜　约 40 分钟 · 1–2 位主讲"),
        ("圆桌对话", "科技企业如何在区域中找到成长坐标　｜　约 45 分钟 · 4–5 位嘉宾"),
        ("项目路演", "企业把需求说清楚，把机会讲明白　｜　每家约 8 分钟 · 5–6 家"),
        ("闭门交流", "面向重点企业与高端客户的定向洽谈　｜　限额邀请制"),
    ]
    for i, (k, v) in enumerate(unit_rows):
        set_cell_shading(units.cell(i, 0), "0F2E5C")
        set_cell_text(units.cell(i, 0), k, size=11, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(units.cell(i, 1), v, size=10.5, color=GREY)

    add_divider_note(doc, "活动关键词｜专业　·　开放　·　精准　·　持续")

    # ===== 04 =====
    add_section_label(doc, "04", "谁应该来到这里")
    add_h2(doc, "让真正有合作可能的人，坐到同一张桌旁")
    add_para(
        doc,
        "我们欢迎愿意在杨浦寻找机会、在复旦链接资源、在产业生态中共同成长的伙伴。",
        color=GREY,
    )

    audiences = [
        ("科技企业", "寻找落地空间、产业协同、政策服务、投融资与高端客户。"),
        ("投资与金融机构", "发现具有成长潜力的项目，建立长期项目观察与服务关系。"),
        ("高校院所与科研团队", "推动技术成果、人才资源与企业应用场景有效连接。"),
        ("产业链伙伴", "围绕供应链、客户、渠道、技术和服务展开合作。"),
        ("园区与专业服务机构", "为企业提供空间、人才、法务、财税、品牌等综合支持。"),
        ("高端客户群", "以专业判断、产业视野与高质量社交，参与更有价值的生态网络。"),
    ]
    aud = doc.add_table(rows=len(audiences), cols=2)
    aud.style = "Table Grid"
    for i, (k, v) in enumerate(audiences):
        set_cell_shading(aud.cell(i, 0), "F5F7FA")
        set_cell_text(aud.cell(i, 0), k, size=10.5, bold=True, color=NAVY)
        set_cell_text(aud.cell(i, 1), v, size=10.5, color=GREY)

    add_para(doc, "参与者将获得什么", size=12, bold=True, color=NAVY, space_before=14, space_after=8)
    gains = [
        ("区域理解", "获得区域机会与政策环境的第一手理解，判断更准，决策更快。"),
        ("合作机会", "与目标伙伴深入交流、验证需求，形成可执行的合作方案。"),
        ("长期平台", "进入由复旦智力、杨浦产业与企业家网络共同支撑的连接平台。"),
    ]
    g = doc.add_table(rows=3, cols=2)
    g.style = "Table Grid"
    for i, (k, v) in enumerate(gains):
        set_cell_shading(g.cell(i, 0), "0F2E5C")
        set_cell_text(g.cell(i, 0), k, size=11, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(g.cell(i, 1), v, size=10.5, color=GREY)

    add_divider_note(doc, "我们寻找的伙伴｜愿意开放资源、认真对话、长期合作——把「单点机会」变成「生态增量」。")

    # ===== 05 =====
    add_section_label(doc, "05", "活动信息")
    add_h2(doc, "报名与合作咨询")

    info = doc.add_table(rows=8, cols=2)
    info.style = "Table Grid"
    info_rows = [
        ("活动时间", "【待定】"),
        ("活动地点", "【待定】"),
        ("活动形式", "主旨分享 · 圆桌对话 · 项目路演 · 闭门交流"),
        ("参与规模", "【待定】人（限额，需审核）"),
        ("参与对象", "科技企业｜投资机构｜高校院所｜产业伙伴｜园区及专业服务机构｜高端客户"),
        ("联系人", "【姓名】"),
        ("联系电话", "【电话】"),
        ("微信 / 公众号", "【微信】　｜　公众号「复联会」"),
    ]
    for i, (k, v) in enumerate(info_rows):
        set_cell_shading(info.cell(i, 0), "F5F7FA")
        set_cell_text(info.cell(i, 0), k, size=10.5, bold=True, color=NAVY)
        set_cell_text(info.cell(i, 1), v, size=10.5, color=GREY)

    add_para(doc, "", space_after=10)
    end = add_para(
        doc,
        "资源不会自己相遇，合作不会自动发生。\n但只要把对的人放在一起，剩下的事情，会比想象中快。",
        size=13,
        bold=True,
        color=WHITE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=8,
        space_after=8,
        line=1.7,
    )
    shade_paragraph(end, "081C3A")

    add_para(
        doc,
        "复旦杨浦科创生态共建计划",
        size=14,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=14,
        space_after=4,
    )
    add_para(
        doc,
        "联合组织：复旦大学住房政策研究中心　｜　杨浦区科技企业联合会",
        size=10,
        color=GREY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_para(
        doc,
        "关注「复联会」，获取活动信息与合作机会",
        size=11,
        bold=True,
        color=GOLD,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"已生成：{OUT}")


if __name__ == "__main__":
    build()
