"""Generate a concise (≈3-page) version of the Guansong (华为鸿蒙智行授权
合作伙伴) Diamond Sponsorship + Official Intelligent Mobility Partner
Agreement. Used solely for showing during Greentown's bid-evaluation, in
matching format with the Tencent Cloud concise version."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# 华为/鸿蒙 brand red, 区分腾讯云蓝
HUAWEI_RED = RGBColor(0xC7, 0x00, 0x0B)
DARK_RED = RGBColor(0x9A, 0x00, 0x08)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)


def _force_font(run, font="宋体"):
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rpr.append(rfonts)


def set_default_font(doc, name="宋体"):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def heading(doc, text, level=1, color=None, center=False, font="黑体"):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(18)
    elif level == 1:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    if color is not None:
        run.font.color.rgb = color
    _force_font(run, font=font)
    return p


def para(doc, text, bold=False, size=10.5, color=DARK, align=None,
         first_indent=True, font="宋体", line_spacing=1.4):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(2)
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    _force_font(run, font=font)
    return p


def clause(doc, num, text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(2)
    run_num = p.add_run(num + "  ")
    run_num.bold = True
    run_num.font.size = Pt(size)
    _force_font(run_num)
    run = p.add_run(text)
    run.font.size = Pt(size)
    _force_font(run)
    return p


def blank(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, color=DARK, size=10.5,
              align_center=False, font="宋体"):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    _force_font(run, font=font)


def build():
    doc = Document()
    set_default_font(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # ===== 抬头 =====
    heading(doc, "2026 人工智能商业化落地与硬核投资破局峰会",
            level=2, color=GREY, center=True)
    heading(doc, "《官方智慧出行伙伴合作协议》",
            level=0, color=HUAWEI_RED, center=True)
    blank(doc)
    para(doc,
         "合同编号：AIBIZ-2026-SP-008      签署地点：上海市      签署日期：2026 年 5 月 8 日",
         first_indent=False, size=10, align="center", color=GREY)
    blank(doc)

    # ===== 协议主体 =====
    para(doc,
         "甲方（主办方 / 组委会）：人工智能商业化落地峰会组委会 / 复旦大学住房政策研究中心 / 上海市杨浦区科技企业联合会（联合主办）",
         bold=True, first_indent=False, size=10.5)
    para(doc,
         "乙方（赞助方）：上海冠松汽车股份有限公司（华为鸿蒙智行授权合作伙伴）",
         bold=True, first_indent=False, size=10.5)
    blank(doc)

    # ===== 鉴于 =====
    para(doc, "鉴于：", bold=True, first_indent=False)
    para(doc,
         "1. 甲方拟于 2026 年 5 月在上海·北外滩核心地标举办「2026 人工智能商业化落地与硬核投资破局峰会」（以下简称「本次峰会」），峰会汇聚顶尖政企智库资源与长三角前沿科技产融生态；")
    para(doc,
         "2. 乙方系国内领先的汽车经销商集团及华为鸿蒙智行授权合作伙伴，希望通过本次峰会向高净值人群展示华为鸿蒙智行系列（尊界、问界等）的智能座舱与高阶智驾体验；")
    para(doc,
         "3. 双方依据《中华人民共和国民法典》《广告法》《反不正当竞争法》《商标法》及其他相关法律法规，本着平等自愿、优势互补、资源共享的原则，就乙方作为本次峰会「官方指定智慧出行独家伙伴」事宜，达成如下协议，以资共同遵守。")
    blank(doc)

    # ===== 第一条 =====
    heading(doc, "第一条  合作级别与赞助标的", level=1, color=HUAWEI_RED)
    clause(doc, "1.1",
           "合作级别：乙方为本次峰会及高净值私局的「钻石赞助商」及「官方指定智慧出行独家伙伴」，享有同一级别下不可被并列、不可被超越的智慧出行类排他权益。")
    clause(doc, "1.2",
           "综合赞助标的额：乙方为本次活动提供专项赞助（含现金及高端接驳车队等重资产实物调用对价），总计金额为人民币（大写）贰拾万元整（小写：￥200,000.00 元），含税。")
    clause(doc, "1.3",
           "支付与对价交付方式：现金部分应于本协议签署并经双方盖章生效后 15 个工作日内一次性汇入甲方指定的对公账户；实物调用对价部分（车队及司机服务）由乙方在峰会执行期间按本协议第二条约定直接交付。甲方在现金部分到账后向乙方开具相应金额的合规发票（发票内容：会议服务费 / 赞助费）。")
    clause(doc, "1.4",
           "资金与资源定向用途：上述 20 万元赞助对价仅限用于且全部消耗于本次峰会当期的专属落地执行，涵盖但不限于：（1）北外滩核心地标的专属江景车辆特装展位费及动线规划费；（2）峰会当期华为鸿蒙智行高端车队的独家调度运营（含车辆折旧成本、全险、专职司机劳务及高端接待标准成本）；（3）峰会现场高净值 VIP 专属试乘试驾动线的安保及综合运营保障支出。")
    blank(doc)

    # ===== 第二条 =====
    heading(doc, "第二条  甲方权益交付（全包制 · 不作分项标价）", level=1, color=HUAWEI_RED)
    para(doc,
         "为对应乙方上述赞助投入，甲方在本次峰会现场及其后续生态互动中，向乙方综合交付包括但不限于以下权益。双方确认：下列权益作为一揽子综合对价整体交付，不再就单项权益进行分项估值与单独定价，亦不作为单项退款依据。")
    clause(doc, "2.1",
           "核心场景曝光：乙方享有本次峰会主视觉背景板、签到处及官方宣传中的「官方指定智慧出行伙伴」专属顶级 logo 露出。")
    clause(doc, "2.2",
           "高端实车特展：甲方在北外滩会场核心区域为乙方划定专属江景实车展位，用于全方位展示乙方华为鸿蒙智行系列（如尊界、问界 M9 等）的智能座舱与高阶智驾体验。")
    clause(doc, "2.3",
           "高净值圈层现场导流：甲方利用自身政商顶级圈层，在峰会期间为乙方定向引荐具备大宗换车及企业集中采购需求的高净值人群（重点覆盖头部房企高管、算力大厂代表、科创企业创始人及顶尖智库学者），并由乙方团队在现场闭门对接。")
    clause(doc, "2.4",
           "长效生态延伸权益（战略附赠，不计入第 1.2 条赞助标的额）：（1）乙方自动成为「见微知海新质商业生态」年度战略理事单位，享有未来半年内组委会举办的其他系列闭门私局的优先参与权；（2）针对后续组委会生态内成功跑通的重资产企业、科创园区及超级个体，如有高端商务用车集中采购或高管换车需求，甲方将优先向乙方及华为鸿蒙智行体系进行线索导入与商业撮合；（3）甲方智库团队在进行「智慧城市与新质人居」相关学术及商业课题调研时，优先将乙方及华为生态作为研究样板与联合发布方。")
    blank(doc)

    # ===== 第三条 =====
    heading(doc, "第三条  双方义务", level=1, color=HUAWEI_RED)
    clause(doc, "3.1",
           "甲方义务：（1）严格按本协议第二条约定，向乙方完整、及时交付各项权益；（2）保障峰会按计划如期举办，并维护峰会的行业影响力与规格档次；（3）在峰会筹备、执行、后续传播全过程中，确保乙方作为「官方指定智慧出行独家伙伴」在智慧出行类别下的排他性与唯一性，不引入与乙方在高端整车展示及智慧出行接驳领域形成直接竞争关系的同级别合作方。")
    clause(doc, "3.2",
           "乙方义务：（1）按本协议约定按时足额支付赞助款项及交付实物调用对价；（2）配合甲方完成 logo、车型简介、试乘试驾动线等权益落地所需的素材提供与确认工作；（3）确保峰会期间投入的车辆、司机均具备合法上路资质与全险保障，承担车辆运营期间的安全与合规责任；（4）不得利用本次峰会平台从事违反国家法律法规及公序良俗的活动。")
    blank(doc)

    # ===== 第四条 =====
    heading(doc, "第四条  知识产权、保密及独立比对声明", level=1, color=HUAWEI_RED)
    clause(doc, "4.1",
           "知识产权：乙方授权甲方在本次峰会的宣传、报道、白皮书、官网、议程手册、媒体通稿、回顾视频等场景中使用乙方「冠松汽车」及与之相关的华为鸿蒙智行授权标识与车型素材；授权范围限于本次峰会及其衍生宣传内容，授权期限至本次峰会结束后 12 个月。")
    clause(doc, "4.2",
           "商业机密保护：本协议第一条体现了甲乙双方最核心的高端重资产调度折算成本与赞助对价构成，未经双方书面同意，任何一方不得向任何第三方（包括但不限于其他赞助商、合作方、媒体及外部审计机构）披露本协议项下的实物折算明细清单、底层财务核算逻辑及打款凭证等敏感信息，保密期限自本协议签订之日起 3 年。")
    clause(doc, "4.3",
           "独立横向比对豁免：为促成更广泛的产融生态建设，甲方有权在对外招商及面向同级别头部企业（如头部商业地产、大型科技巨头、金融集团等）进行横向比对与资质说明时，合法出示本协议并引用本协议第一条项下的合作总金额（贰拾万元）及合作级别（钻石赞助商 + 官方指定智慧出行独家伙伴），以证明本组委会系列活动的顶级市场公允估值与合作门槛。该等引用不视为对本条第 4.2 款保密义务的违反。")
    clause(doc, "4.4",
           "合规承诺：双方承诺，本协议项下的资金往来与权益交付均合法合规，不存在任何商业贿赂、利益输送及违反反不正当竞争法、反商业贿赂相关规定的情形。")
    blank(doc)

    # ===== 第五条 =====
    heading(doc, "第五条  违约责任与不可抗力", level=1, color=HUAWEI_RED)
    clause(doc, "5.1",
           "任何一方未按本协议约定履行义务，给对方造成损失的，应承担相应的违约责任并赔偿对方因此遭受的直接经济损失。")
    clause(doc, "5.2",
           "若因不可抗力（包括但不限于自然灾害、政府管制、重大公共卫生事件等）导致本次峰会延期或无法举办，双方应友好协商，按已发生的实际支出比例处理已支付的赞助款项及已动用的实物对价；剩余款项可顺延用于甲方下一届同级别峰会，或经双方书面同意后退还乙方。")
    blank(doc)

    # ===== 第六条 =====
    heading(doc, "第六条  争议解决与法律适用", level=1, color=HUAWEI_RED)
    clause(doc, "6.1",
           "本协议的签订、履行、解释及争议解决均适用中华人民共和国法律。")
    clause(doc, "6.2",
           "因本协议产生或与本协议有关的任何争议，双方应首先通过友好协商解决；协商不成的，任何一方均有权向甲方所在地有管辖权的人民法院提起诉讼。")
    blank(doc)

    # ===== 第七条 =====
    heading(doc, "第七条  签署与生效", level=1, color=HUAWEI_RED)
    clause(doc, "7.1",
           "生效条件：本协议自甲乙双方法定代表人或授权代表签字并加盖公章（或合同专用章、组委会印章）之日起生效。")
    clause(doc, "7.2",
           "物理要求：本协议一式两份，甲乙双方各执一份，具有同等法律效力。")
    clause(doc, "7.3",
           "附件效力：本协议如有附件、补充协议或经双方书面确认的往来函件（含电子邮件、企业微信、钉钉等电子形式），均为本协议不可分割的组成部分，与本协议具有同等法律效力。")
    blank(doc)

    # ===== 签署页 =====
    para(doc, "（以下无正文，下为签署页）",
         align="center", first_indent=False, size=9, color=GREY)
    blank(doc)

    sig = doc.add_table(rows=5, cols=2)
    sig.style = "Table Grid"
    sig.autofit = False
    for row in sig.rows:
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(8.0)
    cell_text(sig.rows[0].cells[0], "甲方（盖章）：", bold=True, size=11)
    cell_text(sig.rows[0].cells[1], "乙方（盖章）：", bold=True, size=11)
    cell_text(sig.rows[1].cells[0],
              "人工智能商业化落地峰会组委会 /\n复旦大学住房政策研究中心 /\n上海市杨浦区科技企业联合会",
              size=10.5)
    cell_text(sig.rows[1].cells[1],
              "上海冠松汽车股份有限公司\n（华为鸿蒙智行授权合作伙伴）",
              size=10.5)
    cell_text(sig.rows[2].cells[0],
              "\n\n\n（公章 / 合同章位置）\n\n",
              size=10.5, align_center=True)
    cell_text(sig.rows[2].cells[1],
              "\n\n\n（公章 / 合同章位置）\n\n",
              size=10.5, align_center=True)
    cell_text(sig.rows[3].cells[0],
              "授权代表（签字）：______________________",
              size=11)
    cell_text(sig.rows[3].cells[1],
              "授权代表（签字）：______________________",
              size=11)
    cell_text(sig.rows[4].cells[0],
              "日期：2026 年 ____ 月 ____ 日",
              size=11)
    cell_text(sig.rows[4].cells[1],
              "日期：2026 年 ____ 月 ____ 日",
              size=11)

    out = "/workspace/deliverables/冠松集团-官方智慧出行伙伴合作协议(精简版).docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
