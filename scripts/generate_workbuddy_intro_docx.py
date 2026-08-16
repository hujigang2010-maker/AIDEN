#!/usr/bin/env python3
"""生成《WorkBuddy 银行引荐判断备忘录》：要不要引荐中行、上海银行，以及下周五见杨行长怎么打。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


CHINESE_FONT = "宋体"
HEADING_FONT = "黑体"
ACCENT = RGBColor(0x0B, 0x2F, 0x5B)
MUTED = RGBColor(0x5C, 0x6B, 0x7A)


def _set_run_font(run, font_name: str, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _add_paragraph(
    doc: Document,
    text: str = "",
    *,
    font: str = CHINESE_FONT,
    size: float = 12,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: float | None = None,
    space_after: float = 6,
    space_before: float = 0,
    line_spacing: float = 1.5,
    color: RGBColor | None = None,
):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if text:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold, color=color)
    return p


def _add_rich_paragraph(
    doc: Document,
    segments: list[tuple[str, bool]],
    *,
    font: str = CHINESE_FONT,
    size: float = 12,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: float | None = None,
    space_after: float = 6,
    space_before: float = 0,
    line_spacing: float = 1.5,
    color: RGBColor | None = None,
):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    for text, bold in segments:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold, color=color)
    return p


def _heading(doc: Document, text: str, level: int = 1) -> None:
    if level == 1:
        _add_paragraph(
            doc,
            text,
            font=HEADING_FONT,
            size=14,
            bold=True,
            space_before=14,
            space_after=8,
            color=ACCENT,
            line_spacing=1.3,
        )
    else:
        _add_paragraph(
            doc,
            text,
            font=HEADING_FONT,
            size=12,
            bold=True,
            space_before=10,
            space_after=6,
            line_spacing=1.3,
        )


def _body(doc: Document, text: str) -> None:
    _add_paragraph(doc, text, first_line_indent=0.74, space_after=6)


def _bullet(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    if bold_prefix:
        _add_rich_paragraph(
            doc,
            [("• ", False), (bold_prefix, True), (text, False)],
            first_line_indent=0.37,
            space_after=3,
            line_spacing=1.4,
        )
    else:
        _add_paragraph(
            doc,
            f"• {text}",
            first_line_indent=0.37,
            space_after=3,
            line_spacing=1.4,
        )


def build_document(output_path: Path) -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = CHINESE_FONT
    normal_style.font.size = Pt(12)
    rpr = normal_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CHINESE_FONT)

    _add_paragraph(
        doc,
        "WorkBuddy 银行引荐判断备忘录",
        font=HEADING_FONT,
        size=18,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
        line_spacing=1.3,
        color=ACCENT,
    )
    _add_paragraph(
        doc,
        "主线：下周五见泰隆上海分行杨行长　辅线：中国银行、上海银行是否现在引荐",
        font=HEADING_FONT,
        size=11,
        bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
        line_spacing=1.3,
        color=MUTED,
    )
    _add_rich_paragraph(
        doc,
        [("内部材料 · 引荐人自用　", False), ("日期：2026年8月16日　约见日：8月21日（周五）", False)],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
        line_spacing=1.3,
        size=11,
        color=MUTED,
    )

    _heading(doc, "一、要回答的问题")
    _body(
        doc,
        "目前已帮腾讯对接 WorkBuddy 与泰隆银行：东口支行行长已经见过，下个周五（8月21日）要约见上海分行杨行长。"
        "腾讯这边最近思考时间比较长。因此要判断：要不要现在再引荐中国银行、上海银行，让腾讯再聊两家。",
    )
    _body(
        doc,
        "这不是“多介绍两家总没错”的问题，而是三条线会不会互相打架：泰隆这条已经走到分行行长，"
        "腾讯内部还没想清楚，再塞两家银行进去，可能加速，也可能把主线冲散。",
    )

    _heading(doc, "二、一句话结论")
    _add_paragraph(
        doc,
        "可以引荐，但不要本周并行约见。先把下周五杨行长这条主线走实；"
        "中国银行和上海银行先做预热名单和轻量探询，等杨行长会后再决定是否正式对接。"
        "两家里，上海银行更适合做第二主线，中国银行只适合做探索性接触。",
        font=HEADING_FONT,
        size=12,
        bold=True,
        first_line_indent=0.74,
        space_after=8,
        line_spacing=1.5,
        color=ACCENT,
    )
    _body(
        doc,
        "不要用“再找两家银行”去催腾讯。如果他们慢，是因为方案、合规、价格还没包装好，"
        "多见面只会让他们在更大的银行面前显得没准备好。如果他们慢，是因为还缺银行侧真实需求，"
        "那才值得加线，但也要排在杨行长会后。",
    )

    _heading(doc, "三、当前态势")
    _heading(doc, "（一）已经走完的台阶", level=2)
    _bullet(doc, "产品侧：腾讯 WorkBuddy，企业级 AI 应用及技能平台，正在找银行渠道落地。", bold_prefix="产品侧：")
    _bullet(doc, "银行侧：已对接浙江泰隆商业银行；合作方向此前已形成积分商城兑换、开户赠礼、员工份额、联名信用卡等讨论框架。", bold_prefix="银行侧：")
    _bullet(doc, "支行已见面：东口支行行长已经见过，说明一线经营单位接得住、话题成立。", bold_prefix="支行已见面：")
    _bullet(doc, "下一步：8月21日见上海分行杨行长。这是从网点验证上升到分行决策的关键一跳。", bold_prefix="下一步：")
    _bullet(doc, "腾讯节奏：最近思考时间比较长，商务方案、内部授权、报价档位或合规口径可能还没锁。", bold_prefix="腾讯节奏：")

    _heading(doc, "（二）为什么杨行长会比再引荐两家更值钱", level=2)
    _body(
        doc,
        "支行行长能确认“这个东西一线愿不愿意推”；分行行长才能确认“这个东西上海要不要立项”。"
        "泰隆已经付出了两轮洽谈、支行见面、材料准备的沉没成本，下周五是把这些成本兑现成方向的窗口。"
        "中国银行、上海银行现在还停留在“有没有人愿意见面”的阶段，和泰隆不在同一能级。",
    )
    _body(
        doc,
        "对引荐人来说，价值不在于同一周内堆多少局，而在于让腾讯在一家已经热起来的银行面前，"
        "把能说的说清楚、把要的决策要回来。杨行长会如果只变成“再熟悉一下”，后面再开两家，也还是熟悉。",
    )

    _heading(doc, "四、腾讯想得久：先分清是哪一种慢")
    _body(
        doc,
        "引荐前，先用一句话问腾讯：最近想得比较长，主要卡在哪一块？"
        "两种病因对策完全相反，不能混着治。",
    )

    _heading(doc, "病因 A：缺银行侧真实需求（加线有用）", level=2)
    _body(
        doc,
        "表现：内部其实想推，但还在问“银行到底买不买账”“除了泰隆还有没有人愿意谈”。"
        "这时再引荐上海银行、中国银行，能给腾讯证明需求，也能降低“只押泰隆一家”的心理负担。",
    )
    _bullet(doc, "对策：会后正式引荐，先上海银行，再中国银行；本周只预热，不占腾讯会前精力。")

    _heading(doc, "病因 B：内部包装没好（加线有害）", level=2)
    _body(
        doc,
        "表现：价格档位、发票结算、数据合规、品牌授权、谁来拍板还没对齐；对银行只能“再沟通一下”。"
        "这时再开中国银行、上海银行，对方问三句就会露出缺口。大行会把“没准备好”记成第一印象，以后更难约。",
    )
    _bullet(doc, "对策：本周只帮腾讯把杨行长会上的“可说 / 不可说”清单收紧，不新开战场。")

    _body(
        doc,
        "在没问清之前，默认按病因 B 处理：保护主线，不扩面。问清是病因 A 之后，再把辅线从预热升到正式约见。",
    )

    _heading(doc, "五、现在引荐的收益和风险")
    _heading(doc, "（一）收益：为什么会想引荐", level=2)
    _bullet(doc, "降低单点依赖：泰隆若在分行层放缓，腾讯不会觉得整条银行线停了。")
    _bullet(doc, "给腾讯一点外部节奏：有第二家在排队，内部讨论会从“要不要做银行”变成“先做哪家”。")
    _bullet(doc, "验证模式可复制：积分权益、开户赠礼、员工提效如果只有泰隆接，腾讯会怀疑这是特例。")
    _bullet(doc, "引荐人自己的位置：持续给腾讯输送银行资源，而不是只绑在泰隆一单上。")

    _heading(doc, "（二）风险：为什么本周不宜并行约见", level=2)
    _bullet(doc, "冲主线：腾讯会前精力被两家新银行分走，杨行长会上仍是“再想想”。")
    _bullet(doc, "暴露准备不足：大行问采购主体、数据出境、品牌联名，腾讯若答不利落，比不见面更亏。")
    _bullet(doc, "伤泰隆体感：若杨行长侧听到“同时还在见中行、上海银行”，会觉得自己只是样品，不愿拍板。")
    _bullet(doc, "把引荐用成施压：腾讯会读成“你们太慢，我另找人”，关系从帮忙变成催单。")
    _bullet(doc, "能级错配：中国银行决策链长、合规门槛高，现在还不是拿来逼单的对手盘。")

    _heading(doc, "六、三家银行怎么排")
    _heading(doc, "（一）泰隆银行：本周唯一主线", level=2)
    _body(
        doc,
        "城商行、主做中小微，和 WorkBuddy 的客群最贴。已经完成支行见面，材料也做过一轮。"
        "下周五见杨行长，目标不是签 1,500 万，而是拿到分行层的方向：原则继续、指定牵头部门、明确试点范围。",
    )
    _bullet(doc, "本周动作：确认时间、出席名单、一页纸议程；逼腾讯给出会上可承诺的边界。")
    _bullet(doc, "会上要的三件事：谁牵头、先试点哪些网点/客群、下一次商务谁来谈。")
    _bullet(doc, "不要在会上把中国银行、上海银行当作筹码。")

    _heading(doc, "（二）上海银行：最合适的第二主线", level=2)
    _body(
        doc,
        "总部在上海的城商行，决策半径比国有大行短，对本地中小微、零售权益、员工数字化都有现成入口，"
        "和泰隆的合作叙事接近，腾讯不用另做一套完全不同的故事。适合作为“泰隆如果放缓，下一手跟谁谈”的备胎，"
        "也适合作为“模式能不能复制到第二家城商行”的验证。",
    )
    _bullet(doc, "本周：只做轻量探询。确认对方条线（公司金融 / 零售权益 / 金融科技 / 办公室）和合适层级，不约正式会。")
    _bullet(doc, "会后：若杨行长给出正向信号，上海银行仍可慢热，作为备份。若杨行长礼貌但空，则把上海银行升为正式引荐。")
    _bullet(doc, "引荐层级建议：先分行相关部门或支行一把手，不要一上来就总行。")

    _heading(doc, "（三）中国银行：只做探索，不当本周对手盘", level=2)
    _body(
        doc,
        "品牌重、流程长、合规严。WorkBuddy 若还没有标准化的银行渠道方案，见中行很容易变成“了解一下”之后长期无下文。"
        "中行的价值是打开国有大行的认知，不是本周用来制造紧迫感。引荐时必须落到具体单位："
        "上海某家支行、普惠金融、公司金融或银行卡/权益部门，而不是“中国银行”四个字。",
    )
    _bullet(doc, "本周：最多确认“有没有合适的人愿意先聊 30 分钟”，不发正式方案，不拉腾讯高管。")
    _bullet(doc, "会后：仅在腾讯已能讲清产品边界、数据不出行、采购档位之后，再安排探索性沟通。")
    _bullet(doc, "不要承诺中行能复制泰隆的采购规模或联名卡节奏。")

    _heading(doc, "七、建议打法：主线走实，辅线预热")
    _body(doc, "把动作拆成三档，避免“引荐”和“约见”混为一谈。")

    _heading(doc, "（一）本周（8月16日—21日）：只做一件大事", level=2)
    _bullet(doc, "主线：服务好杨行长约见。确认时间地点、双方出席、一页纸材料、腾讯可说边界。")
    _bullet(doc, "辅线：中国银行、上海银行只进“预热名单”。发一条探询，问对方是否方便会后再安排交流。")
    _bullet(doc, "对腾讯：明确告诉他们，本周不安排第二场银行正式会，避免他们分心。")
    _bullet(doc, "问清病因：用一句非指责的话问卡点，是缺银行样本，还是内部方案没齐。")

    _heading(doc, "（二）会后 3 个工作日：按信号升级或按住", level=2)
    _bullet(doc, "若杨行长原则同意继续：泰隆仍是主线；上海银行保持预热；中行继续放着。")
    _bullet(doc, "若杨行长要内部再议、没有牵头人：启动上海银行正式引荐，给腾讯第二条活线。")
    _bullet(doc, "若会上暴露腾讯讲不清价格/合规：暂停一切新引荐，先帮腾讯补一页“银行合作标准说法”。")

    _heading(doc, "（三）不要做的四件事", level=2)
    _bullet(doc, "不要在杨行长会前，把腾讯拉去见中行或上海银行。")
    _bullet(doc, "不要对两家新银行说“腾讯已经要采购了”或报出 1,500 万框架。")
    _bullet(doc, "不要对泰隆说“我们同时在谈中行和上海银行，请尽快拍板”。")
    _bullet(doc, "不要把引荐当催单工具。催单用会前清单，不用新银行。")

    _heading(doc, "八、下周五杨行长约见怎么打")
    _heading(doc, "（一）会议目标（按优先级）", level=2)
    _add_paragraph(doc, "1. 原则同意：把 WorkBuddy 作为中小微客户增值服务方向，继续论证。", first_line_indent=0.37, space_after=3, line_spacing=1.4)
    _add_paragraph(doc, "2. 指定牵头：明确分行侧牵头部门或牵头人，避免会后回到“再研究”。", first_line_indent=0.37, space_after=3, line_spacing=1.4)
    _add_paragraph(doc, "3. 划定试点：同意先在部分网点或客群小范围试点，而不是一次谈全年采购。", first_line_indent=0.37, space_after=3, line_spacing=1.4)
    _add_paragraph(doc, "4. 约下次：确定商务/科技/合规谁来参加下一轮，给出时间窗口。", first_line_indent=0.37, space_after=6, line_spacing=1.4)
    _body(
        doc,
        "东口支行行长已见过，这是会上最有用的事实：不是空对空，而是一线已经验证过话题。"
        "开场用这个铺垫，再请杨行长拍方向。",
    )

    _heading(doc, "（二）建议议程（60—90 分钟）", level=2)
    _add_paragraph(doc, "1. 5 分钟：引荐人说明来意——支行已见面，请分行定方向。", first_line_indent=0.37, space_after=3, line_spacing=1.4)
    _add_paragraph(doc, "2. 15 分钟：腾讯讲 WorkBuddy 是什么、银行能用在哪（积分、开户赠礼、员工提效）。", first_line_indent=0.37, space_after=3, line_spacing=1.4)
    _add_paragraph(doc, "3. 15 分钟：结合泰隆客群，只讲试点怎么做，不讲全年总包。", first_line_indent=0.37, space_after=3, line_spacing=1.4)
    _add_paragraph(doc, "4. 15 分钟：杨行长和分行同事提问。", first_line_indent=0.37, space_after=3, line_spacing=1.4)
    _add_paragraph(doc, "5. 10 分钟：请行长定三件事——是否继续、谁牵头、试点范围。", first_line_indent=0.37, space_after=6, line_spacing=1.4)

    _heading(doc, "（三）可说 / 不可说", level=2)
    _add_rich_paragraph(doc, [("可说：", True)], space_before=4, space_after=2, line_spacing=1.3)
    _bullet(doc, "东口支行行长已交流，一线认为对中小微客户有差异化权益价值。")
    _bullet(doc, "主方向仍是积分兑换、开户/拜访赠 AI 权益；员工应用和联名卡可以研究，不要求当场拍板。")
    _bullet(doc, "建议先试点、后扩量；采购规模以试点效果为准。")
    _bullet(doc, "行外场景用兑换码/License，客户自助激活；不接行内核心系统。")
    _add_rich_paragraph(doc, [("不可说：", True)], space_before=6, space_after=2, line_spacing=1.3)
    _bullet(doc, "不要把 1,500 万说成已定预算或腾讯报价承诺。")
    _bullet(doc, "不要承诺“行业首创”“腾讯联合产品”等未授权口径。")
    _bullet(doc, "不要主动提中国银行、上海银行，更不要用别家催泰隆。")
    _bullet(doc, "不要把行内办公系统对接说成马上能做；这涉及合规，只能列为后续评估。")
    _bullet(doc, "不要请杨行长当场定采购金额。")

    _heading(doc, "（四）会前必须问腾讯的五句话", level=2)
    _add_paragraph(doc, "1. 杨行长若问“你们希望我们做什么”，标准回答是哪一句？", first_line_indent=0.37, space_after=3, line_spacing=1.4)
    _add_paragraph(doc, "2. 首期试点，你们最低希望银行做什么、最高不能承诺什么？", first_line_indent=0.37, space_after=3, line_spacing=1.4)
    _add_paragraph(doc, "3. 价格、发票、结算，会上能不能讲，还是一律说会后书面给？", first_line_indent=0.37, space_after=3, line_spacing=1.4)
    _add_paragraph(doc, "4. 数据存在哪、日志谁能看、是否出域，谁来答、答到哪一步？", first_line_indent=0.37, space_after=3, line_spacing=1.4)
    _add_paragraph(doc, "5. 最近想得比较长，卡点是缺银行样本，还是内部方案没齐？", first_line_indent=0.37, space_after=6, line_spacing=1.4)

    _heading(doc, "九、给三方的口径")
    _heading(doc, "（一）给腾讯", level=2)
    _body(
        doc,
        "泰隆下周五见杨行长是本周唯一正式场。支行已经见过，分行这场是定方向的。"
        "中国银行和上海银行我可以帮你们预热，先确认合适的人，不占你们会前时间。"
        "等泰隆信号出来，再决定要不要正式约。你们这边如果卡在方案和合规，我先不把大行带上场。",
    )

    _heading(doc, "（二）给上海银行（预热，会后再正式引荐）", level=2)
    _body(
        doc,
        "腾讯 WorkBuddy 在看银行渠道，怎么把企业 AI 做成客户权益和员工提效。"
        "上海已经有一家城商行在交流。想先了解你们公司金融或零售权益条线，是否方便做一次业务沟通。"
        "不是采购谈判，先看场景合不合。",
    )

    _heading(doc, "（三）给中国银行（探索性接触）", level=2)
    _body(
        doc,
        "想介绍一个企业 AI 工具，看是否适合作为客户增值或网点服务补充。"
        "先找上海地区具体支行或普惠/公司金融条线做一次轻量交流，不谈总行级合作，也不谈采购额。",
    )

    _heading(doc, "（四）给泰隆（只用于杨行长会，不提另外两家）", level=2)
    _body(
        doc,
        "东口支行已经沟通过。今天请杨行长看两件事：这个方向值不值得在上海试点；如果值得，谁来牵头。"
        "金额、系统对接、联名卡都可以后置，先定方向。",
    )

    _heading(doc, "十、十天行动清单")
    _heading(doc, "8月16日—17日（今日到周一）", level=2)
    _bullet(doc, "向腾讯确认周五出席名单，并问清“思考时间长”的卡点。")
    _bullet(doc, "把会上可说 / 不可说发给腾讯，请他们书面点头。")
    _bullet(doc, "确认杨行长会的时间、地点、分行出席人员。")
    _bullet(doc, "上海银行、中国银行：只整理人选，不发正式邀约。")

    _heading(doc, "8月18日—20日（周二到周四）", level=2)
    _bullet(doc, "准备一页纸：背景、试点建议、请行长定的三件事。")
    _bullet(doc, "与腾讯对一次模拟提问：采购额、数据安全、和谁签合同。")
    _bullet(doc, "对上海银行、中国银行最多发一条预热消息，把正式会放在 8 月 25 日之后。")

    _heading(doc, "8月21日（周五，杨行长会）", level=2)
    _bullet(doc, "开场用东口支行已见面作铺垫，请杨行长定方向。")
    _bullet(doc, "会中盯三件事：牵头人、试点范围、下次谁来。")
    _bullet(doc, "会后当天出半页纪要，发给腾讯和自己留底。")

    _heading(doc, "8月24日—25日（会后第一个工作日窗口）", level=2)
    _bullet(doc, "按会后信号决定：按住辅线，或把上海银行升为正式引荐。")
    _bullet(doc, "中国银行仍保持探索，除非腾讯明确说需要国有大行样本。")

    _heading(doc, "十一、会后决策树")
    _body(doc, "用杨行长会的结果，而不是用焦虑，来决定要不要引荐。")
    _bullet(doc, "原则同意 + 有牵头人：泰隆继续深挖；上海银行预热即可；中行暂缓。", bold_prefix="原则同意 + 有牵头人：")
    _bullet(doc, "态度正向但要内部再议：给泰隆 1—2 周，同时把上海银行约到探索性沟通。", bold_prefix="态度正向但要内部再议：")
    _bullet(doc, "礼貌、无牵头、无下次：正式引荐上海银行；中行只做一人探询。", bold_prefix="礼貌、无牵头、无下次：")
    _bullet(doc, "会上腾讯答不利落：停止新引荐，先补标准说法，再见任何人。", bold_prefix="会上腾讯答不利落：")

    _heading(doc, "十二、结论")
    _body(
        doc,
        "引荐中国银行和上海银行，方向是对的，时间不是现在这一周。"
        "腾讯思考时间长，不能靠把更大的银行推进来催；要先分清他们是缺样本，还是缺方案。",
    )
    _body(
        doc,
        "下周五见杨行长，是目前最接近“银行真正开始立项”的动作。"
        "东口支行已经验证过一线，分行这场要把验证变成方向。"
        "上海银行留作第二主线，中国银行留作品牌型探索。主次一乱，三家都容易变成聊天。",
    )
    _add_paragraph(
        doc,
        "口头一句：本周只打泰隆杨行长；中行和上海银行先预热、会后再引荐；"
        "两家里先上海银行，中国银行不当催单工具。",
        font=HEADING_FONT,
        size=12,
        bold=True,
        first_line_indent=0.74,
        space_before=8,
        space_after=10,
        line_spacing=1.5,
        color=ACCENT,
    )

    _add_paragraph(
        doc,
        "说明：本备忘录依据当面已知进度整理，供引荐人内部判断。"
        "不代表腾讯或任何银行的正式立场；1,500 万等数字仅来自既有讨论框架，不是采购承诺。"
        "杨行长、东口支行均按现有约见口径表述，对外材料以对方名片和职务为准。",
        size=10,
        color=MUTED,
        space_before=12,
        space_after=0,
        line_spacing=1.3,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"已生成：{output_path}")


if __name__ == "__main__":
    out = (
        Path(__file__).resolve().parents[1]
        / "deliverables"
        / "WorkBuddy银行引荐判断备忘录_杨行长约见_20260816.docx"
    )
    build_document(out)
