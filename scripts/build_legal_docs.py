"""
冠松 GS · iDrive Hub 合作协议 Word 文档生成脚本

生成 4 份核心法律文件（草稿口径，供法务定稿前的业务讨论使用）：
1. 链主总部租赁合同（甲档样张）
2. 中介居间服务协议（5 家中介通用）
3. 联合实验室共建协议（3F · 链主 + 我方）
4. 政府专班合作备忘录（区政府 + 我方）

注：所有金额、面积、政策条款均为策划阶段口径，签约前由法务最终定稿。
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

CN_FONT = "WenQuanYi Micro Hei"
EN_FONT = "Calibri"

OUT_DIR = Path(__file__).resolve().parent.parent / "docs" / "legal"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def set_cn_font(run, size=11, bold=False, color=None):
    run.font.name = EN_FONT
    rPr = run._element.get_or_add_rPr()
    for tag in ("ea", "cs"):
        existing = rPr.find(qn(f"w:rFonts"))
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_title(doc, text, size=20):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=True)


def add_subtitle(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=False, color=RGBColor(0x6B, 0x73, 0x80))


def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, size=14, bold=True, color=RGBColor(0x0F, 0x2D, 0x52))


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, size=12, bold=True)


def add_p(doc, text, indent=False, bold=False, italic=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_cn_font(r, size=11, bold=bold)
    r.italic = italic


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_cn_font(r, size=11)


def add_kv_table(doc, kv_list):
    table = doc.add_table(rows=len(kv_list), cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11)
    for i, (k, v) in enumerate(kv_list):
        cell_k = table.rows[i].cells[0]
        cell_v = table.rows[i].cells[1]
        cell_k.text = ""
        cell_v.text = ""
        rk = cell_k.paragraphs[0].add_run(k)
        set_cn_font(rk, size=11, bold=True)
        rv = cell_v.paragraphs[0].add_run(v)
        set_cn_font(rv, size=11)


def add_signature_block(doc, parties):
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=len(parties))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["（盖章）", "法定代表人 / 授权代表：", "签署日期：", ""]
    party_titles = parties
    for j, p in enumerate(party_titles):
        cell = table.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(p)
        set_cn_font(r, size=12, bold=True)
    for i in range(1, 4):
        for j in range(len(parties)):
            cell = table.rows[i].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(headers[i])
            set_cn_font(r, size=11)


def doc_setup(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    rPr.append(rFonts)


# ======================================================================
# 1. 链主总部租赁合同
# ======================================================================
def build_anchor_lease():
    doc = Document()
    doc_setup(doc)
    add_title(doc, "房 屋 租 赁 合 同")
    add_subtitle(doc, "GS · iDrive Hub · 01# 研发楼 · 链主总部（甲档）样张")
    add_subtitle(doc, "草案版本：v1.0 · 仅供业务讨论 · 签约前由法务最终定稿")

    add_p(doc, "")
    add_h2(doc, "合同编号：GS-LH-XXXX-001")

    add_h1(doc, "签约方信息")
    add_kv_table(doc, [
        ("出租方（甲方）", "上海冠松 [项目运营公司全称] 有限公司"),
        ("注册地址", "上海市静安区永和路 [号待补]"),
        ("法定代表人", "[姓名]"),
        ("统一社会信用代码", "[91310106XXXXXXXXXX]"),
        ("承租方（乙方）", "[链主企业全称]"),
        ("注册地址", "[乙方注册地]"),
        ("法定代表人", "[姓名]"),
        ("统一社会信用代码", "[XXXXX]"),
    ])

    add_h1(doc, "鉴于")
    add_p(doc,
          "1. 甲方系上海市静安区永和社区 N070601 单元 075b-07 地块『01# 研发楼』"
          "（地上 9F 共 15,152.75 ㎡，地下 2F 共 6,992.87 ㎡）的合法权属人；",
          indent=True)
    add_p(doc,
          "2. 乙方系智能驾驶产业链链主企业，拟在上海中心城区设立总部级研发中心；",
          indent=True)
    add_p(doc,
          "3. 双方就乙方租赁甲方所属 8F+9F 两整层（共约 3,366 ㎡）"
          "并配套相关产业服务事宜，经平等友好协商，达成本合同。",
          indent=True)

    add_h1(doc, "第一条 · 租赁标的")
    add_kv_table(doc, [
        ("标的名称", "01# 研发楼 · 8F + 9F 整层"),
        ("出租面积", "约 3,366 ㎡（其中套内 ~2,800 ㎡，公摊 ~566 ㎡，按实测交付为准）"),
        ("净层高", "8F 4.2 m / 9F 4.3 m"),
        ("用地性质", "C6 教育科研设计用地"),
        ("装配式 / 绿建", "100% 装配整体式框架核心筒 / 绿建二星"),
        ("配套车位", "B2 专属智驾车位 8–12 个（含充电桩）"),
        ("交付标准", "毛坯（机电、消防、弱电按附件 E 实测）"),
    ])

    add_h1(doc, "第二条 · 租赁期限")
    add_p(doc, "2.1 租赁期共 6 年，自 [起始日] 起至 [结束日] 止，含免租期 15 个月。", indent=True)
    add_p(doc, "2.2 交付日：[YYYY-MM-DD]", indent=True)
    add_p(doc, "2.3 进场日（装修起算）：[YYYY-MM-DD]", indent=True)
    add_p(doc, "2.4 起租日（免租期满）：[YYYY-MM-DD]", indent=True)

    add_h1(doc, "第三条 · 租金与免租")
    add_p(doc, "3.1 起始租金：人民币 5.8 元/㎡·天（不含税不含物业费）；", indent=True)
    add_p(doc, "3.2 物业费：人民币 28 元/㎡·月，按月支付，由乙方承担；", indent=True)
    add_p(doc, "3.3 免租期：自交付日起 15 个月（含装修期）；", indent=True)
    add_p(doc, "3.4 调价机制：每 3 年一调，调幅 = max(CPI, 5%)；", indent=True)
    add_p(doc, "3.5 租金支付：按月支付，每月 5 日前支付当月租金；", indent=True)
    add_p(doc, "3.6 滞纳金：万分之五/日；逾期 30 日以上甲方有权单方解除合同。", indent=True)

    add_h1(doc, "第四条 · 履约保证金")
    add_p(doc, "4.1 履约保证金：6 个月租金（人民币 [金额]）；", indent=True)
    add_p(doc, "4.2 保证金支付：合同签订后 30 日内一次性支付；", indent=True)
    add_p(doc, "4.3 扣减顺序：欠付租金 → 物业费 → 设施损坏赔偿 → 其他违约金；", indent=True)
    add_p(doc, "4.4 退还：合同终止且无未结清款项后 30 日内无息退还。", indent=True)

    add_h1(doc, "第五条 · 装修与装补")
    add_p(doc, "5.1 装补金额：1,000 元/㎡，封顶人民币 336 万元；", indent=True)
    add_p(doc, "5.2 装补支付节奏：", indent=True)
    add_bullet(doc, "进场施工 30%")
    add_bullet(doc, "完工验收 30%")
    add_bullet(doc, "起租日 30%")
    add_bullet(doc, "起租满 12 个月 10%")
    add_p(doc, "5.3 装修方案：乙方装修方案需经甲方书面审批后施工，符合装配式建筑接口要求；", indent=True)
    add_p(doc, "5.4 恢复原状：合同终止时，乙方装修无需恢复原状（甲方接受现状）。", indent=True)

    add_h1(doc, "第六条 · 用途相符性（C6 用地特别约定）")
    add_p(doc,
          "标的位于 C6 教育科研设计用地，乙方仅限将标的用于：研发设计、设计服务、"
          "第三方测试认证、产学研合作、博士后工作站等符合 C6 主导功能的业态，"
          "不得用于：纯商务总部（除经主管部门书面认可外）、制造、仓储、教培、零售。"
          "如政府主管部门提出整改，乙方应配合整改。",
          indent=True)

    add_h1(doc, "第七条 · 转租与分租")
    add_p(doc, "7.1 原则禁止；", indent=True)
    add_p(doc, "7.2 经甲方书面同意，可分租至乙方关联公司；", indent=True)
    add_p(doc, "7.3 不得超出原合同条件分租。", indent=True)

    add_h1(doc, "第八条 · 退出与续约")
    add_p(doc, "8.1 优先续约权：到期前 12 个月乙方可优先续约 5 年；", indent=True)
    add_p(doc, "8.2 提前退租赔偿：", indent=True)
    add_bullet(doc, "合同生效 5 年内：偿还未摊销装补 + 80% 已发生政策返还")
    add_bullet(doc, "合同生效 5 年后：偿还未摊销装补的 50%")
    add_p(doc, "8.3 装修撤场赔偿：乙方装修资产视情况由甲方按净值收购或拆除恢复。", indent=True)

    add_h1(doc, "第九条 · 产业服务包")
    add_p(doc, "服务包详细内容见附件 A。包括但不限于：", indent=True)
    add_bullet(doc, "1F 大堂主背景墙冠名权 5 年")
    add_bullet(doc, "9F 屋顶花园专属使用日 ≥ 2 次/年")
    add_bullet(doc, "嘉定/临港封闭测试场会员代办（终身免费）")
    add_bullet(doc, "1.5 km 静安区路测延伸路段联合申请")
    add_bullet(doc, "共享算力券配额 50 万元/年")
    add_bullet(doc, "数据合规沙盒白名单 + 优先")
    add_bullet(doc, "招聘联运 + 落户绿通 30 个/年")
    add_bullet(doc, "9 月发布会主旨演讲位 + 媒体首发权")

    add_h1(doc, "第十条 · 政策返还（一企一策）")
    add_p(doc,
          "10.1 区级税收留成 80% 三年返、50% 后两年返，"
          "以静安区政府专班书面文件（附件 B）为准；",
          indent=True)
    add_p(doc,
          "10.2 一次性落户奖励、人才公寓配额、L3/L4 试点联合体推荐 等，"
          "以一企一策书面协议为准；",
          indent=True)
    add_p(doc,
          "10.3 政策返还不构成甲方独立担保义务。"
          "如政府未按期兑现，甲方协助乙方追索，但不承担连带支付责任。",
          indent=True)

    add_h1(doc, "第十一条 · 品牌权益")
    add_p(doc, "11.1 楼宇外立面冠名「[链主] · iDrive Tower 01」5 年；", indent=True)
    add_p(doc, "11.2 1F 大堂主背景墙冠名 5 年；", indent=True)
    add_p(doc, "11.3 9 月年度发布会主旨演讲 + 媒体首发权；", indent=True)
    add_p(doc, "11.4 年度站台 ≥ 2 次。", indent=True)

    add_h1(doc, "第十二条 · 合规与数据")
    add_p(doc, "双方共同遵守：", indent=True)
    add_bullet(doc, "《反不正当竞争法》《反垄断法》《反贿赂法》")
    add_bullet(doc, "《个人信息保护法》《数据安全法》")
    add_bullet(doc, "《出口管制法》《外商投资法》")
    add_bullet(doc, "智能网联汽车测试相关法规")

    add_h1(doc, "第十三条 · 不可抗力")
    add_p(doc,
          "因不可抗力（含突发疫情、重大灾害、政府征收、用地用途调整等）"
          "导致合同无法履行，双方互不追究违约责任，但应及时通知对方并采取"
          "合理措施降低损失。",
          indent=True)

    add_h1(doc, "第十四条 · 违约与解除")
    add_p(doc, "14.1 乙方违约情形：", indent=True)
    add_bullet(doc, "无故拖欠租金、物业费超 30 日")
    add_bullet(doc, "未经书面同意擅自转租、改变用途")
    add_bullet(doc, "重大安全 / 合规事件")
    add_p(doc, "14.2 甲方违约情形：", indent=True)
    add_bullet(doc, "未按交付标准交付")
    add_bullet(doc, "未按合同协助办理政策返还、人才公寓等")
    add_bullet(doc, "重大安保 / 物业事故")
    add_p(doc, "14.3 违约金：违约方支付未执行租金的 20% 作为违约金；", indent=True)
    add_p(doc, "14.4 解除条件：单方违约且经书面催告 30 日仍未改正的，守约方可解除合同。", indent=True)

    add_h1(doc, "第十五条 · 通知与送达")
    add_p(doc, "本合同所有书面通知发至以下地址或邮箱视为送达：", indent=True)
    add_kv_table(doc, [
        ("甲方收件地址", "[出租方地址]"),
        ("甲方邮箱", "[contact@guansong.com]"),
        ("乙方收件地址", "[承租方地址]"),
        ("乙方邮箱", "[xxx@xxx.com]"),
    ])

    add_h1(doc, "第十六条 · 争议解决")
    add_p(doc,
          "因本合同发生的任何争议，应首先通过友好协商解决。协商不成的，"
          "提交上海仲裁委员会按其届时有效的仲裁规则进行仲裁，仲裁裁决为终局，"
          "对双方均有约束力。",
          indent=True)

    add_h1(doc, "第十七条 · 附则")
    add_p(doc, "17.1 本合同自双方法定代表人或授权代表签字盖章之日起生效；", indent=True)
    add_p(doc, "17.2 本合同一式四份，双方各执两份；", indent=True)
    add_p(doc, "17.3 附件与本合同具有同等法律效力。", indent=True)

    add_h1(doc, "附件清单")
    add_bullet(doc, "附件 A：产业服务包细则")
    add_bullet(doc, "附件 B：政策返还文件清单（一企一策）")
    add_bullet(doc, "附件 C：交付标准 + 房屋现状照片 + 楼面荷载/供电预留实测")
    add_bullet(doc, "附件 D：装修管理办法（含装配式接口要求）")
    add_bullet(doc, "附件 E：客户入驻须知")

    add_signature_block(doc, ["甲方：上海冠松 [项目运营公司]", "乙方：[链主企业全称]"])

    out = OUT_DIR / "01-合作协议-链主总部租赁合同.docx"
    doc.save(out)
    print(f"  ✓ {out.name}")
    return out


# ======================================================================
# 2. 中介居间服务协议
# ======================================================================
def build_broker_agreement():
    doc = Document()
    doc_setup(doc)
    add_title(doc, "房 屋 租 赁 居 间 服 务 协 议")
    add_subtitle(doc, "GS · iDrive Hub · 01# 研发楼 · 中介渠道通用版")
    add_subtitle(doc, "草案版本：v1.0 · 仅供业务讨论 · 签约前由法务最终定稿")

    add_p(doc, "")
    add_h2(doc, "合同编号：GS-BRK-XXXX")

    add_h1(doc, "签约方信息")
    add_kv_table(doc, [
        ("委托方（甲方）", "上海冠松 [项目运营公司全称] 有限公司"),
        ("居间方（乙方）", "[戴德梁行 / 仲量联行 / 高力国际 / 世邦魏理仕 / 本地精品行 之一]"),
    ])

    add_h1(doc, "第一条 · 服务标的")
    add_p(doc,
          "标的位于上海市静安区永和社区 075b-07 地块『01# 研发楼』"
          "（地上 9F 共 15,152.75 ㎡）的可出租办公面积。"
          "标的具体面积、租金、免租等以甲方当时有效之《报价单》为准。",
          indent=True)

    add_h1(doc, "第二条 · 合作模式")
    add_p(doc, "2.1 非独家合作：甲方有权同时委托其他中介及自行招商；", indent=True)
    add_p(doc,
          "2.2 30 天首报机制：乙方就特定客户首次以书面《客户报备表》登记，"
          "并经甲方书面确认，自确认日起 6 个月内该客户与甲方就标的"
          "签署租赁合同的，佣金归乙方；逾期未签约的，客户回归公盘；",
          indent=True)
    add_p(doc, "2.3 客户冲突：以甲方首次书面确认者为准；", indent=True)
    add_p(doc, "2.4 服务内容：客户开发、需求摸查、踏勘陪同、报价转达、谈判协助、合同签署见证、入驻后回访等。", indent=True)

    add_h1(doc, "第三条 · 佣金标准")
    add_kv_table(doc, [
        ("标准生态客户（合同期 ≥ 3 年）", "成交首月不含税净租金 × 100%"),
        ("合同期 ≥ 5 年", "成交首月不含税净租金 × 120%"),
        ("链主客户（单笔 ≥ 3,000 ㎡）", "一事一议，最高 150%"),
        ("续约（原中介）", "续约首月 × 30%（仅一次）"),
    ])
    add_p(doc, "口径说明：不含税净租金 = 合同租金 − 物业费 − 增值税。", indent=True)

    add_h1(doc, "第四条 · 佣金支付")
    add_p(doc, "4.1 客户进场入驻 + 首月租金到账后 30 个工作日内，甲方一次性支付佣金；", indent=True)
    add_p(doc, "4.2 乙方需开具 6% 现代服务业增值税专用发票；", indent=True)
    add_p(doc,
          "4.3 扣回机制：客户在合同生效 12 个月内退租 / 违约 / 提前终止，"
          "按未履行月数比例扣回佣金；",
          indent=True)

    add_h1(doc, "第五条 · 客户报备与流程")
    add_p(doc, "5.1 乙方使用《客户报备表》（附件 A）发送至 [info@guansong-ihub.com]，甲方 2 个工作日内书面回执；", indent=True)
    add_p(doc, "5.2 踏勘前 24 小时书面预约，甲方陪同；", indent=True)
    add_p(doc,
          "5.3 报价以甲方书面《报价单》为准，乙方不得擅自承诺免租、装补、政策返还等条件。",
          indent=True)

    add_h1(doc, "第六条 · 合规与保密")
    add_bullet(doc, "反贿赂 / 反不正当竞争声明")
    add_bullet(doc, "反商业秘密泄露：客户信息、租金谈判过程仅用于本项目居间")
    add_bullet(doc, "数据合规：客户决策人个人信息按《个人信息保护法》处理")
    add_bullet(doc, "乙方不得在公开渠道使用甲方未授权的项目效果图、租金水平")

    add_h1(doc, "第七条 · 考核与退出")
    add_p(doc, "7.1 季度评比：成交家数、面积、转化率、客户满意度；", indent=True)
    add_p(doc, "7.2 末位中介下季度佣金档下调 20%；连续两季末位则甲方有权终止合作；", indent=True)
    add_p(doc,
          "7.3 任何一方提前 30 日书面通知可终止协议；已报备客户在 6 个月内仍享佣金权益。",
          indent=True)

    add_h1(doc, "第八条 · 争议解决")
    add_p(doc,
          "提交上海仲裁委员会按其届时有效仲裁规则仲裁，仲裁裁决终局。",
          indent=True)

    add_h1(doc, "附件清单")
    add_bullet(doc, "附件 A：客户报备表模板")
    add_bullet(doc, "附件 B：标的报价单（每月更新）")
    add_bullet(doc, "附件 C：佣金结算单模板")

    add_signature_block(doc, ["甲方：上海冠松 [项目运营公司]", "乙方：[中介机构]"])

    out = OUT_DIR / "02-合作协议-中介居间服务协议.docx"
    doc.save(out)
    print(f"  ✓ {out.name}")
    return out


# ======================================================================
# 3. 联合实验室共建协议（3F · 链主 + 我方）
# ======================================================================
def build_joint_lab():
    doc = Document()
    doc_setup(doc)
    add_title(doc, "联合实验室共建合作协议")
    add_subtitle(doc, "GS · iDrive Hub · 3F 联合研发实验室 · 链主 + 我方")
    add_subtitle(doc, "草案版本：v1.0 · 仅供业务讨论 · 签约前由法务最终定稿")

    add_p(doc, "")
    add_h2(doc, "协议编号：GS-LAB-XXXX")

    add_h1(doc, "签约方信息")
    add_kv_table(doc, [
        ("甲方（园区方）", "上海冠松 [项目运营公司全称] 有限公司"),
        ("乙方（链主方）", "[链主企业全称，如：地平线 / 华为车 BU 等]"),
    ])

    add_h1(doc, "鉴于")
    add_p(doc, "1. 甲方运营 GS · iDrive Hub · 01# 研发楼 3F 核心研发层（净层高 5.7 m，约 1,683 ㎡）；", indent=True)
    add_p(doc, "2. 乙方系智能驾驶产业链链主企业，拥有领先的算法/硬件/数据/仿真能力；", indent=True)
    add_p(doc,
          "3. 双方拟在 3F 共建『[名称] · GS 联合研发实验室』，"
          "用于算法-硬件协同验证、Tier1 客户演示、产业生态加速。",
          indent=True)

    add_h1(doc, "第一条 · 实验室基本信息")
    add_kv_table(doc, [
        ("实验室名称", "[链主] · GS 联合研发实验室（暂定）"),
        ("位置", "01# 研发楼 3F"),
        ("面积", "约 1,683 ㎡（含办公 + 实验区 + 演示厅）"),
        ("净层高", "5.7 m"),
        ("楼面荷载", "甲方加固至 ≥ 7.5 kN/㎡（甲方承担）"),
        ("供电预留", "≥ 350 W/㎡"),
        ("成立时间", "[YYYY-MM-DD]"),
        ("合作期限", "5 年（可续 5 年）"),
    ])

    add_h1(doc, "第二条 · 出资与权益")
    add_p(doc, "2.1 出资：双方按 50%:50% 比例共同出资建设实验室硬件设施；", indent=True)
    add_p(doc, "2.2 权益分配：", indent=True)
    add_bullet(doc, "甲方：提供物理空间 + 楼面加固 + 机电预留")
    add_bullet(doc, "乙方：提供算法/硬件设备 + 研发人员 + 品牌冠名")
    add_p(doc, "2.3 命名权：乙方享冠名「[链主] · GS 联合研发实验室」5 年；", indent=True)
    add_p(doc, "2.4 优先调度权：乙方优先使用 60% 时段，剩余 40% 由甲方安排其他客户使用。", indent=True)

    add_h1(doc, "第三条 · 治理结构")
    add_p(doc, "3.1 实验室管理委员会：", indent=True)
    add_bullet(doc, "甲方代表 2 人（项目总监 + 运营总监）")
    add_bullet(doc, "乙方代表 2 人（[岗位] + [岗位]）")
    add_bullet(doc, "管委会主任：双方协商确定，每年轮值")
    add_p(doc, "3.2 决策机制：", indent=True)
    add_bullet(doc, "实验室年度预算 / 重大事项：管委会一致同意")
    add_bullet(doc, "日常运营 / 排期：管委会主任决策")

    add_h1(doc, "第四条 · 财务安排")
    add_p(doc, "4.1 一次性建设投入：人民币 [金额] 万元（双方各 50%）；", indent=True)
    add_p(doc, "4.2 年度运营预算：人民币 [金额] 万元/年（双方各 50%）；", indent=True)
    add_p(doc,
          "4.3 实验室对外承接服务（如对生态企业、Tier1 客户）的收入，"
          "扣除直接成本后按双方出资比例分配。",
          indent=True)

    add_h1(doc, "第五条 · 知识产权")
    add_p(doc, "5.1 各自原有知识产权归各自所有，对方有合理使用权；", indent=True)
    add_p(doc, "5.2 联合实验室期间产生的新知识产权：", indent=True)
    add_bullet(doc, "由乙方主导研发的：归乙方所有，甲方有学术发表署名权")
    add_bullet(doc, "由甲方主导提供的运营/管理工具：归甲方所有")
    add_bullet(doc, "联合产生的：双方共有，商用需双方协商")

    add_h1(doc, "第六条 · 数据与合规")
    add_p(doc, "6.1 数据使用：实验室产生的测试数据按数据合规沙盒规则脱敏处理；", indent=True)
    add_p(doc, "6.2 出口管制：涉及出口管制的算法/硬件，乙方承担合规义务；", indent=True)
    add_p(doc, "6.3 数据出境：与上海数交所合规沙盒对接，按 PIPL 处理。", indent=True)

    add_h1(doc, "第七条 · 退出机制")
    add_p(doc, "7.1 任何一方提前 12 个月书面通知可退出；", indent=True)
    add_p(doc, "7.2 退出时硬件资产按账面净值由继续方收购或拆除；", indent=True)
    add_p(doc, "7.3 命名权同步终止；", indent=True)
    add_p(doc, "7.4 联合知识产权按既定比例分割。", indent=True)

    add_h1(doc, "第八条 · 争议解决")
    add_p(doc,
          "提交上海仲裁委员会按其届时有效仲裁规则仲裁，仲裁裁决终局。",
          indent=True)

    add_signature_block(doc, ["甲方：上海冠松 [项目运营公司]", "乙方：[链主企业]"])

    out = OUT_DIR / "03-合作协议-联合实验室共建协议.docx"
    doc.save(out)
    print(f"  ✓ {out.name}")
    return out


# ======================================================================
# 4. 政府专班合作备忘录（区政府 + 我方）
# ======================================================================
def build_gov_mou():
    doc = Document()
    doc_setup(doc)
    add_title(doc, "战略合作备忘录")
    add_subtitle(doc, "上海市静安区人民政府 × 冠松集团")
    add_subtitle(doc, "GS · iDrive Hub · 智能驾驶研发与总部楼 · 一企一策政府专班")
    add_subtitle(doc, "草案版本：v1.0 · 仅供业务讨论 · 签约前由法务最终定稿")

    add_p(doc, "")
    add_h2(doc, "备忘录编号：GS-MOU-XXXX")

    add_h1(doc, "签约方信息")
    add_kv_table(doc, [
        ("甲方（区政府）", "上海市静安区人民政府"),
        ("代表", "区委书记 / 区长 [姓名]"),
        ("乙方（冠松集团）", "冠松集团 / 上海冠松 [项目运营公司]"),
        ("代表", "集团董事长 / 项目总监 [姓名]"),
    ])

    add_h1(doc, "前言")
    add_p(doc,
          "为加快上海市静安区智能网联汽车产业发展，落实"
          "《上海市智能网联汽车高质量发展行动方案》，依托冠松集团于永和社区"
          "075b-07 地块投资建设的『01# 研发楼』，打造『中国中心城区首个智能驾驶"
          "研发与总部楼』 GS · iDrive Hub，特签订本合作备忘录。",
          indent=True)

    add_h1(doc, "第一条 · 项目定位与目标")
    add_p(doc, "1.1 项目定位：『中国中心城区首个智能驾驶研发与总部楼』；", indent=True)
    add_p(doc, "1.2 客群目标：链主总部 1–2 家 + 核心研发 4–6 家 + 算法软件 6–10 家；", indent=True)
    add_p(doc, "1.3 三年目标：入驻率 92%，链主 ≥ 1 家，入驻企业 12–18 家。", indent=True)

    add_h1(doc, "第二条 · 政府支持承诺")
    add_p(doc, "2.1 成立『GS · iDrive Hub 区级专项工作组』，由区委书记/区长任组长；", indent=True)
    add_p(doc, "2.2 区投促办派驻人员，建立『一窗通办』机制；", indent=True)
    add_p(doc, "2.3 区交警支队协助申请『1.5 km 静安区路测延伸路段』；", indent=True)
    add_p(doc, "2.4 区财政承诺：链主级别税收留成 80% 三年返、50% 后两年返；", indent=True)
    add_p(doc, "2.5 区房管局承诺：人才公寓配额 ≥ 200 套（链主总部）+ ≥ 300 套（生态企业）；", indent=True)
    add_p(doc, "2.6 区人社局承诺：高级研发落户绿通 ≥ 80 个/年；", indent=True)
    add_p(doc, "2.7 区级『iDrive · 静安 10 条』政策包于 2026 年 9 月发布会同步发布。", indent=True)

    add_h1(doc, "第三条 · 冠松承诺")
    add_p(doc, "3.1 启动预算：Y0–Y1 投入不少于 [金额] 亿元；", indent=True)
    add_p(doc, "3.2 团队配置：12 个月内核心团队达 22 人；", indent=True)
    add_p(doc, "3.3 链主签约：T+90d 至少 1 家链主 Term Sheet；", indent=True)
    add_p(doc, "3.4 9 月旗舰发布会：200 人现场 + 5 家以上签约项目；", indent=True)
    add_p(doc, "3.5 三年贡献：累计税收 ≥ [金额] 亿元，带动产业链就业 ≥ [人数] 人。", indent=True)

    add_h1(doc, "第四条 · 协同工作机制")
    add_p(doc, "4.1 联席会议：", indent=True)
    add_bullet(doc, "区领导 / 集团董事长 季度联席会")
    add_bullet(doc, "区投促办主任 / 项目总监 月度协调会")
    add_bullet(doc, "专项工作组成员 周度对接")
    add_p(doc, "4.2 重大事项升级：", indent=True)
    add_bullet(doc, "重大政策诉求 → 区四套班子专题会")
    add_bullet(doc, "重大客户接待 → 区领导/区委书记")
    add_bullet(doc, "重大风险事件 → 联席会议 24 小时内召开")

    add_h1(doc, "第五条 · 信息保护")
    add_p(doc,
          "5.1 双方共享的链主企业信息、商务谈判信息、政策细节等，"
          "按『内部知悉范围最小化』原则管理；",
          indent=True)
    add_p(doc,
          "5.2 任何一方对外发声需事先通知另一方，"
          "以政府新闻通稿/区委宣传部审核稿为口径标准。",
          indent=True)

    add_h1(doc, "第六条 · 期限与变更")
    add_p(doc, "6.1 本备忘录自双方签署之日起生效，期限 5 年；", indent=True)
    add_p(doc, "6.2 任何一方提前 6 个月书面通知可调整本备忘录；", indent=True)
    add_p(doc, "6.3 本备忘录是双方合作的指导性文件，具体事项以另行签订的执行性协议为准。", indent=True)

    add_h1(doc, "第七条 · 争议解决")
    add_p(doc,
          "因本备忘录发生的任何争议，应首先通过联席会议协商解决。"
          "确实无法协商一致的，提交上海仲裁委员会按其届时有效仲裁规则仲裁。",
          indent=True)

    add_signature_block(doc, ["甲方：上海市静安区人民政府", "乙方：冠松集团"])

    out = OUT_DIR / "04-合作协议-政府专班合作备忘录.docx"
    doc.save(out)
    print(f"  ✓ {out.name}")
    return out


if __name__ == "__main__":
    print("生成合作协议 Word 文档：")
    build_anchor_lease()
    build_broker_agreement()
    build_joint_lab()
    build_gov_mou()
    print("✓ 全部完成")
