# -*- coding: utf-8 -*-
"""生成可直接打印的一页纸 Word（发给对方预览）。"""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import content as C

OUT = Path(__file__).resolve().parent.parent / "output" / "港大经管上海中心_联合策划_一页纸.docx"

CN = "微软雅黑"
GREEN = RGBColor(0x00, 0x3D, 0x2E)
GOLD = RGBColor(0xC4, 0xA3, 0x5A)
DARK = RGBColor(0x1A, 0x24, 0x20)
GREY = RGBColor(0x5B, 0x6B, 0x64)


def font(run, size, bold=False, color=DARK, name=CN):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def add(doc, text, size=10.5, bold=False, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT, after=4, before=0, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.15
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    font(r, size, bold, color)
    return p


def build(path: Path | None = None) -> Path:
    path = path or OUT
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.4)
    sec.bottom_margin = Cm(1.2)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

    add(doc, C.THEIR_UNIT + "  ×  联合策划方", 10, False, GOLD, after=2)
    add(doc, C.PROJECT_NAME, 22, True, GREEN, after=2)
    add(doc, C.PROJECT_SUBTITLE + f"    {C.VERSION}", 12, False, GREY, after=8)
    add(doc, C.ONE_LINER, 11, False, DARK, after=8)

    add(doc, "思路", 13, True, GREEN, before=2, after=4)
    add(doc, "不做宣讲会外包。把招生嵌进产业闭门课：来的人谈产业，名单共管，会后由招生官一对一。录取权在港大。", 10.5, after=4)
    for p in C.PRODUCTS:
        add(doc, f"· {p['name']}：{p['desc']}", 10.5, after=2)

    add(doc, "合作怎么成立", 13, True, GREEN, before=8, after=4)
    add(
        doc,
        f"{C.FEE_NAME}  {C.FEE_AMOUNT_CN}（¥{C.FEE_AMOUNT:,}）。协议生效后 {C.FEE_DAYS} 个工作日内一次性支付。"
        f"覆盖 {C.PLAN_DAYS} 天策划、不少于 {C.NAME_LIST_MIN} 人定向名单、首场闭门课（{C.FIRST_EVENT_SIZE}）及会后纪要。",
        10.5,
        after=4,
    )
    add(doc, "不碰学费分成，不承诺录取人数，第二场另签。", 10.5, True, after=4)

    add(doc, "90 天", 13, True, GREEN, before=6, after=4)
    for a, b, c in C.NINETY_DAY:
        add(doc, f"{a}  {b}  {c}", 10, after=2)

    add(doc, "下一步", 13, True, GREEN, before=6, after=4)
    for i, x in enumerate(C.NEXT_STEPS, 1):
        add(doc, f"{i}. {x}", 10.5, after=2)

    add(
        doc,
        f"致 {C.THEIR_CONTACT}  {C.THEIR_TITLE}  {C.THEIR_TEL}  {C.THEIR_EMAIL}  {C.THEIR_ADDR}",
        9,
        False,
        GREY,
        before=10,
        after=2,
    )
    add(doc, f"联合策划方：{C.OUR_PARTIES}", 9, False, GREY, after=0)

    doc.save(path)
    return path


if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else OUT
    p = build(out)
    print(f"已生成 {p}")
