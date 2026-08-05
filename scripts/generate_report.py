# -*- coding: utf-8 -*-
"""生成《赛普客户：房企“十五五”竞争的破局关键——客户服务体系化建设》Word 报告。

用法：
    python3 scripts/generate_report.py [--tocmap toc.json]

两遍目录页码回填：
    1) 不带 tocmap 生成初版 docx → 转 PDF → 提取各章页码写入 toc.json
    2) 带 --tocmap 重新生成，目录带页码
"""
import argparse
import json
import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import report_content as rc

ASSETS = "/workspace/assets"
OUT_DIR = "/workspace/report"
DOCX_NAME = "赛普客户-房企十五五竞争的破局关键-客户服务体系化建设.docx"

NAVY = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xE9, 0x4D, 0x4F)
BLUE = RGBColor(0x3D, 0x5A, 0x80)
GRAY = RGBColor(0x8D, 0x99, 0xAE)
DARKTXT = RGBColor(0x2D, 0x37, 0x48)
MIDTXT = RGBColor(0x4A, 0x55, 0x68)

SERIF_EA = "Noto Serif CJK SC"
SERIF_LA = "Noto Serif"
SANS_EA = "Noto Sans CJK SC"
SANS_LA = "Noto Sans"

PAGE_W = Cm(21.0)
PAGE_H = Cm(29.7)
MARGIN_TB = Cm(2.5)
MARGIN_LR = Cm(2.6)
CONTENT_W = Cm(21.0 - 2.6 * 2)  # 15.8cm


# ---------------------------------------------------------------- 基础工具
def set_run(run, east=SERIF_EA, latin=SERIF_LA, size=10.5, bold=False,
            color=None, italic=False, spacing=None):
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if spacing is not None:  # 字距，单位磅
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(int(spacing * 20)))
        rPr.append(sp)


def para(doc, align=None, before=0, after=6, line=1.5, indent=None,
         left_indent=None, keep_next=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
    if indent is not None:
        pf.first_line_indent = indent
    if left_indent is not None:
        pf.left_indent = left_indent
    pf.keep_with_next = keep_next
    return p


def body_p(doc, text, **kw):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             indent=Cm(0.74), after=kw.pop("after", 8))
    r = p.add_run(text)
    set_run(r, **kw)
    return p


def p_shade(p, color_hex):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def p_border(p, edge="bottom", color="E94D4F", sz=12, space=4):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    pBdr.append(el)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_table_borders(table, color="B7C1CE", sz=4, inside=True):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    edges = ["top", "left", "bottom", "right"] + (["insideH", "insideV"] if inside else [])
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tblPr.append(borders)


def cell_text(cell, text, size=9, bold=False, color=None, east=SANS_EA,
              latin=SANS_LA, align=WD_ALIGN_PARAGRAPH.LEFT, after=2, before=2):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run(r, east=east, latin=latin, size=size, bold=bold, color=color)


def add_page_field(p, size=9, color=MIDTXT):
    r1 = p.add_run()
    set_run(r1, east=SANS_EA, latin=SANS_LA, size=size, color=color)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    r1._element.append(fld1)
    r2 = p.add_run()
    set_run(r2, east=SANS_EA, latin=SANS_LA, size=size, color=color)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    r2._element.append(instr)
    r3 = p.add_run()
    set_run(r3, east=SANS_EA, latin=SANS_LA, size=size, color=color)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r3._element.append(fld2)


def keep_with_next(p):
    p.paragraph_format.keep_with_next = True


# ---------------------------------------------------------------- 页面元素
def h1(doc, text, page_break=True):
    p = para(doc, before=0, after=18, line=1.3, keep_next=True)
    if page_break:
        p.paragraph_format.page_break_before = True
    # 红色小方块 + 标题
    r0 = p.add_run("■ ")
    set_run(r0, east=SANS_EA, latin=SANS_LA, size=20, bold=True, color=RED)
    r = p.add_run(text)
    set_run(r, east=SANS_EA, latin=SANS_LA, size=20, bold=True, color=NAVY)
    p_border(p, edge="bottom", color="1F2A44", sz=10, space=8)
    return p


def h2(doc, text):
    p = para(doc, before=14, after=8, line=1.3, keep_next=True)
    r0 = p.add_run("▍")
    set_run(r0, east=SANS_EA, latin=SANS_LA, size=14, bold=True, color=RED)
    r = p.add_run(text)
    set_run(r, east=SANS_EA, latin=SANS_LA, size=14, bold=True, color=NAVY)
    return p


def h3(doc, text):
    p = para(doc, before=10, after=6, line=1.3, keep_next=True)
    r = p.add_run(text)
    set_run(r, east=SANS_EA, latin=SANS_LA, size=12, bold=True, color=BLUE)
    return p


def lead_p(doc, text):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=2, after=12,
             line=1.5, left_indent=Cm(0.35))
    p_shade(p, "F4F6F8")
    p_border(p, edge="left", color="E94D4F", sz=20, space=4)
    r = p.add_run(text)
    set_run(r, east=SERIF_EA, latin=SERIF_LA, size=11, bold=True, color=NAVY)
    return p


def quote_p(doc, text):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=6, after=12,
             line=1.6, left_indent=Cm(0.8))
    p_border(p, edge="left", color="E94D4F", sz=18, space=6)
    r0 = p.add_run("“ ")
    set_run(r0, east=SERIF_EA, latin=SERIF_LA, size=12, bold=True, color=RED)
    r = p.add_run(text)
    set_run(r, east=SERIF_EA, latin=SERIF_LA, size=11, color=NAVY, italic=False)
    r2 = p.add_run(" ”")
    set_run(r2, east=SERIF_EA, latin=SERIF_LA, size=12, bold=True, color=RED)
    return p


def bullets(doc, items):
    for it in items:
        p = para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=5,
                 left_indent=Cm(0.74))
        r0 = p.add_run("▪ ")
        set_run(r0, east=SANS_EA, latin=SANS_LA, size=10.5, color=RED, bold=True)
        r = p.add_run(it)
        set_run(r, east=SERIF_EA, latin=SERIF_LA, size=10.5, color=DARKTXT)


def img_block(doc, fname, caption):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=2, line=1.0,
             keep_next=True)
    run = p.add_run()
    run.add_picture(os.path.join(ASSETS, fname), width=CONTENT_W)
    cap = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=12, line=1.2)
    r = cap.add_run(caption)
    set_run(r, east=SANS_EA, latin=SANS_LA, size=9, color=MIDTXT)


def table_block(doc, caption, header, rows, widths_cm):
    cap = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=4,
               line=1.2, keep_next=True)
    r = cap.add_run(caption)
    set_run(r, east=SANS_EA, latin=SANS_LA, size=9.5, bold=True, color=NAVY)

    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="B7C1CE", sz=4)
    table.autofit = False
    for j, w in enumerate(widths_cm):
        for row in table.rows:
            row.cells[j].width = Cm(w)
    # 表头
    for j, htext in enumerate(header):
        c = table.rows[0].cells[j]
        shade_cell(c, "1F2A44")
        cell_text(c, htext, size=9.5, bold=True, color=RGBColor(255, 255, 255),
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = table.rows[i + 1].cells[j]
            if i % 2 == 1:
                shade_cell(c, "F4F6F8")
            cell_text(c, val, size=9, color=DARKTXT,
                      align=WD_ALIGN_PARAGRAPH.LEFT if j == len(header) - 1 or len(header) > 2 else WD_ALIGN_PARAGRAPH.LEFT)
    # 表后留白
    para(doc, after=8, line=1.0)
    return table


def callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(table)
    c = table.rows[0].cells[0]
    c.width = CONTENT_W
    shade_cell(c, "F4F6F8")
    tcPr = c._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    el = OxmlElement("w:left")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), "24")
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), "E94D4F")
    borders.append(el)
    tcPr.append(borders)
    p1 = c.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after = Pt(2)
    r = p1.add_run(title)
    set_run(r, east=SANS_EA, latin=SANS_LA, size=10.5, bold=True, color=RED)
    p2 = c.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing = 1.5
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r2 = p2.add_run(text)
    set_run(r2, east=SERIF_EA, latin=SERIF_LA, size=10.5, color=DARKTXT)
    para(doc, after=8, line=1.0)


# ---------------------------------------------------------------- 封面
def cover(doc):
    # 顶部 logo
    p = para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=1.0)
    p.add_run().add_picture(os.path.join(ASSETS, "sap-logo.png"), height=Cm(1.15))
    # 留白
    para(doc, after=0, line=1.0).paragraph_format.space_before = Pt(64)
    # 系列行
    p = para(doc, after=6, line=1.0)
    r = p.add_run(rc.REPORT_SERIES)
    set_run(r, east=SANS_EA, latin=SANS_LA, size=12, bold=True, color=RED, spacing=2)
    p_border(p, edge="bottom", color="E94D4F", sz=8, space=6)
    # 主标题
    p = para(doc, before=18, after=4, line=1.25)
    r = p.add_run(rc.REPORT_TITLE_MAIN)
    set_run(r, east=SANS_EA, latin=SANS_LA, size=32, bold=True, color=NAVY)
    p = para(doc, before=0, after=16, line=1.25)
    r = p.add_run(rc.REPORT_TITLE_SUB)
    set_run(r, east=SANS_EA, latin=SANS_LA, size=22, bold=True, color=RED)
    # 英文题
    p = para(doc, after=24, line=1.2)
    r = p.add_run("THE BREAKTHROUGH KEY OF REAL ESTATE COMPETITION IN THE 15TH FIVE-YEAR PLAN\n— SYSTEMATIC CONSTRUCTION OF CUSTOMER SERVICE")
    set_run(r, east=SANS_EA, latin=SANS_LA, size=8.5, color=GRAY, spacing=1)
    # 摘要引语
    p = para(doc, after=30, line=1.6, left_indent=Cm(0.2))
    r = p.add_run("从“经验驱动”到“体系驱动”——以“14554”客户服务体系蓝图，构筑房企穿越周期的服务护城河")
    set_run(r, east=SERIF_EA, latin=SERIF_LA, size=11.5, color=MIDTXT)
    # 信息表
    info = [("出品机构", rc.REPORT_ORG), ("研究团队", rc.REPORT_AUTHORS),
            ("发布时间", rc.REPORT_DATE), ("报告编号", rc.REPORT_NO)]
    table = doc.add_table(rows=4, cols=2)
    no_borders(table)
    for i, (k, v) in enumerate(info):
        c0, c1 = table.rows[i].cells
        c0.width = Cm(3.2)
        c1.width = Cm(12.6)
        cell_text(c0, k, size=10, bold=True, color=GRAY, before=3, after=3)
        cell_text(c1, v, size=10.5, bold=True, color=NAVY, before=3, after=3)
    para(doc, after=0, line=1.0).paragraph_format.space_before = Pt(44)
    # 底部深藏青色带
    table = doc.add_table(rows=1, cols=1)
    no_borders(table)
    c = table.rows[0].cells[0]
    c.width = CONTENT_W
    shade_cell(c, "1F2A44")
    p1 = c.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(14)
    p1.paragraph_format.space_after = Pt(4)
    p1.add_run().add_picture(os.path.join(ASSETS, "sap-logo-white.png"), height=Cm(0.72))
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    r = p2.add_run("赛普咨询全国统一热线 400-9669-209    ｜    www.chinasap.cn")
    set_run(r, east=SANS_EA, latin=SANS_LA, size=9, color=RGBColor(0xC9, 0xD2, 0xE0))


# ---------------------------------------------------------------- 版权页
def copyright_page(doc):
    para(doc, after=0, line=1.0).paragraph_format.page_break_before = True
    para(doc, after=0, line=1.0).paragraph_format.space_before = Pt(84)
    p = para(doc, after=10, line=1.4)
    r = p.add_run(rc.COPYRIGHT["title"])
    set_run(r, east=SANS_EA, latin=SANS_LA, size=14, bold=True, color=NAVY)
    for line in [rc.COPYRIGHT["org"], rc.COPYRIGHT["authors"],
                 rc.COPYRIGHT["date"], rc.COPYRIGHT["no"]]:
        p = para(doc, after=4, line=1.4)
        r = p.add_run(line)
        set_run(r, east=SANS_EA, latin=SANS_LA, size=10.5, color=DARKTXT)
    para(doc, after=0, line=1.0).paragraph_format.space_before = Pt(26)
    p = para(doc, after=6, line=1.4)
    r = p.add_run("版权与免责声明")
    set_run(r, east=SANS_EA, latin=SANS_LA, size=10.5, bold=True, color=NAVY)
    for line in rc.COPYRIGHT["disclaimer"]:
        p = para(doc, after=5, line=1.45, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        r = p.add_run(line)
        set_run(r, east=SERIF_EA, latin=SERIF_LA, size=9, color=MIDTXT)
    para(doc, after=0, line=1.0).paragraph_format.space_before = Pt(18)
    for line in rc.COPYRIGHT["contact"]:
        p = para(doc, after=3, line=1.3)
        r = p.add_run(line)
        set_run(r, east=SANS_EA, latin=SANS_LA, size=9, color=MIDTXT)


# ---------------------------------------------------------------- 摘要页
def abstract_page(doc):
    p = para(doc, before=0, after=16, line=1.3, keep_next=True)
    p.paragraph_format.page_break_before = True
    r0 = p.add_run("■ ")
    set_run(r0, east=SANS_EA, latin=SANS_LA, size=20, bold=True, color=RED)
    r = p.add_run("摘要")
    set_run(r, east=SANS_EA, latin=SANS_LA, size=20, bold=True, color=NAVY)
    p_border(p, edge="bottom", color="1F2A44", sz=10, space=8)
    for text in rc.ABSTRACT:
        body_p(doc, text, east=SERIF_EA, latin=SERIF_LA, size=10.5, color=DARKTXT)
    # 核心数字条
    para(doc, after=4, line=1.0).paragraph_format.space_before = Pt(6)
    nums = [("30%+", "老业主复购与推荐率"), ("5%—10%", "住宅品牌溢价"),
            ("8%—15%", "二手房价格优势"), ("104", "个标准化服务流程")]
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(table)
    for j, (num, lab) in enumerate(nums):
        c = table.rows[0].cells[j]
        shade_cell(c, "1F2A44")
        cell_text(c, num, size=15, bold=True, color=RGBColor(0xE9, 0x4D, 0x4F),
                  align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=2)
        c2 = table.rows[1].cells[j]
        shade_cell(c2, "1F2A44")
        cell_text(c2, lab, size=8.5, color=RGBColor(0xFF, 0xFF, 0xFF),
                  align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=8)


# ---------------------------------------------------------------- 目录页
def toc_page(doc, tocmap):
    p = para(doc, before=0, after=16, line=1.3)
    p.paragraph_format.page_break_before = True
    r0 = p.add_run("■ ")
    set_run(r0, east=SANS_EA, latin=SANS_LA, size=20, bold=True, color=RED)
    r = p.add_run("目录")
    set_run(r, east=SANS_EA, latin=SANS_LA, size=20, bold=True, color=NAVY)
    p_border(p, edge="bottom", color="1F2A44", sz=10, space=8)

    entries = []
    for blk in rc.BLOCKS:
        if blk[0] == "h1":
            entries.append((1, blk[1]))
        elif blk[0] == "h2":
            entries.append((2, blk[1]))

    for level, title in entries:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(3 if level == 1 else 1)
        pf.space_before = Pt(7 if level == 1 else 0)
        pf.line_spacing = 1.1
        pf.left_indent = Cm(0 if level == 1 else 0.9)
        pf.tab_stops.add_tab_stop(Cm(15.4 - (0 if level == 1 else 0.9)),
                                  WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(title)
        if level == 1:
            set_run(r, east=SANS_EA, latin=SANS_LA, size=11, bold=True, color=NAVY)
        else:
            set_run(r, east=SERIF_EA, latin=SERIF_LA, size=9.5, color=MIDTXT)
        page_no = (tocmap or {}).get(_norm(title), "")
        r2 = p.add_run("\t" + str(page_no))
        set_run(r2, east=SANS_EA, latin=SANS_LA, size=10,
                bold=(level == 1), color=NAVY if level == 1 else MIDTXT)


def _norm(s):
    return "".join(s.split())


# ---------------------------------------------------------------- 关于赛普（末页）
def about_page(doc):
    h1(doc, "关于赛普", page_break=True)
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=14, line=1.0)
    p.add_run().add_picture(os.path.join(ASSETS, "sap-logo.png"), height=Cm(1.3))
    for text in rc.ABOUT_SAP["intro"]:
        body_p(doc, text, east=SERIF_EA, latin=SERIF_LA, size=10.5, color=DARKTXT)
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=16, after=8, line=1.4)
    r = p.add_run(rc.ABOUT_SAP["values"])
    set_run(r, east=SANS_EA, latin=SANS_LA, size=12, bold=True, color=RED)
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=20, line=1.4)
    r = p.add_run(rc.ABOUT_SAP["mission"])
    set_run(r, east=SERIF_EA, latin=SERIF_LA, size=10.5, color=MIDTXT)
    # 联系方式色带
    table = doc.add_table(rows=1, cols=1)
    no_borders(table)
    c = table.rows[0].cells[0]
    c.width = CONTENT_W
    shade_cell(c, "1F2A44")
    p1 = c.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(12)
    p1.paragraph_format.space_after = Pt(4)
    r = p1.add_run("互动与交流")
    set_run(r, east=SANS_EA, latin=SANS_LA, size=11, bold=True, color=RGBColor(255, 255, 255))
    for line in rc.COPYRIGHT["contact"]:
        p = c.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(line)
        set_run(r, east=SANS_EA, latin=SANS_LA, size=9,
                color=RGBColor(0xC9, 0xD2, 0xE0))
    pl = c.add_paragraph()
    pl.paragraph_format.space_after = Pt(10)


# ---------------------------------------------------------------- 页眉页脚
def header_footer(doc):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    # 页眉：logo + 短标题 + 红线
    hp = sec.header.paragraphs[0]
    hp.paragraph_format.tab_stops.add_tab_stop(CONTENT_W, WD_TAB_ALIGNMENT.RIGHT)
    hp.add_run().add_picture(os.path.join(ASSETS, "sap-logo.png"), height=Cm(0.55))
    r = hp.add_run("\t" + "房企“十五五”竞争的破局关键——客户服务体系化建设")
    set_run(r, east=SANS_EA, latin=SANS_LA, size=8, color=GRAY)
    p_border(hp, edge="bottom", color="E94D4F", sz=6, space=2)
    # 页脚：机密声明 + 页码
    fp = sec.footer.paragraphs[0]
    fp.paragraph_format.tab_stops.add_tab_stop(CONTENT_W, WD_TAB_ALIGNMENT.RIGHT)
    r = fp.add_run(f"赛普研究院 · {rc.REPORT_NO} · 仅供客户内部参考")
    set_run(r, east=SANS_EA, latin=SANS_LA, size=8, color=GRAY)
    r = fp.add_run("\t第 ")
    set_run(r, east=SANS_EA, latin=SANS_LA, size=9, color=MIDTXT)
    add_page_field(fp)
    r = fp.add_run(" 页")
    set_run(r, east=SANS_EA, latin=SANS_LA, size=9, color=MIDTXT)
    p_border(fp, edge="top", color="B7C1CE", sz=4, space=2)


# ---------------------------------------------------------------- 主流程
def build(tocmap=None):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = PAGE_W
    sec.page_height = PAGE_H
    sec.top_margin = MARGIN_TB
    sec.bottom_margin = MARGIN_TB
    sec.left_margin = MARGIN_LR
    sec.right_margin = MARGIN_LR

    # 默认样式
    normal = doc.styles["Normal"]
    normal.font.name = SERIF_LA
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), SERIF_EA)

    header_footer(doc)
    cover(doc)
    copyright_page(doc)
    abstract_page(doc)
    toc_page(doc, tocmap)

    for blk in rc.BLOCKS:
        kind = blk[0]
        if kind == "h1":
            h1(doc, blk[1])
        elif kind == "h2":
            h2(doc, blk[1])
        elif kind == "h3":
            h3(doc, blk[1])
        elif kind == "p":
            body_p(doc, blk[1], east=SERIF_EA, latin=SERIF_LA, size=10.5, color=DARKTXT)
        elif kind == "lead":
            lead_p(doc, blk[1])
        elif kind == "quote":
            quote_p(doc, blk[1])
        elif kind == "bullets":
            bullets(doc, blk[1])
        elif kind == "img":
            img_block(doc, blk[1], blk[2])
        elif kind == "table":
            table_block(doc, blk[1], blk[2], blk[3], blk[4])
        elif kind == "callout":
            callout(doc, blk[1], blk[2])
        else:
            raise ValueError(f"unknown block: {kind}")

    about_page(doc)

    os.makedirs(OUT_DIR, exist_ok=True)
    out = os.path.join(OUT_DIR, DOCX_NAME)
    doc.save(out)
    print("saved", out)


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--tocmap", default=None)
    args = ap.parse_args()
    tocmap = None
    if args.tocmap and os.path.exists(args.tocmap):
        tocmap = json.load(open(args.tocmap, encoding="utf-8"))
    build(tocmap)
