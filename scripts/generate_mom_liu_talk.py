#!/usr/bin/env python3
"""生成给妈妈跟刘孝春沟通的口径。全部家属自留。Word/PDF 接话卡不含 70%/30%。"""

from pathlib import Path
import html

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.colors import HexColor, white
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from comm_plan import CN_FONT_PATH
from content import MOM_LIU_FAMILY, MOM_LIU_TALK

CHINESE_FONT = "宋体"
HEADING_FONT = "黑体"
ACCENT = RGBColor(0x0B, 0x2F, 0x5B)
MUTED = RGBColor(0x5C, 0x6B, 0x7A)
RED = RGBColor(0xA6, 0x3D, 0x2F)

MD_NAME = "给妈妈跟刘孝春沟通的口径_2026-08-20.md"
DOCX_NAME = "青岛抚顺路和哈尔滨路路口交通事故_给妈妈跟刘孝春的接话卡_20260820.docx"
PDF_NAME = "青岛抚顺路和哈尔滨路路口交通事故_给妈妈跟刘孝春的接话卡_20260820.pdf"


def _set_run_font(run, font_name: str, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _add_paragraph(
    doc: Document,
    text: str = "",
    *,
    font: str = CHINESE_FONT,
    size: float = 12,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: float | None = None,
    space_after: float = 6,
    space_before: float = 0,
    line_spacing: float = 1.35,
    color: RGBColor | None = None,
):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if text:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold, color=color)
    return p


def render_markdown() -> str:
    n = MOM_LIU_TALK
    lines = [
        f"# {n['title']}",
        "",
        n["subtitle"],
        "",
        f"> {n['file_note']}",
        "",
        "## 他现在在想什么",
        "",
    ]
    for p in n["think"]:
        lines.append(p)
        lines.append("")
    lines.extend(["## 他会不会觉得自己有责任", ""])
    for p in n["responsibility"]:
        lines.append(p)
        lines.append("")
    lines.extend(["## 他这边钱的诉求是什么", ""])
    for item in n["his_money"]:
        lines.append(f"- {item}")
    lines.extend(["", "## 妈妈这趟见面要办成什么", ""])
    for item in n["mom_goal"]:
        lines.append(f"- {item}")
    lines.extend(["", "## 开场就这几句（念完停下）", "", n["opening"], "", "## 他若这样说，就这样接", ""])
    for him, mom in n["replies"]:
        lines.append(f"**他说：**{him}")
        lines.append("")
        lines.append(f"**妈妈说：**{mom}")
        lines.append("")
    lines.extend(["## 当面不要说", ""])
    for item in n["never"]:
        lines.append(f"- {item}")
    lines.extend(
        [
            "",
            "---",
            "",
            "## 家属自留，不要念给刘听",
            "",
            "> 下面这几句只给胡继刚自己看。打印给妈妈的接话卡 Word / PDF 不含本节。",
            "",
        ]
    )
    for p in MOM_LIU_FAMILY:
        lines.append(f"- {p}")
    lines.extend(["", "重新生成：`python3 scripts/build_all.py`。本文不是律师函。", ""])
    return "\n".join(lines)


def build_markdown(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(render_markdown(), encoding="utf-8")


def _write_card_docx(doc: Document) -> None:
    n = MOM_LIU_TALK
    _add_paragraph(
        doc,
        "给妈妈跟刘孝春的接话卡",
        font=HEADING_FONT,
        size=18,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
        line_spacing=1.2,
        color=ACCENT,
    )
    _add_paragraph(
        doc,
        n["subtitle"],
        font=HEADING_FONT,
        size=11,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
        line_spacing=1.2,
        color=MUTED,
    )
    _add_paragraph(doc, n["file_note"], bold=True, color=RED, space_after=8, line_spacing=1.3)

    _add_paragraph(doc, "他大概怎么想（心里有数，不要讲给他听）", font=HEADING_FONT, size=13, bold=True, space_before=6, space_after=4, color=ACCENT, line_spacing=1.2)
    for p in n["think"][:4]:
        _add_paragraph(doc, p, first_line_indent=0.74, space_after=4, size=11)

    _add_paragraph(doc, "他钱上要什么（心里有数，不要讲给他听）", font=HEADING_FONT, size=13, bold=True, space_before=8, space_after=4, color=ACCENT, line_spacing=1.2)
    for item in n["his_money"]:
        _add_paragraph(doc, f"• {item}", first_line_indent=0.37, space_after=2, size=11)

    _add_paragraph(doc, "这趟只要办成三件事", font=HEADING_FONT, size=13, bold=True, space_before=8, space_after=4, color=ACCENT, line_spacing=1.2)
    for item in n["mom_goal"]:
        _add_paragraph(doc, f"• {item}", first_line_indent=0.37, space_after=2, size=11)

    _add_paragraph(doc, "开场就这几句（念完停下）", font=HEADING_FONT, size=13, bold=True, space_before=8, space_after=4, color=ACCENT, line_spacing=1.2)
    _add_paragraph(doc, n["opening"], first_line_indent=0.74, space_after=6, size=11)

    _add_paragraph(doc, "他若这样说，就这样接", font=HEADING_FONT, size=13, bold=True, space_before=8, space_after=4, color=ACCENT, line_spacing=1.2)
    for him, mom in n["replies"]:
        _add_paragraph(doc, f"他说：{him}", bold=True, space_after=1, size=11)
        _add_paragraph(doc, f"妈妈说：{mom}", first_line_indent=0.37, space_after=5, size=11)

    _add_paragraph(doc, "当面不要说", font=HEADING_FONT, size=13, bold=True, space_before=8, space_after=4, color=ACCENT, line_spacing=1.2)
    for item in n["never"]:
        _add_paragraph(doc, f"• {item}", first_line_indent=0.37, space_after=2, size=11, color=RED)

    _add_paragraph(
        doc,
        "本文家属自留。重新生成：python3 scripts/build_all.py。不是律师函。不要给刘孝春看。",
        size=10,
        color=MUTED,
        space_before=10,
    )


def build_document(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = CHINESE_FONT
    normal_style.font.size = Pt(11)
    rpr = normal_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CHINESE_FONT)

    _write_card_docx(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _pdf_font() -> str:
    kwargs = {"subfontIndex": 0} if CN_FONT_PATH.endswith(".ttc") else {}
    pdfmetrics.registerFont(TTFont("CN", CN_FONT_PATH, **kwargs))
    return "CN"


def _pdf_styles(font: str) -> dict[str, ParagraphStyle]:
    navy = HexColor("#0B2F5B")
    muted = HexColor("#5C6B7A")
    red = HexColor("#A63D2F")
    dark = HexColor("#1F2A37")
    return {
        "title": ParagraphStyle("t", fontName=font, fontSize=15, leading=22, textColor=navy, alignment=TA_CENTER, spaceAfter=4),
        "sub": ParagraphStyle("s", fontName=font, fontSize=9.5, leading=14, textColor=muted, alignment=TA_CENTER, spaceAfter=6),
        "warn": ParagraphStyle("w", fontName=font, fontSize=9.5, leading=14, textColor=red, alignment=TA_LEFT, spaceAfter=8),
        "h1": ParagraphStyle("h", fontName=font, fontSize=12, leading=17, textColor=navy, spaceBefore=7, spaceAfter=3),
        "body": ParagraphStyle("b", fontName=font, fontSize=9.5, leading=15, textColor=dark, alignment=TA_JUSTIFY, firstLineIndent=18, spaceAfter=3),
        "bullet": ParagraphStyle("u", fontName=font, fontSize=9.5, leading=14, textColor=dark, leftIndent=10, spaceAfter=2),
        "him": ParagraphStyle("i", fontName=font, fontSize=9.5, leading=14, textColor=navy, spaceBefore=3, spaceAfter=1),
        "mom": ParagraphStyle("m", fontName=font, fontSize=9.5, leading=14, textColor=dark, leftIndent=12, spaceAfter=4),
        "red_bullet": ParagraphStyle("r", fontName=font, fontSize=9.5, leading=14, textColor=red, leftIndent=10, spaceAfter=2),
        "foot": ParagraphStyle("f", fontName=font, fontSize=8, leading=12, textColor=muted, spaceBefore=8),
    }


def _p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(html.escape(str(text)).replace("\n", "<br/>"), style)


def _header_footer(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(HexColor("#0B2F5B"))
    canvas.rect(0, A4[1] - 12 * mm, A4[0], 12 * mm, fill=1, stroke=0)
    canvas.setFillColor(white)
    canvas.setFont("CN", 8)
    canvas.drawString(18 * mm, A4[1] - 8 * mm, "青岛抚顺路和哈尔滨路路口交通事故 · 给妈妈的接话卡 · 家属自留")
    canvas.setFillColor(HexColor("#C4A35A"))
    canvas.rect(0, 0, A4[0], 10 * mm, fill=1, stroke=0)
    canvas.setFillColor(HexColor("#0B2F5B"))
    canvas.setFont("CN", 8)
    canvas.drawString(18 * mm, 4 * mm, "不是律师函 · 不要给刘孝春看 · 当面不谈总赔偿")
    canvas.drawRightString(A4[0] - 18 * mm, 4 * mm, f"第 {doc.page} 页")
    canvas.restoreState()


def build_pdf(output_path: Path) -> None:
    n = MOM_LIU_TALK
    font = _pdf_font()
    styles = _pdf_styles(font)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=16 * mm,
        bottomMargin=14 * mm,
        title="给妈妈跟刘孝春的接话卡",
        author="内部整理",
    )
    story = [
        _p("给妈妈跟刘孝春的接话卡", styles["title"]),
        _p(n["subtitle"], styles["sub"]),
        _p(n["file_note"], styles["warn"]),
        _p("他大概怎么想（心里有数，不要讲给他听）", styles["h1"]),
    ]
    for para in n["think"][:4]:
        story.append(_p(para, styles["body"]))
    story.append(_p("他钱上要什么（心里有数，不要讲给他听）", styles["h1"]))
    for item in n["his_money"]:
        story.append(_p("• " + item, styles["bullet"]))
    story.append(_p("这趟只要办成三件事", styles["h1"]))
    for item in n["mom_goal"]:
        story.append(_p("• " + item, styles["bullet"]))
    story.append(_p("开场就这几句（念完停下）", styles["h1"]))
    story.append(_p(n["opening"], styles["body"]))
    story.append(_p("他若这样说，就这样接", styles["h1"]))
    for him, mom in n["replies"]:
        story.append(_p("他说：" + him, styles["him"]))
        story.append(_p("妈妈说：" + mom, styles["mom"]))
    story.append(_p("当面不要说", styles["h1"]))
    for item in n["never"]:
        story.append(_p("• " + item, styles["red_bullet"]))
    story.append(Spacer(1, 6))
    story.append(_p("本文家属自留。重新生成：python3 scripts/build_all.py。不是律师函。不要给刘孝春看。", styles["foot"]))
    doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "deliverables"
    build_markdown(out / MD_NAME)
    build_document(out / DOCX_NAME)
    build_pdf(out / PDF_NAME)
    print(f"Wrote {out / MD_NAME}")
    print(f"Wrote {out / DOCX_NAME}")
    print(f"Wrote {out / PDF_NAME}")


if __name__ == "__main__":
    main()
