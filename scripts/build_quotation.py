"""Generate a standalone, ready-to-stamp service quotation (报价单) document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


GREEN = RGBColor(0x00, 0x6B, 0x3F)
DARK_GREEN = RGBColor(0x00, 0x4A, 0x2C)
GOLD = RGBColor(0xC8, 0xA2, 0x5B)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _force_font(run, font="宋体"):
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rpr.append(rfonts)


def set_default_font(doc, name="宋体"):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def heading(doc, text, size=20, color=DARK_GREEN, center=True, font="黑体", bold=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    _force_font(run, font=font)
    return p


def para(doc, text, bold=False, size=11, color=DARK, align=None,
         first_indent=True, font="宋体", line_spacing=1.5):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = line_spacing
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    _force_font(run, font=font)
    return p


def blank(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()


def shade(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, color=DARK, size=10.5,
              align_center=False, font="宋体"):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    _force_font(run, font=font)


def kv_table(doc, rows, col_widths=(3.8, 12.5)):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    t.autofit = False
    for row in t.rows:
        row.cells[0].width = Cm(col_widths[0])
        row.cells[1].width = Cm(col_widths[1])
    for i, (k, v) in enumerate(rows):
        cell_text(t.rows[i].cells[0], k, bold=True, size=10.5)
        shade(t.rows[i].cells[0], "E8F1EC")
        cell_text(t.rows[i].cells[1], v, size=10.5)
    return t


def build():
    doc = Document()
    set_default_font(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    # Header
    heading(doc, "服务报价单", size=24, color=DARK_GREEN)
    heading(doc, "（绿城中国 | 绿城·潮鸣外滩 · 晚宴冠名战略合作伙伴 · 10 万元）",
            size=13, color=GREY, bold=False)
    para(doc, "Service Quotation — Dinner Title Strategic Partnership",
         align="center", size=10, color=GREY, first_indent=False)
    blank(doc)

    # Quotation meta
    para(doc, "报价单编号：QT-2026-GTC-潮鸣外滩-001       报价日期：2026 年 5 月 ____ 日       有效期：本报价单签发之日起 10 个自然日内有效",
         first_indent=False, size=10.5)
    blank(doc)

    # Parties
    para(doc, "报价方（甲方 / 服务提供方）", bold=True, first_indent=False, size=12, color=GREEN)
    kv_table(doc, rows=[
        ["主办单位", "北京大学经济学院上海校友会 / 复旦大学住房政策研究中心（联合主办）"],
        ["统一社会信用代码", "________________________________"],
        ["通讯地址", "上海市______区______路______号"],
        ["授权代表 / 职务", "________________ / ____________"],
        ["项目联系人 / 电话", "________________ / ________________"],
        ["电子邮箱", "________________"],
    ])
    blank(doc)

    para(doc, "受报价方（乙方 / 服务采购方）", bold=True, first_indent=False, size=12, color=GREEN)
    kv_table(doc, rows=[
        ["公司名称", "绿城中国控股有限公司 / ____________________________（项目主体）"],
        ["项目品牌", "绿城·潮鸣外滩"],
        ["统一社会信用代码", "________________________________"],
        ["注册地址", "________________________________"],
        ["授权代表 / 职务", "________________ / ____________"],
        ["项目联系人 / 电话", "Aiden（绿城方总对接） / ________________"],
        ["电子邮箱", "________________"],
    ])
    blank(doc)

    # Quotation body
    para(doc, "一、合作活动", bold=True, first_indent=False, size=12, color=GREEN)
    para(doc, "重构与突围 · 2026 人工智能商业化落地与硬核投资破局峰会")
    para(doc, "时间：2026 年 5 月 22 日           地点：上海·北外滩·一滴水           规模：500+ 高净值嘉宾")
    blank(doc)

    para(doc, "二、合作身份", bold=True, first_indent=False, size=12, color=GREEN)
    para(doc, "晚宴冠名战略合作伙伴（Dinner Title Strategic Partner）。")
    blank(doc)

    para(doc, "三、服务费用合计：人民币壹拾万元整（含税）", bold=True, first_indent=False, size=12, color=GREEN)
    para(doc, "小写：￥100,000.00 元；发票内容：会议服务费 / 赞助费；增值税普通发票或专用发票。",
         first_indent=False, size=10.5, color=GREY)
    blank(doc)

    para(doc, "四、服务项目价格拆解", bold=True, first_indent=False, size=12, color=GREEN)

    rows = [
        ["A", "晚宴冠名核心权益（小计：￥50,000）", "", "", "", ""],
        ["A1", "晚宴冠名权全程统称 + 晚宴 LED 大屏 KV 主视觉锁屏", "1 套（以场地像素比为准）", "甲方", "15,000", "15,000"],
        ["A2", "晚宴席卡 logo + 项目主张植入", "全场（90×55mm 双面）", "甲方", "4,000", "4,000"],
        ["A3", "晚宴桌卡 logo 植入（含品牌字样）", "全场（A5 或亚克力）", "活动公司", "4,000", "4,000"],
        ["A4", "晚宴菜单 logo 植入（封面 + 页脚）", "全场（210×285mm 或 285×210mm）", "活动公司", "3,000", "3,000"],
        ["A5", "晚宴期间项目宣传片轮播（含 LED 上屏与音视频运维）", "≤ 60 秒循环", "甲方", "6,000", "6,000"],
        ["A6", "晚宴主持人口播 2 段（开场 + 散场）", "每段约 30 秒", "甲方", "4,000", "4,000"],
        ["A7", "川总专场宣讲位置（含 LED + 话筒×2 + 切换器）", "15min + Q&A 5min", "甲方提供位置 / 设备", "6,000", "6,000"],
        ["A8", "晚宴入场券（主桌 3 + 销售/接待 5）", "8 张 × 1,000 元", "甲方", "1,000", "8,000"],
        ["B", "主会场植入（小计：￥30,000）", "", "", "", ""],
        ["B1", "主背景板项目 logo 位（同钻石级）+ 晚宴冠名身份标识", "1 处", "甲方", "10,000", "10,000"],
        ["B2", "议程手册整版广告 + 项目折页夹页", "整版 A4，出血 3mm", "甲方 / 印刷单位", "6,000", "6,000"],
        ["B3", "主会场宣传片轮播（入场及中场休息）", "16:9 1920×1080 ≤ 60s", "甲方", "4,000", "4,000"],
        ["B4", "500 份手拎袋夹页（项目折页 + 户型图，含印刷+装袋）", "三折页 + 户型单页 × 500 套", "印刷单位 / 活动公司", "20", "10,000"],
        ["C", "现场专属展位（小计：￥12,000）", "", "", "", ""],
        ["C1", "电梯口品牌展位（场地租赁 + 轻包装）", "1 处", "甲方场地 / 活动公司", "6,000", "6,000"],
        ["C2", "展场内品牌展位（场地租赁 + 洽谈圆桌轻包装）", "1 处", "甲方场地 / 活动公司", "6,000", "6,000"],
        ["D", "会后口播与引导（小计：￥2,000）", "", "", "", ""],
        ["D1", "论坛 / 颁奖结束后主持人鸣谢口播 + 动线引导", "1 次主口播 + 引导员 1 名", "甲方", "2,000", "2,000"],
        ["E", "宣发与回执（小计：￥6,000）", "", "", "", ""],
        ["E1", "朋友圈九宫格（≥ 3 张含项目 logo / KV）", "9 张 1080×1080 PNG", "甲方媒体组", "1,500", "1,500"],
        ["E2", "回顾视频片头 / 片尾项目鸣谢", "1 条 mp4，露出 ≤ 3 秒", "甲方", "2,500", "2,500"],
        ["E3", "执行回执（图片 + 合影 + 1 份执行报告）", "7 个自然日内出具", "甲方", "2,000", "2,000"],
    ]

    t = doc.add_table(rows=1 + len(rows) + 1, cols=6)
    t.style = "Table Grid"
    t.autofit = False
    widths_cm = [1.2, 6.0, 4.0, 3.0, 1.8, 1.8]
    for i, w in enumerate(widths_cm):
        for row in t.rows:
            row.cells[i].width = Cm(w)
    headers = ["序号", "服务项 / 物料", "规格 / 数量", "责任方", "单价(元)", "金额(元)"]
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, bold=True, color=WHITE, size=10.5, align_center=True)
        shade(t.rows[0].cells[i], "006B3F")
    for r, row in enumerate(rows):
        cells = t.rows[r + 1].cells
        is_cat = (row[2] == "")
        if is_cat:
            cell_text(cells[0], row[0], bold=True, size=10.5, align_center=True)
            cell_text(cells[1], row[1], bold=True, size=10.5)
            for ci in range(2, 6):
                cell_text(cells[ci], "", size=10.5)
            for ci in range(6):
                shade(cells[ci], "E8F1EC")
        else:
            cell_text(cells[0], row[0], bold=True, size=10, align_center=True)
            cell_text(cells[1], row[1], size=10)
            cell_text(cells[2], row[2], size=10)
            cell_text(cells[3], row[3], size=10)
            cell_text(cells[4], row[4], size=10, align_center=True)
            cell_text(cells[5], row[5], size=10, align_center=True)
    total_cells = t.rows[-1].cells
    cell_text(total_cells[0], "合计", bold=True, color=WHITE, size=11.5, align_center=True)
    cell_text(total_cells[1], "人民币壹拾万元整（含税）", bold=True, color=WHITE, size=11.5)
    cell_text(total_cells[2], "—", color=WHITE, align_center=True)
    cell_text(total_cells[3], "—", color=WHITE, align_center=True)
    cell_text(total_cells[4], "—", color=WHITE, align_center=True)
    cell_text(total_cells[5], "￥100,000.00", bold=True, color=WHITE, size=12, align_center=True)
    for ci in range(6):
        shade(total_cells[ci], "006B3F")
    blank(doc)

    # Notes
    para(doc, "五、报价说明与排除项", bold=True, first_indent=False, size=12, color=GREEN)
    para(doc, "1. 本报价为整体打包价，含税；发票内容：会议服务费 / 赞助费。", size=10.5)
    para(doc, "2. 上述执行项均为甲方对乙方的可交付物，过程中由甲方逐项归档证据材料（现场照片、印刷成品照、屏幕截图、回执报告等），并在大会结束后 7 个自然日内随《赞助权益执行回执》一并提交乙方作为验收依据；无需逐项贴照片验收。", size=10.5)
    para(doc, "3. 经双方协商，本次合作不纳入或已调整的事项：", size=10.5)
    para(doc, "（1）白皮书扉页联合署名 — 不纳入；", size=10.5, first_indent=False)
    para(doc, "（2）媒体通稿「不少于 5 家头部财经/地产媒体」「标题级或副标题级」硬性数量与位置 — 不作硬性承诺；", size=10.5, first_indent=False)
    para(doc, "（3）双校长三角校友产业联盟战略合作伙伴永久入册 + 牌匾交接 — 不纳入；", size=10.5, first_indent=False)
    para(doc, "（4）现场引导员 — 由原 3 名调整为 1 名（含在 D1）；", size=10.5, first_indent=False)
    para(doc, "（5）同行业排他承诺 — 限定为「晚宴冠名战略合作伙伴」身份及「主背景板钻石级 logo 位」，其他级别合作不在排他范围。", size=10.5, first_indent=False)
    para(doc, "4. 下列事项由乙方自办，不计入本报价 100,000 元，亦不构成甲方义务：", size=10.5)
    para(doc, "（1）项目案场接驳（华为尊界 8–10 辆 + 考斯特补位）；", size=10.5, first_indent=False)
    para(doc, "（2）项目案场接待（沙盘 + 样板间，含夜场灯光秀，如适用）；", size=10.5, first_indent=False)
    para(doc, "（3）现场销售人员（电梯口 + 展场内合计 4 人，统一品牌服装与胸卡）；", size=10.5, first_indent=False)
    para(doc, "（4）现场展位升级包装（如背景墙、特装等，在甲方提供的基础包装上叠加）。", size=10.5, first_indent=False)
    para(doc, "5. 付款方式：合同（含本报价单签字盖章后视同确认）签订之日起 5 个工作日内一次性付款至甲方指定账户。", size=10.5)
    para(doc, "6. 发票开具：甲方在款项到账后 10 个工作日内向乙方开具等额合法有效的增值税普通发票或专用发票。", size=10.5)
    para(doc, "7. 本报价单与《附件四·晚宴冠名战略合作伙伴专项合作协议》及其「附表一·服务报价单 & 执行清单」内容一致；本报价单经双方加盖公章（合同章）后，作为本次合作的正式商业要约与验收依据。", size=10.5)
    blank(doc)

    para(doc, "六、收款账户信息", bold=True, first_indent=False, size=12, color=GREEN)
    kv_table(doc, rows=[
        ["收款单位", "________________________________"],
        ["开户银行", "________________________________"],
        ["银行账号", "________________________________"],
        ["税号", "________________________________"],
    ])
    blank(doc, 2)

    # Signature
    para(doc, "七、双方确认（签字 / 盖章）", bold=True, first_indent=False, size=12, color=GREEN)
    sig = doc.add_table(rows=4, cols=2)
    sig.style = "Table Grid"
    sig.autofit = False
    for row in sig.rows:
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(8.0)
    cell_text(sig.rows[0].cells[0], "甲方（盖章）：\n北京大学经济学院上海校友会 /\n复旦大学住房政策研究中心",
              bold=True, size=11)
    cell_text(sig.rows[0].cells[1], "乙方（盖章）：\n绿城中国控股有限公司 /\n____________________________（项目主体）",
              bold=True, size=11)
    cell_text(sig.rows[1].cells[0], "\n\n\n（公章 / 合同章位置）\n\n", size=10.5, align_center=True)
    cell_text(sig.rows[1].cells[1], "\n\n\n（公章 / 合同章位置）\n\n", size=10.5, align_center=True)
    cell_text(sig.rows[2].cells[0], "授权代表（签字）：______________________\n职务：__________________", size=11)
    cell_text(sig.rows[2].cells[1], "授权代表（签字）：______________________\n职务：__________________", size=11)
    cell_text(sig.rows[3].cells[0], "签字日期：______年______月______日", size=11)
    cell_text(sig.rows[3].cells[1], "签字日期：______年______月______日", size=11)

    blank(doc)
    para(doc, "本报价单一式肆份，甲乙双方各执贰份，具有同等法律效力。",
         align="center", first_indent=False, size=10, color=GREY)

    out = "/workspace/deliverables/绿城中国-潮鸣外滩-10万元晚宴冠名服务报价单(加急盖章版).docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
