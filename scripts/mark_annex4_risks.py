#!/usr/bin/env python3
"""以红字标注版附件四为主，对照盖章版主合同与以往方案，标注风险（黄底+黄字）并补充说明（蓝字）。"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

SRC = Path(
    "/workspace/deliverables/"
    "附件四-绿城中国-潮鸣外滩-晚宴冠名战略合作伙伴专项合作协议-按盖章版补充(红字标注).docx"
)
OUT = SRC  # 原位覆盖生成

YELLOW = RGBColor(0xFF, 0xC0, 0x00)  # 偏深黄，黄底上仍可辨认
YELLOW_FILL = "FFFF00"
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xFF, 0x00, 0x00)


def _ensure_rpr(run):
    r_pr = run._r.get_or_add_rPr()
    return r_pr


def paint_run_yellow(run) -> None:
    """风险事项：黄底覆盖 + 黄字。"""
    run.font.color.rgb = YELLOW
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r_pr = _ensure_rpr(run)
    # 同步单元格/字符底纹，避免部分阅读器不显示 highlight
    from docx.oxml import OxmlElement
    shd = r_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        r_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), YELLOW_FILL)


def paint_run_blue(run) -> None:
    """新增补充：蓝色字体。"""
    run.font.color.rgb = BLUE
    run.font.highlight_color = None


def _first_run_style(paragraph):
    size = None
    name = None
    bold = None
    for r in paragraph.runs:
        if size is None and r.font.size:
            size = r.font.size
        if name is None and r.font.name:
            name = r.font.name
        if bold is None and r.bold is not None:
            bold = r.bold
    return size, name, bold


def replace_paragraph_keep_style(paragraph, new_text: str, painter) -> None:
    size, name, bold = _first_run_style(paragraph)
    if paragraph.runs:
        first = paragraph.runs[0]
        first.text = new_text
        if size:
            first.font.size = size
        if name:
            first.font.name = name
            r_fonts = first._r.get_or_add_rPr().get_or_add_rFonts()
            r_fonts.set(qn("w:eastAsia"), name)
        if bold is not None:
            first.bold = bold
        painter(first)
        for r in paragraph.runs[1:]:
            r.text = ""
            # 清空后去掉旧颜色干扰
            painter(r) if False else None
    else:
        run = paragraph.add_run(new_text)
        if size:
            run.font.size = size
        painter(run)


def yellow_paragraph(paragraph) -> None:
    text = paragraph.text
    if not text.strip():
        return
    replace_paragraph_keep_style(paragraph, text, paint_run_yellow)


def yellow_cell(cell) -> None:
    for p in cell.paragraphs:
        if p.text.strip():
            yellow_paragraph(p)


def append_blue_paragraph(paragraph_or_cell, text: str, size_pt: float = 10.5):
    """在段落所属 body 元素后插入蓝色补充段，或在单元格内追加段落。"""
    # 单元格：直接追加
    if hasattr(paragraph_or_cell, "add_paragraph"):
        p = paragraph_or_cell.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size_pt)
        paint_run_blue(run)
        return p

    # 段落：在其后插入同级段落
    paragraph = paragraph_or_cell
    new_p = deepcopy(paragraph._p)
    # 清空克隆内容
    for child in list(new_p):
        if child.tag == qn("w:r") or child.tag == qn("w:hyperlink"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    # 通过 xpath 找到刚插入的元素并包一层 Paragraph 不便，直接写 run XML
    from docx.text.paragraph import Paragraph

    p_obj = Paragraph(new_p, paragraph._parent)
    run = p_obj.add_run(text)
    run.font.size = Pt(size_pt)
    paint_run_blue(run)
    return p_obj


def yellow_row_content_cells(row, start_ci: int = 0) -> None:
    seen = set()
    for ci, cell in enumerate(row.cells):
        if ci < start_ci:
            continue
        # 合并单元格去重
        cell_id = id(cell._tc)
        if cell_id in seen:
            continue
        seen.add(cell_id)
        yellow_cell(cell)


def append_blue_in_cell(cell, text: str, size_pt: float = 9) -> None:
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    paint_run_blue(run)


def find_paragraph(doc: Document, prefix: str):
    for p in doc.paragraphs:
        if p.text.startswith(prefix) or p.text.strip().startswith(prefix):
            return p
    raise KeyError(f"paragraph not found: {prefix!r}")


def main() -> None:
    doc = Document(str(SRC))

    # ---------- 正文风险条款（按文首定位，避免插入蓝字后索引错位） ----------
    # 2.4 同行业排他过宽（以往方案仅限定晚宴冠名 + 钻石 logo 位）
    p24 = find_paragraph(doc, "2.4")
    yellow_paragraph(p24)
    append_blue_paragraph(
        p24,
        "【补充·对照以往方案】排他范围建议收窄为：「晚宴冠名」身份及「主背景板钻石级 logo 位」；"
        "其他级别（黄金/铂金/基础曝光等）不在排他范围。过宽排他在已办结活动场景下难以核实与执行。",
        size_pt=10,
    )

    # 10.3 知识产权/保密违约金「不低于 5 万元」相对 10 万赞助额偏高
    p103 = find_paragraph(doc, "10.3")
    yellow_paragraph(p103)
    append_blue_paragraph(
        p103,
        "【补充·对照盖章版】盖章版将保密/知识产权违约排除在 30% 总上限之外；"
        "但本条「不低于人民币 50,000 元」已达赞助额 50%，建议双方确认是否改为「不超过赞助金额 30% 或双方另行书面约定」，"
        "避免与 10.2 / 盖章版风险上限体系口径冲突。",
        size_pt=10,
    )

    # ---------- 第三条表格：不可达成 / 难核实细项 ----------
    # TABLE 3：第 13 项 500 份手拎袋硬性数量
    t3 = doc.tables[3]
    yellow_row_content_cells(t3.rows[4])  # item 13
    append_blue_in_cell(
        t3.rows[4].cells[-1],
        "【补充】硬性「500 套」易超实际发放量；建议以「按实际参会嘉宾手拎袋发放数量为准、不超过 500 套」验收，避免不可达成数量争议。",
    )

    # TABLE 4：展位仅列 2 处 vs 盖章版 3 个；引导员 3 名；尊界 8–10 辆表述
    t4 = doc.tables[4]
    # 展位两行：与盖章版「3 个展位」不一致
    yellow_row_content_cells(t4.rows[1])
    yellow_row_content_cells(t4.rows[2])
    append_blue_in_cell(
        t4.rows[2].cells[-1],
        "【补充·对照盖章版】盖章版主合同约定现场展位共 3 个；本附件正文仅列电梯口+展场内 2 处。"
        "建议签署前书面确认第 3 处位置或明确「以盖章版 3 个展位为准、位置现场协商」。",
    )
    # 引导员 3 名
    yellow_row_content_cells(t4.rows[3])
    append_blue_in_cell(
        t4.rows[3].cells[3],
        "【补充·对照以往方案】现场引导员由 3 名调整为 1 名（合并至论坛后口播动线）；其余视嘉宾流量灵活补位，不作硬性 3 名承诺。",
    )
    # 尊界车队规模（乙方自办，但写死 8–10 辆 + 往返不低于 60 分钟易成争议点）
    yellow_row_content_cells(t4.rows[4])
    append_blue_in_cell(
        t4.rows[4].cells[3],
        "【补充】车队规模取决于意向人数；建议「尊界/考斯特由乙方视现场意向人数安排并承担费用」，删除或弱化固定 8–10 辆及「往返不低于 60 分钟」刚性表述。"
        "另：附表一写「8 辆」、正文写「8–10 辆」，口径需统一。",
    )

    # TABLE 5：媒体≥3家、永久入册、证书名称
    t5 = doc.tables[5]
    yellow_row_content_cells(t5.rows[1])  # 媒体 ≥3
    append_blue_in_cell(
        t5.rows[1].cells[2],
        "【补充·对照以往方案】媒体通稿不作「≥3 家主流媒体」硬性数量/位置承诺；甲方尽力推送，实际刊发以媒体自主决定为准。",
    )
    yellow_row_content_cells(t5.rows[4])  # 永久入册
    append_blue_in_cell(
        t5.rows[4].cells[1],
        "【补充·对照以往方案】「双校长三角校友产业联盟战略合作伙伴永久入册」以往方案已明确不纳入本次计价与必交付范围；执行机制另议，不作本附件硬性义务。",
    )
    yellow_row_content_cells(t5.rows[6])  # 证书名称
    append_blue_in_cell(
        t5.rows[6].cells[1],
        "【补充·对照盖章版】盖章版证书名称为「2026 年度智慧人居新质资产领军企业 暨卓越战略合作伙伴」；"
        "与本处「主办方之一战略合作伙伴证书」表述不一致，建议统一为盖章版全称。",
    )

    # ---------- 第六条时间节点：5/21 当日完成交稿+确认+下印，不可达成 ----------
    t6 = doc.tables[6]
    for ri in (2, 3, 4):  # 物料交付 / 设计确认 / 印刷下单 同日
        yellow_row_content_cells(t6.rows[ri])
    append_blue_in_cell(
        t6.rows[4].cells[-1],
        "【补充·执行风险】签约日（5/21）当日完成互交物料、一轮确认并全部下印，印刷物流周期客观上不可达成；"
        "且活动已于 5/22 完成，该节点表仅作过程记录时，不得以未满足「同日下印」倒追违约。建议注明「以实际已执行情况为准 / 视同过程节点豁免」。",
    )
    # 「名单对接」与盖章版个保法不交名单冲突
    yellow_cell(t6.rows[7].cells[3])
    append_blue_in_cell(
        t6.rows[7].cells[3],
        "【补充·对照盖章版】盖章版基于个保法不就参会名单作硬性提交承诺；「名单对接」应理解为现场协助促成意向客户，不得解释为必须交付完整嘉宾联络清单。",
    )

    # ---------- 附表一对应风险行 ----------
    t9 = doc.tables[9]
    # row 13: 500份
    yellow_row_content_cells(t9.rows[13])
    append_blue_in_cell(
        t9.rows[13].cells[-1],
        "【补充】数量以实际发放为准（上限 500），不作未达 500 即违约解释。",
    )
    # row 14: 白皮书 — 正文第三节表格缺失该项，且以往方案不纳入
    yellow_row_content_cells(t9.rows[14])
    append_blue_in_cell(
        t9.rows[14].cells[2],
        "【补充·对照以往方案/正文结构】以往方案已明确「白皮书扉页联合署名」不纳入；"
        "且第三节权益表序号由 13 跳至 15，正文未列本项，附表保留易导致验收争议——建议本项打「×」或不纳入必交付。",
    )
    # row 17: 引导员 3 名
    yellow_row_content_cells(t9.rows[17])
    append_blue_in_cell(
        t9.rows[17].cells[3],
        "【补充】引导员调整为 1 名，不作 3 名硬性承诺。",
    )
    # row 20: 媒体≥3
    yellow_row_content_cells(t9.rows[20])
    append_blue_in_cell(
        t9.rows[20].cells[2],
        "【补充】不作「≥3 家主流媒体」硬性刊发承诺。",
    )
    # row 23: 永久入册
    yellow_row_content_cells(t9.rows[23])
    append_blue_in_cell(
        t9.rows[23].cells[2],
        "【补充】以往方案不纳入本次必交付；「永久」表述无法验收，建议打「×」或改为双方另议。",
    )
    # row 25: 证书名称
    yellow_row_content_cells(t9.rows[25])
    append_blue_in_cell(
        t9.rows[25].cells[2],
        "【补充】证书名称与盖章版统一为「2026 年度智慧人居新质资产领军企业 暨卓越战略合作伙伴」。",
    )

    # ---------- 文首总说明（蓝字，统一口径） ----------
    anchor = find_paragraph(doc, "上述权益项的最终执行清单")
    append_blue_paragraph(
        anchor,
        "【补充·风险对照总说明】本附件黄底黄字标注事项，系对照《活动赞助合同》（盖章版）及以往专项方案中"
        "「已删除/弱化（甲方不好实现项）」梳理出的不可达成或难核实风险，主要包括："
        "①白皮书扉页署名；②媒体通稿硬性家数；③校友联盟「永久入册」；④引导员 3 名；"
        "⑤手拎袋 500 套刚性数量；⑥展位数量与盖章版 3 个不一致；⑦5/21 同日交稿确认下印节点；"
        "⑧证书名称口径不一致；⑨过宽同业排他。验收口径建议与盖章版一致：按四大模块整体打包验收，"
        "不以细项逐条未达为由拒付无争议款项；结算与违约上限以盖章版为准。",
        size_pt=10,
    )

    # 附表一引导段落后再补一句验收原则
    for p in doc.paragraphs:
        if p.text.startswith("本附表逐项列示"):
            append_blue_paragraph(
                p,
                "【补充·验收原则】黄标细项如双方未另签确认，不作为单项违约计价依据；"
                "与盖章版或报价单四大模块冲突时，以盖章版/整体打包验收为准。",
                size_pt=10,
            )
            break

    doc.save(str(OUT))
    print(f"saved: {OUT}")


if __name__ == "__main__":
    main()
