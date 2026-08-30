#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""8 月 30 日对方三点回复后：微信稿、再确认函、内部研判。"""

import sys
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
import generate_wangdefeng_proposal as g  # noqa: E402
from generate_external_outputs import _mini_table, _p, set_cell_margins  # noqa: E402

OUT = Path(__file__).resolve().parents[1] / "output"

WECHAT_PATH = OUT / "致赵海_8月30日对方三点后回复稿.txt"
LETTER_PATH = OUT / "致赵海_8月30日合作路径再确认.docx"
INTERNAL_PATH = OUT / "致赵海_8月30日对方三点_内部研判.docx"

WECHAT_BUBBLE_1 = """赵海好。

三点都看到了。中欧三名学员缴费，恭喜，说明主课自己的高端渠道是通的。分流的担心我也理解：9 月我这边不对外铺 299 / 399 / 499，避免价格锚点碰到你们主课。

但 10 月 31 日这场公开讲座，日期不能往后放，也不能改成「先推主课、单课以后再说」。9 月 30 日是公开场的开售节点，不是改期。王老师可以同时是二期亮点；公开场仍是对外讲座，不是二期附属单课。

路径我对了一版，你转给团队看。"""

WECHAT_BUBBLE_2 = """【回复三点，请转达】

一、公开场：10 月 31 日不变；9 月不开售，9 月 30 日起预售
理解主课要价格保护。9 月不对外推出 299 / 399 / 499。
10 月 31 日仍做 2.5–3 小时公开讲座，标准票 499。9 月 30 日是开售节点，不是把公开场改期或取消。
对外不叫「体验课 / 单课」，叫「公开讲座」。499 可抵系统课学费，深度学习者是被导进主课，不是被截走。
若二期也在 10 月 31 日：公开场是对外场次，系统课学员可内部入场，不是把当天让给二期课堂。

二、主课渠道费：同意作附带，按阶梯；不替代公开场
我方独立介绍并缴费的系统课学员，按实收学费：1–5 人 20%；6–10 人 25%；11 人及以上 30%。建议达到该档后，该档全部人数按该比例计算。
此项是公开场之外的转化合作，不是把我这边改成主课招生渠道。人数不作承诺，不替二期冲刺凑班。

三、群：按你说的执行
我企业微信建群、我做群主不转让；你做管理员，内容、老师动态你发。群名用「活动交流群」。续课群同样。群内先只提老师和时间，不对外报价。群用于 10 月 31 日公开场，不改成二期招生群。

公开场仍按联合主办对接场地、组织和渠道。此前确认函四点仍然有效。"""

WECHAT_FALLBACK = """9 月不铺低价票，这点我让。10 月 31 日公开讲座日期不改，9 月 30 日只是开售。主课阶梯 20 / 25 / 30 可以，作为转化，不是替代公开场。群按你说的办。"""

WECHAT_IF_DATE = """二期和公开场可以同一天分层：对外是公开讲座，系统课学员走内部名额。老师当天的对外场次在公开场。二期课堂不要占掉 10 月 31 日对外场。"""


def write_wechat_txt() -> Path:
    text = "\n".join(
        [
            "致赵海｜8月30日对方三点后回复稿（直接复制，分两条发）",
            "说明：先发第一条，等他回「好」或「你发来我转」，再发第二条。",
            "",
            "======== 第一条（对人） ========",
            "",
            WECHAT_BUBBLE_1,
            "",
            "======== 第二条（可转发） ========",
            "",
            WECHAT_BUBBLE_2,
            "",
            "======== 若对方仍要先推主课、单课以后再说 ========",
            "",
            WECHAT_FALLBACK,
            "",
            "======== 若对方说 10 月 31 日二期要用老师 ========",
            "",
            WECHAT_IF_DATE,
            "",
        ]
    )
    WECHAT_PATH.write_text(text, encoding="utf-8")
    return WECHAT_PATH


def _rose_box(doc, text):
    box = doc.add_table(rows=1, cols=1)
    cell = box.cell(0, 0)
    g.shade_cell(cell, g.ROSE_HEX)
    g.set_cell_border(
        cell,
        top={"val": "single", "sz": "6", "color": g.RED_HEX, "space": "0"},
        left={"val": "single", "sz": "16", "color": g.RED_HEX, "space": "0"},
        bottom={"val": "single", "sz": "6", "color": g.RED_HEX, "space": "0"},
        right={"val": "single", "sz": "6", "color": g.RED_HEX, "space": "0"},
    )
    cell.text = ""
    set_cell_margins(cell, top=40, bottom=36, left=80, right=80)
    _p(cell, text, size=8.5, after=0)
    g.set_table_full_width(box, [18.5])
    return box


def build_letter() -> Path:
    doc = g.new_doc()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.4)
    section.footer_distance = Cm(0.35)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("致赵海  ·  可转达团队  ·  2026年8月30日")
    g.set_run_font(run, size=8, color=g.MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("本函为沟通确认文件，不构成合同。最终以双方签署的协议为准。请就下列路径一并书面回复。")
    g.set_run_font(fr, size=7.5, color=g.MUTED)

    g.add_para(
        doc,
        "10 月 31 日王德峰老师活动｜合作路径再确认",
        size=15,
        bold=True,
        color=g.NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0,
        space_after=3,
        line_spacing=1.0,
    )
    g.add_para(
        doc,
        "9 月不开售低价票；10 月 31 日公开讲座日期不变。主课阶梯渠道费为附带，不替代公开场。",
        size=9,
        bold=True,
        color=g.GOLD,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
        line_spacing=1.05,
    )
    g.add_para(
        doc,
        "赵海好。贵方三点收悉。中欧三名学员缴费，祝贺。主课价格保护的考虑，我方同意 9 月不对外推出 299 / 399 / 499。但 9 月 30 日是公开场开售节点，不是改期；10 月 31 日公开讲座照常举办，不改为先推主课、单课以后再说。",
        size=9,
        space_before=0,
        space_after=4,
        line_spacing=1.1,
    )

    g.add_para(doc, "一、公开讲座：日期不变，开售后置", size=11, bold=True, color=g.NAVY, space_before=2, space_after=3, line_spacing=1.0)
    _mini_table(
        doc,
        ["事项", "我方确认"],
        [
            ["举办日期", "2026 年 10 月 31 日，2.5–3 小时公开讲座，日期不改"],
            ["9 月", "不对外开售 299 / 399 / 499，避免冲击主课价格锚点"],
            ["开售节点", "9 月 30 日起预售；9 月 30 日是开售日，不是改期"],
            ["票价与名称", "标准票 499 元；对外称公开讲座，不称体验课或单课"],
            ["与主课关系", "499 可抵系统课学费；深度学习者导入主课，不是分流"],
            ["与二期同日", "公开场为对外场次；系统课学员可内部入场，不让出当天"],
        ],
        [3.6, 14.9],
    )

    g.add_para(doc, "二、主课渠道费：同意阶梯，仅作附带", size=11, bold=True, color=g.NAVY, space_before=7, space_after=3, line_spacing=1.0)
    g.add_para(
        doc,
        "我方独立介绍并缴费的系统课学员，按实收学费阶梯计：1–5 人 20%，6–10 人 25%，11 人及以上 30%。建议满档全额。此项不替代 10 月 31 日公开场，人数不作承诺。",
        size=8.5,
        space_before=0,
        space_after=3,
        line_spacing=1.08,
    )
    _rose_box(
        doc,
        "我方不能接受：取消或改期 10 月 31 日公开讲座；把公开场改成二期附属单课；把我方改成主课招生渠道、只拿阶梯渠道费。王老师可作为二期亮点，但不替代对外公开场。",
    )

    g.add_para(doc, "三、社群：按贵方提议执行", size=11, bold=True, color=g.NAVY, space_before=7, space_after=3, line_spacing=1.0)
    _mini_table(
        doc,
        ["事项", "双方已对齐"],
        [
            ["建群 / 群主", "我方企业微信创建，我方任群主且不转让"],
            ["管理员", "贵方任管理员，内容与老师动态由贵方发"],
            ["群名 / 续课群", "「活动交流群」；续课群同样，先只提老师与时间，不对外报价"],
            ["边界", "用于 10 月 31 日公开场，不改成二期招生群；不得转群主或导出名单另建群"],
        ],
        [3.6, 14.9],
    )

    g.add_para(doc, "四、此前确认函仍有效", size=11, bold=True, color=g.NAVY, space_before=7, space_after=3, line_spacing=1.0)
    next_box = doc.add_table(rows=1, cols=4)
    items = [
        ("① 票价", "标准票 499 元；9 月不开售。"),
        ("② 分成", "公开场按确认函；主课阶梯仅附带。"),
        ("③ 社群", "我方任群主，贵方任管理员。"),
        ("④ 角色", "联合主办，不是主课渠道。"),
    ]
    fills = [g.SOFT_HEX, g.CREAM_HEX, g.GREEN_HEX, g.AMBER_HEX]
    for i, ((title, body), fill) in enumerate(zip(items, fills)):
        c = next_box.cell(0, i)
        c.text = ""
        g.shade_cell(c, fill)
        set_cell_margins(c, top=46, bottom=46, left=60, right=60)
        _p(c, title, size=8.5, bold=True, color=g.NAVY, after=2)
        _p(c, body, size=7.5, after=0)
    g.set_table_full_width(next_box, [4.625, 4.625, 4.625, 4.625])

    g.add_para(
        doc,
        "请一并确认后回复。我方（联合运营方）　　胡继刚　　2026 年 8 月 30 日",
        size=9,
        bold=True,
        color=g.NAVY,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        space_before=8,
        space_after=0,
        line_spacing=1.0,
    )

    doc.save(LETTER_PATH)
    return LETTER_PATH


def build_internal() -> Path:
    doc = g.new_doc()
    g.setup_section(doc, "对方三点回复 · 内部研判（不要发给对方）", confidential=True)

    g.add_para(
        doc,
        "对方三点怎么看，怎么回",
        size=18,
        bold=True,
        color=g.NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    g.add_para(
        doc,
        "只给胡总看。不要把本文、心仪评价、养病细节、中欧是施压发给对方。",
        size=10,
        bold=True,
        color=g.RED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )

    g.add_h1(doc, "一、一句话判断")
    g.add_para(
        doc,
        "表面让了三步：公开场可以留、渠道费加到 30%、群主归你。实际是把你从联合主办收成「9 月先帮他们招二期」。公开场被改成 9 月 30 日以后再说的单课；王老师被收成二期核心亮点。群主是唯一真让步，必须书面锁住。",
    )

    g.add_h1(doc, "二、三点拆开")
    g.add_table(
        doc,
        ["对方说", "实际含义", "你的应对"],
        [
            [
                "公开场可保留，现阶段不宜推出；先主课，9 月 30 日再推单课",
                "不定日期的口头保留。8–9 月你只做渠道。10 月 31 日可能被二期开班占掉",
                "让开售节点：9 月不铺低价。不让日期：10 月 31 日公开讲座不变",
            ],
            [
                "王老师作为二期核心亮点，吸引报整体课程",
                "10 月 31 日变成他们的课，不是你的对外公开场",
                "二期亮点他们可以对系统课学员讲；对外场次仍是公开讲座",
            ],
            [
                "渠道费 1–5 人 20%、5–10 人 25%、10 人以上 30%",
                "还是渠道逻辑。1–5 与 5–10 边界重叠。10 人 × 13000 × 30% 约 3.9 万，换掉公开场不值",
                "可作附带；改成 1–5 / 6–10 / 11+；按实收；人数不承诺",
            ],
            [
                "群主按你说的办 + 中欧三人缴费",
                "群是真让步。中欧是 FOMO，催你用信誉推 13000",
                "群书面锁定。恭喜中欧，证明主课有自己的高端渠道，更不必停公开场",
            ],
        ],
        [4.2, 6.3, 6.0],
    )

    g.add_h1(doc, "三、分流论为什么不成立（内部，勿原文发给对方）")
    g.add_para(
        doc,
        "真会掏 13000 的人，不会因为多一场 2.5 小时讲座就放弃系统课。绝大多数 499 客本来就不会直接报 13000；没有公开场，就没有转化池。中欧三人缴费，说明主课客群是高端学历圈，和公开场不是一拨人。分流只发生在把 499 和 13000 都叫成「王德峰的课」的时候。所以对外改名称、499 可抵学费、9 月不铺低价，已经回应他们的担心。用分流当理由停掉 10 月 31 日，是占场，不是保护主课。",
    )

    g.add_h1(doc, "四、让什么，不让什么")
    g.add_table(
        doc,
        ["可以让", "不能让"],
        [
            ["9 月不对外开售 299 / 399 / 499", "10 月 31 日改期、取消，或让给二期课堂"],
            ["开售放到 9 月 30 日（距活动约一个月）", "先推主课，公开场变成以后再说的单课"],
            ["对外叫公开讲座，不叫体验课 / 单课", "把联合主办改成主课招生渠道"],
            ["499 可抵系统课学费，帮他们转化", "承诺凑 10 人拿 30%"],
            ["主课转化阶梯 20 / 25 / 30 作为附带", "群变成二期招生群"],
        ],
        [8.25, 8.25],
    )

    g.add_h1(doc, "五、怎么发")
    g.add_table(
        doc,
        ["步骤", "做法"],
        [
            ["先发第一条", "恭喜中欧、承认分流担心、让出 9 月不开售。把「改期」挡回去。"],
            ["再发第二条", "可转发。钉死日期、开售节点、渠道是附带、群已同意。"],
            ["需要书面", "再发《合作路径再确认》Word。"],
            ["不要", "不要嘲中欧三人、不要说他们没经验、不要提养病、不要算 3.9 万给他们看。"],
        ],
        [4.0, 12.5],
    )

    g.add_h1(doc, "六、对方再顶怎么说")
    g.add_table(
        doc,
        ["对方说", "你回"],
        [
            ["先帮我们把二期招满再谈公开场。", "9 月不铺低价票已经让了。公开场 10 月 31 日不改，9 月 30 日开售。"],
            ["10 月 31 日二期要用王老师。", "可以同一天分层：对外公开讲座，系统课学员走内部名额。对外场次在公开场。"],
            ["你先推 10 个人，渠道费给你 30%。", "阶梯可以，人数不承诺。我这边不替二期冲刺凑班。"],
            ["低价一出，深度学习者就跑了。", "所以 9 月不开售，对外叫讲座不叫课，499 可抵系统课。这是导流，不是分流。"],
            ["中欧都缴费了，课好卖。", "恭喜。主课走高端渠道是通的。公开场是另一拨人，两件事一起做。"],
        ],
        [6.2, 10.3],
    )

    g.add_para(
        doc,
        "内部文件，勿外传。",
        size=9,
        color=g.RED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=16,
    )

    doc.save(INTERNAL_PATH)
    return INTERNAL_PATH


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    wechat = write_wechat_txt()
    letter = build_letter()
    internal = build_internal()
    print(f"已生成微信稿：{wechat}")
    print(f"已生成确认函：{letter}")
    print(f"已生成内部研判：{internal}")


if __name__ == "__main__":
    main()
