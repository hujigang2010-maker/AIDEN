#!/usr/bin/env python3
"""生成《发展与去留：2026 中美判断备忘录》Word 长文。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


TEAL = RGBColor(0x0F, 0x76, 0x6E)
INK = RGBColor(0x0F, 0x17, 0x2A)
SLATE = RGBColor(0x47, 0x55, 0x69)
ROSE = RGBColor(0xBE, 0x12, 0x31)
AMBER = RGBColor(0xD9, 0x77, 0x06)


def set_run_font(run, *, size=12, bold=False, color=None, name="微软雅黑"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading_cn(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(
            run,
            size=16 if level == 1 else 13,
            bold=True,
            color=TEAL if level == 1 else INK,
        )
    return p


def add_para(doc, text, *, size=11, bold=False, color=None, space_after=8, first_line=True, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=11, color=SLATE)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.4
    return p


def add_bullets(doc, items, *, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=size)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.3


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=INK)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size=10)
    doc.add_paragraph()


def build_docx(output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("发展与去留")
    set_run_font(run, size=22, bold=True, color=INK)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("2026 中美判断备忘录")
    set_run_font(run, size=16, bold=True, color=TEAL)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("默认深耕 · 出海做成期权 · 不归根也可以不斩根")
    set_run_font(run, size=11, color=SLATE)
    meta.paragraph_format.space_after = Pt(4)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("2026 年 8 月 · 供个人决策存档，不构成移民、投资或法律建议")
    set_run_font(run, size=10, color=SLATE)
    date_p.paragraph_format.space_after = Pt(16)

    add_heading_cn(doc, "〇、先给判断，再展开理由", 1)
    add_para(
        doc,
        "留在中国还是去美国，不是一道道德题，也不是一道爱国题。它是一道匹配题：你的能力结构，放在哪一套制度、市场和关系网里更值钱；你要过的日子，哪一边更能把心安养出来。",
    )
    add_para(
        doc,
        "对网络型、场景型、转化型从业者——也就是靠本地信任、产业接口、高校与园区、把新技术接到真实组织里吃饭的人——2026 年的默认答案是：留在中国继续发展，而且是换一种发展；把「去美国」做成有限期期权，而不是斩根式流亡。人工智能的机会在中国抓得住，但抓住的是应用、场景、空间与商业化，不是美西少数实验室里的前沿训练。没有具体岗位、签证路径和薪酬或意义溢价，就谈不上「果断去海外」。",
    )
    add_para(
        doc,
        "对方那段话很动人。它适合已经走在路上的人用来对抗乡愁，不适合还在十字路口的人用来代替判断。真正可取的只有半句：心安之处即是家。前半句「不要回来、不要写信、把这里的一切封存」是战时流亡的操作系统，不是 2026 年有护照、有职业选择的人该签下的人生协议。",
    )

    add_heading_cn(doc, "一、如何看待那段话", 1)
    add_para(doc, "先把原文放在这里，不改一个字：", first_line=True)
    add_quote(doc, "不要回来，不要想念我们，不要回头，不要写信，不要向乡愁屈服，不要落叶归根，而要落地生根，心安之处即是家！")
    add_quote(doc, "请把关于这里的一切都封存起来，别让思念成为你的软肋。只要内心安定，脚下就是你的家。")
    add_quote(doc, "别回头，向前走。不归根，只生根。心安处，即吾乡。")

    add_heading_cn(doc, "1. 它在解决什么", 2)
    add_para(
        doc,
        "移民史、下南洋史、战争流亡史上，都有这一类「斩断嘱咐」。长辈或同行者害怕的不是远行本身，而是人被撕成两半：身体已经离开，心还停在旧地址，于是既走不远，也回不去。所以他们用决绝的修辞，把思念从决策权里拿掉。就心理技术而言，这是有效的。很多人走不稳，不是因为新地方不够好，而是旧地方一直在投票。",
    )
    add_para(
        doc,
        "「落地生根」「心安处即吾乡」也有一层真：家确实不必等于出生地。一个人若能在新的坐标上把工作、关系、身体节奏和意义感重新养出来，那就是家。这半句值得留下。",
    )

    add_heading_cn(doc, "2. 它在回避什么", 2)
    add_para(
        doc,
        "这段话把「走」当成已经完成的决定，只处理走了以后怎么活。它完全没有帮你回答：该不该走、现在走是不是时候、走了之后你还剩什么可交易的能力。把诀别辞当成战略，是最常见的误用。",
    )
    add_para(
        doc,
        "第二处误用，是把家简化成地理选择。家当然可以重建，但职业资本、信任网络、语言习惯、对规则的直觉，都有路径依赖。换护照或换城市，不等于换命运。对实验室科学家，故土是背景；对靠本地关系做转化的人，故土是资产负债表。",
    )
    add_para(
        doc,
        "第三处误用，是把「封存」当成勇敢。「别让思念成为软肋」这句话只对了一半。思念本身不是软肋，失控的思念才是。封存关于这里的一切，看起来像决绝，对网络型从业者却接近自我废武功：你最值钱的东西，往往正是「关于这里的一切」。",
    )

    add_heading_cn(doc, "3. 可取的内核，不可取的协议", 2)
    add_para(
        doc,
        "建议把原话改写成一套能用的纪律，而不是一份永别书：可以走，但要带着根走；可以想家，但不让想家做决定；可以回头看，但每周只看一次；可以写信，把思念变成连接而不是抽签。不必归根，也不必斩根。先在能创造价值的地方生根。心安处即吾乡——但心安要自己养，不能靠切断记忆来换。",
    )
    add_para(
        doc,
        "家是养出来的，不是逃出来的。如果一个人必须把过去全部封存才能迈步，那说明迈步的理由还不够硬，硬的是情绪。情绪可以出发，不能领航。",
    )

    add_heading_cn(doc, "二、现在的中国，适不适合继续发展", 1)
    add_heading_cn(doc, "1. 先看 2026 上半年的事实，不看情绪", 2)
    add_para(
        doc,
        "国家统计局公布：2026 年上半年国内生产总值 69.57 万亿元，同比增长 4.7%。一季度 5.0%，二季度回落到 4.3%。城镇调查失业率均值 5.2%，与上年同期持平。CPI 同比上涨 1.0%，PPI 由负转正至 1.5%。这不是崩盘叙事，也不是高歌猛进。它是一个大体量经济体在外部不确定、内部结构转换中的「还在合理区间」。",
    )
    add_para(
        doc,
        "真正决定个人命运的不是 4.7% 这个平均数，而是双速。冷的一侧：固定资产投资同比下降 5.7%，民间投资下降 8.5%，房地产开发投资下降 18.0%，新建商品房销售面积下降 11.6%、销售额下降 13.6%。热的一侧：规模以上高技术制造业增加值增长 13.3%，数字产品制造业增长 12.3%，装备制造业增长 9.3%；规上工业企业利润增长 18.7%，其中电子行业利润增长 96.9%；信息传输、软件和信息技术服务业增加值增长 10.7%。官方口径称，高端制造、数智经济、现代服务等新动能对上半年增长贡献率超过四成。与人工智能相关的集成电路、智能车载设备等行业保持 30% 以上高增长；新能源汽车零售渗透率连续三个月超过 60%。",
    )
    add_para(
        doc,
        "读法应当很冷静：中国经济没有「结束」，但它已经不允许你用 2010 到 2020 年的方法继续发展。旧引擎还在出清，新动能已经能看见利润，而不只是口号。站在冷的一侧会觉得「不适合再发展」；站在热的一侧会觉得「窗口正在打开」。两种体感都真实，因为它们不在同一条速度上。",
    )

    add_heading_cn(doc, "2. 「适不适合」必须定义「发展」", 2)
    add_para(doc, "把「发展」拆成四种，结论就不含糊：", first_line=True)
    add_bullets(
        doc,
        [
            "旧发展：拿地、加杠杆、等升值。2026 年不适合。房地产仍在深度调整，开发投资与销售继续下滑，风险出清没有结束。",
            "守成发展：躺平等周期、等政策、等风口回头。不适合。等待本身会消耗最好的三到五年。",
            "转化发展：把人工智能、先进制造、数智服务接到具体产业、空间和组织。适合。这正是新动能兑现利润的一层。",
            "网络发展：在上海这类城市做高校、园区、企业、资本之间的接口。适合。接口人的溢价在结构转换期通常上升，而不是下降。",
        ],
    )
    add_para(
        doc,
        "所以，更准确的句子不是「中国不适合发展」，而是「中国不适合用旧剧本发展，仍适合用新动能剧本发展」。对已经把主业绑在增量开发上的人，继续原样做下去会越来越难，必须转向存量、资产管理、产业空间载体。对已经站在 AI 商业化、园区、高校与产业接口上的人，现在离开，等于在热的一侧尚未充分兑现时离场。",
    )

    add_heading_cn(doc, "3. 不适感从哪里来", 2)
    add_para(
        doc,
        "很多人觉得「不适合再待」，往往不是 GDP 数字，而是体感：预期不稳、民营投资偏弱、年轻人就业难、行业话语权从地产转向硬科技、原来有效的关系忽然变现变慢。这些体感应当被承认。承认之后要做的，是区分「宏观不适」和「个人错配」。宏观不适几乎人人都有；个人错配，才是该不该走的理由。如果你的不适来自旧剧本失效，解法是换剧本，不是换国家。如果你的不适来自你真正想做的事在这里做不成，而那边有具体机会，那才进入「走」的讨论。",
    )

    add_heading_cn(doc, "三、人工智能的机会，能不能就地抓住", 1)
    add_heading_cn(doc, "1. 中国抓得住什么", 2)
    add_para(
        doc,
        "2026 年中国人工智能最扎实的机会，不在「再造一个 OpenAI」，而在把已有能力接到最大的应用与制造市场上。上半年电子行业利润近翻倍，直接原因之一就是人工智能与各领域加速融合、算力需求大幅增长。集成电路、智能机器人、智能车载、设备更新、工业软件，都在利润和产量上给出证据。高技术服务业投资也快于整体投资。这说明：钱和订单正在往「AI + 实体」移动，而不只是往概念移动。",
    )
    add_para(
        doc,
        "对转化者，窗口尤其清楚。地方政府要场景和试点，园区要内容和招商，制造企业要降本和质检，高校要产学研接口，会展与空间要「可被看见的 AI」。这些需求不需要你成为顶尖研究者，需要你能把技术、空间、组织和资金翻译成可执行项目。这是中国相对美国更厚的一层土壤。",
    )

    add_heading_cn(doc, "2. 中国抓不住什么，以及不必硬抓什么", 2)
    add_para(
        doc,
        "Frontier 基础模型的资本密度、顶尖研究者密度、芯片与开源话语权，仍高度集中在美西少数机构。没有实验室履历的人，人到了美国，也不会自动进入 OpenAI、Anthropic、Nvidia 的核心圈。把「去美国」等同于「抓住人工智能」，是把地理迷信当成了产业判断。",
    )
    add_para(
        doc,
        "同样不必硬抓的，是纯概念平台、空转 PPT、没有场景的「AI 中台」。机会只认交付。用焦虑代替订单，在哪一国都一样会失败。",
    )

    add_heading_cn(doc, "3. 美国的高地与关隘要分开看", 2)
    add_para(
        doc,
        "美国作为人工智能高地，这一点 2026 年没有过时。基础模型、算力生态、风险资本、英语世界的标准和薪酬锚仍在那里。头部机构仍在全球抢人。如果你已经被这类机构锁定，去，是升级。",
    )
    add_para(
        doc,
        "但 2026 年的关隘同样真实。学生签证、OPT、H-1B、就业类绿卡全线收紧；新 H-1B 的高额费用显著打击中小企业和创业公司的雇佣能力；身份不确定本身会吞噬精力与议价权。公开报道显示，高技能人才已在分流到加拿大、英国、阿联酋等政策更可预期的地方。与此同时，中国也在用类 H-1B 的人才签证与揽才动作对冲。结论很硬：美国仍是高地，但不再是对所有外国人默认敞开的高地。「果断去海外」只对已拿到具体岗位、签证路径和溢价的人成立。",
    )

    add_heading_cn(doc, "四、能力结构，而不是国家口号", 1)
    add_para(
        doc,
        "去留的正确单位不是「中国人 / 美国人」，而是「你的哪一项能力在哪边溢价更高、迁移成本更低」。",
    )
    add_table(
        doc,
        ["能力类型", "在中国的溢价", "在美国的溢价", "迁移成本"],
        [
            ["Frontier 研究 / 训练", "中（少数实验室）", "极高", "中（若已有论文与推荐）"],
            ["大厂工程 / 算法岗", "高，但内卷", "高，但签证卡脖子", "高"],
            ["产业场景转化 / 落地", "很高（政府+制造+园区）", "中低（要重做信用）", "极高"],
            ["高校 / 园区 / 商会网络", "核心资产", "接近归零", "极高"],
            ["不动产 + 产业空间", "存量时代仍有专业位", "需牌照与本地信用", "极高"],
            ["英语内容 / 跨境合作", "稀缺加分", "入场券，不是护城河", "中"],
        ],
    )
    add_para(
        doc,
        "读这张表时不要自欺。如果一个人的核心资产是高校、园区、商会、不动产与产业空间的接口，那么去美国首先不是「寻找机会」，而是「资产归零后再创业」。这可以做，但必须按创业的标准来评估：启动资金、客户从哪里来、三年如何养活自己、失败了如何回来。不能按「那边更先进」的感觉来评估。先进是环境，不是你的利润表。",
    )

    add_heading_cn(doc, "五、三条路径，以及什么情况下才该走", 1)
    add_heading_cn(doc, "1. 路径 A：深耕（默认）", 2)
    add_para(
        doc,
        "把人工智能 × 产业空间 × 高校 / 园区做成一条主航道。停止用情绪阅读宏观，改用线索、试点、订单、复购来阅读自己。把地产旧能力转向存量、资产管理、场景载体，而不是等待房价叙事回归。90 天内形成可对外讲清楚的一条产品线：你卖什么、卖给谁、凭什么是你。",
    )
    add_heading_cn(doc, "2. 路径 B：双根（升级项）", 2)
    add_para(
        doc,
        "不去「移民」，去做 12 到 36 个月的项目驻留、访问、合作。保留中国主体与网络，海外只做增量能力。英语一页纸、跨境项目、对口机构，先于签证。走之前写好返回条款：何时回、带回什么、本地谁在值守。这是成年人的出海，不是少年的出走。",
    )
    add_heading_cn(doc, "3. 路径 C：有条件迁徙（纪律，不是浪漫）", 2)
    add_para(doc, "同时点亮下列条件，才进入「果断」：", first_line=True)
    add_bullets(
        doc,
        [
            "硬条件：已有雇主或机构的书面机会；签证路径清楚，不靠赌抽签；有三年财务缓冲，不把全部流动性押上。",
            "事业条件：去了能做这里做不了的事；不是逃避失败，而是升级能力；走之前本地资产有交接，而非蒸发。",
            "心理条件：能忍受身份不确定与可能的社会降级；不需要靠「封存故乡」才能行动；家人知情同意，不是一个人的诗。",
        ],
    )
    add_para(
        doc,
        "现在更可能的状态是：本地仍有航道，海外尚无具体机会。这落在矩阵里「深耕中国」那一格。那段话可以收藏，但不要让它代替这张清单。",
    )

    add_heading_cn(doc, "六、90 天把判断做成行为", 1)
    add_para(doc, "判断若不落地，会变成新的焦虑来源。建议严格按三个 30 天执行：", first_line=True)
    add_heading_cn(doc, "D1–30 收束", 2)
    add_bullets(
        doc,
        [
            "只留一条主航道。写清：你卖的是什么，卖给谁，凭什么是你。",
            "停掉所有「既要出国又要本地爆」的并行幻想。并行在这个阶段等于两头落空。",
            "把地产旧能力映射到产业空间 / 存量资产管理，列出三件可立刻做的事。",
        ],
    )
    add_heading_cn(doc, "D31–60 验证", 2)
    add_bullets(
        doc,
        [
            "至少推进一个可报价、可试点的 AI 场景项目。没有报价，就还在聊天。",
            "用周复盘看三项：线索、会议、订单，哪一项在动。不动的动作停掉。",
            "宏观新闻降权，客户反馈升权。国家统计局不能替你签合同。",
        ],
    )
    add_heading_cn(doc, "D61–90 期权", 2)
    add_bullets(
        doc,
        [
            "准备一套英文一页纸加两个项目案例。这是期权的最小可行产品。",
            "锁定一到两个跨境合作接口：机构、校友、展会，而不是移民中介。",
            "明确纪律：没有书面机会，就不谈移民时间表。把「去美国」从情绪议题降级为商务议题。",
        ],
    )

    add_heading_cn(doc, "七、结语", 1)
    add_para(
        doc,
        "对方要你不要回来、不要想念、不要回头。我建议你把这句话听成善意，而不是听成命令。善意是：走了以后，不要被乡愁撕开。命令是：为了证明自己够决绝，先把根砍掉。前者是成人，后者是誓言成瘾。",
    )
    add_para(
        doc,
        "2026 年的中国，仍适合一种发展：把人工智能接到真实的产业、空间与组织里。这段机会对转化者比对逃亡者更友好。美国仍是前沿高地，但城门在收；没有具体机会就「果断离开」，是把乡愁问题做成了错误的战略。",
    )
    add_para(
        doc,
        "别回头，向前走。不归根，也可以不斩根。心安处，即吾乡。先在能创造价值的地方，把心安养出来。",
    )

    add_heading_cn(doc, "附录：数据与边界", 1)
    add_para(
        doc,
        "本文宏观数据主要依据国家统计局 2026 年 7 月发布的上半年国民经济运行情况及记者会口径，包括 GDP、投资、房地产、就业、物价、高技术制造与新动能贡献等。美国移民与 AI 人才流向依据 2026 年公开报道（H-1B 政策收紧与费用上升、头部 AI 公司仍在申请、人才向加拿大 / 英国 / 阿联酋分流等）。上述事实会随政策与数据修订变化，使用时请复核最新原文。",
        first_line=True,
    )
    add_para(
        doc,
        "本文是个人决策框架，不是移民法律意见，不是投资建议，也不是对任何国家前途的预言。它只处理一个问题：在 2026 年 8 月这个时点，对网络型、场景型、转化型从业者，把宝贵的三到五年放在哪一边，期望值更高。期望值更高，不等于没有痛苦；留下也不等于躺赢。留下的人必须换剧本。走的人必须带根走。",
        first_line=True,
    )
    add_para(
        doc,
        "重新生成本备忘录：在本分支目录执行 python3 scripts/build_all_stay_or_go.py。配套文件包括 PPT 汇报版与 Excel 决策表。",
        first_line=True,
        color=SLATE,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"已生成：{output_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    build_docx(root / "deliverables" / "发展与去留_2026中美判断备忘录.docx")
