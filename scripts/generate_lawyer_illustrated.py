#!/usr/bin/env python3
"""从《胡志远病例文件》抽出关键页，做成给律师的图文 Word / PDF。不是律师函。"""

from __future__ import annotations

from pathlib import Path

import numpy as np
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.colors import HexColor, white
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from comm_plan import CN_FONT_PATH
from content import QILU_CHART

CHINESE_FONT = "宋体"
HEADING_FONT = "黑体"
ACCENT = RGBColor(0x0B, 0x2F, 0x5B)
MUTED = RGBColor(0x5C, 0x6B, 0x7A)
RED = RGBColor(0xA6, 0x3D, 0x2F)

DOCX_NAME = "青岛抚顺路和哈尔滨路路口交通事故_给律师的病历图文_20260819.docx"
PDF_NAME = "青岛抚顺路和哈尔滨路路口交通事故_给律师的病历图文_20260819.pdf"

ROOT = Path(__file__).resolve().parents[1]
SRC_PDF = ROOT / "deliverables" / "胡志远病例文件.pdf"
PHOTO_DIR = ROOT / "deliverables" / "chart-photos"

# page 为 1-based。只摘律师必看页，不放出院准备清单、签到凭条。
FIGURES = [
    {
        "page": 4,
        "file": "01_住院腕带.jpg",
        "rotate": 90,
        "title": "图1　住院腕带（身份核对）",
        "why": "核对病历姓名、住院号、病区，避免把转写姓名写进对外文书。",
        "points": [
            f"姓名 {QILU_CHART['name']}，男，{QILU_CHART['age']}。",
            f"住院号 {QILU_CHART['inpatient_no']}。病区 {QILU_CHART['ward']}。",
            "过敏药物栏空白。腕带另有防跌倒黄标和药物过敏红标；红标与栏内空白、急诊过敏史「无」不一致，对外以急诊病历为准，不要写成有药物过敏。",
        ],
    },
    {
        "page": 9,
        "file": "02_急诊病历.jpg",
        "title": "图2　急诊骨科初诊病历（事故性质与就诊时间）",
        "why": "这是「车祸伤」的第一份临床记录，可框定事故时间窗。",
        "points": [
            f"就诊时间 {QILU_CHART['er_time']}，科室 {QILU_CHART['er_dept']}，门诊号 {QILU_CHART['outpatient_no']}。",
            f"主诉：{QILU_CHART['chief_er']}。",
            f"生命体征：{QILU_CHART['vitals']}。既往：{QILU_CHART['pmh']}。过敏史无。",
            "初步诊断含车祸伤、右小腿外伤、右踝关节骨折伴脱位。请手足外科赛佳明副主任医师会诊，建议手术并开住院证。",
            QILU_CHART["time_window"],
        ],
    },
    {
        "page": 2,
        "file": "03_破伤风处方.jpg",
        "title": "图3　急诊处方：破伤风人免疫球蛋白",
        "why": "已发生费用的一张票面依据，理赔台账应贴原件。",
        "points": [
            QILU_CHART["tetanus"],
            "临床诊断栏只写「肢体疼痛」，伤情以 CT 和出院记录为准，不要用这四个字对外概括伤情。",
        ],
    },
    {
        "page": 1,
        "file": "04_右小腿CT.jpg",
        "title": "图4　右小腿 CT（检查号 10008056847）——本次事故核心影像",
        "why": "骨伤的影像依据。报告页注明仅供临床参考，不是鉴定意见。",
        "points": [
            "检查时间 2026-08-14 13:02。申请科室急诊骨科。报告/审核医师马驰。",
            "诊断：右侧胫骨远端、腓骨近端骨折并踝关节半脱位；右距骨内缘、外踝撕脱性骨折；软组织肿胀。",
        ],
    },
    {
        "page": 3,
        "file": "05_右足CT.jpg",
        "title": "图5　右足 CT（检查号 10008056848）——不计入本次事故",
        "why": "用来挡住「跟骨粉碎性骨折」的口述。不要写进索赔和评残。",
        "points": [
            "右足组成骨完整，诸关节在位。诊断为退行性变、跟骨骨刺、第一跖骨头囊变。",
            "高血压、糖尿病同样是既往病，出院诊断里有，但不按事故伤向对方索赔。",
        ],
    },
    {
        "page": 8,
        "file": "06_出院记录.jpg",
        "title": "图6　出院记录（最重要：开放伤、手术时间、术式）",
        "why": "比 CT 更重。手术日是 8 月 14 日 19:01，不是 8 月 15 日。",
        "points": [
            f"入院 {QILU_CHART['admit']}，出院 {QILU_CHART['discharge']}，住院 {QILU_CHART['days']}。",
            "出院诊断：右开放性踝关节骨折、右开放性踝关节脱位、右踝部韧带断裂、右踝挫伤并皮肤坏死。后两项既往病不计入事故伤残。",
            f"手术：{QILU_CHART['surgery_time']} {QILU_CHART['surgery']}。",
            f"出院时：{QILU_CHART['out_status']}。{QILU_CHART['fix']}。",
            "带药：欧开 BID、利伐沙班 QD。医嘱继续负压吸引，根据创面二期治疗。通话里的两根/四根钢钉停用。",
            QILU_CHART["signer"],
        ],
    },
    {
        "page": 7,
        "file": "07_术后踝CT.jpg",
        "title": "图7　术后右踝 CT（检查号 10008059257）",
        "why": "证明手术次日固定在位，人还在齐鲁 33 床，不是「已经出院回家」。",
        "points": [
            "检查 2026-08-15 16:26，报告医师管帅。申请科室手足与显微重建外科，床号 33。",
            "影像所见：右踝关节骨折内外固定术后复查，内外固定在位，位线可，周围软组织肿胀。",
        ],
    },
]


def _set_run_font(run, font_name: str, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _p(
    doc: Document,
    text: str = "",
    *,
    font: str = CHINESE_FONT,
    size: float = 12,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    space_after: float = 6,
    space_before: float = 0,
    color: RGBColor | None = None,
    line_spacing: float = 1.35,
):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold, color=color)
    return p


def _crop_ink(im: Image.Image, pad: int = 22, border: int = 36, thresh: int = 80) -> Image.Image:
    arr = np.array(im.convert("L"))
    h, w = arr.shape
    inner = arr[border : h - border, border : w - border]
    ys, xs = np.where(inner < thresh)
    if len(ys) < 800:
        return im
    y0, y1 = int(ys.min() + border), int(ys.max() + border)
    x0, x1 = int(xs.min() + border), int(xs.max() + border)
    y0, x0 = max(0, y0 - pad), max(0, x0 - pad)
    y1, x1 = min(h - 1, y1 + pad), min(w - 1, x1 + pad)
    return im.crop((x0, y0, x1 + 1, y1 + 1))


def extract_photos(src_pdf: Path = SRC_PDF, out_dir: Path = PHOTO_DIR) -> list[Path]:
    import pypdfium2 as pdfium

    if not src_pdf.exists():
        raise FileNotFoundError(src_pdf)
    out_dir.mkdir(parents=True, exist_ok=True)
    pdf = pdfium.PdfDocument(str(src_pdf))
    written: list[Path] = []
    for fig in FIGURES:
        page = pdf[fig["page"] - 1]
        im = page.render(scale=2.4).to_pil().convert("RGB")
        im = _crop_ink(im)
        if fig.get("rotate"):
            im = im.rotate(int(fig["rotate"]), expand=True)
        dest = out_dir / fig["file"]
        im.save(dest, format="JPEG", quality=86, optimize=True)
        written.append(dest)
    return written


def _add_picture(doc: Document, path: Path, max_width_cm: float = 15.4, max_height_cm: float = 18.5) -> None:
    im = Image.open(path)
    w, h = im.size
    width_cm = max_width_cm
    height_cm = width_cm * h / w
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * w / h
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm), height=Cm(height_cm))


def build_document(output_path: Path, photo_dir: Path = PHOTO_DIR) -> None:
    photos = extract_photos(out_dir=photo_dir)
    by_name = {p.name: p for p in photos}

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)

    _p(doc, "交通事故病历图文摘录（供受托律师阅卷）", font=HEADING_FONT, size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=4)
    _p(doc, "青岛市市北区抚顺路和哈尔滨路路口　·　货运三轮 × 美团女骑手　·　伤者胡志远", font=HEADING_FONT, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, space_after=4)
    _p(doc, "本文只给受托律师看原件长什么样。不是律师函，不要发给骑手、站点、保险公司或刘孝春。现阶段不谈总赔偿数额、不报残级。", size=11, color=RED, space_after=10)

    _p(doc, "律师先看这几句", font=HEADING_FONT, size=14, bold=True, color=ACCENT, space_before=4, space_after=6)
    lead = [
        "病历姓名胡志远。门诊号 0002750430，住院号 0001519455，11D 手足与显微重建外科 33 床。",
        "急诊 2026-08-14 12:27 主诉「车祸伤致右踝疼痛肿胀2小时」。入院 14:14，出院 8 月 17 日 08:00，住院 3 天。",
        "手术是 8 月 14 日 19:01，不是 8 月 15 日。8 月 15 日是术后 CT。",
        "出院诊断比 CT 更重：开放性踝关节骨折脱位、韧带断裂、皮肤坏死；已清创并 VSD。固定写外固定架及克氏针。",
        "右足 CT 是骨刺和退变，不计入本次事故。高血压、糖尿病是既往病。",
        "书面事故认定书尚未出具。地点对外写抚顺路和哈尔滨路路口。完整经过说明是另一份 Word/PDF。",
    ]
    for x in lead:
        _p(doc, "• " + x, size=12, space_after=3)

    _p(doc, "以下图片从《胡志远病例文件.pdf》10 页扫描件中抽出，已裁掉空白边。原件请与本稿对照，不要二次拍照传真导致更糊。", size=10.5, color=MUTED, space_before=8, space_after=12)

    for fig in FIGURES:
        path = by_name[fig["file"]]
        _p(doc, fig["title"], font=HEADING_FONT, size=13, bold=True, color=ACCENT, space_before=12, space_after=4)
        _p(doc, "用途：" + fig["why"], size=11, color=RED, space_after=4)
        for x in fig["points"]:
            _p(doc, "• " + x, size=11.5, space_after=2)
        max_h = 19.5 if fig["page"] == 8 else 17.2
        _add_picture(doc, path, max_height_cm=max_h)
        _p(doc, f"原件页码：病例文件第 {fig['page']} 页。文件名 {fig['file']}。", size=9, color=MUTED, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    _p(doc, "仍请律师指导尽快补的", font=HEADING_FONT, size=14, bold=True, color=ACCENT, space_before=8)
    for x in (
        "完整手术记录、麻醉记录、植入物编码（出院记录已有术式，仍缺根数）",
        "齐鲁出院结算票据、费用清单",
        "人民医院住院号",
        "病历正式复印：出院后 10 个工作日，博施楼 1 楼服务中心，带患者身份证",
        "事故认定书、监控复制件",
    ):
        _p(doc, "• " + x, size=12, space_after=3)

    _p(
        doc,
        "重新生成：python3 scripts/generate_lawyer_illustrated.py。3D 示范动画留作待用，不是证据。",
        size=10,
        color=MUTED,
        space_before=12,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _pdf_font() -> str:
    kwargs = {"subfontIndex": 0} if CN_FONT_PATH.endswith(".ttc") else {}
    pdfmetrics.registerFont(TTFont("CN", CN_FONT_PATH, **kwargs))
    return "CN"


def _header_footer(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(HexColor("#0B2F5B"))
    canvas.rect(0, A4[1] - 11 * mm, A4[0], 11 * mm, fill=1, stroke=0)
    canvas.setFillColor(white)
    canvas.setFont("CN", 8)
    canvas.drawString(16 * mm, A4[1] - 7 * mm, "青岛抚顺路和哈尔滨路路口交通事故 · 给律师的病历图文 · 内部阅卷")
    canvas.setFillColor(HexColor("#C4A35A"))
    canvas.rect(0, 0, A4[0], 9 * mm, fill=1, stroke=0)
    canvas.setFillColor(HexColor("#0B2F5B"))
    canvas.setFont("CN", 8)
    canvas.drawString(16 * mm, 3.5 * mm, "不是律师函 · 不要发给骑手或刘孝春")
    canvas.drawRightString(A4[0] - 16 * mm, 3.5 * mm, f"第 {doc.page} 页")
    canvas.restoreState()


def _rl_image(path: Path, max_w=170, max_h=175) -> RLImage:
    im = Image.open(path)
    w, h = im.size
    width = max_w * mm
    height = width * h / w
    if height > max_h * mm:
        height = max_h * mm
        width = height * w / h
    return RLImage(str(path), width=width, height=height)


def build_pdf(output_path: Path, photo_dir: Path = PHOTO_DIR) -> None:
    photos = list(photo_dir.glob("*.jpg"))
    if len(photos) < len(FIGURES):
        extract_photos(out_dir=photo_dir)
    by_name = {p.name: p for p in photo_dir.glob("*.jpg")}
    font = _pdf_font()
    navy = HexColor("#0B2F5B")
    muted = HexColor("#5C6B7A")
    red = HexColor("#A63D2F")
    dark = HexColor("#1F2A37")
    styles = {
        "title": ParagraphStyle("t", fontName=font, fontSize=16, leading=24, textColor=navy, alignment=TA_CENTER, spaceAfter=4),
        "sub": ParagraphStyle("s", fontName=font, fontSize=10, leading=15, textColor=muted, alignment=TA_CENTER, spaceAfter=6),
        "warn": ParagraphStyle("w", fontName=font, fontSize=10, leading=15, textColor=red, alignment=TA_LEFT, spaceAfter=8),
        "h1": ParagraphStyle("h", fontName=font, fontSize=12.5, leading=19, textColor=navy, spaceBefore=8, spaceAfter=4),
        "body": ParagraphStyle("b", fontName=font, fontSize=10.5, leading=16, textColor=dark, alignment=TA_JUSTIFY, spaceAfter=3),
        "cap": ParagraphStyle("c", fontName=font, fontSize=8, leading=12, textColor=muted, alignment=TA_CENTER, spaceAfter=8),
    }
    story = [
        Paragraph("交通事故病历图文摘录（供受托律师阅卷）", styles["title"]),
        Paragraph("抚顺路和哈尔滨路路口 · 伤者胡志远 · 截至 2026-08-19", styles["sub"]),
        Paragraph("不是律师函。不要发给骑手、站点、保险公司或刘孝春。现阶段不谈总赔偿数额。", styles["warn"]),
        Paragraph("律师先看：病历姓名胡志远，住院号 0001519455。手术 2026-08-14 19:01。出院诊断为开放性踝关节骨折脱位、韧带断裂、皮肤坏死，已 VSD。8 月 15 日是术后 CT。右足骨刺不计入本次事故。认定书尚未出具。", styles["body"]),
    ]
    for fig in FIGURES:
        path = by_name[fig["file"]]
        if fig["page"] == 8:
            story.append(PageBreak())
        story.append(Paragraph(fig["title"], styles["h1"]))
        story.append(Paragraph("用途：" + fig["why"], styles["warn"]))
        for x in fig["points"]:
            story.append(Paragraph("• " + x, styles["body"]))
        max_h = 185 if fig["page"] == 8 else 160
        story.append(Spacer(1, 4))
        story.append(_rl_image(path, max_h=max_h))
        story.append(Paragraph(f"原件：病例文件第 {fig['page']} 页", styles["cap"]))
        if fig["page"] in (9, 8, 1):
            story.append(PageBreak())
    story.append(Paragraph("仍缺：完整手术记录和植入物编码、齐鲁结算票据、人民医院住院号、认定书。病历复印出院后 10 个工作日到博施楼 1 楼。", styles["body"]))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=16 * mm,
        bottomMargin=14 * mm,
        title="给律师的病历图文摘录",
        author="内部整理",
    )
    doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)


def main() -> None:
    out = ROOT / "deliverables"
    extract_photos()
    build_document(out / DOCX_NAME)
    build_pdf(out / PDF_NAME)
    print(f"Wrote {out / DOCX_NAME}")
    print(f"Wrote {out / PDF_NAME}")


if __name__ == "__main__":
    main()
