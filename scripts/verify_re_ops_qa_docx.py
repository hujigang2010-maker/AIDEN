#!/usr/bin/env python3
"""校验《房地产开发运营一线答疑》Word 文档结构与关键口径。"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "deliverables" / "房地产开发运营一线答疑.docx"

REQUIRED_HEADING_FRAGMENTS = [
    "房地产开发运营一线答疑",
    "先看结论",
    "第一部分  运营",
    "项目公司",
    "现金流",
    "缴税",
    "表外项目",
    "第二部分  拿地与销售",
    "拿地测算",
    "销售速度与降价",
    "第三部分  公司点评",
    "绿城",
    "华润置地",
    "中海地产",
    "建发",
    "中国金茂",
    "滨江",
    "越秀",
    "保利发展",
    "招商蛇口",
    "附录",
]

REQUIRED_PHRASES = [
    "有限责任",
    "预售资金监管",
    "不得用于购置土地",
    "土地出让金",
    "预计计税毛利率",
    "预征率",
    "多退少补",
    "集团合并纳税",
    "股东借款",
    "合营",
    "联营",
    "无息",
    "净利率",
    "去化",
    "跑冒滴漏",
]


def main() -> int:
    if not DOCX.exists():
        print(f"文件不存在：{DOCX}", file=sys.stderr)
        return 1

    doc = Document(str(DOCX))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    blob = "\n".join(texts)
    tables = doc.tables

    errors: list[str] = []
    if len(texts) < 80:
        errors.append(f"段落过少：{len(texts)}")
    if len(blob) < 12000:
        errors.append(f"正文过短：{len(blob)} 字")
    if len(tables) < 8:
        errors.append(f"表格过少：{len(tables)}")

    for frag in REQUIRED_HEADING_FRAGMENTS:
        if frag not in blob:
            errors.append(f"缺少标题或章节：{frag}")
    for phrase in REQUIRED_PHRASES:
        if phrase not in blob:
            errors.append(f"缺少关键口径：{phrase}")

    # 封面标题应存在
    if texts[0] != "内部讨论稿  ·  一线投资运营口径" and "房地产开发运营一线答疑" not in texts[:5]:
        if "房地产开发运营一线答疑" not in texts[:8]:
            errors.append("封面标题未出现在文首")

    size_kb = DOCX.stat().st_size / 1024
    print(f"文件：{DOCX}")
    print(f"大小：{size_kb:.1f} KB")
    print(f"段落：{len(texts)}")
    print(f"正文字数：{len(blob)}")
    print(f"表格：{len(tables)}")
    print(f"前5段：{texts[:5]}")

    if errors:
        print("校验失败：")
        for e in errors:
            print(f"  - {e}")
        return 1

    print("校验通过。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
