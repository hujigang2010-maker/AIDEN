#!/usr/bin/env python3
"""生成发给合作方的《CGC领事会客厅合作方案修改建议》。"""

from __future__ import annotations

import html as html_lib
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output"
DOCS_DIR = ROOT / "docs"
# 发给对方的稿使用微软雅黑：Windows 可正常显示；跨平台 PDF 请嵌入字体后导出。
FONT_CN = "微软雅黑"
FONT_EN = "Calibri"
NAVY = RGBColor(0x1B, 0x3A, 0x5F)
ORANGE = RGBColor(0xB8, 0x5C, 0x00)
GREEN = RGBColor(0x1F, 0x6B, 0x3A)
GRAY = RGBColor(0x4A, 0x4A, 0x4A)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

MD_LINES: list[str] = []


def set_run_font(run, size=11, bold=False, color=BLACK, font=FONT_CN):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = FONT_EN
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:cs"), FONT_EN)


def shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell, color="D0D5DD") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, *, size=9, bold=False, color=BLACK, align="left", fill=None):
    if fill:
        shade_cell(cell, fill)
    set_cell_border(cell)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_paragraph(doc, text, *, size=11, bold=False, color=BLACK, space_after=8, first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.35
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading_cn(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0)
    size = 16 if level == 1 else 13
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=NAVY)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "1B3A5F")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_bullet(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run("• ")
    set_run_font(run, size=11, bold=True, color=NAVY)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True, color=BLACK)
        r2 = p.add_run(text)
        set_run_font(r2, size=11, color=BLACK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=BLACK)
    return p


def add_callout(doc, title, body, fill="FFF6E8", border="E8C48A", title_color=ORANGE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, border)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(title)
    set_run_font(r1, size=12, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.line_spacing = 1.3
    r2 = p2.add_run(body)
    set_run_font(r2, size=11, color=BLACK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, size=9, bold=True, color=WHITE, align="center", fill="1B3A5F")
    for r_idx, row in enumerate(rows, start=1):
        fill = "F7F9FC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            set_cell_text(table.cell(r_idx, c_idx), val, size=9, align="left", fill=fill)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, size=9, color=GRAY)


def md(text: str = "") -> None:
    MD_LINES.append(text)


def md_table(headers, rows) -> None:
    md("| " + " | ".join(headers) + " |")
    md("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in rows:
        md("| " + " | ".join(str(x).replace("\n", "<br>") for x in row) + " |")
    md()


def markdown_to_html(md_text: str) -> str:
    def inline(s: str) -> str:
        s = html_lib.escape(s)
        s = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", s)
        return s.replace("&lt;br&gt;", "<br>")

    lines = md_text.splitlines()
    body: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# "):
            body.append(f"<h1>{inline(line[2:])}</h1>")
        elif line.startswith("## "):
            body.append(f"<h2>{inline(line[3:])}</h2>")
        elif line.startswith("### "):
            body.append(f"<h3>{inline(line[4:])}</h3>")
        elif line.startswith("> "):
            body.append(f"<blockquote>{inline(line[2:])}</blockquote>")
        elif line.startswith("| ") and i + 1 < len(lines) and lines[i + 1].startswith("| ---"):
            rows = []
            while i < len(lines) and lines[i].startswith("| "):
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            i -= 1
            header, data = rows[0], rows[2:]
            th = "".join(f"<th>{inline(c)}</th>" for c in header)
            trs = [f"<tr>{''.join(f'<td>{inline(c)}</td>' for c in row)}</tr>" for row in data]
            body.append(f"<table><thead><tr>{th}</tr></thead><tbody>{''.join(trs)}</tbody></table>")
        elif line.startswith("- "):
            items = []
            while i < len(lines) and lines[i].startswith("- "):
                items.append(f"<li>{inline(lines[i][2:])}</li>")
                i += 1
            i -= 1
            body.append("<ul>" + "".join(items) + "</ul>")
        elif line.strip():
            body.append(f"<p>{inline(line)}</p>")
        i += 1
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>CGC领事会客厅合作方案修改建议</title>
<style>
body {{ font-family: "WenQuanYi Micro Hei", "微软雅黑", sans-serif; color:#1a1a1a;
  max-width: 980px; margin: 32px auto 80px; padding: 0 28px; line-height: 1.65; background:#f4f1ea; }}
.sheet {{ background:#fff; padding: 40px 48px 56px; box-shadow: 0 8px 28px rgba(27,58,95,.08); }}
.kicker {{ color:#6b7280; font-size:12px; text-align:right; }}
h1 {{ color:#1B3A5F; font-size:26px; margin: 8px 0 6px; text-align:center; }}
blockquote {{ background:#FFF6E8; border-left:4px solid #B85C00; padding:12px 16px; margin:16px 0 24px; }}
h2 {{ color:#1B3A5F; border-bottom:2px solid #1B3A5F; padding-bottom:6px; margin-top:32px; font-size:19px; }}
h3 {{ color:#1B3A5F; font-size:15px; margin-top:20px; }}
p {{ font-size:15px; }}
table {{ border-collapse: collapse; width:100%; margin: 10px 0 22px; font-size:13px; }}
th {{ background:#1B3A5F; color:#fff; padding:8px; }}
td {{ border:1px solid #D0D5DD; padding:8px; vertical-align:top; }}
tr:nth-child(even) td {{ background:#F7F9FC; }}
ul {{ padding-left: 1.2em; }}
li {{ margin: 6px 0; }}
strong {{ color:#1B3A5F; }}
</style>
</head>
<body>
<div class="sheet">
<div class="kicker">合作沟通稿 · 供下一版修订 · 2026-09-05</div>
{chr(10).join(body)}
</div>
</body></html>
"""


def build() -> tuple[Path, Path, Path]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    MD_LINES.clear()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("CGC领事会客厅合作方案修改建议  ·  2026年9月5日")
    set_run_font(hr, size=8, color=GRAY)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("供合作沟通与下一版修订使用，不构成合同或承诺")
    set_run_font(fr, size=8, color=GRAY)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(4)
    r = t.add_run("CGC领事会客厅合作方案修改建议")
    set_run_font(r, size=20, bold=True, color=NAVY)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(12)
    sr = st.add_run("森马上海国际运营中心  ·  创智汇  ·  尚9·一滴水")
    set_run_font(sr, size=11, color=GRAY)

    md("# CGC领事会客厅合作方案修改建议")
    md()
    md("**森马上海国际运营中心 · 创智汇 · 尚9·一滴水**")
    md()
    md("> 供合作沟通与下一版修订使用，不构成合同或承诺。2026年9月5日。")
    md()

    add_heading_cn(doc, "一、总体反馈")
    md("## 一、总体反馈")
    md()
    add_paragraph(
        doc,
        "感谢提供森马、创智汇及尚9·一滴水三份合作方案。整体合作思路有吸引力，也结合不同场景做了定位："
        "CGC提供领事、海外商协会和企业网络，合作方提供场地、企业组织或餐饮场景；"
        "免费沙龙解决基础信息需求，深度出海服务再市场化收费。"
        "三个场景的差异化方向也已成形——森马侧重潮玩与IP，创智汇侧重数字科创，一滴水侧重中阿商务宴请。",
    )
    md("感谢提供三份合作方案。整体合作思路有吸引力，场景定位清楚：CGC提供网络，合作方提供场地与场景；免费沙龙解决基础信息，深度服务再收费。")
    md()
    add_callout(
        doc,
        "下一版请把方案写成可评估、可验收的实施方案",
        "当前文本更接近资源合作提案，还不是完整策划案。建议在补齐双方投入、具体交付、验收方式和预算之前，"
        "暂不作为签约、拨款或启动装修的依据。下一版重点把“谁投入什么、交付什么、如何验收”写实，方便各方内部评估和推进。",
    )
    md("**下一版请把方案写成可评估、可验收的实施方案。** 当前更接近资源合作提案。补齐投入、交付、验收和预算之前，暂不作为签约、拨款或启动装修的依据。")
    md()

    add_caption(doc, "表1  建议的合作定位（供下一版按场景分别撰写，不必三案齐推）")
    md("**表1  建议的合作定位**")
    md()
    headers = ["方案", "建议定位", "建议节奏"]
    rows = [
        ["创智汇", "最适合作为长期产业旗舰", "先借用现有会议空间试点；若需专属点位，面积上限约 300㎡，不再按 1000㎡ 报批"],
        ["森马", "适合潮玩、IP垂直基地，但专属场地条件尚未成熟", "先用现有或可交付公共空间试运营，再决定是否交付 300—500㎡"],
        ["一滴水", "最适合低成本活动型试点", "先做 3—6 个月、3—5 场按场次合作，挂牌费分期并与履约挂钩"],
    ]
    add_table(doc, headers, rows, col_widths=[2.8, 5.6, 8.2])
    md_table(headers, rows)

    add_heading_cn(doc, "二、先分清：CGC 的权限价值，不是驻沪总领馆")
    md("## 二、先分清：CGC 的权限价值，不是驻沪总领馆")
    md()
    add_paragraph(
        doc,
        "下一版请把“上海总领事俱乐部（CGC）”和“各国驻上海总领事馆”分开写，不要让合作方误读成官方领事机构入驻。"
        "上海没有一个叫“上海总领事馆”的统一机构，而是七十余国在沪设立的驻沪总领馆；"
        "CGC 是民间俱乐部及国内运营主体承办的交流平台，可以邀请领事参加活动，但不能行使领馆职权，也不能代表任何国家或领馆发言。",
    )
    md("请把 CGC 与各国驻沪总领馆分开写。上海没有统一的“上海总领事馆”；CGC 是民间平台，不能行使领馆职权。")
    md()
    add_caption(doc, "表1-1  CGC 权限价值与驻沪总领馆的区别（请写入方案和合同口径）")
    md("**表1-1  CGC 权限价值与驻沪总领馆的区别**")
    md()
    headers = ["对照项", "各国驻沪总领馆", "CGC（上海总领事俱乐部）"]
    rows = [
        [
            "法律地位",
            "外国驻华官方领事机构，依据《维也纳领事关系公约》及中国法律履行领事职务",
            "民间国际交流平台，由国内运营主体承办；不是外交机构，也不是领馆内设部门",
        ],
        [
            "权限范围",
            "领事保护、签证与证件、官方经贸文化联络等法定职权，对派遣国负责",
            "无权办签证、领事认证、照会或审批；不能代表任何国家、领馆或政府作出承诺",
        ],
        [
            "对合作方的真实价值",
            "官方渠道和领区事务，须走外办及馆方程序，不因挂牌而自动获得",
            "圈层对接、活动组织、非正式介绍和宴会/沙龙场景。价值在“请得到人、办得成场”，不在“有审批权”",
        ],
        [
            "官员出席",
            "按外事纪律和馆方安排，出席不等于对商业项目背书",
            "可邀请领事或商务官员作为嘉宾，须一事一请、书面确认；方案不得保证到场或常态化坐镇",
        ],
        [
            "对外表述",
            "使用馆名、国徽、职衔须经馆方同意，并符合外事口径",
            "只可写“CGC品牌授权 / 民间交流基地挂牌”；禁止“官方授权”“总领馆办公室”“领事馆驻点”",
        ],
        [
            "方案中应删除的承诺",
            "不得把馆方尚未确认的事项写成已落实",
            "不得写“对接各国中央政府”“解读并降低法律风险”“保证项目审批或成交”",
        ],
    ]
    add_table(doc, headers, rows, col_widths=[2.8, 7.0, 6.8])
    md_table(headers, rows)
    add_paragraph(
        doc,
        "因此，挂牌费和对接服务购买的是 CGC 的组织能力和人脉转介，不是驻沪总领馆的官方权限。"
        "领事是否出席、以何种身份出席，只能按场确认；转化结果仍取决于企业和项目本身。",
    )
    md("挂牌费购买的是 CGC 的组织能力和人脉转介，不是驻沪总领馆的官方权限。领事出席须按场确认。")
    md()

    add_heading_cn(doc, "三、三份方案请统一补齐的内容")
    md("## 三、三份方案请统一补齐的内容")
    md()

    add_heading_cn(doc, "（一）合作主体、授权与可核验业绩", 2)
    md("### （一）合作主体、授权与可核验业绩")
    md()
    add_paragraph(
        doc,
        "请补充 CGC 签约及运营主体材料，便于合作方法务和财务内部评估。活动存在可以说明平台有过公开痕迹，"
        "但不能替代法律主体、履约能力和转化案例的尽调。请勿把“认识领事”直接等同于可承诺的商业资源。",
    )
    md("请补充签约主体材料。活动存在不能替代法律主体、履约能力和转化案例尽调。请勿把“认识领事”等同于可承诺的商业资源。")
    md()
    add_bullet(doc, "签约法人全称、统一社会信用代码、开票及收款主体。", bold_prefix="主体：")
    add_bullet(doc, "“CGC / 领事会客厅”品牌权属证明及对合作方的授权范围、期限和终止后的拆除义务。", bold_prefix="授权：")
    add_bullet(doc, "专职团队名单、岗位分工和对接人。", bold_prefix="团队：")
    add_bullet(doc, "近三年可核验活动清单（时间、主题、嘉宾层级、参会规模、形成的书面成果）。", bold_prefix="业绩：")
    add_bullet(doc, "领事或领馆商务人员实际出席的证明材料；客户案例及项目转化数据（线索、复谈、MOU 或成交，可脱敏）。", bold_prefix="转化：")
    md("- **主体：** 法人全称、统一社会信用代码、开票及收款主体。")
    md("- **授权：** 品牌权属及授权范围、期限、终止后拆除义务。")
    md("- **团队：** 专职名单、分工和对接人。")
    md("- **业绩：** 近三年可核验活动清单。")
    md("- **转化：** 领事实际出席证明；脱敏后的线索、复谈、MOU 或成交。")
    md()

    add_heading_cn(doc, "（二）年度交付、职责与可验收指标", 2)
    md("### （二）年度交付、职责与可验收指标")
    md()
    add_paragraph(
        doc,
        "现稿对合作方投入写得很重（免租空间、全额装修、公用费用、每年 10 万元及额外活动预算，或餐厅挂牌费加免费场地），"
        "对 CGC 的表述多为“开展、协助、导入、促成、优先落地”，缺少最低活动数、嘉宾标准、有效线索数和转化结果。"
        "请明确年度活动场次、服务企业数量、有效项目线索认定标准、会后跟进机制，以及双方项目负责人和职责分工，"
        "使场地、装修和经费投入有对应的交付和验收标准。",
    )
    md("现稿合作方投入很重，CGC 承诺多为不可验收表述。请补场次、企业数、线索标准、跟进机制和双方负责人。")
    md()
    add_caption(doc, "表2  建议写入下一版的可验收指标（可按场景微调，但须有数字、定义和统计口径）")
    md("**表2  建议写入下一版的可验收指标**")
    md()
    headers = ["指标", "建议写法", "验收材料"]
    rows = [
        ["对公活动场次", "试点期和年度分别列出最低场次", "活动通知、签到、纪要、现场照片"],
        ["嘉宾标准", "领事/领馆商务人员、商协会负责人、企业决策人的场次或比例", "嘉宾名单及身份说明（可脱敏）"],
        ["服务企业数量", "公益场覆盖家次；深度服务另计", "报名表、参会企业清单"],
        ["有效线索", "有主体、有对接意向、有跟进人的条目数", "线索台账，定期抄送合作方"],
        ["复谈 / MOU / 成交", "会后 30 日内复谈数；书面意向或成交另列", "跟进记录、意向书（可脱敏）"],
        ["招商或项目导入", "优先导流的规则、条数和权属", "月报；线索去向书面同步"],
        ["宴会或场地增量", "一滴水可写包场场次和毛利增量；园区可写公共空间利用率", "结算单、档期表"],
        ["满意度与合规", "参会满意度；零重大合规、舆情或安全事故", "问卷、合规检查记录"],
    ]
    add_table(doc, headers, rows, col_widths=[3.4, 6.8, 6.4])
    md_table(headers, rows)

    add_heading_cn(doc, "（三）预算、费用承担与挂牌费内涵", 2)
    md("### （三）预算、费用承担与挂牌费内涵")
    md()
    add_paragraph(
        doc,
        "每年 10 万元挂牌费只是很小一部分。请补齐整体财务安排，至少给出成本上限和承担主体，便于对方过会。"
        "建议按以下口径测算三年总成本，并分项列出上限、付款节点和采购机制：",
    )
    md("10 万元挂牌费只是很小一部分。请按三年总成本口径补齐上限、付款节点和采购机制。")
    md()
    add_callout(
        doc,
        "建议采用的成本口径",
        "三年总成本 ＝ 场地机会成本 ＋ 装修及设备 ＋ 物业能耗网络 ＋ 活动预算 ＋ 人员、安保、保险 ＋ 挂牌费用。"
        "请同时说明：挂牌费包含哪些服务、不包含哪些服务；活动费用如何审批和结算；"
        "是否有单场活动成本标准；收入预测或至少保本所需的最低场次。",
        fill="EAF6EE",
        border="A9CDB6",
        title_color=GREEN,
    )
    md("**成本口径：** 三年总成本＝场地机会成本＋装修设备＋物业能耗＋活动预算＋人员安保保险＋挂牌费用。并说明挂牌费内涵、活动费用结算和单场成本。")
    md()

    add_heading_cn(doc, "（四）免费服务与后续收费业务的边界", 2)
    md("### （四）免费服务与后续收费业务的边界")
    md()
    add_paragraph(
        doc,
        "分层服务可以保留，但须避免形成“合作方出资做公益活动—CGC 获得线索—再向企业收费”且权责不清的路径。请在下一版写明：",
    )
    md("分层服务可保留，但须写清公益获客与收费业务的隔离。")
    md()
    add_bullet(doc, "企业自愿选择付费服务，公益场不捆绑销售。", bold_prefix="自愿：")
    add_bullet(doc, "收费项目公开价目或报价原则，与补贴活动分开合同、分开开票、分开账目。", bold_prefix="公开与分账：")
    add_bullet(doc, "客户名单和活动数据的使用需经授权，合作方对属地企业数据的知情和限制。", bold_prefix="数据：")
    add_bullet(doc, "投诉受理主体；深度服务失误是否由 CGC 独立承担责任。", bold_prefix="投诉：")
    add_bullet(doc, "合作方是否参与市场化收益；如不参与，须写明不分成、也不为收费项目背书。", bold_prefix="收益：")
    md("- **自愿：** 公益场不捆绑销售。")
    md("- **公开与分账：** 分开合同、开票和账目。")
    md("- **数据：** 名单使用需授权。")
    md("- **投诉：** 明确受理和责任主体。")
    md("- **收益：** 写明合作方是否分成。")
    md()

    add_heading_cn(doc, "（五）涉外表述、慈善资金与合同底线", 2)
    md("### （五）涉外表述、慈善资金与合同底线")
    md()
    add_paragraph(
        doc,
        "“民间交流、不承担官方外交职能”需要落到合同条款，并与上文表1-1保持一致。请统一对外口径：",
    )
    md("民间定位须写入合同，不能只出现在方案开头。")
    md()
    add_bullet(doc, "不代表任何领馆或政府部门背书；不保证官员出席、项目审批或商业成交。", bold_prefix="不背书、不保证：")
    add_bullet(doc, "一律改为“CGC品牌授权 / 民间交流基地挂牌”，删除“官方授权”“官方外交机构”等表述。", bold_prefix="品牌用语：")
    add_bullet(doc, "涉及公开义拍、慈善晚会募捐的，须由取得公开募捐资格的慈善组织开展或与其正式合作，募集资金进入该组织账户，并履行备案、公开等程序。", bold_prefix="慈善：")
    add_bullet(doc, "补充反商业贿赂、个人信息与数据、肖像与宣传、IP、安保保险及退出后的标识拆除、资料交接。", bold_prefix="其他条款：")
    md("- **不背书、不保证：** 不代表领馆或政府；不保证官员出席、审批或成交。")
    md("- **品牌用语：** 统一为 CGC 品牌授权，删除官方授权。")
    md("- **慈善：** 公开募捐须由有资格的慈善组织开展或正式合作，资金入其账户并备案公开。")
    md("- **其他条款：** 反商业贿赂、数据、肖像、IP、安保保险、退出交接。")
    md()

    add_heading_cn(doc, "四、分方案修改意见")
    md("## 四、分方案修改意见")
    md()

    add_heading_cn(doc, "（一）森马方案", 2)
    md("### （一）森马方案")
    md()
    add_paragraph(
        doc,
        "潮玩、消费品牌和 IP 出海与森马方向有协同，适合作为垂直基地来写，但当前专属场地条件尚未成熟，风险也最大。请重点修改：",
    )
    md("潮玩与 IP 方向有协同，但专属场地条件尚未成熟。")
    md()
    add_bullet(doc, "正文“2000㎡免租专属场地”与后文“1000㎡专属空间”不一致，请统一，并附功能分区面积表。该数字直接影响投资审批。", bold_prefix="面积：")
    add_bullet(
        doc,
        "浙江森马服饰股份有限公司 2026 年半年度报告仍将“上海国际运营中心项目”列为在建工程，期末账面价值约 7.83 亿元，工程累计投入约占预算 52.19%。请核实并补充实际竣工、验收、可交付使用的空间范围和时间，不宜按已全面投入运营的园区来承诺大面积专属场地。",
        bold_prefix="交付条件：",
    )
    add_bullet(doc, "2000㎡、全额装修加三年合作明显超前。全文仅“每年不少于 3 场青年活动”较明确，请升格为完整年历和整体 KPI。", bold_prefix="投入节奏：")
    add_bullet(doc, "建议先用现有或届时可使用的公共空间做 6—12 个月试运营；达到利用率和转化门槛后，再讨论交付 300—500㎡，而不是一步到位。", bold_prefix="建议改法：")
    add_bullet(doc, "会客厅名称建议突出潮玩、IP 与消费出海，避免与创智汇同时主打“AI 会客厅”。", bold_prefix="主题：")
    md("- **面积：** 统一 2000㎡ / 1000㎡，并附分区表。")
    md("- **交付条件：** 2026 年半年报仍列为在建工程，累计投入约占预算 52.19%。请补充可交付空间和时间。")
    md("- **投入节奏：** 大面积免租加全额装修超前；3 场青年活动不足以对应三年投入。")
    md("- **建议改法：** 先公共空间试运营 6—12 个月，达标后再谈 300—500㎡。")
    md("- **主题：** 突出潮玩与 IP，避免与创智汇抢 AI 会客厅。")
    md()

    add_heading_cn(doc, "（二）创智汇方案", 2)
    md("### （二）创智汇方案")
    md()
    add_paragraph(
        doc,
        "三份之中，创智汇最适合作为长期产业旗舰来设计：大创智的文化科技方向、周边高校和园区场景，与数字科创国际对接更匹配。"
        "也正因涉及国资物业、装修、补助和公共产业服务，合规要求最高。专属场地请按最多约 300㎡ 来写，不再按 1000㎡ 设计和报批。请补充：",
    )
    md("创智汇最适合做长期产业旗舰，但国资合规要求最高。专属场地上限约 300㎡。")
    md()
    add_bullet(doc, "原文 1000㎡ 请下调。专属点位最多约 300㎡，用于小型展陈、洽谈和 20—50 人路演即可；高峰或大型活动继续使用园区现有会议室、路演厅，按次预约。", bold_prefix="面积上限：")
    add_bullet(doc, "国有物业使用、装修投资、补助或采购所适用的决策、审计和资产管理程序，以及合同主体（北岛科技 / 杨浦科创集团）。面积越小，过会越容易。", bold_prefix="国资程序：")
    add_bullet(doc, "CGC 最低交付量（场次、嘉宾层级、线索和月报）。", bold_prefix="对价：")
    add_bullet(doc, "园区公共投入与 CGC 向企业收费如何隔离，是否允许在补贴空间内销售深度服务。", bold_prefix="隔离：")
    add_bullet(doc, "三年退出或提前终止时，装修资产、数据、品牌物料和客户关系如何交接。", bold_prefix="退出：")
    add_bullet(doc, "与园内已有“AI+IP 产业创新中心”等载体的分工，避免重复建设观感。", bold_prefix="存量协同：")
    add_bullet(doc, "建议下一版改为：先借用现有会议空间试点；达标后再议不超过约 300㎡ 的专属点位，不立即按 1000㎡ 装修交付。", bold_prefix="建议改法：")
    md("- **面积上限：** 专属点位最多约 300㎡，不再按 1000㎡ 报批；大型活动用园区公共空间。")
    md("- **国资程序：** 决策、采购、审计、资产管理及合同主体。")
    md("- **对价：** CGC 最低交付量。")
    md("- **隔离：** 公共投入与收费业务。")
    md("- **退出：** 装修、数据、品牌、客户交接。")
    md("- **存量协同：** 与园内 AI+IP 中心分工。")
    md("- **建议改法：** 先借用会议空间试点，达标后再议不超过约 300㎡。")
    md()

    add_heading_cn(doc, "（三）尚9·一滴水方案", 2)
    md("### （三）尚9·一滴水方案")
    md()
    add_paragraph(
        doc,
        "轻资产、按场次合作，最容易低成本试错，也最容易形成直接宴会收入。请先改正事实和口径，再谈挂牌：",
    )
    md("轻资产最适合试点，但须先改正地址、授权口径和与餐厅主业无关的表述。")
    md()
    add_bullet(
        doc,
        "方案写“黄浦区外滩及陆家嘴滨江核心区域”。公开信息显示，尚9·一滴水江景中餐厅位于上海市虹口区东大名路500号（北外滩国际客运中心一带）。建议改为“上海市虹口区东大名路500号”，地理描述可用“临黄浦江，隔江眺望陆家嘴并可观外滩”。",
        bold_prefix="地址：",
    )
    add_bullet(doc, "一边写不承担官方外交职能，一边写“CGC官方授权”。请统一为“CGC品牌授权”。", bold_prefix="口径：")
    add_bullet(doc, "“反哺企业主在阿联酋的海外地产业务”与餐厅经营无直接关系，请从正式方案中删除，避免借涉外资源为私人业务背书的观感。", bold_prefix="表述：")
    add_bullet(doc, "“实际控制人长期深耕阿联酋”如无公开可核验资料，请改为待核实或删除，不作为合作依据。", bold_prefix="待核实：")
    add_bullet(doc, "只做 3—6 个月、3—5 场活动的按场次试点；采用临时联合标识；挂牌费分期支付并与履约挂钩。请同时附餐费/场地优惠表、档期预约和高峰占用规则。", bold_prefix="建议改法：")
    md("- **地址：** 改为虹口区东大名路 500 号；地理描述改为临江眺望陆家嘴、可观外滩。")
    md("- **口径：** “CGC官方授权”改为“CGC品牌授权”。")
    md("- **表述：** 删除迪拜地产业务赋能。")
    md("- **待核实：** 实控人长期深耕阿联酋如无依据则删除。")
    md("- **建议改法：** 3—6 个月、3—5 场按场次试点；临时标识；挂牌费分期挂钩履约。")
    md()

    add_heading_cn(doc, "五、建议的试运营安排")
    md("## 五、建议的试运营安排")
    md()
    add_paragraph(
        doc,
        "合作节奏上，建议先设置试运营，约定阶段目标、付款节点及未达标时的调整或退出方式，再确定长期合作和空间投入。",
    )
    md("先试运营，再确定长期合作和空间投入。")
    md()
    add_caption(doc, "表3  建议的试点节奏")
    md("**表3  建议的试点节奏**")
    md()
    headers = ["阶段", "动作", "通过后再做"]
    rows = [
        ["第 1—2 周", "签署活动合作备忘录（先不签物业无偿使用）；明确口径、档期和取消规则", "锁定试点场次"],
        ["试点期", "一滴水 3—5 场；园区各至少 1 场可核验闭门沙龙。嘉宾须可核验，不做“拟邀请”", "形成线索台账和复盘"],
        ["试点验收", "对照表2指标：出席层级、企业主体、纪要时限、对方书面意愿", "未达标则调整或退出，不启动装修"],
        ["空间条款", "创智汇专属点位上限约 300㎡；森马仍建议先公共空间、达标后再议", "12 个月空间试点 + 续约条件"],
    ]
    add_table(doc, headers, rows, col_widths=[2.8, 8.4, 5.4])
    md_table(headers, rows)

    add_heading_cn(doc, "六、文档规范")
    md("## 六、文档规范")
    md()
    add_paragraph(
        doc,
        "森马与创智汇文本结构高度相近，模板感较强。请在下一版让各案独立回答“与其他点有何不同”。三份均请补：版本日期、编制人、决策摘要、预算表和 KPI 表。"
        "正式外发 PDF 时请嵌入中文字体（如微软雅黑），避免跨平台出现方框或文字丢失。",
    )
    md("请降低模板感，补版本日期、编制人、决策摘要、预算表和 KPI 表；外发 PDF 请嵌入中文字体。")
    md()

    add_caption(doc, "表4  下一版提交清单")
    md("**表4  下一版提交清单**")
    md()
    headers = ["序号", "请提交的材料或章节"]
    rows = [
        ["1", "CGC 主体、授权、开票收款和专职团队附件"],
        ["2", "近三年活动清单、领事出席证明、脱敏转化案例"],
        ["3", "6 个月试运营方案、阶段目标、付款节点和未达标退出机制"],
        ["4", "双方职责矩阵和项目负责人名单"],
        ["5", "活动产品表、年度日历和单场标准流程"],
        ["6", "可验收 KPI（见表2）及统计口径"],
        ["7", "三年总成本、装修和活动预算上限、挂牌费内涵"],
        ["8", "免费活动与收费业务隔离规则（合同、发票、数据、投诉、是否分成）"],
        ["9", "品牌、涉外、反商业贿赂、数据隐私、肖像、IP、慈善资金、安保保险和退出条款"],
        ["10", "CGC 与驻沪总领馆的口径对照（见表1-1），删除官方授权和官方权限承诺"],
        ["11", "三案各自的事实勘误（创智汇面积上限约 300㎡、森马面积与在建条件、一滴水地址）"],
    ]
    add_table(doc, headers, rows, col_widths=[1.6, 15.0])
    md_table(headers, rows)

    add_paragraph(
        doc,
        "补齐上述内容后，方案的说服力和落地可操作性会更强。我们愿在下一版基础上，继续商议试点场次、验收标准和合作文本。",
        first_line=True,
    )
    md("补齐后愿继续商议试点场次、验收标准和合作文本。")
    md()

    md_path = DOCS_DIR / "CGC领事会客厅合作方案修改建议.md"
    md_path.write_text("\n".join(MD_LINES).rstrip() + "\n", encoding="utf-8")
    docx_path = OUTPUT_DIR / "CGC领事会客厅合作方案修改建议.docx"
    doc.save(docx_path)
    html_path = OUTPUT_DIR / "CGC领事会客厅合作方案修改建议.html"
    html_path.write_text(markdown_to_html(md_path.read_text(encoding="utf-8")), encoding="utf-8")
    return docx_path, md_path, html_path


if __name__ == "__main__":
    docx_path, md_path, html_path = build()
    print(f"docx: {docx_path}")
    print(f"md:   {md_path}")
    print(f"html: {html_path}")
