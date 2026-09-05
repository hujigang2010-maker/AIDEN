# -*- coding: utf-8 -*-
"""生成《联合策划服务合作协议》建议稿 Word。"""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import content as C

OUT = Path(__file__).resolve().parent.parent / "output" / "港大经管上海中心_联合策划服务合作协议_建议稿.docx"

CN_FONT = "宋体"
HEAD_FONT = "黑体"
GREEN = RGBColor(0x00, 0x3D, 0x2E)


def set_run_font(run, font_name: str, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def p(
    doc,
    text="",
    *,
    font=CN_FONT,
    size=12,
    bold=False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_indent=None,
    space_after=6,
    space_before=0,
    color=None,
):
    para = doc.add_paragraph()
    para.alignment = align
    pf = para.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.5
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    if text:
        run = para.add_run(text)
        set_run_font(run, font, size, bold=bold, color=color)
    return para


def rich(doc, segments, *, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=0.74, space_after=6):
    para = doc.add_paragraph()
    para.alignment = align
    pf = para.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.5
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    for text, bold in segments:
        run = para.add_run(text)
        set_run_font(run, CN_FONT, size, bold=bold)
    return para


def h(doc, text):
    return p(
        doc,
        text,
        font=HEAD_FONT,
        size=14,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=12,
        space_after=8,
        color=GREEN,
    )


def body(doc, text):
    return p(doc, text, first_indent=0.74)


def shade_cell(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill})
    tcPr.append(shd)


def build(path: Path | None = None) -> Path:
    path = path or OUT
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    normal = doc.styles["Normal"]
    normal.font.name = CN_FONT
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CN_FONT)

    p(doc, "联合策划服务合作协议", font=HEAD_FONT, size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    p(doc, "（建议稿 · 前期费用）", font=HEAD_FONT, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, color=GREEN)
    p(doc, f"{C.PROJECT_NAME}  ·  {C.VERSION}  ·  {C.DATE_CN}", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    rich(
        doc,
        [("甲方（联合策划与承办方）：", True), (C.OUR_PARTIES, False)],
        first_indent=0,
        space_after=2,
    )
    rich(
        doc,
        [("甲方指定结算主体：", True), (C.OUR_SETTLEMENT, False)],
        first_indent=0,
        space_after=2,
    )
    rich(
        doc,
        [("乙方（合作方）：", True), (f"{C.THEIR_LEGAL}（执行机构：{C.THEIR_UNIT}）", False)],
        first_indent=0,
        space_after=2,
    )
    rich(
        doc,
        [("乙方授权联系人：", True), (f"{C.THEIR_CONTACT}  {C.THEIR_TITLE}  {C.THEIR_TEL}  {C.THEIR_EMAIL}", False)],
        first_indent=0,
        space_after=2,
    )
    p(doc, f"乙方联系地址：{C.THEIR_ADDR}", size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)

    body(
        doc,
        f"鉴于{C.MEETING_DATE}双方在乙方上海中心交流：乙方明确合作主赛道为企业出海，"
        "合作层级应从经管学院维度展开，不限于 EMBA；乙方指定联系人负责华东华中高端教育招生与市场，"
        "经管学院层事项需中心主任支持。鉴于甲方具备上海校友与产业活动平台、出海主题活动与定向高管组织能力。"
        "双方就「外滩·出海课堂」联合策划及前期服务订立本协议。"
        "本协议为建议稿，供乙方院务/法务审核，不构成已生效合同。",
    )

    h(doc, "第一条　合作目的与原则")
    body(
        doc,
        "1.1 双方联合策划以出海为主题的闭门课与高规格活动场景：乙方输出学位品牌、场地、出海路径与招生流程，"
        "甲方输出议题、定向高管邀约、总领事或企业家场组织及现场统筹。来宾因出海议题到场，会后由乙方招生接口跟进，录取权始终在乙方。",
    )
    body(
        doc,
        "1.2 合作层级按乙方交流口径分为三层：层 A 为 EMBA/高端教育粘性场（本期主交付）；"
        "层 B 为邀请乙方上海中心主任出席一次甲方高规格出海活动（含在前期费用内，主任是否出席以其公务为准）；"
        "层 C 为出海高管课联合推广、原创谷及海外模块对接（本期仅出备忘，不作为验收）。",
    )
    body(
        doc,
        "1.3 本协议性质为联合策划服务，不是招生代理分成、不是联合办学、不是学位项目合作。"
        "任何一方不得对外宣称对方为其从属机构，或暗示存在联合学位、互认学分等安排。",
    )
    body(
        doc,
        "1.4 商业结构从简：乙方仅向甲方支付一笔前期联合策划服务费；本协议不约定学费、报名费、"
        "留位费分成，也不按录取人头计酬。",
    )

    h(doc, "第二条　合作内容")
    body(doc, "2.1 合作产品名称为「外滩·出海课堂」，首场举办地点原则上为乙方上海中心（外滩SOHO F栋），主题为企业出海（港股路径、跨境与产业决策），具体题目双方可改，赛道不改。")
    body(
        doc,
        f"2.2 本协议服务周期为自生效日起 {C.PLAN_DAYS} 日。甲方在周期内完成联合策划，落地首场出海闭门课（规模 {C.FIRST_EVENT_SIZE}），"
        "并完成一次对乙方上海中心主任出席甲方高规格出海活动的正式邀约。",
    )
    body(doc, "2.3 甲方交付范围如下：")
    for i, item in enumerate(C.FEE_COVERS, 1):
        body(doc, f"（{i}）{item}")
    body(doc, "2.4 下列事项不属于本协议标的，须另行书面确认：")
    for i, item in enumerate(C.FEE_NOT_COVERED, 1):
        body(doc, f"（{i}）{item}")
    body(
        doc,
        "2.5 首场主题、日期由双方在生效后 7 日内书面确认。建议主题见配套策划方案，乙方有权调整，"
        "但不得因此免除前期费用支付义务。",
    )
    body(
        doc,
        "2.6 主任体验：甲方提供一场已具备总领事、高质量企业家或同等级别嘉宾的出海主题活动，"
        "乙方指定联系人负责邀请上海中心主任。主任因公务未能出席的，只要甲方已发出正式书面邀约且活动如期举办，即视为本项交付完成。",
    )
    body(
        doc,
        "2.7 层 C 事项（本部企业出海高级管理课程联合推广、原创谷、海外模块企业参访）由双方在周期内形成书面备忘，"
        "不作为本期验收条件，亦不自动产生费用。",
    )

    h(doc, "第三条　双方权利与义务")
    body(doc, "3.1 乙方应当：")
    for i, item in enumerate(C.ROLES["港大经管上海中心"], 1):
        body(doc, f"（{i}）{item}")
    body(doc, "3.2 甲方应当：")
    for i, item in enumerate(C.ROLES["联合策划方"], 1):
        body(doc, f"（{i}）{item}")
    body(
        doc,
        "3.3 乙方指定潘嘉琰为日常接口人（华东华中 EMBA 招生与市场），变更接口人应提前 3 个工作日书面通知。"
        "甲方指定项目经理一名，于生效后 3 个工作日内书面告知乙方。经管学院层事项（主任出席、品牌授权超出招生口径的）"
        "由接口人转呈中心主任，接口人无义务越权承诺。",
    )
    body(
        doc,
        "3.4 录取、面试、学位授予、学费标准及退费等学院内部事项，均由乙方依其规章独立决定，"
        "甲方不得干预，亦不得向来宾作出录取或奖学金承诺。",
    )

    h(doc, "第四条　费用与支付（核心条款）")
    rich(
        doc,
        [
            ("4.1 乙方同意向甲方指定结算主体支付", False),
            (f"{C.FEE_NAME}，金额为{C.FEE_AMOUNT_CN}（¥{C.FEE_AMOUNT:,}）", True),
            ("。该费用为包干性质，覆盖第二条所列交付，不含本协议明确排除的项目。", False),
        ],
    )
    body(
        doc,
        f"4.2 支付方式：本协议生效之日起 {C.FEE_DAYS} 个工作日内，乙方将上述费用一次性支付至甲方指定结算主体账户。"
        "甲方在收到款项后 5 个工作日内开具合法有效增值税发票（税率与票种以结算主体资质为准，签署前书面确认）。",
    )
    body(
        doc,
        "4.3 未在约定期限内全额到账的，甲方有权暂停策划与邀约，并将首场档期顺延；逾期超过 15 个工作日仍未到账的，"
        "甲方有权书面解除本协议，已产生的实际成本由乙方在 10 个工作日内据实结算，且不少于费用总额的 30%。",
    )
    body(
        doc,
        "4.4 费用到账前，甲方无义务提交完整定向名单或对外发出正式邀请。双方可进行档期与主题沟通。",
    )
    body(
        doc,
        "4.5 除乙方书面要求增加协议外项目外，甲方不得再行加收本协议项下费用。第二场及以后场次，"
        "建议执行费为人民币陆万捌仟元整/场，以双方确认单为准，不自动生效。",
    )
    body(
        doc,
        "4.6 收款账户（签署时填写）：开户名称________；开户银行________；账号________；"
        "纳税人识别号________。账户信息与发票抬头不一致的，以双方书面确认为准。",
    )

    h(doc, "第五条　验收")
    body(
        doc,
        "5.1 策划验收：甲方提交客群画像、邀约话术及不少于 80 人的定向名单初稿，乙方在 5 个工作日内书面提出修改意见，"
        "逾期视为通过。名单为工作名单，不保证每一人均出席。",
    )
    body(
        doc,
        "5.2 首场验收：完成场地布置、签到、出海主题议程执行，并在结束后 7 日内提交纪要与 A/B/C 分级名单。"
        "到场人数受不可抗力、乙方改期或乙方未按期确认嘉宾影响的，不视为甲方违约。",
    )
    body(
        doc,
        "5.3 主任体验验收：以甲方发出的书面邀约（含时间、地点、嘉宾级别说明）及活动签到或现场影像为据。"
        "主任未出席不构成甲方违约。",
    )
    body(
        doc,
        "5.4 验收关注组织与交付，不将录取人数、学费进账、主任是否到场作为违约条件。策划方案所列 KPI 为工作目标，"
        "除「提交纪要与分级名单」及「完成主任书面邀约」外，不构成违约条款。",
    )

    h(doc, "第六条　名单、知识产权与品牌")
    body(
        doc,
        "6.1 因本合作形成的来宾名单、意向分级及联系方式，由双方共管。任何一方不得向第三方转让、出售，"
        "或用于与本合作无关的推销。协议终止后，共管名单仅可用于本合作已启动的招生跟进，期限不超过 12 个月。",
    )
    body(
        doc,
        "6.2 各方原有品牌、校徽、商标、课程内容仍归各方所有。一方使用另一方名称、校徽、场地照片，"
        "须事先书面同意，并符合对方视觉规范。",
    )
    body(
        doc,
        "6.3 本合作项下新产生的策划文案、议程、纪要，双方可在合作目的范围内使用；对外公开发布须征得对方同意。"
        "港大课程内容、招生政策口径以乙方最新官方文本为准。",
    )

    h(doc, "第七条　保密与合规")
    body(
        doc,
        "7.1 双方对合作中知悉的未公开信息、名单、商业条款负有保密义务，期限至协议终止后两年。"
        "因法律法规或监管要求必须披露的除外。",
    )
    body(
        doc,
        "7.2 活动内容应遵守国家教育、广告、外事及港澳相关规定。甲方不代理签证、不承诺海外身份、"
        "不使用「保录取」「内部名额」等表述。双方不得将违规跨境结算、地下钱庄、逃税或未经许可的虚拟货币业务"
        "作为活动主题或对外宣传口径。出海议题限于合法的企业国际化、港股路径、跨境贸易合规与金融基础设施。",
    )
    body(
        doc,
        "7.3 若乙方内部规定限制使用校外招生中介，双方确认本协议项下甲方身份为联合策划与活动承办，"
        "不以中介佣金方式结算。如院方要求调整表述，可签订补充协议，不改变前期费用金额。",
    )

    h(doc, "第八条　期限、变更与解除")
    body(doc, f"8.1 本协议自双方签署且前期费用到账之日起生效，有效期 {C.PLAN_DAYS} 日。到期自动终止，除非双方书面续签。")
    body(doc, "8.2 主题、日期、嘉宾等执行事项可用邮件或加盖印章的确认单变更，不构成对第四条费用条款的修改。")
    body(
        doc,
        "8.3 一方严重违约，经书面催告 10 个工作日仍未纠正的，另一方可解除协议。因乙方原因取消首场且不同意改期的，"
        "前期费用不予退还；因甲方原因未能举办首场且无法在周期内改期的，甲方应在 15 个工作日内退还费用的 50%。",
    )
    body(doc, "8.4 不可抗力导致无法履约的，双方协商改期；仍无法履行的，据实结算已发生成本后解除，互不追究。")

    h(doc, "第九条　违约责任与争议")
    body(doc, "9.1 除本协议另有约定外，违约方应赔偿因此造成的直接损失，不包括预期招生收入或学费损失。")
    body(
        doc,
        "9.2 因本协议引起的争议，双方先协商；协商不成的，提交甲方指定结算主体所在地有管辖权的人民法院诉讼。"
        "如结算主体签署前尚未确定，则提交上海市黄浦区人民法院。",
    )
    body(doc, "9.3 本协议适用中华人民共和国法律。涉及香港大学内部规章的，以不违反内地强制性规定为前提参照执行。")

    h(doc, "第十条　其他")
    body(doc, "10.1 本协议一式肆份，甲乙双方各执贰份，具有同等效力。可签署电子印章或扫描件，与原件同效。")
    body(
        doc,
        "10.2 未尽事宜可签补充协议。配套《联合策划方案》为执行参考，与本协议冲突时以本协议为准。"
        "金额与期限以本协议正文为准。",
    )
    body(
        doc,
        "10.3 乙方签署主体以香港大学经管学院认可的有权机构为准（学院本部、上海中心或指定法人）。"
        "甲方签署时须同时明确结算主体全称、统一社会信用代码及开票信息。",
    )
    p(
        doc,
        "（以下无正文，为签署页）",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=16,
        space_after=16,
    )

    table = doc.add_table(rows=12, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    left = [
        "甲方（联合策划与承办方）",
        C.OUR_PARTIES,
        "指定结算主体：",
        "统一社会信用代码：",
        "授权代表（签字）：",
        "职务：",
        "日期：　　年　　月　　日",
        "联系人：",
        "电话：",
        "邮箱：",
        "开户行：",
        "账号：",
    ]
    right = [
        "乙方（合作方）",
        f"{C.THEIR_LEGAL}",
        f"执行机构：{C.THEIR_UNIT}",
        "有权签署机构：",
        "授权代表（签字）：",
        "职务：",
        "日期：　　年　　月　　日",
        f"联系人：{C.THEIR_CONTACT}",
        f"电话：{C.THEIR_TEL}",
        f"邮箱：{C.THEIR_EMAIL}",
        f"地址：{C.THEIR_ADDR}",
        "盖章：",
    ]
    for i, (a, b) in enumerate(zip(left, right)):
        cell0, cell1 = table.cell(i, 0), table.cell(i, 1)
        if i == 0:
            shade_cell(cell0, "003D2E")
            shade_cell(cell1, "003D2E")
        for cell, txt in ((cell0, a), (cell1, b)):
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            run = para.add_run(txt)
            set_run_font(run, CN_FONT, 10.5, bold=(i == 0), color=(RGBColor(0xFF, 0xFF, 0xFF) if i == 0 else None))

    p(
        doc,
        "附件：前期费用覆盖范围与 90 天交付清单（见配套 Excel）；联合策划方案 PPT 为执行参考。",
        size=10.5,
        space_before=16,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    p(
        doc,
        "声明：本稿为商务建议文本，不构成法律意见。涉及香港大学内部授权、发票资质与跨境支付的，签署前应由双方有权部门审核。",
        size=10.5,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    doc.save(path)
    return path


if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else OUT
    pth = build(out)
    print(f"已生成 {pth}")
