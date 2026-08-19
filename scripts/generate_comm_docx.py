#!/usr/bin/env python3
"""生成肇事方沟通方案与影像分析 Word。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from comm_plan import (
    BANS,
    COMM_GOAL,
    CT_PLAIN,
    DOC_SUB,
    DOC_TITLE,
    FLOW,
    LAWYER_NOW,
    MED_FOR_TALK,
    MED_SUMMARY,
    OPENING_SCRIPT,
    RIDER_REPLIES,
    SCHEMES,
    SHOW_LIST,
)
from content import (
    ACCIDENT,
    ASR_CORRECTIONS,
    CLOUD_FILM,
    CT_REPORTS,
    DISABILITY,
    HOSPITAL_CHOICE,
    HOSPITAL_TRACK,
    INJURY_NOT_THIS_ACCIDENT,
    INJURY_THIS_ACCIDENT,
    MONEY,
    PARTIES,
    TALKING_POINTS,
)
from generate_docx import (
    ACCENT,
    CHINESE_FONT,
    HEADING_FONT,
    MUTED,
    RED,
    _add_paragraph,
    _body,
    _bullet,
    _heading,
    _table,
)

MUST_GET = [
    "微信群：骑手 + 美团站点 + 平台保险 + 家属（胡继刚）",
    "保险对接人姓名、电话、保单/工号；是否必须先有认定书才能垫付",
    "市北区人民医院（抚顺路院区）后续费用认不认；齐鲁已发生费用认不认",
    "已发生费用处理路径：谁先垫、转到哪、要哪些发票",
]


def build_comm_document(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = CHINESE_FONT
    normal_style.font.size = Pt(12)
    rpr = normal_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CHINESE_FONT)

    _add_paragraph(
        doc,
        DOC_TITLE,
        font=HEADING_FONT,
        size=18,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
        line_spacing=1.3,
        color=ACCENT,
    )
    _add_paragraph(
        doc,
        DOC_SUB,
        font=HEADING_FONT,
        size=11,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
        line_spacing=1.3,
        color=MUTED,
    )
    _add_paragraph(
        doc,
        "内部使用 · 非律师函 · 非伤残鉴定书 · 依据为齐鲁医院已审核 CT",
        font=CHINESE_FONT,
        size=10.5,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
        color=MUTED,
    )

    _heading(doc, "一、先把「体检报告」说清楚")
    for line in MED_SUMMARY:
        _bullet(doc, line)
    _body(
        doc,
        f"医院：{CLOUD_FILM['hospital']}。云影像患者：{CLOUD_FILM['patient_display']}，"
        f"{CLOUD_FILM['sex']}，{CLOUD_FILM['age']} 岁。"
        f"{CLOUD_FILM['disclaimer']} 报告只能证明伤了什么、做了什么手术，不能当残级，也不能当入职体检结论。",
    )
    _body(doc, PARTIES["injured"] + " " + PARTIES["family"])
    _body(doc, PARTIES["other"] + " " + PARTIES["police"])
    _body(doc, f"地点：{ACCIDENT['place']}。过程：{ACCIDENT['process']} {ACCIDENT['police_range']}")

    _heading(doc, "二、三份 CT 对照分析（对外口径）")
    _table(
        doc,
        ["检查", "白话结论", "对骑手怎么说", "索赔"],
        [[c["title"], c["plain"], c["to_rider"], c["claim"]] for c in CT_PLAIN],
        col_widths=[4.2, 4.8, 4.8, 2.2],
    )
    _table(
        doc,
        ["日期", "检查号", "项目 / 科室", "影像诊断（报告原文）", "本次事故"],
        [
            [
                r["date"],
                r["no"],
                f"{r['item']}\n{r['place']}（{r['status']}）",
                "；".join(r["diagnosis"]),
                "计入" if r["accident_related"] else "不计入",
            ]
            for r in CT_REPORTS
        ],
        col_widths=[2.2, 2.6, 4.4, 5.0, 1.8],
    )
    _heading(doc, "本次事故计入的伤", 2)
    for x in INJURY_THIS_ACCIDENT:
        _bullet(doc, x)
    _heading(doc, "不计入本次事故", 2)
    for x in INJURY_NOT_THIS_ACCIDENT:
        _bullet(doc, x)
    _heading(doc, "通话里必须改口的说法", 2)
    _table(
        doc,
        ["不要再说", "报告实际写的"],
        [[a, b] for a, b in ASR_CORRECTIONS],
        col_widths=[7.0, 9.0],
    )
    _body(doc, DISABILITY["now"] + " " + DISABILITY["cannot_grade"])
    _body(doc, DISABILITY["when"] + " " + DISABILITY["standard"])
    _body(doc, HOSPITAL_TRACK["now"])
    _body(doc, HOSPITAL_TRACK["not_rehab"])
    _body(doc, HOSPITAL_TRACK["two_tracks"])
    _body(doc, MONEY["paid_family_816"] + " " + MONEY["paid_815_conflict"])
    _body(doc, MONEY["insurance"])
    _body(doc, HOSPITAL_CHOICE["headline"])

    _heading(doc, "三、跟肇事者沟通：这一场到底要什么")
    for line in COMM_GOAL:
        _bullet(doc, line)
    _heading(doc, "这一场必须拿到", 2)
    for x in MUST_GET:
        _bullet(doc, x)
    _heading(doc, "对骑手的四句态度", 2)
    for x in TALKING_POINTS["to_rider"]:
        _bullet(doc, x)
    _heading(doc, "伤情对外只准这样说", 2)
    for x in MED_FOR_TALK:
        _bullet(doc, x)

    _heading(doc, "四、建议开场（可照着念）")
    _body(doc, OPENING_SCRIPT)
    _body(
        doc,
        "念完停住，等她接。不要自己往下加残级、全责、现金数字。第一次由胡继刚一人出面，岳父场外看稿，律师不到场。",
    )

    _heading(doc, "五、对方常见接话，怎么接")
    _table(
        doc,
        ["对方可能说", "建议你怎么接"],
        [[a, b] for a, b in RIDER_REPLIES],
        col_widths=[5.5, 10.5],
    )

    _heading(doc, "六、这一场绝对不能说")
    for x in BANS:
        _bullet(doc, x, color=RED)
    for x in TALKING_POINTS["do_not_say"]:
        _bullet(doc, x, color=RED)

    _heading(doc, "七、要不要请律师（完整判断）")
    _body(doc, LAWYER_NOW["answer"])
    _heading(doc, "现在不请诉讼律师的原因", 2)
    for x in LAWYER_NOW["why_not"]:
        _bullet(doc, x)
    _body(doc, LAWYER_NOW["why_yes_internal"])
    _heading(doc, "出现下面任一情形，再对外用律师", 2)
    for x in LAWYER_NOW["triggers"]:
        _bullet(doc, x)
    _body(
        doc,
        "对外用律师时：先发律师函催保险进场和指定医院，再等认定书后起诉。"
        "执行盯美团平台险，不指望骑手个人财产。"
        "不要用「走一般程序、拘留」威胁对方——己方三轮无正规号牌，一般程序可能把无牌上路、甚至无证驾驶的处罚砸回来。",
    )

    _heading(doc, "八、三套完整方案，按情况选用")
    for sch in SCHEMES:
        _heading(doc, sch["name"], 2)
        _bullet(doc, sch["use"], bold_prefix="何时用：")
        _bullet(doc, sch["steps"], bold_prefix="怎么走：")
        _bullet(doc, sch["lawyer"], bold_prefix="律师：")
        _bullet(doc, sch["risk"], bold_prefix="风险：")
        _bullet(doc, sch["dont"], bold_prefix="不做：")

    _heading(doc, "九、接下来沟通顺序（12 步）")
    _body(
        doc,
        "按天推进，不要跳步。第 1–7 步现在就走方案甲；卡住再进方案乙或丙。"
        "每一步「拿到什么才算过」没拿到，就不要假装过关。",
    )
    _table(
        doc,
        ["步", "何时", "谁出面 / 找谁", "做什么", "过关", "过不了 / 律师"],
        [
            [
                str(s["n"]),
                s["when"],
                f"{s['who']} → {s['to']}",
                s["do"],
                s["get"],
                f"{s['fail']}；律师：{s['lawyer']}",
            ]
            for s in FLOW
        ],
        col_widths=[1.2, 2.2, 3.0, 4.0, 2.8, 2.8],
    )

    _heading(doc, "十、现场决策树")
    for title, text in HOSPITAL_CHOICE["tree"]:
        _bullet(doc, text, bold_prefix=title + "：")
    _body(
        doc,
        "认定书出具后，交警不再参与赔多少钱。已发生医疗费按责任比例报保险；伤残等治疗终结再鉴定。"
        "协商不成只能起诉。评残走司法鉴定、用途写道路交通事故，不要走人社局工伤通道。",
    )

    _heading(doc, "十一、这一场可以给对方看的材料")
    for x in SHOW_LIST:
        _bullet(doc, x)

    _heading(doc, "十二、使用边界")
    _body(
        doc,
        "本文是家属内部沟通脚本和影像对照，不是医学鉴定、不是司法鉴定、不是律师函、不是向法院提交的诉状。"
        "对外出示时，以医院已审核报告原文、发票原件、交警认定书为准。"
        "费用两套口述数字冲突，对外只说以发票为准、还在发生。",
    )
    _add_paragraph(
        doc,
        "配套：同名 PDF；《肇事方沟通流程表》Excel（12 步可勾选、对方接话、律师触发、出示清单）。伤情底稿见备忘录与伤残简报。",
        font=CHINESE_FONT,
        size=10.5,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12,
        color=MUTED,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print("Wrote", output_path)


if __name__ == "__main__":
    out = Path(__file__).resolve().parent.parent / "deliverables" / "青岛抚顺路和哈尔滨路路口交通事故_肇事方沟通方案与体检分析_20260817.docx"
    build_comm_document(out)
