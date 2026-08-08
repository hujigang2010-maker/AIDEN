#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《能量话语体系对照表》Excel 与说明文档 Word。

五种话语：物理学能量、生物学能量、心理学能量（意志力/心流）、
道家真气、灵性能量——标明可互译处与不可通约处。
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output"
XLSX_PATH = OUT_DIR / "能量话语体系对照表.xlsx"
DOCX_PATH = OUT_DIR / "能量话语体系对照表_说明.docx"

SYSTEMS = [
    "物理学能量",
    "生物学能量",
    "心理学能量",
    "道家真气",
    "灵性能量",
]

# ---------------------------------------------------------------------------
# 样式
# ---------------------------------------------------------------------------

THIN = Border(
    left=Side(style="thin", color="B0B0B0"),
    right=Side(style="thin", color="B0B0B0"),
    top=Side(style="thin", color="B0B0B0"),
    bottom=Side(style="thin", color="B0B0B0"),
)
HEADER_FILL = PatternFill("solid", fgColor="1A3A5C")
HEADER_FONT = Font(name="微软雅黑", size=11, bold=True, color="FFFFFF")
TITLE_FONT = Font(name="微软雅黑", size=14, bold=True, color="1A3A5C")
SUB_FONT = Font(name="微软雅黑", size=10, color="555555")
CELL_FONT = Font(name="微软雅黑", size=10)
BOLD_FONT = Font(name="微软雅黑", size=10, bold=True)
WRAP = Alignment(wrap_text=True, vertical="center", horizontal="left")
CENTER = Alignment(wrap_text=True, vertical="center", horizontal="center")

FILL_STRONG = PatternFill("solid", fgColor="C6EFCE")  # 强可互译
FILL_PARTIAL = PatternFill("solid", fgColor="FFEB9C")  # 部分/隐喻
FILL_WEAK = PatternFill("solid", fgColor="FCE4D6")  # 弱/类比
FILL_NONE = PatternFill("solid", fgColor="FFC7CE")  # 不可通约
FILL_SELF = PatternFill("solid", fgColor="D9E2F3")  # 自身
FILL_ALT = PatternFill("solid", fgColor="F2F5F8")
FILL_NOTE = PatternFill("solid", fgColor="FFF8E7")


def style_header_row(ws, row, cols):
    for c in range(1, cols + 1):
        cell = ws.cell(row=row, column=c)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = CENTER
        cell.border = THIN


def style_body(ws, start_row, end_row, cols, alt=True):
    for r in range(start_row, end_row + 1):
        for c in range(1, cols + 1):
            cell = ws.cell(row=r, column=c)
            cell.font = CELL_FONT
            cell.alignment = WRAP
            cell.border = THIN
            if alt and r % 2 == 0 and cell.fill.fgColor is None or (
                alt and r % 2 == 0 and (not cell.fill.fgColor or cell.fill.fgColor.rgb in ("00000000", None))
            ):
                if not cell.fill or cell.fill.fgColor is None or str(getattr(cell.fill.fgColor, "rgb", "")) in (
                    "00000000",
                    "None",
                ):
                    cell.fill = FILL_ALT


def set_widths(ws, widths):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w


def apply_fill_by_level(cell, level: str):
    mapping = {
        "强可互译": FILL_STRONG,
        "部分可译": FILL_PARTIAL,
        "弱类比": FILL_WEAK,
        "不可通约": FILL_NONE,
        "自身": FILL_SELF,
    }
    if level in mapping:
        cell.fill = mapping[level]


# ---------------------------------------------------------------------------
# Sheet 1：总览对照
# ---------------------------------------------------------------------------

OVERVIEW_HEADERS = [
    "对照维度",
    "物理学能量",
    "生物学能量",
    "心理学能量（意志力/心流）",
    "道家真气",
    "灵性能量",
]

OVERVIEW_ROWS = [
    [
        "核心定义",
        "做功的能力；守恒的标量；可用焦耳度量",
        "维持生命活动的代谢可用能（ATP、营养氧化）",
        "主观的可用心理资源：注意力、自制力、动机与心流体验（多为隐喻）",
        "人身可感、可练、可导的精微之气；贯通精—气—神",
        "超越个体的神圣/宇宙/微妙之力（传统异质，含圣灵、普拉那、气场等）",
    ],
    [
        "本体论地位",
        "客观物理量；科学实在论默认项",
        "生物化学过程；可还原到物理化学，但以有机体功能为框架",
        "心理学构念/体验描述；一般不主张字面能量守恒",
        "传统宇宙论中的实体—过程混合；非 SI 物理量",
        "形而上学/信仰/体验性存在；通常不可被实验室操作定义",
    ],
    [
        "典型单位/度量",
        "焦耳(J)、电子伏特(eV)、卡路里(cal)",
        "千卡(kcal)、ATP 分子数、基础代谢率(BMR)、VO₂",
        "多为行为指标与自评量表；无统一能量单位",
        "主观体感、脉象、传统辨证；无国际单位",
        "一般无科学度量；偶有争议性「气场摄影」等伪测",
    ],
    [
        "守恒与耗散",
        "严格守恒（封闭系）；形式可转化",
        "对有机体开放系：摄入—消耗—储存；最终服从热力学",
        "「耗尽」叙事常见（自我损耗），但证据与机制有争议；非物理守恒",
        "可补、可泄、可滞、可化；讲究调养而非焦耳守恒",
        "依传统而异：可被恩赐、传递、污染或升华；非热力学语言",
    ],
    [
        "传递与转化",
        "功、热、辐射、场；质量—能量等价",
        "消化、呼吸、磷酸化；化学能→机械能/热能",
        "休息、睡眠、激励、心流条件匹配；多为调节隐喻",
        "呼吸、导引、存想、周天运转；经络为通道叙事",
        "仪式、祈祷、加持、冥想、传承；通道多为象征或神秘",
    ],
    [
        "关键人物/理论",
        "焦耳、亥姆霍兹、爱因斯坦（E=mc²）",
        "克雷布斯循环、米切尔化学渗透、代谢组学",
        "弗洛伊德（力比多，历史隐喻）、鲍迈斯特（自我损耗）、契克森米哈赖（心流）",
        "《黄帝内经》、内丹传统、气功导引文献",
        "各宗教神秘主义、瑜伽普拉那、当代「新时代」综合话语",
    ],
    [
        "可检验性",
        "高：可重复实验与精密测量",
        "高：生化/生理测量；临床与运动科学",
        "中：可测行为与神经相关，但「能量」本身常不可操作化",
        "低—中：练习效应可测（呼吸、自主神经），「气」本体难证伪/难证实",
        "低：依赖信仰与私人体验；公共可检验性弱",
    ],
    [
        "话语功能",
        "解释自然、工程与宇宙结构",
        "解释生命维持、疾病、运动与营养",
        "解释动机、倦怠、专注与峰值体验",
        "指导修身、养生、武术与宇宙—人身同构",
        "提供意义、救赎、联结与超越叙事",
    ],
    [
        "常见误用",
        "把「感觉有能量」说成焦耳守恒被破坏",
        "把卡路里神话化；忽视行为与环境",
        "把隐喻当物理实体（「脑电量用完了」）",
        "用焦耳/电磁直接等同真气；或医学替代延误",
        "用科学辞藻包装不可检验主张；贩卖「高频能量」",
    ],
]


def _fill_overview(ws):
    ws["A1"] = "能量话语体系对照表 · 总览"
    ws["A1"].font = TITLE_FONT
    ws.merge_cells("A1:F1")
    ws["A2"] = (
        "说明：五种话语共用「能量」一词，但指称对象、度量方式与可检验性不同。"
        "本表用于澄清概念边界，而非做价值裁决。"
    )
    ws["A2"].font = SUB_FONT
    ws.merge_cells("A2:F2")
    ws.row_dimensions[1].height = 24
    ws.row_dimensions[2].height = 36

    for col, h in enumerate(OVERVIEW_HEADERS, 1):
        ws.cell(row=4, column=col, value=h)
    style_header_row(ws, 4, 6)

    for i, row in enumerate(OVERVIEW_ROWS):
        r = 5 + i
        for c, val in enumerate(row, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.font = BOLD_FONT if c == 1 else CELL_FONT
            cell.alignment = WRAP
            cell.border = THIN
            if i % 2 == 1:
                cell.fill = FILL_ALT
        ws.row_dimensions[r].height = 72

    set_widths(ws, [16, 28, 28, 32, 28, 32])
    ws.freeze_panes = "B5"
    ws.print_title_rows = "4:4"


def build_overview_sheet(wb: Workbook):
    ws = wb.create_sheet("01_总览对照")
    _fill_overview(ws)


# ---------------------------------------------------------------------------
# Sheet 2：互译矩阵
# ---------------------------------------------------------------------------

# 上三角语义：行→列 的可译程度与一句话理由
# level: 强可互译 / 部分可译 / 弱类比 / 不可通约 / 自身
TRANSLATION_MATRIX = {
    ("物理学能量", "物理学能量"): ("自身", "同一话语内部，单位与定律自洽。"),
    ("物理学能量", "生物学能量"): (
        "强可互译",
        "卡路里/ATP 可还原为焦耳与化学势；生物学能量是物理能量在生命系统中的组织形态。",
    ),
    ("物理学能量", "心理学能量"): (
        "弱类比",
        "「脑耗能」可谈葡萄糖与 ATP，但意志力/心流不等于可守恒的焦耳预算。",
    ),
    ("物理学能量", "道家真气"): (
        "不可通约",
        "真气无 SI 操作定义；导引练习的生理效应≠真气即电磁或热能。",
    ),
    ("物理学能量", "灵性能量"): (
        "不可通约",
        "神圣/微妙之力不进入热力学账本；用焦耳「证明」灵性是范畴错误。",
    ),
    ("生物学能量", "物理学能量"): (
        "强可互译",
        "代谢最终服从热力学；营养学卡路里可换算焦耳（组织层次不同）。",
    ),
    ("生物学能量", "生物学能量"): ("自身", "同一话语内部，代谢框架自洽。"),
    ("生物学能量", "心理学能量"): (
        "部分可译",
        "睡眠、血糖、疲劳影响自制与专注；但心流/意志力仍有目标、意义与技能匹配等不可还原维度。",
    ),
    ("生物学能量", "道家真气"): (
        "弱类比",
        "呼吸与自主神经调节有生理对应；经络—真气本体仍超出现行生物学模型。",
    ),
    ("生物学能量", "灵性能量"): (
        "不可通约",
        "生化代谢无法翻译「恩典/气场」；最多描述冥想的生理副作用。",
    ),
    ("心理学能量", "物理学能量"): (
        "弱类比",
        "仅在「大脑耗能」窄义上挂钩；自我损耗≠能量守恒定律的心理学版。",
    ),
    ("心理学能量", "生物学能量"): (
        "部分可译",
        "可部分落到疲劳、激素、唤醒水平；「心理能量」整体仍是构念。",
    ),
    ("心理学能量", "心理学能量"): ("自身", "同一话语内部（且内部流派已多义）。"),
    ("心理学能量", "道家真气"): (
        "弱类比",
        "专注、体感、情绪调节与修炼体验有家族相似；通道与宇宙论不可互换。",
    ),
    ("心理学能量", "灵性能量"): (
        "弱类比",
        "心流、敬畏、高峰体验可被灵性话语再描述，但「被加持」主张超出心理学。",
    ),
    ("道家真气", "物理学能量"): (
        "不可通约",
        "把真气写成「生物电/量子」是修辞嫁接，不是互译。",
    ),
    ("道家真气", "生物学能量"): (
        "弱类比",
        "可承认练习改变呼吸与循环；不能把丹田气库存成 ATP 当量。",
    ),
    ("道家真气", "心理学能量"): (
        "弱类比",
        "调心调息与注意力训练可对话；「气沿经络」不可被意志力量表替代。",
    ),
    ("道家真气", "道家真气"): ("自身", "同一传统内部（亦有门派差异）。"),
    ("道家真气", "灵性能量"): (
        "部分可译",
        "同属精微身体/修炼家族；但道门宇宙论≠普世「宇宙能量」或一神论圣灵。",
    ),
    ("灵性能量", "物理学能量"): (
        "不可通约",
        "范畴错误：神圣叙事不服从能量守恒实验协议。",
    ),
    ("灵性能量", "生物学能量"): (
        "不可通约",
        "生理相关物（如放松反应）≠灵性本体被测量。",
    ),
    ("灵性能量", "心理学能量"): (
        "弱类比",
        "可用敬畏、意义、自我超越作心理学转述，会丢失救赎/神圣维度。",
    ),
    ("灵性能量", "道家真气"): (
        "部分可译",
        "部分修炼与气脉叙事可对照；不可把真气、普拉那、圣灵无差别互换。",
    ),
    ("灵性能量", "灵性能量"): ("自身", "名义统一，传统内部高度异质。"),
}


def build_matrix_sheet(wb: Workbook):
    ws = wb.create_sheet("02_互译矩阵")
    ws["A1"] = "两两互译矩阵（行 → 列）"
    ws["A1"].font = TITLE_FONT
    ws.merge_cells("A1:F1")
    ws["A2"] = (
        "图例：绿=强可互译｜黄=部分可译｜橙=弱类比｜红=不可通约｜蓝=自身。"
        "「可译」指概念可无损或近似映射，不是价值高低。"
    )
    ws["A2"].font = SUB_FONT
    ws.merge_cells("A2:F2")

    ws.cell(row=4, column=1, value="从 \\ 到")
    for i, name in enumerate(SYSTEMS, 2):
        ws.cell(row=4, column=i, value=name)
    style_header_row(ws, 4, 6)
    ws.cell(row=4, column=1).fill = HEADER_FILL

    for r, src in enumerate(SYSTEMS, 5):
        ws.cell(row=r, column=1, value=src).font = BOLD_FONT
        ws.cell(row=r, column=1).alignment = CENTER
        ws.cell(row=r, column=1).border = THIN
        ws.cell(row=r, column=1).fill = FILL_ALT
        for c, dst in enumerate(SYSTEMS, 2):
            level, reason = TRANSLATION_MATRIX[(src, dst)]
            cell = ws.cell(row=r, column=c, value=f"【{level}】\n{reason}")
            cell.font = CELL_FONT
            cell.alignment = WRAP
            cell.border = THIN
            apply_fill_by_level(cell, level)
        ws.row_dimensions[r].height = 90

    # 图例区
    ws["A11"] = "互译等级定义"
    ws["A11"].font = TITLE_FONT
    legends = [
        ("强可互译", "双方核心量可用共同操作定义或可还原换算，损失可忽略。"),
        ("部分可译", "有稳定的局部对应（机制、相关物、实践效果），但整体不可还原。"),
        ("弱类比", "修辞或体验上相似，强行等同会产生范畴错误。"),
        ("不可通约", "本体论、度量与真值条件无法共享；翻译即改写或抹除。"),
    ]
    for i, (lv, desc) in enumerate(legends):
        r = 12 + i
        c1 = ws.cell(row=r, column=1, value=lv)
        c2 = ws.cell(row=r, column=2, value=desc)
        c1.font = BOLD_FONT
        c2.font = CELL_FONT
        c1.alignment = CENTER
        c2.alignment = WRAP
        apply_fill_by_level(c1, lv)
        for col in range(1, 7):
            ws.cell(row=r, column=col).border = THIN
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=6)
        ws.row_dimensions[r].height = 28

    set_widths(ws, [14, 22, 22, 24, 22, 24])
    ws.freeze_panes = "B5"


# ---------------------------------------------------------------------------
# Sheet 3：可互译清单 / 不可通约清单
# ---------------------------------------------------------------------------

TRANSLATABLE = [
    [
        "物理学 ↔ 生物学",
        "强可互译",
        "卡路里与焦耳换算；ATP 水解的自由能；体温与散热",
        "把「生命力」直接写成违反热力学第二定律的神秘盈余",
    ],
    [
        "生物学 → 心理学（局部）",
        "部分可译",
        "睡眠债、低血糖、炎症与疲劳对自制/注意的影响；运动提升心境",
        "宣称「只要补糖就能补满意志力库存」",
    ],
    [
        "心理学 → 生物学（局部）",
        "部分可译",
        "慢性压力的生理代价；心流时的自主神经与注意网络变化（相关）",
        "把心流体验还原为「就是多巴胺」并取消现象学描述",
    ],
    [
        "道家真气 ↔ 灵性（家族相似）",
        "部分可译",
        "精微身体、呼吸门控、修炼次第、身心一如的实践语法",
        "把真气、普拉那、圣灵、宇宙振动当成同一可互换实体",
    ],
    [
        "道家练习 → 生理/心理效应",
        "弱类比—部分可译",
        "调息改善 HRV/焦虑；站桩的本体感觉与注意训练",
        "用效应证明「经络气」已在物理学中被测到",
    ],
    [
        "灵性体验 → 心理学转述",
        "弱类比",
        "敬畏、自我消融、意义感、皈依叙事的心理功能",
        "转述后仍声称已穷尽「神圣」本身",
    ],
]

INCOMMENSURABLE = [
    [
        "物理学 ↔ 灵性能量",
        "不可通约",
        "一方要求可重复测量与守恒；另一方以恩典/启示/微妙场为真",
        "「用量子力学证明疗愈能量」类话术",
    ],
    [
        "物理学 ↔ 道家真气（本体层）",
        "不可通约",
        "焦耳/场方程 vs 精气神与经络通道；无共同操作定义",
        "「丹田是电池、气是生物电」式硬译",
    ],
    [
        "生物学 ↔ 灵性能量（本体层）",
        "不可通约",
        "代谢网络不包含「被祝圣」谓词",
        "补充剂广告套用「提升灵性振动频率」",
    ],
    [
        "心理学字面能量 ↔ 物理守恒",
        "不可通约",
        "自我损耗不是封闭系能量守恒定理",
        "「今天意志力焦耳用完了」当作科学陈述",
    ],
    [
        "各灵性传统内部异质项",
        "内部亦常不可通约",
        "圣灵、气场、昆达里尼、加持力的权威来源与修法不同",
        "「新时代」大熔炉式混用导致语义崩塌",
    ],
]


def build_lists_sheet(wb: Workbook):
    ws = wb.create_sheet("03_可译与不可通约")
    ws["A1"] = "可互译条目与不可通约条目"
    ws["A1"].font = TITLE_FONT
    ws.merge_cells("A1:D1")

    ws["A3"] = "一、可以（近似）互译或局部对接的方向"
    ws["A3"].font = Font(name="微软雅黑", size=12, bold=True, color="1F7A3D")
    headers = ["话语对", "等级", "可译内容（保留什么）", "翻译时必须丢掉/禁止什么"]
    for c, h in enumerate(headers, 1):
        ws.cell(row=4, column=c, value=h)
    style_header_row(ws, 4, 4)

    for i, row in enumerate(TRANSLATABLE):
        r = 5 + i
        for c, val in enumerate(row, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.font = CELL_FONT
            cell.alignment = WRAP
            cell.border = THIN
            if c == 2:
                apply_fill_by_level(cell, val.split("—")[0] if "—" in val else val)
        ws.row_dimensions[r].height = 48

    start = 5 + len(TRANSLATABLE) + 2
    ws.cell(row=start, column=1, value="二、完全或实质上不可通约的方向").font = Font(
        name="微软雅黑", size=12, bold=True, color="B01F24"
    )
    for c, h in enumerate(headers, 1):
        ws.cell(row=start + 1, column=c, value=h)
    style_header_row(ws, start + 1, 4)

    for i, row in enumerate(INCOMMENSURABLE):
        r = start + 2 + i
        for c, val in enumerate(row, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.font = CELL_FONT
            cell.alignment = WRAP
            cell.border = THIN
            if c == 2:
                apply_fill_by_level(cell, "不可通约")
        ws.row_dimensions[r].height = 48

    set_widths(ws, [28, 14, 48, 42])


# ---------------------------------------------------------------------------
# Sheet 4：术语对照
# ---------------------------------------------------------------------------

TERMS = [
    ["能量 / energy", "焦耳可度量的守恒量", "代谢可用能", "心理资源隐喻", "真气/气", "灵力/普拉那/圣灵等"],
    ["力 / force", "矢量，F=ma", "肌力、收缩力", "意志力（隐喻）", "气力、劲", "灵力、权能"],
    ["场 / field", "物理场（电磁等）", "少用；偶指微环境", "场域/情境（社会心理）", "气场（体感氛围）", "气场/能量场（微妙）"],
    ["守恒", "定律", "质量—能量在开放系中的收支", "库存隐喻（易误导）", "固精养气（规范性）", "恩典不可被「库存化」或反之"],
    ["转化", "功热互变等", "化学能→热能/功", "动机重定向、情绪调节", "炼精化气、炼气化神", "转化/圣化/觉醒"],
    ["耗尽", "可用能减少（仍守恒）", "糖原耗竭、过度训练", "自我损耗、倦怠", "气虚、气泄", "灵性枯干（传统隐喻）"],
    ["流动", "能流、功率", "血液/代谢流", "心流(flow)", "气机条达、周天", "能量流动（新时代常用）"],
    ["堵塞", "罕用（非术语）", "循环/代谢障碍", "心理阻滞、反刍", "气滞、经络不通", "气场堵塞、负面能量"],
    ["充电/补给", "储能", "进食、休息、营养", "休息、奖励、意义感", "采气、服气、食补药补", "祷告、仪式、加持"],
    ["测量", "仪器+单位", "实验室/可穿戴", "量表+行为", "体感+四诊", "信仰见证/争议仪器"],
]


def build_terms_sheet(wb: Workbook):
    ws = wb.create_sheet("04_术语对照")
    ws["A1"] = "高频词在五种话语中的含义漂移"
    ws["A1"].font = TITLE_FONT
    ws.merge_cells("A1:F1")
    ws["A2"] = "同一中文词跨列不可默认同义；阅读时请先问：它在这一列是否可操作、可测量、可证伪？"
    ws["A2"].font = SUB_FONT
    ws.merge_cells("A2:F2")

    headers = ["词语", "物理学", "生物学", "心理学", "道家真气", "灵性"]
    for c, h in enumerate(headers, 1):
        ws.cell(row=4, column=c, value=h)
    style_header_row(ws, 4, 6)

    for i, row in enumerate(TERMS):
        r = 5 + i
        for c, val in enumerate(row, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.font = BOLD_FONT if c == 1 else CELL_FONT
            cell.alignment = WRAP
            cell.border = THIN
            if i % 2 == 1:
                cell.fill = FILL_ALT
        ws.row_dimensions[r].height = 40

    set_widths(ws, [16, 22, 20, 22, 22, 26])
    ws.freeze_panes = "B5"


# ---------------------------------------------------------------------------
# Sheet 5：使用指南
# ---------------------------------------------------------------------------

GUIDE_ROWS = [
    ["适用场景", "科普写作、课程设计、跨领域对话仲裁、内容选题（避免把隐喻当定律）"],
    ["推荐读法", "先读「01_总览对照」→「02_互译矩阵」→「03_可译与不可通约」→「04_术语对照」"],
    ["写作原则1", "凡声称跨话语「已被科学证明」，检查是否仅证明了生理相关物，而非本体同一。"],
    ["写作原则2", "心理学与日常口语中的「能量」默认视为隐喻，除非给出操作定义。"],
    ["写作原则3", "道家与灵性话语可尊重其内在规范与体验报告，但不要偷偷换成 SI 单位。"],
    ["写作原则4", "生物学是物理学与心理学之间最稳的桥梁；仍不要用它吞并真气/灵性本体。"],
    ["争议提示", "自我损耗理论存在复制危机；经络实质争议未决——表中已按「弱/部分」处理，不作终审。"],
    ["颜色图例", "绿强可互译 / 黄部分可译 / 橙弱类比 / 红不可通约 / 蓝自身"],
]


def build_guide_sheet(wb: Workbook):
    ws = wb.create_sheet("00_使用说明")
    ws["A1"] = "能量话语体系对照表"
    ws["A1"].font = TITLE_FONT
    ws.merge_cells("A1:B1")
    ws["A2"] = "副标题：物理学 · 生物学 · 心理学（意志力/心流）· 道家真气 · 灵性能量"
    ws["A2"].font = SUB_FONT
    ws.merge_cells("A2:B2")
    ws["A3"] = "目标：标明哪些可以互译，哪些完全不可通约——让「能量」一词停止走私概念。"
    ws["A3"].font = Font(name="微软雅黑", size=10, bold=True, color="1A3A5C")
    ws.merge_cells("A3:B3")

    ws.cell(row=5, column=1, value="项目").font = HEADER_FONT
    ws.cell(row=5, column=2, value="内容").font = HEADER_FONT
    ws.cell(row=5, column=1).fill = HEADER_FILL
    ws.cell(row=5, column=2).fill = HEADER_FILL
    for c in (1, 2):
        ws.cell(row=5, column=c).border = THIN
        ws.cell(row=5, column=c).alignment = CENTER

    for i, (k, v) in enumerate(GUIDE_ROWS):
        r = 6 + i
        ws.cell(row=r, column=1, value=k).font = BOLD_FONT
        ws.cell(row=r, column=2, value=v).font = CELL_FONT
        for c in (1, 2):
            ws.cell(row=r, column=c).alignment = WRAP
            ws.cell(row=r, column=c).border = THIN
            if i % 2 == 1:
                ws.cell(row=r, column=c).fill = FILL_ALT
        ws.row_dimensions[r].height = 36

    note_row = 6 + len(GUIDE_ROWS) + 1
    ws.cell(row=note_row, column=1, value="生成方式").font = BOLD_FONT
    ws.cell(
        row=note_row,
        column=2,
        value="python3 scripts/generate_energy_discourse_comparison.py",
    ).font = CELL_FONT
    ws.cell(row=note_row + 1, column=1, value="配套文档").font = BOLD_FONT
    ws.cell(
        row=note_row + 1,
        column=2,
        value="output/能量话语体系对照表_说明.docx（论述版，含判断标准与例析）",
    ).font = CELL_FONT
    for r in (note_row, note_row + 1):
        for c in (1, 2):
            ws.cell(row=r, column=c).border = THIN
            ws.cell(row=r, column=c).alignment = WRAP
            ws.cell(row=r, column=c).fill = FILL_NOTE

    set_widths(ws, [16, 80])


def build_xlsx():
    wb = Workbook()
    # 去掉默认空表，按固定顺序建表
    wb.remove(wb.active)
    build_guide_sheet(wb)
    build_overview_sheet(wb)
    build_matrix_sheet(wb)
    build_lists_sheet(wb)
    build_terms_sheet(wb)

    order = ["00_使用说明", "01_总览对照", "02_互译矩阵", "03_可译与不可通约", "04_术语对照"]
    for i, name in enumerate(order):
        wb.move_sheet(name, offset=i - wb.sheetnames.index(name))
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    wb.save(XLSX_PATH)


# ---------------------------------------------------------------------------
# DOCX 说明
# ---------------------------------------------------------------------------


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def add_para(
    doc,
    text,
    *,
    size=12,
    bold=False,
    space_before=0,
    space_after=6,
    first_line=False,
    color=None,
    align=None,
    font="宋体",
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, name=font, size=size, bold=bold, color=color)
    return p


def add_h(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12}
    return add_para(
        doc,
        text,
        size=sizes.get(level, 12),
        bold=True,
        space_before=14 if level == 1 else 10,
        space_after=6,
        color=RGBColor(0x1A, 0x3A, 0x5C),
        font="黑体",
    )


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    add_para(
        doc,
        "能量话语体系对照表",
        size=22,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
        color=RGBColor(0x1A, 0x3A, 0x5C),
        font="黑体",
    )
    add_para(
        doc,
        "物理学 · 生物学 · 心理学（意志力/心流）· 道家真气 · 灵性能量",
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
        color=RGBColor(0x55, 0x55, 0x55),
    )
    add_para(
        doc,
        "说明文档（与 Excel 对照表配套）",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
        color=RGBColor(0x66, 0x66, 0x66),
    )

    add_h(doc, "一、问题：同一个词，五套账本", 1)
    add_para(
        doc,
        "日常汉语里，「能量」可以指焦耳，也可以指卡路里，也可以指「今天状态不好、能量不足」，"
        "还可以指气功里的真气，或灵修话语中的气场与加持。五套用法共用一个能指，却并不共享同一套"
        "本体论、度量与真值条件。结果是：跨领域对话时常把隐喻走私成定律，或把信仰命题伪装成实验结论。",
        first_line=True,
    )
    add_para(
        doc,
        "本对照表的任务不是裁定谁「更真」，而是回答两个更可操作的问题：哪些说法可以互译（保留核心"
        "推断力），哪些说法完全不可通约（翻译即改写）。配套 Excel 给出矩阵与术语漂移表；本文给出判断标准与例析。",
        first_line=True,
    )

    add_h(doc, "二、五套话语各自在说什么", 1)

    portraits = [
        (
            "1. 物理学能量",
            "它是可做功的守恒标量，有单位、有测量协议、有转化定律。在封闭系中总量守恒，形式可变。"
            "这是工程与自然科学默认的「硬核」能量。把「我感觉被掏空」直接写成违反守恒，是范畴错误。",
        ),
        (
            "2. 生物学能量",
            "它落在代谢：营养氧化、ATP、体温与做功。它强烈可还原到物理化学，但叙述单位是有机体的"
            "维持与适应。卡路里是物理学能量在生命账本中的记账方式——因此物理↔生物是五对关系里互译最强的一对。",
        ),
        (
            "3. 心理学能量（意志力/心流）",
            "意志力、动机、注意力与心流，主要是构念与体验描述。「自我损耗」曾用能量隐喻组织实验，"
            "但复制与机制仍有争议；心流强调技能—挑战匹配与现象学，而非焦耳预算。可以说大脑耗能（生物），"
            "不能说意志力服从热力学守恒（物理）。",
        ),
        (
            "4. 道家真气",
            "真气属于传统身体—宇宙论：可体感、可修炼、可沿经络叙事运行，与精气神、内丹、导引相连。"
            "练习可产生可测的生理与心理效应，但「效应可测」不等于「真气＝电磁/热能」。在本体层，它与 SI 物理量不可通约。",
        ),
        (
            "5. 灵性能量",
            "这是最异质的一列：圣灵、普拉那、气场、昆达里尼、「宇宙高频能量」等权威来源不同。"
            "共同特征是超越日常物理操作定义，服务意义、救赎或觉醒。用实验室仪器「证明灵性」或用灵性"
            "否定焦耳，两边都在跨账本透支。",
        ),
    ]
    for title, body in portraits:
        add_h(doc, title, 2)
        add_para(doc, body, first_line=True)

    add_h(doc, "三、互译的判准（四级）", 1)
    add_para(
        doc,
        "Excel「互译矩阵」使用四级标签。判定时同时看三项：是否有共同操作定义、推断是否可迁移、"
        "翻译后损失的是否为该话语的核心承诺。",
        first_line=True,
    )
    levels = [
        "强可互译：可换算或可还原，损失可忽略。典型：物理焦耳 ↔ 生物卡路里/ATP。",
        "部分可译：局部机制或相关物稳定对应，整体不可还原。典型：血糖/睡眠 ↔ 自制表现；真气修炼 ↔ 呼吸—自主神经效应。",
        "弱类比：体验或修辞相似，强行等同会出错。典型：心流 ↔ 气机条达；敬畏 ↔ 灵性临在。",
        "不可通约：本体论与真值条件无法共享。典型：物理守恒 ↔ 灵性恩典；焦耳 ↔ 真气本体。",
    ]
    for line in levels:
        add_para(doc, "· " + line, space_after=4)

    add_h(doc, "四、哪些可以互译", 1)
    add_para(
        doc,
        "物理与生物之间，互译几乎是默认科研实践：营养学、运动生理学、体温调节都以热力学为底。"
        "需要警惕的不是互译本身，而是把「生命力」写成神秘的反熵盈余。",
        first_line=True,
    )
    add_para(
        doc,
        "生物与心理之间，只宜做局部对接：疲劳、炎症、睡眠、唤醒水平会影响意志与注意；"
        "心流也有可观察的生理相关。但目标表征、意义框架、技能—挑战匹配，无法被 ATP 计数穷尽。",
        first_line=True,
    )
    add_para(
        doc,
        "道家与灵性之间，存在「精微身体—修炼」的家族相似，可对照呼吸门控、次第与身心一如的实践语法。"
        "可译的是实践结构，不是把真气、普拉那与圣灵当成同一可加总的实体。",
        first_line=True,
    )
    add_para(
        doc,
        "道家/灵性实践指向生理与心理效应时，属于「效应层部分可译」：可以研究 HRV、焦虑量表、疼痛阈值；"
        "不可倒推「本体已被物理学捕获」。",
        first_line=True,
    )

    add_h(doc, "五、哪些完全不可通约", 1)
    add_para(
        doc,
        "物理能量与灵性能量：一方靠可重复测量与守恒，另一方靠启示、恩典或微妙场。把量子辞藻贴到疗愈广告上，"
        "不是科普，是借权威。",
        first_line=True,
    )
    add_para(
        doc,
        "物理能量与道家真气（本体层）：没有共同的操作定义，就没有互译，只有比喻。"
        "「丹田是电池」可以是教学隐喻，不能是物理命题。",
        first_line=True,
    )
    add_para(
        doc,
        "心理学的字面化能量与物理守恒：自我损耗叙事若被读成「脑内焦耳用完」，就越界了。"
        "允许说资源有限，不允许偷运热力学。",
        first_line=True,
    )
    add_para(
        doc,
        "灵性话语内部亦常不可通约：传统的权威、罪与救、空性与位格神，并不共享同一谓词体系。"
        "「宇宙能量」大熔炉会制造虚假共识。",
        first_line=True,
    )

    add_h(doc, "六、三个快速鉴别句式", 1)
    checks = [
        "它有单位吗？——有焦耳/千卡，更靠近物理/生物；只有体感与见证，则在真气/灵性侧。",
        "它守恒吗？——声称守恒却无封闭系定义，多半是隐喻或修辞。",
        "证据证明了哪一层？——证明「练了之后更放松」≠证明「气沿某经络以某速率流动」。",
    ]
    for line in checks:
        add_para(doc, "· " + line, space_after=4)

    add_h(doc, "七、对 AI 内容与跨界写作的用法", 1)
    add_para(
        doc,
        "选题上，可用本表做「概念清洁」：先标明话语坐标，再展开案例。叙事上，允许隐喻，但隐喻要贴标签。"
        "仲裁上，遇到「科学已证明气场」类句子，拆成：测到了什么相关物、主张了什么本体、两者是否被非法等同。",
        first_line=True,
    )
    add_para(
        doc,
        "配套文件：output/能量话语体系对照表.xlsx（总览、互译矩阵、可译与不可通约清单、术语漂移）。"
        "重新生成：python3 scripts/generate_energy_discourse_comparison.py",
        first_line=True,
    )

    add_para(
        doc,
        "（完）",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=18,
        color=RGBColor(0x88, 0x88, 0x88),
        size=10,
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def main():
    build_xlsx()
    build_docx()
    print(f"已生成：{XLSX_PATH}")
    print(f"已生成：{DOCX_PATH}")


if __name__ == "__main__":
    main()
