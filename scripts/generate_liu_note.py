#!/usr/bin/env python3
"""生成给刘孝春核对的垫付、欠薪与护理说明。Word/PDF 打印给刘；Markdown 另附家属自留。"""

from pathlib import Path
import html

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.colors import HexColor, white
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from comm_plan import CN_FONT_PATH
from content import LIU_NOTE, LIU_NOTE_FAMILY

CHINESE_FONT = "宋体"
HEADING_FONT = "黑体"
ACCENT = RGBColor(0x0B, 0x2F, 0x5B)
MUTED = RGBColor(0x5C, 0x6B, 0x7A)
RED = RGBColor(0xA6, 0x3D, 0x2F)

MD_NAME = "给刘孝春的垫付欠薪与护理说明_2026-08-19.md"
DOCX_NAME = "青岛抚顺路和哈尔滨路路口交通事故_给刘孝春的垫付与护理说明_20260819.docx"
PDF_NAME = "青岛抚顺路和哈尔滨路路口交通事故_给刘孝春的垫付与护理说明_20260819.pdf"


def _set_run_font(run, font_name: str, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _add_paragraph(
    doc: Document,
    text: str = "",
    *,
    font: str = CHINESE_FONT,
    size: float = 12,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: float | None = None,
    space_after: float = 6,
    space_before: float = 0,
    line_spacing: float = 1.4,
    color: RGBColor | None = None,
):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if text:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold, color=color)
    return p


def render_print_markdown() -> str:
    n = LIU_NOTE
    lines = [
        f"# {n['title']}",
        "",
        n["subtitle"],
        "",
        f"> {n['file_note']}",
        "",
        "## 这份纸要办的三件事",
        "",
    ]
    for i, item in enumerate(n["purpose"], 1):
        lines.append(f"{i}. {item}")
    lines.extend(["", "## 已经发生的事实", ""])
    for p in n["facts"]:
        lines.append(p)
        lines.append("")
    lines.extend(["## 三万元垫付收据必须写清的五句话", ""])
    for i, item in enumerate(n["receipt_five"], 1):
        lines.append(f"{i}. {item}")
    lines.extend(["", "## 未结工钱另写一张", ""])
    for item in n["wage_points"]:
        lines.append(f"- {item}")
    lines.extend(["", "## 护理费：医院一对一，合同加发票", ""])
    for item in n["nurse_points"]:
        lines.append(f"- {item}")
    lines.extend(["", "## 今天当场做", ""])
    for item in n["do_now"]:
        lines.append(f"- {item}")
    lines.extend(["", "## 现在不要签、不要写", ""])
    for item in n["do_not"]:
        lines.append(f"- {item}")
    lines.extend(["", "## 收据模板（医疗费垫付）", "", "```", *n["receipt_tpl"], "```", ""])
    lines.extend(["## 欠薪确认模板", "", "```", *n["wage_tpl"], "```", ""])
    lines.extend(["## 收据模板（护理费垫付，按次另开）", "", "```", *n["nurse_tpl"], "```", ""])
    lines.append("重新生成：`python3 scripts/build_all.py`。本文不是律师函。")
    lines.append("")
    return "\n".join(lines)


def render_markdown() -> str:
    fam = LIU_NOTE_FAMILY
    extra = [
        "",
        "---",
        "",
        f"## {fam['banner']}",
        "",
        "> 以下内容只写在本 Markdown。打印给刘孝春的 Word / PDF 不含本节，当面也不要念。",
        "",
    ]
    for p in fam["paras"]:
        extra.append(p)
        extra.append("")
    extra.append("内部测算不要发到群里，也不要说给骑手。")
    extra.append("")
    return render_print_markdown() + "\n".join(extra)


def build_markdown(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(render_markdown(), encoding="utf-8")


def _write_print_body_docx(doc: Document) -> None:
    n = LIU_NOTE
    _add_paragraph(
        doc,
        n["title"],
        font=HEADING_FONT,
        size=18,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
        line_spacing=1.3,
        color=ACCENT,
    )
    _add_paragraph(
        doc,
        n["subtitle"],
        font=HEADING_FONT,
        size=11,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
        line_spacing=1.3,
        color=MUTED,
    )
    _add_paragraph(doc, n["file_note"], bold=True, color=RED, space_after=10, line_spacing=1.35)

    _add_paragraph(doc, "这份纸要办的三件事", font=HEADING_FONT, size=14, bold=True, space_before=8, space_after=6, color=ACCENT, line_spacing=1.3)
    for item in n["purpose"]:
        _add_paragraph(doc, f"• {item}", first_line_indent=0.37, space_after=3)

    _add_paragraph(doc, "已经发生的事实", font=HEADING_FONT, size=14, bold=True, space_before=10, space_after=6, color=ACCENT, line_spacing=1.3)
    for p in n["facts"]:
        _add_paragraph(doc, p, first_line_indent=0.74, space_after=5)

    _add_paragraph(doc, "三万元垫付收据必须写清的五句话", font=HEADING_FONT, size=14, bold=True, space_before=10, space_after=6, color=ACCENT, line_spacing=1.3)
    for i, item in enumerate(n["receipt_five"], 1):
        _add_paragraph(doc, f"{i}. {item}", first_line_indent=0.37, space_after=3)

    _add_paragraph(doc, "未结工钱另写一张", font=HEADING_FONT, size=14, bold=True, space_before=10, space_after=6, color=ACCENT, line_spacing=1.3)
    for item in n["wage_points"]:
        _add_paragraph(doc, f"• {item}", first_line_indent=0.37, space_after=3)

    _add_paragraph(doc, "护理费：医院一对一，合同加发票", font=HEADING_FONT, size=14, bold=True, space_before=10, space_after=6, color=ACCENT, line_spacing=1.3)
    for item in n["nurse_points"]:
        _add_paragraph(doc, f"• {item}", first_line_indent=0.37, space_after=3)

    _add_paragraph(doc, "今天当场做", font=HEADING_FONT, size=14, bold=True, space_before=10, space_after=6, color=ACCENT, line_spacing=1.3)
    for item in n["do_now"]:
        _add_paragraph(doc, f"• {item}", first_line_indent=0.37, space_after=3)

    _add_paragraph(doc, "现在不要签、不要写", font=HEADING_FONT, size=14, bold=True, space_before=10, space_after=6, color=ACCENT, line_spacing=1.3)
    for item in n["do_not"]:
        _add_paragraph(doc, f"• {item}", first_line_indent=0.37, space_after=3, color=RED)

    _add_paragraph(doc, "收据模板（医疗费垫付）", font=HEADING_FONT, size=14, bold=True, space_before=10, space_after=6, color=ACCENT, line_spacing=1.3)
    for line in n["receipt_tpl"]:
        _add_paragraph(doc, line, size=11, space_after=2, line_spacing=1.25)

    _add_paragraph(doc, "欠薪确认模板", font=HEADING_FONT, size=14, bold=True, space_before=10, space_after=6, color=ACCENT, line_spacing=1.3)
    for line in n["wage_tpl"]:
        _add_paragraph(doc, line, size=11, space_after=2, line_spacing=1.25)

    _add_paragraph(doc, "收据模板（护理费垫付，按次另开）", font=HEADING_FONT, size=14, bold=True, space_before=10, space_after=6, color=ACCENT, line_spacing=1.3)
    for line in n["nurse_tpl"]:
        _add_paragraph(doc, line, size=11, space_after=2, line_spacing=1.25)

    _add_paragraph(
        doc,
        "本文根据 2026 年 8 月 14–19 日材料整理。重新生成：python3 scripts/build_all.py。本文不是律师函。",
        size=10.5,
        color=MUTED,
        space_before=12,
    )


def build_document(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = CHINESE_FONT
    normal_style.font.size = Pt(12)
    rpr = normal_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CHINESE_FONT)

    _write_print_body_docx(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _pdf_font() -> str:
    kwargs = {"subfontIndex": 0} if CN_FONT_PATH.endswith(".ttc") else {}
    pdfmetrics.registerFont(TTFont("CN", CN_FONT_PATH, **kwargs))
    return "CN"


def _pdf_styles(font: str) -> dict[str, ParagraphStyle]:
    navy = HexColor("#0B2F5B")
    muted = HexColor("#5C6B7A")
    red = HexColor("#A63D2F")
    dark = HexColor("#1F2A37")
    return {
        "title": ParagraphStyle("t", fontName=font, fontSize=15, leading=22, textColor=navy, alignment=TA_CENTER, spaceAfter=4),
        "sub": ParagraphStyle("s", fontName=font, fontSize=9.5, leading=14, textColor=muted, alignment=TA_CENTER, spaceAfter=6),
        "warn": ParagraphStyle("w", fontName=font, fontSize=9.5, leading=15, textColor=red, alignment=TA_LEFT, spaceAfter=8),
        "h1": ParagraphStyle("h", fontName=font, fontSize=12, leading=18, textColor=navy, spaceBefore=8, spaceAfter=4),
        "body": ParagraphStyle("b", fontName=font, fontSize=10, leading=16, textColor=dark, alignment=TA_JUSTIFY, firstLineIndent=20, spaceAfter=4),
        "bullet": ParagraphStyle("u", fontName=font, fontSize=10, leading=15, textColor=dark, leftIndent=12, spaceAfter=2),
        "red_bullet": ParagraphStyle("r", fontName=font, fontSize=10, leading=15, textColor=red, leftIndent=12, spaceAfter=2),
        "tpl": ParagraphStyle("p", fontName=font, fontSize=9.5, leading=14, textColor=dark, leftIndent=8, spaceAfter=1),
        "foot": ParagraphStyle("f", fontName=font, fontSize=8, leading=12, textColor=muted, spaceBefore=8),
    }


def _p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(html.escape(str(text)).replace("\n", "<br/>"), style)


def _header_footer(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(HexColor("#0B2F5B"))
    canvas.rect(0, A4[1] - 12 * mm, A4[0], 12 * mm, fill=1, stroke=0)
    canvas.setFillColor(white)
    canvas.setFont("CN", 8)
    canvas.drawString(18 * mm, A4[1] - 8 * mm, "青岛抚顺路和哈尔滨路路口交通事故 · 给刘孝春的垫付与护理说明")
    canvas.setFillColor(HexColor("#C4A35A"))
    canvas.rect(0, 0, A4[0], 10 * mm, fill=1, stroke=0)
    canvas.setFillColor(HexColor("#0B2F5B"))
    canvas.setFont("CN", 8)
    canvas.drawString(18 * mm, 4 * mm, "不是律师函 · 不是了结协议 · 现阶段不谈总赔偿数额")
    canvas.drawRightString(A4[0] - 18 * mm, 4 * mm, f"第 {doc.page} 页")
    canvas.restoreState()


def build_pdf(output_path: Path) -> None:
    n = LIU_NOTE
    font = _pdf_font()
    styles = _pdf_styles(font)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=15 * mm,
        title=n["title"],
        author="内部整理",
    )
    story = [
        _p(n["title"], styles["title"]),
        _p(n["subtitle"], styles["sub"]),
        _p(n["file_note"], styles["warn"]),
        _p("这份纸要办的三件事", styles["h1"]),
    ]
    for item in n["purpose"]:
        story.append(_p("• " + item, styles["bullet"]))
    story.append(_p("已经发生的事实", styles["h1"]))
    for para in n["facts"]:
        story.append(_p(para, styles["body"]))
    story.append(_p("三万元垫付收据必须写清的五句话", styles["h1"]))
    for i, item in enumerate(n["receipt_five"], 1):
        story.append(_p(f"{i}. {item}", styles["bullet"]))
    story.append(_p("未结工钱另写一张", styles["h1"]))
    for item in n["wage_points"]:
        story.append(_p("• " + item, styles["bullet"]))
    story.append(_p("护理费：医院一对一，合同加发票", styles["h1"]))
    for item in n["nurse_points"]:
        story.append(_p("• " + item, styles["bullet"]))
    story.append(_p("今天当场做", styles["h1"]))
    for item in n["do_now"]:
        story.append(_p("• " + item, styles["bullet"]))
    story.append(_p("现在不要签、不要写", styles["h1"]))
    for item in n["do_not"]:
        story.append(_p("• " + item, styles["red_bullet"]))
    story.append(_p("收据模板（医疗费垫付）", styles["h1"]))
    for line in n["receipt_tpl"]:
        story.append(_p(line, styles["tpl"]))
    story.append(_p("欠薪确认模板", styles["h1"]))
    for line in n["wage_tpl"]:
        story.append(_p(line, styles["tpl"]))
    story.append(_p("收据模板（护理费垫付，按次另开）", styles["h1"]))
    for line in n["nurse_tpl"]:
        story.append(_p(line, styles["tpl"]))
    story.append(Spacer(1, 6))
    story.append(_p("本文根据 2026 年 8 月 14–19 日材料整理。重新生成：python3 scripts/build_all.py。本文不是律师函。", styles["foot"]))
    doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "deliverables"
    build_markdown(out / MD_NAME)
    build_document(out / DOCX_NAME)
    build_pdf(out / PDF_NAME)
    print(f"Wrote {out / MD_NAME}")
    print(f"Wrote {out / DOCX_NAME}")
    print(f"Wrote {out / PDF_NAME}")


if __name__ == "__main__":
    main()
