# -*- coding: utf-8 -*-
"""
生成《住房即服务：医疗、养老与生产效率提升的 2030 展望》白皮书。

署名：河南大学住房政策研究中心
运行：python3 scripts/build_charts.py && python3 scripts/build_whitepaper.py
输出：dist/河南大学住房政策研究中心_住房即服务_医疗养老与提效服务2030展望白皮书.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DIST = ROOT / "dist"
CHARTS = ROOT / "whitepaper" / "assets" / "charts"
DIST.mkdir(parents=True, exist_ok=True)
OUT = DIST / "河南大学住房政策研究中心_住房即服务_医疗养老与提效服务2030展望白皮书.docx"

NAVY = RGBColor(0x1A, 0x2A, 0x4A)
DARK = RGBColor(0x1B, 0x2A, 0x3A)
GRAY = RGBColor(0x5A, 0x64, 0x6E)
ACCENT = RGBColor(0x8B, 0x1A, 0x1A)
OCHRE = RGBColor(0xC4, 0x5C, 0x26)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "8B1A1A"
TABLE_ALT_BG = "F7F4EF"


def set_run_font(run, size=12, bold=False, color=DARK, font="宋体", east_asia=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman" if font in ("宋体", "楷体", "黑体") else font)
    rFonts.set(qn("w:hAnsi"), "Times New Roman" if font in ("宋体", "楷体", "黑体") else font)
    rFonts.set(qn("w:eastAsia"), east_asia or font)


def set_cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="C5CDD6"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_page_number(paragraph):
    run1 = paragraph.add_run("— ")
    set_run_font(run1, size=9, color=GRAY, font="宋体")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    set_run_font(run, size=9, color=GRAY, font="Times New Roman")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    r.append(fld_end)
    run2 = paragraph.add_run(" —")
    set_run_font(run2, size=9, color=GRAY, font="宋体")


def setup_header_footer(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run("河南大学住房政策研究中心  ·  住房即服务白皮书（2030 展望）")
    set_run_font(run, size=9, color=NAVY, font="宋体")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)


def add_cover_line(doc, text, size=12, bold=False, color=DARK, font="宋体", space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, font=font)
    return p


def add_heading_cn(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    size = 16 if level == 1 else 13
    color = NAVY if level == 1 else ACCENT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color, font="黑体")
    return p


def add_body(doc, text, first_indent=True, bold=False, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=DARK, font="宋体")
    return p


def add_quote(doc, text, source=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(text)
    set_run_font(run, size=11, italic=True, color=NAVY, font="楷体")
    if source:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(10)
        p2.paragraph_format.right_indent = Cm(0.5)
        r2 = p2.add_run(source)
        set_run_font(r2, size=10, color=GRAY, font="楷体")


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Cm(0.74)
        run = p.add_run("●  " + item)
        set_run_font(run, size=12, color=DARK, font="宋体")


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=10, color=GRAY, font="楷体")
    return p


def add_image(doc, name, caption, width_cm=15.4):
    path = CHARTS / name
    if not path.exists():
        add_body(doc, f"（配图未生成：{name}，请先运行 python3 scripts/build_charts.py）", first_indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=WHITE, font="黑体")
        set_cell_shading(cell, TABLE_HEADER_BG)
        set_cell_borders(cell, "8B1A1A")
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, size=10, color=DARK, font="宋体")
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)
            set_cell_borders(cell)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_page_break(doc):
    doc.add_page_break()


def build_cover(doc):
    for _ in range(2):
        doc.add_paragraph()
    add_cover_line(doc, "河南大学住房政策研究中心", size=16, bold=True, color=NAVY, font="黑体", space_after=4)
    add_cover_line(
        doc,
        "Housing Policy Research Center, Henan University",
        size=11,
        color=GRAY,
        font="Times New Roman",
        space_after=18,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("HPRC 研 究 报 告  ·  2026 年 第 1 号")
    set_run_font(run, size=12, bold=True, color=ACCENT, font="黑体")

    add_cover_line(doc, "住房即服务", size=28, bold=True, color=NAVY, font="黑体", space_before=10, space_after=8)
    add_cover_line(
        doc,
        "医疗、养老与生产效率提升的空间载体\n面向 2030 年的白皮书",
        size=18,
        bold=True,
        color=DARK,
        font="楷体",
        space_after=16,
    )
    add_cover_line(
        doc,
        "Housing as a Service:\nHealthcare, Elderly Care, and Productivity-Enhancing Services toward 2030",
        size=11,
        color=GRAY,
        font="Times New Roman",
        space_after=26,
    )
    add_cover_line(
        doc,
        "结合“十五五”开局形势、人工智能+行动与 2026 年最新科技进展的综合研究",
        size=11,
        color=DARK,
        font="宋体",
        space_after=8,
    )
    add_cover_line(doc, "2026 年 8 月", size=14, bold=True, color=NAVY, font="黑体", space_before=22, space_after=6)
    add_cover_line(doc, "开封 · 郑州", size=12, color=DARK, font="宋体", space_after=6)
    add_page_break(doc)


def build_statement(doc):
    add_heading_cn(doc, "说  明", level=1)
    add_body(
        doc,
        "本白皮书由河南大学住房政策研究中心组织撰写。中心以住房制度、社区配套与城市更新为研究主线，关注居住空间如何承载医疗、养老和生产效率提升等公共服务与市场服务。选择这一题目，不是把住房政策研究“跨界”到卫生、民政或工业部门的专业领地，而是确认一个正在发生的事实：医疗服务、养老服务和提效服务，最终都要落到一套房子、一个社区、一座园区里才能被居民和企业使用。",
    )
    add_body(
        doc,
        "2026 年是“十五五”开局之年，也是具身智能被产业界称为“量产元年”“应用元年”的年份。国务院《关于深入实施“人工智能+”行动的意见》、国家卫生健康委等部门《关于促进和规范“人工智能+医疗卫生”应用发展的实施意见》、中办国办《关于加快建立长期护理保险制度的意见》、国务院《城市更新“十五五”规划》，以及网信部门《智能体规范应用与创新发展实施意见》，在不到一年时间里密集落地。与此同时，世界人工智能大会（WAIC 2026）上的康养机器人、郑州高新区下线的“河南造”人形机器人、3C 产线连续作业实证，使技术叙事第一次有了可核对的现场证据。住房政策如果仍把这些变化视为“远景附录”，将错过 2026—2030 年最关键的制度窗口。",
    )
    add_body(
        doc,
        "白皮书所依据的事实材料，主要来自政府部门公开发布的规划、意见、统计公报与标准，以及 2025 年下半年至 2026 年 8 月的公开新闻报道。文中涉及的企业产品、产线数据、融资规模等均转引自公开报道，供政策讨论与学术研究参考，不构成对任何企业或投资标的的背书。展望 2030 年的部分采用情景分析，明确区分“已公布的政策目标”“可观察的技术进展”和“本中心的判断”，避免把愿望写成预测。",
    )
    add_body(
        doc,
        "河南是观察样本，也是政策对象。作为人口大省、粮食大省和制造业门类较为完整的经济大省，河南同时面对老龄化加速、劳动力供给趋紧、房地产转向存量更新、县域城乡服务不均衡等多重约束。把全国政策翻译成可在中原城市、县城和乡村落地的空间安排，是本中心写作本文的出发点。",
    )
    add_body(doc, "文责自负。欢迎学界、产业界与政策部门批评指正。")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("河南大学住房政策研究中心\n2026 年 8 月")
    set_run_font(run, size=12, color=DARK, font="楷体")
    add_page_break(doc)


def build_toc(doc):
    add_heading_cn(doc, "目  录", level=1)
    items = [
        "摘要：2026 年的五个判断与 2030 年的一条主线",
        "一、导论：住房政策为何必须同时讨论医疗、养老与提效",
        "二、当下形势：三重压力在居住空间里交汇",
        "三、科技与制度的交汇：2025—2026 年的关键进展",
        "四、医疗服务展望 2030：从医院能力到家庭可达",
        "五、养老服务展望 2030：从床位思维到支付—空间—机器人闭环",
        "六、生产效率服务展望 2030：工厂、社区与住房如何共同提效",
        "七、住房作为三类服务的空间操作系统",
        "八、四个情景与河南含义",
        "九、政策建议",
        "附录一  2025—2030 年重要政策与技术节点",
        "附录二  主要公开资料来源",
        "附录三  术语简释",
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(item)
        set_run_font(run, size=12, color=DARK, font="宋体")
    add_page_break(doc)


def build_abstract(doc):
    add_heading_cn(doc, "摘要：2026 年的五个判断与 2030 年的一条主线", level=1)
    add_body(
        doc,
        "住房正在从“资产与居住容器”转向“医疗、养老与提效服务的操作系统”。这一转变不是修辞，而是 2026 年政策、技术与人口结构三股力量同时加压的结果。国务院把人工智能全面赋能高质量发展的时间表写到 2030 年：新一代智能终端、智能体应用普及率到 2027 年超 70%、到 2030 年超 90%。卫生健康部门同步给出更硬的行业目标：到 2030 年基层诊疗智能辅助应用基本实现全覆盖，二级以上医院普遍开展医学影像智能辅助诊断和临床诊疗智能辅助决策。养老一侧，医养结合促进行动要求到 2027 年底基本实现县域全覆盖；长期护理保险则被明确为用三年左右时间建制、到 2028 年底全国基本全面覆盖。城市更新“十五五”规划又把完整社区、社区嵌入式服务设施、好房子和存量闲置房屋补齐民生短板，写成到 2030 年必须交出进展的空间工程。四条政策轨道在居住空间汇合。",
    )
    add_body(doc, "本白皮书将观察收敛为五个判断：")
    add_bullets(
        doc,
        [
            "服务供给的瓶颈，已经从“缺机构、缺床位”转向“缺可进入家庭和社区的稳定服务包”。医疗和养老的主战场是居家与社区，提效服务的主战场是工厂、仓储、物业和最后一公里。住房与社区是它们共同的物理接口。",
            "2026 年技术拐点真实存在，但不等于 2026 年全面替代人力。WAIC 上的康养具身、腾讯“小六”按摩机器人、智元 3C 产线连续作业、郑州众擎 T800 量产下线，证明工程化开始发生；规模化入户、规模化上门护理仍需支付、标准、安全和空间改造同步。",
            "河南是三类服务需求同时爆发的典型省份：常住人口近 9800 万，城镇化率刚过 59%，人口自然增长率已转为负，房地产投资连续下行，工业机器人等新产品产量高增。人口规模、老龄化、劳动力约束和存量住房更新，在同一套社区里叠加。",
            "2030 年的基准情景不是“家家有人形保姆、厂厂无人化”，而是分层渗透：工业与仓储率先形成可复制作业单元；康养机构与社区驿站形成人机协同服务包；家庭端以健康监测、远程问诊、家庭养老床位和轻量化辅助为主。",
            "住房政策的任务，是把三类服务写进户型、公区、完整社区、老旧小区改造、保障性住房和产业园区配套，而不是等待技术成熟后再被动适应。机器人可达性、数字家庭、急救响应、数据权利和支付接口，应成为 2026—2030 年住房制度的新基建。",
        ],
    )
    add_body(
        doc,
        "一条主线贯穿全文：以住房和社区为底座，把“病有所医、老有所养、产有所效”写成同一套空间制度。医疗解决可达与连续，养老解决照护与支付，提效解决劳动力约束下的服务供给。三者分开规划，必然在小区门口互相打架；三者以住房为操作系统，才可能在 2030 年形成可复制的中原方案。",
    )
    add_body(
        doc,
        "关键词：住房即服务；社区嵌入式设施；医养结合；长期护理保险；人工智能+；具身智能；完整社区；生产效率；2030 展望；河南",
        first_indent=False,
    )


def build_ch1(doc):
    add_heading_cn(doc, "一、导论：住房政策为何必须同时讨论医疗、养老与提效", level=1)
    add_heading_cn(doc, "（一）问题的提出", level=2)
    add_body(
        doc,
        "传统住房政策处理土地、金融、保障、市场调控和居住福利。医疗归卫生健康，养老归民政与医保，生产效率归工业和科技。部门分割在文件上清晰，在居民生活里并不存在。一个失能老人是否住得下去，取决于户型能不能进轮椅和护理车、社区有没有助餐和日间照料、家庭病床能否上门、长护险能否支付、物业能否协助呼叫。一条产线能不能把机器人用起来，取决于厂房布局，也取决于工人是否住得近、生活服务是否把通勤和照料负担降下来。住房是这些服务发生的场所，也是这些服务失败时矛盾最集中的场所。",
    )
    add_body(
        doc,
        "2026 年的新变化在于：服务不再只由人和机构提供，开始由软件智能体和具身机器人分担。国务院《人工智能+》行动明确提出，探索推广人人可享的高水平居民健康助手，发挥人工智能在精神慰藉陪伴、养老托育助残中的作用，并拓展人工智能在“好房子”全生命周期的应用。网信部门 2026 年 5 月发布的智能体实施意见，把医疗健康、商业服务中的家政养老托育助残、智能制造和城市治理中的房屋管理，写成可以部署的应用方向。技术一旦进入住宅和社区，住房规范就必须回答门宽、电梯、网络、隐私、责任和公平问题。",
    )
    add_heading_cn(doc, "（二）“住房即服务”的分析框架", level=2)
    add_body(
        doc,
        "本白皮书使用“住房即服务”（Housing as a Service）作为分析框架。它包含三层含义。第一，住房是服务接入点：居民获得医疗、养老、家政、物流、物业的第一入口，越来越多发生在家门口和户门内。第二，住房是服务的生产空间：家庭养老床位、家庭病床、远程办公、社区工坊、园区配套宿舍，使住宅和社区同时成为劳动与照护的现场。第三，住房是服务的制度接口：规划配建、设施权属、运营补贴、长护险和医保支付、数据授权，都要通过住房和社区才能闭环。",
    )
    add_image(doc, "chart01_architecture.png", "图1  住房作为三类服务的空间操作系统（2030 目标形态）")
    add_body(
        doc,
        "在这一框架下，医疗服务回答的是“人在家中能否被连续照护和及时救治”，养老服务回答的是“失能与半失能状态能否在熟悉环境中维持尊严”，提效服务回答的是“在劳动力减少的约束下，工厂、社区和家庭的必要劳动如何被稳定完成”。三者共享同一套空间约束：存量小区、新建住宅、产业园区和县域乡村的房屋条件差异极大。政策如果只鼓励技术下乡、入户、进厂，却不改造空间，技术会停在展厅。",
    )
    add_heading_cn(doc, "（三）研究方法与边界", level=2)
    add_body(
        doc,
        "方法上，本文采取“政策文本对照 + 公开统计 + 2026 年科技现场观察 + 住房空间推演”。政策文本以 2025 年 8 月至 2026 年 8 月中央和部门文件为主；统计以河南省 2024 年国民经济和社会发展统计公报等官方数据为主；科技现场以 WAIC 2026、郑州机器人量产报道、工业产线实证和已发布的行业标准为主。住房空间推演把上述事实翻译为户型、公区、配建、更新和治理议题。",
    )
    add_body(
        doc,
        "边界需要事先声明。本中心不替代临床医学评价、不预测某家人形机器人公司的市占率，也不把河南的局部试点写成全国平均水平。文中“渗透率”“压力指数”等示意图，仅用于比较情景和讨论优先级。对尚未被官方统计确认的企业数据，一律标注为公开报道。",
    )


def build_ch2(doc):
    add_heading_cn(doc, "二、当下形势：三重压力在居住空间里交汇", level=1)
    add_heading_cn(doc, "（一）全国：存量时代的服务缺口", level=2)
    add_body(
        doc,
        "中国城市发展已从大规模增量扩张转向存量提质。国务院《城市更新“十五五”规划》提出，到 2030 年城市更新行动取得重要进展，开发建设方式转型初见成效；规划明确改造城镇老旧小区约 11.5 万个、城中村约 4000 个，并加快建设完整社区、完善社区嵌入式服务设施、推动一刻钟便民生活圈扩围升级。这意味着未来五年住房工作的主战场，是既有小区如何补齐养老、托育、健康和便民功能，而不是继续以新房销售规模衡量居住进步。",
    )
    add_body(
        doc,
        "与空间转向并行的是人口转向。公开报道援引的全国数据表明，截至 2025 年底，全国 60 岁及以上老年人口已超过 3.1 亿，失能、半失能老年人规模超过 4000 万。家庭小型化、空巢化使“一人失能、全家失衡”成为高频风险。医疗服务的矛盾，不再只是大医院号源，而是慢病管理、康复、安宁疗护和上门服务能否在社区连续提供。生产效率的矛盾，也不再只是资本深化，而是劳动年龄人口下降后，制造、物流、家政和照护岗位如何维持供给。",
    )
    add_heading_cn(doc, "（二）河南：大省约束下的典型样本", level=2)
    add_body(
        doc,
        "河南省统计局和国家统计局河南调查总队发布的《2024 年河南省国民经济和社会发展统计公报》显示：2024 年全省地区生产总值 63589.99 亿元，增长 5.1%；年末常住人口 9785 万人，城镇常住人口 5795 万人，城镇化率 59.22%；人口出生率 7.78‰，死亡率 7.89‰，自然增长率 -0.11‰。河南已经进入人口自然负增长，城镇化仍有上升空间，城乡服务能力差距不会自动消失。",
    )
    add_body(
        doc,
        "住房市场同步转入调整。2024 年全省房地产开发投资 3908.41 亿元，下降 7.5%；新建商品房销售面积 6202.80 万平方米，下降 11.0%。与此同时，城镇保障性安居工程住房基本建成 41.73 万套、新开工 76.79 万套。增量减速、保障加力、存量更新，构成河南住房政策的新三角。南阳市住建部门探索“急救保命 + 生活服务 + 中医康养”的居家养老住房试点，把急救响应压缩到约 4 分钟，说明地方已经意识到：住房产品若不能内嵌健康服务，去化与居住品质都会失分。",
    )
    add_body(
        doc,
        "医疗卫生资源规模不小，但结构偏“机构内”。2024 年末全省医疗卫生机构 85539 个，床位 76.21 万张，卫生技术人员 88.61 万人；全年总诊疗 6.84 亿人次，其中基层医疗卫生机构诊疗 3.98 亿人次。基层已经承担可观诊疗量，但家庭病床、上门巡诊、社区康复和智慧健康管理仍不足以覆盖快速增长的老年需求。社会保障方面，职工医保参保 1462.40 万人，城乡居民医保参保 8266.49 万人，覆盖面广；真正决定居家照护能否运转的，将是长护险建制后的服务供给，而不是参保人数本身。",
    )
    add_body(
        doc,
        "生产效率一侧，河南同时具备“必须提效”和“能够提效”的条件。2024 年规模以上工业增加值增长 8.1%，高技术制造业增长 12.0%，工业机器人等新产品产量增长 15.6%；农村劳动力转移就业总量仍达 3064.97 万人。劳动密集型产业的用工成本上升与技能人才不足，正在倒逼智能制造。郑州引入众擎、智元、优必选等具身智能企业，2026 年 7 月众擎 T800 在郑州高新区量产下线，公开报道称基地导入自动化后整体生产效率提升 40%。河南 41 个工业大类，是人工智能落地的试验场，也是职住服务必须跟上的压力源。",
    )
    add_image(doc, "chart05_henan_pressure.png", "图5  河南作为观察样本：五类结构性压力同时存在")
    add_heading_cn(doc, "（三）社区设施：有场所，缺服务闭环", level=2)
    add_body(
        doc,
        "河南在设施覆盖上并不落后。公开报道显示，近年全省街道综合养老服务设施和社区养老服务场所已实现较高覆盖，智慧养老平台录入大量老年人信息，新建小区养老服务设施“四同步”经过专项清理后达标率明显提升。问题在于：有场所不等于有稳定运营，有平台不等于有上门服务，有配建面积不等于面积被真正用于医养而不是改作他用。完整社区试点、老旧小区补建、家庭养老床位和“社区 + 物业 + 养老服务”，仍处于从覆盖走向闭环的中途。",
    )
    add_body(
        doc,
        "《河南省深化养老服务改革发展三年行动计划》要求，2026 年 6 月底前市、县级完成养老服务设施专项规划并与国土空间总体规划、详细规划衔接，同时制定统筹利用存量资产建设普惠养老设施的政策。时间表本身说明：空间供给仍是短板，而且短板必须在“十五五”前期补上，否则 2030 年的智能服务将没有落地的房间。",
    )


def build_ch3(doc):
    add_heading_cn(doc, "三、科技与制度的交汇：2025—2026 年的关键进展", level=1)
    add_body(
        doc,
        "把 2030 年写成可讨论的对象，前提是承认 2025—2026 年已经出现一组可核对的节点。它们不是“未来可能发生”，而是“已经写进文件或已经出现在产线和展馆”。",
    )
    add_image(doc, "chart02_timeline.png", "图2  2025—2030 关键政策与技术节点")
    add_heading_cn(doc, "（一）人工智能成为国家行动，医疗与养老被点名", level=2)
    add_body(
        doc,
        "2025 年 8 月，国务院印发《关于深入实施“人工智能+”行动的意见》（国发〔2025〕11 号）。文件提出：到 2027 年率先实现人工智能与六大重点领域广泛深度融合，新一代智能终端、智能体等应用普及率超 70%；到 2030 年人工智能全面赋能高质量发展，上述普及率超 90%，智能经济成为重要增长极。民生条款直接点到医疗、养老和住房：推广居民健康助手，有序推动辅助诊疗、健康管理、医保服务应用；发挥人工智能在精神慰藉、养老托育助残中的作用；拓展人工智能在“好房子”全生命周期的应用。",
    )
    add_body(
        doc,
        "2025 年 11 月，国家卫生健康委等五部门发布《关于促进和规范“人工智能+医疗卫生”应用发展的实施意见》。到 2027 年，要形成一批临床专病专科垂直大模型和智能体应用，基层诊疗智能辅助和患者就诊智能服务广泛应用；到 2030 年，基层诊疗智能辅助应用基本实现全覆盖，推动二级以上医院普遍开展医学影像智能辅助诊断、临床诊疗智能辅助决策，标准规范体系基本完善。文件把健康管理、养老和托育智能服务单独列出，并强调人工智能赋能而不替代。对住房政策而言，这意味着 2030 年的社区卫生站和家庭健康终端，将是国家考核的应用现场，而不仅是地方创新试点。",
    )
    add_heading_cn(doc, "（二）智能体进入物理世界，康养与制造同时被打开", level=2)
    add_body(
        doc,
        "2026 年 5 月 8 日，网信部门发布《智能体规范应用与创新发展实施意见》。智能体被定义为具备自主感知、记忆、决策、交互与执行能力的智能系统。意见在产业发展中部署生产管理智能体、工艺优化与缺陷识别；在商业服务中明确探索通过具身智能体提供低成本家政、养老、托育、助残等服务；在医疗健康中部署影像分析、诊断推理、预问诊、报告解析、药品与病历管理；在城市治理中支撑智能建造、房屋管理和基础设施安全运行。安全条款同步强调：用户对智能体自主决策享有知情权和最终决策权，防范老年人沉迷与情感依赖，医疗等领域可制定强制性标准。",
    )
    add_body(
        doc,
        "2026 年世界人工智能大会把这一方向从文件推进到可参观的现场。产业观察普遍将 2026 年称为具身智能应用元年，评价尺度从“能不能动”转向连续作业、异常恢复和场景适配成本。腾讯发布具身智能基座模型，示范机器人“小六”学习中医按摩手法，团队公开表示希望放到养老院真实业务场景继续迭代；其 HyVLA-0.5 在日化工厂面对高混合、小批量产线，公开报道称作业成功率高于 95%、节拍快于 6 秒/件。傅利叶等康养方向企业则强调：优先把成熟能力落到康复训练、认知训练等高频场景，由真实场景反向驱动产品。这些案例的政策含义是：康养不再是机器人的“终局故事”，而是 2026 年已经开始的工程化赛道。",
    )
    add_heading_cn(doc, "（三）工业提效给出可核对的产线证据", level=2)
    add_body(
        doc,
        "公开报道称，智元精灵 G2 在 3C 工业产线完成 8 小时实证，联合端侧算力平台把现场调试周期大幅压缩，换型重训时间以小时计。人形机器人自动化产线被报道为可实现约每 30 分钟一台组装下线，较传统手工模式生产效率提升超过 50%。与此同时，专业人士提醒：量产元年不等于商业化元年，成本、场景适配和稳定性仍是鸿沟，全行业规模复制可能还需要一到两年。本中心采纳这一冷静判断：2030 年之前，工业场景会率先形成付费闭环，但不会在所有行业同步铺开。",
    )
    add_body(
        doc,
        "对河南尤其重要的是制造端。2026 年 7 月 24 日，众擎 T800 在郑州高新区云智科技园基地首批下线发往全球，公开报道称一期约 8000 平方米全闭环产线、生产效率提升 40%，出厂前完成数十项质检与工况模拟。智元中部具身智能产业基地、优必选相关布局也已落子郑州。河南省还发布人工智能应用场景机会清单，覆盖产业、民生与治理。这意味着“提效服务”对河南不是外生概念，而是正在进入本地厂房和本地就业结构的现实。",
    )
    add_heading_cn(doc, "（四）支付、标准与空间制度同步补位", level=2)
    add_body(
        doc,
        "没有支付的服务只能停留在演示。2026 年 3 月，中办、国办印发《关于加快建立长期护理保险制度的意见》，要求用三年左右时间基本建立覆盖全民、统筹城乡的长期护理保险制度，到 2028 年底全国基本全面覆盖。起步阶段保障重度失能人员，费率总体控制在 0.3% 左右；国家统一长护服务项目目录纳入 36 项，鼓励使用居家和社区护理并在支付比例上适当倾斜。居家护理一旦可报销，家庭养老床位、物业上门、嵌入式机构和护理机器人租赁，才具备可持续的需求侧基础。",
    )
    add_body(
        doc,
        "标准方面，国际电工委员会已发布由中国牵头制定的 IEC 63310《互联家庭环境下使用的主动辅助生活机器人性能准则》；国内《医养结合健康管理服务标准》（WS/T 876—2026）等卫生行业标准于 2026 年 9 月 1 日起施行。空间方面，《城市更新“十五五”规划》要求充分利用存量闲置房屋和低效用地优先补齐民生设施，支持有条件的医疗卫生机构开展医养结合，推动物业服务向养老、托育、家政、健康等领域延伸。科技、支付、标准、空间四条线第一次在同一时间窗口对齐。这是展望 2030 年时最重要的制度事实。",
    )
    add_table(
        doc,
        ["领域", "2026 年可观察状态", "文件或现场给出的 2030 前后目标"],
        [
            ["人工智能总体", "智能体实施意见发布，展会进入工程化", "2030 年智能终端/智能体普及率超 90%"],
            ["医疗卫生", "垂直模型与辅助诊疗加速进院", "2030 年基层智能辅助基本全覆盖"],
            ["养老照护", "医养结合促进行动推进，康养机器人试运行", "2027 年县域医养结合基本覆盖；2028 年长护险基本全覆盖"],
            ["城市住房", "完整社区、嵌入式设施、好房子成为更新主线", "2030 年城市更新取得重要进展，存量服务设施补齐"],
            ["工业提效", "头部产线实证与河南造人形机器人下线", "规模复制取决于成本、稳定性和场景开放"],
        ],
        col_widths=[2.6, 6.2, 6.4],
    )
    add_caption(doc, "表1  三类服务及相关制度在 2026 年与 2030 年的对照")


def build_ch4(doc):
    add_heading_cn(doc, "四、医疗服务展望 2030：从医院能力到家庭可达", level=1)
    add_heading_cn(doc, "（一）2030 年的基准图像", level=2)
    add_body(
        doc,
        "若卫生健康部门目标按期实现，2030 年的医疗服务将呈现三层结构。基层：社区卫生服务中心、乡镇卫生院和村卫生室普遍配备智能辅助诊疗，慢病管理、预问诊、报告解析和健康画像成为常规工具，居民电子健康档案更大范围向个人开放。医院：二级以上医院普遍使用影像智能辅助和临床决策支持，手术排程、药品管理和病历质控由智能体分担事务性负荷。家庭与社区：健康监测设备、家庭病床、上门巡诊和急救响应，把医院能力延伸到户门内。住房政策关心的是第三层——没有合适的空间和网络，前两层能力无法被老人和慢性病人使用。",
    )
    add_quote(
        doc,
        "到 2030 年，基层诊疗智能辅助应用基本实现全覆盖，推动实现二级以上医院普遍开展医学影像智能辅助诊断、临床诊疗智能辅助决策等人工智能技术应用。",
        source="——国家卫生健康委等五部门《关于促进和规范“人工智能+医疗卫生”应用发展的实施意见》",
    )
    add_heading_cn(doc, "（二）社区卫生站成为调度中枢，而不是“小医院”", level=2)
    add_body(
        doc,
        "过去社区卫生机构常被要求“什么都做一点”，结果是能力不足又重复建设。2030 年更合理的定位是调度中枢：向上连接县域医共体和区域医疗中心的专科智能体，向下调度家庭病床、护理员、药剂配送和急救。智能体可以承担分诊、慢病随访、用药提醒和异常预警，但处方权、高风险决策和伦理判断必须留在执业人员手中。这与“赋能而不替代”的政策定位一致。",
    )
    add_body(
        doc,
        "对住宅和小区，这意味着配建不能只给“卫生用房”四个字。需要可停放护理车辆的回车空间、无障碍坡道、可放置自助检测设备的半公共空间、稳定的网络和电源、以及在突发事件中可被急救使用的门禁协同。南阳试点把急救响应作为住房产品竞争力，方向正确；2030 年应把它从个别项目标准提升为完整社区和“好房子”的可选甚至必选指标，尤其是老年人口密度高的小区。",
    )
    add_heading_cn(doc, "（三）中原路径：县域医共体 + 中医药 + 数字家庭", level=2)
    add_body(
        doc,
        "河南医疗服务的难点在县域和农村。人口外出、村医老化、检查能力不足，使智能辅助的边际效用可能高于大城市三级医院。2030 年的中原方案，应把紧密型县域医共体作为智能应用的主容器：县级医院的专科模型和影像能力，通过基层医生工作站到达乡镇和村庄；居民在家中完成血压、血糖、心电等监测，异常由医共体值班智能体和值班医生共同响应。中医药是河南的文化与产业资源，南阳等地已在社区嵌入中医诊室。智能体可以用于经典方剂知识检索、体质辨识辅助和康复指导，但必须严守安全与循证边界，避免把未经评价的生成内容写成“处方”。",
    )
    add_heading_cn(doc, "（四）必须提前约束的风险", level=2)
    add_bullets(
        doc,
        [
            "数字鸿沟：不会使用智能设备的高龄老人被排除在“全覆盖”之外。家庭和社区必须保留人工值守入口。",
            "责任划分：辅助诊断出错时，模型提供方、医院、基层医生和家属之间的责任链需在 2030 年前形成可执行规则。",
            "数据与隐私：户内摄像头、生物体征和电子病历一旦联通，住宅将成为高敏感数据空间。授权应可撤回，默认最小必要。",
            "资源虹吸：智能应用若只强化大医院效率，可能进一步虹吸基层病人。考核应同时看基层解决率和上转合理性。",
        ],
    )


def build_ch5(doc):
    add_heading_cn(doc, "五、养老服务展望 2030：从床位思维到支付—空间—机器人闭环", level=1)
    add_heading_cn(doc, "（一）2030 年仍将以居家社区为主", level=2)
    add_body(
        doc,
        "中国老年人的居住选择不会在四年内颠覆。2030 年的主流仍是居家为基础、社区为依托、机构为专业支撑、医养相结合。变化在于：居家不再等于家庭成员独自承担全部照护；社区不再只有棋牌室和偶尔的义诊；机构不再只比拼床位数。长护险全国建制、家庭养老床位、嵌入式机构和具身辅助，将把“在家养老”从道德号召变成可购买、可报销、可监管的服务。",
    )
    add_body(
        doc,
        "2025 年 10 月启动的医养结合促进行动，要求到 2027 年底结合老龄化程度，推动每个县（市、区、旗）至少有 1 家医疗卫生机构或养老机构直接开展医养结合服务，基本实现县域全覆盖。对河南 100 多个县（市、区）而言，这是硬任务。2030 年应在此基础上进一步要求：县域医养结合机构能够向社区和家庭延伸服务，而不是把老年人集中到城关镇的一栋楼里。",
    )
    add_heading_cn(doc, "（二）机器人进入康养的合理顺序", level=2)
    add_body(
        doc,
        "本中心不支持“2026 年机器人全面替代护工”的叙事。合理顺序是：机构场景先承担高强度、重复性、可标准化的任务（转运辅助、夜间巡视、康复训练、药品分发、环境清洁）；社区驿站承担问询、陪伴、助餐协助、健康监测和紧急呼叫；家庭端以跌倒监测、用药提醒、轻量陪伴和远程家属可视为主，复杂护理仍由护理员完成，机器人作为辅助。腾讯团队以养老场景驱动触觉与力控研究，本身说明：与人接触的任务容错率极低，必须用真实机构数据迭代，不能靠展会演示跨越。",
    )
    add_body(
        doc,
        "IEC 63310 已经把互联家庭环境下辅助生活机器人的功能类别写出框架，包括紧急预警、家属医护通信、家务娱乐照护、外出辅助和信息管理。2030 年的住房验收和适老化改造，可以把其中与建筑相关的条款转译为地方标准：地面平整与门宽是否允许机器人通过，插座与充电位是否安全，网络是否稳定，紧急制动与人工接管是否可用。机器人友好住宅，首先是对人友好的住宅。",
    )
    add_heading_cn(doc, "（三）支付闭环决定服务能否活过试点期", level=2)
    add_body(
        doc,
        "长护险到 2028 年底基本全覆盖，意味着 2030 年河南应已运行两到三年。起步保障重度失能，36 项服务目录以生活照护为主、兼顾基础医疗护理，并鼓励居家和社区。这对住房政策是利好：支付会引导服务进入家庭，从而要求更多适老化改造、家庭养老床位和社区嵌入式设施。风险同样清楚：若评估标准不统一、服务人员不足、虚假服务套保，支付会推高费用却推不高质量。智能监管、服务轨迹核验和家属评价，应与机器人、传感器同步设计，而不是事后打补丁。",
    )
    add_heading_cn(doc, "（四）“社区 + 物业 + 养老”需要空间权属改革", level=2)
    add_body(
        doc,
        "中央文件已鼓励物业向养老、托育、家政、健康延伸。实践中的卡点往往不是意愿，而是房子：配建用房被挪用、权属不清、消防验收困难、运营面积被分割出租。河南三年行动计划把存量资产转型的规划调整、不动产登记、改扩建和消防验收列为 2026 年必须处理的事项，方向准确。2030 年应形成可复制规则：小区配建养老用房不得擅自改变用途；老旧小区补建优先使用闲置会所、底商和低效办公；街道级设施与社区级设施功能分层，避免每个小区都建一个小而全却不可运营的“中心”。",
    )


def build_ch6(doc):
    add_heading_cn(doc, "六、生产效率服务展望 2030：工厂、社区与住房如何共同提效", level=1)
    add_heading_cn(doc, "（一）提效服务不只发生在车间", level=2)
    add_body(
        doc,
        "提高生产效率，通常被理解为工业机器人、工业互联网和工艺优化。这没有错，但不完整。对一个人口大省，效率损失大量发生在车间之外：工人通勤过长、照料负担重导致缺勤、社区物流“最后一百米”反复搬运、物业和家政劳动密集却难标准化。智能体实施意见把生产管理、工艺检测、仓储配售、导引清洁和家政养老助残并列，已经暗示：提效是一条从工厂延伸到生活服务的谱系。住房政策进入这条谱系，是因为职住关系、宿舍与公寓、园区配套和社区物流场地，直接决定机器人与智能调度能否降低综合成本。",
    )
    add_heading_cn(doc, "（二）2030 年工业场景：人机协同而不是无人车间神话", level=2)
    add_body(
        doc,
        "基准情景下，2030 年河南的电子信息、装备、汽车零部件、食品和部分物流仓储，将出现一批可复制的人机协同单元：拆垛、上料、质检、分拣、巡检由具身设备承担，换型与异常处理由技术工人加智能体完成。头部企业的 8 小时连续作业、小时级换型、40%—50% 的局部效率提升，有望在更多产线被部分复制，但不会变成全部工业门类的平均值。食品、超硬材料、能源原材料等河南优势产业，对卫生、粉尘、高温的要求不同，需要分行业打开场景，而不是推广一种人形方案。",
    )
    add_body(
        doc,
        "住房和园区政策应配套三件事。一是产业社区：在机器人产业基地和智能工厂周边提供可负担租赁住房、倒班宿舍、托育和社区医疗，降低综合用工成本。二是实训与数据空间：把部分旧厂房更新为具身智能实训场和采集中心，既盘活存量，又服务本地企业。三是中小企业共享：多数河南制造业企业买不起专线机器人团队，需要园区级的“机器人即服务”和共享产线，避免效率红利只留在龙头企业。",
    )
    add_heading_cn(doc, "（三）生活性提效：把照护和物流从家庭隐形劳动中解放出来", level=2)
    add_body(
        doc,
        "生产效率的社会含义，是减少为维持基本生活而被迫投入的过长时间。2030 年，社区配送机器人、智能快递柜、物业巡检、助餐集中制备、家庭清洁辅助，会比全自主家务人形更早形成稳定供给。它们看起来是民生，实际上是劳动参与率政策：当家庭中的中年劳动力不必为送餐、陪护、取药反复请假，工厂和服务业的有效工时才会上升。这是住房社区服务与工业效率之间被低估的连接。",
    )
    add_heading_cn(doc, "（四）农业与县域：不能只复制城市车间方案", level=2)
    add_body(
        doc,
        "河南粮食产量 2024 年为 6719.37 万吨，农业仍然是基本盘。智能体实施意见已部署农技指导、病虫害诊断、智能农机与智慧大棚。2030 年县域提效，应把冷链节点、农忙互助、农村幸福院和乡镇卫生院放在同一张空间图上：农机和无人机提高田间效率，乡镇服务中心提高生活效率，避免农村年轻人流出后只留下无法被服务覆盖的老年村落。住房政策在县域的任务，是把闲置农房、乡镇供销社和学校腾退空间，依法改建为可运营的服务据点，而不是按城市小区标准硬套电梯与地下车库。",
    )


def build_ch7(doc):
    add_heading_cn(doc, "七、住房作为三类服务的空间操作系统", level=1)
    add_heading_cn(doc, "（一）从“配建面积”到“可运营接口”", level=2)
    add_body(
        doc,
        "过去十年，新建小区配建养老、社区用房的主要考核是面积达标。2030 年应改考四个接口是否成立：物理接口（无障碍、门宽、电梯、回车、充电与网络）、服务接口（卫生站、助餐、日间照料、家政驿站能否实际运营）、支付接口（医保、长护险、政府购买、个人付费能否在同一场所结算）、数据接口（健康档案、呼叫、门禁、设备告警在授权下互通）。没有这四个接口，配建面积只是规划图上的色块。",
    )
    add_table(
        doc,
        ["空间层次", "医疗服务", "养老服务", "提效服务"],
        [
            ["套内住房", "体征监测、家庭病床、急救可达、数字家庭", "适老化、家庭养老床位、防跌倒", "家政辅助、远程办公、充电与网络"],
            ["单元与小区", "无障碍、门禁协同急救、自助检测点", "助餐点、日间照料、活动空间", "快递与无人配送停靠、物业巡检路径"],
            ["社区 / 街道", "社区卫生站、康复与医共体接点", "嵌入式机构、养老顾问", "社区商业、共享家政、技能培训"],
            ["园区 / 县城", "职工健康管理、紧急医疗", "职住平衡减少照料冲突", "机器人产线、共享制造、倒班公寓"],
        ],
        col_widths=[2.8, 4.2, 4.2, 4.0],
    )
    add_caption(doc, "表2  住房空间层次与三类服务的对应关系")
    add_heading_cn(doc, "（二）存量更新是主战场", level=2)
    add_body(
        doc,
        "河南大量老年人住在没有电梯、地面高差大、底商用途混乱的老旧小区。城市更新“十五五”把老旧小区改造、完整社区和存量闲置房屋补民生短板并列，对河南是主要政策工具。2030 年若仍把智能医疗和养老机器人优先放进新建高价住宅，会制造新的服务不平等。本中心主张：中央和省级更新资金应把“可进入的医养服务包”作为老旧小区验收的加分项，而不是只验收管网和墙面。加装电梯、平整入户、增设助餐和卫生用房，是机器人服务的前置条件，也是没有机器人时老人同样需要的条件。",
    )
    add_heading_cn(doc, "（三）好房子的全生命周期，应写进三类服务", level=2)
    add_body(
        doc,
        "“好房子”若只被理解为更好的精装和隔音，会低估政策窗口。国务院人工智能+行动已要求拓展人工智能在好房子全生命周期的应用。2030 年的好房子标准，至少应可选择配置：适老化与儿童安全、家庭健康监测预留、紧急呼叫、可扩展的弱电与边缘算力、家政和护理人员临时工作空间、以及面向未来设备的门宽与荷载。保障性住房和安置房不能被排除在这些选项之外，否则公共住房将成为技术红利的洼地。",
    )
    add_heading_cn(doc, "（四）数据治理就是住房治理", level=2)
    add_body(
        doc,
        "当住宅安装跌倒监测、门磁、摄像头和健康手环，物业、医护、设备商和平台会同时看见家庭内部状态。这已经超出传统物业管理范围。2030 年前应明确：居民是数据权利主体；社区公共安全数据与户内健康数据分类授权；老年人的情感陪护智能体不得诱导消费和制造依赖；发生算法伤害时有本地可及的投诉和关停通道。智能体实施意见已提出决策权限、行为围栏和老年人沉迷风险，住房和物业主管部门应把它写成小区公约和物业合同条款，而不是留在网信文件里。",
    )


def build_ch8(doc):
    add_heading_cn(doc, "八、四个情景与河南含义", level=1)
    add_image(doc, "chart04_scenarios.png", "图4  面向 2030 的四个情景：支付能力 × 工程落地")
    add_image(doc, "chart03_penetration.png", "图3  人机协同服务分层渗透：2026 对照 2030（示意）")
    add_heading_cn(doc, "（一）基准情景：分层渗透（本中心主情景）", level=2)
    add_body(
        doc,
        "政策大体按期推进，技术按工程化速度而不是宣传速度扩散。到 2030 年：工业与仓储的人机协同在电子、物流等场景形成规模；康养机构普遍使用监测、巡视和康复辅助，少量机构部署接触式护理机器人；社区驿站成为健康监测、助餐、呼叫和护理员调度节点；家庭端以可报销的居家护理、家庭养老床位和轻量化设备为主，全自主家务人形仍集中在支付能力较强、户型条件较好的家庭。基层智能辅助诊疗基本覆盖，但使用深度城乡有别。城市更新取得重要进展，完整社区在设区市中心城区较明显，县城和老旧小区仍参差。",
    )
    add_heading_cn(doc, "（二）加速情景：支付、标准与量产同步", level=2)
    add_body(
        doc,
        "若长护险筹资到位、虚假服务被有效遏制、机器人成本持续下降、地方把社区场景真正打开，2030 年可能出现“服务包进小区”的加速：街道级设施连锁化运营，物业成为居家护理的组织者，河南造机器人在本地康养和工厂形成闭环。这一情景的关键不是更多发布会，而是县一级有可运营主体和可审计资金。",
    )
    add_heading_cn(doc, "（三）分化情景：城区样板与服务断崖", level=2)
    add_body(
        doc,
        "这是需要警惕的现实风险。郑州、洛阳等中心城市和新建社区形成智慧医养样板，老旧小区、县城和农村幸福院停留在“有房间、无服务”。智能医疗全覆盖在统计上完成，在高龄、失能、低教育人群中无效。住房价格与服务可及性重新绑定，公共住房和城中村成为空白。分化情景在政治上最危险，因为它看起来完成了指标，却扩大了不平等。",
    )
    add_heading_cn(doc, "（四）迟滞情景：安全、支付或治理停摆", level=2)
    add_body(
        doc,
        "若发生严重人机接触伤害、数据泄露或套保丑闻，应用可能被急刹。技术停留在展厅，小区改造中的充电桩、弱电和用房成为沉没成本。迟滞并不可耻，它提醒政策必须把安全、伦理和审计做在规模化之前。本中心认为，宁可基准情景下慢一点，也不应用迟滞换来的不信任去透支 2030 年之后的十年。",
    )
    add_heading_cn(doc, "（五）对河南的含义", level=2)
    add_body(
        doc,
        "河南同时具备走向基准甚至部分加速的条件，也具备滑向分化的条件。条件在于：人口规模提供应用场景，制造业提供提效需求，中医药和医养结合有地方实践，机器人产业开始在郑州集聚。风险在于：县域财政和专业人才不足，存量小区改造复杂，农村老龄化更重而支付能力更弱。2030 年的河南方案，必须是“中心城市示范 + 县域可复制的简化包”，而不是把沿海都市的全套智能社区原样搬到黄淮平原。",
    )


def build_ch9(doc):
    add_heading_cn(doc, "九、政策建议", level=1)
    add_body(doc, "建议面向河南省及同类中部省份，按 2026—2030 年可执行的原则提出，供住房城乡建设、民政、卫生健康、医保、工业和信息化及高校科研机构讨论。")
    add_heading_cn(doc, "建议一：把三类服务写入详细规划，而不是停留在专项规划附录", level=2)
    add_body(
        doc,
        "在市县国土空间详细规划中，明确社区卫生、嵌入式养老、助餐、急救通道和物流停靠的点位与兼容用途。2026 年已要求完成养老服务设施专项规划并与详规衔接，应把医疗和社区物流一并纳入，避免三类设施争抢同一处底商。对挪用配建用房的项目，建立整改和收回机制。",
    )
    add_heading_cn(doc, "建议二：改造标准增加“服务可达”和“机器人可达”", level=2)
    add_body(
        doc,
        "在老旧小区改造、完整社区和“好房子”指引中，增加门宽、高差、电梯轿厢、无障碍通道、充电与弱电预留、急救门禁协同等指标。这些指标首先服务于轮椅、担架和护理车，同时为辅助机器人预留条件。保障性住房同步适用，防止公共住房掉队。",
    )
    add_heading_cn(doc, "建议三：开放社区与机构作为中试场景，但设置安全红线", level=2)
    add_body(
        doc,
        "支持公办养老机构、社区驿站、县域医共体和产业园区成为智能体与具身设备的中试点，对接河南省人工智能应用场景机会清单。中试必须具备人工接管、责任保险、数据最小化采集和老年人知情同意。接触式护理未达到安全评价前，不得以“智慧养老”名义替代夜间值班人员。",
    )
    add_heading_cn(doc, "建议四：以长护险为杠杆，打通家庭养老床位、物业上门和机构延伸", level=2)
    add_body(
        doc,
        "在长护险推进过程中，优先将居家和社区护理项目落地到已完成适老化改造的家庭和嵌入式设施。鼓励连锁化运营主体同时承接机构床位和周边上门，形成“一个运营者、多层空间”。对物业企业延伸养老服务，明确资质、培训和意外责任，避免无资质兼职造成伤害。",
    )
    add_heading_cn(doc, "建议五：产业园区同步配置职住与健康服务包", level=2)
    add_body(
        doc,
        "在郑州及省内机器人、电子信息、装备产业集聚区，把租赁住房、倒班公寓、托育、职业健康和社区医疗作为园区验收内容。推广“机器人即服务”，让中小制造企业以租赁和共享方式提高效率，降低对一次性资本开支的依赖。",
    )
    add_heading_cn(doc, "建议六：建立住宅健康数据的分类授权与本地投诉通道", level=2)
    add_body(
        doc,
        "由网信、住建、卫生健康联合制定社区智能设备数据指引：户内数据默认本地处理、最小上报；公共区域安防与健康数据分开；老年人陪伴类应用禁止诱导消费。街道设立关停和投诉入口，使技术治理能被居民够到。",
    )
    add_heading_cn(doc, "建议七：把人才培养写进高校和职业院校的空间实践", level=2)
    add_body(
        doc,
        "落实医养照护与管理、智慧健康养老等专业建设要求，推动医学、护理、建筑、城乡规划、公共管理和人工智能交叉培养。河南大学及省内职业院校可与社区、医共体、养老机构和制造企业共建实训，避免只培养会考试不会上门的人才。到 2030 年，县域至少应具备“规划 + 护理 + 运维”的基层团队，而不是只依赖省会专家短期驻点。",
    )
    add_heading_cn(doc, "建议八：为县域和农村准备简化版服务包", level=2)
    add_body(
        doc,
        "不把城市智能社区作为唯一模板。农村幸福院、乡镇卫生院、家庭医生签约和助餐，应构成 2030 年县域的标配简化包：可负担、可维护、可在断电和弱网条件下降级运行。智能辅助作为增强项，而不是能否获得基本服务的前提。",
    )
    add_heading_cn(doc, "结语", level=2)
    add_body(
        doc,
        "2030 年并不遥远。从 2026 年夏天望去，只隔一个“十五五”的主体施工期。医疗、养老和生产效率提升，看起来分属三个部门，实际上会在同一栋住宅、同一个社区门口相遇。住房政策若仍把自己理解为“盖什么样的房子、卖什么样的房子”，会错过这一轮公共服务和技术扩散。住房政策若把自己理解为“服务如何进入空间、空间如何使服务可持续”，则可能在人口大省里走出一条可核对、可复制、对老年人友好、对产业有用的中原路径。",
    )
    add_body(
        doc,
        "本中心的立场可以收为一句话：让技术服务于居住，而不是让居住迁就技术；让三类服务进入家庭，而不让家庭独自承担三类风险。展望 2030 年，成功的标志不是机器人出现在宣传片里，而是一个失能老人能在自己的社区被连续照护，一名基层医生能获得可靠辅助，一条产线能在少用工的约束下稳定运行，并且这些改进同样发生在老旧小区和县城，而不仅发生在新城。",
    )


def build_appendix(doc):
    add_heading_cn(doc, "附录一  2025—2030 年重要政策与技术节点", level=1)
    add_table(
        doc,
        ["时间", "节点", "与本白皮书的关系"],
        [
            ["2025-08", "国务院《人工智能+》行动意见", "给出 2027/2030 普及率与好房子、医疗、养老条款"],
            ["2025-10", "医养结合促进行动", "2027 年底县域医养结合基本覆盖"],
            ["2025-11", "AI+医疗卫生实施意见", "2030 年基层智能辅助基本全覆盖"],
            ["2026-03", "长护险制度意见", "2028 年底全国基本全覆盖，居家社区支付倾斜"],
            ["2026-05", "智能体实施意见；城市更新“十五五”规划", "打开家政养老具身与完整社区空间工程"],
            ["2026-07", "WAIC 2026；郑州众擎 T800 下线", "康养工程化与河南造量产的现场证据"],
            ["2026-09", "WS/T 876—2026 等标准施行", "医养结合服务可检查、可培训"],
            ["2027", "智能体普及率超 70%；医养结合行动收官年", "检验社区设施是否真正运营"],
            ["2028", "长护险基本全覆盖", "居家护理从试点转为制度"],
            ["2030", "AI 全面赋能；基层智能辅助全覆盖；城市更新重要进展", "本白皮书展望的目标年"],
        ],
        col_widths=[2.4, 6.4, 6.4],
    )
    add_caption(doc, "表3  关键节点一览")

    add_heading_cn(doc, "附录二  主要公开资料来源", level=1)
    add_bullets(
        doc,
        [
            "河南省统计局、国家统计局河南调查总队：《2024 年河南省国民经济和社会发展统计公报》，2025 年 4 月。",
            "国务院：《关于深入实施“人工智能+”行动的意见》（国发〔2025〕11 号）。",
            "国家卫生健康委等：《关于促进和规范“人工智能+医疗卫生”应用发展的实施意见》，2025 年 11 月。",
            "国家卫生健康委、民政部等：《关于开展医养结合促进行动的通知》，2025 年 10 月。",
            "中共中央办公厅、国务院办公厅：《关于加快建立长期护理保险制度的意见》，2026 年 3 月。",
            "国务院：《城市更新“十五五”规划》，2026 年 5 月。",
            "中央网信办：《智能体规范应用与创新发展实施意见》，2026 年 5 月 8 日。",
            "IEC 63310《互联家庭环境下使用的主动辅助生活机器人性能准则》；WS/T 876—2026《医养结合健康管理服务标准》。",
            "河南省人民政府及相关部门：养老服务体系和康养产业规划、深化养老服务改革发展三年行动计划、具身智能产业发展行动计划等公开文件。",
            "公开报道：WAIC 2026 康养与具身智能观察、腾讯“小六”机器人、智元工业产线实证、郑州众擎 T800 下线、南阳居家养老住房试点、河南社区养老设施覆盖等（详见 scripts 生成所用资料摘编）。",
        ],
    )
    add_body(
        doc,
        "上述来源均为公开信息。企业效率、成功率、下线规模等数据以报道当时披露为准，可能随后续官方或企业更新而调整。",
        first_indent=True,
    )

    add_heading_cn(doc, "附录三  术语简释", level=1)
    add_table(
        doc,
        ["术语", "在本白皮书中的用法"],
        [
            ["住房即服务", "把住房视为医疗、养老、提效服务的接入点、生产空间和制度接口"],
            ["完整社区", "基本公共服务、便民商业、健身与公共空间相对齐备的社区形态"],
            ["社区嵌入式设施", "插入既有社区、小规模多功能的养老或医养服务空间"],
            ["家庭养老床位", "把专业照护标准延伸到居家环境，并由机构提供上门服务"],
            ["长护险", "为失能人员基本生活照料及相关医疗护理提供服务或资金保障的社会保险"],
            ["智能体", "具备感知、记忆、决策、交互与执行能力的智能系统"],
            ["具身智能", "通过身体与环境交互完成任务的智能，常以机器人本体为载体"],
            ["人机协同服务包", "人员、设备、空间、支付和监管共同构成的可运营服务单元"],
            ["机器人可达性", "门宽、高差、电梯、网络和安全制动等使辅助设备能够进入并工作的空间条件"],
        ],
        col_widths=[3.4, 11.8],
    )
    add_caption(doc, "表4  术语简释")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("（全文完）")
    set_run_font(run, size=12, color=GRAY, font="楷体")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("河南大学住房政策研究中心")
    set_run_font(run2, size=11, color=NAVY, font="黑体")


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "宋体")

    setup_header_footer(doc)
    build_cover(doc)
    build_statement(doc)
    build_toc(doc)
    build_abstract(doc)
    build_ch1(doc)
    build_ch2(doc)
    build_ch3(doc)
    build_ch4(doc)
    build_ch5(doc)
    build_ch6(doc)
    build_ch7(doc)
    build_ch8(doc)
    build_ch9(doc)
    build_appendix(doc)

    doc.save(str(OUT))
    print(f"已生成：{OUT}")


if __name__ == "__main__":
    main()
