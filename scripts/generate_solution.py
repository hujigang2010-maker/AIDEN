#!/usr/bin/env python3
"""生成完整解决方案 Word + 赔付测算 Excel。家属内部测算件，不得给骑手、刘孝春或保险看。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

CHINESE_FONT = "宋体"
HEADING_FONT = "黑体"
ACCENT = RGBColor(0x0B, 0x2F, 0x5B)
MUTED = RGBColor(0x5C, 0x6B, 0x7A)
RED = RGBColor(0xA6, 0x3D, 0x2F)

DOCX_NAME = "青岛抚顺路和哈尔滨路路口交通事故_完整解决方案_20260822.docx"
XLSX_NAME = "青岛抚顺路和哈尔滨路路口交通事故_赔付测算表_20260822.xlsx"

WARN = (
    "家属内部测算件。所有金额是估算区间，以发票、医嘱、认定书、鉴定意见和保单条款为准。"
    "不是律师函。不要给骑手、刘孝春、站点或任何保险公司看。对外仍然不报总数、不报残级。"
)

# ---------------------------------------------------------------------------
# 一、责任与赔付分担
# ---------------------------------------------------------------------------
SPLIT_INTRO = [
    "书面认定书尚未出具。交警看监控排除任何一方全责，剩两种可能：外卖主责、三轮次责（内部按 70%/30% 测算）；或双方同等（50%/50%）。下面两种都算，认定书出来套哪种用哪种。",
    "责任只在胡志远（三轮使用人）和美团骑手之间划分。刘孝春不上认定书，但作为车主（无牌、灯坏的车交人使用）和接活的老板，要在三轮这一方的内部缺口里分担一部分（内部按其过错系数 k 估 50%–70% 争取）。",
    "爸爸自己要担的那部分（缺口里刘不担的余数），不是白白认掉——去翻爸爸名下有没有泰康等个人保险：意外医疗险可以报销自负的医疗费（报销型，不能和美团险重复报同一张发票）；意外伤残给付按残级给钱，给付型，不看责任比例，可以和交通事故赔偿叠加；有住院津贴的按天另领。",
    "未结工钱约 5000 元在赔付体系之外，100% 由刘孝春支付，不乘任何比例。",
]

SPLIT_ROWS = [
    ["层级", "谁出钱", "出什么", "口径"],
    ["第 1 层", "医保", "住院医疗费先行结算部分", "已用；理赔时扣除已报金额，不重复"],
    ["第 2 层", "美团平台险（骑手侧）", "总损失 L × 骑手责任比例（70% 或 50%）", "以保单限额为上限；保单号、限额待拉群确认"],
    ["第 3 层", "刘孝春（车主+雇主）", "己方缺口 G × k（内部按 k=50%–70% 争取）", "G = L × 己方比例（30% 或 50%）；协商不成随雇员受害之诉主张"],
    ["第 4 层", "泰康等爸爸个人险", "爸爸自负部分 G × (1−k) 中的医疗费（报销型）；伤残给付按残级（给付型，全额不打折）", "先找保单：险种、限额、免赔额、报案时限；意外医疗不能与美团险重复报同一发票"],
    ["第 5 层", "起诉执行", "以上谈不拢的差额", "起诉骑手（视情况追加平台）+ 刘孝春；执行盯保险，不指望骑手个人"],
    ["单列", "刘孝春", "欠薪约 5000 元", "100% 支付，不乘比例，不进 L"],
]

# ---------------------------------------------------------------------------
# 二、费用估算（单位：元；内部区间）
# ---------------------------------------------------------------------------
COST_ROWS = [
    # 项目, 低, 高, 计算口径, 备注
    ["医疗费（已发生）", 40000, 40000, "家属口径 4 万，发票逐笔核清中", "与刘垫付 3 万现金分开记；扣医保已报部分"],
    ["医疗费（后续二期）", 15000, 30000, "拆外固定/取克氏针住院 + 换药 + 可能植皮", "皮肤坏死若需植皮取高值"],
    ["护理费", 15000, 35000, "300 元/天 × 住院约 30–50 天 + 出院后医嘱 30–60 天", "医院推荐一对一已请；必须合同+发票+护理记录"],
    ["住院伙食补助", 3000, 5000, "约 100 元/天 × 住院 30–50 天", "按青岛国家机关出差伙食标准"],
    ["营养费", 2000, 5000, "有医嘱或鉴定支持的天数 × 30–50 元/天", "没有医嘱不主张"],
    ["误工费（伤者本人）", 0, 30000, "约 5000 元/月 × 休息期 4–6 个月", "64 岁已退休，必须证明给刘干活的实际收入；证明不了为 0。胡继刚工资不算"],
    ["交通费", 1000, 8000, "救护车 160 元 + 家属必要往返", "全部要票"],
    ["住宿费", 0, 5000, "外地家属必要住宿", "有票且必要才报"],
    ["鉴定费", 1500, 2500, "伤残鉴定一次", "有的判决由败诉方担"],
    ["残疾辅助器具", 500, 2000, "拐杖、轮椅等", "有医嘱才报"],
]

DISABILITY_BASE = 71703  # 内部测算基数：城镇居民人均可支配收入口径（元/年）
DISABILITY_YEARS = 16    # 64 岁：20 年 −（64−60）
GRADE_ROWS = [
    # 情形, 残疾赔偿金, 精神抚慰金, 说明
    ["评不上残级", 0, 0, "开放性骨折不必然构成伤残；以踝关节最终功能为准"],
    ["十级", DISABILITY_BASE * DISABILITY_YEARS * 0.1, 5000, "71703 × 16 年 × 0.1 ≈ 11.5 万；踝关节功能部分丧失时常见争取方向"],
    ["九级", DISABILITY_BASE * DISABILITY_YEARS * 0.2, 10000, "71703 × 16 年 × 0.2 ≈ 22.9 万；须功能丧失更重，由鉴定说了算"],
]

# 三情景总账（取费用中值 + 各残级）
def _scenario_totals():
    cost_low = sum(r[1] for r in COST_ROWS)
    cost_high = sum(r[2] for r in COST_ROWS)
    cost_mid = (cost_low + cost_high) / 2
    out = []
    for grade, dis, jing, _ in GRADE_ROWS:
        l_mid = cost_mid + dis + jing
        row = {
            "grade": grade,
            "L": l_mid,
            "meituan70": l_mid * 0.7,
            "gap30": l_mid * 0.3,
            "liu30_k": (l_mid * 0.3 * 0.5, l_mid * 0.3 * 0.7),
            "self30": (l_mid * 0.3 * 0.3, l_mid * 0.3 * 0.5),
            "meituan50": l_mid * 0.5,
            "gap50": l_mid * 0.5,
            "liu50_k": (l_mid * 0.5 * 0.5, l_mid * 0.5 * 0.7),
            "self50": (l_mid * 0.5 * 0.3, l_mid * 0.5 * 0.5),
        }
        out.append(row)
    return cost_low, cost_high, cost_mid, out


# ---------------------------------------------------------------------------
# 三、理赔流程与时间线
# ---------------------------------------------------------------------------
TIMELINE_ROWS = [
    ["节点", "预计时间", "做什么", "完成标志"],
    ["1 固定证据", "本周（8 月下旬）", "3 万收据五句话；欠薪 5000 另写；护工合同发票；骑手拉群；监控回执；泰康等个人保单找出来并报案", "四张纸齐 + 群建立 + 泰康报案号"],
    ["2 书面认定书", "1–2 周内", "催陈师傅给预计出具日；拿到后先拍照发岳父，再决定签不签；岳父看监控定简易/一般", "认定书原件在手，比例落定（70/30 或 50/50）"],
    ["3 第一轮理赔", "认定书后 2–4 周", "美团平台险按比例赔已发生医疗费、护理费、交通费；材料按保险清单交；泰康意外医疗报自负部分", "首笔理赔到账或书面拒赔理由"],
    ["4 持续治疗", "伤后 0–3 个月", "人民医院换药、VSD 管理、外固定维护；发票、护理记录逐月归档", "创面愈合、外固定稳定"],
    ["5 二次住院", "伤后约 6–12 周", "拆外固定架/取克氏针，住院约 5–7 天；若植皮另加 1–2 周", "二次手术出院记录、发票"],
    ["6 治疗终结", "伤后约 3–6 个月", "临床稳定、医生认可终结；收齐完整手术记录、植入物清单", "复查报告显示骨愈合"],
    ["7 伤残鉴定", "治疗终结后启动，约 1 个月出", "司法鉴定机构、道路交通事故用途；右足骨刺不送评", "鉴定意见书（有级或无级）"],
    ["8 总账协商", "鉴定后 1–2 个月", "按票 + 残级 + 认定书比例列请求：美团险 → 刘孝春 → 泰康伤残给付；欠薪单独收", "调解协议或明确谈不拢"],
    ["9 起诉（如需）", "协商不成后；一审约 6–12 个月", "起诉骑手（视情况追加平台）+ 刘孝春；岳父审代理合同", "判决/调解书 + 执行到账"],
]

HOSPITAL_ROWS = [
    ["阶段", "时间", "时长", "说明"],
    ["齐鲁医院（已发生）", "2026-08-14 至 08-17", "3 天", "急诊 + 手术（8-14 19:01 切开复位内外固定、清创、VSD）"],
    ["市北区人民医院（进行中）", "2026-08-17 起", "估 2–4 周", "隔日换药、VSD、外固定维护；出院看创面"],
    ["二次住院（拆外固定/取克氏针）", "伤后约 6–12 周", "估 5–7 天", "医嘱定；若皮肤坏死需植皮，另加 1–2 周"],
    ["合计住院", "—", "估 30–50 天", "护理费、伙食补助按此计"],
    ["治疗终结", "伤后约 3–6 个月", "—", "之后才能评残；总流程顺利 6–9 个月，诉讼则 12–18 个月"],
]

TAIKANG_ROWS = [
    ["核对项", "问什么", "为什么"],
    ["保单是否存在", "爸爸名下泰康（或其他公司）意外险、医疗险、住院津贴险；翻保单、问业务员、查泰康 App/95522", "自负部分能不能转出去，先看有没有单"],
    ["报案时限", "事故已发生 8 天，立即报案拿报案号", "多数条款要求及时报案，拖久了扯皮"],
    ["意外医疗限额与免赔", "限额多少、免赔额多少、是否限医保内用药、报销比例", "报销型：报爸爸自负的医疗费，不能和美团险重复报同一张发票"],
    ["伤残给付条款", "按《人身保险伤残评定标准》几级给百分之几", "给付型：按残级给钱，不看责任比例，可与交通事故赔偿叠加；注意保险伤残标准与司法鉴定标准不同，可能要两次鉴定"],
    ["住院津贴", "有没有、每天多少、限多少天", "按住院天数另领，与其他赔付不冲突"],
    ["材料清单", "发票原件给谁（原件只有一套，先给赔付大头的）", "美团险和泰康都要发票时，用分割单或复印件加盖章"],
]

FORBID = [
    "本件所有数字不对骑手、刘孝春、保险任何一方出口。有人问一共要多少：费用还在发生，等认定书和治疗稳定后按票算。",
    "残级没鉴定之前，对外不说十级九级；对泰康报案只报事故和伤情，不自报残级。",
    "不签一次性了结；刘的每一笔钱都写垫付收据。",
    "美团侧约 70%（或 50%）找平台险，不让刘替骑手补，也不自己认掉。",
    "误工费证据（微信派活记录、工钱转账、市场证人）现在就收集，否则这项就是 0。",
]


def _set_font(run, font, size, bold=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rf)
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(attr), font)


def _para(doc, text="", font=CHINESE_FONT, size=11.5, bold=False, color=None,
          align=WD_ALIGN_PARAGRAPH.LEFT, indent=None, before=0, after=5, spacing=1.35):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = spacing
    if indent is not None:
        pf.first_line_indent = Cm(indent)
    if text:
        _set_font(p.add_run(text), font, size, bold, color)
    return p


def _h(doc, text, size=14):
    _para(doc, text, font=HEADING_FONT, size=size, bold=True, color=ACCENT, before=10, after=6, spacing=1.25)


def _table(doc, rows, widths=None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(str(val))
            _set_font(run, CHINESE_FONT, 10, bold=(i == 0))
    return t


def _wan(v):
    return f"{v / 10000:.1f} 万"


def _wan_range(pair):
    return f"{pair[0] / 10000:.1f}–{pair[1] / 10000:.1f} 万"


def build_document(output_path: Path) -> None:
    cost_low, cost_high, cost_mid, scenarios = _scenario_totals()
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)

    _para(doc, "青岛抚顺路和哈尔滨路路口交通事故：完整解决方案", font=HEADING_FONT, size=17, bold=True,
          color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, spacing=1.2)
    _para(doc, "责任分担 · 理赔流程与时间线 · 费用与赔付估算 · 截至 2026-08-22", font=HEADING_FONT, size=10.5,
          color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=6, spacing=1.2)
    _para(doc, WARN, bold=True, color=RED, after=8)

    _h(doc, "一、责任与赔付分担（含泰康路径）")
    for p in SPLIT_INTRO:
        _para(doc, p, indent=0.7)
    _table(doc, SPLIT_ROWS)

    _h(doc, "二、理赔流程、顺序与时间节点")
    _para(doc, "理赔顺序固定为：医保 → 美团平台险（按认定书比例）→ 刘孝春分担己方缺口 → 泰康等个人险补自负医疗并另领伤残给付 → 差额起诉执行。欠薪单独找刘收。", indent=0.7)
    _table(doc, TIMELINE_ROWS)
    _para(doc, "总时长：顺利（协商成）约 6–9 个月；进入诉讼约 12–18 个月。", bold=True, before=6)

    _h(doc, "三、住院时间与是否二次住院")
    _table(doc, HOSPITAL_ROWS)
    _para(doc, "是否二次住院：是，基本确定要。外固定架和克氏针要拆（常见伤后 6–12 周），届时住院约 5–7 天；皮肤坏死若需植皮再加。都以主治医嘱为准。", indent=0.7, before=6)

    _h(doc, "四、各项费用估算（内部区间，以票为准）")
    header = ["项目", "低（元）", "高（元）", "计算口径", "备注"]
    rows = [header] + [[r[0], f"{r[1]:,}", f"{r[2]:,}", r[3], r[4]] for r in COST_ROWS]
    rows.append(["常规费用小计", f"{cost_low:,}", f"{cost_high:,}", f"中值约 {_wan(cost_mid)}", "不含残疾赔偿金和精神抚慰金"])
    _table(doc, rows)

    _para(doc, "伤残相关（评残后才有，内部测算基数 71703 元/年 × 16 年）：", bold=True, before=8, after=4)
    g_rows = [["情形", "残疾赔偿金", "精神抚慰金", "说明"]]
    for grade, dis, jing, note in GRADE_ROWS:
        g_rows.append([grade, _wan(dis) if dis else "0", f"{jing:,} 元" if jing else "0", note])
    _table(doc, g_rows)

    _h(doc, "五、三种情形的总账测算（费用取中值）")
    s_rows = [["残级情形", "总损失 L", "外卖主责：美团 70%", "己方缺口 30%", "刘孝春担（k=50–70%）", "爸爸自负→泰康", "同等：美团 50%", "刘孝春担", "爸爸自负→泰康"]]
    for sc in scenarios:
        s_rows.append([
            sc["grade"], _wan(sc["L"]),
            _wan(sc["meituan70"]), _wan(sc["gap30"]), _wan_range(sc["liu30_k"]), _wan_range(sc["self30"]),
            _wan(sc["meituan50"]), _wan_range(sc["liu50_k"]), _wan_range(sc["self50"]),
        ])
    _table(doc, s_rows)
    _para(doc, "读法举例：若外卖主责且评上十级，总损失约 " + _wan(scenarios[1]["L"]) +
          "，美团侧约 " + _wan(scenarios[1]["meituan70"]) + "（受保单限额约束），向刘孝春争取约 " +
          _wan_range(scenarios[1]["liu30_k"]) + "，爸爸自负约 " + _wan_range(scenarios[1]["self30"]) +
          "——这块用泰康意外医疗报销 + 伤残给付去补。泰康伤残给付另算，不占上表。", indent=0.7, before=6)

    _h(doc, "六、泰康（个人险）核对清单")
    _table(doc, TAIKANG_ROWS)

    _h(doc, "七、纪律（贯穿全程）")
    for x in FORBID:
        _para(doc, "• " + x, color=RED, after=3)

    _para(doc, "本件依据 2026-08-14 至 08-22 病历、口述与家属确认整理。重新生成：python3 scripts/build_all.py。", size=9.5, color=MUTED, before=10)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------
HEAD_FILL = PatternFill("solid", fgColor="0B2F5B")
WARN_FILL = PatternFill("solid", fgColor="FDECEA")


def _sheet(wb, name, title, rows, widths):
    ws = wb.create_sheet(name)
    ws.cell(1, 1, title).font = Font(name="黑体", size=12, bold=True, color="0B2F5B")
    ws.cell(2, 1, WARN).font = Font(name="宋体", size=9, color="A63D2F")
    ws.cell(2, 1).fill = WARN_FILL
    start = 4
    for j, val in enumerate(rows[0], 1):
        c = ws.cell(start, j, val)
        c.font = Font(name="黑体", size=10, bold=True, color="FFFFFF")
        c.fill = HEAD_FILL
        c.alignment = Alignment(vertical="center", wrap_text=True)
    for i, row in enumerate(rows[1:], start + 1):
        for j, val in enumerate(row, 1):
            c = ws.cell(i, j, val)
            c.font = Font(name="宋体", size=10)
            c.alignment = Alignment(vertical="top", wrap_text=True)
    for j, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(j)].width = w
    ws.freeze_panes = ws.cell(start + 1, 1)
    return ws


def build_workbook(output_path: Path) -> None:
    cost_low, cost_high, cost_mid, scenarios = _scenario_totals()
    wb = Workbook()
    wb.remove(wb.active)

    _sheet(wb, "责任与赔付分担", "谁出钱、出什么、按什么口径", SPLIT_ROWS, [10, 24, 42, 46])

    _sheet(wb, "理赔流程时间线", "理赔顺序与时间节点（顺利 6–9 个月；诉讼 12–18 个月）", TIMELINE_ROWS, [16, 18, 52, 30])

    _sheet(wb, "住院与治疗时长", "住院时间、二次住院与治疗终结", HOSPITAL_ROWS, [26, 22, 14, 50])

    cost_rows = [["项目", "低（元）", "高（元）", "计算口径", "备注"]]
    for r in COST_ROWS:
        cost_rows.append([r[0], r[1], r[2], r[3], r[4]])
    cost_rows.append(["常规费用小计", cost_low, cost_high, f"中值约 {cost_mid:,.0f} 元", "不含残疾赔偿金/精神抚慰金"])
    for grade, dis, jing, note in GRADE_ROWS:
        cost_rows.append([f"残疾赔偿金（{grade}）", dis, dis, "71703 × 16 年 × 残级系数", note])
        if jing:
            cost_rows.append([f"精神抚慰金（{grade}）", jing, jing, "青岛法院尺度内部估", "平台险常不赔，向侵权人主张"])
    _sheet(wb, "费用估算明细", "各项费用估算（以票、医嘱、鉴定为准）", cost_rows, [24, 12, 12, 40, 40])

    s_rows = [["残级情形", "总损失 L（元）", "美团 70%", "缺口 30%", "刘担 k=50%", "刘担 k=70%", "自负→泰康（低）", "自负→泰康（高）", "美团 50%", "缺口 50%", "刘担 50% 情形 k=50%", "自负（同等,低）"]]
    for sc in scenarios:
        s_rows.append([
            sc["grade"], round(sc["L"]), round(sc["meituan70"]), round(sc["gap30"]),
            round(sc["liu30_k"][0]), round(sc["liu30_k"][1]),
            round(sc["self30"][0]), round(sc["self30"][1]),
            round(sc["meituan50"]), round(sc["gap50"]),
            round(sc["liu50_k"][0]), round(sc["self50"][0]),
        ])
    s_rows.append(["欠薪（体系外）", 5000, 0, 0, 5000, 5000, 0, 0, 0, 0, 5000, 0])
    _sheet(wb, "三情景总账", "外卖主责 70/30 与同等 50/50 两套测算（费用取中值）", s_rows, [16, 14, 12, 12, 12, 12, 14, 14, 12, 12, 16, 14])

    _sheet(wb, "泰康保单核对", "爸爸个人险核对清单（先找保单、立即报案）", TAIKANG_ROWS, [16, 52, 46])

    forbid_rows = [["序号", "纪律"]] + [[i, x] for i, x in enumerate(FORBID, 1)]
    _sheet(wb, "纪律", "贯穿全程", forbid_rows, [8, 100])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(output_path))


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "deliverables"
    build_document(out / DOCX_NAME)
    build_workbook(out / XLSX_NAME)
    print(f"Wrote {out / DOCX_NAME}")
    print(f"Wrote {out / XLSX_NAME}")


if __name__ == "__main__":
    main()
