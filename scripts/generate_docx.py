#!/usr/bin/env python3
"""生成交通事故处理备忘录 Word。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from content import (
    ACCIDENT,
    ASR_CORRECTIONS,
    BOSS,
    BOSS_ITEMS,
    BOSS_NOT,
    GAPS,
    CALLS,
    CLOUD_FILM,
    CONCLUSION_BULLETS,
    CT_REPORTS,
    DATE_LABEL,
    DISABILITY,
    HOSPITAL_TRACK,
    ACCIDENT_FOUR,
    TRANSFER_THREE,
    NOW_EIGHT,
    INJURY_NOT_THIS_ACCIDENT,
    INJURY_THIS_ACCIDENT,
    LEGAL_REMINDER,
    MONEY,
    PARTIES,
    PROCEDURES,
    QINGXIAN,
    SOURCE_PACK,
    STAGE_FORBID,
    STAGE_PRINCIPLE,
    STAGE_STEPS,
    SUBTITLE,
    TALKING_POINTS,
    TITLE,
)

CHINESE_FONT = "宋体"
HEADING_FONT = "黑体"
ACCENT = RGBColor(0x0B, 0x2F, 0x5B)
MUTED = RGBColor(0x5C, 0x6B, 0x7A)
RED = RGBColor(0xA6, 0x3D, 0x2F)
GREEN = RGBColor(0x2F, 0x6B, 0x4F)


def _set_run_font(run, font_name: str, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _add_paragraph(
    doc: Document,
    text: str = "",
    *,
    font: str = CHINESE_FONT,
    size: float = 12,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: float | None = None,
    space_after: float = 6,
    space_before: float = 0,
    line_spacing: float = 1.5,
    color: RGBColor | None = None,
):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if text:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold, color=color)
    return p


def _add_rich(doc: Document, segments, **kwargs):
    p = _add_paragraph(doc, "", **kwargs)
    font = kwargs.get("font", CHINESE_FONT)
    size = kwargs.get("size", 12)
    color = kwargs.get("color")
    for text, bold in segments:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold, color=color)
    return p


def _heading(doc: Document, text: str, level: int = 1) -> None:
    if level == 1:
        _add_paragraph(doc, text, font=HEADING_FONT, size=14, bold=True, space_before=14, space_after=8, color=ACCENT, line_spacing=1.3)
    else:
        _add_paragraph(doc, text, font=HEADING_FONT, size=12, bold=True, space_before=10, space_after=6, line_spacing=1.3)


def _body(doc: Document, text: str) -> None:
    _add_paragraph(doc, text, first_line_indent=0.74, space_after=6)


def _bullet(doc: Document, text: str, *, bold_prefix: str | None = None, color: RGBColor | None = None) -> None:
    if bold_prefix:
        p = _add_rich(
            doc,
            [("• ", False), (bold_prefix, True), (text, False)],
            first_line_indent=0.37,
            space_after=3,
            line_spacing=1.4,
            color=color,
        )
        return p
    return _add_paragraph(doc, f"• {text}", first_line_indent=0.37, space_after=3, line_spacing=1.4, color=color)


def _set_cell_text(cell, text, *, bold=False, size=10.5, color=None, font=CHINESE_FONT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    _set_run_font(run, font, size, bold=bold, color=color)


def _shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = tcPr.makeelement(qn("w:shd"), {})
        tcPr.append(shd)
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")


def _table(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, h, bold=True, size=10.5, color=RGBColor(0xFF, 0xFF, 0xFF), font=HEADING_FONT)
        _shade(cell, "0B2F5B")
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            _set_cell_text(cell, val, size=10.5)
            if r_i % 2 == 1:
                _shade(cell, "F3F6FA")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    _add_paragraph(doc, "", space_after=8)
    return table


def build_document(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.6)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = CHINESE_FONT
    normal_style.font.size = Pt(12)
    rpr = normal_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CHINESE_FONT)

    _add_paragraph(doc, TITLE, font=HEADING_FONT, size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, line_spacing=1.3, color=ACCENT)
    _add_paragraph(doc, SUBTITLE, font=HEADING_FONT, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, line_spacing=1.3, color=MUTED)
    _add_paragraph(doc, DATE_LABEL + "  ·  材料：17 份沟通记录原文合集 + 口径对照 + 齐鲁医院云影像 3 份 CT", font=CHINESE_FONT, size=10.5, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, color=MUTED)

    _heading(doc, "一、先把三句话钉死")
    for item in CONCLUSION_BULLETS:
        _bullet(doc, item)
    _body(doc, "这份备忘录只做事实整理和行动清单。它不是医学鉴定书，也不是律师意见书。伤残等级必须由有资质的司法鉴定机构在治疗终结后出具。")

    _heading(doc, "现阶段办理顺序：先证据，后总账")
    _add_paragraph(doc, STAGE_PRINCIPLE, bold=True, color=RED, space_after=8)
    _heading(doc, "贯穿禁止", 2)
    for item in STAGE_FORBID:
        _bullet(doc, item, color=RED)
    _heading(doc, "分步清单（紧急程度）", 2)
    stage_rows = [
        [s["id"], s["urgency"], s["when"], s["name"], s["who"], s["done"]]
        for s in STAGE_STEPS
    ]
    _table(doc, ["序号", "缓急", "时限", "事项", "谁来办", "完成标志"], stage_rows)
    _heading(doc, "每一项怎么办理", 2)
    for s in STAGE_STEPS:
        _bullet(doc, s["how"], bold_prefix=f"{s['id']} {s['name']}（{s['urgency']}）：")
        _add_paragraph(doc, "为什么现在做：" + s["why"], size=10.5, color=MUTED, space_after=6, first_line_indent=0.37)

    _heading(doc, "沟通记录去重：转院三点、事故四条、立即八事")
    _body(doc, "这些材料大量重复。读原件只为核对细节；对外和行动只看下面原文清单。")
    _heading(doc, "转院与保险：核心只三点", 2)
    for t in TRANSFER_THREE:
        _bullet(doc, t["what"], bold_prefix=f"{t['n']}. ")
    _heading(doc, "事故经过、责任与对方沟通", 2)
    for a in ACCIDENT_FOUR:
        _bullet(doc, a["text"])
    _heading(doc, "立即八事", 2)
    for x in NOW_EIGHT:
        _bullet(doc, x["what"], bold_prefix=f"{x['n']}. ")
    _heading(doc, "八事：谁来办、为什么", 2)
    eight_rows = [[str(x["n"]), x["who"], x["what"], x["why"]] for x in NOW_EIGHT]
    _table(doc, ["序号", "谁", "做什么", "为什么"], eight_rows)

    _heading(doc, "二、谁受伤：只有 64 岁男性胡某")
    _body(doc, f"医院：{CLOUD_FILM['hospital']}。云影像患者信息：{CLOUD_FILM['patient_display']} / {CLOUD_FILM['sex']} / {CLOUD_FILM['age']}岁。三份已审核 CT 的检查号分别为 10008056847（右小腿）、10008056848（右足）、10008059257（术后右踝）。")
    _bullet(doc, PARTIES["injured"])
    _bullet(doc, PARTIES["family"])
    _bullet(doc, PARTIES["other"])
    _bullet(doc, PARTIES["police"])
    _bullet(doc, PARTIES["boss"])
    _body(doc, "结论：伤残如果以后评得上，也只可能评在伤者胡某本人。家属、对接人、骑手都没有本次伤残材料。")

    _heading(doc, "三、伤情鉴定：以齐鲁医院 CT 为准，不以通话口述为准")
    _body(doc, CLOUD_FILM["source_note"] + CLOUD_FILM["disclaimer"])
    _heading(doc, "（一）本次交通事故造成的损伤", 2)
    for item in INJURY_THIS_ACCIDENT:
        _bullet(doc, item, color=RED)
    _heading(doc, "（二）右足 CT 显示的非本次外伤", 2)
    for item in INJURY_NOT_THIS_ACCIDENT:
        _bullet(doc, item, color=GREEN)
    _heading(doc, "（三）三份报告原文要点", 2)
    rows = []
    for r in CT_REPORTS:
        rows.append(
            [
                r["date"],
                r["item"] + "\n检查号 " + r["no"],
                r["place"],
                r["findings"],
                "；".join(r["diagnosis"]),
                r["role"],
            ]
        )
    _table(
        doc,
        ["日期", "项目", "科室", "影像所见", "影像诊断", "怎么用"],
        rows,
    )

    _heading(doc, "（四）通话口述必须立刻改口的地方", 2)
    for wrong, right in ASR_CORRECTIONS:
        _bullet(doc, right, bold_prefix=wrong + " → ")

    _heading(doc, "四、谁伤残：现在谁都还没有残级")
    _bullet(doc, DISABILITY["who"], bold_prefix="评谁：")
    _bullet(doc, DISABILITY["now"], bold_prefix="现在：")
    _bullet(doc, DISABILITY["when"], bold_prefix="何时评：")
    _bullet(doc, DISABILITY["standard"], bold_prefix="用什么标准：")
    _bullet(doc, DISABILITY["what_counts"], bold_prefix="评什么：")
    _add_paragraph(doc, DISABILITY["cannot_grade"], bold=True, color=RED, space_after=8)
    _heading(doc, "评残前只做准备、不报价", 2)
    for item in DISABILITY["prep"]:
        _bullet(doc, item)
    _body(doc, "医学含义（给家属内部看，不要拿去跟骑手谈级数）：这是右踝关节骨折合并半脱位，加上胫骨远端、腓骨近端骨折，已经手术内外固定。属于较重的下肢创伤。治疗终结后如果踝关节活动明显受限、行走距离或负重受影响，通常具备启动评残的医学基础。残级高低看日后功能，不看急诊片子有多碎。右足骨刺不要塞进评残申请。")

    _heading(doc, "五、事故与责任：听交警视频，不听「全责」口述")
    _bullet(doc, ACCIDENT["place"], bold_prefix="地点：")
    _bullet(doc, ACCIDENT["time_note"], bold_prefix="时间：")
    _bullet(doc, ACCIDENT["process"], bold_prefix="过程：")
    _heading(doc, "交警指出的三轮车问题", 2)
    for item in ACCIDENT["tri_faults"]:
        _bullet(doc, item)
    _heading(doc, "交警指出的外卖骑手问题", 2)
    for item in ACCIDENT["rider_faults"]:
        _bullet(doc, item)
    _add_paragraph(doc, ACCIDENT["police_range"], bold=True, space_after=8)
    _body(doc, "8 月 15–16 日家属内部一度按「对方全责」讨论。8 月 17 日交警已用监控否定该前提。对外口径立即改为：排除全责，争取主次（三轮次责、外卖主责），保底同等责任。")

    _heading(doc, "六、简易程序还是一般程序")
    _bullet(doc, PROCEDURES["simple"], bold_prefix="简易：")
    _bullet(doc, PROCEDURES["general"], bold_prefix="一般：")
    _bullet(doc, PROCEDURES["police_advice"], bold_prefix="本次怎么选：")
    _bullet(doc, PROCEDURES["license"], bold_prefix="驾驶资格：")
    _bullet(doc, PROCEDURES["after_letter"], bold_prefix="认定书：")
    _body(doc, "同意任何责任比例或程序之前，必须让处理交通事故的本地律师先看监控和车辆材料。有驾驶证后，一般程序的无证风险下降；无青岛市号牌的风险仍在。家属不自己拍板走简易。")

    _heading(doc, "七、转院与保险：核心只三点")
    _body(doc, "转院、报销类录音大量重复。读原件只为核对细节；行动只看这三点。人已到市北区人民医院，家属确认可以赔、比例待定。")
    for t in TRANSFER_THREE:
        _bullet(doc, t["how"] + "完成标志：" + t["done"], bold_prefix=f"{t['n']}. {t['what']}：")
    _bullet(doc, HOSPITAL_TRACK["settle_rule"], bold_prefix="结算铁律：")
    _bullet(doc, HOSPITAL_TRACK["no_motor_insurer"], bold_prefix="找哪家保险：")
    _bullet(doc, HOSPITAL_TRACK["family_rule"], bold_prefix="家属口径：")
    _heading(doc, "明确不做", 2)
    for item in HOSPITAL_TRACK["not_do"]:
        _bullet(doc, item)
    _body(doc, "手术在齐鲁三甲手足与显微重建外科完成。不要另找其他三甲碰运气。医院分叉细节不再展开，见立即八事第 3、4 条。")

    _heading(doc, "八、钱从哪里来")
    _bullet(doc, MONEY["paid_family_816"], bold_prefix="费用口径：")
    _bullet(doc, MONEY["paid_815_conflict"], bold_prefix="垫付现金：")
    _bullet(doc, MONEY["pad_receipt"], bold_prefix="收据怎么写：")
    _bullet(doc, MONEY["insurance"], bold_prefix="保险：")
    _bullet(doc, MONEY["claims_split"], bold_prefix="项目：")
    _bullet(doc, MONEY["ambulance"], bold_prefix="救护车：")
    _heading(doc, "法律关系提醒：不必先经法院确认才生效", 2)
    for item in LEGAL_REMINDER:
        _bullet(doc, item)
    _heading(doc, "刘孝春：车主/劳务，不是替美团补差额的人", 2)
    _bullet(doc, BOSS["who"], bold_prefix="身份：")
    _bullet(doc, BOSS["facts"], bold_prefix="当天：")
    _bullet(doc, BOSS["labor"], bold_prefix="有没有劳动合同：")
    _bullet(doc, BOSS["traffic"], bold_prefix="交警怎么划：")
    _bullet(doc, BOSS["owner"], bold_prefix="车主过错：")
    _bullet(doc, BOSS["not_meituan"], bold_prefix="不替美团赔：")
    _bullet(doc, BOSS["yes_own_side"], bold_prefix="可能分担哪一块：")
    _bullet(doc, BOSS["now"], bold_prefix="现在怎么做：")
    _bullet(doc, BOSS["ratio"], bold_prefix="30% 怎么理解：")
    _bullet(doc, BOSS["formula"], bold_prefix="怎么切：")
    _heading(doc, "刘孝春应当分担的项目清单", 2)
    item_rows = [[x["name"], x["now"], x["how"], x["liu"], x["no"]] for x in BOSS_ITEMS]
    _table(doc, ["项目", "现在能否主张", "计算口径", "刘孝春怎么担", "不要搞错"], item_rows)
    _heading(doc, "刘孝春不应承担的", 2)
    for item in BOSS_NOT:
        _bullet(doc, item)
    _heading(doc, "待补充信息（请直接补事实）", 2)
    _body(doc, "下列空白没有就不要编进 3D 或对外口径。有新事实发文字即可写入。")
    gap_rows = [[g["cat"], g["have"], g["need"], g["who"], g["u"]] for g in GAPS]
    _table(doc, ["类别", "已经掌握", "还缺什么", "找谁要", "缓急"], gap_rows)

    _heading(doc, "九、立即八事")
    _body(doc, "与前文「沟通记录去重」中的立即八事为同一张表，此处不重复。转院类录音不再逐份展开。")

    _heading(doc, "十、口径卡")
    for title, items in (
        ("对交警", TALKING_POINTS["to_police"]),
        ("对骑手/平台", TALKING_POINTS["to_rider"]),
        ("对医院", TALKING_POINTS["to_hospital"]),
        ("对保险公司", TALKING_POINTS["to_insurer"]),
        ("绝对不要说", TALKING_POINTS["do_not_say"]),
    ):
        _heading(doc, title, 2)
        for item in items:
            _bullet(doc, item)

    _heading(doc, "十一、录音来源")
    call_rows = [[c["date"], c["title"], c["mins"], c["use"]] for c in CALLS]
    _table(doc, ["日期", "录音标题", "时长", "本备忘录如何使用"], call_rows)
    _body(doc, "得到大脑转写有口误和串台。凡与齐鲁医院已审核 CT 冲突的，一律以 CT 为准；凡与交警监控结论冲突的，一律以交警为准。")
    _body(doc, f"{SOURCE_PACK['asr_date']} 另归档《{SOURCE_PACK['asr_file']}》和《{SOURCE_PACK['asr_guide']}》。{SOURCE_PACK['asr_note']}{SOURCE_PACK['video_status']}")

    _heading(doc, "十二、" + QINGXIAN["title"])
    for item in QINGXIAN["points"]:
        _bullet(doc, item)

    _heading(doc, "十三、使用限制")
    _body(doc, "本文根据 2026 年 8 月 15–19 日通话及转写整理，并核对用户提供的齐鲁医院云影像检查列表。影像所见和诊断摘自医院已审核报告原文。伤残部分只说明「尚未鉴定、现在不能定级」，不构成评残意见。二甲报销以保险公司确认为准，不以朋友咨询为准。赔偿金额、责任比例以交警认定书、保险合同和法院/鉴定机构为准。重新生成：python3 scripts/build_all.py。3D 复原视频本次不重做，留作待用。")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    build_document(root / "deliverables" / "青岛红枫路交通事故_伤情与处理备忘录_20260817.docx")
