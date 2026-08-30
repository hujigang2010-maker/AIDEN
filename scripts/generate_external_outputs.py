#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成对外完整方案 + 一页纸方案。

融合《10月31日活动联合运营合作方案_双版本》对外部分，
并写入基于 2026 年宏观与可支配收入的票价建议。
"""

import sys
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
import generate_wangdefeng_proposal as g  # noqa: E402

OUT = Path(__file__).resolve().parents[1] / "output"


def set_cell_margins(cell, top=40, bottom=40, left=70, right=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for edge, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def compact_table(doc, headers, rows, widths, *, header_fill=g.NAVY_HEX, size=9, pad=30):
    table = g.add_table(doc, headers, rows, widths, header_fill=header_fill)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=pad, bottom=pad, left=60, right=60)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    run.font.size = Pt(size)
    return table


def build_external_full():
    doc = g.new_doc()
    g.setup_section(doc, "10 月 31 日王德峰老师活动 · 联合运营合作方案（对外沟通版）")

    g.add_para(doc, "", size=10, space_after=10)
    g.add_para(doc, "PROJECT BRIEF", size=11, bold=True, color=g.GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    g.add_para(doc, "10 月 31 日王德峰老师活动", size=22, bold=True, color=g.NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    g.add_para(doc, "联合运营合作方案", size=26, bold=True, color=g.NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    g.add_para(doc, "对外沟通版", size=14, bold=True, color=g.GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    g.add_quote_box(
        doc,
        "合作定位",
        "双方共同成立 10 月 31 日活动独立项目：合作方提供老师及内容 / IP 资源，我方负责商业化、组织、渠道与现场交付；双方依据资源贡献、获客来源、成本承担及履约责任进行结算。",
    )
    g.add_para(
        doc,
        "角色：联合主办 / 联合运营方  ·  非渠道代理",
        size=12,
        bold=True,
        color=g.NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    g.add_para(
        doc,
        "版本日期：2026 年 8 月 14 日　　文件属性：商业合作讨论稿｜非最终合同",
        size=10,
        color=g.MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    g.add_para(
        doc,
        "活动名称、地点、容量等尚未完整确定，正文以“待确认”标示。分成比例与票价为建议，最终以正式合同为准。",
        size=9.5,
        color=g.MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.add_page_break()

    g.add_h1(doc, "一、合作目标")
    g.add_bullets(
        doc,
        [
            "以安全、合规、专业的标准完成 10 月 31 日活动，兼顾品牌口碑与商业结果。",
            "通过清晰分工和透明结算，让老师 / IP 价值与运营执行价值均得到合理回报。",
            "以本次活动作为一期样板，成功后优先探讨后续人文、哲学及相关知识产品合作。",
        ],
    )

    g.add_h1(doc, "二、合作原则")
    g.add_table(
        doc,
        ["原则", "具体含义"],
        [
            ["身份对等", "我方作为联合主办 / 联合运营方参与项目，不仅承担票务渠道职能。"],
            ["按贡献分配", "区分客户来源、招商来源、成本承担和实际履约，分别核算。"],
            ["风险共担", "不设置与项目确定性不匹配的大额前置买断；重大风险按约定分担。"],
            ["数据透明", "统一票务及对账口径，双方可查验报名、退款、赠票和渠道数据。"],
            ["长期保护", "明确客户、社群、企业及赞助资源的归属和保护期，避免绕单。"],
        ],
        [4.2, 12.2],
    )

    g.add_h1(doc, "三、双方分工")
    g.add_table(
        doc,
        ["工作模块", "合作方主要职责", "我方主要职责"],
        [
            ["核心资源", "落实老师出席；确认主题、内容边界、活动形式及授权链", "根据市场反馈协助产品化、场次设计与商业包装"],
            ["场地与制作", "配合确认老师端要求", "场地筛选与谈判、舞美视听、物料、供应商和现场统筹"],
            ["营销与销售", "提供可用的内容素材及自有渠道支持", "营销方案、海报文案、票务、社群、企业客户及渠道管理"],
            ["招商与商务", "配合权益确认和必要的老师端审核", "招商产品设计、客户开发、谈判、签约及权益交付"],
            ["客户与现场", "老师接待及内容环节配合", "报名咨询、开票 / 退款协同、签到、安保、应急及复盘"],
            ["财务与数据", "核对对方来源订单及应结款项", "建立统一台账，按周期出具销售、成本和结算报告"],
        ],
        [3.4, 6.5, 6.5],
    )

    g.add_h1(doc, "四、推荐商务结构：按来源分成")
    g.add_para(
        doc,
        "以下为用于启动讨论的建议比例。双方可结合最终成本责任、资源投入及独家程度，在区间内确认。不要笼统谈“利润五五分”或“流水五五分”，应按照收入来源、谁承担成本、谁负责履约分别结算。",
        first_line=True,
    )
    g.add_table(
        doc,
        ["收入来源", "我方建议比例", "合作方建议比例", "识别口径"],
        [
            ["我方独立带来的个人报名", "60%–70%", "30%–40%", "我方专属链接 / 二维码 / 邀请码及可核验线索"],
            ["合作方独立带来的个人报名", "20%–30%", "70%–80%", "合作方专属链接 / 二维码 / 邀请码"],
            ["双方共同渠道报名", "50%", "50%", "联合投放、共同活动或无法合理单独归因"],
            ["我方引入的赞助", "70%–80%", "20%–30%", "我方首次引荐并主导成交、交付"],
            ["合作方引入的赞助", "20%–30%", "70%–80%", "合作方首次引荐并主导成交"],
            ["我方开发的企业 / 团体客户", "60%–70%", "30%–40%", "我方开发、谈判并维护的组织客户"],
        ],
        [4.6, 3.2, 3.4, 5.2],
        emphasize_col=1,
    )
    g.add_quote_box(
        doc,
        "建议书面起点",
        "我方独立开发的报名和客户，建议落在 60% : 40%。该比例覆盖获客、活动运营、场地统筹、商务、客户服务、现场执行及风险承担，不是渠道佣金。合作方自有渠道带来的报名，合作方拿大头。",
        fill=g.SOFT_HEX,
        title_color=g.NAVY,
    )

    g.add_h1(doc, "五、票价建议（供双方确认）")
    g.add_para(
        doc,
        "本建议针对上海、北京、深圳、杭州等一线及强一线城市，活动时长约 2.5–3 小时的公开人文讲座。核心判断：主力票价定在 499 元，整体平均实收客单价控制在 500–550 元。不建议把 699 元作为全场基础票价，也不建议普通票超过 599 元。场地、容量确定后锁定最终票档。",
        first_line=True,
    )
    g.add_quote_box(
        doc,
        "对外主价格",
        "用 299 元和 399 元降低第一次购买门槛，用 499 元做主力收入，用 699 元和 999 元承接王德峰老师的核心粉丝与高净值客户。499 元是最适合对外传播的标准价格。",
    )

    g.add_h2(doc, "5.1 一线及强一线城市推荐票档")
    g.add_table(
        doc,
        ["票档", "建议价格", "座位占比", "功能"],
        [
            ["限量引流票", "299 元", "10%", "早鸟、后区座位，制造开售速度"],
            ["早鸟普通票", "399 元", "15%", "开售前 7–10 天限量"],
            ["标准票", "499 元", "40%", "核心销售票档，对外传播主价格"],
            ["优选票", "699 元", "25%", "前区座位、赠书或资料"],
            ["VIP 票", "999 元", "10%", "前排、专属问答、签名书或合影权益"],
        ],
        [3.6, 3.2, 2.8, 6.8],
        emphasize_col=1,
    )
    g.add_para(
        doc,
        "按上表座位占比，票面加权均价约为 559 元。考虑早鸟、团购、渠道优惠、退款等因素，实际收到手里的平均客单价大约在 500–530 元。目标是把平均实收客单价控制在 500–550 元。",
        first_line=True,
    )

    g.add_h2(doc, "5.2 为什么 499 元比较合适")
    g.add_para(
        doc,
        "国家统计局 2026 年上半年：全国居民人均可支配收入 22,981 元，中位数 19,036 元；城镇居民人均可支配收入 30,126 元，中位数 26,389 元。换算下来，城镇居民月均可支配收入中位数约 4,398 元。",
        first_line=True,
    )
    g.add_table(
        doc,
        ["票价", "约占城镇月可支配收入中位数", "对购买决策的含义"],
        [
            ["399 元", "约 9.1%", "认真想一想会买，适合早鸟"],
            ["499 元", "约 11.3%", "需要认真决策的文化消费，适合做主力"],
            ["699 元", "约 15.9%", "进入粉丝型、高收入型区间，不宜做全场基础价"],
            ["999 元", "约 22.7%", "必须配套前排、问答、签名或合影等可见权益"],
        ],
        [3.4, 5.8, 7.2],
    )
    g.add_para(
        doc,
        "同期全国居民人均教育文化娱乐支出只有 1,572 元，城镇居民为 1,952 元，折合每月约 262 元和 325 元。一张 499 元门票，已经相当于普通城镇居民约一个半月的教育文化娱乐预算。2026 年上半年城镇居民消费支出名义增长 3.0%、实际增长 2.0%，社会消费品零售增长也比较温和。消费者并非完全不消费，但对非刚需项目会明显比较价格和体验。",
        first_line=True,
    )
    g.add_para(
        doc,
        "因此，对一次约 3 小时、没有学历证书或长期服务的公开人文讲座而言，499 元已经到位；699 元以上应作为优选 / VIP，而不是普通票。",
        first_line=True,
    )

    g.add_h2(doc, "5.3 老师 IP 与价格的边界")
    g.add_para(
        doc,
        "王德峰老师的 IP 可以支撑更高价格，但要区分“公开讲座”和“深度课程”。公开资料中，相关长期课程存在 19,800 元、25,800 元等价格，通常是多天或数月的系统课程，面向企业家、高管及深度学习人群，不能直接证明单场讲座应该卖 999 元以上。市场上也存在免费或低价公开课信息，说明普通票必须让人一眼看到“这一场为什么值得买”。",
        first_line=True,
    )
    g.add_para(doc, "如果 10 月 31 日只是“老师讲 3 小时”，没有以下附加价值，普通票卖 699 元会比较吃力：", first_line=True)
    g.add_bullets(
        doc,
        [
            "全新主题或具有稀缺性的内容；",
            "较长问答和现场互动；",
            "签名书、讲义或独家资料；",
            "前排、小范围交流或合影；",
            "高规格场地和完整舞台体验；",
            "明确的“本年度少有公开场次”；",
            "企业家、人文社群等高质量观众圈层。",
        ],
    )
    g.add_para(
        doc,
        "如果没有赠书、问答、合影、独家内容等权益，VIP 不要超过 999 元。如果只是普通阶梯教室或会议厅讲座，主力价格应进一步降到 399 元。",
        first_line=True,
    )

    g.add_h2(doc, "5.4 不同城市的建议")
    g.add_table(
        doc,
        ["城市类型", "主力标准票", "推荐票档"],
        [
            ["上海、北京、深圳、杭州", "499 元", "299 / 399 / 499 / 699 / 999"],
            ["南京、苏州、广州、成都等强二线", "399 元", "299 / 399 / 599 / 799"],
            ["普通二三线城市", "299 元", "199 / 299 / 499 / 699"],
            ["企业家定向、小规模精品场", "999 元以上", "999 / 1,599 / 2,999，须含深度交流与圈层服务"],
        ],
        [5.6, 3.6, 7.2],
        emphasize_col=1,
    )

    g.add_h2(doc, "5.5 票价不能为了覆盖买断费而硬抬")
    g.add_para(
        doc,
        "假设项目总固定成本仍达 25 万元，平均实收客单价按 520 元计，扣除票务平台、退款、税费和渠道成本后，单个付费用户实际贡献可能只有 460–480 元。保本大约需要 250,000 ÷ 470 ≈ 532 名付费用户。若预计只能卖出场馆容量的 80%，场馆至少需要约 665 座，才有机会单靠票务覆盖 25 万元成本。",
        first_line=True,
    )
    g.add_para(
        doc,
        "这意味着：票价不能为了覆盖 10–18 万元买断费而硬抬到 699–999 元。成本结构不合理，不应该让消费者通过高票价替主办方承担。若是 500 座场馆、总成本 25 万元，即使平均实收 520 元，也几乎没有利润空间。这样的项目必须满足至少一个条件：取消大额买断；将总固定成本压到 12–15 万元以内；获得 5–10 万元赞助；引入企业团购；场馆达到 650–800 座；或将活动升级为半天 / 全天精品，再提高客单价。",
        first_line=True,
    )

    g.add_h2(doc, "5.6 先小规模预售，再锁死标准价")
    g.add_quote_box(
        doc,
        "测试方法",
        "确定票价前，先释放 50 张 399 元早鸟票。若 48 小时内售出 70% 以上，再维持 499 元标准价；若售出不足 30%，说明市场真实接受价更接近 299–399 元，而不是 699 元。",
        fill=g.SOFT_HEX,
        title_color=g.NAVY,
    )

    g.add_h2(doc, "5.7 晚宴、企业包桌作为附加产品（待确认）")
    g.add_para(
        doc,
        "若另设晚间私享或企业包桌，应与 3 小时公开讲座分开定价，不把晚宴成本摊进 499 元标准票。晚宴私享可单列 1,280–1,680 元；企业包桌（8–10 人）1.2 万–1.8 万 / 桌。个人票负责场面和口碑，企业桌和赞助负责覆盖场地与老师相关成本。",
        first_line=True,
    )

    g.add_h1(doc, "六、备选商务结构")
    g.add_h2(doc, "6.1 备选 A：阶梯式票务分成")
    g.add_para(
        doc,
        "不支付 10–18 万元固定买断费；按活动累计票务实收收入分段计算，既保障冷启动期运营投入，也让项目成功时 IP 方获得更高回报。",
        first_line=True,
    )
    g.add_table(
        doc,
        ["累计票务实收区间", "我方比例", "合作方比例"],
        [
            ["0–10 万元（含）", "70%", "30%"],
            ["10–30 万元（含）部分", "60%", "40%"],
            ["30 万元以上部分", "50%", "50%"],
        ],
        [7.0, 4.7, 4.7],
        header_fill="3D5A40",
    )
    g.add_para(
        doc,
        "“实收收入”建议定义为已收取且超过约定退款观察期的口径，并在合同中明确平台手续费、退款及税费的处理方式。",
        first_line=True,
        size=10.5,
        color=g.MUTED,
    )

    g.add_h2(doc, "6.2 备选 B：小额基础保障 + 高比例分成")
    g.add_para(
        doc,
        "如合作方确需基础保障，可讨论 3–5 万元以内的老师 / 授权相关基础成本，但须与明确的交付节点、退款机制、独家权益及销售分成绑定。",
        first_line=True,
    )
    g.add_table(
        doc,
        ["触发节点", "建议支付比例", "前置条件"],
        [
            ["合同生效", "不超过 20%", "老师书面确认、授权链及基本活动方案齐备"],
            ["达到可开售条件", "不超过 30%", "场地 / 票务 / 宣传素材通过确认，可正式售票"],
            ["老师完成活动", "余额 50% 以上", "按约出席并完成约定内容，无重大违约"],
        ],
        [4.2, 4.2, 8.0],
    )
    g.add_quote_box(
        doc,
        "不建议结构",
        "由我方单独支付 10–18 万元固定授权费，并同时承担场地、宣传、制作、人员、退款等全部经营风险。该结构需在完整独家权利和充分历史数据支持下另行评估。",
        fill=g.ROSE_HEX,
        title_color=g.RED,
    )

    g.add_h1(doc, "七、成本、结算与数据机制")
    g.add_bullets(
        doc,
        [
            "预算先批：活动启动前由双方确认预算表，超预算支出须经双方书面确认。",
            "成本归属清晰：按“谁提出、谁受益、谁确认”界定承担方；共同成本按约定比例分摊。",
            "统一台账：票务、企业客户、赞助、退款、赠票、渠道费和开票信息进入同一项目台账。",
            "定期对账：建议开售后每周对账，活动结束后 10 个工作日内完成初步结算。",
            "现金流隔离：如条件允许，使用项目专用收款账户或双方均可查验的票务后台。",
        ],
    )

    g.add_h1(doc, "八、关键权利与风险条款")
    g.add_table(
        doc,
        ["条款主题", "建议约定"],
        [
            ["老师确认与取消", "老师出席、内容和时长须书面确认；因老师端原因取消 / 延期，已付授权或保障款 100% 退还，并明确已发生不可回收成本的承担。"],
            ["宣传授权", "明确姓名、肖像、简介、海报、短视频、直播 / 录播片段的可用范围、渠道、期限及审批时限。"],
            ["独家与冲突", "明确活动在地域、时间、主题和渠道上的独家范围，并约定冲突活动的处理。"],
            ["客户与数据", "报名数据仅为本项目及经同意的后续转化使用；数据访问、保管、隐私合规及删除机制清晰。"],
            ["资源保护 / 防绕单", "各自引入的企业、赞助商及渠道设 12 个月保护期，未经引入方书面同意不得绕开成交。"],
            ["社群与二次转化", "微信群 / 社群管理权、后续课程推广权和收益分配另行明确，不因本次合作自动转移。"],
            ["内容成果", "活动照片、视频、录音、课件及剪辑内容的著作权、使用权和商业化权利分别约定。"],
            ["安全与合规", "场地报批、消防、安保、广告宣传、发票税务、个人信息保护等责任到人。"],
        ],
        [4.4, 12.0],
    )

    g.add_h1(doc, "九、项目推进建议")
    g.add_table(
        doc,
        ["步骤", "事项", "目标结果"],
        [
            ["1", "确认合作结构及授权代表", "书面确认联合运营身份，必要时签署保密协议"],
            ["2", "合作方提供老师确认与授权链", "出席书面确认、内容边界、历史活动数据"],
            ["3", "我方提交场地、票价、营销初稿", "确认 299 / 399 / 499 / 699 / 999，并安排 50 张早鸟测试"],
            ["4", "确认归因、预算、后台和结算", "收入来源可核验，成本可审批"],
            ["5", "签署正式协议及附件", "达到可开售条件后统一对外发布"],
            ["6", "按周复盘销售与风险", "活动结束后完成结算及二期合作评估"],
        ],
        [2.0, 6.4, 8.0],
    )

    g.add_h1(doc, "十、需双方确认的开放事项")
    g.add_table(
        doc,
        ["事项", "待确认内容"],
        [
            ["活动基础信息", "正式名称、日期 / 时段、城市、场地、容量、形式、主题、时长"],
            ["票务", "最终票档、赠票上限、早鸟 / 团购规则、退改政策、开票主体"],
            ["IP 与内容", "老师确认文件、授权主体、宣传素材、审核流程、录音录像权限"],
            ["商务", "最终分成、税费口径、成本分担、回款节点、独家范围"],
            ["运营", "项目负责人、审批时限、重大事项决策机制、应急预案"],
            ["长期合作", "后续课程优先合作权、客户二次转化及收益安排"],
        ],
        [4.4, 12.0],
    )

    g.add_h1(doc, "十一、给合作方的一页共识")
    g.add_bullets(
        doc,
        [
            "请把我方看成联合运营方，而不是卖票渠道。",
            "10 月 31 日用按来源分成，而不是 10–18 万元买断。",
            "我方独立开发的报名建议 60% : 40%，我方引入的赞助建议 70% : 30%；合作方自己的报名，合作方拿大头。",
            "一线城市 3 小时公开讲座：299 / 399 / 499 / 699 / 999，主力 499 元，平均实收 500–530 元。",
            "这一场是第一期。做成了，再一起做后续哲学、人文和相关知识产品。",
        ],
    )
    g.add_quote_box(
        doc,
        "下一步建议",
        "双方先用一次 60–90 分钟工作会议确认合作结构、票价方向和开放事项，再由双方指定负责人将共识转化为正式合同与项目排期。",
    )
    g.add_para(
        doc,
        "（本文件为沟通方案，不构成要约。最终合作内容以双方签署的协议为准。正式签署前建议由专业律师结合交易主体及所在地法规审阅。）",
        size=9.5,
        color=g.MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12,
    )

    path = OUT / "王德峰老师10月31日活动_联合运营合作方案_对外完整版.docx"
    doc.save(path)
    return path


def _p(cell, text, *, size=8.5, bold=False, color=g.INK, after=1, before=0, align=None):
    if cell.paragraphs and not cell.paragraphs[0].text and len(cell.paragraphs) == 1:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    g.set_run_font(run, size=size, bold=bold, color=color)
    return p


def _mini_table(doc, headers, rows, widths, header_fill=g.NAVY_HEX, size=8):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        g.shade_cell(cell, header_fill)
        set_cell_margins(cell, top=28, bottom=28, left=50, right=50)
        _p(cell, header, size=size, bold=True, color=g.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    for r_idx, row in enumerate(rows):
        fill = "FFFFFF" if r_idx % 2 == 0 else g.ROW_ALT_HEX
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            g.shade_cell(cell, fill)
            set_cell_margins(cell, top=24, bottom=24, left=50, right=50)
            _p(
                cell,
                str(val),
                size=size,
                bold=c_idx == 0,
                color=g.NAVY if c_idx == 1 else g.INK,
                align=WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT,
                after=0,
            )
        g.prevent_row_split(table.rows[r_idx + 1])
    g.set_table_full_width(table, widths)
    return table


def build_onepager():
    doc = g.new_doc()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.header_distance = Cm(0.4)
    section.footer_distance = Cm(0.35)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("对外沟通 · 一页纸  ·  讨论稿  ·  2026年8月14日")
    g.set_run_font(run, size=8, color=g.MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("本页供合作方快速对齐。细节见《联合运营合作方案（对外完整版）》。不构成要约，最终以合同为准。")
    g.set_run_font(fr, size=7.5, color=g.MUTED)

    g.add_para(
        doc,
        "10 月 31 日王德峰老师活动｜联合运营合作方案",
        size=15,
        bold=True,
        color=g.NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0,
        space_after=2,
        line_spacing=1.0,
    )
    g.add_para(
        doc,
        "合作方提供老师及内容 / IP  ·  我方负责商业化、组织、渠道与现场交付  ·  按来源分成，不买断",
        size=9,
        bold=True,
        color=g.GOLD,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
        line_spacing=1.0,
    )

    box = doc.add_table(rows=1, cols=1)
    cell = box.cell(0, 0)
    g.shade_cell(cell, g.CREAM_HEX)
    g.set_cell_border(
        cell,
        top={"val": "single", "sz": "6", "color": g.GOLD_HEX, "space": "0"},
        left={"val": "single", "sz": "16", "color": g.GOLD_HEX, "space": "0"},
        bottom={"val": "single", "sz": "6", "color": g.GOLD_HEX, "space": "0"},
        right={"val": "single", "sz": "6", "color": g.GOLD_HEX, "space": "0"},
    )
    cell.text = ""
    set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
    _p(cell, "一句话定位", size=8, bold=True, color=g.GOLD, after=1)
    _p(
        cell,
        "我方不是票务渠道，而是联合主办 / 联合运营方。10 月 31 日作为独立项目和第一期样板：按收入来源、成本承担和履约责任结算，不做 10–18 万元一次性买断。",
        size=9,
        after=1,
    )
    g.set_table_full_width(box, [18.6])

    g.add_para(doc, "1. 双方分工", size=10.5, bold=True, color=g.NAVY, space_before=5, space_after=2, line_spacing=1.0)
    _mini_table(
        doc,
        ["模块", "合作方", "我方"],
        [
            ["IP / 内容", "老师出席、主题、授权链", "产品化、场次与商业包装"],
            ["场地 / 现场", "配合老师端要求", "场地、制作、人员、交付"],
            ["销售 / 招商", "自有渠道与素材", "票务、企业客户、赞助、社群"],
            ["财务 / 数据", "核对本方来源订单", "统一台账、对账、结算报告"],
        ],
        [3.4, 7.6, 7.6],
    )

    g.add_para(doc, "2. 按来源分成（建议起点）", size=10.5, bold=True, color=g.NAVY, space_before=5, space_after=2, line_spacing=1.0)
    _mini_table(
        doc,
        ["收入来源", "我方", "合作方"],
        [
            ["我方独立报名 / 企业客户", "60%", "40%"],
            ["合作方独立报名", "25%", "75%"],
            ["共同渠道", "50%", "50%"],
            ["我方引入赞助", "70%", "30%"],
            ["合作方引入赞助", "25%", "75%"],
        ],
        [8.6, 5.0, 5.0],
        header_fill=g.NAVY_HEX,
    )
    g.add_para(
        doc,
        "备选：票务实收 0–10 万部分 70%:30%，10–30 万部分 60%:40%，30 万以上 50%:50%。如需保障，可讨论 3–5 万元以内、分期、取消全退，不建议大额买断。",
        size=8,
        color=g.MUTED,
        space_before=3,
        space_after=3,
        line_spacing=1.08,
    )

    g.add_para(doc, "3. 票价建议（一线城市 · 2.5–3 小时公开讲座）", size=10.5, bold=True, color=g.NAVY, space_before=5, space_after=2, line_spacing=1.0)
    g.add_para(
        doc,
        "城镇月可支配收入中位数约 4,398 元。499 元约占 11.3%，适合做主力；699 元约占 15.9%，不宜做全场基础价。票面加权约 559 元，预计平均实收 500–530 元。",
        size=8.5,
        space_before=0,
        space_after=3,
        line_spacing=1.08,
    )
    _mini_table(
        doc,
        ["票档", "价格", "占比", "功能"],
        [
            ["限量引流票", "299 元", "10%", "后区，制造开售速度"],
            ["早鸟普通票", "399 元", "15%", "开售前 7–10 天限量"],
            ["标准票", "499 元", "40%", "对外传播主价格"],
            ["优选票", "699 元", "25%", "前区 + 赠书 / 资料"],
            ["VIP 票", "999 元", "10%", "前排、问答、签名或合影"],
        ],
        [3.8, 3.2, 2.4, 9.2],
        header_fill="3D5A40",
    )
    g.add_para(
        doc,
        "先放 50 张 399 元早鸟：48 小时售出 ≥70% 则维持 499 标准价；不足 30% 则市场更接近 299–399 元。晚宴 / 企业桌另计，不摊进标准票。",
        size=8,
        color=g.MUTED,
        space_before=3,
        space_after=3,
        line_spacing=1.08,
    )

    g.add_para(doc, "4. 必须写入合同的底线", size=10.5, bold=True, color=g.NAVY, space_before=5, space_after=2, line_spacing=1.0)
    _mini_table(
        doc,
        ["主题", "约定要点"],
        [
            ["老师确认 / 取消", "书面确认出席；老师端取消则已付保障款 100% 退还"],
            ["宣传授权", "姓名、肖像、海报、短视频可用范围和审批时限"],
            ["数据与防绕单", "项目新客户共有；各自引入资源设 12 个月保护期"],
            ["结算", "统一台账，开售后按周对账，结束后 10 个工作日内初结"],
        ],
        [4.2, 14.4],
    )

    g.add_para(doc, "5. 建议立即确认的三件事", size=10.5, bold=True, color=g.NAVY, space_before=5, space_after=2, line_spacing=1.0)
    next_box = doc.add_table(rows=1, cols=3)
    items = [
        ("① 身份", "书面确认我方为联合运营方，按来源分成，不采用 10–18 万买断。"),
        ("② 票价", "按 299 / 399 / 499 / 699 / 999 启动，主力 499 元；先测 50 张早鸟。"),
        ("③ 授权", "一周内提供老师出席书面确认、授权链，达到可开售讨论条件。"),
    ]
    fills = [g.SOFT_HEX, g.CREAM_HEX, g.GREEN_HEX]
    for i, ((title, body), fill) in enumerate(zip(items, fills)):
        c = next_box.cell(0, i)
        c.text = ""
        g.shade_cell(c, fill)
        set_cell_margins(c, top=50, bottom=50, left=70, right=70)
        _p(c, title, size=9, bold=True, color=g.NAVY, after=2)
        _p(c, body, size=8, after=1)
    g.set_table_full_width(next_box, [6.2, 6.2, 6.2])

    path = OUT / "王德峰老师10月31日活动_联合运营合作方案_一页纸.docx"
    doc.save(path)
    return path


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    full = build_external_full()
    one = build_onepager()
    print(f"已生成对外完整版：{full}")
    print(f"已生成一页纸：{one}")


if __name__ == "__main__":
    main()
