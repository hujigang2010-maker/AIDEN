"""按盖章版《活动赞助合同》文字信息补充《晚宴冠名战略合作伙伴专项合作协议》，改动红字标注，格式不变。"""
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn

SRC = Path("/home/ubuntu/.cursor/projects/workspace/uploads/____-____-_________________2_5973.docx")
OUT = Path(
    "/workspace/deliverables/"
    "附件四-绿城中国-潮鸣外滩-晚宴冠名战略合作伙伴专项合作协议-按盖章版补充(红字标注).docx"
)
RED = RGBColor(0xFF, 0x00, 0x00)


def replace_paragraph_text(paragraph, new_text, color=RED, bold=None):
    size = None
    name = None
    first_bold = None
    for r in paragraph.runs:
        if size is None and r.font.size:
            size = r.font.size
        if name is None and r.font.name:
            name = r.font.name
        if first_bold is None and r.bold is not None:
            first_bold = r.bold
    if paragraph.runs:
        first = paragraph.runs[0]
        first.text = new_text
        if size:
            first.font.size = size
        if name:
            first.font.name = name
            rFonts = first._r.get_or_add_rPr().get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), name)
        first.font.color.rgb = color
        if bold is not None:
            first.bold = bold
        elif first_bold is not None:
            first.bold = first_bold
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        run = paragraph.add_run(new_text)
        run.font.color.rgb = color
        if bold is not None:
            run.bold = bold


def set_cell_text(cell, new_text, color=RED):
    if not cell.paragraphs:
        return
    replace_paragraph_text(cell.paragraphs[0], new_text, color=color)
    for p in cell.paragraphs[1:]:
        for r in p.runs:
            r.text = ""


def main():
    if not SRC.exists():
        raise SystemExit(f"source missing: {SRC}")
    doc = Document(str(SRC))

    replace_paragraph_text(doc.paragraphs[4], "协议编号：GTH-2026-潮鸣外滩-SP-012-附件四")
    replace_paragraph_text(doc.paragraphs[5], "签订地点：上海市虹口区")
    replace_paragraph_text(doc.paragraphs[6], "签订日期：2026年5月21日")
    replace_paragraph_text(
        doc.paragraphs[9],
        "1. 甲方已于 2026 年 5 月 22 日在上海·北外滩·一滴水主办「2026 人工智能商业化落地与硬核投资破局峰会」（以下简称「本次峰会」）；",
    )
    replace_paragraph_text(
        doc.paragraphs[10],
        "2. 乙方系上海绿城泓盛建设发展有限公司及其关联项目「绿城·潮鸣外滩」的项目主体/合作品牌方，愿意作为本次峰会「晚宴冠名战略合作伙伴」提供赞助；",
    )
    replace_paragraph_text(
        doc.paragraphs[11],
        "3. 双方已就赞助合作的核心商业要素达成一致，并已签署《活动赞助合同》（合同编号：GTH-2026-潮鸣外滩-SP-012，以下简称「主协议」/「盖章版活动赞助合同」）；",
    )

    t0 = doc.tables[0]
    set_cell_text(t0.rows[1].cells[1], "（盖章版未载明，签署时补填）")
    set_cell_text(t0.rows[2].cells[1], "上海市杨浦区国泰路11号复旦科技园副楼一楼1008室")
    set_cell_text(t0.rows[3].cells[1], "朱震 / 法定代表人")
    set_cell_text(t0.rows[4].cells[1], "胡继刚  联系电话：13262607888")
    set_cell_text(t0.rows[5].cells[1], "（盖章版未载明）")

    t1 = doc.tables[1]
    set_cell_text(t1.rows[0].cells[1], "上海绿城泓盛建设发展有限公司（项目主体）")
    set_cell_text(t1.rows[2].cells[1], "91310107MA1G0H6X8C")
    set_cell_text(t1.rows[3].cells[1], "上海市虹口区四川北路1688号2415室（集中登记地）")
    set_cell_text(t1.rows[4].cells[1], "刘继涛 / 法定代表人")
    set_cell_text(t1.rows[5].cells[1], "笪浩")
    set_cell_text(t1.rows[6].cells[1], "18678408669")
    set_cell_text(t1.rows[7].cells[1], "（盖章版未载明）")

    replace_paragraph_text(
        doc.paragraphs[50],
        "（4）乙方指派代表专场宣讲 PPT（16:9 1920×1080，≤ 25 张，控制 15 分钟内）；",
    )
    replace_paragraph_text(
        doc.paragraphs[62],
        "7.1  鉴于本次峰会已于2026年5月22日实际完成，双方确认采用后置一次性付款：甲方应向乙方提交执行报告（含现场图片、嘉宾合影、媒体链接等代表性证据材料）；"
        "乙方应在收到执行报告后3个工作日内书面提出具体异议，逾期未提出书面异议的，视为验收合格；验收合格后15个工作日内，乙方将全部赞助款项人民币100,000元一次性汇入甲方指定账户。"
        "若本协议生效日期晚于上述付款期限届满之日，则付款期限相应顺延至本协议生效后15个工作日内。甲方指定账户如下：",
    )
    replace_paragraph_text(
        doc.paragraphs[63],
        "7.2  甲方系民办非企业单位，按小规模纳税人简易计税方式适用1%增值税征收率，在款项到账后10个工作日内向乙方开具等额合法有效的增值税普通发票（不开具增值税专用发票）；"
        "合同总价款100,000元含税不变（不含税价款约99,009.90元，增值税税金约990.10元）。"
        "发票内容以实际交易内容、税务资质及国家税法规定为准，常见可适用科目包括「会议服务费」/「会务咨询费」/「赞助费」。"
        "乙方不得以非实质性的发票形式问题拒付款项；仅就发票真伪、税号及票面金额一致性等实质性问题享有合理异议权，"
        "并应于收到发票之日起3个工作日内书面提出，逾期未提出视为发票接收无异议。",
    )

    t7 = doc.tables[7]
    set_cell_text(t7.rows[0].cells[1], "上海市杨浦区科技企业联合会")
    set_cell_text(t7.rows[1].cells[1], "中国农业银行股份有限公司上海营口支行")
    set_cell_text(t7.rows[2].cells[1], "03354200040012533")
    set_cell_text(t7.rows[3].cells[1], "（盖章版未载明甲方/主办方税号，开票时按甲方税务登记信息填写）")

    replace_paragraph_text(
        doc.paragraphs[77],
        "10.1  乙方因自身原因无故逾期支付赞助款项的，每逾期一日按逾期应付未付金额的千分之一向甲方支付违约金；"
        "逾期超过15个自然日的，甲方有权解除本协议并要求乙方一次性支付全部应付未付款项及违约金。",
    )
    replace_paragraph_text(
        doc.paragraphs[78],
        "10.2  甲方未按本协议第三条/第六条/附表一约定提供权益且无正当理由的，应在双方确认的未履行权益范围内承担相应责任，"
        "或在双方协商一致后以同等价值的下一届同类活动权益补偿。"
        "【风险对齐盖章版】除因甲方故意或重大过失导致核心权益无法交付外，甲方就本协议项下未履行权益向乙方承担的退款、违约金及赔偿总额"
        "不超过本协议赞助金额的30%（即不超过人民币30,000元）；甲方已提交执行报告并经视同验收的，乙方就个别细项的异议不影响无争议部分款项的支付。"
        "乙方拟从应付未付款项中扣除违约金或赔偿金的，应经甲方书面确认或经生效裁判文书确认后方可抵扣。",
    )
    replace_paragraph_text(
        doc.paragraphs[92],
        "13.2  因本协议引起的或与本协议有关的任何争议，双方应首先友好协商解决；协商不成的，任何一方均有权向甲方所在地（即上海市杨浦区）有管辖权的人民法院提起诉讼。",
    )
    replace_paragraph_text(
        doc.paragraphs[95],
        "14.1  本附件四作为主协议（《活动赞助合同》，合同编号：GTH-2026-潮鸣外滩-SP-012）的组成部分，与主协议具有同等法律效力。"
        "本附件与主协议正文条款不一致的，涉及款项支付、发票开具、验收与视同验收、违约责任上限、单方扣款限制等风险与结算事项，"
        "以盖章版《活动赞助合同》约定为准；涉及晚宴冠名专项权益细项、物料植入与执行节点的，以本附件及附表一约定为准。"
        "本附件未尽事宜，适用主协议约定及双方另行签订的书面补充协议。",
    )

    t8 = doc.tables[8]
    set_cell_text(t8.rows[1].cells[1], "上海绿城泓盛建设发展有限公司")
    set_cell_text(t8.rows[4].cells[0], "职务：法定代表人/授权代表")
    set_cell_text(t8.rows[4].cells[1], "职务：法定代表人/授权代表")
    set_cell_text(t8.rows[5].cells[0], "签字日期：2026年5月21日")
    set_cell_text(t8.rows[5].cells[1], "签字日期：2026年5月21日")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"saved: {OUT}")


if __name__ == "__main__":
    main()
