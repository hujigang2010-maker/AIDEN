#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《胡继刚｜百度百科词条草稿》Word 文档（首版精简提交稿）。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "胡继刚_百度百科词条草稿.docx"


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, *, size=12, bold=False, space_before=0, space_after=6, first_line=False, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading_custom(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12}
    colors = {
        1: RGBColor(0x1A, 0x3A, 0x5C),
        2: RGBColor(0x1A, 0x3A, 0x5C),
        3: RGBColor(0x33, 0x33, 0x33),
    }
    return add_para(
        doc,
        text,
        size=sizes.get(level, 12),
        bold=True,
        space_before=12 if level == 1 else 8,
        space_after=6,
        color=colors.get(level),
    )


def add_mixed_para(doc, segments, *, size=12, space_before=0, space_after=6, first_line=False):
    """segments: list of (text, bold)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        pf.first_line_indent = Cm(0.74)
    for text, bold in segments:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
    return p


def set_cell_text(cell, text, *, bold=False, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    # 封面标题
    add_para(doc, "胡继刚", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, color=RGBColor(0x1A, 0x3A, 0x5C))
    add_para(
        doc,
        "百度百科词条草稿（第一版·精简提交稿）",
        size=14,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
        color=RGBColor(0x33, 0x33, 0x33),
    )
    add_para(
        doc,
        "定位：复旦大学住房政策研究中心秘书长、上海市杨浦区科技企业联合会执行会长",
        size=10.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
        color=RGBColor(0x66, 0x66, 0x66),
    )
    add_para(
        doc,
        "说明：本稿整合政府官网、央广网、人民网、中新网及正规机构官网的公开报道，仅收录姓名与职务、活动可交叉印证的陈述。可直接用于百度百科创建/新增义项。灰色编辑说明部分请勿粘贴提交。",
        size=10.5,
        space_after=10,
        color=RGBColor(0x55, 0x55, 0x55),
    )

    # ========== 可提交正文 ==========
    add_heading_custom(doc, "一、可直接提交正文", level=1)

    add_heading_custom(doc, "词条名", level=2)
    add_para(doc, "胡继刚", size=12, space_after=4)
    add_para(
        doc,
        "如百度百科提示已有同名人物，选择「新增义项」，义项名称建议填写：复旦大学住房政策研究中心秘书长",
        size=10.5,
        color=RGBColor(0x55, 0x55, 0x55),
        space_after=8,
    )

    add_heading_custom(doc, "概述", level=2)
    add_para(
        doc,
        "胡继刚，复旦大学住房政策研究中心秘书长、上海市杨浦区科技企业联合会执行会长。曾参与不动产资产管理、新经济增长、人工智能商业化及中欧创新合作等主题的行业交流活动。〔1〕〔2〕〔3〕〔4〕〔5〕",
        first_line=True,
        space_after=8,
    )

    add_heading_custom(doc, "基本信息", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ("项目", "内容", "参考资料")
    rows = [
        ("中文名", "胡继刚", "—"),
        ("主要职务", "复旦大学住房政策研究中心秘书长", "〔1〕〔2〕〔3〕〔5〕"),
        ("社会职务", "上海市杨浦区科技企业联合会执行会长", "〔1〕〔4〕"),
    ]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            set_cell_text(table.rows[r].cells[c], val)
    add_para(
        doc,
        "首版暂不填写：国籍、出生日期、籍贯、学历、毕业院校、职称、职业、主要研究方向。现有报道可证明参与相关议题，尚不足以等同于稳定的个人研究方向；“国籍”若被要求提供直接来源，可不填。",
        size=10.5,
        color=RGBColor(0x55, 0x55, 0x55),
        space_before=6,
        space_after=8,
    )

    add_heading_custom(doc, "人物经历", level=2)
    experiences = [
        "2024年6月，胡继刚以复旦大学住房政策研究中心兼复旦MBA不动产资产管理协会秘书长身份，与上海市锦天城律师事务所高级合伙人顾晓共同主持“他山之石——中国投资者的海外不动产战略布局”主题分享活动。〔5〕",
        "2025年5月，胡继刚以复旦大学住房政策研究中心秘书长身份参加“2025全球新经济增长引擎峰会”，在圆桌对话环节与证券、经济及资产管理领域嘉宾围绕房地产市场、资本布局与产业创新等议题展开交流。〔2〕〔3〕",
        "2026年3月，胡继刚以上海市杨浦区科技企业联合会执行会长身份参加“北欧创新国际会客厅”揭牌活动，并围绕中欧创新合作机制建设、科创生态联动与企业国际化发展等议题进行交流。〔4〕",
        "2026年5月，胡继刚以复旦大学住房政策研究中心秘书长、上海市杨浦区科技企业联合会执行会长身份参加“2026人工智能商业化落地与硬核投资破局峰会”，参与“从算力引擎到新质资产（AI全产业链的商业化实战）”圆桌对话。〔1〕〔6〕",
    ]
    for text in experiences:
        add_para(doc, text, first_line=True, space_after=6)

    add_heading_custom(doc, "社会职务", level=2)
    duties = [
        "复旦大学住房政策研究中心秘书长。〔1〕〔2〕〔3〕〔5〕〔6〕",
        "复旦MBA不动产资产管理协会秘书长。〔5〕",
        "上海市杨浦区科技企业联合会执行会长。〔1〕〔4〕",
    ]
    for text in duties:
        add_para(doc, text, first_line=True, space_after=4)

    add_para(
        doc,
        "说明：首版不再单列「社会活动」章节，以免与「人物经历」重复；活动事实已按时间写入人物经历。",
        size=10.5,
        color=RGBColor(0x55, 0x55, 0x55),
        space_before=4,
        space_after=10,
    )

    # ========== 参考资料 ==========
    add_heading_custom(doc, "二、参考资料（提交时逐条挂脚注）", level=1)

    refs = [
        {
            "no": "〔1〕",
            "title": "探寻AI时代的超级个体与新质风口，2026人工智能商业化落地与硬核投资破局峰会在沪落幕",
            "source": "央广网",
            "date": "2026年5月26日",
            "url": "https://www.cnr.cn/shanghai/tt/20260526/t20260526_527636297.shtml",
            "support": "明确记载“复旦大学住房政策研究中心秘书长、上海市杨浦区科技企业联合会执行会长胡继刚”；记载其参与“从算力引擎到新质资产（AI全产业链的商业化实战）”圆桌对话，并以主办方代表身份发表观点。",
            "insert": "概述第一句后；人物经历第四段后；两项现任社会职务后。",
            "quote": "在第二场圆桌对话“从算力引擎到新质资产（AI全产业链的商业化实战）”中，复旦大学住房政策研究中心秘书长、上海市杨浦区科技企业联合会执行会长胡继刚与五位行业嘉宾……作为此次峰会的主办方代表，复旦大学住房政策研究中心秘书长、上海市杨浦区科技企业联合会执行会长胡继刚表示……",
        },
        {
            "no": "〔2〕",
            "title": "2025全球新经济增长引擎峰会在沪落幕",
            "source": "人民网",
            "date": "2025年5月30日",
            "url": "https://sh.people.com.cn/n2/2025/0530/c134768-41245071.html",
            "support": "明确记载“复旦大学住房政策研究中心秘书长胡继刚”参加圆桌对话，议题为房地产市场、资本布局与产业创新。",
            "insert": "概述后；人物经历第二段后；研究中心秘书长职务后。",
            "quote": "在圆桌对话环节，复旦大学住房政策研究中心秘书长胡继刚与中信建投证券房地产及建筑首席分析师竺劲、长城证券首席经济学家汪毅、泓塬资产董事长王维军围绕房地产市场、资本布局与产业创新展开讨论。",
        },
        {
            "no": "〔3〕",
            "title": "“创新应变·蓄力远航”2025全球新经济增长引擎峰会在沪举行",
            "source": "中新网上海",
            "date": "2025年5月26日",
            "url": "https://www.sh.chinanews.com.cn/chanjing/2025-05-26/136263.shtml",
            "support": "与人民网交叉印证秘书长身份及2025年峰会圆桌交流经历。",
            "insert": "人物经历第二段后（与〔2〕并用）。",
            "quote": "在圆桌对话环节，复旦大学住房政策研究中心秘书长胡继刚对话中信建投证券房地产及建筑首席分析师竺劲先生、长城证券首席经济学家汪毅先生以及泓塬资产董事长王维军先生，并重点围绕房地产市场、资本布局与产业创新的主题展开了探讨。",
        },
        {
            "no": "〔4〕",
            "title": "“北欧创新国际会客厅”落户我区 打造中欧科创合作新枢纽",
            "source": "上海市杨浦区人民政府门户网站",
            "date": "2026年3月10日",
            "url": "https://www.shyp.gov.cn/shypq/xwzx-bmdt/20260310/501286.html",
            "support": "明确记载“执行会长胡继刚”，以及2026年3月揭牌活动中的交流议题。",
            "insert": "概述第一句后；人物经历第三段后；执行会长职务后。",
            "quote": "区科技企业联合会会长王晓明、执行会长胡继刚……也围绕中欧创新合作机制建设、科创生态联动与企业国际化发展等议题进行了交流。",
        },
        {
            "no": "〔5〕",
            "title": "锦天城“他山之石—中国投资者的海外不动产战略布局”主题分享活动圆满举行",
            "source": "上海市锦天城律师事务所官网",
            "date": "2024年6月21日（报道活动日期：2024年6月19日）",
            "url": "https://www.allbrightlaw.com/CN/10482/ec4ff32d459f0010.aspx",
            "support": "明确记载主持人身份，以及“复旦大学住房政策研究中心兼复旦MBA不动产资产管理协会秘书长”公开职务。",
            "insert": "人物经历第一段后；研究中心秘书长、MBA协会秘书长职务后。",
            "quote": "复旦大学住房政策研究中心兼复旦MBA不动产资产管理协会秘书长胡继刚先生以及锦天城律师事务所高级合伙人顾晓律师担任本次活动主持人。",
        },
        {
            "no": "〔6〕",
            "title": "2026人工智能商业化落地与硬核投资破局峰会探寻AI时代的超级个体与新质风口",
            "source": "中新网上海",
            "date": "2026年5月26日",
            "url": "https://www.sh.chinanews.com.cn/chanjing/2026-05-26/146727.shtml",
            "support": "与央广网交叉印证秘书长身份及2026年人工智能峰会主办方代表公开发言；该文未同时写明执行会长职务，执行会长请以〔1〕〔4〕为准。",
            "insert": "人物经历第四段后（与〔1〕并用）。",
            "quote": "作为主办方代表，复旦大学住房政策研究中心秘书长胡继刚表示……",
        },
    ]

    for ref in refs:
        add_mixed_para(
            doc,
            [(f"{ref['no']} ", True), (ref["title"], True)],
            size=11,
            space_before=8,
            space_after=2,
        )
        add_para(doc, f"来源：{ref['source']}　发布时间：{ref['date']}", size=10.5, space_after=2)
        add_para(doc, f"链接：{ref['url']}", size=10.5, space_after=2)
        add_para(doc, f"支持内容：{ref['support']}", size=10.5, space_after=2)
        add_para(doc, f"建议插入位置：{ref['insert']}", size=10.5, space_after=2)
        add_para(doc, f"原文摘录：{ref['quote']}", size=10.5, space_after=4, color=RGBColor(0x44, 0x44, 0x44))

    # ========== 创建填写 ==========
    add_heading_custom(doc, "三、创建时填写内容", level=1)
    add_para(doc, "词条名称：胡继刚", space_after=4)
    add_para(doc, "义项名称（如需新增义项）：复旦大学住房政策研究中心秘书长", space_after=4)
    add_para(doc, "分类：人物（其他人物或相近行业人物分类）", space_after=4)
    add_para(
        doc,
        "提交理由：根据政府门户网站、中央及主流媒体、正规机构官网的公开报道，创建人物基础词条。内容仅收录有可靠公开来源支持的职务和公开活动经历。",
        space_after=8,
    )

    add_heading_custom(doc, "粘贴顺序", level=2)
    steps = [
        "创建人物词条，或在同名提示下选择「新增义项」。",
        "粘贴「概述」。",
        "添加「人物经历」「社会职务」两个一级目录。",
        "按上文插入位置逐条添加参考资料；同一段可挂两条交叉来源。",
        "只填写有来源的基本信息栏字段。",
        "预览时检查每个职务和每段经历末尾是否都有脚注。",
        "删除宣传性、评价性词语后提交。",
    ]
    for i, s in enumerate(steps, 1):
        add_para(doc, f"{i}. {s}", size=11, space_after=3)

    # ========== 编辑说明 ==========
    add_heading_custom(doc, "四、编辑说明（勿提交）", level=1)

    add_heading_custom(doc, "本版相对旧稿的主要优化", level=2)
    opts = [
        "概述改为“职务先行”一句定位，删除“产业运营及科技创新服务领域从业者”“长期参与/长期从事”等缺少连续公开履历支撑的概括，降低审核难度。",
        "补回并核验2025年全球新经济增长引擎峰会经历：人民网、中新网上海均可检索到实名与秘书长职务。",
        "纠正2026年人工智能峰会圆桌名称：以央广网为准，写为“从算力引擎到新质资产（AI全产业链的商业化实战）”，不再使用旧稿中的“从不良资产到新质资产”等不准确表述。",
        "纠正协办单位等活动背景信息：首版正文不展开主办协办名单，避免与不同报道表述差异纠缠；需要时以央广网、中新网原文为准。",
        "基本信息栏仅保留中文名、主要职务、社会职务；删除职业、研究方向等易被要求补证的字段。",
        "社会职务保留锦天城官网可证的“复旦MBA不动产资产管理协会秘书长”；不单列社会活动，避免与人物经历重复。",
        "参考资料由4条增至6条，优先顺序：政府网站 → 央广网/人民网/中新网 → 机构官网。",
    ]
    for i, s in enumerate(opts, 1):
        add_para(doc, f"{i}. {s}", size=10.5, space_after=3)

    add_heading_custom(doc, "已核验来源", level=2)
    verified = [
        "央广网〔1〕：双职务实名 + 2026年5月22日峰会圆桌与主办方代表发言。",
        "人民网〔2〕：秘书长实名 + 2025年峰会圆桌议题。",
        "中新网上海〔3〕：秘书长实名 + 2025年峰会圆桌，与人民网交叉印证。",
        "杨浦区政府官网〔4〕：执行会长实名 + 2026年3月会客厅揭牌交流。",
        "锦天城官网〔5〕：2024年6月主持活动 + 住房政策研究中心兼MBA协会秘书长职务。",
        "中新网上海〔6〕：秘书长实名 + 2026年峰会主办方代表发言，与央广网交叉印证。",
    ]
    for s in verified:
        add_para(doc, f"· {s}", size=10.5, space_after=2)

    add_heading_custom(doc, "首版暂不收录", level=2)
    excludes = [
        "宝龙商办总部战略发展部投资拓展总监；杨浦宝龙旭辉广场项目总经理等企业职务。",
        "万科、融创、金科、绿地等完整职业履历。",
        "高级工程师等专业职称（存在同名公示，需进一步证明为同一人）。",
        "福布斯相关奖项或荣誉、具体招商面积、签约金额、去化成果等经营数据。",
        "“专家”“著名”“资深”“领军人物”“推动”“赋能”“引领”“打造”等评价性或宣传性表述。",
        "出生日期、籍贯、家庭情况、手机号、微信、二维码、商业项目推介等。",
    ]
    for s in excludes:
        add_para(doc, f"· {s}", size=10.5, space_after=2)

    add_heading_custom(doc, "风险与策略", level=2)
    risks = [
        "同名消歧：存在多位“胡继刚”。概述第一句必须立即标明两项公开职务，义项名建议用“复旦大学住房政策研究中心秘书长”。",
        "利益相关：本人账号为自己创建词条可能触发利益相关审核；更稳妥可由同事或第三方提交。",
        "成功率：政府官网＋央媒＋主流媒体＋机构官网支撑的精简版通过可能性较高；一次写满职业履历通过率明显下降。",
        "后续补充：词条通过后，再分批增加有公开来源的职业经历、专业资质、代表性项目和图片。",
    ]
    for s in risks:
        add_para(doc, f"· {s}", size=10.5, space_after=2)

    add_heading_custom(doc, "纯文本粘贴备用（无格式）", level=2)
    add_para(doc, "【概述】", bold=True, size=11, space_before=4, space_after=2)
    add_para(
        doc,
        "胡继刚，复旦大学住房政策研究中心秘书长、上海市杨浦区科技企业联合会执行会长。曾参与不动产资产管理、新经济增长、人工智能商业化及中欧创新合作等主题的行业交流活动。〔1〕〔2〕〔3〕〔4〕〔5〕",
        size=10.5,
        space_after=6,
    )
    add_para(doc, "【人物经历】", bold=True, size=11, space_after=2)
    for text in experiences:
        add_para(doc, text, size=10.5, space_after=4)
    add_para(doc, "【社会职务】", bold=True, size=11, space_before=4, space_after=2)
    for text in duties:
        add_para(doc, text, size=10.5, space_after=3)

    add_para(
        doc,
        "（完）核验日期以本稿生成为准；提交前请再次打开各链接确认页面仍可访问、原文未被修订。",
        size=10.5,
        color=RGBColor(0x66, 0x66, 0x66),
        space_before=12,
        space_after=0,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"已生成：{OUT}")
    return OUT


if __name__ == "__main__":
    build()
