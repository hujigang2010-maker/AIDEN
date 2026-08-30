#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成一页 A4 确认函，可直接发给合作方。"""

import sys
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
import generate_wangdefeng_proposal as g  # noqa: E402
from generate_external_outputs import _mini_table, _p, set_cell_margins  # noqa: E402

OUT = Path(__file__).resolve().parents[1] / "output"
PATH = OUT / "10月31日王德峰老师活动_合作事项确认函.docx"


def build_letter():
    doc = g.new_doc()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.4)
    section.footer_distance = Cm(0.35)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("致合作方  ·  一页确认  ·  2026年8月19日")
    g.set_run_font(run, size=8, color=g.MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("本函为沟通确认文件，不构成合同。最终以双方签署的协议为准。请就下列四点一并书面回复。")
    g.set_run_font(fr, size=7.5, color=g.MUTED)

    g.add_para(
        doc,
        "10 月 31 日王德峰老师活动｜合作事项确认函",
        size=15,
        bold=True,
        color=g.NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0,
        space_after=3,
        line_spacing=1.0,
    )
    g.add_para(
        doc,
        "票价、分成、社群权限请一并确认，不单边调整其中一项。分成若下调，须对等让利；官方活动群由我方任群主。",
        size=9,
        bold=True,
        color=g.GOLD,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
        line_spacing=1.05,
    )
    g.add_para(
        doc,
        "感谢贵方持续沟通。现就 8 月 14 日讨论事项书面确认我方意见，请一并回复后，双方再进入场地、票务与开售准备。",
        size=9,
        space_before=0,
        space_after=4,
        line_spacing=1.1,
    )

    g.add_para(doc, "一、票价：标准票维持 499 元", size=11, bold=True, color=g.NAVY, space_before=2, space_after=3, line_spacing=1.0)
    g.add_para(
        doc,
        "2.5–3 小时公开讲座，对外主价格维持 499 元，不改为 599 元主力。599 元可作为优选 / 前区票；999 元保留为 VIP，须配套前排、问答、签名或合影。299、399 票档保留。",
        size=8.5,
        space_before=0,
        space_after=3,
        line_spacing=1.08,
    )
    _mini_table(
        doc,
        ["票档", "我方确认", "说明"],
        [
            ["限量引流票", "299 元，保留", "后区、限量，用于开售节奏"],
            ["早鸟普通票", "399 元，保留", "开售前 7–10 天限量"],
            ["标准票（对外主价格）", "499 元", "不改为 599 元主力"],
            ["优选 / 前区票", "599 或 699 元", "前区座位、赠书或资料"],
            ["VIP 票", "999 元，保留", "须配套前排、问答、签名或合影"],
        ],
        [4.4, 4.4, 9.7],
    )

    g.add_para(doc, "二、分成：要么两边都不调，要么两边一起调", size=11, bold=True, color=g.NAVY, space_before=7, space_after=3, line_spacing=1.0)
    g.add_para(
        doc,
        "我方客户此前为 60% : 40%。贵方希望下调至 55%，但贵方客户仍为 70% : 30%，属于单边调整，我方不能接受。请在方案 A、方案 B 中选择其一。",
        size=8.5,
        space_before=0,
        space_after=3,
        line_spacing=1.08,
    )
    _mini_table(
        doc,
        ["收入来源", "方案 A 我方", "方案 A 贵方", "方案 B 我方", "方案 B 贵方"],
        [
            ["我方独立报名 / 企业客户", "60%", "40%", "55%", "45%"],
            ["贵方独立报名", "30%", "70%", "40%", "60%"],
            ["双方共同渠道", "50%", "50%", "50%", "50%"],
            ["我方引入赞助", "70%", "30%", "70%", "30%"],
            ["贵方引入赞助", "30%", "70%", "30%", "70%"],
        ],
        [5.0, 3.15, 3.15, 3.15, 3.15],
        header_fill="3D5A40",
    )
    box = doc.add_table(rows=1, cols=1)
    cell = box.cell(0, 0)
    g.shade_cell(cell, g.ROSE_HEX)
    g.set_cell_border(
        cell,
        top={"val": "single", "sz": "6", "color": g.RED_HEX, "space": "0"},
        left={"val": "single", "sz": "16", "color": g.RED_HEX, "space": "0"},
        bottom={"val": "single", "sz": "6", "color": g.RED_HEX, "space": "0"},
        right={"val": "single", "sz": "6", "color": g.RED_HEX, "space": "0"},
    )
    cell.text = ""
    set_cell_margins(cell, top=40, bottom=36, left=80, right=80)
    _p(cell, "我方不能接受：我方客户 55% : 45%，同时贵方客户仍为 70% : 30%。分成微调的前提是标准票维持 499 元、官方群由我方任群主；否则按方案 A 执行。", size=8.5, after=0)
    g.set_table_full_width(box, [18.5])

    g.add_para(doc, "三、官方社群：我方任群主，贵方任管理员", size=11, bold=True, color=g.NAVY, space_before=7, space_after=3, line_spacing=1.0)
    _mini_table(
        doc,
        ["事项", "我方确认"],
        [
            ["群名称 / 建群", "「王德峰老师 10 月 31 日活动交流群」；我方企业微信创建"],
            ["权限", "我方任群主且不转让；贵方任管理员，负责内容答疑"],
            ["边界", "不得擅自解散、转群主、导出名单另建群，或在群内单独销售其他课程"],
            ["时间", "本函事项书面确认后再建正式群；咨询群同样由我方任群主"],
        ],
        [3.6, 14.9],
    )

    g.add_para(doc, "四、文字、头像与 logo", size=11, bold=True, color=g.NAVY, space_before=7, space_after=3, line_spacing=1.0)
    g.add_para(
        doc,
        "可按老师授权调整文案、头像与 logo，双方书面确认后对外。主视觉、售票页、海报和社群头像须保留联合主办双方露出，不改为其他课程附属品牌。",
        size=8.5,
        space_before=0,
        space_after=4,
        line_spacing=1.08,
    )

    g.add_para(doc, "请贵方一并确认", size=11, bold=True, color=g.NAVY, space_before=2, space_after=3, line_spacing=1.0)
    next_box = doc.add_table(rows=1, cols=4)
    items = [
        ("① 票价", "标准票 499 元；599 仅作优选票。"),
        ("② 分成", "选方案 A，或选方案 B（两边同步调）。"),
        ("③ 社群", "我方企微建群并任群主，贵方任管理员。"),
        ("④ 视觉", "双方书面确认后对外，保留联合主办。"),
    ]
    fills = [g.SOFT_HEX, g.CREAM_HEX, g.GREEN_HEX, g.AMBER_HEX]
    for i, ((title, body), fill) in enumerate(zip(items, fills)):
        c = next_box.cell(0, i)
        c.text = ""
        g.shade_cell(c, fill)
        set_cell_margins(c, top=46, bottom=46, left=60, right=60)
        _p(c, title, size=8.5, bold=True, color=g.NAVY, after=2)
        _p(c, body, size=7.5, after=0)
    g.set_table_full_width(next_box, [4.625, 4.625, 4.625, 4.625])

    g.add_para(
        doc,
        "我方（联合运营方）　　2026 年 8 月 19 日",
        size=9,
        bold=True,
        color=g.NAVY,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        space_before=8,
        space_after=0,
        line_spacing=1.0,
    )

    doc.save(PATH)
    return PATH


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    path = build_letter()
    print(f"已生成一页确认函：{path}")


if __name__ == "__main__":
    main()
