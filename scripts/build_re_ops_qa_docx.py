#!/usr/bin/env python3
"""生成《房地产开发运营一线答疑》Word 文档。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


CHINESE_FONT = "宋体"
HEADING_FONT = "黑体"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
ACCENT_DARK = RGBColor(0x0D, 0x2B, 0x4A)
GRAY = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ROW_ALT = "F3F6FA"
HEADER_BG = "1F4E79"
CALL_BG = "EEF3F8"
WARN_BG = "FFF4E5"

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "deliverables" / "房地产开发运营一线答疑.docx"


def _set_run_font(
    run,
    font_name: str,
    size: float,
    bold: bool = False,
    color: RGBColor | None = None,
    italic: bool = False,
) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_cell_borders(cell, color: str = "C5D0DC") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 9.5,
    color: RGBColor | None = None,
    font: str = CHINESE_FONT,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15
    run = p.add_run(text)
    _set_run_font(run, font, size, bold=bold, color=color)


def _add_paragraph(
    doc: Document,
    text: str = "",
    *,
    font: str = CHINESE_FONT,
    size: float = 11,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent: float | None = None,
    space_after: float = 8,
    space_before: float = 0,
    line_spacing: float = 1.45,
    color: RGBColor | None = None,
    italic: bool = False,
):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if text:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold, color=color, italic=italic)
    return p


def _add_rich(
    doc: Document,
    segments: list[tuple],
    *,
    size: float = 11,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent: float | None = 0.74,
    space_after: float = 8,
    line_spacing: float = 1.45,
):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    for seg in segments:
        text = seg[0]
        bold = seg[1] if len(seg) > 1 else False
        color = seg[2] if len(seg) > 2 else None
        font = seg[3] if len(seg) > 3 else CHINESE_FONT
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold, color=color)
    return p


def _add_heading(doc: Document, text: str, level: int) -> None:
    sizes = {1: 16, 2: 13.5, 3: 12}
    spaces_before = {1: 16, 2: 14, 3: 10}
    p = _add_paragraph(
        doc,
        text,
        font=HEADING_FONT,
        size=sizes[level],
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=spaces_before[level],
        space_after=8,
        line_spacing=1.3,
        color=ACCENT if level > 1 else ACCENT_DARK,
    )
    p.style = doc.styles[f"Heading {level}"]
    for run in p.runs:
        _set_run_font(run, HEADING_FONT, sizes[level], bold=True, color=ACCENT if level > 1 else ACCENT_DARK)


def _add_bullet(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Cm(0.75)
    pf.first_line_indent = Cm(-0.4)
    pf.space_after = Pt(4)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.4
    run = p.add_run("• ")
    _set_run_font(run, CHINESE_FONT, 11, bold=True, color=ACCENT)
    if bold_lead:
        run = p.add_run(bold_lead)
        _set_run_font(run, CHINESE_FONT, 11, bold=True)
        run = p.add_run(text)
        _set_run_font(run, CHINESE_FONT, 11)
    else:
        run = p.add_run(text)
        _set_run_font(run, CHINESE_FONT, 11)


def _callout(doc: Document, title: str, body: str, fill: str = CALL_BG) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    _set_cell_borders(cell, "1F4E79" if fill == CALL_BG else "C47B2B")
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    p1.paragraph_format.space_before = Pt(4)
    r = p1.add_run(title)
    _set_run_font(r, HEADING_FONT, 10.5, bold=True, color=ACCENT_DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.line_spacing = 1.35
    r2 = p2.add_run(body)
    _set_run_font(r2, CHINESE_FONT, 10.5)
    _add_paragraph(doc, "", space_after=8)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, HEADER_BG)
        _set_cell_borders(cell, "1F4E79")
        _set_cell_text(cell, h, bold=True, size=9.5, color=WHITE, font=HEADING_FONT, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                _set_cell_shading(cell, ROW_ALT)
            _set_cell_borders(cell)
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 or len(val) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
            _set_cell_text(cell, val, size=9, align=align)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    _add_paragraph(doc, "", space_after=6)


def _set_heading_styles(doc: Document) -> None:
    for i, size in ((1, 16), (2, 13.5), (3, 12)):
        style = doc.styles[f"Heading {i}"]
        style.font.name = HEADING_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT_DARK if i == 1 else ACCENT
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), HEADING_FONT)


def build_document(output_path: Path) -> None:
    doc = Document()
    _set_heading_styles(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.6)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    normal = doc.styles["Normal"]
    normal.font.name = CHINESE_FONT
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CHINESE_FONT)

    core = doc.core_properties
    core.title = "房地产开发运营一线答疑"
    core.subject = "项目公司、现金流、税务、表外合作、拿地销售与典型房企点评"
    core.category = "房地产投资运营"
    core.comments = "一线投资运营口径答疑，供内部讨论使用"

    # —— 封面 ——
    _add_paragraph(doc, "内部讨论稿  ·  一线投资运营口径", font=HEADING_FONT, size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=18, space_before=36)
    _add_paragraph(doc, "房地产开发运营一线答疑", font=HEADING_FONT, size=22, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT_DARK, space_after=8, line_spacing=1.2)
    _add_paragraph(doc, "项目公司 · 现金流 · 税务 · 表外合作 · 拿地销售 · 典型房企", font=HEADING_FONT, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=18)
    _add_paragraph(doc, "面向：投资、运营、财务、投研讨论", font=CHINESE_FONT, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_after=2)
    _add_paragraph(doc, "2026年8月", font=CHINESE_FONT, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_after=24)

    _callout(
        doc,
        "阅读说明",
        "本文按国内商品房开发的通行做法作答，口径以“上市公司集团总部—城市公司—项目公司”三级架构为默认场景。税率、预征率、资金监管以全国规则为骨架，具体城市细则会有差异。文中数字是教学级示意，不是某一家公司的真实账套。不构成投资建议、法律意见或税务鉴证。",
    )

    _add_heading(doc, "先看结论：八个问题的一线答案", 1)
    _add_paragraph(
        doc,
        "讨论地产运营，最容易把“法律形式”和“经济实质”混在一起。项目公司能隔离的是诉讼主体和清算单位，隔离不了担保、股东借款和保交楼；发债和开发贷在纸面上都不能拿地，真正拿地的是总部资金池通过增资或股东借款下去的钱；税盾看起来很多，但困在项目公司里，集团用不出去。下面八条是全文的压缩版。",
        first_line_indent=0.74,
    )

    _add_table(
        doc,
        ["问题", "一线结论（先记住这一句）"],
        [
            ["1. 项目公司", "为了合资、融资、清算和出表，不是为了让上市公司“出事就跑”。极端风险下法律上能挡一部分，经济上挡不住。"],
            ["2. 现金流", "拿地主要靠自有资金/股东借款。公募债、开发贷原则上都不能缴土地出让金；销售回款先进项目监管户，不是直接进总部。"],
            ["3. 税务", "土增、所得、增值都是“预缴+清算/结转”。退税法律上有、实务上难。亏损税盾主要是企业所得税，且很难跨项目用。"],
            ["4. 表外", "合营多是对等优质合作，联营多是小股、别人操盘或出表残留。无息往来本质是劣后资本，合作方爆雷时最先亏。"],
            ["5. 拿地", "10%净利率多按周边现价测算，压力测试常只做3%–5%。开盘价跌10%再叠加去化差，基本把利润打穿。"],
            ["6. 销售降价", "一年能卖得差不多，通常不必大幅降价。一线少见明面大降价，多见配赠、精装减配、渠道费。"],
            ["7. 绿城", "产品力强、利润释放弱，漏洞在投资纪律、成本超支、合资稀释和结算关闭，而不只在营销。"],
            ["8. 九家对比", "中海、华润管理最稳；建发、滨江销售最锐；绿城、金茂产品亮、利润钝；保利、越秀、招商是国企资源型。"],
        ],
        col_widths=[3.2, 12.8],
    )

    # ============================================================
    # 第一部分
    # ============================================================
    _add_heading(doc, "第一部分  运营", 1)

    _add_heading(doc, "一、项目公司：几乎每个项目都单独立项公司，原理是什么？", 2)

    _add_heading(doc, "1. 这样设计到底在解决什么问题", 3)
    _add_paragraph(
        doc,
        "地产开发不是“一个法人做全国生意”，而是“一块地、一个法人、一套账”。项目公司（俗称项目公司、项目公司SPV）是行业基础设施，不是财务花活。一线投委会、银行、合作方、税务局，默认的对话单位都是项目公司，不是上市公司本部。",
        first_line_indent=0.74,
    )
    _add_paragraph(doc, "这样做的真实原因，按重要性排序，通常是下面六条。", first_line_indent=0.74)

    _add_bullet(doc, "土地竞买和证照主体必须是独立法人。招拍挂、土地出让合同、四证、预售证，都挂在项目公司名下。一块地换一个主体，权属清晰，后续股权转让、合作、退出都好做。", bold_lead="拿地与证照主体。")
    _add_bullet(doc, "开发贷、按揭监管、土地抵押，银行认的是“这宗地+这个项目公司”。总部信用再强，贷款合同、抵押物、放款专户也落在项目公司。", bold_lead="融资抵押单位。")
    _add_bullet(doc, "地产很少100%独资。对等合作、小股操盘、代建+少数股权、地方政府平台入股，都必须有一个股权比例清楚的盒子。项目公司就是这个盒子。", bold_lead="合资与跟投的容器。")
    _add_bullet(doc, "土地增值税按项目（清算单位）计税，企业所得税也按独立纳税人申报。一个法人混多个项目，成本和分摊会打成一锅粥，清算和对账都会失控。", bold_lead="税务清算单位。")
    _add_bullet(doc, "并表、合营、联营、出表，都靠持股比例和治理安排切换。三道红线之后，大量项目被做成50%合营，本质是把负债和货值从并表口径里拿出去。", bold_lead="并表与出表开关。")
    _add_bullet(doc, "法律形式上，项目公司破产，股东以出资额为限承担责任。这是教科书理由，在中国地产里只是“部分有效”，下文会拆开讲。", bold_lead="有限责任。")

    _add_heading(doc, "2. 单个项目极端风险，上市公司躲得开吗？", 3)
    _add_paragraph(
        doc,
        "短答案：法律上能挡掉一部分合同债务和项目公司层面的诉讼，经济上、声誉上、资金上通常挡不住。投研如果把“项目公司有限责任”理解成“雷炸在项目里、上市平台没事”，会严重误判。",
        first_line_indent=0.74,
    )

    _add_table(
        doc,
        ["风险类型", "项目公司能否隔离", "一线实际情况"],
        [
            ["项目公司自身经营债务（无担保）", "原则上可以", "供应商、总包往往要求集团或城市公司担保、承诺付款，隔离被戳穿。"],
            ["开发贷、并购贷", "很难", "银行几乎必然要求控股股东担保、差额补足或流动性支持。"],
            ["已投入的股东借款/资本金", "不能", "钱已经下去了，项目烂尾就是真金白银损失，不是或有负债。"],
            ["预售监管与保交楼", "不能", "出险后集团必须垫资完工，这是政治任务和品牌底线，不是可选项。"],
            ["交叉违约与评级", "不能", "一个项目逾期，可能触发集团融资文件的交叉违约、评级下调。"],
            ["品牌与渠道", "不能", "业主维权、停工停贷，会迅速传导到同城其他盘的去化。"],
            ["刑事、行政责任", "基本不能", "违法用地、质量事故、预售资金挪用，监管会追到实控人和集团。"],
        ],
    )

    _add_paragraph(
        doc,
        "所以，项目公司真正隔离得比较干净的，是“没有担保、没有股东借款、已经把股权卖掉、也不再使用集团品牌”的少数项目。这类项目在头部房企里并不多。对上市公司而言，极端风险的损失顺序通常是：先亏掉已经投进去的资本金和往来款，再代偿担保，再垫资保交楼，最后才是并表减值。股票和债券投资人感受到的，是集团现金流和信用，不是项目公司的有限责任条款。",
        first_line_indent=0.74,
    )
    _callout(
        doc,
        "投资含义",
        "看一家房企“项目出险会不会打到集团”，不要先看股权比例，先看三张网：担保网、往来款网、品牌/保交楼责任网。三张网还在，有限责任就是纸面安慰。",
        fill=WARN_BG,
    )

    _add_heading(doc, "3. 单个项目亏损的税盾，能不能给其他项目用？", 3)
    _add_paragraph(
        doc,
        "短答案：原则上不能。中国企业所得税没有普遍的集团合并纳税。每个项目公司是独立纳税人，亏损留在自己账上，结转以后年度（一般企业最长五年），不能拿去冲减另一个项目公司的利润。",
        first_line_indent=0.74,
    )
    _add_paragraph(doc, "这是地产税务结构里最容易被低估的一点，也是“项目公司化”的代价。", first_line_indent=0.74)
    _add_bullet(doc, "A项目亏2亿、B项目赚2亿，集团合并报表可能打平，但税务上B照样按应纳税所得缴25%，A的亏损只能等A自己以后赚钱再抵。两个项目分属两个法人，税局不认“集团整体”。", bold_lead="不能横向对冲。")
    _add_bullet(doc, "总部可以亏损，用来抵总部自己的应纳税所得（管理费、品牌费、利息收入等），但总部往往不是主要利润中心，能吃掉的税盾有限。", bold_lead="总部亏损不等于集团税盾。")
    _add_bullet(doc, "吸收合并、把亏损公司并进盈利公司，理论上能把税盾带过去，但要满足特殊性税务处理、商业实质、反避税审查，地产项目公司还背着土地、贷款和监管账户，合并成本极高，实务中很少为了税盾去并。", bold_lead="合并转移很难。")
    _add_bullet(doc, "权益法下，上市公司报表可以确认对合营/联营的投资损失，那是会计亏损，不是税前可扣亏损。税盾仍困在被投资企业。", bold_lead="会计亏损≠税务亏损。")
    _add_paragraph(
        doc,
        "因此：单个项目亏损形成的税盾，主要是该项目公司自己的企业所得税亏损弥补，基本无法给其他项目用。这会系统性抬高“有盈有亏”组合的集团实际税负——报表利润平、现金税负高。2022年以来大量项目由盈转亏，很多集团出现“会计亏损很大、但仍在预缴所得税和土增税”的撕裂，根源就在这里。",
        first_line_indent=0.74,
    )

    # ---- Q2 ----
    _add_heading(doc, "二、现金流：发债、贷款、股东借款怎么流动，回款去哪？", 2)

    _add_heading(doc, "1. 发债的钱能不能拿地？", 3)
    _add_paragraph(
        doc,
        "短答案：公募公司债、企业债、中票，纸面上基本不能用于购置土地。募集说明书里通常会写“不直接或间接用于购置土地 / 土地储备”。海外债相对灵活，但仍受承诺函、评级机构和债权人约束。能拿地的，是发债到位后置换出来的自有资金，以及本来就不限用途的股东资金。",
        first_line_indent=0.74,
    )
    _add_paragraph(
        doc,
        "2016年沪深交易所对房地产公司债分类监管之后，“募集资金不得用于购置土地”成为标准承诺。此后无论是偿还到期债务、补充营运资金还是项目建设，审核口径都不接受“缴土地出让金”。银行间中票、发改企业债同样。2020年三道红线之后，能发公募债的房企更少，用途更集中在借新还旧。",
        first_line_indent=0.74,
    )
    _add_paragraph(
        doc,
        "一线怎么操作？合规路径是：债券进上市公司或财务公司专户，用于还债、补流动资金、付工程款；拿地另走自有资金。灰色路径是“置换”——发债还掉以前的股东借款或补总部资金池，总部再拿“自有资金”去缴地价。监管和中介机构近年对“间接用于拿地”查得更紧，但资金进了集团资金池之后，要完全证明每一分钱没有去拿地，本身就很难。投研需要区分：制度上不能、账上可能置换、出了事算违规使用募集资金。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "2. 银行贷款有没有可能用来拿地？", 3)
    _add_paragraph(
        doc,
        "短答案：开发贷不能。监管明确禁止金融机构发放专门用于缴纳土地出让金的贷款。开发贷的前提通常是四证齐全、项目资本金到位（住房开发项目资本金比例多年按20%或以上掌握，银行实务常要求30%左右先投入），放款进工程监管户，按形象进度支付建安。土地已经买完，才谈得上开发贷。",
        first_line_indent=0.74,
    )
    _add_paragraph(doc, "可能和“拿地”沾边的银行资金，只有这几类，而且都越来越窄：", first_line_indent=0.74)
    _add_bullet(doc, "并购贷可以用于收购项目公司股权，间接收购土地。这是目前相对干净的“用银行钱拿地”路径，但要符合真实并购、资本金和集中度要求，不能做成“假股权真融资”。", bold_lead="并购贷款。")
    _add_bullet(doc, "历史上土地储备贷款主要给地方政府土储机构，不是开发商。开发商前融（信托、资管、股东借款通道）才是过去拿地的主力，2017–2021年被大幅压缩。", bold_lead="前融/非标。")
    _add_bullet(doc, "经营性物业贷、供应链融资、个人经营贷流入拿地，属于违规挪用，不是正规渠道。", bold_lead="挪用。")
    _add_bullet(doc, "白名单、保交房专项借款，用途锁死在续建和交付，更不可能去拿地。", bold_lead="保交房资金。")
    _add_paragraph(
        doc,
        "所以，一个项目从拍地到开工，地价几乎必须是股东真金白银（含销售回款上收后再下拨）。这也是为什么土储多、开工慢的公司，会把大量现金锁在土地上，报表“现金充足”和“能拿来还债的现金”不是一回事。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "3. 不同融资方式，钱怎么从总部流到项目？", 3)
    _add_paragraph(
        doc,
        "集团司库的默认结构是“总部归集、项目申拨”。融资工具进账的第一落点不同，落到项目公司的形式却高度收敛：不是增资，就是股东借款（往来款）。很少有债券或银行贷款直接以总部为借款人、再现金赠与项目公司——那样既不合法，也算不清权益。",
        first_line_indent=0.74,
    )

    _add_table(
        doc,
        ["融资方式", "第一落点", "流向项目公司的常见形式", "项目端能否直接拿去缴地价"],
        [
            ["公募债/中票", "上市公司或发债主体专户", "还集团债务，腾出自有资金；或增资/股东借款（须符合募集用途）", "直接：否。置换后的自有资金：看合规。"],
            ["海外债", "境外SPV，再贷给境内", "股东借款、内保外贷资金下沉", "条款松一些，仍常限制购地。"],
            ["开发贷", "项目公司贷款专户", "不经过总部，银行直接放给项目", "否，只能付工程。"],
            ["并购贷", "收购主体（常是城市公司或项目公司）", "支付股权对价", "间接可以（买项目公司=买地）。"],
            ["股东借款", "总部资金池→项目公司往来", "最灵活的下沉通道", "可以，这是缴地价的主力。"],
            ["增资扩股", "计入项目公司实收资本", "资本金，银行认可的自有投入", "可以，且对开发贷资本金认定最干净。"],
            ["合作方投入", "合作方→项目公司", "按股比同步出资", "可以，常与股东借款搭配。"],
        ],
    )

    _add_paragraph(
        doc,
        "增资和股东借款怎么选？投委会和财务的本能是：能增资则增资，因为银行认资本金；能借款则借款，因为以后好抽回来，也不稀释合作方。现实是混合使用——先按股比实缴一部分资本金满足四证和贷款要求，缺口用股东借款补。股东借款大量是无息或低息、无抵押、无固定期限，本质是集团给项目的劣后资金。合作项目里，如果只有你在借钱给项目公司、对方不跟贷，你就在给对方加杠杆。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "4. 销售回款是进项目公司，还是进总部？", 3)
    _add_paragraph(
        doc,
        "短答案：先进入项目公司名下的预售资金监管账户，不是买家把钱打到上市公司。能上收到总部的，只是监管规则允许拨付之后的“可动用部分”，以及按揭放款进监管户之后满足节点的余额。2022年保交楼之后，这条通道显著收窄，是本轮房企流动性危机的核心机制之一。",
        first_line_indent=0.74,
    )
    _add_paragraph(doc, "一条完整的回款链，大致是：", first_line_indent=0.74)
    _add_bullet(doc, "客户认购，定金/首付打入项目公司监管账户（或售楼专户再转入监管户）。")
    _add_bullet(doc, "按揭银行放款，同样进入监管账户，不是集团一般户。")
    _add_bullet(doc, "住建/银行按工程进度、节点、留存比例审核拨付。付总包、材料、税费，优先在项目端完成。")
    _add_bullet(doc, "满足监管留存后的超额部分，经审批可转至项目公司一般户，再通过归还股东借款、分红、资金归集上收到总部或财务公司。")
    _add_bullet(doc, "现房销售、车位、商办等不受预售监管或监管较弱的回款，上收会快一些，这也是为什么房企在危机期拼命推现房和车位。")
    _add_paragraph(
        doc,
        "投研看“经营性现金流”时，最常见的误读是把签约金额当成可动用现金。一线财务看的是三个数：签约、回笼、可上收。三者在上行周期差距不大；在强监管周期，签约100、回笼80、可上收可能只有30–50。总部拿去还债、发工资、缴下一幅地的，是这最后一截。",
        first_line_indent=0.74,
    )
    _callout(
        doc,
        "资金流动一句话",
        "拿地：总部→增资/股东借款→项目公司→土地局。建安：银行开发贷+监管户回款→总包。还债：监管户可动用资金上收总部→债券/贷款到期户。方向不同，账户不同，不能混用。",
    )

    # ---- Q3 ----
    _add_heading(doc, "三、地产开发缴税的完整过程：预缴、清算、退税、税盾", 2)
    _add_paragraph(
        doc,
        "开发商的税，不是交房那一天一次性算清楚，而是“预售时先按规则预缴，完工清算/汇算时多退少补”。三个主税种——增值税、土地增值税、企业所得税——都按这个节奏走，但退税的现实概率差得很远。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "1. 时间轴：从拿地到清算，税是怎么叠上去的", 3)
    _add_table(
        doc,
        ["阶段", "增值税", "土地增值税", "企业所得税", "其他"],
        [
            ["拿地/开工", "土地价款作为进项/差额扣除的备查，本身不缴增值税", "无", "无销售收入则无预缴", "契税、印花税、城镇土地使用税"],
            ["预售回款", "一般计税项目按预收款的3%预缴", "按当地预征率预征（见下）", "按预计计税毛利率计算预计毛利，计入当期所得预缴", "附加税随增值税预缴"],
            ["完工交房", "纳税义务发生，按9%一般计税（或老项目5%简易）正式申报，预缴可抵", "继续预征，直至清算受理", "产品完工后按实际计税成本调整，多退少补在汇算中体现", "房产税（自持）"],
            ["土增清算", "无直接联动，但清算收入用不含税口径", "正式清算，预征抵减应缴，多退少补", "清算补缴的土增税可在所得税前扣除，常造成所得税波动", "—"],
            ["尾盘/注销", "留抵税额结转或（符合条件时）申请退还", "尾盘按清算结论或核定征收", "亏损结转；注销时未弥补亏损通常作废", "土地使用税至交房/过户"],
        ],
    )

    _add_heading(doc, "2. 增值税：预缴3%，交房时再算正式税", 3)
    _add_paragraph(
        doc,
        "2016年营改增后，新项目一般计税，税率几经调整后为9%。土地成本可以从销售额中差额扣除（向政府支付的土地价款），这是开发商增值税税负能压下来的关键。简易计税5%主要留给老项目。",
        first_line_indent=0.74,
    )
    _add_paragraph(
        doc,
        "预缴：收到预收款时，一般计税项目按预收款的3%预缴增值税（不是按9%）。交房（合同约定交房日或实际交房孰早）才发生纳税义务，再按差额计税方法算应纳税额，已经预缴的部分抵减。进项税来自建安、材料、顾问费等，进项多的项目，正式申报时可能出现应纳税额小于已预缴，形成留抵。",
        first_line_indent=0.74,
    )
    _add_paragraph(
        doc,
        "退税可能大吗？增值税对开发商很少是“把预缴的现金退回来”。更常见的是留抵结转，以后的销项再吃。增量留抵退税政策对房地产行业限制较多，实操中不要按“能退现金”去模型。预缴大于应纳税的部分，优先抵以后属期，注销时才可能谈到退还，而项目公司注销又极慢。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "3. 土地增值税：预征多少，最后如何清算、能否退税", 3)
    _add_paragraph(
        doc,
        "土增税是开发商最“重”、也最有博弈空间的税。预征是为了把税先收上来，清算才是算总账。",
        first_line_indent=0.74,
    )
    _add_rich(
        doc,
        [
            ("预征计税依据（2025年起更明确的口径）：", True),
            ("预收款÷（1+增值税税率或征收率）。例如一般计税收到1.09亿预收款，计征依据=1亿元。预征税额=计征依据×当地预征率。", False),
        ],
    )
    _add_paragraph(
        doc,
        "预征率由各省在国家下限之上自定，且按业态区分。2024年11月税务总局将预征率下限下调0.5个百分点：东部1.5%、中部和东北1%、西部0.5%（保障性住房除外）。下调前东部下限长期是2%。一线实操里，普通住宅常见1.5%–2%，非普通住宅/改善盘2%–3%，商办和车位3%–5%，保障房多为0。同一城市、同一项目，住宅和商办预征率可以不同。",
        first_line_indent=0.74,
    )
    _add_paragraph(
        doc,
        "清算触发（国税发〔2006〕187号等）：项目全部竣工并销售完毕；转让在建；直接转让土地使用权；已竣工且转让面积达可售面积85%以上（或剩余已出租/自用）；拿预售证满三年仍未卖完；注销税务登记等。税务机关也可指定清算。受理清算后停止预征，之后的尾盘按清算结论申报。",
        first_line_indent=0.74,
    )
    _add_paragraph(doc, "清算公式的骨架：", first_line_indent=0.74)
    _add_bullet(doc, "应税收入 − 扣除项目 = 增值额。扣除项目包括：土地成本、开发成本、开发费用、与转让有关的税金（不含增值税）、财政部规定的加计20%（从事开发的纳税人）。")
    _add_bullet(doc, "增值率 = 增值额 / 扣除项目。普通标准住宅增值率未超过20%的，可免税（这是住宅项目最重要的“政策税盾”）。")
    _add_bullet(doc, "超20%后实行四级超率累进：30%、40%、50%、60%，并有速算扣除。高溢价项目清算税负可以非常惊人；低溢价或亏损项目清算应缴可能低于已预征。")
    _add_paragraph(
        doc,
        "多退少补：清算应缴小于已预征，理论上应退税；大于已预征，则补税。补税很常见，退税才是讨论的焦点。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "4. 企业所得税：预计毛利率预缴，完工再调整", 3)
    _add_paragraph(
        doc,
        "核心文件是国税发〔2009〕31号。未完工产品的预售收入，不能等交房才进所得，而要按预计计税毛利率算出预计毛利，并入当期应纳税所得额预缴。完工后结算计税成本，再按实际毛利调整。",
        first_line_indent=0.74,
    )
    _add_paragraph(
        doc,
        "预计计税毛利率下限：省会/计划单列市城区及郊区不得低于15%；地级市不得低于10%；其他地区不得低于5%；经济适用房、限价房、危改房不得低于3%。各地可上浮，热点城市实际执行经常在15%–20%甚至更高。注意：这是税务口径的毛利率，不是公司投委会的销售净利率，两者不能混用。",
        first_line_indent=0.74,
    )
    _add_paragraph(
        doc,
        "示意：某省会住宅项目预售不含税收入10亿元，当地预计毛利率15%，则当年要确认预计毛利1.5亿元（还要加减期间费用、纳税调整）。即便这个项目投委会测算净利率只有8%，甚至已经预感到要亏，预缴阶段仍可能按15%毛利缴所得税。这就是下行周期里“项目要亏、税还在缴”的来源。",
        first_line_indent=0.74,
    )
    _add_paragraph(
        doc,
        "完工调整后，若实际毛利低于已按预计毛利缴税的水平，会调减所得；若该年度汇算出现亏损，可结转以后年度弥补。项目公司如果再也没有新盘，这笔亏损往往只能等到注销时消失，变不成现金。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "5. 现实中退税可能性大吗？亏损税盾主要是所得税吗？", 3)
    _add_table(
        doc,
        ["税种", "法律上能否退/抵", "一线现实", "对现金流的意义"],
        [
            ["增值税", "预缴抵正式税；留抵结转；符合条件的留抵退税", "开发商拿到大额现金退税的概率低，以留抵为主", "不要把留抵当可动用现金"],
            ["土地增值税", "清算多退少补", "补税容易、退税难。审核1–3年常见，成本被核减后“该退的变成该补”。近年部分城市为纾困加快清算，但仍不顺畅", "模型里退税概率应打折，甚至按0"],
            ["企业所得税", "汇算多缴可抵下期或申请退还；亏损结转", "有后续项目的公司可慢慢吃回来；单项目公司多缴后申请退还，税局很慢、很严", "税盾困在法人实体内"],
            ["土地增值税（普宅20%免税）", "满足条件不征", "这是“免掉”不是“退回”。预征过的部分，要靠清算退出来，仍然难", "比现金退税更真实的是少缴"],
        ],
    )
    _add_paragraph(
        doc,
        "为什么土增税退税难？不是条例写了“多退少补”就算数。清算审核会抠发票备注栏、成本分摊、利息扣除、车位配套、精装是否并入、关联方建安是否偏高。审核人员有“少退多补”的激励，地方财政在下行期更不愿意掏出现金。结果是：增值率下行后，企业觉得自己超缴了，税局觉得你成本不实，双方在清算报告上拉锯几年。能谈成的，经常是“不退现金、抵以后项目”——但以后项目在另一个项目公司，还是那句：法人独立，抵不了。",
        first_line_indent=0.74,
    )
    _add_paragraph(
        doc,
        "亏损项目的税盾，主要就是企业所得税的亏损弥补，外加土增税可能免征或低税。增值税是流转税，项目亏钱不会自动变成增值税负的“盾”。契税、土使用税更没有亏损抵减。务必再强调一次：这枚所得税税盾默认困在项目公司，集团报表上的亏损不能直接拿去抵别的项目的应纳税所得。",
        first_line_indent=0.74,
    )
    _callout(
        doc,
        "税务建模的老实做法",
        "预售期：按当地预征率计土增、按3%计增值税预缴、按预计毛利率计所得税预缴，三项都当现金流出。清算期：土增退税概率给0–30%的折，不要按100%退回；所得税多缴优先当成以后年度抵税，不要当年回流。跨项目税盾，默认按零。",
        fill=WARN_BG,
    )

    # ---- Q4 ----
    _add_heading(doc, "四、表外项目：合营、联营怎么评估，为什么一个亏一个赚？", 2)

    _add_heading(doc, "1. 合营和联营分别是什么，风险怎么看", 3)
    _add_paragraph(
        doc,
        "会计上：合营是共同控制，重大事项需要一致同意；联营是重大影响，通常持股20%–50%但说了不算。两者在国际财务报告准则和中国企业会计准则下，对上市公司都按权益法核算，不逐行并表。所以它们是“表外”的负债和货值，只以“对合营企业/联营企业的投资”一个余额出现在资产端，利润端只看“投资收益”。",
        first_line_indent=0.74,
    )
    _add_paragraph(
        doc,
        "对地产投资人，这是三道红线时代最重要的报表游戏之一：把持股做到50%或以下、把治理写成共同控制，并表负债下去了，合作方的贷款也不进你的有息债务。货值说明会上仍可能把这些项目的销售算进“全口径”，利润却只按股比进来。评估风险，必须把表外加回来。",
        first_line_indent=0.74,
    )
    _add_paragraph(doc, "一线评估表外项目，建议至少看这七项：", first_line_indent=0.74)
    _add_bullet(doc, "股比、董事会席位、一票否决事项、谁有操盘权和品牌。只有分红权、没有操盘权的，风险特征完全不同。", bold_lead="控制权。")
    _add_bullet(doc, "土地成本、剩余货值、去化、售价与周边二手倒挂幅度。表外项目的质量往往比并表项目更两极。", bold_lead="项目质地。")
    _add_bullet(doc, "项目公司自己的开发贷、供应链欠款、预售监管缺口。这些债务你不并表，不代表你不用救。", bold_lead="项目公司杠杆。")
    _add_bullet(doc, "你对项目公司、合作方的担保、差额补足、流动性支持。这是隐藏的或有负债。", bold_lead="增信。")
    _add_bullet(doc, "其他应收款里对合营/联营的借款，账龄、利息、抵押、是否已减值。这是已经出去的真金白银。", bold_lead="往来款。")
    _add_bullet(doc, "合作方是谁：央企互保、地方国企、民企、城投，爆雷概率和救助意愿差一个数量级。", bold_lead="对手方。")
    _add_bullet(doc, "承诺回购、对赌、保收益、劣后。名股实债的，应按债务评估，不应按股权。", bold_lead="协议条款。")

    _add_heading(doc, "2. 为什么很多联营亏钱、合营赚钱？", 3)
    _add_paragraph(
        doc,
        "这不是会计魔法，是“项目筛选 + 控制权 + 出表动机”叠在一起的结果。权益法本身不会让合营变赚钱、联营变亏钱。看到这种分化，优先用下面几条解释。",
        first_line_indent=0.74,
    )
    _add_bullet(doc, "合营多为50/50或接近对等，合作对象经常是另一家央企、地方龙头或资源方。双方都要在自己的报表里体现投资收益，挑地会更狠，治理会更对称，质量和利润天然更好看。", bold_lead="选择偏差。")
    _add_bullet(doc, "联营往往是小股（20%–40%）、别人操盘，或你把控股权卖掉之后剩下的残值。别人的盘，你进场时可能已经是高价地、尾盘、商办、车位。品牌输出、小股操盘的管理费进了你的轻资产业务，项目利润却可能很薄，权益法还要按股比认亏。", bold_lead="小股与残值。")
    _add_bullet(doc, "为了降并表负债，好项目也可能被做成合营；为了把问题项目“送出表”，持股会降到不并表的区间，于是问题项目集中出现在联营。出表不等于出险，只是换了一个会计抽屉。", bold_lead="出表动机。")
    _add_bullet(doc, "合营亏损时，两边都有动力补充资金、共管销售；联营亏损时，操盘方可能先保自己的现金流，少数股东只能跟着减值。控制权决定了亏钱之后谁有权自救。", bold_lead="亏损后的处置权。")
    _add_bullet(doc, "部分联营是持有型商业、城市运营、产业配套，折旧和利息会把账面利润压成负数，但合营住宅开发在交房年可以一次释放利润。业态不同，不要直接比“赚/亏”。", bold_lead="业态差。")
    _add_paragraph(
        doc,
        "反例也很多：有的合营是“各怀鬼胎”的高价地，对等持股一样能亏穿；有的联营是你小股跟进的核心盘，照样贡献厚利润。所以正确读法不是“联营=坏、合营=好”，而是：联营组合更可能混进你说了不算的项目和出表残留，需要按项目清单拆开，不能看一个投资收益总数。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "3. 互相借钱、无息无抵押，有没有风险？", 3)
    _add_paragraph(
        doc,
        "有，而且是表外风险里最实在的一块。无息、无抵押、无固定期限的往来，经济实质是股东给项目或给合作方的劣后资本。上行周期它是润滑剂，下行周期它是炸药。",
        first_line_indent=0.74,
    )
    _add_bullet(doc, "项目公司破产或被查封时，股东往来排在开发贷、按揭监管、工程款之后，回收率经常接近零。无抵押=没有优先权。", bold_lead="清偿顺序垫底。")
    _add_bullet(doc, "你单边放款、合作方不按股比跟贷，等于你在给对方的股权加杠杆。对方只出小钱，分享利润；你出大钱，承担流动性。这是最典型的隐性利益输送。", bold_lead="不对等出资。")
    _add_bullet(doc, "往来可以让一个“不并表”的项目在经济上仍被你输血。审计要判断是否构成实质控制；投资人则应把它加回净负债。只看三道红线档位会低估杠杆。", bold_lead="实质并表。")
    _add_bullet(doc, "无息在税务上可能被核定利息、涉及转移定价；在公司治理上，则是抽贷、停贷的扳机——总部一紧，项目立刻停工。", bold_lead="可随时抽贷。")
    _add_bullet(doc, "借给合作方（而不是项目公司）的款项，对手方风险更纯粹。民企合作方爆雷后，这类应收在2022–2024年出现过集中计提。", bold_lead="借给合作方。")
    _add_paragraph(
        doc,
        "评估方法很具体：把“对合营/联营企业的其他应收款”按对手方、有无担保、账龄、是否逾期拆开，再把担保余额加回去，得到“表外净暴露”。超过集团账面现金一定比例（例如30%–50%）的，这项就应该进入信用评级的核心假设，而不是附注里的一行。",
        first_line_indent=0.74,
    )

    # ============================================================
    # 第二部分
    # ============================================================
    _add_heading(doc, "第二部分  拿地与销售", 1)

    _add_heading(doc, "五、拿地测算：10%净利率是按现价算的吗？跌价后会不会亏？", 2)

    _add_heading(doc, "1. 投委会那张测算表，价格假设从哪来", 3)
    _add_paragraph(
        doc,
        "行业里说的“净利率10%才拿地”，指的是项目全周期静态或准动态测算的税后利润 / 含税或不含税销售收入（各家分母略有不同，比较时要统一）。它不是年化ROE，也不是权益内部收益率。很多公司会同时看：销售净利率、权益IRR、峰值资金、回正周期。2023–2026年，央国企投委会把销售净利率底线放在8%–12%很常见，10%是一个好记的口头禅。",
        first_line_indent=0.74,
    )
    _add_paragraph(doc, "价格假设，一线默认是“现在”，不是“三年后还要涨”：", first_line_indent=0.74)
    _add_bullet(doc, "同盘段竞品近3–6个月成交价、一房一价、折扣后实际成交，不是挂牌价。")
    _add_bullet(doc, "周边二手成交（注意二手含装修、家具，和新建精装的口径差）。")
    _add_bullet(doc, "限价/备案价天花板。一线和强二线很多地块，测算价格不是你想卖多少，而是政府让你卖多少。")
    _add_bullet(doc, "未来两年的推盘结构：开盘引流价、主售价、尾盘、车位、商铺，分别给价，再加权。")
    _add_paragraph(
        doc,
        "会不会考虑3%–5%的下跌？正规投委会有敏感度：售价±5%、建安±5%、去化周期+6/12个月。有的公司把“现价下浮3%”作为基案，把现价作为乐观案。但10%的净利率底线，往往是基案过线就过会，并不要求“跌5%之后仍有10%”。也就是说，10%里通常只含了薄薄一层价格安全垫，不是为了扛10%的房价下跌而准备的。",
        first_line_indent=0.74,
    )
    _add_paragraph(
        doc,
        "上行周期还有一个坏习惯：用“开盘后涨价”把均价做高，或用车位、储藏室、精装溢价把综合单价做漂亮。下行周期这些加成先没，净利率比表上掉得更快。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "2. 开盘价跌10%、去化也不好，项目是不是必亏？", 3)
    _add_paragraph(
        doc,
        "对一个“基案净利率约10%”的项目，售价下跌10%，在成本刚性和周期拉长的双重作用下，大概率从微利变成亏损，至少变成“会计微利、权益回报不可接受”。不是绝对的数学必亏，但作为投资判断可以按“利润被打穿”处理。",
        first_line_indent=0.74,
    )
    _add_paragraph(doc, "用一个教学级结构说明弹性（数字是结构，不是某盘）：", first_line_indent=0.74)

    _add_table(
        doc,
        ["项目", "基案（净利率10%）", "售价−10%、去化仍快", "售价−10%、去化慢一年"],
        [
            ["销售收入", "100", "90", "90"],
            ["土地+契税等", "35", "35（沉没）", "35"],
            ["建安/精装/配套", "28", "28", "28"],
            ["费用（销售、管理）", "8", "7.5（略降）", "9（渠道加码）"],
            ["财务费用（资本化/费用化）", "4", "4", "7"],
            ["增值税附加+土增+所得税（示意）", "15", "13", "12（税基下降，但预缴已出去）"],
            ["税后利润（示意）", "10", "2.5", "−1"],
            ["净利率", "10%", "约3%", "转负"],
        ],
    )
    _add_paragraph(
        doc,
        "关键机制有三条。第一，土地是沉没成本，建安下调的空间很小（最多减配、改毛坯、砍园林），所以售价下跌几乎按“1:1”打在利润上。第二，去化变差会增加资金占用，财务费用是第二把刀，往往比售价本身更伤权益IRR。第三，税不能按利润同比下降来理解：预缴已经按高价格发生，退税又慢，现金税负下降滞后于利润下降。",
        first_line_indent=0.74,
    )
    _add_paragraph(
        doc,
        "还有一个更狠的情况：限价盘“价格不能涨、成本不能降”，利润薄得像纸；一旦周边二手跌破限价，去化也会坏，变成“限价没有保护利润，只限制了涨价”。2021年后不少二线改善盘就是这种结构。",
        first_line_indent=0.74,
    )
    _callout(
        doc,
        "经验规则",
        "销售净利率约10%的项目，售价每降3个点，净利率大约降2.5–3.5个点（视税和费用）。跌10%而不延长周期，利润大概还剩0–4个点；再叠加去化差、渠道费、降价损失，就可以按亏损来预案。投委会如果只做+0/−5%价格带，对10%下跌是没有准备的。",
    )

    _add_heading(doc, "六、销售速度与降价：大概什么关系，管理层看哪些指标？", 2)

    _add_heading(doc, "1. 一年能卖得差不多，还会不会大降价？", 3)
    _add_paragraph(
        doc,
        "通常不会。定价的第一目标不是“卖最贵”，而是“在计划周期内把货值换成现金”。如果开盘去化达到60%–70%，年内能把可售货值消化到70%以上，说明价格已经位于市场接受带，再大幅降价是在破坏剩余货值，也会惹恼老业主。这时更常见的是小幅一房一价调整、渠道点位微调、赠送车位或装修升级，而不是公开八折。",
        first_line_indent=0.74,
    )
    _add_paragraph(
        doc,
        "反过来，需要大幅降价的项目，几乎都先出现“去化失败”：开盘去化低于40%，三个月仍明显低于供货计划，库存年龄往12个月以上走，而集团又恰好要现金。降价是结果，不是风格。",
        first_line_indent=0.74,
    )
    _add_paragraph(doc, "一个粗的对应关系（因城市而异，只用于建立直觉）：", first_line_indent=0.74)

    _add_table(
        doc,
        ["去化状态", "常见价格动作", "管理层心理"],
        [
            ["开盘去化≥70%，年内可售去化≥80%", "坚持价单，顶多一房一价微调", "不要为最后10%的货去砸盘"],
            ["开盘50%上下，月度去化尚可", "节点促销、渠道加点、赠送，明面折扣有限", "用时间换价格，但开始盯现金流"],
            ["开盘<40%，连续两季不达计划", "实质性降价5%–10%，或改产品、改精装口径", "保回笼、保交付、保不爆监管户"],
            ["库存>18个月，或资金链已紧", "10%以上的实际成交价下调，工抵、团购、渠道高点位", "价格服从生存"],
        ],
    )

    _add_heading(doc, "2. 管理层决定降价时，真正看哪些指标？", 3)
    _add_paragraph(doc, "不是看“利润率还够不够10%”——那是拿地时的语言。到了销售期，决策函数几乎变成现金和风险。", first_line_indent=0.74)
    _add_bullet(doc, "开盘去化、近四周认购、来访转认购、渠道与自然到访比。这是需求温度计。", bold_lead="去化。")
    _add_bullet(doc, "监管账户缺口、工程付款节点、三个月内到期的供应链和贷款。这是必须卖的理由。", bold_lead="项目现金流。")
    _add_bullet(doc, "集团层面的到期债券、三道红线档位、信用评级展望。项目为集团输血的压力，会迫使好盘也降价。", bold_lead="集团流动性。")
    _add_bullet(doc, "竞品近期成交价、二手倒挂、同城兄弟盘是否先降。降价有传染性。", bold_lead="市场参照。")
    _add_bullet(doc, "备案价、限价、一房一价系统能不能改下来。一线经常是“想降降不了”或“明着不能降、暗着配赠”。", bold_lead="政策空间。")
    _add_bullet(doc, "交房时间、质量承诺、已售业主维权风险。大降价会引发老业主索赔，法律和舆情成本要进决策。", bold_lead="已售业主。")
    _add_bullet(doc, "剩余货值结构：刚需小户型还是大平层、车位去化、商铺。结构差的盘，降价弹性也差。", bold_lead="货值结构。")
    _add_bullet(doc, "合作方是否同意。合营盘降价要过董事会，常常是“一方要现金、一方要利润”的拉锯。", bold_lead="治理。")
    _add_paragraph(
        doc,
        "所以会看到一种刺眼的现象：利润率还过得去的盘，因为集团要还债，照样降；已经亏的盘，因为限价和业主维权，明面降不下去，只能停工或改分期。价格不是利润管理工具，是流动性管理工具。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "3. 大型房企在一线城市，是不是很少降价？", 3)
    _add_paragraph(
        doc,
        "明面的、全盘公开的大降价确实不多，但“不降价”不等于“价格坚挺”。更准确的说法是：一线核心区改善盘，名义备案价稳定；实际成交价通过结构、配赠和渠道在动。郊区、大户型、商办、以及2021–2022高价地，实际成交价下调并不罕见。",
        first_line_indent=0.74,
    )
    _add_bullet(doc, "一线限价、一房一价、公证摇号的制度，限制了公开打折的空间。国企为主力，也不愿做全城价格锚的破坏者。")
    _add_bullet(doc, "核心区土地贵、改善客对总价敏感但对“打折盘”也敏感，品牌房企更怕老业主维权。")
    _add_bullet(doc, "实际让利：精装减配或改毛坯、送车位/家具、渠道费从1%加到3%–5%、工抵房、团购、付款方式优惠。成交净价可以下来一截，而“项目降价”新闻可以没有。")
    _add_bullet(doc, "2023–2025年，北京、上海、深圳、广州都出现过郊区盘、大面积段的实质性降价或变相降价。不要用核心区几个热盘代替一线全体。")
    _add_paragraph(
        doc,
        "投研如果只用“有没有降价通告”判断一线项目质量，会漏掉渠道点和配赠里的暗降。看实际成交价、与备案价的差、以及二手倒挂，比看新闻可靠。",
        first_line_indent=0.74,
    )

    # ============================================================
    # 第三部分
    # ============================================================
    _add_heading(doc, "第三部分  公司点评", 1)

    _add_heading(doc, "七、绿城：利润为什么释放不出来，管理漏洞在哪？", 2)
    _add_paragraph(
        doc,
        "绿城的公开形象是产品力顶级，园林、精装、园区生活是行业教材。但资本市场长期的抱怨也稳定：结算慢、利润薄、少数股东损益厚、股权结构复杂，账面利润和销售规模长期对不上。这不是单一环节的事故，是一套互相加强的结构。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "1. 利润释放不出来，先分清是“还没到”还是“根本没有”", 3)
    _add_bullet(doc, "精装改善盘从预售到竣工结算，比毛坯刚需长一年甚至更久。利润在交房结算时才进表，销售火不等于当年有利润。", bold_lead="结算周期长。")
    _add_bullet(doc, "大量项目是合资、代建+少数股权、或并表但少数股东占比高。规模在全口径，利润被合作方切走。绿城管理（代建）的利润表和绿城中国（开发）不能混着夸。", bold_lead="少数股东与代建。")
    _add_bullet(doc, "历史上不低的土地成本和始终偏高的产品成本，把毛利空间先吃掉。再叠融资成本和结算滞后，净利率很难看。", bold_lead="高成本结构。")
    _add_bullet(doc, "成本关闭慢：竣工后签证、结算、索赔拉很久，项目公司迟迟不能把成本钉死，财务只能高估成本、低估利润，或反过来在某一年集中释放/计提，造成利润“脉冲”。", bold_lead="成本关闭。")
    _add_paragraph(
        doc,
        "所以“释放不出来”有两种：一种是利润在项目公司里，还没走到结算；一种是测算时以为有10个点，做完只剩2个点甚至没有。绿城两种都有，市场更担心第二种。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "2. 管理上真正难管的跑冒滴漏", 3)
    _add_paragraph(
        doc,
        "绿城的漏洞，不太像某些高周转民企那种“营销飞单、恶打价格战”，更像高端改善开发商的典型病：前端投资软、中端产品变更多、后端结算弱，再加上合资结构把管理半径拉到失控边缘。",
        first_line_indent=0.74,
    )
    _add_bullet(doc, "投委会纪律。产品口碑会转化成“这个城市应该有绿城”的拿地冲动。地价略贵、配建略多、自持略重，基案净利率靠乐观售价和精装溢价撑着。这一步错了，后面运营再勤奋也补不回土地账。", bold_lead="投资。")
    _add_bullet(doc, "方案阶段的园林、大堂、精装、会所标准，到了现场变成“不能比上一个盘差”。设计变更和现场签证是成本黑洞。高端盘的变更，单方可以多出几百到上千元，几个点的净利率就此消失。", bold_lead="产品变更。")
    _add_bullet(doc, "总包、精装、园林、材料，若区域自主权大、集采执行弱，折扣和指定分包会变成利益链。越是“品质”作为不可挑战的话语，采购越难被财务挑战。", bold_lead="采购与合约。")
    _add_bullet(doc, "合资项目里，谁的队伍、用谁的集采、营销费谁批、降价谁点头，全是扯皮。绿城品牌输出时，对方可能用绿城的标准花绿城不认的成本；绿城操盘时，对方可能卡住增资和降价。跑冒滴漏发生在缝隙里。", bold_lead="合资界面。")
    _add_bullet(doc, "高端盘渠道点位、业主关系和示范区运营本身就贵。去化一慢，点位上浮，费用率失控。", bold_lead="营销费用。")
    _add_bullet(doc, "跟投、关联方施工、过往股东体系内的资源方，都可能让定价不是纯市场价。不一定违法，但会让集团利润变薄。", bold_lead="关联交易。")
    _add_bullet(doc, "项目公司长期不关闭，未决签证、质保金、诉讼、尾盘车位挂在账上，总部无法把现金和利润收干净。", bold_lead="项目关闭。")
    _add_paragraph(
        doc,
        "代建业务其实是绿城相对健康的一块：轻资产、品牌溢价、费收相对清晰。把代建的增长理解成开发利润的替代，会高估集团可分配利润；把代建的管理能力理解成开发投资纪律，也会误判。两套能力不是自动迁移的。",
        first_line_indent=0.74,
    )
    _callout(
        doc,
        "看绿城，抓三个对账单",
        "全口径销售 vs 权益销售 vs 并表结算；投委会净利率 vs 竣工结算净利率；集采合约价 vs 竣工结算价。三组差收窄了，利润才会真正释放。只听产品口碑，会持续高估。",
        fill=WARN_BG,
    )

    _add_heading(doc, "八、九家公司的管理能力与销售能力：一个可用的对照", 2)
    _add_paragraph(
        doc,
        "以下是一线投资、合作、同行交流里比较稳定的印象，不是财务预测，也不是道德评价。管理能力侧重投资纪律、成本、治理、利润转化；销售能力侧重品牌、定价、渠道、去化。打分是相对排序，用于讨论，不是量化模型。",
        first_line_indent=0.74,
    )

    _add_table(
        doc,
        ["公司", "管理（投资/成本/治理）", "销售（品牌/去化）", "一句话"],
        [
            ["华润置地", "极强", "极强", "开发+商业双轮，投资克制，产品稳定，利润质量高。"],
            ["中海地产", "极强", "强", "成本与利润的行业标杆，拿地极挑，销售不靠叫卖。"],
            ["建发", "强（偏周转）", "极强", "渠道和节奏一流，规模起来快，利润厚度要打折看。"],
            ["金茂", "中上", "强（高端）", "金茂府产品力突出，城市运营和个别高价地拖累利润。"],
            ["绿城", "中（产品强、投资弱）", "强（改善）", "会做房子，不善把房子做成利润。"],
            ["滨江", "强（区域）", "极强（浙江）", "杭州体系销冠级操盘，全国化能力明显弱于区域能力。"],
            ["越秀", "强（稳健）", "中上", "广州国企节奏，土储有地铁合作亮点，爆发力一般。"],
            ["保利发展", "强", "极强", "品牌信任和全国渠道是资产，组织偏大，利润中等偏好。"],
            ["招商蛇口", "中上", "强", "资源禀赋好（蛇口/城市更新），开发利润释放不够利落。"],
        ],
    )

    _add_heading(doc, "1. 华润置地", 3)
    _add_paragraph(
        doc,
        "管理上最接近“可以当教材”的综合开发商之一。投资口径偏核心城市和核心板块，很少为了规模去赌边缘地。成本不是中海那种极致抠，但产品标准稳定，返工和变更相对少。商业（万象系）反哺品牌和现金流，使它在下行周期仍有操盘从容度。销售上，“润”系列识别度高，一线和强二线去化不靠极端折扣。弱点是规模诉求与利润诉求的平衡，以及商业重资产对资本的占用。把华润当“管理能力上限”的参照，误差不大。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "2. 中海地产", 3)
    _add_paragraph(
        doc,
        "管理能力的第一关键词是成本与克制。拿地可以连续几个季度几乎不进场，宁缺毋滥；建造端的目标成本、合约规划、变更管理，是同行去挖人、去抄流程的对象。利润率长期好于多数央企，不是因为卖得更贵，而是因为买地更便宜、浪费更少。销售能力强，但品牌调性偏“价值”而非“情绪”，在某些网红盘、改善审美赛道上不如绿城、滨江、金茂府热闹。投研容易低估它的去化：它不需要热闹，只要定价对着目标客群的支付能力。风险在于过度克制可能错过结构行情，以及个别城市更新/收并购的整合。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "3. 建发", 3)
    _add_paragraph(
        doc,
        "销售能力非常突出：高周转、强渠道、快开盘、快回笼，在长三角和福建能把货推出去。管理上是“销售驱动的运营”，投资决策快，拿地曾经非常积极，这在上行期是优点，在2021年后变成利润和土储质量的压力。看建发，不要用看中海的利润率去要求它，而要看它的周转、权益比、合作盘占比和表外。它能卖，但“卖得贵、卖得剩利润”不是它的第一标签。与建发合作，界面通常专业、节奏快；要盯的是对赌、操盘权和资本金同步到位。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "4. 中国金茂", 3)
    _add_paragraph(
        doc,
        "金茂府是少数能在改善盘里形成全国溢价的产品系列，销售端在对的城市很能打。管理的分裂在于：住宅开发一套逻辑，城市运营/持有一套逻辑。后者货值故事大、现金流回收慢，容易让集团看起来“有资产、缺利润”。投资端有过溢价拿地、主题地块配建负担重的案例。评价金茂，要把府系住宅和城市运营拆开：前者管理与销售都不弱，后者决定了利润什么时候能进现金。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "5. 绿城", 3)
    _add_paragraph(
        doc,
        "销售能力在改善客群中是第一阵营，示范区、园区、口碑能转化为溢价。管理能力则明显分叉：产品管理强，投资与成本管理弱，合资治理复杂。中交入主后，资金和土储来源更国企化，但组织磨合、关联界面、利润转化并没有自动变成中海。代建（绿城管理）的销售和管理口碑，不能直接记到开发表上。若只选“谁会做产品”，绿城靠前；若选“谁会把利润做出来”，它不靠前。细节见第七问。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "6. 滨江", 3)
    _add_paragraph(
        doc,
        "把“区域深耕”做到极致的样本。在杭州及浙江，产品、定价、渠道、政府关系和客户口碑形成闭环，销售能力可以给到一线最高档。管理上，决策半径短、创始人/专业团队对市场的感觉好，投资很少离开自己懂的水位。一旦离开浙江，品牌溢价和去化效率都会下降，这是它和保利、华润的差距。利润能力好于绿城这类全国改善商，因为更少为品牌而溢价拿地。风险就是单一区域和单一客群：杭州一冷，报表就冷。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "7. 越秀", 3)
    _add_paragraph(
        doc,
        "广州国企的典型气质：稳健、不抢戏、资源型。与地铁等股东的合作，能拿到别人拿不到的地，这是管理能力在“投资入口”上的优势，而不是造房子造得更花。销售在大湾区有基础，全国声量一般，去化靠城市深耕而不是全国渠道机器。商业和持有能平滑周期，但也降低周转。把它放在“不会出事、也不会突然爆发”的篮子里，符合大多数同行的体感。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "8. 保利发展", 3)
    _add_paragraph(
        doc,
        "销售能力是全国级的：品牌即信任，覆盖深，渠道全，下行期“国企、能交付”本身就是最强卖点。管理能力属于第一阵营但不是最尖的那个——投资纪律好过激进民企，细腻程度不如中海；组织大，区域公司水平参差，个别高价地、个别城市会拖后腿。利润中等偏好，规模是它的盾。看保利，重点看结构：核心城市货值占比、合作盘权益、表外担保，而不是它会不会卖。它会卖。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "9. 招商蛇口", 3)
    _add_paragraph(
        doc,
        "资源禀赋在这九家里非常独特：蛇口、前海、城市更新、产业园区，不是普通招拍挂开发商能复制的。销售端有品牌，核心盘不愁；但城市更新和综合开发的节奏天然慢，利润释放不利落，管理上要同时驾驭开发、持有、园区、邮轮等板块，注意力分散。和华润比，商业运营的品牌密度不如万象系；和中海比，开发利润率不够锐。适合当“资源+城市更新”标的来理解，不适合当高周转开发商来要求。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "10. 放在一张图里怎么用", 3)
    _add_paragraph(
        doc,
        "如果只记住两个轴：横轴销售能力，纵轴管理/利润转化。右上角是华润、中海（中海更偏上、华润更偏右）；右中是保利、建发、滨江（建发、滨江更右，利润轴略低）；中上是越秀；中部是金茂、招商蛇口；产品右移、利润下移的是绿城。这个相对位置比绝对分数有用。合作选对手时：要利润和对账清晰选中海、华润；要去化和渠道选建发、滨江、保利；要产品溢价选绿城、金茂、滨江；要资源和慢变量选越秀、招商。",
        first_line_indent=0.74,
    )

    # ============================================================
    # 附录：把四个“专家必答”再压缩一次
    # ============================================================
    _add_heading(doc, "附录  四个“专家必答”的压缩口径", 1)
    _add_paragraph(
        doc,
        "来问里后半部分把四个问题又列了一遍。这里用更短的口径再答一次，便于直接贴进纪要。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "必答1：缴税完整过程", 2)
    _add_paragraph(
        doc,
        "预售时三税齐预缴：增值税按预收款3%；土增税按不含税口径乘当地预征率（东部普通住宅现在很多落在1.5%左右，商办更高）；企业所得税按预计计税毛利率（省会下限15%）计算预计毛利后缴25%。交房后增值税按9%一般计税（可扣土地价款）正式申报。项目达清算条件后做土增税清算，四级超率累进，普宅增值率≤20%可免税，与预征多退少补。所得税完工后按实际计税成本调整。退税：法律上土增和所得都可以退，实务中土增退税难、所得税多缴以结转抵扣为主、增值税以留抵为主。亏损税盾主要是企业所得税，且困在项目公司，不能给其他项目用。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "必答2：发债和贷款能不能拿地", 2)
    _add_paragraph(
        doc,
        "公募债募集资金不得用于购置土地；开发贷不得用于缴土地出让金，且通常要求四证和资本金先到位。能拿地的正规资金是自有资金、股东增资和股东借款；比较干净的银行路径只有并购贷（买项目公司股权）。行业存在“发债置换出自有资金再去拿地”的操作，属于合规灰区，模型上不要把债券当成土地款来源。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "必答3：项目公司原理、风险隔离、回款路径", 2)
    _add_paragraph(
        doc,
        "单独立项公司是为了证照、抵押、合资、税务清算和并表开关，不只是为了有限责任。极端风险下，上市公司躲不开担保代偿、股东借款损失、保交楼垫资和交叉违约。销售回款先进项目公司预售监管账户，满足节点后才能上收总部；不是客户直接打给上市公司。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "必答4：合营/联营风险与盈亏分化", 2)
    _add_paragraph(
        doc,
        "把表外项目当并表项目一样拆：控制权、货值、项目负债、担保、往来、对手方、对赌。联营更容易亏，是因为小股、别人操盘、出表残留和业态更杂，不是权益法的锅。无息无抵押往来是劣后资本，合作方或项目爆雷时回收率极低，应加回净暴露，不能因为“无息”就当成无风险。",
        first_line_indent=0.74,
    )

    _add_heading(doc, "附：讨论时建议避开的六个常见误判", 2)
    _add_bullet(doc, "“有项目公司，所以上市平台没有项目风险。”——先看担保和往来。")
    _add_bullet(doc, "“集团亏损很大，所以明年不用缴税。”——预缴和法人独立会打你的脸。")
    _add_bullet(doc, "“账上现金很多，所以能还债、能拿地。”——先扣监管户和土地款。")
    _add_bullet(doc, "“合营赚钱、联营亏钱，所以合营都是好资产。”——按项目清单拆。")
    _add_bullet(doc, "“投委会净利率10%，所以跌10%还保本。”——成本刚性，不保本。")
    _add_bullet(doc, "“一线没降价新闻，所以一线盘很健康。”——看暗降和二手倒挂。")

    _add_paragraph(
        doc,
        "（完）本文供内部讨论。具体项目仍应以该城市当时有效的预征率文件、预售资金监管细则、贷款合同和合资协议为准。",
        first_line_indent=0.74,
        color=GRAY,
        space_before=12,
        italic=True,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == "__main__":
    build_document(OUTPUT)
    print(f"已生成：{OUTPUT}")
