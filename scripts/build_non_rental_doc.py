"""
非租金收入解读 · Word 文档生成脚本（投决会备答版）
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

CN_FONT = "WenQuanYi Micro Hei"
EN_FONT = "Calibri"
OUT_DIR = Path(__file__).resolve().parent.parent / "docs" / "phase4-commercial"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def set_cn_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = EN_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_title(doc, text, size=20):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cn_font(p.add_run(text), size=size, bold=True,
                color=RGBColor(0x0F, 0x2D, 0x52))


def add_subtitle(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cn_font(p.add_run(text), size=size, color=RGBColor(0x6B, 0x73, 0x80))


def add_h1(doc, text):
    p = doc.add_paragraph()
    set_cn_font(p.add_run(text), size=15, bold=True,
                color=RGBColor(0x0F, 0x2D, 0x52))


def add_h2(doc, text):
    p = doc.add_paragraph()
    set_cn_font(p.add_run(text), size=13, bold=True,
                color=RGBColor(0x1F, 0x6F, 0xEB))


def add_h3(doc, text):
    p = doc.add_paragraph()
    set_cn_font(p.add_run(text), size=12, bold=True)


def add_p(doc, text, indent=False, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    set_cn_font(p.add_run(text), size=11, bold=bold,
                italic=italic, color=color)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_cn_font(p.add_run(text), size=11)


def add_table(doc, headers, rows, header_fill=RGBColor(0x0F, 0x2D, 0x52),
              col_widths=None):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, cw in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = cw

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        set_cn_font(cell.paragraphs[0].add_run(h),
                    size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # 设置背景色
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "0F2D52")
        tcPr.append(shd)

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            set_cn_font(cell.paragraphs[0].add_run(str(val)), size=10)


def doc_setup(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    rPr.append(rFonts)


def build():
    doc = Document()
    doc_setup(doc)

    # 封面
    add_title(doc, "非租金收入构成解读")
    add_subtitle(doc, "服务平台佣金 + 后市场协同分成 · 产业逻辑 / 6+6 子项 / 投决会备答")
    add_subtitle(doc, "GS · iDrive Hub · 01# 研发楼")
    add_subtitle(doc, "v1.2.1 · 投决会汇报版")
    doc.add_paragraph()

    # === 一、为什么单独解读 ===
    add_h1(doc, "一、为什么这两项要单独解读")

    add_h3(doc, "1.1 财务模型位置")
    add_table(doc,
              ["收入项", "Y1", "Y2", "Y3", "Y3 占比"],
              [
                  ["租金 + 物业", "712", "1,637", "2,348", "41%"],
                  ["服务平台佣金", "80", "250", "600", "10%"],
                  ["1F+2F 冠名", "100", "300", "500", "9%"],
                  ["后市场协同分成", "50", "150", "300", "5%"],
                  ["政策返还", "80", "600", "1,500", "26%"],
                  ["基金管理费", "0", "300", "500", "9%"],
                  ["总收入", "1,022", "3,237", "5,748", "100%"],
              ])
    doc.add_paragraph()

    add_h3(doc, "1.2 关键观察 · 这两项决定项目生死")
    add_p(doc, "若仅靠租金 + 物业，Y3 收入 2,348 万 < 总成本 4,380 万 → "
               "EBITDA 必定亏损 −2,032 万。", indent=True, bold=True,
          color=RGBColor(0xD0, 0x4A, 0x4A))
    add_p(doc, "加上这两项 + 冠名 + 政策 + 基金合计 3,400 万，Y3 EBITDA 才能"
               "转正 +1,368 万。", indent=True, bold=True,
          color=RGBColor(0x2F, 0xA3, 0x6F))
    add_p(doc, "→ 这两项不是『锦上添花』，而是『项目能否赚钱』的决定性因子。",
          indent=True, bold=True)

    doc.add_paragraph()

    # === 二、服务平台佣金 ===
    add_h1(doc, "二、服务平台佣金（Y3 600 万）· 议价权变现")

    add_h3(doc, "2.1 本质定义")
    add_p(doc, "服务平台佣金 = 园区作为『智驾产业服务中台』，统一对接外部供应商，"
               "将服务转售给入驻企业（链主+生态），赚取的价差/分成佣金。",
          indent=True, italic=True)
    add_p(doc, "简单说：园区不是『卖楼』，而是把入驻企业聚合后的议价权变现，"
               "把服务以『团购价/集采价』分发，吃中间利差。", indent=True)

    add_h3(doc, "2.2 Y3 600 万构成 · 6 大子项")
    add_table(doc,
              ["子项", "Y3 (万)", "业务模式", "关键合作方", "抽成档"],
              [
                  ["A. 算力服务转售", "250", "园区年采购 GPU 时长，按算力券分销吃价差",
                   "阿里云 / 华为云 / 上海超算", "价差 8%–15%"],
                  ["B. 政府事务服务费", "150", "一企一策协调、牌照代办、L3/L4 试点推荐、人才落户",
                   "区投促办 / 市经信委 / 市公安", "单项 5–20 万"],
                  ["C. 联合实验室分成", "80", "3F 联合实验室对外承接 Tier1 演示/联调",
                   "地平线 / 链主企业", "出资比例分成"],
                  ["D. 招聘联运", "60", "高级研发岗位渠道分成 + 落户绿通服务费",
                   "光辉国际 / Robert Walters", "单笔 8%–15%"],
                  ["E. 法务/IP/咨询分成", "60", "律所/会计师事务所驻点 · 按客户介绍分成",
                   "君合 / 方达 / 毕马威", "单笔 5%–10%"],
              ])

    add_h3(doc, "2.3 三年爬坡逻辑")
    add_p(doc, "Y1 (35% 入驻率) → 80 万", bold=True)
    add_p(doc, "  └ 刚起步、链主未签约，主要靠政策代办与小规模算力券", indent=True)
    add_p(doc, "Y2 (70%) → 250 万", bold=True)
    add_p(doc, "  └ 链主入驻 1 家、生态 9 家，算力 GMV 起量", indent=True)
    add_p(doc, "Y3 (92%) → 600 万", bold=True)
    add_p(doc, "  └ 14 家入驻全负荷，3F 实验室对外服务起量", indent=True)

    add_h3(doc, "2.4 商业模式特点")
    add_bullet(doc, "议价权变现：园区以『1 链主 + 14 家生态』的体量议价，单户企业拿不到这个价")
    add_bullet(doc, "边际成本低：1 个 IT/数据经理 + 1 个法务对接，年增量人力成本 ≤ 100 万")
    add_bullet(doc, "强生态依附：Y3 600 万中 70% 与入驻量直接挂钩")
    add_bullet(doc, "政策合规：政府事务代办不能跨越『代办』边界，不承诺审批结果，不收回扣")

    add_h3(doc, "2.5 实操路径 · 4 个里程碑")
    add_table(doc,
              ["里程碑", "时间", "关键动作", "负责人"],
              [
                  ["M3", "启动 + 3 个月", "与 2–3 家算力供应商签框架协议", "招商总监 + IT 经理"],
                  ["M5", "启动 + 5 个月", "D 栋驻点律所/会计师事务所签年度合作", "BD 总监 + 法务"],
                  ["M8", "启动 + 8 个月", "3F 联合实验室开放，承接首批外部 Tier1 客户",
                   "运营总监 + 测试场协同主任"],
                  ["M12", "启动 + 12 个月", "服务平台 SOP 跑通，月度 GMV 突破 100 万元", "项目总监"],
              ])

    add_h3(doc, "2.6 风险与对冲")
    add_table(doc,
              ["风险", "对冲手段"],
              [
                  ["单一供应商绑架（如阿里云涨价）", "同时签 2–3 家供应商，每家不超过 40%"],
                  ["GMV 不达预期", "Y1 保守预估（80 万 = 入驻率 × 同业基准），Y3 才到 600 万"],
                  ["政府代办合规边界", "法务月度审查 + 不承诺审批结果 + 收费明码标价"],
                  ["客户挑战『为什么不直接找供应商』",
                   "服务包价 ≤ 直采价 95% + 1 对 1 服务体验差异化"],
              ])

    doc.add_paragraph()

    # === 三、后市场协同分成 ===
    add_h1(doc, "三、后市场协同分成（Y3 300 万）· 冠松独家壁垒")

    add_h3(doc, "3.1 本质定义")
    add_p(doc, "后市场协同分成 = 冠松集团的汽车后市场资源（4S/保险/二手车/融资租赁/车队等）"
               "与园区企业的智驾产品/数据进行商业化协同，产生的协同收入分成。",
          indent=True, italic=True)
    add_p(doc, "★ 这是除冠松外无法复制的独家壁垒。其他园区即使复制了"
               "『独栋 + 政策』，也复制不了冠松集团 30+ 年汽车后市场积累的"
               "真实数据/牌照/网点资源。",
          indent=True, bold=True, color=RGBColor(0xC9, 0xA2, 0x4A))

    add_h3(doc, "3.2 冠松集团 6 大可协同资源")
    add_table(doc,
              ["资源", "规模", "对智驾的价值"],
              [
                  ["① 4S 经销网络", "华东 60+ 网点 · 鸿蒙智行/主流品牌",
                   "真实事故/维修/智驾故障数据自然产生"],
                  ["② 保险事业部", "智驾保险 / 定损 / 理赔大数据",
                   "智驾保险产品需要的专家定损能力"],
                  ["③ 二手智驾车", "检测 / 翻新 / 流通",
                   "残值评估 / 智驾包激活流转数据"],
                  ["④ 融资租赁子公司", "汽车融资租赁牌照",
                   "测试车队融资利率比市场低 1.5–2%"],
                  ["⑤ 冠松车队（试运营）", "现役 100+ 车",
                   "真实上海典型场景数据采集"],
                  ["⑥ 冠松产业基金", "战投部 + 5 亿规模",
                   "LP/GP 双模式 + 跟投权 + 资源捆绑"],
              ])

    add_h3(doc, "3.3 Y3 300 万构成 · 6 大子项")
    add_table(doc,
              ["子项", "Y3 (万)", "业务模式", "与谁分成"],
              [
                  ["A. 数据闭环分成", "120", "4S 真实事故/维修/智驾数据脱敏后通过数交所沙盒订阅给链主",
                   "园区合规撮合方 · 抽 10–15%"],
                  ["B. 保险定损协同", "80", "智驾事故/NCAP 保险产品 · 冠松定损 + 园区算法 = 联合保险",
                   "与平安/人保/太保分成"],
                  ["C. 二手智驾车认证", "30", "智驾包激活流转 + 残值评估",
                   "园区认证中台抽成"],
                  ["D. 测试车队融资", "30", "链主测试车队（50–200 辆/家）冠松融资租赁提供融资",
                   "撮合 + 冠松内部分成"],
                  ["E. 冠松车队数据采集", "30", "冠松营运车队改装为数据采集车",
                   "与链主直接分成"],
                  ["F. 体验店流量分成", "10", "1F 大堂智驾后市场体验店 · 园区企业产品分销",
                   "抽成 GMV"],
              ])

    add_h3(doc, "3.4 四道护城河（不可复制壁垒）")
    add_bullet(doc, "数据真实性壁垒：4S 网络真实数据自然产生，竞争对手买不到、造不出")
    add_bullet(doc, "保险牌照壁垒：智驾保险定损依赖于持牌保险经纪/公估机构合作，冠松已具备")
    add_bullet(doc, "资金成本壁垒：冠松融资租赁子公司利率比市场低 1.5–2%")
    add_bullet(doc, "数据合规壁垒：通过上海数交所合规沙盒处理过的数据 = 链主能合规使用，"
                    "对外资链主（华为/百度/SHEIN 等）极为关键")

    add_h3(doc, "3.5 实操路径 · 5 个里程碑")
    add_table(doc,
              ["里程碑", "时间", "关键动作", "负责人"],
              [
                  ["M2", "启动 + 2 个月", "与冠松集团各子公司签《资源协同框架协议》",
                   "集团董事长 + 法务"],
                  ["M4", "启动 + 4 个月", "与上海数交所对接，建立脱敏数据合规沙盒",
                   "GR 总监 + 法务"],
                  ["M6", "启动 + 6 个月", "第一笔数据订阅订单（建议从地平线/Momenta 试点）",
                   "项目总监"],
                  ["M9", "启动 + 9 个月", "与平安/人保/太保至少 1 家签订智驾保险定损合作",
                   "BD 总监"],
                  ["M12", "启动 + 12 个月", "完成首个链主测试车队的融资租赁交付", "项目总监"],
              ])

    add_h3(doc, "3.6 风险与对冲")
    add_table(doc,
              ["风险", "对冲手段"],
              [
                  ["冠松集团子公司不愿『真协同』（怕利益冲突）",
                   "M2 框架协议落地，写入子公司 KPI 考核；财务模型保守按『内部转移定价 + 10% 让利』"],
                  ["数据合规出问题",
                   "100% 通过上海数交所沙盒；4S 数据脱敏后处理；法务年框 + 内审 SOP"],
                  ["保险产品监管不批",
                   "双轨：直接保险定损分成 + 间接数据服务合作"],
                  ["链主不付费订阅数据",
                   "M6 试点期：先免费给地平线，换其推荐 + 联名背书；M9 起转付费"],
              ])

    doc.add_paragraph()

    # === 四、投决会三大应答 ===
    add_h1(doc, "四、汇报应答口径（投决会 + 政府汇报）")

    add_h3(doc, "Q1：服务平台 600 万 / 后市场 300 万 是怎么算出来的？")
    add_p(doc, "答：每项都有 5–6 个明确子项，每个子项有可计算的市场基准价：", indent=True)
    add_bullet(doc, "服务平台 600 万：算力转售 250（年采购 3,000 万 × 8% 价差）+ "
                    "政府代办 150（30 家 × 5 万均价）+ 实验室 80（5 个项目 × 16 万）+ "
                    "招聘 60（20 个高级岗 × 3 万）+ 法务 60（30 家 × 2 万）")
    add_bullet(doc, "后市场 300 万：数据订阅 120（5 个链主级 × 24 万均价）+ "
                    "保险 80（与 1–2 家保司分成）+ 其他 100（C/D/E/F 累计）")

    add_h3(doc, "Q2：万一冠松集团内部协同推不动？")
    add_p(doc, "答：写入 v1.2 的合作协议 Word 中（《资源协同框架协议》），且我们的财务模型"
               "保守按『内部转移定价 + 10% 让利』测算。即使冠松内部协同效率打 70% 折扣，"
               "后市场分成 Y3 仍能达到 210 万（300 × 70%），影响整体 EBITDA 仅 −90 万。",
          indent=True)

    add_h3(doc, "Q3：数据合规是不是定时炸弹？")
    add_p(doc, "答：① 100% 通过上海数交所合规沙盒处理，零原始数据外流；"
               "② 法务年框 + 月度审查 + 客户分级；"
               "③ 外资链主（如华为/百度/外资 ADAS）使用单独通道；"
               "④ 若新法规出台，72 小时内启动暂停 SOP。", indent=True)

    doc.add_paragraph()

    # === 五、对标 ===
    add_h1(doc, "五、产业园非租金收入占比对标")
    add_table(doc,
              ["园区", "非租金占比", "构成", "备注"],
              [
                  ["张江集成电路设计园", "~25%", "IC IP 分成 + 共享 EDA + 测试服务", "头部产业园样板"],
                  ["北京中关村智造大街", "~18%", "技术服务 + 孵化分成", "—"],
                  ["上海大零号湾", "~20%", "高校转化分成 + 服务", "—"],
                  ["杭州梦想小镇", "~30%", "含基金管理费 + 投资收益", "—"],
                  ["★ GS · iDrive Hub 目标", "24% (Y3 1,400/5,748)", "服务+冠名+后市场+基金",
                   "在头部产业园合理区间"],
              ])

    add_h2(doc, "结论")
    add_p(doc, "GS · iDrive Hub 的非租金收入目标 24% 处于头部产业园合理区间，"
               "且依托『冠松后市场独家壁垒』+ 『智驾产业链中台议价权』两条不可复制路径，"
               "可信度高于纯文创/纯科技园区类同行。",
          indent=True, bold=True, color=RGBColor(0x2F, 0xA3, 0x6F))

    out = OUT_DIR / "07c-非租金收入解读-投决会备答版.docx"
    doc.save(out)
    print(f"✓ Word 写入：{out}")
    return out


if __name__ == "__main__":
    build()
