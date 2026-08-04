#!/usr/bin/env python3
"""
以红字标注版附件四为底稿，对照盖章版落实用户确认的修改建议。
规则：
- 原有内容尽量保留；
- 新增内容：红色字体；
- 修改内容：蓝色字体；
- 盖章版《活动赞助合同》本身不改动。
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

SRC = Path(
    "/home/ubuntu/.cursor/projects/workspace/uploads/"
    "___-____-____-________________-_____________ec7d.docx"
)
OUT = Path(
    "/workspace/deliverables/"
    "附件四-绿城中国-潮鸣外滩-晚宴冠名战略合作伙伴专项合作协议-按盖章版补充(红字标注).docx"
)

RED = RGBColor(0xFF, 0x00, 0x00)
BLUE = RGBColor(0x00, 0x70, 0xC0)


def _style_from(paragraph):
    size = name = bold = None
    for r in paragraph.runs:
        if size is None and r.font.size:
            size = r.font.size
        if name is None and r.font.name:
            name = r.font.name
        if bold is None and r.bold is not None:
            bold = r.bold
    return size, name, bold


def paint(run, color: RGBColor, size=None, name=None, bold=None):
    if size:
        run.font.size = size
    if name:
        run.font.name = name
        r_fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), name)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = color
    # 清除可能残留的高亮/底纹，避免与红/蓝标注冲突
    r_pr = run._r.get_or_add_rPr()
    hl = r_pr.find(qn("w:highlight"))
    if hl is not None:
        r_pr.remove(hl)
    shd = r_pr.find(qn("w:shd"))
    if shd is not None:
        r_pr.remove(shd)


def set_paragraph_text(paragraph, text: str, color: RGBColor):
    size, name, bold = _style_from(paragraph)
    if paragraph.runs:
        first = paragraph.runs[0]
        first.text = text
        paint(first, color, size=size, name=name, bold=bold)
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        run = paragraph.add_run(text)
        paint(run, color, size=size, name=name, bold=bold)


def set_cell_text(cell, text: str, color: RGBColor):
    if not cell.paragraphs:
        return
    set_paragraph_text(cell.paragraphs[0], text, color)
    for p in cell.paragraphs[1:]:
        for r in p.runs:
            r.text = ""


def append_colored_paragraph_after(paragraph, text: str, color: RGBColor, size_pt: float = 10.5):
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    p_obj = Paragraph(new_p, paragraph._parent)
    run = p_obj.add_run(text)
    paint(run, color, size=Pt(size_pt))
    return p_obj


def append_colored_in_cell(cell, text: str, color: RGBColor, size_pt: float = 9):
    p = cell.add_paragraph()
    run = p.add_run(text)
    paint(run, color, size=Pt(size_pt))
    return p


def find_paragraph(doc: Document, prefix: str):
    for p in doc.paragraphs:
        if p.text.startswith(prefix):
            return p
    raise KeyError(prefix)


def unique_cells(row):
    seen = set()
    out = []
    for cell in row.cells:
        cid = id(cell._tc)
        if cid in seen:
            continue
        seen.add(cid)
        out.append(cell)
    return out


def main() -> None:
    doc = Document(str(SRC))

    # ========== 1. 笪浩联系方式修改 → 蓝字 ==========
    t1 = doc.tables[1]
    set_cell_text(t1.rows[6].cells[1], "18502131428", BLUE)

    # ========== 2. 2.4 排他收窄 → 蓝字修改 ==========
    p24 = find_paragraph(doc, "2.4")
    set_paragraph_text(
        p24,
        "2.4  甲方确认在本次峰会期间，不再向同一行业（中国房地产开发与销售、品质改善型住宅项目）的其他企业授予"
        "「晚宴冠名」身份及「主背景板钻石级 logo 位」；其他级别（黄金/铂金/基础曝光等）不在排他范围。"
        "如需引入同行业其他同级以上赞助商，应事先以书面形式告知乙方并取得乙方书面同意。",
        BLUE,
    )
    append_colored_paragraph_after(
        p24,
        "【新增说明】排他范围已按实际可执行口径收窄；过宽的「同级以上对外宣传身份」排他在已办结活动场景下难以核实与执行，故不再采用。",
        RED,
        size_pt=10,
    )

    # ========== 第三条权益表 ==========
    t2 = doc.tables[2]  # 晚宴权益 1-9

    # 4. 桌卡 → 晚宴入口欢迎立牌（未实际制作桌卡）蓝字修改
    set_cell_text(
        t2.rows[4].cells[1],
        "晚宴入口欢迎立牌植入项目 logo + 「绿城·潮鸣外滩 · 晚宴冠名」字样",
        BLUE,
    )
    set_cell_text(
        t2.rows[4].cells[2],
        "1 座（立式 KT/亚克力，建议 60×180cm 或现场同等规格；原「桌卡/桌号牌」未实际制作，改为此项）",
        BLUE,
    )

    # 3. 菜单 → 晚宴签到台品牌立牌（未做菜单）蓝字修改
    set_cell_text(
        t2.rows[5].cells[1],
        "晚宴签到台品牌立牌植入项目 logo（含晚宴冠名字样）",
        BLUE,
    )
    set_cell_text(
        t2.rows[5].cells[2],
        "1 套（签到台立牌/背景露出；原「晚宴菜单封面整版+内页页脚」未实际制作，改为此项）",
        BLUE,
    )

    # 5. 晚宴入场券 8 张 → 实际出席 2 位 蓝字修改
    set_cell_text(
        t2.rows[9].cells[2],
        "合计 2 张 / 实际出席 2 人（原约定主桌 3 人 + 销售/接待 5 人共 8 张，以晚宴实际出席为准）",
        BLUE,
    )

    # 主会场表
    t3 = doc.tables[3]
    # 6. 议程手册整版 → 半页 蓝字
    set_cell_text(
        t3.rows[2].cells[1],
        "议程手册广告（半页）+ 项目折页夹页",
        BLUE,
    )
    set_cell_text(
        t3.rows[2].cells[2],
        "半页 A4（约 210×142.5mm），出血 3mm，CMYK 300dpi",
        BLUE,
    )

    # 7. 500 份手拎袋 → 按实际发放 蓝字
    set_cell_text(
        t3.rows[4].cells[1],
        "手拎袋植入项目折页 + 195/285/310 户型图",
        BLUE,
    )
    set_cell_text(
        t3.rows[4].cells[2],
        "三折页 297×210mm + 户型单页；按实际参会嘉宾手拎袋发放数量为准，不超过 500 套",
        BLUE,
    )
    append_colored_in_cell(
        t3.rows[4].cells[-1],
        "【新增】不作未达 500 套即违约解释；验收以实际发放数量及代表性影像为准。",
        RED,
    )

    # 展位/引导/车队
    t4 = doc.tables[4]
    # 8. 展位与盖章版 3 个对齐：蓝字改规格说明 + 红字新增第3处确认
    set_cell_text(
        t4.rows[1].cells[3] if len(unique_cells(t4.rows[1])) > 3 else t4.rows[1].cells[-2],
        "1 处，含品牌桌台轻包装、易拉宝×2、户型 KT 板、销售 2 人（盖章版现场展位共 3 个之一）",
        BLUE,
    )
    # table 4 has weird merge: cells 0=序号, 1/2=权益, 3=规格, 4=责任方
    cells_r1 = unique_cells(t4.rows[1])
    cells_r2 = unique_cells(t4.rows[2])
    # find 规格列：通常倒数第二
    set_cell_text(
        t4.rows[1].cells[3],
        "1 处，含品牌桌台轻包装、易拉宝×2、户型 KT 板、销售 2 人（盖章版现场展位共 3 个之一）",
        BLUE,
    )
    set_cell_text(
        t4.rows[2].cells[3],
        "1 处，含洽谈圆桌、易拉宝×2、户型 KT 板、销售 2 人（盖章版现场展位共 3 个之一）",
        BLUE,
    )
    append_colored_in_cell(
        t4.rows[2].cells[-1],
        "【新增·对照盖章版】盖章版约定现场展位共 3 个；本附件正文列电梯口+展场内 2 处，"
        "第 3 处位置由双方现场协商确定，或以盖章版「3 个展位、位置现场协商」为准。",
        RED,
    )

    # 9/20. 引导员 3→1 蓝字
    set_cell_text(
        t4.rows[3].cells[3],
        "1 次主口播 + 现场引导员 1 名（其余视嘉宾流量灵活补位，不作硬性 3 名承诺）",
        BLUE,
    )

    # 10. 车队统一 8 辆，弱化往返 60 分钟 蓝字
    set_cell_text(
        t4.rows[4].cells[1],
        "项目案场接驳：华为尊界 8 辆 + 考斯特补位",
        BLUE,
    )
    # merged duplicate cell
    if t4.rows[4].cells[2].text.strip():
        set_cell_text(
            t4.rows[4].cells[2],
            "项目案场接驳：华为尊界 8 辆 + 考斯特补位",
            BLUE,
        )
    set_cell_text(
        t4.rows[4].cells[3],
        "由乙方视现场意向人数安排并承担费用（口径统一为尊界 8 辆 + 考斯特补位；不以「往返不低于 60 分钟」为硬性验收条件）",
        BLUE,
    )

    # 媒体表
    t5 = doc.tables[5]
    # 11/21. 媒体不作 ≥3 家硬性承诺 蓝字
    set_cell_text(
        t5.rows[1].cells[2],
        "甲方尽力推送；实际刊发数量、位置与口径以媒体自主决定为准，不作「≥3 家主流媒体」硬性承诺",
        BLUE,
    )
    # 12/22. 永久入册不纳入 蓝字
    set_cell_text(
        t5.rows[4].cells[1],
        "双校长三角校友产业联盟战略合作伙伴入册（本次不纳入必交付）",
        BLUE,
    )
    set_cell_text(
        t5.rows[4].cells[2],
        "不纳入本次计价与必交付；执行机制双方另议（「永久」表述不作验收标准）",
        BLUE,
    )
    append_colored_in_cell(
        t5.rows[4].cells[-1],
        "【新增】本项建议在附表一确认列打「×」，或双方书面另议后再执行。",
        RED,
    )
    # 13/23. 证书名称统一盖章版 蓝字
    set_cell_text(
        t5.rows[6].cells[1],
        "「2026 年度智慧人居新质资产领军企业 暨卓越战略合作伙伴」证书 / 奖项",
        BLUE,
    )
    set_cell_text(t5.rows[6].cells[2], "1 份（与盖章版主合同名称一致）", BLUE)

    # ========== 时间节点表 ==========
    t6 = doc.tables[6]
    # 同步桌卡/菜单名称变更
    set_cell_text(
        t6.rows[2].cells[3],
        "晚宴 KV / 入口欢迎立牌 / 签到台立牌 / 席卡设计稿初稿；晚宴 LED 像素比反馈",
        BLUE,
    )
    set_cell_text(
        t6.rows[4].cells[3],
        "入口欢迎立牌、签到台立牌、易拉宝、户型折页等物料制作（以实际已执行情况为准）",
        BLUE,
    )
    # 15. 执行风险说明 红字新增
    append_colored_in_cell(
        t6.rows[4].cells[-1],
        "【新增·执行风险】签约日（5/21）当日完成互交物料、一轮确认并全部下印，印刷物流周期客观上不可达成；"
        "且活动已于 5/22 完成，该节点表仅作过程记录，不得以未满足「同日下印」倒追违约；"
        "以实际已执行情况为准 / 视同过程节点豁免。",
        RED,
    )
    # 10. 时间表尊界统一 8 辆
    set_cell_text(
        t6.rows[7].cells[2],
        "尊界 8 辆 + 考斯特 + 案场接待",
        BLUE,
    )
    # 16. 名单对接澄清
    set_cell_text(
        t6.rows[7].cells[3],
        "现场引导员 + 现场协助促成意向客户对接（不就完整嘉宾联络清单作硬性提交承诺）",
        BLUE,
    )
    append_colored_in_cell(
        t6.rows[7].cells[3],
        "【新增·对照盖章版】盖章版基于个保法不就参会名单作硬性提交承诺；「名单对接」不得解释为必须交付完整嘉宾联络清单。",
        RED,
    )

    # ========== 10.3 违约金口径 蓝字修改 ==========
    p103 = find_paragraph(doc, "10.3")
    set_paragraph_text(
        p103,
        "10.3  任何一方违反第八条（知识产权）、第九条（保密）的，应向守约方支付违约金；"
        "违约金数额不超过本协议赞助金额的 30%（即不超过人民币 30,000 元），或由双方另行书面约定；"
        "造成实际损失超过违约金的，超过部分应另行赔偿。"
        "本条与第 10.2 条及盖章版《活动赞助合同》违约责任上限体系保持口径一致。",
        BLUE,
    )
    append_colored_paragraph_after(
        p103,
        "【新增说明】原「不低于人民币 50,000 元」已达赞助额 50%，与 10.2 / 盖章版 30% 风险上限体系冲突，故调整为「不超过赞助金额 30% 或双方另行书面约定」。",
        RED,
        size_pt=10,
    )

    # ========== 14. 风险对照总说明 红字新增 ==========
    anchor = find_paragraph(doc, "上述权益项的最终执行清单")
    append_colored_paragraph_after(
        anchor,
        "【新增·风险对照总说明】本附件已对照《活动赞助合同》（盖章版）及以往专项方案中"
        "「已删除/弱化（甲方不好实现项）」对不可达成或难核实事项作口径调整，主要包括："
        "①白皮书扉页署名不纳入；②媒体通稿不作硬性家数承诺；③校友联盟「永久入册」不纳入必交付；"
        "④引导员调整为 1 名；⑤手拎袋按实际发放（上限 500）；⑥展位数量与盖章版 3 个对齐确认；"
        "⑦5/21 同日交稿确认下印节点视同过程豁免；⑧证书名称与盖章版统一；⑨同业排他收窄为晚宴冠名+钻石 logo 位；"
        "⑩未实际制作的菜单/桌卡改换为签到台立牌/入口欢迎立牌；⑪晚宴入场以实际出席 2 人为准；⑫议程手册改为半页。"
        "验收口径与盖章版一致：按四大模块整体打包验收，不以细项逐条未达为由拒付无争议款项；结算与违约上限以盖章版为准。",
        RED,
        size_pt=10,
    )

    # ========== 附表一同步 ==========
    t9 = doc.tables[9]
    # 桌卡→入口欢迎立牌
    set_cell_text(t9.rows[4].cells[1], "晚宴入口欢迎立牌 logo 植入", BLUE)
    set_cell_text(t9.rows[4].cells[2], "晚宴入口欢迎立牌 logo 植入", BLUE)
    set_cell_text(t9.rows[4].cells[3], "1 座（原桌卡/桌号牌改为此项）", BLUE)
    # 菜单→签到台立牌
    set_cell_text(t9.rows[5].cells[1], "晚宴签到台品牌立牌 logo 植入", BLUE)
    set_cell_text(t9.rows[5].cells[2], "晚宴签到台品牌立牌 logo 植入", BLUE)
    set_cell_text(t9.rows[5].cells[3], "1 套（原晚宴菜单改为此项）", BLUE)
    # 入场券
    set_cell_text(t9.rows[9].cells[3], "合计 2 张 / 实际出席 2 人（原约定 8 张，以实际为准）", BLUE)
    # 议程手册半页
    set_cell_text(t9.rows[11].cells[1], "议程手册广告（半页）+ 折页夹页", BLUE)
    set_cell_text(t9.rows[11].cells[2], "议程手册广告（半页）+ 折页夹页", BLUE)
    set_cell_text(t9.rows[11].cells[3], "半页 A4，出血 3mm", BLUE)
    # 18. 500 实际发放
    set_cell_text(t9.rows[13].cells[1], "手拎袋夹页（折页 + 户型图）", BLUE)
    set_cell_text(t9.rows[13].cells[2], "手拎袋夹页（折页 + 户型图）", BLUE)
    set_cell_text(
        t9.rows[13].cells[3],
        "按实际发放（上限 500）；不作未达 500 即违约解释",
        BLUE,
    )
    append_colored_in_cell(t9.rows[13].cells[-1], "【新增】数量以实际发放为准。", RED)
    # 19. 白皮书不纳入
    set_cell_text(t9.rows[14].cells[1], "白皮书扉页联合署名（本次不纳入）", BLUE)
    set_cell_text(t9.rows[14].cells[2], "白皮书扉页联合署名（本次不纳入）", BLUE)
    set_cell_text(t9.rows[14].cells[3], "不纳入必交付（建议确认列打「×」）", BLUE)
    append_colored_in_cell(
        t9.rows[14].cells[2],
        "【新增】以往方案已明确不纳入；正文第三节权益表序号由 13 跳至 15 亦未列本项，附表保留易致验收争议。",
        RED,
    )
    # 展位说明
    append_colored_in_cell(
        t9.rows[16].cells[-1],
        "【新增】与盖章版 3 个展位对齐；第 3 处位置现场协商。",
        RED,
    )
    # 20. 引导员 1 名
    set_cell_text(t9.rows[17].cells[3], "1 次主口播 + 引导员 1 名", BLUE)
    # 车队 8 辆
    set_cell_text(t9.rows[18].cells[1], "项目案场接驳（华为尊界 8 辆 + 考斯特补位）", BLUE)
    set_cell_text(t9.rows[18].cells[2], "项目案场接驳（华为尊界 8 辆 + 考斯特补位）", BLUE)
    set_cell_text(t9.rows[18].cells[3], "由乙方视意向人数安排并承担费用（统一 8 辆口径）", BLUE)
    # 21. 媒体
    set_cell_text(t9.rows[20].cells[1], "媒体通稿（不作硬性家数承诺）", BLUE)
    set_cell_text(t9.rows[20].cells[2], "媒体通稿（不作硬性家数承诺）", BLUE)
    set_cell_text(t9.rows[20].cells[3], "甲方尽力推送，刊发以媒体自主决定为准", BLUE)
    # 22. 永久入册
    set_cell_text(t9.rows[23].cells[1], "校友产业联盟入册（本次不纳入必交付）", BLUE)
    set_cell_text(t9.rows[23].cells[2], "校友产业联盟入册（本次不纳入必交付）", BLUE)
    set_cell_text(t9.rows[23].cells[3], "不纳入；建议打「×」或双方另议", BLUE)
    # 23. 证书
    set_cell_text(
        t9.rows[25].cells[1],
        "「2026 年度智慧人居新质资产领军企业 暨卓越战略合作伙伴」证书",
        BLUE,
    )
    set_cell_text(
        t9.rows[25].cells[2],
        "「2026 年度智慧人居新质资产领军企业 暨卓越战略合作伙伴」证书",
        BLUE,
    )
    set_cell_text(t9.rows[25].cells[3], "1 份（与盖章版一致）", BLUE)

    # 附表一验收原则 红字
    for p in doc.paragraphs:
        if p.text.startswith("本附表逐项列示"):
            append_colored_paragraph_after(
                p,
                "【新增·验收原则】已调整/不纳入细项（蓝字修改、红字新增说明）如双方未另签确认，"
                "不作为单项违约计价依据；与盖章版或报价单四大模块冲突时，以盖章版/整体打包验收为准。",
                RED,
                size_pt=10,
            )
            break

    # 正文 5.5 已是 8 辆，保持；若有 8–10 再改
    p55 = find_paragraph(doc, "5.5")
    if "8–10" in p55.text or "8-10" in p55.text:
        set_paragraph_text(
            p55,
            "5.5  乙方应自行承担华为尊界 8 辆 + 考斯特补位接驳、项目案场接待、销售人员差旅及现场展位包装升级等乙方直接承办环节的相关费用，该等费用不计入本协议第 2.2 条之赞助金额。",
            BLUE,
        )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"saved: {OUT}")


if __name__ == "__main__":
    main()
