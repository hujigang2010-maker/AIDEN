# -*- coding: utf-8 -*-
"""将《住房产业三链融合白皮书》Markdown 文稿转换为 Word（docx）。

- 封面含复旦大学住房政策研究中心 logo 与文稿编号
- 章节标题、表格、图表、引用块均按正式白皮书样式排版
- 页眉页脚含中心名称、文稿编号与页码
"""
import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE = "/workspace/whitepaper"
MD = os.path.join(BASE, "住房产业三链融合白皮书.md")
OUTPUT = os.path.join(BASE, "住房产业三链融合白皮书.docx")

BLUE = RGBColor(0x0E, 0x4E, 0x9B)
RED = RGBColor(0xC8, 0x10, 0x2E)
GRAY = RGBColor(0x6B, 0x6F, 0x78)


def set_font(run, cn="宋体", en="Times New Roman", size=12, bold=False,
             color=None, italic=False):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    r = run._element.rPr
    rfonts = r.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = r.makeelement(qn("w:rFonts"), {})
        r.append(rfonts)
    rfonts.set(qn("w:eastAsia"), cn)


def add_rich(par, text, size=12, cn="宋体", color=None, base_bold=False):
    for seg in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            set_font(par.add_run(seg[2:-2]), cn="黑体", size=size, bold=True,
                     color=color)
        elif seg.startswith("*") and seg.endswith("*") and len(seg) > 2:
            set_font(par.add_run(seg[1:-1]), cn="楷体", size=size, italic=True,
                     color=color)
        else:
            set_font(par.add_run(seg), cn=cn, size=size, bold=base_bold,
                     color=color)


def add_page_number(paragraph):
    run1 = paragraph.add_run("— ")
    set_font(run1, size=9, color=GRAY)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    set_font(r, size=9, color=GRAY)
    r._element.append(fld1)
    r2 = paragraph.add_run()
    set_font(r2, size=9, color=GRAY)
    r2._element.append(instr)
    r3 = paragraph.add_run()
    set_font(r3, size=9, color=GRAY)
    r3._element.append(fld2)
    run2 = paragraph.add_run(" —")
    set_font(run2, size=9, color=GRAY)


def shade_cell(cell, fill="DCE6F1"):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


doc = Document()
sec = doc.sections[0]
sec.top_margin, sec.bottom_margin = Cm(2.6), Cm(2.6)
sec.left_margin, sec.right_margin = Cm(2.8), Cm(2.8)
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.different_first_page_header_footer = True

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("复旦大学住房政策研究中心  ·  FDU-HPRC-WP-2026-03"),
         cn="楷体", size=9, color=GRAY)
# 首页页眉留空
sec.first_page_header.paragraphs[0].text = ""

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_number(footer)
sec.first_page_footer.paragraphs[0].text = ""

# ---------------- 封面 ----------------
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
p.add_run().add_picture(os.path.join(BASE, "assets/logo_fudan_hprc.png"),
                        width=Cm(14.2))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(48)
set_font(p.add_run("住房产业“三链融合”白皮书"), cn="黑体", en="Arial",
         size=28, bold=True, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
set_font(p.add_run("供应链、硬件与软件协同的 2030 图景"),
         cn="楷体", size=16, color=GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
set_font(p.add_run("—— 从“造房子”到“造产品、管资产、服务生活” ——"),
         cn="楷体", size=12.5, color=GRAY)

for txt, sz, sp in [("中心研究文稿 · 第三号", 15, 56),
                    ("文稿编号：FDU-HPRC-WP-2026-03", 12.5, 6),
                    ("复旦大学住房政策研究中心", 16, 28),
                    ("Housing Policy Research Center, Fudan University", 11, 2),
                    ("二〇二六年八月 · 上海", 13, 18)]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(sp)
    bold = txt.startswith(("中心研究", "复旦大学"))
    set_font(p.add_run(txt), cn="黑体" if bold else "宋体", size=sz, bold=bold,
             color=BLUE if bold else GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
set_font(p.add_run("学术性  ·  中立性  ·  公益性"), cn="楷体", size=11,
         color=RED)

doc.add_page_break()

# ---------------- 正文解析 ----------------
lines = open(MD, encoding="utf-8").read().split("\n")
i = 0
first_h1_skipped = False
h1_count = 0
while i < len(lines):
    line = lines[i].rstrip()

    if (not line or line == "---" or line.startswith("<p align")
            or line.startswith("<img")):
        i += 1
        continue

    if re.match(r"^\*\*(编制单位|文稿系列|文稿编号|成稿日期)\*\*", line):
        i += 1
        continue

    # 文稿里的二级大标题（封面副标题），跳过已在封面呈现的
    if line.startswith("## ——"):
        i += 1
        continue

    if line.startswith("# "):
        title = line[2:].strip()
        if not first_h1_skipped:
            first_h1_skipped = True
            i += 1
            continue
        if h1_count > 0:
            doc.add_page_break()
        h1_count += 1
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(14)
        set_font(p.add_run(title), cn="黑体", en="Arial", size=18, bold=True,
                 color=BLUE)
        i += 1
        continue

    if line.startswith("### "):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        set_font(p.add_run(line[4:].strip()), cn="黑体", en="Arial", size=13,
                 bold=True, color=BLUE)
        i += 1
        continue

    if line.startswith("## "):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        set_font(p.add_run(line[3:].strip()), cn="黑体", en="Arial", size=14.5,
                 bold=True)
        i += 1
        continue

    m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", line)
    if m:
        path = os.path.join(BASE, m.group(2))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        if os.path.exists(path):
            p.add_run().add_picture(path, width=Cm(15.2))
        else:
            set_font(p.add_run(f"[缺图：{m.group(2)}]"), size=10, color=RED)
        i += 1
        continue

    # 图注
    if line.startswith("*图") and line.endswith("*"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        set_font(p.add_run(line.strip("*")), cn="楷体", size=9.5, italic=True,
                 color=GRAY)
        i += 1
        continue

    if line.startswith("> "):
        buf = []
        while i < len(lines) and lines[i].startswith(">"):
            buf.append(lines[i].lstrip("> ").strip())
            i += 1
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.right_indent = Cm(0.4)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.3
        add_rich(p, " ".join(b for b in buf if b), size=11.5, cn="楷体",
                 color=BLUE)
        continue

    if line.startswith("|"):
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                rows.append(cells)
            i += 1
        if rows:
            ncol = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncol)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                for ci in range(ncol):
                    cell = table.cell(ri, ci)
                    txt = row[ci] if ci < len(row) else ""
                    txt = re.sub(r"\*\*([^*]+)\*\*", r"\1", txt)
                    cp = cell.paragraphs[0]
                    set_font(cp.add_run(txt), cn="黑体" if ri == 0 else "宋体",
                             size=9.5, bold=(ri == 0))
                    if ri == 0:
                        shade_cell(cell)
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(6)
        continue

    if line.startswith("- "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.3
        set_font(p.add_run("• "), cn="黑体", size=12, color=BLUE)
        add_rich(p, line[2:].strip(), size=12)
        i += 1
        continue

    m = re.match(r"^(\d+)\.\s+(.*)", line)
    if m:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.3
        set_font(p.add_run(f"{m.group(1)}. "), cn="黑体", size=12, bold=True,
                 color=BLUE)
        add_rich(p, m.group(2).strip(), size=12)
        i += 1
        continue

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.38
    add_rich(p, line, size=12)
    i += 1

doc.save(OUTPUT)
print("saved:", OUTPUT)
