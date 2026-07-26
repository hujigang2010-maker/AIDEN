# -*- coding: utf-8 -*-
"""
生成《复旦科技园创业孵化基地2026年度工作报告》Word 文档。

内容来源：
- 2023年上海市创业孵化示范基地工作报告 PPT 结构
- 《复旦科技园创业孵化基地2026年度工作报告》原稿（覆盖 2025 年工作）

运行：python3 scripts/build_2026_annual_report.py
输出：dist/复旦科技园创业孵化基地2026年度工作报告.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DIST = ROOT / "dist"
DIST.mkdir(parents=True, exist_ok=True)
OUT = DIST / "复旦科技园创业孵化基地2026年度工作报告.docx"

NAVY = RGBColor(0x0A, 0x2F, 0x6B)
DARK = RGBColor(0x1B, 0x2A, 0x3A)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)


def set_run_font(run, size=12, bold=False, color=DARK, font="宋体", east_asia=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), east_asia or font)


def add_title(doc, text, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=NAVY, font="黑体")
    return p


def add_heading_cn(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    size = 16 if level == 1 else 14
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=NAVY, font="黑体")
    return p


def add_body(doc, text, first_indent=True, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=bold, color=DARK, font="宋体")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.left_indent = Cm(0.74)
        run = p.add_run("● " + item)
        set_run_font(run, size=12, color=DARK, font="宋体")


def add_kpi_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.left_indent = Cm(0.74)
    r1 = p.add_run(f"{label}：")
    set_run_font(r1, size=12, bold=True, color=NAVY, font="宋体")
    r2 = p.add_run(value)
    set_run_font(r2, size=12, bold=False, color=DARK, font="宋体")


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    add_title(doc, "复旦科技园创业孵化基地", size=20)
    add_title(doc, "2026年度工作报告", size=22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("复旦大学国家大学科技园\n2026年6月")
    set_run_font(run, size=12, color=DARK, font="楷体")

    add_heading_cn(doc, "目  录", level=1)
    toc = [
        "一、载体基本情况",
        "二、2025年工作成效",
        "三、与人社部门合作紧密度",
        "四、特色与亮点",
        "五、综合效益与下一步工作",
    ]
    for i, item in enumerate(toc, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{item}")
        set_run_font(run, size=12, color=DARK, font="宋体")

    # ------------------------------------------------------------------
    add_heading_cn(doc, "一、载体基本情况", level=1)
    add_body(
        doc,
        "近年来，复旦科技园创业孵化基地紧扣国家大学科技园优化重塑主线，充分发挥高校科研资源、人才资源和创新生态优势，持续完善创业孵化服务体系，提升科技成果转化承载能力。2025年，基地围绕“成果发现－概念验证－创业孵化－产业培育”创新服务链条，聚焦科技成果转化、创新人才培养、科技企业成长和区域协同发展，推动一批高校科技成果加速转化、一批创新创业项目落地成长、一批科技型企业集聚发展，为区域创新能力提升和产业高质量发展注入新动能。",
    )
    add_body(
        doc,
        "复旦大学国家大学科技园创建于2000年，是科技部、教育部联合认定的首批国家大学科技园。复旦科技园创业孵化基地依托复旦大学国家大学科技园建设运营。基地位于杨浦区国泰路11号，地处复旦大学邯郸校区东南侧，紧邻五角场商圈，场地面积3.4万平方米，其中孵化面积2.56万平方米，是高校校区、科技基地、公共社区“三区联动”的创新实践基地。基地秉持“转化科技、服务社会、汇聚智慧、共创未来”的发展理念，着力打造高校科技成果转化的重要渠道、创新创业生态建设的重要阵地和区域经济发展新动能的重要平台。",
    )
    add_body(
        doc,
        "经过二十余年建设发展，基地已形成以科技成果转化、创业孵化、产业培育和创新人才培养为核心功能的综合性创新创业服务体系，构建“众创空间＋孵化器＋加速器＋产业基地”全链条科创矩阵，并持续完善创业苗圃、复翼互联众创空间、技术转移服务平台、海外人才科创基地等孵化功能板块。",
    )
    add_body(doc, "主要资质与荣誉包括：", first_indent=True)
    add_bullets(
        doc,
        [
            "国家高新技术创业服务中心、国家高校学生科技创业实习基地",
            "上海市创业孵化示范基地、上海市海聚英才创新创业示范基地",
            "上海市文明单位、上海市科技创新创业服务站、上海市知识产权示范基地",
            "曾获国家科技计划（火炬计划）实施20周年全国先进服务机构、中国高校孵化器十强等",
        ],
    )
    add_body(
        doc,
        "截至2025年，基地存续注册企业1500余家，科技型企业占比80%；累计培育高新技术企业115家，专精特新小巨人企业1家，上海市专精特新中小企业28家，上海市科技小巨人（含培育）企业8家，杨浦区科技小巨人（含培育）企业21家，各类资本市场挂牌上市企业14家。初步形成集成电路、软件和信息技术、节能环保、新材料新能源、生物医药等特色产业集群。",
    )

    # ------------------------------------------------------------------
    add_heading_cn(doc, "二、2025年工作成效", level=1)
    add_body(
        doc,
        "2025年，复旦科技园创业孵化基地围绕创业孵化实效、就业带动、政策落实、专业服务、活动组织、品牌宣传等方面，不断提升服务能力，推动更多创新创业项目落地成长。",
    )

    add_heading_cn(doc, "（一）创业孵化实效", level=2)
    add_body(
        doc,
        "基地以科技成果转化和创新创业项目培育为核心，围绕高校师生及校友创业团队和科技型初创企业成长需求，建立覆盖项目发现、资源匹配、创业辅导和企业培育的全过程孵化机制。",
    )
    add_body(
        doc,
        "一是加强创业项目挖掘。基地依托复旦大学学科优势和科研资源，聚焦人工智能、集成电路、生物医药、新材料等重点领域，加强与院系、科研团队和创新创业平台协同。2025年，基地累计对接复旦大学16个院系，新增入库成果21项，推动基地企业依托高校开展科技成果转化12项；深度服务益臻新能源、波达医疗、眸深智能、复泓智研等11家复旦科创衍生企业，在知识产权布局、政策申报、融资对接、企业注册落地等方面提供定制化支持。",
    )
    add_body(
        doc,
        "二是提升企业培育能力。2025年，基地新增入孵创业团队26家，当年孵化成功率100%；新增入孵初创组织23家，年末存活率100%；2025年末在孵创业实体159家。围绕企业成长需求，建立科技型中小企业、高新技术企业、科技领军企业梯度培育机制，全年新增高新技术企业（含复审）21家、杨浦区科技小巨人（含培育）企业2家、“3310”企业3家。",
    )
    add_body(
        doc,
        "三是培育创业典型。2025年，基地共计产生创业评选优秀创业组织9家。眸深智能荣获全国颠覆性技术创新大赛总决赛最高奖及多项赛事奖项；波达医疗在国家级与市级科创赛事中取得优异成绩，并联合复旦大学团队获批上海市2025年度关键技术研发计划项目；复旦大学本科生团队“赤子青山乡村旅游研学共富平台”项目荣获中国国际大学生创新大赛“青年筑梦红色之旅”赛道创业组金奖；3个创业团队获得上海市大学生科技创业基金复旦大学基金支持；音书科技荣获2025年度“天使基金优秀雏鹰企业”称号。",
    )

    add_heading_cn(doc, "（二）创业带动就业", level=2)
    add_body(
        doc,
        "基地通过创业空间支持、专业服务赋能和创新资源链接，发挥创业带动就业作用。2025年末，基地在孵创业组织在岗人数324人，同比增长57.28%；在孵创业组织新增就业118人。基地分层分类开展创新创业教育，通过创业培训、项目路演、企业实训等方式，提高青年人才创新创业能力，为高校毕业生、青年创业者和专业技术人才提供创新创业就业机会。",
    )

    add_heading_cn(doc, "（三）创业政策落实与融资服务", level=2)
    add_body(
        doc,
        "基地积极落实市、区各类创新创业扶持政策，加强政策服务和金融资源链接。2025年，基地新增高新技术企业（含复审）21家，杨浦区科技小巨人（含培育）企业2家、“3310”企业3家；成功推荐9家创业组织享受人社部门创业扶持政策；帮扶2家创业组织获得人社部门创业担保贷款支持。",
    )
    add_body(
        doc,
        "在科技金融方面，基地与山东工研基金、禹泽资本、小村创投、创业接力基金、杨浦科创集团等投资机构建立合作机制，与上海农商行、招商银行、中国银行等金融机构开展融资服务合作；参与管理上海市大学生科技创业基金复旦分基金，联动复旦科创母基金及五支前孵化专项资金。2025年，基地帮扶创业组织获得投融资支持累计约9700余万元。",
    )

    add_heading_cn(doc, "（四）创业孵化服务", level=2)
    add_body(
        doc,
        "一是强化载体空间支持。复旦科技园综合更新与功能提升项目成功获批张江专项发展资金重点项目。基地场地面积3.41万平方米，其中孵化面积2.56万平方米，占比75%。2025年，基地共计向49家创业组织提供租金减免等场地支持。",
    )
    add_body(
        doc,
        "二是强化专业服务供给。现有签约入驻第三方服务机构41个，全年提供服务超过2000家次；落实创业帮扶指标40家。复煜概念验证中心于2025年8月投入运营，围绕新材料、人工智能、集成电路等领域开展首批项目征集，从60余个项目中筛选出4个项目进入验证阶段，提供资金、载体、应用场景、科研支撑、投融资对接、运营协助等全方位服务。",
    )
    add_body(
        doc,
        "三是强化创业导师支持。2025年新增创业导师16人，形成由高校教师、企业管理者、投资机构专家等组成的多元化创业辅导力量，为创业组织提供技术、管理、资本、市场等方面支持。",
    )

    add_heading_cn(doc, "（五）创新创业活动组织", level=2)
    add_body(
        doc,
        "基地以活动为载体，完善“复・创课堂”“复旦科技园企业沙龙”等双创课程体系；顺利举办上海市创业培训“马兰花计划”复旦大学创业意识激发期（GYB）培训班；联合复旦大学创新创业学院开发“理论＋实践”创新创业课程。同时组织开展人工智能产业应用路演、新材料研发与产业应用研讨会、“无掩膜光刻应用于微纳器件加工”产学研沙龙等系列活动，并依托“复旦科创大赛”“复旦之星”“创·在上海”“杨浦‘科创之星’”等品牌赛事以赛引才。",
    )
    add_body(
        doc,
        "2025年，基地组织举办创新创业培训、创业沙龙、项目路演、企业实训等创新创业活动47场，其中3场与人社部门合作开展；推荐各类创新创业赛事参赛项目100余项；推荐创业组织和项目参加人社类活动/赛事12项。",
    )

    add_heading_cn(doc, "（六）创业媒体宣传", level=2)
    add_body(
        doc,
        "基地积极参与复旦科技园品牌焕新工程，协助完成园区官方网站、微信公众号等宣传阵地升级迭代。重点围绕眸深智能、波达医疗等科技企业成长案例及创新创业赛事优秀项目加强典型宣传。2025年，基地通过官方网站、微信公众号发布创新创业信息稿34篇；向人社部门供稿11篇，其中被“海纳百创”“乐业杨浦”等宣传媒介采纳10篇。",
    )

    # ------------------------------------------------------------------
    add_heading_cn(doc, "三、与人社部门合作紧密度", level=1)
    add_body(
        doc,
        "基地持续深化与市、区人社部门及就业促进机构的协同联动，把政策宣传、活动联办、见习实训、就业促进纳入日常孵化服务体系，提升人社创业扶持政策在高校创业场景中的落地效率。",
    )
    add_heading_cn(doc, "（一）政策宣传与对接落实", level=2)
    add_body(
        doc,
        "基地围绕科技企业成长需求，提供高新技术企业认定、科技项目申报、“3310”企业申报、人社扶持申报等政策辅导服务。2025年，成功推荐9家创业组织享受人社部门创业扶持政策，帮扶2家创业组织获得人社部门创业担保贷款支持，推动创业启动、场地、社保、融资等政策精准惠及在孵组织。",
    )
    add_heading_cn(doc, "（二）联合开展创业活动与培训", level=2)
    add_body(
        doc,
        "2025年，基地与人社部门合作开展创新创业活动3场，并推荐创业组织和项目参加人社类活动/赛事12项。上海市创业培训“马兰花计划”复旦大学创业意识激发期（GYB）培训班顺利举办，强化创业意识启发与项目方向筛选。基地持续发挥创业见习与就业促进功能，通过校园招聘、企业实训和岗位对接，帮助青年人才进入创新创业生态。",
    )
    add_heading_cn(doc, "（三）宣传供稿与品牌协同", level=2)
    add_body(
        doc,
        "基地主动向人社部门供稿11篇，其中10篇被“海纳百创”“乐业杨浦”等媒介采纳，形成政策落地、活动联办、宣传联动的闭环协作机制。",
    )

    # ------------------------------------------------------------------
    add_heading_cn(doc, "四、特色与亮点", level=1)
    add_body(
        doc,
        "复旦科技园创业孵化基地立足国家大学科技园功能定位，聚焦高校科技成果转化“最初一公里”，以成果转化为主线、以创业孵化为核心、以生态建设为支撑，逐步形成具有高校特色的成果转化型创业孵化模式。",
    )

    add_heading_cn(doc, "（一）打造特色品牌，持续提升创业孵化影响力", level=2)
    add_body(
        doc,
        "基地依托复旦大学学科、科研和人才优势，不断完善覆盖创业辅导、赛事辅导、技术转移、政策咨询、知识产权、法律财税、投融资对接、人才引育、市场拓展等的一站式服务网络。现有签约入驻第三方服务机构41家，全年提供专业服务超过2000家次；2025年新增签约创业导师16人，组织举办创新创业活动47场。",
    )
    add_body(doc, "活动案例：复旦科创大赛", bold=True, first_indent=True)
    add_body(
        doc,
        "复旦科创大赛由复旦大学、虹口区人民政府、锦江国际集团共同指导，复旦科技园创业孵化基地是主要承办单位之一，承担赛事组织、项目征集遴选、创业辅导培训、评审组织协调、资源对接及赛后孵化培育等工作。大赛吸引全国456个项目报名，涵盖信息技术、集成电路、生命健康、新材料等重点领域，20支团队进入总决赛。创意组突出基础研究与前沿探索，创业组聚焦场景转化与商业落地，多个项目已具备产业化条件，并与投资机构、高校转化平台达成初步合作。",
    )

    add_heading_cn(doc, "（二）创新服务模式，构建高校成果转化型创业孵化体系", level=2)
    add_body(
        doc,
        "基地坚持成果转化前移，形成“成果发现－概念验证－创业孵化－产业培育”全链条服务路径。2025年累计对接复旦大学16个院系，新增入库科技成果21项，推动基地企业依托高校开展科技成果转化12项。复煜概念验证中心正式投入运营并遴选4个项目进入验证阶段，有效提升早期科技成果承接能力。",
    )
    add_body(doc, "孵化案例：波达医疗·推动超快超声技术产业化", bold=True, first_indent=True)
    add_body(
        doc,
        "上海波达医疗科技有限公司由复旦大学信息学院教师团队创办，于2022年入驻基地，聚焦超快超声成像技术成果产业化。基地为企业提供创业辅导、政策申报、资源对接、产业协同等全周期服务，协助申报杨浦区“3310”计划B类项目，推动企业在海聚英才大赛中荣获二等奖，并帮助其入驻复旦科技园湾谷园区、争取租金减免与社保补贴、对接思南基金等投资资源，助力企业获得飞图创投数千万元融资，公司估值突破亿元，实现从高校科研成果向产业化企业的快速成长。",
    )

    add_heading_cn(doc, "（三）深化区域协同合作，构建开放融合创新生态", level=2)
    add_body(
        doc,
        "作为复旦大学与杨浦区政府校地合作的重要载体，基地积极推动学校科研成果、人才资源与区域产业需求精准对接。2025年探索通过“财政引导资金＋社会资本接续＋产业协同落地”的拨投结合模式，帮助复旦成果转化企业益臻新能源成功对接杨浦区科技成果转化支持政策；基地企业实现区级税收贡献1.37亿元。",
    )
    add_body(
        doc,
        "基地与祖泉研究院、复旦大学技术转移中心、同济大学技术转移中心、张江磁谷、上海国际化工创新中心、迈科技、国科新研等创新主体建立长期合作关系；参与发起“北欧创新国际会客厅”，举办“丝路华章·大使领创计划”出海沙龙、“复创·大咖说”等活动，接待巴西、俄罗斯等国际参访团，持续提升开放合作水平和国际影响力。",
    )

    # ------------------------------------------------------------------
    add_heading_cn(doc, "五、综合效益与下一步工作", level=1)
    add_heading_cn(doc, "（一）综合效益", level=2)
    add_kpi_line(doc, "经济效益", "2025年新增注册企业190家（科技型企业143家）；新增高新技术企业（含复审）21家、杨浦区科技/双创“小巨人”企业2家、“3310”企业3家；企业新增知识产权318项（发明专利42项）；区级税收贡献1.37亿元。")
    add_kpi_line(doc, "社会效益", "推动科技成果转化12项；新增创业导师16人；举办创新创业培训11场、创新创业活动37场；推荐赛事项目100余项。复旦科技园在上海市创新创业载体绩效评价中获评优良，复翼互联众创空间获评优秀；在杨浦区科技园区考评中获评优秀。")
    add_kpi_line(doc, "绿色发展", "完成景观照明提升、公共空间改造、设施设备更新等项目；综合更新与功能提升项目获批张江专项；复旦科技园大厦通过上海市“无废城市细胞”建设评估，获评区级“无废楼宇”。")

    add_heading_cn(doc, "（二）下一步工作打算", level=2)
    add_body(
        doc,
        "下一阶段，基地将继续发挥国家大学科技园功能优势，围绕成果转化“最初一公里”深化概念验证与早期孵化，强化与人社部门在政策落实、创业培训、就业促进方面的协同，持续提升创业孵化服务能力和区域贡献度，为更多高校科技成果转化和青年创新创业提供支撑。",
    )
    add_body(
        doc,
        "未来，复旦科技园创业孵化基地将深化高校、政府、产业、资本、社会多方协同，持续提升创业孵化服务能力和成果转化承载能力，为区域创新发展和产业升级贡献更大力量。",
    )

    doc.save(OUT)
    print(f"已生成：{OUT}")


if __name__ == "__main__":
    build()
