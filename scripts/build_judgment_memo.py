#!/usr/bin/env python3
"""生成《2026.9.21 首尔 Finwise 中韩论坛赞助合作判断备忘录》。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output"
FONT_CN = "WenQuanYi Micro Hei"
FONT_EN = "Calibri"
NAVY = RGBColor(0x1B, 0x3A, 0x5F)
RED = RGBColor(0xB4, 0x23, 0x18)
ORANGE = RGBColor(0xB8, 0x5C, 0x00)
GREEN = RGBColor(0x1F, 0x6B, 0x3A)
GRAY = RGBColor(0x4A, 0x4A, 0x4A)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)


def set_run_font(run, size=11, bold=False, color=BLACK, font=FONT_CN):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = FONT_EN
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:cs"), FONT_EN)


def shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell, color="D0D5DD") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, *, size=10, bold=False, color=BLACK, align="left", fill=None):
    if fill:
        shade_cell(cell, fill)
    set_cell_border(cell)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_paragraph(doc, text, *, size=11, bold=False, color=BLACK, space_after=8, first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.35
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading_cn(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0)
    size = 16 if level == 1 else 13
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=NAVY)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "1B3A5F")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_bullet(doc, text, *, bold_prefix=None, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run("• ")
    set_run_font(run, size=11, bold=True, color=NAVY)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True, color=color)
        r2 = p.add_run(text)
        set_run_font(r2, size=11, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=color)
    return p


def add_callout(doc, title, body, fill="FDECEC", title_color=RED):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, "E4B4B4" if fill.startswith("FD") else "C5D5C8")
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(title)
    set_run_font(r1, size=12, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.line_spacing = 1.3
    r2 = p2.add_run(body)
    set_run_font(r2, size=11, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_table(doc, headers, rows, col_widths=None, header_fill="1B3A5F"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        set_cell_text(
            table.cell(0, i),
            h,
            size=9,
            bold=True,
            color=RGBColor(0xFF, 0xFF, 0xFF),
            align="center",
            fill=header_fill,
        )
    for r_idx, row in enumerate(rows, start=1):
        fill = "F7F9FC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            align = "center" if c_idx > 0 else "left"
            set_cell_text(table.cell(r_idx, c_idx), val, size=9, align=align, fill=fill)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return table


def build() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # 页眉
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("内部判断备忘录  ·  仅供决策使用  ·  2026-08-22")
    set_run_font(hr, size=8, color=GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("结论基于公开材料交叉核对。嘉宾与主办背书均须书面核实后再对外引用。")
    set_run_font(fr, size=8, color=GRAY)

    # 标题
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(4)
    tr = t.add_run("要不要帮他拉赞助？")
    set_run_font(tr, size=22, bold=True, color=NAVY)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(2)
    sr = st.add_run("2026.9.21 首尔 Finwise 中韩火炬产业投资论坛")
    set_run_font(sr, size=14, bold=True, color=NAVY)

    st2 = doc.add_paragraph()
    st2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st2.paragraph_format.space_after = Pt(12)
    sr2 = st2.add_run("赞助合作可行性与报价判断备忘录")
    set_run_font(sr2, size=12, color=GRAY)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(14)
    mr = meta.add_run(
        "拟合作名义：复旦大学住房政策研究中心  /  上海市杨浦区科技企业联合会\n"
        "对方条件：现金赞助以美元（或人民币兑换美元）呈现，抽成 20%–30%"
    )
    set_run_font(mr, size=10, color=GRAY)

    add_callout(
        doc,
        "一句话结论：现有条件，不建议答应。",
        "票价 18,888 元偏高但还卖得动给少数人；赞助官价 3–10 万 USDT 对实体企业基本收不到。"
        "20%–30% 抽成看起来诱人，本质是把卖不动的标价风险转给你。"
        "用复旦住房政策研究中心和杨浦科企联的名义去推 USDT 标价、晚宴还带 Token 经济专场，机构风险远大于佣金。"
        "若人情上必须给回应：只能做“有书面背书、人民币对公、佣金按实收、剔除虚拟货币专场、价格腰斩”的试水，且不要用高校名义。",
    )

    add_heading_cn(doc, "一、先回答三个问题", 1)

    add_table(
        doc,
        ["问题", "判断", "把握"],
        [
            [
                "要不要按对方条件答应？",
                "不答应。不是“先做做看”，是条件本身不成立。",
                "高",
            ],
            [
                "18,888 元贵不贵？",
                "对真实政企团偏贵；标价 59,800 是虚锚。心理价 0.8–1.2 万（不含机票）。",
                "高",
            ],
            [
                "3–10 万 USDT 赞助收不收得到？",
                "按官价几乎收不到。钻石 10 万美金对本场规格是幻想价。",
                "高",
            ],
            [
                "20%–30% 抽成值不值？",
                "费率偏高，说明对方自己也卖不动。按实收 15% 且到账后付，才谈得上。",
                "高",
            ],
            [
                "用两个官方名义合不合适？",
                "不合适。这是否决项，不是谈判细节。",
                "高",
            ],
        ],
        col_widths=[4.2, 9.6, 2.4],
    )

    add_heading_cn(doc, "二、先看清这是什么活动", 1)
    add_paragraph(
        doc,
        "对外包装是“工信部火炬中心、驻韩大使馆主办的中韩政企投资论坛”。对内看产品结构，是两场气质相反的活动被缝在同一天：下午卖上海/韩国政企背书，晚上卖 Token 经济与加密圈意见领袖。",
    )
    add_paragraph(
        doc,
        "承办方 Finwex / 纷睿的公开主业是 AI+Web3、数据资产发行、Yomirgo（$YGO）和 IFIC/FINWISE 系列峰会，创始人尤立长期在区块链峰会圈层活动。宣传落点不在市政府或大使馆官网，而在 PANews、Odaily、BlockBeats 这类加密媒体。活动行报名主体写的是 Finwex AI，不是火炬中心，也不是大使馆。",
    )

    add_heading_cn(doc, "公开口径对不上，这是最大的销售风险", 2)
    add_bullet(doc, "PANews / Odaily：工信部火炬中心、驻韩大使馆主办，Finwex AI 承办。", bold_prefix="主办身份：")
    add_bullet(doc, "BlockBeats：中韩火炬创新园、Finwex 主办；嘉宾栏目前只挂出戴兵、朴允圭、尤立。", bold_prefix="另一套说法：")
    add_bullet(doc, "指导单位写火炬中心和驻韩大使馆，主办栏被裁切；银牌标价 30,000 USDT。", bold_prefix="PDF 手册：")
    add_bullet(doc, "组织者是 Finwex AI。早鸟截止，宣传稿写 8 月 15 日，售票页写 9 月 1 日。", bold_prefix="活动行：")
    add_bullet(doc, "你拿到的议程有副市长卢山、临港吴晓华、得物、同传主持周岚；PDF 日程结构和人名并不一致，且多处“待定”。", bold_prefix="议程：")
    add_paragraph(
        doc,
        "截至 2026 年 8 月 22 日，中国驻韩国大使馆官网、工信部/火炬中心公开渠道，检索不到这场 9 月 21 日首尔论坛的主办公告。同主题、同规格的真正官方活动，通常会有使馆活动稿或部门新闻。现在只有商业媒体通稿，说明“国家级主办”尚未被公开背书，对外招商时不能当既成事实来用。",
        first_line=True,
    )

    add_heading_cn(doc, "下午政企、晚上 Token，买家是两群合不来的人", 2)
    add_paragraph(
        doc,
        "下午可以讲临港、得物、西井、XREAL、智元机器人；晚上圆桌是 Midu、韩锋、Eric Gu、郭宏才、Mikey Yeung（Yomirgo，与承办方生态相关）。中国实体企业和区级联合会能介绍的会员，看到 Token、USDT、加密 KOL，会后退。加密项目愿意为“和副市长同台”付费，但副市长、大使、市商务委领导通常不会和虚拟货币专场出现在同一份对外议程里。结果是：两头的核心卖点，都可能在临场被抽掉。",
    )
    add_paragraph(
        doc,
        "手册还把 SouLink 写成 2026 年成立的新项目，并安排“视频生态 OS 战略发布”。若招商对象随后发现发布方、路演项目与承办方/中间人存在关联，会认为自己买的不是中立平台，而是项目方自办发布会。这会直接摧毁溢价。",
    )

    add_heading_cn(doc, "三、18,888 元贵不贵？", 1)
    add_paragraph(
        doc,
        "先把产品拆开。活动行在售三档：下午活动票早鸟 188 元（原价 388）；论坛+投资晚宴 VIP 早鸟 2,888 元（原价 8,888）；三天参访团早鸟 18,888 元（原价 59,800）。参访团含韩国当地住宿、餐饮、交通，不含机票和签证。限 20 席，单笔最多 8 张。",
    )

    add_table(
        doc,
        ["对照项", "市场价格带", "本场报价", "评价"],
        [
            [
                "上海 AI 大会论坛票",
                "约 300–800 元",
                "VIP 2,888 元",
                "偏贵，但闭门晚宴可部分解释",
            ],
            [
                "协会组织韩国 4 天商务考察",
                "约 0.8–1.5 万（含或不含机票）",
                "1.8888 万不含机票",
                "加上机票约 2.2–2.6 万，属高价团",
            ],
            [
                "含国会/财阀接待的高端团",
                "约 1.5–3 万，兑现才站得住",
                "原价 5.98 万",
                "原价是虚锚，不能当价值证明",
            ],
            [
                "机票+签证（需自理）",
                "约 2,500–5,000 元",
                "未包含",
                "真实总成本明显高于票面",
            ],
        ],
        col_widths=[4.0, 4.6, 3.6, 4.0],
    )

    add_paragraph(
        doc,
        "结论：18,888 元不是“便宜到该抢”，也不是完全卖不掉。它是高价商务团。只有在“国会、NIPA、LG/NAVER 是对口部门负责人座谈，而不是参观走廊和展览馆”的前提下，这个价才对得上。原价 59,800、立减 40,912，是会务销售的经典高开低走，不能拿来对会员说“已经打两折了”。宣传写限 20 席，下单却允许 8 张，说明稀缺是话术，不是刚性约束。",
    )
    add_paragraph(
        doc,
        "对杨浦科技企业联合会会员：能接受的成交带更接近 8,000–12,000 元（不含机票）。对需要“出访照片、签约背景板、投资晚宴”做对外叙事的项目，18,888 有少数人会买，但不要指望批量。对复旦住房政策研究中心的学术网络：这个价格和主题都不匹配，不该开口。",
    )

    add_heading_cn(doc, "四、赞助官价能不能收到？", 1)
    add_paragraph(
        doc,
        "PDF 报价以 USDT 计价：钻石 100,000、白金 80,000、金牌 40,000、银牌 30,000。按 1 美元约 7.2 元人民币，大约是 72 万 / 58 万 / 29 万 / 22 万人民币。你提供的权益表有一版银牌写成 20,000 USDT，与 PDF 不一致——连价目表都还没锁死，说明这是招商清单，不是已售套餐。",
    )

    add_table(
        doc,
        ["档位", "官价 USDT", "约合人民币", "核心权益", "按官价成交可能"],
        [
            ["钻石", "100,000", "约 72 万", "下午 15 分钟演讲、晚宴冠名、8 席、参访 2 席", "极低"],
            ["白金", "80,000", "约 58 万", "晚宴 15 分钟演讲、联合主办、5 席、参访 2 席", "极低"],
            ["金牌", "40,000", "约 29 万", "下午 15 分钟演讲、3 席、参访 1 席", "低，且多半要打折"],
            ["银牌", "30,000", "约 22 万", "晚宴 15 分钟演讲、2 席、无参访团", "低"],
        ],
        col_widths=[2.4, 2.6, 2.8, 6.2, 2.2],
    )

    add_paragraph(
        doc,
        "为什么贵？因为权益是按“国家级双边论坛”标的，活动体量却是酒店半日会+晚宴+小团出访。金牌 4 万美金只附赠 1 个参访席（零售约 2,600 美金），其余都是 Logo、展位和 15 分钟演讲。真正能支撑 10 万美金演讲费的，是 Token2049 那种上万人、全球转播的加密年会（其白金赞助公开案例约 40 万美金），或者 WAIC 这种国家级展会。本场都不是。",
    )
    add_paragraph(
        doc,
        "上海普通产业展会/垂直论坛的金牌常见报价是 10–30 万人民币，银牌 8–20 万。学术或行业年会钻石往往只有 8–20 万人民币。把本场银牌标到 22 万人民币以上，已经对标国内中型展会金牌，但现场曝光、媒体和观众质量明显弱于那些展会。",
    )

    add_heading_cn(doc, "分客群看，谁可能付钱", 2)
    add_bullet(
        doc,
        "杨浦、临港、张江的制造业、硬科技、传统VC，对标的是人民币发票、政府活动、合规传播。USDT 报价本身就会让对方停下。按官价成交可能接近零。",
        bold_prefix="实体企业：",
    )
    add_bullet(
        doc,
        "愿意为“政府背书照片+亚洲 Token 叙事”付营销费。可能谈的是 1–3 万美金，而不是 8–10 万。他们要的是同台，不是联合会介绍信。",
        bold_prefix="海外加密项目：",
    )
    add_bullet(
        doc,
        "需要路演和晚宴座位，更可能要免费票或打折票，而不是 20 万人民币赞助。",
        bold_prefix="被点名的独角兽：",
    )
    add_paragraph(
        doc,
        "综合判断：按官价收到钻石/白金的概率很低；收到 1 单银牌或金牌的前提，是大幅议价，且买家来自加密营销预算，不是来自科企联会员。你若用两个体制内名义去找后者，会把机构信用花在最不匹配的商品上。",
    )

    add_heading_cn(doc, "五、合适的价格带（建议你心里用这张表）", 1)
    add_paragraph(
        doc,
        "下面不是去跟对方“砍价砍到他亏本”，而是你对外报价时自己要守的成交带。超过上沿，就按卖不动处理，不要为了冲佣金硬推。人民币按 7.2 折算，仅作对照。",
    )

    add_table(
        doc,
        ["产品", "对方官价", "对实体企业可成交", "对加密项目可成交", "建议你开口"],
        [
            ["下午票", "188 / 原价 388", "赠票或 188", "不敏感", "不要当招商主体"],
            ["论坛+晚宴 VIP", "2,888 / 原价 8,888", "1,200–2,000", "2,000–2,888", "2,000 左右"],
            ["三天参访团", "18,888 / 原价 59,800", "8,000–12,000 不含机票", "12,000–18,888", "1.2 万，含机票另说"],
            ["银牌赞助", "2–3 万 USDT", "8–12 万人民币", "0.8–1.5 万美金", "10 万人民币 / 1.2 万美金"],
            ["金牌赞助", "4 万 USDT", "15–25 万人民币", "1.5–2.5 万美金", "18 万人民币 / 2 万美金"],
            ["白金赞助", "8 万 USDT", "30–40 万，且副市长书面确认", "2.5–4 万美金", "先不要卖这一档"],
            ["钻石赞助", "10 万 USDT", "基本无", "4–6 万美金且要冠名兑现", "不做"],
        ],
        col_widths=[3.0, 3.2, 3.6, 3.4, 3.0],
    )

    add_callout(
        doc,
        "佣金怎么算才合理",
        "行业里成熟展会的代理费常见 8%–15%。对方给 20%–30%，不是因为你特别值，是因为标价虚高、临近活动、需要你用机构牌子去背书。"
        "只接受“实际到账金额”的 15%–20%，人民币对公支付，款到后 7 个工作日内结佣。"
        "不要接受按官价计提、活动后分账、打 USDT 到个人钱包、用门票或股权抵佣金。"
        "按官价 10 万美金抽 30% 是幻影收入；按实收 1.2 万美金抽 20%，才是可能拿到手的钱。",
        fill="FFF6E8",
        title_color=ORANGE,
    )

    add_heading_cn(doc, "六、用两个名义合作，为什么是否决项", 1)
    add_paragraph(
        doc,
        "这不是“风险提示一下仍可做”，而是做了会把你自己和机构放在说不清的位置上。",
    )
    add_bullet(
        doc,
        "这是住房政策学术机构，公开宗旨是学术性、中立性、公益性。用它的名字去卖韩国 AI 论坛赞助和 Token 晚宴，业务上完全不对口。没有学校或中心书面授权，就是擅自使用校名。一旦赞助出纠纷、嘉宾缺席或涉及虚拟货币舆论，问责会落到中心和你个人。",
        bold_prefix="复旦大学住房政策研究中心：",
    )
    add_bullet(
        doc,
        "介绍会员参加合规的产业考察，勉强说得通，但必须理事会或秘书处书面同意，钱走对公，发票、合同、宣传口径受控。不能由个人按 20%–30% 抽成。更不能把联合会 Logo 用在 USDT 报价和 Token 经济专场上。国内对虚拟货币宣传的监管红线并没有因为“AI”三个字消失。",
        bold_prefix="杨浦区科技企业联合会：",
    )
    add_bullet(
        doc,
        "现金以美元呈现、报价单位是 USDT，对高校和社团是外汇与收费合规问题。个人收稳定币再“帮忙拉赞助”，外观上像地下佣金。",
        bold_prefix="收款币种：",
    )
    add_bullet(
        doc,
        "议程或手册里若已出现你或你所在项目（例如路演、发布、联合创始人），再以高校/联合会中立平台名义招商，属于关联交易外观。被赞助方事后核对，会认为被误导。",
        bold_prefix="关联关系：",
    )
    add_bullet(
        doc,
        "副市长、大使、火炬中心、得物负责人目前都还停留在“将出席/待定”。你用机构名义发出去的微信，法律上近似代为陈述。人没来，退款和声誉由介绍方先扛。",
        bold_prefix="嘉宾未锁定：",
    )
    add_paragraph(
        doc,
        "一句话：个人身份帮朋友问问，和机构身份背书卖赞助，是两件事。对方要的其实是后者。前者佣金有限；后者你付不起。",
    )

    add_heading_cn(doc, "七、三种做法，只建议后两种里的一种", 1)

    add_heading_cn(doc, "方案 A：拒绝招商（推荐）", 2)
    add_paragraph(
        doc,
        "礼貌肯定活动方向，明确不能用两个机构名义做赞助代理，也不接受 USDT 结算的分成。若对方确有政企考察的成色，可以个人或企业身份考虑是否买票核实，但不承担销售指标。这是对机构、对会员、对你自己最干净的做法。",
        first_line=True,
    )

    add_heading_cn(doc, "方案 B：有条件试水（仅在人情必须给台阶时）", 2)
    add_paragraph(doc, "全部满足再开口，缺一条就停：", first_line=True)
    add_bullet(doc, "火炬中心或驻韩使馆的指导/支持函；上海市领导、韩方领导的书面确认或组委会盖章嘉宾名单，不能是 PPT。")
    add_bullet(doc, "中国大陆可核验公司主体、人民币对公账户、增值税发票、书面赞助合同。不收、不转 USDT。")
    add_bullet(doc, "佣金 15%–20%，基数为实际到账，款到后人民币支付。写进合同，不接受口头。")
    add_bullet(doc, "禁用复旦住房政策研究中心名义。科企联若出面，须内部书面批准，只做 AI 产业对接，宣传材料删除 Token/USDT/加密 KOL。")
    add_bullet(doc, "对外只报上一节的成交带。先尝试 1 家银牌或 2 张参访票，收到钱再决定是否继续。")
    add_bullet(doc, "合同写明：嘉宾降级、参访改为旅游点、晚宴改成项目方自办酒会，赞助方可全额退或按比例降级。")
    add_bullet(doc, "你不在对方招商材料上署名“复旦/联合会独家代理”。介绍人身份最多写个人姓名。")

    add_heading_cn(doc, "方案 C：不要选", 2)
    add_paragraph(
        doc,
        "不核实嘉宾，直接用两个机构的微信名片和公众号去群发；按 10 万美金钻石去跟企业开口；个人钱包收 USDT；承诺“和副市长合影、和 NAVER 高层闭门”；把 SouLink 发布和联合会招商混在一起。这几条任何一条踩上，后面的麻烦都大过佣金。",
        first_line=True,
    )

    add_heading_cn(doc, "八、如果拒绝，可以这样说", 1)
    add_paragraph(
        doc,
        "“方向上我理解你们想做中韩 AI 对接。但我这边能用的两个名义，一个是高校政策研究机构，一个是区级企业联合会，都不适合代理以美元/USDT 计价、含 Token 经济专场的商业赞助。会员企业也更认人民币对公和已确认的政务行程。等嘉宾书面确认、赞助改人民币合同之后，如有纯产业考察名额，我可以帮着问问有没有人愿意自己买票，但不做赞助分成，也不用机构名义背书。”",
        first_line=True,
    )
    add_paragraph(
        doc,
        "如果对方坚持“先帮我找两家钻石/白金”：不要转发他们的价目表。价目表本身会把你锁进虚假锚点。",
        first_line=True,
    )

    add_heading_cn(doc, "九、你现在就可以向对方要的材料", 1)
    add_paragraph(doc, "在说“帮不帮”之前，先要这 8 样。要不到，答案就是不帮。", first_line=True)
    add_table(
        doc,
        ["序号", "要什么", "为什么要"],
        [
            ["1", "主办/指导单位公函或授权书", "核验“火炬”“大使馆”是指导还是借用"],
            ["2", "已确认嘉宾的书面名单（含职务）", "副市长、大使不能停留在“将出席”"],
            ["3", "中国签约主体营业执照与开票信息", "不能是个人或海外钱包"],
            ["4", "人民币赞助合同模板+退款条款", "嘉宾变更如何处理"],
            ["5", "佣金补充协议（按实收、人民币、节点）", "避免活动后空转"],
            ["6", "9 月 22–23 日参访的对方确认函", "国会、NIPA、LG、NAVER 分别是谁接待"],
            ["7", "已售票和已签约赞助的匿名统计", "判断是热销还是仅有价目表"],
            ["8", "晚宴完整议程（是否保留 Token 专场）", "决定联合会能不能沾边"],
        ],
        col_widths=[1.6, 7.4, 7.2],
    )

    add_heading_cn(doc, "十、最终判断", 1)
    add_paragraph(
        doc,
        "这不是“价格高一点、佣金高一点，值不值得冲一把”的问题。这是一场商业峰会在使用尚未被公开文件钉死的政务品牌，报价按加密年会的上沿来标，销售对象却希望你用上海高校和区级联合会去找实体企业。中间的错配，就是你的风险。",
        first_line=True,
    )
    add_paragraph(
        doc,
        "18,888 元：高，但不是笑话价；不要用 59,800 证明它便宜。合适的会员价在 8,000–12,000 元（不含机票）。",
        first_line=True,
    )
    add_paragraph(
        doc,
        "赞助：官价 3–10 万 USDT 收不到。心里按银牌约 10 万人民币、金牌约 18 万人民币来，白金以上先不要卖。20%–30% 抽成只在“实收到账、人民币、书面”时才有意义，否则是纸面数字。",
        first_line=True,
    )
    add_paragraph(
        doc,
        "要不要答应：不答应现有条件。复旦住房政策研究中心的名义不要用。杨浦科企联除非走内部决策、剔除虚拟货币内容、人民币对公，否则也不要用。人情需要台阶，就按方案 B 要材料；材料不齐，就停。",
        first_line=True,
    )

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(16)
    nr = note.add_run(
        "方法说明：本备忘录交叉核对了活动 PDF《FINWISE AI SUMMIT SEOUL》、活动行售票页、PANews/Odaily/BlockBeats 通稿、Finwex 官网，以及大使馆/政府部门公开检索结果。"
        "价格对照取自国内产业展会、学术会议公开赞助档和韩国商务考察团常见收费带。"
        "嘉宾与主办关系以书面文件为准，不以宣传稿为准。"
    )
    set_run_font(nr, size=9, color=GRAY)

    out = OUTPUT_DIR / "2026-09-21-Finwise首尔论坛-赞助合作判断备忘录.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    path = build()
    print(f"已生成：{path}")
    print(f"大小：{path.stat().st_size} bytes")
