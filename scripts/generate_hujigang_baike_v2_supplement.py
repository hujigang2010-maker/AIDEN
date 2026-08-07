#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《胡继刚｜百度百科第二版增补稿》Word 文档。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "胡继刚_百度百科第二版增补稿.docx"


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def add_para(
    doc,
    text,
    *,
    size=12,
    bold=False,
    space_before=0,
    space_after=6,
    first_line=False,
    color=None,
    align=None,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading_custom(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12}
    colors = {
        1: RGBColor(0x1A, 0x3A, 0x5C),
        2: RGBColor(0x1A, 0x3A, 0x5C),
        3: RGBColor(0x33, 0x33, 0x33),
    }
    return add_para(
        doc,
        text,
        size=sizes.get(level, 12),
        bold=True,
        space_before=12 if level == 1 else 8,
        space_after=6,
        color=colors.get(level),
    )


def add_mixed_para(doc, segments, *, size=12, space_before=0, space_after=6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for text, bold in segments:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
    return p


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    add_para(
        doc,
        "胡继刚",
        size=22,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
        color=RGBColor(0x1A, 0x3A, 0x5C),
    )
    add_para(
        doc,
        "百度百科第二版增补稿（可粘贴）",
        size=14,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_para(
        doc,
        "适用范围：第一版已通过后的小幅增补。本版只做三件事：①新增杨浦区政府官网交叉来源；②补强2025年经历中“组织方代表”表述；③配图说明。不新增企业履历、职称、获奖与评价性词语。",
        size=10.5,
        space_after=10,
        color=RGBColor(0x55, 0x55, 0x55),
    )

    # 一、相对第一版改了什么
    add_heading_custom(doc, "一、相对第一版改了什么", level=1)
    changes = [
        "人物经历·2025年全球新经济增长引擎峰会：在原有圆桌交流表述后，补半句“并以主要组织方代表身份就峰会主题及住房政策研究中心后续工作方向发表观点”，脚注仍用中新网上海〔3〕（该文后半段已有原文，无需新开栏目）。",
        "人物经历·2026年人工智能峰会：脚注新增杨浦区政府官网〔9〕，与央广网〔1〕、中新网〔6〕、上观新闻〔7〕交叉印证“主持圆桌”及执行会长身份。",
        "社会职务·执行会长：脚注增加〔9〕。",
        "配图：建议上传1张正式照；可选再加1张活动照。正文结构不变。",
        "不改：概述主句、基本信息栏、社会职务条目本身、不单列社会活动、不写宝龙/万科等企业履历。",
    ]
    for i, s in enumerate(changes, 1):
        add_para(doc, f"{i}. {s}", size=10.5, space_after=3)

    # 二、可粘贴正文（第二版）
    add_heading_custom(doc, "二、可直接粘贴正文（第二版）", level=1)

    add_heading_custom(doc, "概述（可保持第一版不变）", level=2)
    add_para(
        doc,
        "胡继刚，复旦大学住房政策研究中心秘书长、上海市杨浦区科技企业联合会执行会长。曾参与不动产资产管理、新经济增长、人工智能商业化及中欧创新合作等主题的行业交流活动。〔1〕〔2〕〔3〕〔4〕〔5〕〔7〕〔8〕〔9〕",
        first_line=True,
        space_after=8,
    )

    add_heading_custom(doc, "基本信息（建议仍只保留三项）", level=2)
    add_para(doc, "中文名：胡继刚", space_after=2)
    add_para(doc, "主要职务：复旦大学住房政策研究中心秘书长　〔1〕〔2〕〔3〕〔5〕〔7〕〔8〕〔9〕", space_after=2)
    add_para(doc, "社会职务：上海市杨浦区科技企业联合会执行会长　〔1〕〔4〕〔7〕〔9〕", space_after=8)

    add_heading_custom(doc, "人物经历（整段替换或按标注改）", level=2)

    add_mixed_para(
        doc,
        [("【保持】", True), (" 2024年6月段：与第一版相同。", False)],
        size=10.5,
        space_after=2,
    )
    add_para(
        doc,
        "2024年6月，胡继刚以复旦大学住房政策研究中心兼复旦MBA不动产资产管理协会秘书长身份，与上海市锦天城律师事务所高级合伙人顾晓共同主持“他山之石——中国投资者的海外不动产战略布局”主题分享活动。〔5〕",
        first_line=True,
        space_after=6,
    )

    add_mixed_para(
        doc,
        [("【本版修改】", True), (" 2025年峰会段：补“主要组织方代表”半句。", False)],
        size=10.5,
        space_after=2,
    )
    add_para(
        doc,
        "2025年5月，胡继刚以复旦大学住房政策研究中心秘书长身份参加“2025全球新经济增长引擎峰会”，在圆桌对话环节与证券、经济及资产管理领域嘉宾围绕房地产市场、资本布局与产业创新等议题展开交流，并以主要组织方代表身份就峰会主题及住房政策研究中心后续工作方向发表观点。〔2〕〔3〕",
        first_line=True,
        space_after=4,
    )
    add_para(
        doc,
        "改动说明：半句依据中新网上海原文“作为此次活动的主要组织方代表，复旦大学住房政策研究中心秘书长胡继刚表示……”；勿据此单开“主要研究方向”信息栏。",
        size=10.5,
        color=RGBColor(0x55, 0x55, 0x55),
        space_after=6,
    )

    add_mixed_para(
        doc,
        [("【保持】", True), (" 2025年5月28日工商联段：与第一版相同。", False)],
        size=10.5,
        space_after=2,
    )
    add_para(
        doc,
        "2025年5月28日，胡继刚以复旦大学住房政策研究中心秘书长身份出席上海市工商联房地产商会主办的“资产管理分会成立大会暨2025不动产资产管理高质量发展论坛”，并代表该中心参与相关战略合作仪式。〔8〕",
        first_line=True,
        space_after=6,
    )

    add_mixed_para(
        doc,
        [("【保持】", True), (" 2026年3月北欧会客厅段：与第一版相同。", False)],
        size=10.5,
        space_after=2,
    )
    add_para(
        doc,
        "2026年3月，胡继刚以上海市杨浦区科技企业联合会执行会长身份参加“北欧创新国际会客厅”揭牌活动，并围绕中欧创新合作机制建设、科创生态联动与企业国际化发展等议题进行交流。〔4〕",
        first_line=True,
        space_after=6,
    )

    add_mixed_para(
        doc,
        [("【本版增强】", True), (" 2026年AI峰会段：正文可不变，脚注加〔9〕。", False)],
        size=10.5,
        space_after=2,
    )
    add_para(
        doc,
        "2026年5月，胡继刚以复旦大学住房政策研究中心秘书长、上海市杨浦区科技企业联合会执行会长身份参加“2026人工智能商业化落地与硬核投资破局峰会”，主持“从算力引擎到新质资产（AI全产业链的商业化实战）”圆桌对话，并以主办方代表身份发表观点。〔1〕〔6〕〔7〕〔9〕",
        first_line=True,
        space_after=8,
    )

    add_heading_custom(doc, "社会职务", level=2)
    add_para(doc, "复旦大学住房政策研究中心秘书长。〔1〕〔2〕〔3〕〔5〕〔6〕〔7〕〔8〕〔9〕", first_line=True, space_after=3)
    add_para(doc, "复旦MBA不动产资产管理协会秘书长。〔5〕", first_line=True, space_after=3)
    add_para(doc, "上海市杨浦区科技企业联合会执行会长。〔1〕〔4〕〔7〕〔9〕", first_line=True, space_after=8)

    # 三、新增参考资料〔9〕
    add_heading_custom(doc, "三、本版新增参考资料〔9〕", level=1)
    add_mixed_para(
        doc,
        [("〔9〕 ", True), ("AI新质风口到来 共议超级个体成长路径", True)],
        size=11,
        space_before=4,
        space_after=2,
    )
    add_para(doc, "来源：上海市杨浦区人民政府门户网站", size=10.5, space_after=2)
    add_para(doc, "发布时间：2026年5月29日", size=10.5, space_after=2)
    add_para(
        doc,
        "链接：https://www.shyp.gov.cn/shypq/xwzx-bmdt/20260529/506469.html",
        size=10.5,
        space_after=2,
    )
    add_para(
        doc,
        "支持内容：明确记载“复旦大学住房政策研究中心秘书长、区科技企业联合会执行会长胡继刚”主持第二场圆桌论坛，聚焦AI全产业链商业化实战；与央广网、上观新闻交叉印证。",
        size=10.5,
        space_after=2,
    )
    add_para(
        doc,
        "建议插入位置：概述第一句后；人物经历2026年AI峰会段后；执行会长职务后；秘书长职务后（可选）。",
        size=10.5,
        space_after=2,
    )
    add_para(
        doc,
        "原文摘录：第二场圆桌论坛由复旦大学住房政策研究中心秘书长、区科技企业联合会执行会长胡继刚主持，聚焦AI全产业链商业化实战拆解落地方法论。",
        size=10.5,
        color=RGBColor(0x44, 0x44, 0x44),
        space_after=8,
    )

    add_heading_custom(doc, "既有参考资料〔3〕补充用法（不换链接）", level=2)
    add_para(
        doc,
        "〔3〕中新网上海《“创新应变·蓄力远航”2025全球新经济增长引擎峰会在沪举行》除支持圆桌交流外，另支持“主要组织方代表”表述。",
        size=10.5,
        space_after=2,
    )
    add_para(
        doc,
        "原文摘录：作为此次活动的主要组织方代表，复旦大学住房政策研究中心秘书长胡继刚表示，此次峰会“创新应变·蓄力远航”的主题，正是基于全球产业链重构与科技革命的深度交汇……未来，复旦大学住房政策研究中心的重点工作也将不局限于传统不动产领域，而是着力围绕构建“资产-数据-碳汇”协同创新金融化体系……等全新方向而展开。",
        size=10.5,
        color=RGBColor(0x44, 0x44, 0x44),
        space_after=4,
    )
    add_para(
        doc,
        "注意：百科正文只写“就……发表观点”，不要把“资产-数据-碳汇”等展开成个人研究方向或成就。",
        size=10.5,
        color=RGBColor(0x88, 0x33, 0x00),
        space_after=8,
    )

    # 四、配图
    add_heading_custom(doc, "四、配图说明（建议一并提交）", level=1)
    add_para(doc, "头图（必选优先）", bold=True, size=11, space_after=2)
    add_para(doc, "· 1张正面半身正式照；免冠、清晰、背景简洁、无商业水印、无二维码。", size=10.5, space_after=2)
    add_para(doc, "· 图注：胡继刚", size=10.5, space_after=2)
    add_para(doc, "· 来源：本人提供（或活动方供图，须有使用权）", size=10.5, space_after=6)

    add_para(doc, "活动图（可选）", bold=True, size=11, space_after=2)
    add_para(doc, "· 1张能辨认本人的公开活动现场照（如2026年AI峰会主持/圆桌）。", size=10.5, space_after=2)
    add_para(doc, "· 图注：胡继刚在2026人工智能商业化落地与硬核投资破局峰会现场", size=10.5, space_after=2)
    add_para(doc, "· 来源：可挂央广网〔1〕、上观新闻〔7〕或杨浦区政府〔9〕相关报道；或本人/活动方供图。", size=10.5, space_after=6)

    add_para(
        doc,
        "暂不上传：生活照、大合影（本人难辨认）、名片/二维码图、带明显商业水印图。",
        size=10.5,
        color=RGBColor(0x55, 0x55, 0x55),
        space_after=8,
    )

    # 五、提交操作
    add_heading_custom(doc, "五、百科后台操作顺序", level=1)
    steps = [
        "进入已通过的「胡继刚」词条，点击编辑。",
        "将2025年峰会经历段替换为本版【本版修改】全文。",
        "在2026年AI峰会经历段、执行会长职务后新增参考资料〔9〕（杨浦区政府2026-05-29）。",
        "确认〔3〕仍挂在2025年经历段（圆桌＋组织方代表共用一条即可）。",
        "上传头图（及可选活动图），填写图注与来源。",
        "预览核对脚注后提交。提交理由示例：补充政府官网交叉来源、完善2025年公开活动表述，并添加人物照片；不新增未经公开来源支持的履历内容。",
    ]
    for i, s in enumerate(steps, 1):
        add_para(doc, f"{i}. {s}", size=11, space_after=3)

    # 六、纯文本
    add_heading_custom(doc, "六、纯文本粘贴备用", level=1)
    add_para(doc, "【人物经历·仅改动段】", bold=True, size=11, space_after=2)
    add_para(
        doc,
        "2025年5月，胡继刚以复旦大学住房政策研究中心秘书长身份参加“2025全球新经济增长引擎峰会”，在圆桌对话环节与证券、经济及资产管理领域嘉宾围绕房地产市场、资本布局与产业创新等议题展开交流，并以主要组织方代表身份就峰会主题及住房政策研究中心后续工作方向发表观点。〔2〕〔3〕",
        size=10.5,
        space_after=6,
    )
    add_para(doc, "【人物经历·脚注增强段】", bold=True, size=11, space_after=2)
    add_para(
        doc,
        "2026年5月，胡继刚以复旦大学住房政策研究中心秘书长、上海市杨浦区科技企业联合会执行会长身份参加“2026人工智能商业化落地与硬核投资破局峰会”，主持“从算力引擎到新质资产（AI全产业链的商业化实战）”圆桌对话，并以主办方代表身份发表观点。〔1〕〔6〕〔7〕〔9〕",
        size=10.5,
        space_after=6,
    )
    add_para(doc, "【社会职务·执行会长】", bold=True, size=11, space_after=2)
    add_para(
        doc,
        "上海市杨浦区科技企业联合会执行会长。〔1〕〔4〕〔7〕〔9〕",
        size=10.5,
        space_after=8,
    )

    add_heading_custom(doc, "七、本版明确不做", level=1)
    for s in [
        "不写宝龙、万科、融创等企业职务（公开实名关联仍不足）。",
        "不写高级工程师、获奖、招商业绩。",
        "不单列社会活动，不写评价性称谓。",
        "不把一次公开发言写成“主要研究方向”。",
        "不加联系方式与商业推介。",
    ]:
        add_para(doc, f"· {s}", size=10.5, space_after=2)

    add_para(
        doc,
        "（完）提交前请再次打开〔9〕链接确认页面可访问。",
        size=10.5,
        color=RGBColor(0x66, 0x66, 0x66),
        space_before=10,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"已生成：{OUT}")
    return OUT


if __name__ == "__main__":
    build()
