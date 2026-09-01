#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成给同浦汇的业务承接策划案 Word 全文。"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
import proposal_data as D

NAVY = RGBColor(0x0E, 0x22, 0x40)
GOLD = RGBColor(0xC9, 0xA2, 0x27)
INK = RGBColor(0x24, 0x30, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TEAL = RGBColor(0x1F, 0x6B, 0x5C)

FONT = "微软雅黑"
OUT = Path(__file__).resolve().parents[1] / "output" / "同浦汇_30场活动与科技企业服务中心筹备_业务承接策划案.docx"


def set_run_font(run, size=11, bold=False, color=INK, font=FONT):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:cs"), font)


def shade(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell, color="C9A227"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def p(doc, text, size=11, bold=False, color=INK, align="left", space_after=8, space_before=0, first_line=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = 1.28
    if first_line is not None:
        para.paragraph_format.first_line_indent = Cm(first_line)
    para.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return para


def h1(doc, text):
    para = p(doc, text, size=16, bold=True, color=NAVY, space_before=16, space_after=10)
    # bottom border via paragraph border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "C9A227")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def h2(doc, text):
    return p(doc, text, size=13, bold=True, color=TEAL, space_before=12, space_after=6)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.clear()
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.2
    run = para.add_run(text)
    set_run_font(run, size=11, color=INK)
    return para


def table(doc, headers, rows, col_widths=None, header_fill="0E2240"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h)
        set_run_font(run, size=10, bold=True, color=WHITE)
        shade(cell, header_fill)
        set_cell_border(cell)
    for r_i, row in enumerate(rows):
        fill = "F4EBD3" if r_i % 2 else "FFFFFF"
        for c_i, val in enumerate(row):
            cell = t.rows[r_i + 1].cells[c_i]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            set_run_font(run, size=10, color=INK)
            shade(cell, fill)
            set_cell_border(cell, "E7DDC6")
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    # 页眉页脚
    header = sec.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("杨浦区科技企业服务中心 × 同浦汇　·　业务承接策划案")
    set_run_font(run, size=9, color=GOLD)

    footer = sec.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(D.CONFIDENTIAL + "　　")
    set_run_font(run, size=8, color=INK)
    run2 = fp.add_run("第 ")
    set_run_font(run2, size=8, color=INK)
    # PAGE field
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r_page = fp.add_run()
    set_run_font(r_page, size=8, color=INK)
    r_page._r.append(fld1)
    r_page._r.append(instr)
    r_page._r.append(fld2)
    run3 = fp.add_run(" 页")
    set_run_font(run3, size=8, color=INK)

    # 封面
    p(doc, "提交对象", size=12, color=GOLD, align="center", space_before=40)
    p(doc, D.DOC_FOR, size=18, bold=True, color=NAVY, align="center", space_after=16)
    p(doc, D.DOC_TITLE, size=22, bold=True, color=NAVY, align="center", space_after=4)
    p(doc, D.DOC_SUBTITLE, size=20, bold=True, color=GOLD, align="center", space_after=20)
    p(doc, "承接范围：30 场可核验活动  ＋  科技企业服务中心挂牌筹备", size=12, align="center")
    p(doc, f"新赛道：{D.NEW_POSITIONING}", size=12, align="center")
    p(doc, f"政策包装：{D.POLICY_PACKAGING}", size=12, align="center", space_after=18)
    p(doc, f"提出方：{D.DOC_FROM}", size=12, align="center")
    p(doc, f"联合：{D.DOC_COFROM}", size=12, align="center")
    p(doc, f"学术支持：{D.DOC_SUPPORT}", size=12, align="center")
    p(doc, f"{D.DOC_DATE}　｜　执行周期 {D.DOC_PERIOD}", size=12, align="center", space_after=12)
    p(doc, D.CONFIDENTIAL, size=10, color=GOLD, align="center", space_after=24)

    p(doc, "致同浦汇", size=14, bold=True, color=NAVY)
    p(
        doc,
        "昨天（8 月 31 日）关于创智汇赛道调整的交流已经把方向说清楚：园区不再以原 AI+IP 内容线作为主叙事，"
        "而是转向智能建造与建筑产业出海，并用现代服务业、新一代信息技术的口径完成政策包装。"
        "合同口径已经调整，运营成本已经发生，审核方工程部门需要尽快看到一份可执行、可核验、可过审的答卷。",
        align="justify",
        first_line=0.74,
    )
    p(
        doc,
        "据此，杨浦区科技企业服务中心提出：承接同浦汇在本项目上的两项具体业务——"
        "一是把 30 场活动按新赛道办完；二是把科技企业服务中心筹备到可挂牌、可申报。"
        "同浦汇继续做创智汇面前的产业招服入口和客户关系主人，不把园区接口让出去。"
        "本文件是给贵司的承接策划案，便于内部对齐后转交审核材料，不作为对园区或投资方的单方承诺函。",
        align="justify",
        first_line=0.74,
        space_after=16,
    )

    h1(doc, "一、昨日共识：为什么现在要承接")
    p(doc, "以下内容按 8 月 31 日交流整理，作为双方共同事实，不另做行业预测承诺。", align="justify")
    for i, t in enumerate(D.YESTERDAY, 1):
        bullet(doc, f"{i}. {t}")
    p(
        doc,
        f"一句话：旧定位「{D.OLD_POSITIONING}」让位给新定位「{D.NEW_POSITIONING}」。"
        "30 场年包的商务结构可以沿用，主题线和服中心筹备必须按新赛道重写。",
        align="justify",
        space_before=8,
    )

    h1(doc, "二、承接范围与不承接边界")
    h2(doc, "2.1 承接什么")
    table(
        doc,
        ["工作包", "具体内容"],
        [[a, b] for a, b in D.TAKEOVER],
        col_widths=[6.5, 10.5],
    )
    h2(doc, "2.2 明确不承接、不承诺")
    for t in D.NOT_TAKEOVER:
        bullet(doc, t)
    p(
        doc,
        "这样写，是为了让同浦汇在面对园区和审核方时站得住：该承诺的是场次、人数、月报和筹备节点；"
        "不承诺的是租金去化、外企到场、补贴获批和领事出席。",
        align="justify",
        first_line=0.74,
    )

    h1(doc, "三、角色分工")
    table(
        doc,
        ["主体", "在本承接方案中的责任"],
        [[a, b] for a, b in D.ROLES],
        col_widths=[5.0, 12.0],
    )
    p(doc, f"同浦汇联系人：{D.PARTIES['联系人']}", bold=True)
    p(doc, f"空间业主：{D.PARTIES['业主']}。审核相关方：{D.PARTIES['过审方']}。", align="justify")

    h1(doc, "四、工作包 A：30 场活动全案")
    h2(doc, "4.1 锁版商务口径（对园区不变）")
    table(
        doc,
        ["项目", "口径"],
        [
            ["产品", D.COMMERCIAL["活动年包"]],
            ["周期", D.DOC_PERIOD],
            ["人数 / 负责人", f"{D.EVENT_STANDARD['人数']}；{D.EVENT_STANDARD['负责人占比']}"],
            ["触达", D.EVENT_STANDARD["触达"]],
            ["付款", D.COMMERCIAL["付款"]],
            ["验收", D.COMMERCIAL["挂钩"]],
            ["门票赞助", D.COMMERCIAL["门票赞助"]],
        ],
        col_widths=[4.0, 13.0],
    )
    h2(doc, "4.2 新赛道下的六条线")
    table(
        doc,
        ["线条", "场次", "时间窗", "作用"],
        [[a, str(b), c, d] for a, b, c, d in D.THEMES],
        col_widths=[4.0, 2.0, 3.5, 7.5],
    )
    p(doc, "合计 30 场。原 ChinaJoy 内容主线不再作为年度叙事；科技能力保留为智能建造的工具，而不是园区主业。")

    h2(doc, "4.3 逐场排期")
    table(
        doc,
        ["编号", "月份", "线条", "活动名称", "本场作用"],
        [[a, b, c, d, e] for a, b, c, d, e in D.EVENTS],
        col_widths=[2.0, 2.4, 2.4, 5.6, 4.6],
    )

    h2(doc, "4.4 单场标准与转化闭环")
    bullet(doc, "执行节奏：" + D.EVENT_STANDARD["节奏"])
    bullet(doc, "转化闭环：" + D.EVENT_STANDARD["转化"])
    bullet(doc, "交付物：" + D.EVENT_STANDARD["交付"])
    p(
        doc,
        "同浦汇在闭环中多承担客户跟踪和回访台账；服中心重心放在下一场策划与现场。"
        "回访纪要只向园区交摘要知会，不要求园区逐户回访，也不再使用「关联度 100% 核验表」。",
        align="justify",
        first_line=0.74,
    )

    h2(doc, "4.5 不进年包的加购项")
    table(
        doc,
        ["项目", "内容", "计价原则"],
        [[a, b, c] for a, b, c in D.EXTRA_PRICE],
        col_widths=[4.0, 6.5, 6.5],
    )

    h1(doc, "五、工作包 B：科技企业服务中心筹备")
    h2(doc, "5.1 筹备目标")
    p(
        doc,
        "90 天内，把「上海市杨浦区科技企业服务中心」在创智汇做成可挂牌、可辅导、可申报的服务入口。"
        "挂牌仪式另计价；筹备过程不另向同浦汇收取费用。政策收益按第八章分成。",
        align="justify",
        first_line=0.74,
    )
    h2(doc, "5.2 九十天节奏")
    table(
        doc,
        ["窗口", "主题", "交付"],
        [[a, b, c] for a, b, c in D.CENTER_90],
        col_widths=[3.0, 3.5, 10.5],
    )
    h2(doc, "5.3 空间产品化")
    table(
        doc,
        ["空间", "口径"],
        [
            ["项目", f"{D.SPACE['项目']}　｜　{D.SPACE['区位']}　｜　{D.SPACE['面积']}"],
            ["3F", D.SPACE["3F"] + "　→　智能建造办公、装备体验、辅导接待"],
            ["5F", D.SPACE["5F"] + "　→　出海展陈、模块化样品、新材料、培训沙龙"],
            ["专题展区", D.SPACE["展厅专题"]],
            ["租金 / 物业", f"办公 {D.SPACE['办公租金']}　｜　物业 {D.SPACE['物业']}　｜　待租 {D.SPACE['待租']}"],
        ],
        col_widths=[4.0, 13.0],
    )
    p(
        doc,
        "租金 3.3 元高于周边，故本承接方案不做租赁对赌、不做租赁必要性。"
        "建议园区给予 1–3 个月免租，由同浦汇与园区线下谈，不写入服中心对同浦汇的承诺。",
        align="justify",
        first_line=0.74,
    )
    h2(doc, "5.4 政策适配（上限测算，非保证获批）")
    bullet(doc, "载体年上限：" + D.POLICY_CAP["载体年上限"])
    bullet(doc, "活动年上限：" + D.POLICY_CAP["活动年上限"])
    bullet(doc, "十年三部分合计上限：" + D.POLICY_CAP["十年三部分上限"])
    bullet(doc, "前置条件：" + D.POLICY_CAP["前置条件"])
    p(
        doc,
        "对企业：引导将智能施工、绿色节能、数字化设计写入经营范围，自行申报高新技术企业；"
        "服中心提供辅导 SOP，成功费以到账为准。"
        "对品牌：联合复旦住房政策研究中心开题白皮书，战略合作签约与官方授牌作为加购。"
        "零碳体验馆可借鉴广州越秀路径，用中小学研学形成不完全依赖补贴的运营入口。",
        align="justify",
        first_line=0.74,
    )

    h1(doc, "六、给审核方的三句口径")
    bullet(doc, "我们不是在招「付不起租金的施工队来租办公室」。")
    bullet(doc, "我们是在用科技企业服务中心，把智能建造产品、模块化建筑和绿色低碳能力组织成可出海的集群。")
    bullet(doc, "业主是杨浦科创集团；政府对接使用科技与产业服务口径，不以「中建」名义包装。")
    p(
        doc,
        "这三句直接对应昨天会上最容易被误解的三点：行业支付能力、出海集群目标、名义风险。"
        "同浦汇转交材料时建议原样保留。",
        align="justify",
        first_line=0.74,
    )

    h1(doc, "七、可核验指标")
    table(
        doc,
        ["指标", "标准"],
        [[a, b] for a, b in D.KPI],
        col_widths=[4.0, 13.0],
    )

    h1(doc, "八、商务与内部结算")
    h2(doc, "8.1 对园区")
    p(doc, "活动年包 30 万元、付款节点、不做租金对赌等口径见第四章。招商佣金按 2 个月净租金（首年不重复），由园区销售闭环触发。")
    h2(doc, "8.2 同浦汇 × 服中心（建议，供确认）")
    table(
        doc,
        ["收入类型", "建议分成", "说明"],
        [
            ["活动年包 30 万", "服中心 70% / 同浦汇 30%", D.COMMERCIAL["内部结算建议"]],
            ["政策申报收益", "同浦汇 38% / 服中心 62%", D.COMMERCIAL["政策分成"]],
            ["服中心筹备", "不另向同浦汇收费", D.COMMERCIAL["筹备费"]],
            ["出海 / 领事 / 挂牌", "一事一议", "与年包、政策分成分开签署"],
        ],
        col_widths=[4.0, 5.5, 7.5],
    )
    p(
        doc,
        "两套分成互不混用。活动年包对应执行劳动；政策收益对应牌照与申报主体。"
        "确认前，不把 38/62 写成活动包分成，也不把 70/30 写成政策分成。",
        align="justify",
        first_line=0.74,
    )

    h1(doc, "九、下一步（请同浦汇确认）")
    for i, t in enumerate(D.NEXT_STEPS, 1):
        bullet(doc, f"{i}. {t}")
    p(
        doc,
        "五件事书面确认后，服中心可在启动款到账后 14 日内交付细化执行手册，并按 Excel 台账启动 8–9 月场次。",
        align="justify",
        first_line=0.74,
        space_before=8,
    )

    h1(doc, "十、配套文件")
    bullet(doc, "PPT：《同浦汇_30场活动与科技企业服务中心筹备_业务承接策划案.pptx》（会面汇报，16 页）")
    bullet(doc, "Excel：《同浦汇_30场活动与科技企业服务中心筹备_执行台账.xlsx》（排期、分工、90 天、商务、风险）")
    bullet(doc, "本文 Word 为可打印、可批注的完整策划案正文。")

    p(doc, "（正文完）", align="center", space_before=18, color=GOLD)
    p(doc, f"{D.DOC_FROM}", align="center", bold=True, color=NAVY)
    p(doc, D.DOC_DATE, align="center")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Word 已写入 {OUT}")
    return OUT


if __name__ == "__main__":
    build()
