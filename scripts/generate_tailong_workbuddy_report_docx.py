#!/usr/bin/env python3
"""生成《泰隆银行上海分行与腾讯 WorkBuddy 合作推进情况汇报》Word 文稿。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


CHINESE_FONT = "宋体"
HEADING_FONT = "黑体"
ACCENT = RGBColor(0x0B, 0x2F, 0x5B)


def _set_run_font(run, font_name: str, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _add_paragraph(
    doc: Document,
    text: str = "",
    *,
    font: str = CHINESE_FONT,
    size: float = 12,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: float | None = None,
    space_after: float = 6,
    space_before: float = 0,
    line_spacing: float = 1.5,
    color: RGBColor | None = None,
):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if text:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold, color=color)
    return p


def _add_rich_paragraph(
    doc: Document,
    segments: list[tuple[str, bool]],
    *,
    font: str = CHINESE_FONT,
    size: float = 12,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: float | None = None,
    space_after: float = 6,
    space_before: float = 0,
    line_spacing: float = 1.5,
):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    for text, bold in segments:
        run = p.add_run(text)
        _set_run_font(run, font, size, bold=bold)
    return p


def _heading(doc: Document, text: str, level: int = 1) -> None:
    if level == 1:
        _add_paragraph(
            doc,
            text,
            font=HEADING_FONT,
            size=14,
            bold=True,
            space_before=14,
            space_after=8,
            color=ACCENT,
            line_spacing=1.3,
        )
    else:
        _add_paragraph(
            doc,
            text,
            font=HEADING_FONT,
            size=12,
            bold=True,
            space_before=10,
            space_after=6,
            line_spacing=1.3,
        )


def _body(doc: Document, text: str) -> None:
    _add_paragraph(doc, text, first_line_indent=0.74, space_after=6)


def _bullet(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    if bold_prefix:
        _add_rich_paragraph(
            doc,
            [("• ", False), (bold_prefix, True), (text, False)],
            first_line_indent=0.37,
            space_after=3,
            line_spacing=1.4,
        )
    else:
        _add_paragraph(
            doc,
            f"• {text}",
            first_line_indent=0.37,
            space_after=3,
            line_spacing=1.4,
        )


def build_document(output_path: Path) -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = CHINESE_FONT
    normal_style.font.size = Pt(12)
    rpr = normal_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CHINESE_FONT)

    # 标题
    _add_paragraph(
        doc,
        "泰隆银行上海分行与腾讯 WorkBuddy 合作推进情况汇报",
        font=HEADING_FONT,
        size=18,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
        line_spacing=1.3,
        color=ACCENT,
    )
    _add_rich_paragraph(
        doc,
        [("汇报对象：", True), ("泰隆银行上海分行行长", False)],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
        line_spacing=1.3,
    )
    _add_rich_paragraph(
        doc,
        [("汇报日期：", True), ("2026年8月3日", False)],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=14,
        line_spacing=1.3,
    )

    # 一、总体情况
    _heading(doc, "一、总体情况")
    _body(
        doc,
        "我行已与腾讯 WorkBuddy 团队开展两轮业务洽谈。双方初步形成共识：以服务中小微企业为切入点，"
        "将企业级AI工具纳入我行客户权益体系，通过“积分商城兑换、开户赠礼、客户联合推广”等方式，"
        "为中小微客户提供差异化增值服务，助力我行新增对公客户、客户活跃及存款沉淀。",
    )
    _body(
        doc,
        "在此基础上，建议同步研究“行员采购 + 行外采购”双轨份额安排：一是面向行内，联手打造提升办公效率的内部系统与小程序，"
        "并以福利形式向员工发放AI使用权益，提升行员个人AI应用能力；二是面向行外，继续推进客户权益合作，并研究发行泰隆银行与腾讯WorkBuddy联名信用卡，"
        "将AI权益与信用卡获客、激活及消费场景进一步结合。",
    )
    _body(
        doc,
        "目前，合作方向和基本模式已经明确，但采购规模、价格体系、技术开发、安全合规、结算开票、联名卡产品方案及内部审批等事项仍需进一步论证。"
        "建议原则同意继续推进，先形成试点方案，经合规、科技、财务、信用卡等部门评估后实施。",
    )

    # 二、合作背景
    _heading(doc, "二、合作背景")
    _body(
        doc,
        "我行在上海经营约16年，已积累近100万客户，主要服务中小微企业。当前银行存款、贷款、理财、结算等传统产品同质化程度较高，"
        "单纯依靠利率、价格和担保方式较难形成明显差异。",
    )
    _body(
        doc,
        "我行经营考核的核心之一是持续增加对公客户。客户数量扩大后，可进一步带动存款、贷款、结算、代发工资、理财等综合业务。"
        "因此，需要引入更贴近企业经营需求的非金融增值服务，提高开户吸引力和存量客户黏性。",
    )
    _body(
        doc,
        "WorkBuddy定位为企业级AI应用及技能平台。据腾讯团队介绍，目前平台已有约2.8万家企业客户，企业版日活跃用户超过10万，"
        "可为企业提供内容生产、数据分析、知识管理、行业研究等AI应用。其企业版产品价格相对较低，适合作为中小微客户低成本体验AI应用的入口。",
    )

    # 三、拟议合作模式
    _heading(doc, "三、拟议合作模式")
    _body(
        doc,
        "拟议合作模式按“行外客户合作、行内员工合作、联名信用卡”三条主线推进，采购份额相应拆分为行外采购与行员采购两部分，"
        "便于分别核算投入产出并分步审批落地。",
    )

    _heading(doc, "（一）面向外部客户：积分、开户与联合推广", level=2)
    _add_rich_paragraph(doc, [("1. 积分权益商城合作", True)], space_before=4, space_after=2, line_spacing=1.3)
    _body(
        doc,
        "由我行集中采购WorkBuddy账号、使用额度或兑换权益，上架我行积分权益商城，支持客户使用积分兑换或直接购买。",
    )
    _body(
        doc,
        "该模式能够丰富我行积分消耗场景，提高积分使用率。后续如WorkBuddy以商户形式入驻我行合作商城并在我行开立结算账户，"
        "相关交易资金可在我行体系内结算，有望形成一定的资金沉淀和结算业务。"
        "具体商城主体、资金路径、账户开立及收入结算方式，仍需财务、法务及商城运营方进一步确认。",
    )
    _add_rich_paragraph(doc, [("2. 开户及客户拜访赠礼", True)], space_before=4, space_after=2, line_spacing=1.3)
    _body(
        doc,
        "客户经理在拓展新客户或维护重点客户时，可将一定期限的WorkBuddy企业账号或AI使用额度作为增值权益赠送，替代部分传统实物礼品。"
        "客户取得兑换码后，可扫码注册并输入License，自助完成激活，原则上不需要客户经理人工开通和维护。"
        "该方式具有一定的新颖性和实用性，可用于提升客户拜访体验和开户转化率。",
    )
    _add_rich_paragraph(doc, [("3. 联合推广中小微客户", True)], space_before=4, space_after=2, line_spacing=1.3)
    _body(
        doc,
        "依托我行线下网点、客户经理、电销及面销团队，面向现有中小微客户推广WorkBuddy。"
        "腾讯团队负责产品、技术及使用支持，我行负责客户触达和场景组织。"
        "我行上海地区拥有130余个经营网点，并具备较强的客户经理和地推能力，可先选择制造业、商贸服务业等AI需求相对明确的客群进行试点。",
    )

    _heading(doc, "（二）面向行内员工：补充员工采购份额", level=2)
    _body(
        doc,
        "建议在总体采购中单独预留员工份额，用于提升行内办公效率与员工个人AI使用能力。该部分原则上与行外客户采购分账核算，"
        "以便明确费用归属、使用范围和合规边界。",
    )
    _add_rich_paragraph(
        doc,
        [("1. 提升行内办公系统效率（初步方案）", True)],
        space_before=4,
        space_after=2,
        line_spacing=1.3,
    )
    _body(
        doc,
        "由我行与WorkBuddy联手打造提升办公效率的内部系统及小程序，面向行内办公场景提供内容起草、会议纪要、知识检索、客户服务辅助、"
        "业务材料整理等AI能力，降低重复性事务性工作负担，提高网点与部门协同效率。",
    )
    _body(
        doc,
        "该方向涉及内部系统对接、权限管理、数据安全及信息科技立项，须经科技、合规、数据安全等部门专项评估后推进；"
        "首期建议以非核心、非敏感办公场景试点，严格限定数据出入边界。",
    )
    _add_rich_paragraph(
        doc,
        [("2. 提升员工个人AI使用能力（福利发放）", True)],
        space_before=4,
        space_after=2,
        line_spacing=1.3,
    )
    _body(
        doc,
        "以员工福利形式向行员发放WorkBuddy个人或轻量企业版账号、使用额度或兑换码，鼓励员工学习并掌握AI工具，"
        "提升数字化办公与客户服务能力。发放对象、额度标准、考核挂钩方式及是否计入福利费，需由人力资源、财务及合规部门共同确认。",
    )

    _heading(doc, "（三）泰隆银行与腾讯WorkBuddy联名信用卡", level=2)
    _body(
        doc,
        "建议原则同意研究发行泰隆银行与腾讯WorkBuddy联名信用卡。对外以联名卡作为获客与品牌合作载体，"
        "持卡客户可获得AI权益包、积分加倍、开卡礼或续费折扣等差异化权益；对内可与员工福利、客户权益、商城兑换形成联动。",
    )
    _body(
        doc,
        "联名卡方案尚属初步设想，卡种定位、权益结构、分润结算、品牌授权、发卡规模及风控规则，"
        "需由信用卡中心会同零售、科技、合规、法务等部门论证后，再纳入正式产品规划。",
    )

    _heading(doc, "（四）后续延伸方向", level=2)
    _body(doc, "如首期合作效果较好，可进一步研究：")
    _bullet(doc, "在WorkBuddy技能平台开设“泰隆银行专区”，展示我行结算、融资、国际业务及财资管理等服务入口。")
    _bullet(doc, "根据不同行业客户需求，联合开发企业经营类AI技能。")
    _bullet(doc, "将客户对AI工具的需求与我行开户、结算、代发工资等金融服务进一步结合。")
    _bullet(doc, "将联名信用卡与内部办公小程序、员工福利发放、客户权益兑换打通，形成“行内提效—员工赋能—客户获客”闭环。")

    # 四、采购方案设想
    _heading(doc, "四、采购方案设想")
    _body(
        doc,
        "腾讯团队在第二轮沟通中提出，可以围绕较大规模合作设计季度或半年度套餐，不采用单个账号按198元零散销售的模式。"
        "建议总体采购按“行外采购 + 行员采购”双轨拆分：行外份额主要用于客户积分兑换、开户赠礼与联合推广；"
        "行员份额主要用于内部办公提效工具与员工福利发放。联名信用卡相关权益可与上述份额联动配置，也可单独测算。",
    )
    _body(doc, "目前讨论中的框架为：")
    _bullet(doc, "总体合作规模可按最高约1,500万元进行方案设计；")
    _bullet(doc, "分三阶段推进，每阶段约500万元；")
    _bullet(doc, "首期金额可在300万至500万元范围内研究，并明确其中行外与行员份额比例；")
    _bullet(doc, "根据实际兑换率、激活率、客户使用率、员工使用率及新增客户/发卡效果，决定后续采购规模。")
    _body(
        doc,
        "需要说明的是，1,500万元属于谈判过程中的方案设想，尚未形成我行正式预算或采购承诺。"
        "首期采购金额也应以内部立项、采购审批、价格谈判及试点测算结果为准。",
    )
    _body(doc, "建议腾讯团队至少提供300万元、500万元两个首期档位，并分别说明：")
    _bullet(doc, "可采购的账号、License及Token数量，以及行外/行员可分配份额；")
    _bullet(doc, "权益有效期及激活后的使用期限；")
    _bullet(doc, "阶梯折扣及价格保护机制；")
    _bullet(doc, "未激活权益的处理方式；")
    _bullet(doc, "后续扩容、续费及退款规则；")
    _bullet(doc, "三阶段付款条件和验收标准；")
    _bullet(doc, "内部办公系统/小程序定制开发范围、周期与报价（如需）；")
    _bullet(doc, "联名信用卡权益包方案及结算分润建议。")

    # 五、预期价值
    _heading(doc, "五、预期价值")
    _heading(doc, "1. 带动新增对公客户", level=2)
    _body(
        doc,
        "将AI工具作为开户或客户拜访权益，有望形成区别于同业传统礼品和金融产品的服务特色，提高客户接触和开户转化效果。",
    )
    _heading(doc, "2. 提升存量客户黏性", level=2)
    _body(
        doc,
        "通过积分兑换和企业AI服务，增加我行与企业老板、财务及经营管理人员的互动触点，为后续存款、贷款、结算等业务创造机会。",
    )
    _heading(doc, "3. 促进积分活跃", level=2)
    _body(
        doc,
        "AI账号及使用额度可丰富现有视频会员、生活消费类权益之外的企业服务权益，提高企业客户积分使用意愿。",
    )
    _heading(doc, "4. 带动结算与资金沉淀", level=2)
    _body(
        doc,
        "如商城交易、商户结算及后续续费资金通过我行账户完成，可形成结算流水和一定的资金沉淀。实际规模需结合交易路径进一步测算。",
    )
    _heading(doc, "5. 提升行内办公效率与员工能力", level=2)
    _body(
        doc,
        "通过内部系统/小程序提效与员工福利发放，增强行员数字化工具使用能力，间接提升客户服务专业度与网点运营效率。",
    )
    _heading(doc, "6. 带动联名信用卡获客与品牌曝光", level=2)
    _body(
        doc,
        "联名信用卡可将AI权益与办卡、激活、消费场景结合，扩大零售获客触点，并强化我行与腾讯生态的品牌联动。",
    )
    _heading(doc, "7. 形成科技服务品牌", level=2)
    _body(
        doc,
        "通过与腾讯团队开展AI服务合作，有利于塑造我行“金融服务+企业经营赋能”的品牌形象，强化我行服务中小微企业的市场特色。",
    )

    # 六、风险及待解决事项
    _heading(doc, "六、风险及待解决事项")

    _heading(doc, "（一）安全合规风险", level=2)
    _body(
        doc,
        "面向外部客户的WorkBuddy合作，目前拟采用公有云SaaS模式，原则上不接入我行内部网络，也不获取我行内部系统权限。"
        "但客户对话内容、企业资料、管理员日志权限等仍涉及数据安全和隐私保护，不能仅凭商务沟通认定合规。",
    )
    _body(
        doc,
        "若推进行内办公系统/小程序合作，则合规要求显著提高，可能涉及内部办公数据、员工账号权限及与行内系统的接口对接，"
        "必须由法务、合规、信息科技和数据安全部门开展专项审查，明确：",
    )
    _bullet(doc, "数据采集范围、存储地点和保存期限；")
    _bullet(doc, "企业租户之间的数据隔离机制；")
    _bullet(doc, "腾讯及平台管理员的数据访问权限；")
    _bullet(doc, "对话日志的开启、关闭、导出和删除规则；")
    _bullet(doc, "数据泄露后的责任承担和应急机制；")
    _bullet(doc, "客户授权、隐私政策及风险提示方式；")
    _bullet(doc, "行内使用场景的数据分类分级、脱敏要求及是否允许出域。")
    _body(
        doc,
        "首期建议：行外客户权益与员工个人福利账号可先行论证；涉及行内办公系统/小程序的，须完成专项合规评估后方可试点，"
        "不得用于处理客户敏感信息或核心业务信息。",
    )

    _heading(doc, "（二）采购及成本风险", level=2)
    _body(
        doc,
        "目前大额采购方案主要基于我行客户规模和推广能力，尚缺少实际兑换率、激活率、持续使用率及获客转化率数据。"
        "若一次性采购规模过大，可能出现权益闲置。",
    )
    _body(
        doc,
        "建议采用分期采购、分批交付和效果验收机制，并约定未激活License的有效期、替换及顺延政策。",
    )

    _heading(doc, "（三）技术交付风险", level=2)
    _body(
        doc,
        "批量兑换码生成、客户自主注册、实名认证、分组管理、额度充值、续费及停用等功能仍需腾讯技术团队确认现有能力及开发周期。",
    )
    _body(
        doc,
        "在正式采购前，应完成完整流程测试，确保百万级客户触达情况下系统能够稳定运行。",
    )

    _heading(doc, "（四）结算及税务风险", level=2)
    _body(
        doc,
        "商城运营主体、采购发票类型、客户购买发票、腾讯与商城运营方的结算关系仍在沟通。"
        "应由财务、税务及法务部门确认资金链路和开票安排后再落地。",
    )

    _heading(doc, "（五）宣传口径与联名品牌风险", level=2)
    _body(
        doc,
        "“腾讯联合产品”“行业首创”“联名信用卡”等表述需取得对方正式品牌授权并经过我行宣传及信用卡中心审核，"
        "避免因品牌使用或效果承诺不准确产生风险。",
    )

    _heading(doc, "（六）联名信用卡产品风险", level=2)
    _body(
        doc,
        "联名卡涉及产品准入、权益兑付、分润结算、投诉处理及合作期限安排，尚需与信用卡中心专项论证，"
        "不宜在商务框架未定前作出发卡规模或权益承诺。",
    )

    # 七、初步共识
    _heading(doc, "七、两轮谈判形成的初步共识与补充方案")
    _body(doc, "已达成共识：")
    consensus = [
        "优先推进积分商城兑换和开户赠礼两类外部客户场景。",
        "原则上采用兑换码或License方式，实现客户自主注册激活。",
        "暂不印制实体卡（权益卡），待确定电子卡或纸质卡形式后再安排。",
        "采购方案按照季度或半年度套餐及阶梯价格设计。",
        "腾讯团队提供包含账号、Token、价格、付款节奏及技术交付的完整方案。",
        "双方暂定于8月5日（周三）下午与分行领导进一步沟通。",
        "如条件成熟，优先在上海选择部分客户开展试点，再根据效果扩大范围。",
    ]
    for i, item in enumerate(consensus, 1):
        _add_paragraph(doc, f"{i}. {item}", first_line_indent=0.37, space_after=3, line_spacing=1.4)

    _body(doc, "初步方案（本次补充，尚待确认）：")
    prelim = [
        "在采购中单列员工份额，采用“行员采购 + 行外采购”双轨安排。",
        "行内方向一：与WorkBuddy联手打造提升办公效率的内部系统及小程序。",
        "行内方向二：以福利形式向员工发放AI账号或额度，提升个人AI使用能力。",
        "对外研究发行泰隆银行与腾讯WorkBuddy联名信用卡，将AI权益与发卡获客结合。",
    ]
    for i, item in enumerate(prelim, 1):
        _add_paragraph(doc, f"{i}. {item}", first_line_indent=0.37, space_after=3, line_spacing=1.4)

    # 八、建议决策事项
    _heading(doc, "八、建议行长决策事项")
    _body(doc, "建议本次重点请示五项：")

    _add_rich_paragraph(
        doc,
        [("1. 是否原则同意继续推进。", True)],
        space_before=4,
        space_after=2,
        line_spacing=1.3,
    )
    _body(doc, "建议同意将该项目作为中小微客户增值服务及行内数字化提效创新方向继续论证。")

    _add_rich_paragraph(
        doc,
        [("2. 是否同意开展上海地区试点。", True)],
        space_before=4,
        space_after=2,
        line_spacing=1.3,
    )
    _body(doc, "建议选取部分网点、客户经理和目标行业开展首期客户侧试点，设定明确的兑换、激活、使用及获客指标。")

    _add_rich_paragraph(
        doc,
        [("3. 是否同意补充员工采购份额并论证行内应用。", True)],
        space_before=4,
        space_after=2,
        line_spacing=1.3,
    )
    _body(
        doc,
        "建议原则同意按“行员采购 + 行外采购”拆分份额；员工福利发放可先行论证，"
        "内部办公系统/小程序须经科技与合规评估后试点。",
    )

    _add_rich_paragraph(
        doc,
        [("4. 是否原则同意研究联名信用卡方案。", True)],
        space_before=4,
        space_after=2,
        line_spacing=1.3,
    )
    _body(doc, "建议授权信用卡中心会同相关部门论证泰隆银行与腾讯WorkBuddy联名信用卡产品方案，不作为立即发卡承诺。")

    _add_rich_paragraph(
        doc,
        [("5. 是否同意启动跨部门评估并授权继续商务谈判。", True)],
        space_before=4,
        space_after=2,
        line_spacing=1.3,
    )
    _body(
        doc,
        "建议明确牵头部门，组织公司金融、零售权益、信用卡、人力资源、科技、合规、法务、财务和采购等部门共同评估；"
        "以最高1,500万元作为方案测算框架，不作为采购承诺，首期规模另行审批。",
    )

    # 九、下一步
    _heading(doc, "九、下一步工作建议")
    next_steps = [
        "要求腾讯团队在8月5日会谈前提交完整的产品及商务方案，并区分行外/行员份额。",
        "明确首期300万元、500万元两个采购档位及对应权益，单列员工福利与客户权益数量。",
        "完成兑换码、注册、激活、分组、续费和数据管理的流程演示。",
        "就内部办公系统/小程序提出场景清单与合规边界，征求科技、合规意见。",
        "启动联名信用卡可行性预研，明确卡种定位、权益包与品牌授权路径。",
        "启动法务、合规、科技、财务、采购、人力及信用卡中心预审。",
        "设计首期试点方案，明确目标客户、试点网点、员工发放范围和量化指标。",
        "将后续采购与试点效果挂钩，达到约定指标后再启动第二、三阶段采购。",
    ]
    for i, item in enumerate(next_steps, 1):
        _add_paragraph(doc, f"{i}. {item}", first_line_indent=0.37, space_after=3, line_spacing=1.4)

    # 十、结论
    _heading(doc, "十、汇报结论")
    _body(
        doc,
        "总体看，该合作与我行服务中小微企业、增加对公客户、提升行内办公效率及建设差异化增值服务的方向较为契合，"
        "双方资源具有互补性，具备进一步推进价值。",
    )
    _body(
        doc,
        "但目前仍处于商务方案和可行性论证阶段，尤其是大额采购依据、行内外份额拆分、安全合规、技术交付、"
        "结算开票、联名信用卡产品方案及客户转化效果尚未完全确定。",
    )
    _add_paragraph(
        doc,
        "建议行长原则同意立项论证并授权继续谈判，同步论证员工份额与联名信用卡方案，先试点、后扩量，"
        "以实际客户转化、员工使用及发卡效果作为后续采购依据。",
        font=HEADING_FONT,
        size=12,
        bold=True,
        first_line_indent=0.74,
        space_before=8,
        space_after=10,
        line_spacing=1.5,
        color=ACCENT,
    )

    _add_paragraph(
        doc,
        "口头汇报一句话概括：",
        font=HEADING_FONT,
        size=12,
        bold=True,
        space_before=8,
        space_after=4,
    )
    _add_paragraph(
        doc,
        "两轮洽谈后，双方已明确以“积分商城兑换+开户赠送AI权益”为外部客户主方向，并建议补充员工采购份额"
        "（内部办公提效系统/小程序 + 员工AI福利）及联名信用卡方案；建议原则同意推进，但1,500万元目前只是方案框架，"
        "需先完成合规审查和小范围试点，再根据实际效果分阶段采购。",
        first_line_indent=0.74,
        space_after=6,
        line_spacing=1.5,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"已生成：{output_path}")


if __name__ == "__main__":
    out = Path(__file__).resolve().parents[1] / "deliverables" / "泰隆银行上海分行与腾讯WorkBuddy合作推进情况汇报_20260803.docx"
    build_document(out)
