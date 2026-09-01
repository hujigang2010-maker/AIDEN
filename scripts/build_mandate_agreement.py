"""
冠松 01# 研发楼 · 招商服务委托协议（Word）

性质：甲方（冠松）独家委托乙方（招商操盘方）替其完成招商，
      不是顾问咨询备忘录，也不是中介居间协议。

重新生成：python3 scripts/build_mandate_agreement.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

CN_FONT = "WenQuanYi Micro Hei"
EN_FONT = "Calibri"

OUT_DIR = Path(__file__).resolve().parent.parent / "docs" / "legal"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def set_cn_font(run, size=11, bold=False, color=None):
    run.font.name = EN_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_title(doc, text, size=20):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=True)


def add_subtitle(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=False, color=RGBColor(0x6B, 0x73, 0x80))


def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, size=14, bold=True, color=RGBColor(0x0F, 0x2D, 0x52))


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, size=12, bold=True)


def add_p(doc, text, indent=False, bold=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_cn_font(r, size=11, bold=bold)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_cn_font(r, size=11)


def add_kv_table(doc, kv_list):
    table = doc.add_table(rows=len(kv_list), cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11)
    for i, (k, v) in enumerate(kv_list):
        cell_k = table.rows[i].cells[0]
        cell_v = table.rows[i].cells[1]
        cell_k.text = ""
        cell_v.text = ""
        rk = cell_k.paragraphs[0].add_run(k)
        set_cn_font(rk, size=11, bold=True)
        rv = cell_v.paragraphs[0].add_run(v)
        set_cn_font(rv, size=11)


def add_grid(doc, header, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            table.columns[i].width = Cm(w)
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_cn_font(r, size=10, bold=True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            set_cn_font(r, size=10)


def add_signature_block(doc, parties):
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=len(parties))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["（盖章）", "法定代表人 / 授权代表：", "签署日期：", ""]
    for j, p in enumerate(parties):
        cell = table.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(p)
        set_cn_font(r, size=12, bold=True)
    for i in range(1, 4):
        for j in range(len(parties)):
            cell = table.rows[i].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(headers[i])
            set_cn_font(r, size=11)


def doc_setup(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    rPr.append(rFonts)


def build_leasing_mandate():
    doc = Document()
    doc_setup(doc)

    add_title(doc, "招 商 服 务 委 托 协 议")
    add_subtitle(doc, "独家全案操盘 · 冠松 01# 研发楼")
    add_subtitle(doc, "GS · iDrive Hub · 草案 v1.0 · 签约前由法务定稿")

    add_p(doc, "")
    add_p(
        doc,
        "重要说明：本协议不是顾问咨询合同，也不是中介居间合同。"
        "甲方委托乙方作为独家招商操盘方，替甲方完成 01# 研发楼的招商落地"
        "（名单攻坚、带看、谈判、签约配合），以生效租赁合同为准考核。"
        "方括号内为待填项；费用数字为建议口径，签署前双方书面确认。",
        bold=True,
    )

    add_h2(doc, "合同编号：GS-MANDATE-XXXX")

    add_h1(doc, "签约方信息")
    add_kv_table(doc, [
        ("委托方（甲方）", "上海冠松 [项目运营公司全称] 有限公司"),
        ("角色", "01# 研发楼合法权属人 / 招商委托方"),
        ("注册地址", "上海市静安区永和路 [号待补]"),
        ("法定代表人", "[姓名]"),
        ("指定对接人", "[项目负责人姓名 / 职务 / 手机 / 邮箱]"),
        ("受托方（乙方）", "[招商操盘方全称]"),
        ("角色", "独家招商操盘服务提供方"),
        ("注册地址", "[乙方注册地]"),
        ("法定代表人", "[姓名]"),
        ("项目经理（驻场）", "[姓名 / 手机 / 邮箱]"),
    ])

    add_h1(doc, "鉴于")
    add_p(
        doc,
        "1. 甲方系上海市静安区永和社区 N070601 单元 075b-07 地块"
        "『01# 研发楼』（地上 9F 共 15,152.75 ㎡，地下 2F 共 6,992.87 ㎡，"
        "用地性质为 C6 教育科研设计用地）的合法权属人，拟将该楼由原整车展厅思路"
        "转为高新研发及汽车配套研发业态并对外招商；",
        indent=True,
    )
    add_p(
        doc,
        "2. 乙方具备产业园区及大宗租赁招商操盘能力，愿意接受甲方委托，"
        "以独家全案操盘方式替甲方完成招商，而不仅限于出具方案或提供咨询意见；",
        indent=True,
    )
    add_p(
        doc,
        "3. 双方确认：出具研究报告、案例对标、汇报材料，均不视为本协议项下"
        "招商服务已完成；完成的唯一业务标准是达成符合本协议约定的生效租赁合同。",
        indent=True,
    )
    add_p(doc, "双方经平等协商，订立本协议如下。", indent=True)

    add_h1(doc, "第一条 · 定义")
    add_p(doc, "1.1 标的：指 01# 研发楼可对外出租的研发办公空间。", indent=True)
    add_p(
        doc,
        "1.2 可出租净面积：双方暂按约 8,300 ㎡ 计（1F+2F 大堂/展厅由甲方自留或冠名，"
        "不纳入招商考核分母；最终以双方书面确认的面积表为准）。",
        indent=True,
    )
    add_p(
        doc,
        "1.3 生效租赁合同：甲方与承租人正式签署、双方盖章，且承租人已支付履约保证金"
        "或首期租金（以先到者为准）的房屋租赁合同。仅有口头意向、会议纪要、"
        "Term Sheet 或未付款的草签，均不构成生效租赁合同。",
        indent=True,
    )
    add_p(
        doc,
        "1.4 链主/锚定租户：单户签约面积 ≥ 1,500 ㎡，或承租整层，"
        "或经双方书面确认为甲档链主总部的承租人。",
        indent=True,
    )
    add_p(
        doc,
        "1.5 有效意向：承租人已完成现场带看，并出具加盖公章或授权代表签署的"
        "书面意向函 / Term Sheet，意向面积与主要商务条款齐备。",
        indent=True,
    )
    add_p(
        doc,
        "1.6 入驻率：累计生效租赁合同签约面积 ÷ 可出租净面积。",
        indent=True,
    )

    add_h1(doc, "第二条 · 委托性质与排他")
    add_p(
        doc,
        "2.1 甲方在本协议期限内，独家委托乙方作为标的之全案招商操盘方。"
        "乙方以自己的专业能力替甲方招商，但租赁合同的出租方仍为甲方，"
        "乙方无权以甲方名义签署租赁合同，除非甲方另行出具书面授权委托书。",
        indent=True,
    )
    add_p(
        doc,
        "2.2 排他范围：委托期内，甲方不得另行委托任何第三方就同一标的从事"
        "全案招商策划、操盘、统筹。中介分销渠道可以保留，但须纳入乙方统一报备与排期管理。",
        indent=True,
    )
    add_p(
        doc,
        "2.3 甲方自行成交：甲方自有客户未经乙方书面报备而成交的，不向乙方支付成功佣金；"
        "但仍应按约支付月度操盘费。甲方应在接触该等客户后 3 个工作日内书面告知乙方，避免撞单。",
        indent=True,
    )
    add_p(
        doc,
        "2.4 本协议与甲方另行签署的《房屋租赁居间服务协议》（中介渠道）并行："
        "中介佣金由甲方按居间协议直接支付中介；乙方不截留中介佣金。"
        "乙方负责统管报备、避免同一客户被多家中介重复报备。",
        indent=True,
    )

    add_h1(doc, "第三条 · 标的概况")
    add_kv_table(doc, [
        ("楼宇", "01# 研发楼 · 地上 9F / 地下 2F"),
        ("地址", "上海市静安区永和社区 N070601 单元 075b-07"),
        ("用地", "C6 教育科研设计用地（准入硬约束）"),
        ("规模", "地上 15,152.75 ㎡ · 地下 6,992.87 ㎡ · 建筑高度 44.95 m"),
        ("层高", "1F 5.7 m / 2F 6.3 m / 3F 5.7 m / 4F 5.4 m / 5–8F 4.2 m / 9F 4.3 m"),
        ("可出租净面积", "约 8,300 ㎡（考核分母，以面积确认表为准）"),
        ("建议落位", "8–9F 链主总部 · 3–4F 核心研发 · 6–7F 算法/软件"),
        ("自留", "1F+2F 大堂及展厅由甲方自留或用于冠名，默认不对外招整车展厅"),
    ])

    add_h1(doc, "第四条 · 委托事项（乙方必须做的事）")
    add_p(doc, "4.1 乙方接受委托后，应亲自组织并完成下列工作，不得以『已出方案』替代：", indent=True)
    add_bullet(doc, "客群锁定：按约定业态建立不少于 200 家目标企业库，明确决策人与跟进责任人")
    add_bullet(doc, "名单攻坚：对重点客户完成电话/微信建联、拜访、痛点确认，形成可核验的接触记录")
    add_bullet(doc, "现场带看：组织看房、讲解楼宇与政策口径、收集反馈，每次带看 48 小时内出纪要")
    add_bullet(doc, "商务谈判：在授权范围内报价、起草 Term Sheet / 租赁合同商务条款，推动法务闭环")
    add_bullet(doc, "中介统管：遴选、培训、任务量考核、报备冲突裁决，对甲方只保留一套客户真相")
    add_bullet(doc, "政府协同：协助甲方准备投促材料、陪同接待、整理『一企一策』诉求清单（对外承诺须甲方确认）")
    add_bullet(doc, "活动执行：策划并落地招商活动、小型闭门会、发布会招商环节（硬成本由甲方承担）")
    add_bullet(doc, "签约配合：从意向到盖章到保证金到账，盯节点、补材料、对接口径")
    add_p(doc, "4.2 下列事项明确不属于『已完成招商』，不得作为收费完成依据：", indent=True)
    add_bullet(doc, "仅出具研究报告、案例对标、PPT、咨询意见")
    add_bullet(doc, "仅召开内部讨论会、务虚会、参观考察")
    add_bullet(doc, "仅提供名单而无有效接触记录")
    add_bullet(doc, "仅有口头意向或未付款的草签")

    add_h1(doc, "第五条 · 业态准入与禁止（写进合同，不再口头）")
    add_p(doc, "5.1 允许招入的业态：", indent=True)
    add_bullet(doc, "人工智能、集成电路设计及测试展示、智能机器人等研发设计类")
    add_bullet(doc, "生物医药研发、医疗器械研发（须满足 C6 及环评/特种设备等法定要求）")
    add_bullet(doc, "汽车配套研发（电控、传感、支架、软件算法、测试认证等，非整车销售）")
    add_p(doc, "5.2 禁止招入的业态：", indent=True)
    add_bullet(doc, "整车 4S 店、整车展厅、汽车零售卖场")
    add_bullet(doc, "纯商业零售、教培、仓储物流、生产制造")
    add_bullet(doc, "不符合 C6 教育科研设计用地主导功能的其他用途")
    add_p(
        doc,
        "5.3 灰色业态：人才短租/公寓仅可作为配建或地下辅助，不得作为地上主业态对外招商；"
        "启动前须完成 C6 合规评估，并经甲方书面同意。未经同意不得向客户承诺可办公寓。",
        indent=True,
    )
    add_p(
        doc,
        "5.4 用途相符：所有承租人合同用途条款须符合 C6 用地及本条。乙方起草文本时应嵌入该约束。",
        indent=True,
    )

    add_h1(doc, "第六条 · 招商授权（红黄绿）")
    add_p(
        doc,
        "6.1 乙方仅可在下列授权内对外报价或书面承诺。超线必须休会，书面请示甲方指定对接人；"
        "触及底线须集团董事长（或甲方书面指定的最终决策人）批准。",
        indent=True,
    )
    add_grid(
        doc,
        ["条款", "绿区（可直接报）", "黄区（48 小时内批复）", "红区（董事长批）"],
        [
            ["起始租金（元/㎡·天，不含税不含物业）", "≥ 6.5", "5.8–6.5（不含 6.5）", "< 5.8；底线 5.0"],
            ["免租期（月，含装修）", "≤ 9", "9–15（不含 9）", "> 15；底线 24"],
            ["装补（元/㎡）", "≤ 600", "600–1,000", "> 1,000；底线 1,500"],
            ["合同期（年）", "≥ 8", "6–8（不含 8）", "< 6；底线 5"],
            ["履约保证（月租金）", "≥ 6", "3–6（不含 6）", "< 3"],
        ],
        col_widths=[5.5, 3.5, 4.0, 3.5],
    )
    add_p(doc, "")
    add_p(
        doc,
        "6.2 绝对禁止（无论谁批准，乙方亦不得代为承诺）："
        "现金回扣、私人佣金、未获政府书面文件的政策返还/落户/人才公寓配额、"
        "包牌照、保 GMV、任何形式的个人对赌。",
        indent=True,
    )
    add_p(
        doc,
        "6.3 黄区请示：甲方指定对接人应在收到完整书面请示后 48 小时内书面答复；"
        "逾期未答复视为不同意，乙方不得按黄区条件对外承诺。",
        indent=True,
    )
    add_p(
        doc,
        "6.4 授权可调整：甲方可以书面通知上调或收紧绿区。通知到达次日起对新报价生效，"
        "不影响已发出且仍在有效期内的书面报价。",
        indent=True,
    )

    add_h1(doc, "第七条 · 工作节奏、人员与报告")
    add_p(
        doc,
        "7.1 项目经理：乙方指定 1 名项目经理，委托期内保持稳定；"
        "每周现场办公不少于 3 个工作日（法定节假日除外），对甲方指定对接人直报。",
        indent=True,
    )
    add_p(
        doc,
        "7.2 核心组：不少于 3 人（项目经理 + 招商 + 材料/商务）。试跑 90 天可先 2 人到位、第 2 个月补齐。"
        "人员名单于进场 5 个工作日内书面提交。",
        indent=True,
    )
    add_p(doc, "7.3 报告：", indent=True)
    add_bullet(doc, "周报：每周五 18:00 前书面提交漏斗（新增线索、带看、谈判、阻塞点、下周动作）")
    add_bullet(doc, "月会：每月第一个工作周召开月度经营会，当面或视频，时长不少于 60 分钟")
    add_bullet(doc, "季报：每季度书面复盘签约、入驻率、授权使用、风险")
    add_p(
        doc,
        "7.4 客户档案：全部报备、带看、报价记录归甲方所有。委托终止后 10 个工作日内完整移交，"
        "乙方可留存副本供合规备查，但不得用于与本项目相竞争的用途。",
        indent=True,
    )

    add_h1(doc, "第八条 · 考核指标（招到什么算完成）")
    add_p(
        doc,
        "8.1 过程与结果指标如下。过程指标用于节奏管理；结果指标用于续约与 M6 退出判断。"
        "未达 M12 目标本身不自动构成根本违约，但影响续约及第九条所述费用安排。",
        indent=True,
    )
    add_grid(
        doc,
        ["时点", "过程指标（必须可见）", "结果指标"],
        [
            [
                "M3",
                "目标库 ≥ 200 家；深度接触 ≥ 8 家；正式带看 ≥ 5 场；周报无断档",
                "至少 3 份有效书面意向",
            ],
            [
                "M6",
                "中介渠道启动并完成报备制度；至少 1 场闭门招商活动",
                "至少 1 份链主/锚定 Term Sheet，或累计有效意向面积 ≥ 2,000 ㎡",
            ],
            [
                "M9",
                "漏斗周会连续召开；重点客户一页纸齐备",
                "累计生效签约面积 ≥ 3,000 ㎡（约 36% 入驻率）",
            ],
            [
                "M12",
                "客户档案与 SOP 可移交",
                "累计生效签约面积 ≥ 3,735 ㎡，或入驻率 ≥ 45%",
            ],
        ],
        col_widths=[2.2, 7.4, 6.4],
    )
    add_p(doc, "")
    add_p(
        doc,
        "8.2 M6 检验：若第 8.1 条 M6 结果指标未达成，且主要原因不可归责于甲方未配合"
        "（见第十三条），甲方有权在 M6 届满后 15 日内书面通知提前终止本协议，"
        "无需支付终止日后的月度操盘费；已发生的成功佣金仍应按约支付；启动费不退。",
        indent=True,
    )
    add_p(
        doc,
        "8.3 因甲方未在授权时限内批复、拒绝提供看房条件、董事长级拜访经书面预约后无故缺席、"
        "或甲方自行否决已符合绿区条件的成交导致指标未达的，不启动 8.2 条退出权，"
        "且甲方不得据此主张乙方未尽责。",
        indent=True,
    )

    add_h1(doc, "第九条 · 服务费用（原口径 · 柔化）")
    add_p(
        doc,
        "9.1 收费方向为：启动费 15 万元 + 月度 8 万元 + 首年租金 8%。"
        "下列柔化只减轻叠加，不改变这三项数字。金额以附件 E 为准。",
        indent=True,
    )

    add_h2(doc, "（一）启动费")
    add_p(
        doc,
        "9.2 启动费人民币 [壹拾伍万元整]。本合同签署后 10 个工作日内一次性支付。"
        "用于进场筹备、材料及第一个月现场服务。启动费可全额抵扣后续应付成功佣金，不退。",
        indent=True,
    )

    add_h2(doc, "（二）月度操盘费")
    add_p(
        doc,
        "9.3 月度操盘费人民币 [捌万元整]/月。自进场日后的第 2 个自然月起付"
        "（第一个月已含在启动费中），每月 5 日前支付当月。"
        "当月若产生应付成功佣金，该月月度费从该笔佣金中抵扣。",
        indent=True,
    )
    add_p(doc, "9.4 启动费与月度费不含：", indent=True)
    add_bullet(doc, "第三方中介居间佣金（按居间协议另付；乙方中介单佣金见 9.9）")
    add_bullet(doc, "闭门会/活动场地与餐饮")
    add_bullet(doc, "样板间装修及工程改造")
    add_bullet(doc, "须由甲方主体出具的政府文件、保证金、税费")

    add_h2(doc, "（三）成功佣金")
    add_p(
        doc,
        "9.5 计算基数 = 生效租赁合同项下承租人应付的首年租金总额"
        "（不含增值税、不含物业费、扣减免租期对应的租金）。",
        indent=True,
    )
    add_p(
        doc,
        "9.6 费率：常规租户与链主/锚定租户一律为基数 × 8%。不上浮。",
        indent=True,
    )
    add_p(
        doc,
        "9.7 支付时点：租赁合同生效且甲方收到履约保证金或首期租金后 15 个工作日内。"
        "启动费抵扣在第一笔成功佣金中优先扣减。",
        indent=True,
    )
    add_p(
        doc,
        "9.8 退租扣回：承租人于起租日后 12 个月内因可归责于招商质量的原因导致合同解除的，"
        "乙方按未履行月数比例退还已收成功佣金。因甲方违约、房屋瑕疵、政策变化、"
        "客户自身战略收缩导致的退租，不扣回。",
        indent=True,
    )
    add_p(
        doc,
        "9.9 中介成交：甲方按居间协议向中介支付居间佣金；"
        "乙方向甲方收取的成功佣金按第 9.6 条的 50% 计（即基数 × 4%），"
        "不按满额 8% 计。未经报备的中介成交，乙方不收取该笔成功佣金。",
        indent=True,
    )
    add_p(
        doc,
        "9.10 税费：本条金额均为未税价，增值税及附加按法定另付，乙方开具合法发票。",
        indent=True,
    )
    add_p(
        doc,
        "9.11 试跑 90 天期满未转正的：之后不产生月度费；已发生成功佣金照付；启动费不退，仍可抵扣已发生佣金。",
        indent=True,
    )

    add_h1(doc, "第十条 · 客户报备与撞单")
    add_p(
        doc,
        "10.1 乙方应以书面《客户报备表》（附件 C）向甲方报备。完整报备（企业全称、决策人、"
        "接触时间、意向面积）后，该客户由乙方锁定 90 日。期满未推进至带看或有效意向的，锁定解除。",
        indent=True,
    )
    add_p(
        doc,
        "10.2 锁定期内该客户成交（含甲方或中介促成），按第 2.3、9.6、9.9 条计付。"
        "同一客户多渠道冲突，以首次完整报备为准。",
        indent=True,
    )
    add_p(
        doc,
        "10.3 禁止虚假报备。经查实虚假报备的，该客户锁定无效，并按第十八条处理。",
        indent=True,
    )

    add_h1(doc, "第十一条 · 甲方配合义务")
    add_p(doc, "甲方应：", indent=True)
    add_bullet(doc, "指定唯一商务对接人，并保证黄区请示 48 小时内有书面答复")
    add_bullet(doc, "提供看房条件、钥匙/门禁、物业陪同、实测图纸与面积表")
    add_bullet(doc, "链主级客户现场拜访，按书面预约安排董事长或授权决策人出席")
    add_bullet(doc, "法务在收到完整合同草案后 10 个工作日内反馈修订意见")
    add_bullet(doc, "不向已由乙方报备锁定的客户另行给出低于绿区的报价")
    add_bullet(doc, "对外政策口径以甲方确认稿为准，不要求乙方口头加码")

    add_h1(doc, "第十二条 · 乙方义务与禁止行为")
    add_p(doc, "12.1 乙方应勤勉、专业、如实汇报漏斗，不得隐瞒重大谈判破裂原因。", indent=True)
    add_p(doc, "12.2 乙方禁止：", indent=True)
    add_bullet(doc, "向承租人或中介收取任何好处费、回扣、顾问费、代币或收益权")
    add_bullet(doc, "超出第六条授权对外承诺租金、免租、装补或政府政策")
    add_bullet(doc, "虚假报备、截留客户、将本项目线索用于竞争性楼宇")
    add_bullet(doc, "未经书面授权以甲方名义签署租赁合同或收款")
    add_bullet(doc, "在甲方不知情时承诺公寓、整车展厅或其他禁止业态")
    add_p(
        doc,
        "12.3 利益冲突：乙方如同时操盘与本标的直接竞争的静安/中心城区同类楼宇，应事先书面披露；"
        "甲方有权决定是否继续独家委托。",
        indent=True,
    )

    add_h1(doc, "第十三条 · 知识产权")
    add_p(
        doc,
        "13.1 为履行本协议专门制作的招商材料、客户档案、报价单、一页纸，著作权归甲方，"
        "仅限用于本项目。",
        indent=True,
    )
    add_p(
        doc,
        "13.2 乙方原有的方法论、模板库、案例研究框架仍归乙方。甲方可在本项目范围内使用。",
        indent=True,
    )
    add_p(
        doc,
        "13.3 公开案例研究双方均可在隐去对方商业秘密后作非商业阐述，但不得暗示对方为对方背书。",
        indent=True,
    )

    add_h1(doc, "第十四条 · 保密")
    add_p(
        doc,
        "14.1 双方对报价底线、客户名单、谈判记录、财务测算互负保密义务，期限至本协议终止后 3 年。",
        indent=True,
    )
    add_p(
        doc,
        "14.2 对外发声（媒体、政府汇报、社交媒体）涉及本项目定位、租金、入驻客户名称的，"
        "须经甲方书面同意。已公开信息除外。",
        indent=True,
    )

    add_h1(doc, "第十五条 · 期限、续约与退出")
    add_p(
        doc,
        "15.1 委托期限 12 个月，自双方盖章之日起算（『进场日』以附件 E 记载为准）。",
        indent=True,
    )
    add_p(
        doc,
        "15.2 期满前 30 日，双方书面确认是否续约 6 个月。续约期月度操盘费、成功佣金按届时附件 E 执行；"
        "如未另签，成功佣金对委托期内已报备并在期满后 90 日内生效的合同仍然适用。",
        indent=True,
    )
    add_p(doc, "15.3 提前终止：", indent=True)
    add_bullet(doc, "按第 8.2 条 M6 检验终止")
    add_bullet(doc, "甲方提前 30 日书面通知无理由终止：须付至通知期满之日的月度费，已发生成功佣金照付，并支付相当于 1 个月月度操盘费的收尾费用")
    add_bullet(doc, "乙方提前 30 日书面通知终止：须完成客户档案移交，当月月度费按日计至移交完成")
    add_bullet(doc, "一方根本违约经书面催告 10 个工作日未改正：守约方可立即终止")
    add_p(
        doc,
        "15.4 终止不影响已经产生的成功佣金、保密、知识产权与争议条款。",
        indent=True,
    )

    add_h1(doc, "第十六条 · 违约责任")
    add_p(
        doc,
        "16.1 甲方逾期支付启动费或月度操盘费超过 15 日的，乙方可书面暂停现场服务；"
        "超过 30 日的，乙方可终止本协议，并要求甲方支付欠费及按欠费日万分之五计的滞纳金。",
        indent=True,
    )
    add_p(
        doc,
        "16.2 乙方违反第 12.2 条禁止行为的，甲方可立即终止本协议，要求返还该行为发生后已付的"
        "月度操盘费，并保留追偿因该行为导致的客户索赔、政府处罚等损失。",
        indent=True,
    )
    add_p(
        doc,
        "16.3 乙方无正当理由连续两周未提交周报、或项目经理连续 10 个工作日无法联系且未指定替手的，"
        "甲方可书面催告；催告后 5 个工作日仍未改正的，可按第 15.3 条立即终止，启动费不退。",
        indent=True,
    )

    add_h1(doc, "第十七条 · 通知与送达")
    add_kv_table(doc, [
        ("甲方收件地址", "[出租方地址]"),
        ("甲方对接邮箱", "[contact@guansong.com]"),
        ("乙方收件地址", "[受托方地址]"),
        ("乙方项目经理邮箱", "[xxx@xxx.com]"),
    ])
    add_p(
        doc,
        "书面通知以专人送达、邮件回执或电子邮件进入对方指定邮箱视为送达。"
        "授权请示、报备、周报均可以电子邮件形式作出。",
        indent=True,
    )

    add_h1(doc, "第十八条 · 争议解决")
    add_p(
        doc,
        "因本协议发生的任何争议，应首先由双方指定对接人友好协商。协商不成的，"
        "提交上海仲裁委员会按其届时有效的仲裁规则仲裁，仲裁地上海，仲裁裁决为终局，"
        "对双方均有约束力。",
        indent=True,
    )

    add_h1(doc, "第十九条 · 附则")
    add_p(doc, "19.1 本协议自双方法定代表人或授权代表签字盖章之日起生效。", indent=True)
    add_p(doc, "19.2 本协议一式四份，双方各执两份，具有同等法律效力。", indent=True)
    add_p(
        doc,
        "19.3 附件为本协议组成部分。附件 E（费用确认单）与正文不一致时，就费用事项以附件 E 为准。",
        indent=True,
    )
    add_p(
        doc,
        "19.4 本协议未尽事宜，由双方另行签订补充协议。补充协议不得改变第二条独家操盘性质，"
        "除非双方明确书面同意改为非独家。",
        indent=True,
    )

    add_h1(doc, "附件清单")
    add_bullet(doc, "附件 A：月度服务清单（乙方每周/每月必须交付的动作）")
    add_bullet(doc, "附件 B：报价授权表（与第六条一致，可单独更新）")
    add_bullet(doc, "附件 C：客户报备表")
    add_bullet(doc, "附件 D：周报必备字段")
    add_bullet(doc, "附件 E：费用与进场日确认单（签署时填写）")

    add_h1(doc, "附件 A · 月度服务清单")
    add_grid(
        doc,
        ["频率", "动作", "交付物", "甲方可见的完成标准"],
        [
            ["每周", "漏斗推进与阻塞升级", "周报", "周五 18:00 前邮件送达，字段齐全"],
            ["每周", "重点客户动作", "接触记录", "至少 5 次有效外联（电话/拜访/带看）"],
            ["每两周", "现场带看或线上路演", "带看纪要", "纪要含异议点与下一步"],
            ["每月", "经营会", "月报 + 会议纪要", "对照第八条指标红黄灯"],
            ["每月", "中介例会（渠道启动后）", "任务量与报备冲突表", "甲方可抽查"],
            ["按需", "绿区报价 / 黄区请示", "报价单或请示单", "不超授权"],
            ["按需", "Term Sheet / 合同商务稿", "Word 草案", "用途条款含 C6 约束"],
        ],
        col_widths=[2.2, 4.0, 4.0, 6.0],
    )

    add_p(doc, "")
    add_h1(doc, "附件 B · 报价授权表")
    add_p(doc, "以第六条表格为准。甲方更新授权时，以盖章或指定邮箱发出的新表替代本附件。")

    add_h1(doc, "附件 C · 客户报备表（字段）")
    add_bullet(doc, "报备时间 / 报备人")
    add_bullet(doc, "企业全称 / 统一社会信用代码（能取得时）")
    add_bullet(doc, "业态细分 / 是否属于禁止业态")
    add_bullet(doc, "决策人姓名职务及联系方式")
    add_bullet(doc, "意向楼层与面积")
    add_bullet(doc, "来源（乙方自拓 / 甲方自有 / 中介名称）")
    add_bullet(doc, "目前阶段（建联 / 带看 / 意向 / TS / 合同）")
    add_bullet(doc, "下次动作与日期")

    add_h1(doc, "附件 D · 周报必备字段")
    add_bullet(doc, "本周新增线索数 / 累计库内企业数")
    add_bullet(doc, "带看场次与客户名")
    add_bullet(doc, "谈判中客户：当前条款 vs 授权区间")
    add_bullet(doc, "本周阻塞点（授权 / 看房 / 政府口径 / 客户内部）")
    add_bullet(doc, "对照第八条：M3/M6/M9/M12 进度红绿灯")
    add_bullet(doc, "下周三件最重要的事（必须具体到人）")

    add_h1(doc, "附件 E · 费用与进场日确认单")
    add_p(doc, "签署时填写，未填写则以正文方括号建议口径为准。")
    add_kv_table(doc, [
        ("进场日", "[YYYY-MM-DD]"),
        ("签法", "□ 方案 B 先 90 天（建议）  □ 方案 A 12 个月独家"),
        ("启动费（未税）", "人民币 [150,000] 元，含第一个月月度，可抵扣成功佣金"),
        ("月度操盘费（未税）", "人民币 [80,000] 元/月，自第 2 个自然月起；当月佣金可抵扣"),
        ("成功佣金", "首年租金基数 × [8]%（链主不上浮）"),
        ("中介成交乙方佣金", "首年租金基数 × [4]%"),
        ("甲方指定对接人", "[姓名 / 手机 / 邮箱]"),
        ("乙方项目经理", "[姓名 / 手机 / 邮箱]"),
        ("最终决策人（红区）", "[董事长或授权人]"),
    ])

    add_signature_block(
        doc,
        ["甲方：上海冠松 [项目运营公司]", "乙方：[招商操盘方全称]"],
    )

    out = OUT_DIR / "05-合作协议-招商服务委托协议.docx"
    doc.save(out)
    print(f"  ✓ {out.name}")
    return out


if __name__ == "__main__":
    print("生成招商服务委托协议：")
    build_leasing_mandate()
    print("✓ 完成")
