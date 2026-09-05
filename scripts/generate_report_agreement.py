# -*- coding: utf-8 -*-
"""给港大经管上海中心审阅的合作协议。按合同书面语写，不写内部分析。"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "output" / "港大经管上海中心_联合策划服务合作协议.docx"

CN_FONT = "宋体"
HEAD_FONT = "黑体"
GREEN = RGBColor(0x00, 0x3D, 0x2E)


def set_run_font(run, font_name: str, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def p(doc, text="", *, font=CN_FONT, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=None, space_after=6, space_before=0, color=None):
    para = doc.add_paragraph()
    para.alignment = align
    pf = para.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.5
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    if text:
        run = para.add_run(text)
        set_run_font(run, font, size, bold=bold, color=color)
    return para


def rich(doc, segments, *, size=12, first_indent=0.74, space_after=6):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.5
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    for text, bold in segments:
        run = para.add_run(text)
        set_run_font(run, CN_FONT, size, bold=bold)
    return para


def h(doc, text):
    return p(doc, text, font=HEAD_FONT, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=8, color=GREEN)


def body(doc, text):
    return p(doc, text, first_indent=0.74)


def shade_cell(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill})
    tcPr.append(shd)


def build(path: Path | None = None) -> Path:
    path = path or OUT
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    normal = doc.styles["Normal"]
    normal.font.name = CN_FONT
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CN_FONT)

    p(doc, "联合策划服务合作协议", font=HEAD_FONT, size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    p(doc, "（供审阅）", font=HEAD_FONT, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14, color=GREEN)

    rich(doc, [("甲方：", True), ("复旦大学住房政策研究中心、上海市杨浦区科技企业联合会", False)], first_indent=0, space_after=2)
    rich(doc, [("甲方指定结算主体：", True), ("【签署时填写】", False)], first_indent=0, space_after=2)
    rich(doc, [("乙方：", True), ("香港大学经管学院（执行机构：香港大学经管学院上海中心）", False)], first_indent=0, space_after=2)
    rich(doc, [("乙方联系人：", True), ("潘嘉琰　　电话：(86) 180 1860 6086　　邮箱：jyanpan@hku.hk", False)], first_indent=0, space_after=2)
    p(doc, "乙方地址：上海市人民路336号外滩SOHO F栋", align=WD_ALIGN_PARAGRAPH.LEFT, space_after=12)

    body(
        doc,
        "甲乙双方拟在上海就乙方EMBA招生宣传、企业出海主题专题课、以及邀请乙方上海中心主任出席相关活动，作为一次合作办理。为明确各自工作、费用和交付事项，订立本协议。本协议签署前供乙方审阅，签署后生效。",
    )

    h(doc, "第一条　合作事项")
    body(
        doc,
        "1.1 双方就乙方EMBA在华东、华中的招生宣传进行合作。甲方按乙方告知的常见门槛准备拟邀请名单（人数不少于八十人），并从中组织三十至四十人到乙方上海中心参加专题课。课后由乙方招生人员与来宾单独沟通。录取、面试、学位由乙方按学院规定办理，甲方不得代为承诺，双方不按学费或录取人数分成。",
    )
    body(
        doc,
        "1.2 上述专题课以企业出海、赴港上市为主题，地点在乙方上海中心（外滩SOHO F栋）。乙方安排场地和老师，甲方负责邀请和现场。本场专题课即为第1.1款招生工作的现场，不另办招生说明会，也不就本场专题课另收费。",
    )
    body(
        doc,
        "1.3 甲方另安排一场出海主题活动（可含总领事或企业负责人出席的场次），请乙方联系人转达，邀请乙方上海中心主任出席。该项与第1.1款、第1.2款同属一次合作，含在第三条费用之内，不另收费、不单开报价。主任是否出席，以乙方安排为准。甲方发出书面邀请、活动如期举办，即完成本项工作。",
    )
    body(
        doc,
        "1.4 港大本部出海课程的推广、原创谷、学生在沪实习参访等事项，双方可另行商议，不列入本协议交付，也不另收费。",
    )
    body(
        doc,
        "1.5 本协议不是招生代理，不按学费或录取人数分成，也不是联合办学。任何一方不得称对方为自己的下属机构，也不得对外宣称双方合办学位。",
    )

    h(doc, "第二条　各自工作")
    body(doc, "2.1 乙方负责：提供上海中心场地和名义；安排老师或校友；书面告知EMBA初筛条件；课后招生跟进；协助邀请中心主任；决定录取和学位。")
    body(doc, "2.2 甲方负责：九十天内的筹备；按门槛准备拟邀请名单；外滩专题课现场；请主任出席的那一场活动；课后七日内提交纪要和有意向人员说明；开具发票、收取费用。")
    body(
        doc,
        "2.3 乙方指定潘嘉琰为日常联系人。甲方指定项目负责人一名，于本协议生效后三个工作日内书面告知乙方。超出招生宣传范围、需中心主任决定的事项，由联系人转达，联系人无义务代为答应。",
    )

    h(doc, "第三条　费用与支付")
    rich(
        doc,
        [
            ("3.1 乙方向甲方指定结算主体支付前期工作费用", False),
            ("人民币捌万捌仟元整（¥88,000）", True),
            ("。该费用为包干，覆盖第一条所列EMBA招生宣传、出海专题课、邀请中心主任出席三项工作，不就其中任何一项单独加收。", False),
        ],
    )
    body(
        doc,
        "3.2 本协议生效后十个工作日内一次付清。甲方收到款项后五个工作日内开具增值税发票。票种和税率以结算主体资质为准，签署前书面确认。",
    )
    body(
        doc,
        "3.3 逾期未付的，甲方可以暂停筹备，日期顺延。逾期超过十五个工作日仍未付清的，甲方可以书面解除本协议，乙方按已发生费用结算，且不少于费用总额的百分之三十。",
    )
    body(
        doc,
        "3.4 费用到账前，甲方没有义务提交完整名单或对外发出正式邀请。双方可以先商定日期。",
    )
    body(
        doc,
        "3.5 本协议范围以外的工作，须另行书面确认。第二场及以后，建议每场人民币陆万捌仟元整，另签确认单后才生效。",
    )
    body(
        doc,
        "3.6 收款账户于签署时填写：开户名称________；开户银行________；账号________；纳税人识别号________。",
    )

    h(doc, "第四条　交付与确认")
    body(
        doc,
        "4.1 甲方提交拟邀请名单初稿，人数不少于八十人，并按乙方告知的EMBA常见门槛初筛。乙方在五个工作日内提出书面意见，逾期视为没有意见。名单供工作使用，不保证每一个人都到场。",
    )
    body(
        doc,
        "4.2 专题课完成后七日内，甲方提交纪要、来宾情况说明，以及有意向继续了解EMBA的人员说明。因不可抗力、乙方改期或乙方未按时确定老师，影响人数的，不视为甲方违约。",
    )
    body(
        doc,
        "4.3 请主任出席一项，以甲方书面邀请为凭。主任未出席，不视为甲方违约，也不因此减收或另收费用。",
    )
    body(
        doc,
        "4.4 录取人数、学费收入不作为确认条件。",
    )

    h(doc, "第五条　名单与名义")
    body(
        doc,
        "5.1 因本协议形成的来宾名单由双方共同使用，不得提供给无关第三方，也不得用于与本次合作无关的推销。本协议结束后十二个月内，仅可用于本次已开始的招生跟进。",
    )
    body(
        doc,
        "5.2 使用对方名称、校徽、场地照片，须事先书面同意，并符合对方规定。港大课程内容和招生办法，以乙方公布的文本为准。",
    )

    h(doc, "第六条　期限、变更与解除")
    body(doc, "6.1 本协议自双方签署且第三条费用到账之日起生效，有效期九十日。期满结束，续做须另签。")
    body(doc, "6.2 日期、题目、老师等安排，可用邮件或盖章确认单变更，不改变第三条费用。")
    body(
        doc,
        "6.3 一方严重违约，书面催告十个工作日仍不改正的，另一方可解除。因乙方原因取消专题课且不同意改期的，费用不退。因甲方原因未能举办专题课、且九十日内无法改期的，甲方应在十五个工作日内退还费用的百分之五十。",
    )
    body(doc, "6.4 因不可抗力不能履行的，双方协商改期；仍不能履行的，按已发生费用结算后解除。")

    h(doc, "第七条　其他")
    body(
        doc,
        "7.1 双方对未公开的信息和名单保密，至本协议结束后两年。法律法规要求披露的除外。活动内容应遵守教育、广告、外事和港澳有关规定。甲方不代理签证，不承诺录取。",
    )
    body(
        doc,
        "7.2 因本协议发生争议，先协商；协商不成的，由甲方指定结算主体所在地有管辖权的人民法院处理。结算主体尚未确定的，由上海市黄浦区人民法院管辖。适用中华人民共和国法律。",
    )
    body(doc, "7.3 本协议一式四份，双方各执两份。电子印章或扫描件与原件相同。未尽事宜可签补充协议。随附的汇报材料仅供说明，与本协议不一致的，以本协议为准。")
    body(doc, "7.4 乙方签署主体以香港大学经管学院认可的有权机构为准。甲方签署时须写明结算主体全称、统一社会信用代码和开票信息。")

    p(doc, "（以下无正文，为签署页）", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=16, space_after=16)

    table = doc.add_table(rows=11, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left = [
        "甲方",
        "复旦大学住房政策研究中心、上海市杨浦区科技企业联合会",
        "指定结算主体：",
        "统一社会信用代码：",
        "授权代表（签字）：",
        "职务：",
        "日期：　　年　　月　　日",
        "联系人：",
        "电话：",
        "开户行：",
        "账号：",
    ]
    right = [
        "乙方",
        "香港大学经管学院",
        "执行机构：香港大学经管学院上海中心",
        "有权签署机构：",
        "授权代表（签字）：",
        "职务：",
        "日期：　　年　　月　　日",
        "联系人：潘嘉琰",
        "电话：(86) 180 1860 6086",
        "邮箱：jyanpan@hku.hk",
        "盖章：",
    ]
    for i, (a, b) in enumerate(zip(left, right)):
        cell0, cell1 = table.cell(i, 0), table.cell(i, 1)
        if i == 0:
            shade_cell(cell0, "003D2E")
            shade_cell(cell1, "003D2E")
        for cell, txt in ((cell0, a), (cell1, b)):
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            run = para.add_run(txt)
            set_run_font(run, CN_FONT, 10.5, bold=(i == 0), color=(RGBColor(0xFF, 0xFF, 0xFF) if i == 0 else None))

    doc.save(path)
    return path


if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else OUT
    print(f"已生成 {build(out)}")
