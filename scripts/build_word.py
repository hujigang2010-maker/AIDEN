"""
生成《上海商办楼宇与产业园区市场深度报告》Word 文档（精装版）
易居研究院 × 复旦大学住房政策研究中心 联合课题组
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree
import os

OUTPUT_DIR = "/workspace/上海商办楼宇与产业园区市场深度报告"
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ======================== Styling helpers ========================

NAVY = (16, 42, 91)
BLUE = (42, 91, 168)
TEAL = (30, 119, 142)
GOLD = (200, 145, 58)
ORANGE = (224, 106, 44)
GRAY = (110, 115, 124)
DARK = (26, 34, 46)
LIGHT_BLUE = (233, 240, 251)


doc = Document()

# 页面设置（横向更宽留白以杂志感）
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)

# 默认正文字体
style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def set_run_font(run, name="宋体", size=11, bold=False, color=None,
                 italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def shade_paragraph(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def set_paragraph_border(paragraph, position="bottom", color="C8913A",
                         size=8):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bd = OxmlElement(f"w:{position}")
    bd.set(qn("w:val"), "single")
    bd.set(qn("w:sz"), str(size))
    bd.set(qn("w:space"), "1")
    bd.set(qn("w:color"), color)
    pBdr.append(bd)


def add_p(text, size=11, indent_first=True, color=DARK, bold=False,
          space_after=4, line_spacing=1.55, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if indent_first:
        p.paragraph_format.first_line_indent = Pt(22)
    if align is not None:
        p.alignment = align
    return p


def add_centered(text, size=14, bold=False, color=GRAY, font="黑体",
                 space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, name=font, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)


def add_cover_title(text, size=34, color=NAVY, space_after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, name="黑体", size=size, bold=True, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(4)


def add_h1(text):
    p = doc.add_paragraph()
    # 段前色块感
    r0 = p.add_run("  ")
    set_run_font(r0, name="黑体", size=18)
    r = p.add_run(text)
    set_run_font(r, name="黑体", size=20, bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(10)
    set_paragraph_border(p, "bottom", color="C8913A", size=12)
    return p


def add_h2(text):
    p = doc.add_paragraph()
    r = p.add_run("▎ ")
    set_run_font(r, name="黑体", size=14, bold=True, color=GOLD)
    r2 = p.add_run(text)
    set_run_font(r2, name="黑体", size=14, bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def add_h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, name="黑体", size=12, bold=True, color=BLUE)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_bullet(text, level=0, color=DARK, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.45
    r1 = p.add_run("●  ")
    set_run_font(r1, name="黑体", size=9, color=BLUE, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color)


def add_table(headers, rows, col_widths=None, header_fill="102A5B",
              alt_fill="E9F0FB"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, name="黑体", size=10, bold=True, color=(255, 255, 255))
        shade_cell(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if ci == 0
                           else WD_ALIGN_PARAGRAPH.LEFT)
            r = p.add_run(str(val))
            set_run_font(r, size=10, color=DARK)
            if ri % 2 == 1:
                shade_cell(cell, alt_fill)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_callout(text, fill="E9F0FB", color=NAVY, bold=True, size=11):
    """ 高亮提示框 """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shade_cell(cell, fill)
    # set left border accent
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("left",):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "24")
        b.set(qn("w:color"), "C8913A")
        tc_borders.append(b)
    tc_pr.append(tc_borders)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def page_break():
    doc.add_page_break()


def add_section_divider(part_num, en, cn, accent_color=BLUE):
    """ 章节大隔页 """
    page_break()
    # 留白
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"PART  {part_num}")
    set_run_font(r, name="黑体", size=22, bold=True, color=GOLD)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    r = p.add_run(en)
    set_run_font(r, name="黑体", size=12, color=GRAY)
    p.paragraph_format.space_after = Pt(20)

    p = doc.add_paragraph()
    r = p.add_run(cn)
    set_run_font(r, name="黑体", size=32, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(10)
    set_paragraph_border(p, "bottom", color="C8913A", size=18)


def add_header_footer():
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("上海商办楼宇与产业园区市场深度报告")
    set_run_font(r, size=9, color=NAVY)
    tab = p.add_run("\t\t")
    set_run_font(tab, size=9)
    r2 = p.add_run("易居研究院 × 复旦大学住房政策研究中心")
    set_run_font(r2, size=9, color=GRAY)
    set_paragraph_border(p, "bottom", color="D8DEEA", size=4)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("— ")
    set_run_font(r, size=9, color=GRAY)
    # page number field
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = fp.add_run()
    set_run_font(run, size=10, bold=True, color=NAVY)
    run._r.append(fld_begin); run._r.append(instr)
    run._r.append(fld_sep); run._r.append(fld_end)
    r2 = fp.add_run(" —")
    set_run_font(r2, size=9, color=GRAY)


add_header_footer()


# ============================== 封面 ==============================

# 顶部金色装饰
top_p = doc.add_paragraph()
top_p.paragraph_format.space_after = Pt(0)
set_paragraph_border(top_p, "bottom", color="C8913A", size=36)

for _ in range(3):
    doc.add_paragraph()

# JOINT REPORT 小标识
add_centered("JOINT  REPORT  ·  2026", size=12, color=GOLD,
             font="黑体", space_after=2)
add_centered("易居研究院  ×  复旦大学住房政策研究中心", size=14,
             color=GRAY, space_after=20)

# 主标题
add_cover_title("上海商办楼宇与产业园区", size=36, color=NAVY,
                space_after=0)
add_cover_title("市  场  深  度  报  告", size=40, color=NAVY,
                space_after=20)

# 装饰横线
mid = doc.add_paragraph()
mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_border(mid, "bottom", color="C8913A", size=18)
r = mid.add_run("                ")
set_run_font(r, size=12)

add_centered("——  全域空间供给 · 企业需求迁徙 · 产业载体运营研究  ——",
             size=15, color=BLUE, font="黑体", space_after=4)
add_centered(
    "Shanghai Commercial Office & Industrial Park Market — In-depth Report",
    size=10, color=GRAY, font="黑体", space_after=24)

# 研发内容
add_centered("研发方案  ·  报告大纲  ·  数据调用建议  ·  执行指引",
             size=14, color=NAVY, font="黑体", space_after=4)
add_centered("（讨论稿）", size=12, color=GRAY, font="宋体",
             space_after=40)

# 底部主体
for _ in range(2):
    doc.add_paragraph()
add_centered("联合发布主体", size=11, color=GRAY, font="宋体",
             space_after=4)
add_centered("易居房地产研究院", size=16, bold=True, color=NAVY,
             font="黑体", space_after=2)
add_centered("×", size=14, bold=True, color=GOLD, font="黑体",
             space_after=2)
add_centered("复旦大学住房政策研究中心", size=16, bold=True, color=NAVY,
             font="黑体", space_after=14)
add_centered("联合课题组：上海商办楼宇与产业园区市场研究课题组",
             size=11, color=GRAY, font="宋体", space_after=4)

# 底部金色装饰
bottom_p = doc.add_paragraph()
bottom_p.paragraph_format.space_before = Pt(40)
set_paragraph_border(bottom_p, "top", color="C8913A", size=36)


# ============================== 目录 ==============================
page_break()
add_h1("目录  ·  Contents")
toc_items = [
    ("前言", "联合发布与项目背景"),
    ("第一章", "上海商办楼宇与产业园区市场发展背景"),
    ("第二章", "上海全域商办楼宇与产业园区供给格局"),
    ("第三章", "上海商办与产业园区租金、空置和价格体系"),
    ("第四章", "入驻企业画像与产业需求结构"),
    ("第五章", "企业迁徙与产业流动趋势"),
    ("第六章", "供需匹配与招商机会分析"),
    ("第七章", "市场趋势预测"),
    ("第八章", "政策建议与市场应用"),
    ("第九章", "成果体系与商业化延展"),
    ("附篇一", "数据体系与字段设计（六大数据库）"),
    ("附篇二", "核心指标体系（五大指数）"),
    ("附篇三", "数据调用、采集与合规建议"),
    ("附篇四", "分工协同机制与工作流程"),
    ("附篇五", "首期试点报告落地建议"),
    ("附篇六", "关键难点与解决方案"),
    ("附篇七", "最终成果包与商业化产品矩阵"),
    ("结语", "研究价值与长期愿景"),
]
for tag, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.8
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT,
                           WD_TAB_LEADER.DOTS)
    r0 = p.add_run(tag + "  ")
    set_run_font(r0, name="黑体", size=11, bold=True, color=GOLD)
    r = p.add_run(title)
    set_run_font(r, size=11, color=NAVY)
    p.add_run("\t")
    rn = p.add_run("· · ·")
    set_run_font(rn, size=10, color=GRAY)


# ============================== 前言 ==============================
add_section_divider("前言", "PREFACE", "联合发布与项目背景")

add_h1("前言：联合发布与项目背景")
add_h2("一、项目定位")

add_h3("1. 报告名称")
add_p("正式名称：《上海商办楼宇与产业园区市场深度报告》。")
add_p("副标题（可根据发布场景选择）：")
add_bullet("全域空间供给、企业需求迁徙与产业载体运营研究；")
add_bullet("上海商办与产业空间供需格局、企业流动及资产运营决策报告。")
add_p("联合发布主体：易居房地产研究院 × 复旦大学住房政策研究中心。可根据合作结构挂"
      "“联合课题组：上海商办楼宇与产业园区市场研究课题组”。")

add_h3("2. 报告核心定位")
add_p("本报告拟打造成为一份面向上海全域商办楼宇与产业园区市场的系统性研究成果，重点突"
      "破传统市场报告只关注核心商务区甲级写字楼、只分析租金和空置率的局限，转向：")
for t in [
    "全域覆盖：覆盖上海16区，逐步细化至街道、镇级颗粒度；",
    "全类型载体覆盖：写字楼、产业园区、科创园区、孵化器、商务园区等；",
    "供需双侧研究：既研究楼宇园区供给，也研究企业入驻、搬迁、扩租、缩租需求；",
    "产业视角嵌入：分析产业集聚、产业迁徙、产业补链机会；",
    "招商和资管导向：最终服务政府招商、园区运营、资产盘活、企业选址和投资决策。",
]:
    add_bullet(t)

add_h3("3. 报告目标读者")
add_table(
    ["读者类别", "典型角色", "核心使用场景"],
    [
        ["政府与产业招商部门",
         "市/区两级招商主管、街镇招商中心、开发区/功能区管委会",
         "空间治理 · 精准招商 · 楼宇与园区经济政策制定"],
        ["楼宇和园区运营方",
         "商办业主、园区运营商、城市更新操盘方、国企平台公司",
         "资产定位 · 租金策略 · 招商策略 · 运营服务设计"],
        ["资产持有方与投资机构",
         "REITs/类REITs资产方、商办投资人、资管公司、地产基金",
         "竞品对标 · 估值 · 改造 · 资本化路径"],
        ["企业选址与扩张决策方",
         "科创/专精特新/总部/生产性服务业企业",
         "选址 · 扩租 · 缩租 · 迁址决策"],
        ["行业研究与咨询机构",
         "房地产研究机构、产业咨询、金融研究部门、城市更新服务商",
         "市场对标 · 行业研究 · 政策咨询"],
    ],
    col_widths=[3.4, 5.6, 6.4],
)

add_h2("二、项目背景与时代意义")
for t in [
    "当前上海正处于城市更新向纵深推进、新质生产力加快培育、“五个中心”建设能级持续提升的"
    "重要阶段。商办楼宇与产业园区作为承载产业经济、集聚创新要素、支撑城市功能的重要空间"
    "载体，其市场运行逻辑已发生深刻变化。",
    "过去研究多集中于核心商务区甲级写字楼，关注租金、空置率、新增供应等传统资产指标。"
    "随着城市发展进入存量时代，市场关注点正从“增量开发”转向“存量盘活”，从“资产租赁”"
    "转向“产业运营”，从“楼宇价值”转向“企业需求、产业集聚和空间效率”。",
    "企业需求端正在系统性重构。办公选址不再单纯追求地段、形象和总部标签，而是更加关注综"
    "合成本、产业协同、政策赋能、空间灵活性、人才可达性和运营服务能力。",
    "在此背景下，研发一份覆盖上海全域、兼顾供需两侧、嵌入产业视角、面向招商与资管决策的"
    "深度报告，具有明显的战略意义与现实价值。",
]:
    add_p(t)

add_callout("本报告拟打造为：上海商办楼宇与产业园区市场年度旗舰研究产品 + "
            "行业标准化研究样本 + 城市空间治理决策参考")

add_h2("三、报告核心价值")
add_table(
    ["维度", "核心价值"],
    [
        ["政府端", "摸清存量底数；识别空置/招商压力与供需错配；追踪企业迁徙、产业集聚和外溢；"
                  "支撑城市更新、楼宇经济政策制定与精准招商"],
        ["行业端", "打破“重核心区、轻全域；重供给、轻需求；重静态、轻动态”的局限，"
                  "建立上海商办与产业园区市场的标准化研究框架"],
        ["资产端", "判断区域竞争格局；对标竞品；明确目标产业与企业画像；"
                  "优化租金、招商、运营、改造与资本化策略"],
        ["企业端", "比较综合承租成本；判断产业生态与政策环境；识别适配空间；"
                  "降低选址风险，提升空间配置效率"],
        ["联合课题品牌", "形成“专业市场研究 + 高校智库 + 数据技术能力”的组合优势，"
                       "打造行业影响力的长期研究品牌"],
    ],
    col_widths=[3.0, 12.4],
)


# ============================== 第一章 ==============================
add_section_divider("01", "MARKET  CONTEXT",
                    "上海商办与产业空间逻辑正在重构")

add_h1("第一章  上海商办楼宇与产业园区市场发展背景")

add_h2("1.1 上海城市发展阶段变化")
for t in [
    "上海城市更新进入深水区，存量空间盘活成为城市发展的核心命题；",
    "新质生产力培育对产业空间提出新的承载要求；",
    "“五个中心”建设推动商务办公、科创研发、总部经济和产业服务能级升级；",
    "传统增量开发逻辑向存量盘活、产业赋能、精细运营转变；",
    "商办楼宇与产业园区从“空间载体”升级为“产业生态平台”。",
]:
    add_bullet(t)

add_h2("1.2 商办与产业园区市场运行逻辑变化")
for t in [
    "从增量开发转向存量运营；",
    "从租金竞争转向产业服务竞争；",
    "从地段价值转向综合生态价值；",
    "从单一办公功能转向办公、研发、展示、交流、孵化、资本服务复合功能；",
    "从楼宇招商转向产业招商和企业生命周期服务；",
    "写字楼与产业园区边界趋于模糊，CBD、高新区、产业社区、城市更新片区进入新的竞合格局。",
]:
    add_bullet(t)

add_h2("1.3 企业办公与产业空间需求变化")
add_p("企业选址逻辑从“地段优先、形象优先、总部标签优先”，转向“成本可控、产业协同、政策"
      "支持、交通便利、空间灵活、服务完善、人才可达”。")
add_table(
    ["企业类型", "需求重点"],
    [
        ["科创企业", "政策、人才、上下游生态、孵化与资本服务"],
        ["专精特新企业", "成本控制、技术平台、行业资质、政策兑现"],
        ["总部企业", "品牌展示、资源链接、稳定性、运营效率"],
        ["生产性服务业", "客户可达性、交通枢纽、商务配套"],
        ["硬科技/制造企业", "复合型研发办公、测试与厂办结合、产业链协同"],
        ["中小成长企业", "灵活面积、灵活租期、租金可承受、政策可达"],
    ],
    col_widths=[4.0, 11.4],
)

add_h2("1.4 当前市场研究体系的不足")
for t in [
    "研究范围偏窄：多聚焦核心商务区甲级写字楼，对乙级写字楼、商务园区、产业园区、孵化器、"
    "街镇载体覆盖不足；",
    "数据颗粒度不足：多以板块为单位，缺乏街道、镇、楼宇、园区、企业层面的精细化数据；",
    "需求端研究不足：对企业行业、规模、成长阶段、迁徙、扩租缩租等行为研究不够；",
    "落地应用不足：大量市场报告难以转化为招商清单、运营动作、资产策略和企业选址依据。",
]:
    add_bullet(t)

add_h2("1.5 本报告的研究突破")
for t in [
    "从核心区样本走向上海全域覆盖；",
    "从写字楼走向商办楼宇与产业园区全类型载体；",
    "从租金空置走向供需双侧研究；",
    "从静态市场描述走向企业迁徙和产业流动监测；",
    "从宏观研究走向政府招商、资产运营和企业选址应用。",
]:
    add_bullet(t)


# ============================== 第二章 ==============================
add_section_divider("02", "SUPPLY  LANDSCAPE",
                    "上海全域商办楼宇与产业园区供给格局")

add_h1("第二章  上海全域商办楼宇与产业园区供给格局")

add_h2("2.1 全市商办与产业载体总体规模")
add_p("建议建立全市基础底数指标体系，主要包括：")
for t in [
    "全市商办楼宇数量、产业园区数量、经营性办公及产业载体数量；",
    "总建筑面积、可租赁面积、当前可招商面积；",
    "新增供应面积、存量载体面积；",
    "载体分布密度（每平方公里载体数）及行政区/街镇结构。",
]:
    add_bullet(t)

add_h2("2.2 按行政区划分的供给格局（16区）")
add_p("以上海16区为一级统计单元，每个区分析：载体数量、建筑面积、可租赁面积、代表性楼"
      "宇/园区、主要物业类型、主导产业、租金水平、空置压力、招商活跃度、未来供应压力。")
add_table(
    ["序号", "行政区", "功能定位概览（示例）"],
    [
        ["01", "黄浦区", "传统CBD、金融与专业服务、消费与文创"],
        ["02", "静安区", "高端商务、总部经济、消费服务"],
        ["03", "徐汇区", "AI、传媒、滨江总部"],
        ["04", "长宁区", "虹桥涉外商务、数字贸易"],
        ["05", "普陀区", "数字经济、城市更新载体"],
        ["06", "虹口区", "北外滩金融与航运服务"],
        ["07", "杨浦区", "在线新经济、AI、科创服务"],
        ["08", "浦东新区", "陆家嘴、前滩、张江、金桥、外高桥、临港"],
        ["09", "闵行区", "虹桥商务、紫竹、生命健康"],
        ["10", "宝山区", "南大智慧城、新材料"],
        ["11", "嘉定区", "智能汽车、高端制造"],
        ["12", "松江区", "G60科创走廊、高端制造"],
        ["13", "青浦区", "西虹桥、长三角一体化示范区"],
        ["14", "奉贤区", "东方美谷、临港南桥"],
        ["15", "金山区", "化工、新材料、临港金山"],
        ["16", "崇明区", "生态产业、海洋经济"],
    ],
    col_widths=[1.5, 3.0, 10.9],
)

add_h2("2.3 按街道/镇维度进行精细化拆解")
add_p("建议构建“行政区 — 街道/镇 — 楼宇/园区 — 企业”的四级数据库。这是报告的差异化亮"
      "点。街道/镇层面重点统计：")
for t in [
    "楼宇园区数量、载体总面积；",
    "空置面积、平均租金、入驻企业数量；",
    "主导产业、企业净流入情况；",
    "招商活跃度、供需匹配度。",
]:
    add_bullet(t)

add_callout("四级数据库：行政区 → 街道/镇 → 楼宇/园区 → 企业 — 报告核心差异化亮点")

add_h2("2.4 按物业类型划分供给结构")
add_table(
    ["物业类型", "典型租金区间（示意）", "典型客户群"],
    [
        ["超甲级写字楼", "高", "外资总部、金融机构、专业服务"],
        ["甲级写字楼", "中高", "总部企业、金融、专业服务、外资"],
        ["乙级写字楼", "中", "成长型企业、生产性服务业"],
        ["普通商务楼宇", "中低", "中小企业、贸易类、本地服务业"],
        ["总部办公园区", "中高", "总部企业、龙头企业"],
        ["科创园区", "中", "科创企业、研发型企业"],
        ["产业园区", "低-中", "高端制造、硬科技、生物医药等"],
        ["孵化器/众创空间", "低", "初创企业、专精特新培育"],
        ["城市更新改造载体", "弹性", "主题型、垂直型产业楼宇"],
        ["国企平台持有载体", "弹性", "区属产业承载、招商配套"],
    ],
    col_widths=[4.0, 3.5, 7.9],
)

add_h2("2.5 重点商务与产业板块")
add_p("建议在区/街镇基础上选择如下重点板块进行专项分析：")
panels = ["陆家嘴", "前滩", "张江科学城", "临港新片区", "徐汇滨江", "北外滩",
          "大虹桥", "杨浦滨江", "五角场", "南京西路", "人民广场", "漕河泾",
          "紫竹", "金桥", "外高桥", "嘉定汽车城", "松江G60科创走廊",
          "青浦西虹桥", "宝山南大智慧城", "奉贤东方美谷"]
add_p("、".join(panels) + "。共 20 大重点板块。")
add_p("每个板块建议分析：板块定位、核心载体、租金区间、空置水平、主导产业、企业结构、产"
      "业集聚度、未来供应压力、招商机会。")


# ============================== 第三章 ==============================
add_section_divider("03", "RENT  ·  VACANCY  ·  COST",
                    "上海商办与产业园区租金、空置和价格体系")

add_h1("第三章  上海商办与产业园区租金、空置和价格体系")

add_h2("3.1 全市租金水平分析（分层统计）")
for t in [
    "全市平均报价租金、全市平均成交租金；",
    "核心商务区/次核心商务区/产业园区/乙级及以下楼宇租金；",
    "16区租金对比、不同物业等级租金对比；",
    "重点板块租金对比与板块内细分租金。",
]:
    add_bullet(t)

add_h2("3.2 租金分化特征")
for t in [
    "核心区租金是否继续承压；",
    "非核心区租金压力是否扩散；",
    "产业园区租金韧性；",
    "新兴板块租金成长性；",
    "高品质楼宇与普通楼宇价差；",
    "同一区域内不同楼宇间的价格分化。",
]:
    add_bullet(t)

add_h2("3.3 空置率分析")
for t in [
    "全市整体空置率、各区/各街镇空置率；",
    "各物业类型空置率；",
    "重点板块空置率；",
    "重点楼宇/园区空置表现；",
    "新增供应导致的去化压力测算。",
]:
    add_bullet(t)

add_h2("3.4 空置面积结构分析")
add_table(
    ["结构维度", "释义", "招商资管含义"],
    [
        ["小面积空置", "≤300㎡", "适配初创、小微团队"],
        ["中面积空置", "300-1000㎡", "适配成长型、中型企业"],
        ["大面积连续空置", "1000-3000㎡", "适配总部、研发中心"],
        ["整层空置", "单层连续可用", "整层招商溢价潜力"],
        ["多层连续空置", "≥2层", "适配总部、产业基地"],
        ["整栋待招商", "全楼可整体租赁/合作", "适合产业主题改造、联合运营"],
        ["长期空置", "≥12个月未去化", "需重新定位或改造"],
        ["新增空置", "近3个月新增", "短期波动信号"],
    ],
    col_widths=[3.0, 4.5, 7.9],
)

add_h2("3.5 企业综合承租成本分析")
add_callout("综合承租成本 = 租金 + 物业费 + 停车费 + 装修摊销 + 搬迁成本 + "
            "通勤成本 − 政策补贴",
            fill="102A5B", color=(255, 255, 255))
for t in [
    "名义租金与实际成本的差异；",
    "不同区域企业综合承租成本比较；",
    "写字楼与产业园区成本差异；",
    "政策补贴对企业选址的影响；",
    "企业降本型迁移的主要路径。",
]:
    add_bullet(t)


# ============================== 第四章 ==============================
add_section_divider("04", "DEMAND  &  ENTERPRISES",
                    "入驻企业画像与产业需求结构")

add_h1("第四章  入驻企业画像与产业需求结构")

add_h2("4.1 全市入驻企业总体分布")
for t in [
    "各区/街镇/楼宇/园区入驻企业数量；",
    "重点楼宇入驻企业密度、产业园区企业密度；",
    "企业行业分布、规模分布、成立年限分布。",
]:
    add_bullet(t)

add_h2("4.2 行业结构分析（17 大重点行业）")
industries = ["人工智能", "集成电路", "生物医药", "智能驾驶", "新能源汽车",
              "机器人", "低空经济", "软件信息服务", "金融服务", "专业服务",
              "文化传媒", "数字贸易", "跨境电商", "高端制造",
              "生产性服务业", "总部经济", "专精特新企业"]
add_p("、".join(industries) + "。")

add_h2("4.3 不同行业空间需求特征（示例）")
add_table(
    ["行业", "典型空间需求", "选址偏好", "关注因素"],
    [
        ["人工智能", "中大型办公 + 研发展示", "徐汇滨江·杨浦·张江·临港",
         "人才/算力/资本/政策"],
        ["生物医药", "研发办公 + 实验空间", "张江·临港·闵行·奉贤",
         "实验条件/审批/产业链"],
        ["智能驾驶", "办公 + 研发 + 测试", "嘉定·浦东·杨浦·临港",
         "测试场景/整车/政策"],
        ["专业服务", "中小型高品质办公", "黄浦·静安·陆家嘴·北外滩",
         "客户可达/品牌"],
        ["高端制造", "研发办公 + 厂办", "嘉定·松江·临港·宝山",
         "成本/物流/产业配套"],
        ["金融服务", "中大型品牌办公", "陆家嘴·外滩·前滩",
         "监管/品牌/客户"],
    ],
    col_widths=[2.5, 3.5, 4.6, 4.8],
)

add_h2("4.4 企业规模与空间需求")
for t in [
    "初创企业、成长期企业、中型企业、大型企业；",
    "总部型企业、上市公司、拟上市企业；",
    "专精特新企业、高新技术企业；",
    "外资企业、国央企子公司。",
]:
    add_bullet(t)
add_p("分析维度：典型面积需求、租期偏好、租金承受能力、区域偏好、扩租潜力、迁址可能性、"
      "政策支持需求。")

add_h2("4.5 重点企业清单体系（9 类）")
for t in [
    "高成长企业清单、近期融资企业清单；",
    "专精特新企业清单、高新技术企业清单；",
    "上市及拟上市企业清单；",
    "迁址可能性企业清单、扩租潜力企业清单、缩租风险企业清单；",
    "跨区迁移企业清单、重点产业链补链企业清单。",
]:
    add_bullet(t)


# ============================== 第五章 ==============================
add_section_divider("05", "MIGRATION  &  FLOW",
                    "企业迁徙与产业流动趋势")

add_h1("第五章  企业迁徙与产业流动趋势")

add_h2("5.1 企业迁徙总体格局")
for t in [
    "企业从哪里迁出、迁到哪里、迁徙时间；",
    "所属行业、企业规模；",
    "面积变化、成本变化；",
    "迁徙原因推测（扩张/降本/总部升级/产业集聚/政策导向/被动搬迁/缩租/注册地变更）。",
]:
    add_bullet(t)

add_h2("5.2 跨行政区迁徙分析")
for t in [
    "核心区企业向次核心区或郊区迁移；",
    "郊区企业向核心功能区集聚；",
    "成长型企业向产业集聚区迁移；",
    "降本型 / 升级型 / 政策导向型 / 被动搬迁型迁移分别识别。",
]:
    add_bullet(t)

add_h2("5.3 重点产业迁徙规律")
add_table(
    ["产业", "重点流向", "代表板块"],
    [
        ["人工智能", "向滨江与科创带集聚", "徐汇滨江、杨浦、张江、临港"],
        ["智能驾驶", "整车 + 软件协同布局", "嘉定汽车城、浦东、杨浦、临港"],
        ["生物医药", "研发与制造分离布局", "张江、临港、闵行、奉贤"],
        ["集成电路", "向产业链核心区集聚", "张江、临港、嘉定"],
        ["机器人/硬科技", "向成本与产业空间外溢", "嘉定、松江、宝山、临港"],
        ["金融/专业服务", "仍偏好核心区", "陆家嘴、北外滩、南京西路"],
        ["文化传媒", "向滨江与新兴板块迁移", "徐汇滨江、杨浦滨江"],
        ["总部型企业", "向品牌资产/滨江总部带集聚",
         "陆家嘴、前滩、北外滩、徐汇滨江"],
    ],
    col_widths=[3.0, 5.5, 6.9],
)

add_h2("5.4 区域吸引力研判")
add_callout("区域吸引力指数 = f(企业净流入 + 重点产业流入 + 高成长企业流入 + "
            "租金性价比 + 政策支持 + 交通可达 + 产业配套 + 空间适配 + 资本服务)")
for t in [
    "企业净流入数量、重点产业企业流入、高成长企业流入；",
    "租金性价比、政策强度；",
    "交通可达性、产业配套成熟度；",
    "空间供给适配度、资本和服务机构集聚度。",
]:
    add_bullet(t)

add_h2("5.5 企业流失风险区域识别")
for t in [
    "企业迁出数量较多区域；",
    "空置率持续上升区域；",
    "租金竞争力不足、产业定位模糊、老旧楼宇集中、政策与服务支撑不足的区域。",
]:
    add_bullet(t)


# ============================== 第六章 ==============================
add_section_divider("06", "MATCHING  &  OPPORTUNITIES",
                    "供需匹配与招商机会分析")

add_h1("第六章  供需匹配与招商机会分析")

add_h2("6.1 载体供给与企业需求匹配")
for t in [
    "面积匹配、租金匹配；",
    "产业匹配、政策匹配；",
    "区位匹配、空间功能匹配；",
    "企业成长阶段匹配。",
]:
    add_bullet(t)

add_h2("6.2 重点楼宇竞争力评价模型（7 维）")
add_table(
    ["评价维度", "核心指标", "建议权重"],
    [
        ["区位交通", "地铁距离、主干路、机场高铁可达性", "15%"],
        ["资产品质", "建筑品质、标准层面积、层高、装修、智能化", "15%"],
        ["租金表现", "报价租金、成交租金、租金弹性", "15%"],
        ["企业结构", "入驻企业质量、行业集中度、龙头数量", "15%"],
        ["运营服务", "物业服务、企业服务、活动运营", "15%"],
        ["政策资源", "补贴政策、产业资质、招商主管", "10%"],
        ["去化能力", "空置率、成交周期、租户稳定性", "15%"],
    ],
    col_widths=[3.5, 8.4, 3.5],
)

add_h2("6.3 产业园区竞争力评价")
for t in [
    "主导产业清晰度、产业链完整度；",
    "孵化服务能力、公共技术平台；",
    "研发空间适配度、政策兑现能力；",
    "产业基金/资本资源、校企合作、产学研协同。",
]:
    add_bullet(t)

add_h2("6.4 招商机会识别（三类机会）")
add_h3("1) 区域机会")
for t in [
    "哪些区域适合承接核心区外溢企业；",
    "哪些区域适合发展 AI、智能驾驶、生物医药等新兴产业；",
    "哪些区域适合导入总部企业；",
    "哪些区域适合做成本型办公承接。",
]:
    add_bullet(t)
add_h3("2) 产业机会")
for t in [
    "AI产业链补链机会；",
    "智能驾驶上下游企业导入机会；",
    "生物医药研发服务企业导入机会；",
    "专精特新企业集聚机会。",
]:
    add_bullet(t)
add_h3("3) 企业线索机会")
for t in [
    "近期融资企业；",
    "新增招聘明显企业；",
    "注册地址变更企业；",
    "扩租可能企业、迁址可能企业；",
    "产业链上下游企业。",
]:
    add_bullet(t)

add_h2("6.5 资产运营建议")
for t in [
    "租金调整建议、招商产业定位建议；",
    "空间改造建议、企业服务补强建议；",
    "政策资源嫁接建议、运营活动建议；",
    "资产品牌重塑建议、资本化与证券化准备建议。",
]:
    add_bullet(t)


# ============================== 第七章 ==============================
add_section_divider("07", "TREND  FORECAST",
                    "市场趋势预测")

add_h1("第七章  市场趋势预测")

add_h2("7.1 租金趋势预测")
for t in [
    "核心商务区租金走势、次核心区租金走势；",
    "乙级写字楼、产业园区、新兴板块租金走势；",
    "老旧楼宇租金调整压力。",
]:
    add_bullet(t)

add_h2("7.2 空置率趋势预测")
for t in [
    "高空置板块、新增供应压力板块；",
    "去化较快板块、产业需求支撑较强板块；",
    "存量改造压力较大板块。",
]:
    add_bullet(t)

add_h2("7.3 产业需求趋势")
add_p("重点关注：人工智能、智能驾驶、机器人、低空经济、生物医药、集成电路、数字贸易、跨"
      "境电商、绿色低碳、科技服务、生产性服务业。")

add_h2("7.4 企业空间策略变化")
for t in [
    "企业倾向更灵活的办公面积；",
    "总部企业重视品牌与资源链接；",
    "成长型企业更关注政策和成本；",
    "硬科技企业更需要复合型产业空间；",
    "部分传统服务业继续缩租或降本迁移；",
    "科创企业对“产业社区+办公+研发+服务”需求增强。",
]:
    add_bullet(t)

add_h2("7.5 未来市场机会")
add_table(
    ["机会场域", "机会方向"],
    [
        ["核心区", "总部经济、金融、专业服务、外资机构"],
        ["滨江区", "AI、数字经济、总部展示、科技服务"],
        ["张江/临港", "硬科技、生物医药、高端制造、智能驾驶"],
        ["郊区产业空间", "成本型外溢、研发制造结合、产业链集聚"],
        ["老旧商办更新", "主题楼宇、垂直产业楼宇、科创服务空间"],
    ],
    col_widths=[3.5, 11.9],
)


# ============================== 第八章 ==============================
add_section_divider("08", "POLICY  &  APPLICATION",
                    "政策建议与市场应用")

add_h1("第八章  政策建议与市场应用")

add_h2("8.1 对政府部门的建议")
for t in [
    "建立上海全域商办与产业空间动态数据库；",
    "以街道/镇为单位监测楼宇空置和企业迁徙；",
    "将企业迁徙数据纳入产业空间治理体系；",
    "对高空置区域实施精准招商和功能更新；",
    "推动楼宇经济、园区经济和城市更新联动；",
    "建立产业空间分类治理机制；",
    "以数据支撑产业政策精准施策。",
]:
    add_bullet(t)

add_h2("8.2 对园区和楼宇运营方的建议")
for t in [
    "从“租赁招商”转向“产业招商”；",
    "建立企业画像和招商漏斗；",
    "动态监测竞品租金和空置；",
    "优化租金和免租策略；",
    "引入企业服务、政策服务和资本服务；",
    "强化产业主题包装、对老旧载体进行场景化改造。",
]:
    add_bullet(t)

add_h2("8.3 对资产持有方的建议")
for t in [
    "重新评估资产定位、建立竞品对标体系；",
    "关注真实成交租金而非单一挂牌租金；",
    "通过产业主题提升资产溢价；",
    "对高空置资产采取降价、改造、联合运营组合策略；",
    "为资产证券化、REITs或资本运作储备数据基础。",
]:
    add_bullet(t)

add_h2("8.4 对企业选址方的建议")
for t in [
    "建立综合承租成本模型；",
    "综合比较租金、政策、交通、产业生态和人才可达性；",
    "根据企业发展阶段选择不同类型空间；",
    "成长型企业关注扩租弹性；总部型企业关注品牌与资源链接；",
    "研发型企业关注空间功能与产业配套。",
]:
    add_bullet(t)


# ============================== 第九章 ==============================
add_section_divider("09", "DELIVERABLES  MATRIX",
                    "成果体系与商业化延展")

add_h1("第九章  成果体系与商业化延展")
add_p("建议将本项目从单份报告升级为长期产品体系：")

add_h3("9.1 月度市场监测简报")
add_p("租金/空置/企业迁徙/热门板块/招商机会/市场预警。")

add_h3("9.2 季度深度报告")
add_p("区域市场分析、产业需求分析、企业迁徙分析、资产运营建议、趋势预测。")

add_h3("9.3 年度白皮书")
add_p("全域格局 + 年度租金空置 + 年度迁徙趋势 + 年度产业演化 + 招商机会 + 政策建议。")

add_h3("9.4 区域专项报告")
for t in [
    "《杨浦区商办楼宇与产业园区市场报告》；",
    "《浦东新区产业空间与企业迁徙报告》；",
    "《徐汇滨江AI产业空间研究报告》；",
    "《嘉定智能汽车产业载体研究报告》；",
    "《临港新片区硬科技产业空间研究报告》。",
]:
    add_bullet(t)

add_h3("9.5 楼宇与园区诊断报告")
add_p("面向业主/运营方/国企平台：竞品分析、租金定位、空置诊断、目标企业画像、招商策略、"
      "改造与运营提升建议。")

add_h3("9.6 企业选址服务产品")
add_p("面向企业：区域推荐、楼宇推荐、租金对标、政策匹配、成本测算、产业生态分析。")


# ============================== 附篇一 ==============================
add_section_divider("附篇一", "DATA  ARCHITECTURE",
                    "数据体系与字段设计（六大数据库）")

add_h1("附篇一  数据体系与字段设计（六大数据库）")

databases = [
    ("1. 楼宇与园区基础数据库",
     ["载体名称", "物业类型", "所属行政区", "所属街道/镇", "详细地址", "经纬度",
      "总建筑面积", "可租赁面积", "标准层面积", "层数", "竣工时间", "物业等级",
      "产权方", "运营方", "物业公司", "交通条件", "地铁距离", "停车位数量",
      "配套商业", "招商联系人", "官方网站或招商页面"],
     "高德/百度/腾讯地图POI、AOI；各区政府公开园区名录；上海市产业园区名录；楼宇/"
     "园区官网与招商页；房地产中介挂牌平台；公开新闻稿；实地调研补充。"),
    ("2. 租金与空置数据库",
     ["报价租金", "成交租金区间", "物业费", "可租面积", "空置面积", "空置楼层",
      "空置套数", "最小可租面积", "最大连续可租面积", "装修状态", "免租期",
      "付款方式", "租期要求", "更新时间", "信息来源"],
     "仲量联行、世邦魏理仕、戴德梁行、高力国际、第一太平戴维斯公开报告；"
     "易居研究院历史数据库；好租、点点租、58同城、安居客、房天下、楼盘网商办频道、"
     "办公伙伴类平台；楼宇/园区招商公众号；中介报价单；电话询价；实地踩盘；"
     "客户真实成交反馈。建议租金区分：公开挂牌、中介报价、实际成交。"),
    ("3. 企业入驻数据库",
     ["企业名称", "统一社会信用代码", "注册资本", "成立时间", "注册地址",
      "实际办公地址", "所在楼宇/园区", "所属行业", "细分赛道", "员工规模",
      "融资轮次", "是否专精特新", "是否高新技术企业", "是否上市/拟上市",
      "经营状态", "企业简介", "关联母公司/子公司"],
     "国家企业信用信息公示系统；上海市市场监督管理局；天眼查、企查查、爱企查、启信宝；"
     "Wind/iFinD/Choice 企业库；各区重点企业名录、高新企业、专精特新、小巨人企业"
     "名单；上市公司公告；融资新闻；园区入驻企业公示；楼宇水牌识别；企业官网与公众号；"
     "招聘平台地址信息；高德/百度企业POI；实地调研。注意区分注册地址与实际办公地址。"),
    ("4. 企业迁徙数据库",
     ["企业名称", "原办公地址", "原楼宇/园区", "原街道/镇", "新办公地址",
      "新楼宇/园区", "新街道/镇", "迁徙时间", "迁徙类型", "面积变化", "租金变化",
      "行业属性", "企业规模", "迁徙原因推测"],
     "工商地址变更记录；企业官网搬迁公告；公众号迁址通知；招聘平台办公地址变化；"
     "地图POI地址变化；园区入驻新闻；楼宇签约新闻；中介成交记录；招商部门线索；"
     "物业访谈；企业访谈。判断标准：至少两个以上来源交叉验证。"),
    ("5. 产业标签数据库",
     ["企业名称", "一级行业", "二级行业", "三级赛道", "战略性新兴产业属性",
      "上海重点产业属性", "新质生产力方向", "产业链核心环节", "上下游关系",
      "主要产品", "主要客户", "技术方向", "融资情况", "专利数量", "资质荣誉"],
     "企业官网与招聘；融资新闻；国家知识产权局/专利数据库；启信宝/企查查产业标签；"
     "政府重点产业企业名单；各区/园区产业图谱；上市公司公告；产业研究报告；"
     "微信公众号与新闻报道。"),
    ("6. 政策与配套数据库",
     ["区域政策", "街镇招商政策", "租金/装修/人才补贴", "税收奖励", "科技项目支持",
      "产业基金", "落户政策", "人才公寓", "公共服务平台", "周边交通", "商业配套",
      "教育医疗资源", "会议展示空间"],
     "上海市/各区政府官网；各区投促办；各区科委、经委、商务委；街镇招商主管部门；"
     "园区招商手册；产业政策汇编；官方公众号；新闻发布会资料；政策申报平台。"),
]

for name, fields, sources in databases:
    add_h2(name)
    add_h3("核心字段")
    add_p("、".join(fields) + "。")
    add_h3("数据来源建议")
    add_p(sources)


# ============================== 附篇二 ==============================
add_section_divider("附篇二", "INDEX  SYSTEM",
                    "核心指标体系（五大指数）")

add_h1("附篇二  核心指标体系（五大指数）")
indices = [
    ("1. 区域市场景气指数",
     ["租金变化", "空置率变化", "新增成交面积", "新增企业数量", "企业净流入数量",
      "招商去化速度"],
     "判断区域市场冷热；识别短期市场波动；辅助政府和资产方进行预警。"),
    ("2. 楼宇/园区竞争力指数",
     ["区位交通", "租金性价比", "空间品质", "产业集聚", "企业质量", "服务能力",
      "政策资源", "去化表现"],
     "对重点楼宇和园区横向比较，形成资产诊断模型，支撑招商和资管建议。"),
    ("3. 企业需求热度指数",
     ["新注册企业数量", "融资企业数量", "招聘岗位增长", "扩租企业数量",
      "新租企业数量", "重点产业企业活跃度"],
     "判断企业需求强弱，识别产业增长方向，输出招商线索。"),
    ("4. 产业集聚度指数",
     ["同行业企业数量", "龙头企业数量", "上下游企业完整度", "专精特新企业数量",
      "高新技术企业数量", "产业链协同程度"],
     "判断区域主导产业，识别产业集群成熟度，支撑产业招商与补链强链。"),
    ("5. 供需匹配度指数",
     ["载体面积与企业需求匹配度", "租金与企业承受能力匹配度",
      "区域产业定位与企业属性匹配度", "政策与企业发展阶段匹配度",
      "空间产品与企业使用场景匹配度"],
     "判断哪些楼宇适配哪些企业，哪些园区适配哪些产业，支撑精准招商与企业选址。"),
]
for name, sub, use in indices:
    add_h2(name)
    add_h3("指标构成")
    add_p("、".join(sub) + "。")
    add_h3("用途")
    add_p(use)


# ============================== 附篇三 ==============================
add_section_divider("附篇三", "DATA  ACCESS  &  COMPLIANCE",
                    "数据调用、采集与合规建议")

add_h1("附篇三  数据调用、采集与合规建议")
add_h2("一、五项数据调用原则")
for t in [
    "公开合规原则：优先使用公开数据、授权数据、合作数据和调研数据；",
    "多源交叉原则：关键字段至少两个以上来源验证；",
    "动态更新原则：建立月度更新机制；",
    "分级可信原则：A级（官方/企业公告/实地/合作方）、B级（机构报告/招商文件/物业访谈）、"
    "C级（公开挂牌/中介口径/平台抓取）、D级（网络零散信息，仅作线索参考）；",
    "人工校验原则：AI和工具用于提效，关键样本由研究员和线下调研校验。",
]:
    add_bullet(t)

add_h2("二、四阶段数据采集分期")
add_table(
    ["阶段", "覆盖范围", "阶段目标", "建议成果"],
    [
        ["第一阶段 中心城区试点",
         "黄浦、静安、徐汇、长宁、普陀、虹口、杨浦",
         "建立字段与匹配规则；建立租金、空置采集口径；建立企业迁徙识别模型",
         "《上海中心城区商办楼宇与产业园区市场试点报告》"],
        ["第二阶段 浦东新区补全",
         "陆家嘴、前滩、张江、金桥、外高桥、世纪公园、临港、唐镇、康桥、周浦",
         "补齐高端商务、科创研发与硬科技产业空间样本",
         "浦东商务与产业空间补全数据库"],
        ["第三阶段 重点产业区补全",
         "闵行、嘉定、松江、青浦、宝山、奉贤；漕河泾、紫竹、嘉定汽车城、G60、南大、"
         "东方美谷、西虹桥",
         "覆盖上海主要产业空间；新兴产业迁徙与产业链空间布局",
         "重点产业区专题数据库"],
        ["第四阶段 全域覆盖",
         "16区全覆盖",
         "形成全域动态数据库；支持月度/季度/年度产品",
         "上海商办楼宇与产业园区全域动态数据库"],
    ],
    col_widths=[3.5, 4.0, 4.5, 3.4],
)

add_h2("三、合规性要点")
for t in [
    "优先使用公开数据；不采集个人隐私；不采集非公开商业秘密；",
    "企业联系方式谨慎处理；",
    "报告以统计性、汇总性结果为主；",
    "对单一企业敏感信息避免未经确认披露；",
    "对外发布前进行合规审查。",
]:
    add_bullet(t)


# ============================== 附篇四 ==============================
add_section_divider("附篇四", "JOINT  WORKFLOW",
                    "分工协同机制与工作流程")

add_h1("附篇四  分工协同机制与工作流程")

add_h2("一、易居研究院  建议职责")
for t in [
    "商办市场研究框架设计、统计口径制定；",
    "楼宇/园区/租金/空置/去化等市场指标研究；",
    "区域和板块市场分析；",
    "资产运营与招商策略研究；",
    "市场访谈和线下调研组织；",
    "报告撰写、可视化、季度/年度趋势研判；",
    "面向资产方、园区方、政府部门的应用转化；",
    "对外发布与行业传播。",
]:
    add_bullet(t)

add_h2("二、复旦大学住房政策研究中心  建议职责")
for t in [
    "城市更新和空间治理理论框架支持；",
    "产业空间与城市功能关系研究；",
    "政策评价体系设计、政府决策建议撰写；",
    "数据合规与公共政策视角把关；",
    "城市住房、就业、通勤、产业空间协同研究；",
    "课题学术背书、专家研讨会组织；",
    "研究成果向政府端转化；",
    "年度白皮书及政策建议章节联合撰写。",
]:
    add_bullet(t)

add_h2("三、数据技术与 AI 采集方  建议职责")
for t in [
    "楼宇/园区 POI/AOI 数据采集；",
    "企业工商数据采集与清洗；",
    "企业地址识别和匹配；",
    "租金挂牌数据抓取；",
    "招聘、新闻、公众号、公告数据监测；",
    "企业迁徙线索识别；",
    "数据标准化处理与数据库搭建；",
    "基础指标计算、数据更新与接口维护。",
]:
    add_bullet(t)

add_h2("四、联合工作机制（四类例会）")
add_table(
    ["机制", "频率", "重点议题"],
    [
        ["周度数据进度会", "周", "数据采集进度、字段问题、样本校验、异常数据处理"],
        ["双周研究模型评审会", "双周", "指标体系、区域口径、租金/空置/迁徙模型"],
        ["月度报告选题会", "月", "重点区域、重点产业、重点企业迁徙案例、市场热点"],
        ["季度成果复盘会", "季度", "数据准确性、客户反馈、报告结构优化、商业化产品规划"],
    ],
    col_widths=[3.8, 2.0, 9.6],
)


# ============================== 附篇五 ==============================
add_section_divider("附篇五", "PILOT  ROLLOUT",
                    "首期试点报告落地建议")

add_h1("附篇五  首期试点报告落地建议")
add_h2("一、首期名称")
add_p("《上海中心城区商办楼宇与产业园区市场试点报告》；或：《上海中心城区商办与产业空间"
      "供需格局及企业迁徙研究》。")

add_h2("二、首期覆盖范围（七大中心城区）")
add_p("黄浦、静安、徐汇、长宁、普陀、虹口、杨浦。")

add_h2("三、首期目录建议")
chapters = [
    ("第一章  中心城区商办与产业载体市场概览",
     ["七区载体总量；七区存量面积；七区租金水平；七区空置水平；七区主导产业。"]),
    ("第二章  重点区域供给格局",
     ["黄浦：传统商务、金融服务和专业服务；",
      "静安：高端商务、总部经济和消费服务；",
      "徐汇：AI、数字经济、传媒和滨江总部；",
      "长宁：虹桥商务、涉外商务和数字贸易；",
      "普陀：数字经济、城市更新和新兴商务载体；",
      "虹口：北外滩金融、航运和总部办公；",
      "杨浦：在线新经济、AI、科创服务和高校资源转化。"]),
    ("第三章  租金、空置与市场分化",
     ["各区租金对比；各区空置率对比；",
      "核心板块与非核心板块分化；",
      "写字楼与产业园区分化；老旧载体压力分析。"]),
    ("第四章  企业需求和入驻画像",
     ["行业分布；企业规模；空间需求；租金承受能力；重点产业企业分布。"]),
    ("第五章  企业迁徙和区域流动",
     ["企业净流入区域；企业流出区域；降本型、升级型、产业集聚型迁移。"]),
    ("第六章  重点板块招商机会",
     ["徐汇滨江、北外滩、杨浦滨江、五角场、苏河湾、长宁虹桥、普陀真如。"]),
    ("第七章  市场趋势与策略建议",
     ["租金趋势；空置趋势；产业需求趋势；",
      "政府招商建议；楼宇运营建议；资产调价建议。"]),
]
for title, items in chapters:
    add_h3(title)
    for it in items:
        add_bullet(it)


# ============================== 附篇六 ==============================
add_section_divider("附篇六", "KEY  CHALLENGES",
                    "关键难点与解决方案")

add_h1("附篇六  关键难点与解决方案")
nodes = [
    ("难点一：实际办公地址难以准确识别",
     "工商注册地址 + 地图POI + 招聘地址 + 企业官网 + 公众号公告 + 楼宇水牌 + 物业访谈 + "
     "实地确认，多源交叉校验，避免将注册地址误认为实际办公地址。"),
    ("难点二：成交租金数据难获取",
     "挂牌租金作为基础参考；中介报价作为市场校准；真实成交作为核心依据；通过物业访谈、经"
     "纪人访谈、企业访谈积累成交数据库；报告中明确区分报价租金与成交租金。"),
    ("难点三：空置率口径不统一",
     "区分物理空置、招商空置、隐性空置三类；以“招商可租面积”口径为主，物理空置和隐性空置"
     "作辅助判断。"),
    ("难点四：企业迁徙存在误判",
     "迁徙判断至少需要两个以上信号（工商地址变化、招聘地址变化、官网/公众号公告、地图POI、"
     "园区入驻新闻、线下确认）。"),
    ("难点五：数据合规",
     "优先使用公开数据；不采集个人隐私；不采集非公开商业秘密；企业联系方式谨慎处理；以汇"
     "总统计与趋势分析为主；对单一企业敏感信息避免未经确认披露；对外发布前进行合规审查。"),
]
for title, txt in nodes:
    add_h2(title)
    add_p(txt)


# ============================== 附篇七 ==============================
add_section_divider("附篇七", "DELIVERABLES",
                    "最终成果包与商业化产品矩阵")

add_h1("附篇七  最终成果包与商业化产品矩阵")
add_table(
    ["类别", "成果", "用途与对象"],
    [
        ["主报告", "《上海商办楼宇与产业园区市场深度报告》",
         "联合发布、行业传播、政策建议"],
        ["数据附录",
         "楼宇/园区清单、区域指标表、产业企业分布表、租金空置表、迁徙样本表",
         "数据透明披露、合作方对接"],
        ["可视化图册",
         "楼宇/园区分布图、租金与空置热力图、企业迁徙流向图、产业集聚图、"
         "重点板块对比图、招商机会地图",
         "汇报、传播、政府决策"],
        ["PPT汇报版", "简版30页 / 标准版50页 / 深度版80+页",
         "政府、园区、楼宇、资产方汇报"],
        ["数据看板",
         "区域市场看板、楼宇园区看板、企业迁徙看板、产业热度看板、招商线索看板",
         "持续监测与商业化输出"],
        ["延伸产品",
         "月度简报、季度报告、年度白皮书、区域专项、楼宇/园区诊断、企业选址服务",
         "长期产品体系与商业化"],
    ],
    col_widths=[3.0, 7.0, 5.4],
)


# ============================== 结语 ==============================
add_section_divider("结语", "VISION",
                    "研究价值与长期愿景")
add_h1("结语  研究价值与长期愿景")
for t in [
    "本报告拟由易居研究院与复旦大学住房政策研究中心联合研发，围绕上海商办楼宇与产业园区市"
    "场的全域空间供给、企业需求变化、租赁价格体系、企业迁徙路径、产业集聚规律和资产运营"
    "策略展开系统研究。",
    "与传统聚焦核心商务区甲级写字楼的市场报告不同，本报告将研究视角下沉至行政区、街道、"
    "楼宇、园区和企业层面，重点构建“空间载体—入驻企业—产业标签—租赁行为—迁徙趋势—招"
    "商机会”的一体化研究框架，力求真实反映上海商办和产业空间市场的供需格局与运行逻辑。",
    "报告成果不仅可为政府部门制定产业空间政策、推进城市更新、开展精准招商提供决策依据，"
    "也可为楼宇园区运营方、资产持有方、投资机构和企业选址方提供可执行的数据支撑和市场判"
    "断。通过持续研发和动态更新，该报告有望成为上海商办楼宇与产业园区市场的长期观察工具、"
    "行业标准化研究样本和城市空间治理决策参考。",
]:
    add_p(t)

add_callout("打造为：上海商办楼宇与产业园区市场年度旗舰研究产品  +  "
            "行业标准化研究样本  +  城市空间治理决策参考",
            fill="102A5B", color=(255, 255, 255))

doc.add_paragraph()
add_centered("—  讨论稿 · 仅供联合课题组内部研讨使用  —",
             size=10, color=GRAY, font="宋体")


out = os.path.join(OUTPUT_DIR, "上海商办楼宇与产业园区市场深度报告.docx")
doc.save(out)
print(f"Word 文档已生成：{out}")
