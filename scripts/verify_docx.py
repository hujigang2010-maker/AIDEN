#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""校验《房地产项目开发运营一线实务问答》结构与关键口径是否齐全。"""

from pathlib import Path
import sys

from docx import Document

DOC = Path(__file__).resolve().parents[1] / "deliverables" / "房地产项目开发运营一线实务问答.docx"

REQUIRED_HEADINGS = [
    "一、项目公司制度",
    "二、现金流使用",
    "三、地产开发缴税全流程",
    "四、表外项目",
    "五、拿地测算",
    "六、去化与降价",
    "七、绿城为何利润释放不出来",
    "八、九家房企",
    "附录A",
    "附录B",
    "附录C",
]

REQUIRED_PHRASES = [
    "一项目一公司",
    "人格混同",
    "税盾",
    "开发贷",
    "四证",
    "预售资金监管",
    "股东借款",
    "3%",
    "预计计税毛利率",
    "国税发〔2009〕31 号",
    "普通住宅",
    "合营",
    "联营",
    "无息",
    "净利率",
    "华润置地",
    "中海",
    "建发",
    "金茂",
    "绿城",
    "滨江",
    "越秀",
    "保利发展",
    "招商蛇口",
    "设计变更",
]


def main() -> int:
    if not DOC.exists():
        print(f"缺失文件：{DOC}", file=sys.stderr)
        return 1
    d = Document(DOC)
    paras = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    blob = "\n".join(paras)
    for row in d.tables:
        for cell in row._cells:
            blob += "\n" + cell.text
    missing_h = [h for h in REQUIRED_HEADINGS if not any(h in t for t in paras)]
    missing_p = [p for p in REQUIRED_PHRASES if p not in blob]
    n_tables = len(d.tables)
    n_chars = len(blob)
    print(f"文件：{DOC}")
    print(f"段落：{len(paras)}  表格：{n_tables}  文本约：{n_chars} 字")
    if missing_h:
        print("缺标题：", missing_h)
        return 1
    if missing_p:
        print("缺口径：", missing_p)
        return 1
    if n_tables < 20:
        print(f"表格过少：{n_tables}")
        return 1
    if n_chars < 12000:
        print(f"正文过短：{n_chars}")
        return 1
    print("校验通过")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
