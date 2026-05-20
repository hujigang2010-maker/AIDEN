"""Generate the Greentown China dinner-naming sponsorship agreement (Annex 4)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


GREEN = RGBColor(0x00, 0x6B, 0x3F)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)


def _force_font(run, font="宋体"):
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rpr.append(rfonts)


def set_default_font(doc, name="宋体"):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def heading(doc, text, level=1, color=None, center=False, font="黑体"):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(20)
    elif level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    if color is not None:
        run.font.color.rgb = color
    _force_font(run, font=font)
    return p


def para(doc, text, bold=False, size=11, color=DARK, align=None,
         first_indent=True, font="宋体", line_spacing=1.5):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = line_spacing
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    _force_font(run, font=font)
    return p


def clause(doc, num, text):
    """A numbered clause with hanging indent."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(0)
    run_num = p.add_run(num + "  ")
    run_num.bold = True
    run_num.font.size = Pt(11)
    _force_font(run_num)
    run = p.add_run(text)
    run.font.size = Pt(11)
    _force_font(run)
    return p


def blank(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=DARK, size=10.5,
                  align_center=False, font="宋体"):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    _force_font(run, font=font)


def add_kv_table(doc, rows, col_widths=(4.0, 12.5)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(col_widths[0])
        row.cells[1].width = Cm(col_widths[1])
    for i, (k, v) in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], k, bold=True, size=10.5)
        shade_cell(table.rows[i].cells[0], "E8F1EC")
        set_cell_text(table.rows[i].cells[1], v, size=10.5)
    return table


def add_table(doc, headers, rows, col_widths=None, header_fill="006B3F",
              header_color=RGBColor(0xFF, 0xFF, 0xFF)):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True,
                      color=header_color, size=10.5, align_center=True)
        shade_cell(table.rows[0].cells[i], header_fill)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            set_cell_text(table.rows[r + 1].cells[c], str(val), size=10)
    return table


def build():
    doc = Document()
    set_default_font(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

    # ====== 抬头 ======
    heading(doc, "附件四", level=0, color=GREEN, center=True)
    heading(
        doc,
        "晚宴冠名战略合作伙伴专项合作协议",
        level=0, center=True,
    )
    heading(
        doc,
        "（绿城中国 | 绿城·潮鸣外滩 · 10 万元赞助）",
        level=2, color=GREY, center=True,
    )
    para(
        doc,
        "Annex 4 — Dinner Title Strategic Partnership Sponsorship Agreement",
        align="center", size=10, color=GREY, first_indent=False,
    )
    blank(doc)

    para(doc, "协议编号：__________________________", first_indent=False, size=10.5)
    para(doc, "签订地点：上海市______区______路______号", first_indent=False, size=10.5)
    para(doc, "签订日期：______年______月______日", first_indent=False, size=10.5)
    blank(doc)

    # ====== 前言 ======
    para(
        doc,
        "鉴于：",
        bold=True, first_indent=False, size=11,
    )
    para(
        doc,
        "1. 甲方拟于 2026 年 5 月 22 日在上海·北外滩·一滴水主办「重构与突围 — 2026 人工智能商业化落地与硬核投资破局峰会」（以下简称「本次峰会」）；",
    )
    para(
        doc,
        "2. 乙方系绿城中国控股有限公司及其下属/关联项目「绿城·潮鸣外滩」的合作品牌方，愿意作为本次峰会「晚宴冠名战略合作伙伴」提供赞助；",
    )
    para(
        doc,
        "3. 双方已就赞助合作的核心商业要素达成一致，并已签署《2026 人工智能商业化落地与硬核投资破局峰会赞助合作协议》（以下简称「主协议」）；",
    )
    para(
        doc,
        "4. 本协议作为主协议之《附件四·晚宴冠名战略合作伙伴专项合作协议》，与主协议正文及其他附件具有同等法律效力，对本次「晚宴冠名战略合作伙伴」专项合作的具体权利义务、物料植入、执行节点、费用结算、签字盖章等事项作出专项约定。",
    )
    blank(doc)

    para(
        doc,
        "双方本着平等、自愿、诚实信用的原则，依据《中华人民共和国民法典》《中华人民共和国广告法》《中华人民共和国反不正当竞争法》《中华人民共和国商标法》《中华人民共和国著作权法》《中华人民共和国个人信息保护法》《中华人民共和国数据安全法》《中华人民共和国网络安全法》等相关法律法规，就乙方赞助本次峰会「晚宴冠名战略合作伙伴」事宜达成如下协议，以资共同遵守。",
    )
    blank(doc)

    # ====== 双方主体 ======
    heading(doc, "一、双方主体信息", level=1)

    para(doc, "甲方（主办方）：", bold=True, first_indent=False)
    add_kv_table(
        doc,
        rows=[
            ["主办单位", "北京大学经济学院上海校友会 / 复旦大学住房政策研究中心（联合主办）"],
            ["统一社会信用代码", "________________________________"],
            ["通讯地址", "上海市______区______路______号"],
            ["授权代表 / 职务", "________________ / ____________"],
            ["项目联系人", "________________  联系电话：________________"],
            ["电子邮箱", "________________"],
        ],
    )
    blank(doc)

    para(doc, "乙方（赞助方 / 晚宴冠名战略合作伙伴）：", bold=True, first_indent=False)
    add_kv_table(
        doc,
        rows=[
            ["公司名称", "绿城中国控股有限公司 / ____________________________（项目主体）"],
            ["项目品牌", "绿城·潮鸣外滩"],
            ["统一社会信用代码", "________________________________"],
            ["注册地址", "________________________________"],
            ["授权代表 / 职务", "________________ / ____________"],
            ["项目联系人", "Aiden（绿城方总对接） / ________________"],
            ["联系电话", "________________"],
            ["电子邮箱", "________________"],
        ],
    )
    blank(doc)

    # ====== 二、合作内容与赞助级别 ======
    heading(doc, "二、合作内容与赞助级别", level=1)
    clause(
        doc, "2.1",
        "经双方协商一致，乙方以「晚宴冠名战略合作伙伴（Dinner Title Strategic Partner）」身份参与本次峰会的合作。该级别由甲方在主协议「钻石赞助」基础上，根据乙方需求叠加「晚宴独家冠名」「项目专场宣讲」「现场展位 + 接驳到访」「媒体宣发 + 长效圈层」四大模块，进行专项定制。",
    )
    clause(
        doc, "2.2",
        "乙方本次专项赞助金额：人民币（大写）壹拾万元整（小写：￥100,000.00 元）。",
    )
    clause(
        doc, "2.3",
        "乙方在本次峰会中享有的具体权益、规格、数量及交付节点，详见本附件第三条及附表《晚宴冠名战略合作伙伴专项权益执行清单》。",
    )
    clause(
        doc, "2.4",
        "甲方确认在本次峰会期间不再向同一行业（房地产开发与销售、品质改善型住宅项目）的其他企业授予「晚宴冠名」或同级以上对外宣传身份；如需引入同行业其他赞助商，应事先以书面形式告知乙方并取得乙方书面同意。",
    )
    blank(doc)

    # ====== 三、具体权益与物料植入 ======
    heading(doc, "三、具体权益与物料植入", level=1)
    para(
        doc,
        "甲乙双方就「晚宴冠名战略合作伙伴」具体权益及物料植入约定如下，由双方在《附件四》之「附表一」项下逐项确认并签字：",
    )

    heading(doc, "（一）晚宴冠名权益", level=2)
    add_table(
        doc,
        headers=["序号", "权益内容", "数量 / 规格", "责任方"],
        rows=[
            ["1", "晚宴冠名权：「晚宴冠名战略合作伙伴：绿城中国 | 绿城·潮鸣外滩」全程统称权（含主持人口播、KV 字样、媒体口径）", "全场", "甲方"],
            ["2", "晚宴 KV 主视觉及 LED 大屏品牌锁屏（含晚宴入场、餐间、散场各场景）", "1 套（建议 16:9 1920×1080，最终以场地像素比为准）", "甲方设计 / 乙方提供素材"],
            ["3", "席卡（晚宴座位卡）植入项目 logo + 品牌主张", "全场所有席卡（90×55mm 双面）", "甲方制作"],
            ["4", "桌卡（桌号牌）植入项目 logo + 「绿城·潮鸣外滩 · 晚宴冠名」字样", "全场所有桌卡（A5 折立 148×210mm 或亚克力 200×150mm）", "活动公司印刷"],
            ["5", "晚宴菜单植入项目 logo（封面整版 + 内页页脚）", "全场所有菜单（对开 A4 210×285mm 或单页 285×210mm）", "活动公司印刷"],
            ["6", "晚宴期间项目宣传片轮播", "≤ 60 秒，全晚循环播放", "甲方上屏 / 乙方提供素材"],
            ["7", "主持人口播：开场鸣谢 + 川总宣讲引荐 + 散场鸣谢", "共 3 段，每段约 30 秒", "甲方主持 / 乙方共写口径"],
            ["8", "晚宴前项目专场宣讲（绿城方川总主讲）", "时长 15 分钟 + Q&A 5 分钟；时间 18:15–18:30（最终以现场议程为准）", "乙方主讲 / 甲方主持引荐"],
            ["9", "晚宴主桌 / 入场券", "合计 8 张（其中主桌 3 人 + 销售/接待 5 人，最终名单 5 月 19 日前互相确认）", "甲方"],
        ],
        col_widths=[1.2, 8.5, 5.0, 3.0],
    )
    blank(doc)

    heading(doc, "（二）主会场植入权益", level=2)
    add_table(
        doc,
        headers=["序号", "权益内容", "数量 / 规格", "责任方"],
        rows=[
            ["10", "主背景板「钻石级」logo 位 + 「晚宴冠名」副标题位", "1 处", "甲方"],
            ["11", "议程手册广告（整版）+ 项目折页夹页", "整版 A4 210×285mm，出血 3mm，CMYK 300dpi", "甲方设计 / 乙方提供素材"],
            ["12", "主会场宣传片轮播（嘉宾入场及中场休息时段）", "16:9 1920×1080，≤ 60 秒", "甲方上屏 / 乙方提供素材"],
            ["13", "500 份手拎袋植入项目折页 + 285/310 户型图", "三折页 297×210mm + 户型单页 × 500 套", "印刷单位 / 活动公司装袋"],
            ["14", "白皮书扉页联合署名「晚宴冠名战略合作伙伴 绿城中国 · 绿城·潮鸣外滩」", "1 处", "甲方"],
        ],
        col_widths=[1.2, 8.5, 5.0, 3.0],
    )
    blank(doc)

    heading(doc, "（三）现场展位与项目到访权益", level=2)
    add_table(
        doc,
        headers=["序号", "权益内容", "数量 / 规格", "责任方"],
        rows=[
            ["15", "现场专属品牌展位（电梯口）", "1 处，含品牌桌台轻包装、易拉宝×2、户型 KT 板、销售 2 人", "乙方派驻 / 活动公司展位包装"],
            ["16", "现场专属品牌展位（展场内）", "1 处，含洽谈圆桌、易拉宝×2、户型 KT 板、销售 2 人", "乙方派驻 / 活动公司展位包装"],
            ["17", "论坛 / 颁奖结束后主持人口播 + 现场引导员，引导意向嘉宾前往项目案场", "1 次主口播 + 现场引导员 3 名", "甲方"],
            ["18", "项目案场接驳：华为尊界 8–10 辆 + 考斯特补位", "由乙方安排并承担费用，往返时长不低于 60 分钟", "乙方"],
            ["19", "项目案场接待：售楼处沙盘 + 样板间参观（含夜场灯光秀，如适用）", "由乙方安排并承担费用", "乙方"],
        ],
        col_widths=[1.2, 8.5, 5.0, 3.0],
    )
    blank(doc)

    heading(doc, "（四）媒体宣发与长效权益", level=2)
    add_table(
        doc,
        headers=["序号", "权益内容", "数量 / 规格", "责任方"],
        rows=[
            ["20", "媒体通稿 — 标题级或副标题级出现「绿城·潮鸣外滩」+ 晚宴冠名身份", "主流财经 / 地产媒体不少于 5 家", "甲方媒体组"],
            ["21", "朋友圈九宫格 — 至少 3 张含项目 logo / KV", "9 张 1080×1080 PNG", "甲方媒体组"],
            ["22", "回顾视频片头 / 片尾鸣谢 + 川总宣讲与项目参观画面剪入", "1 条 1920×1080 mp4", "甲方"],
            ["23", "双校长三角校友产业联盟战略合作伙伴永久入册 + 牌匾交接", "1 次（永久入册）", "甲方"],
            ["24", "执行回执：现场图片、媒体链接、嘉宾合影、媒体监测", "1 份（大会结束后 7 个自然日内出具）", "甲方"],
        ],
        col_widths=[1.2, 8.5, 5.0, 3.0],
    )
    blank(doc)

    para(
        doc,
        "上述权益项的最终执行清单以双方签字盖章的「附表一·晚宴冠名战略合作伙伴专项权益执行清单」为准。如执行过程中需调整，应经双方书面确认（含电子邮件或加盖业务章的书面文件）。",
    )
    blank(doc)

    # ====== 四、甲方义务 ======
    heading(doc, "四、甲方义务", level=1)
    clause(doc, "4.1", "甲方应按本协议第三条载明的权益项、规格、数量及交付节点，保质保量交付乙方所享有的全部权益。")
    clause(doc, "4.2", "甲方应在本协议签订后 3 个工作日内向乙方提供：（1）物料规格清单（含 logo / KV / 宣传片 / PPT / 折页等技术参数与提交规范）；（2）晚宴 LED 屏 / 投影准确像素比及刷新率参数；（3）专人对接物料交付与上线（含微信群、邮件双通道）。")
    clause(doc, "4.3", "甲方应在峰会结束后 7 个自然日内向乙方出具《赞助权益执行回执》，并附现场图片、媒体链接、嘉宾合影、媒体监测报告等证明材料。")
    clause(doc, "4.4", "甲方应妥善保管乙方提供的物料、商标及其他商业资料，未经乙方书面同意，不得用于本次峰会之外的任何用途。")
    clause(doc, "4.5", "甲方对其所发布的媒体通稿、朋友圈九宫格、回顾视频等传播物料的合法合规性负责，发布前应将涉及乙方品牌名称、项目名称及核心宣传口径的内容提交乙方品牌部书面确认。")
    clause(doc, "4.6", "甲方应根据《中华人民共和国广告法》《中华人民共和国反不正当竞争法》之规定，在任何宣传物料中真实陈述项目名称及合作身份，不得作出禁止性、误导性或与房地产开发销售管理规定相悖的表述。")
    blank(doc)

    # ====== 五、乙方义务 ======
    heading(doc, "五、乙方义务", level=1)
    clause(doc, "5.1", "乙方应按本协议第七条约定的时间与方式足额支付赞助款项。")
    clause(doc, "5.2", "乙方应在本协议签订后 3 个工作日内向甲方交付权益落地所需物料，包括但不限于：")
    para(doc, "（1）矢量项目 logo（AI / EPS / SVG，无白底）+ PNG 透明底 4 套尺寸；")
    para(doc, "（2）项目品牌简介（≤ 200 字）；")
    para(doc, "（3）项目宣传片（≤ 60 秒，1920×1080 mp4，≤ 200MB）；")
    para(doc, "（4）川总专场宣讲 PPT（16:9 1920×1080，≤ 25 张，控制 15 分钟内）；")
    para(doc, "（5）285/310 户型折页源文件（CMYK，300 dpi）；")
    para(doc, "（6）易拉宝设计稿 / 现场销售物料清单。")
    clause(doc, "5.3", "乙方保证其提供的所有 logo、商标、文字、图片、视频、PPT 及其他物料不侵犯任何第三方知识产权、商业秘密、肖像权、人格权及其他合法权益；如因乙方提供物料导致甲方或第三方主张权利，由乙方独立承担全部法律责任并赔偿甲方因此遭受的全部直接损失（含但不限于诉讼费、律师费、合理调查费等）。")
    clause(doc, "5.4", "乙方应遵守峰会现场及晚宴的相关秩序与流程安排，配合甲方完成议程互动；乙方在主会场及晚宴现场的宣讲内容应符合相关法律法规，不得作出与房地产广告管理规定相悖的承诺。")
    clause(doc, "5.5", "乙方应自行承担华为尊界 8–10 辆 + 考斯特补位接驳、项目案场接待、销售人员差旅及现场展位包装升级等乙方直接承办环节的相关费用，该等费用不计入本协议第 2.2 条之赞助金额。")
    clause(doc, "5.6", "乙方应在乙方自有渠道（含官方微信、视频号、销售案场）二次宣发本次合作内容时，使用甲方提供的统一宣传口径，并确保使用的甲方名称、商标、嘉宾形象等元素已取得甲方书面同意。")
    blank(doc)

    # ====== 六、物料交付与时间节点 ======
    heading(doc, "六、物料交付与执行时间节点", level=1)
    para(
        doc,
        "双方共同确认以下关键时间节点。如任一方因不可抗力以外的原因延迟，应提前 24 小时书面通知对方并协商解决：",
    )
    add_table(
        doc,
        headers=["日期", "节点", "乙方动作", "甲方 / 活动公司动作"],
        rows=[
            ["5/18", "确认合作框架及金额", "盖章合同回传", "出具盖章合同 + 银行账户"],
            ["5/19", "物料源文件交付（甲乙互交）", "logo / 项目宣传片 / 川总 PPT / 户型折页 / 品牌简介 / 易拉宝稿", "晚宴 KV / 桌卡 / 菜单 / 席卡设计稿初稿；晚宴 LED 像素比反馈"],
            ["5/20", "设计稿确认", "确认一轮，提出修改意见", "完成印刷物料终稿"],
            ["5/21", "印刷下单 & 物流", "—", "桌卡、菜单、易拉宝、户型折页全部下印"],
            ["5/22 上午", "现场布展", "销售 4 人到场对接展位包装", "活动公司完成展位、KV、LED 调试"],
            ["5/22 全天", "大会执行", "执行 + 陪同 + 案场对接", "执行 + 主持口播 + 宣传片轮播"],
            ["5/22 20:30+", "项目参观（「会后专场」方案 A）", "尊界 8–10 辆 + 考斯特 + 案场接待", "现场引导员 + 名单对接"],
            ["5/23–5/29", "宣发 & 回执", "同步乙方自媒体矩阵", "媒体通稿 / 九宫格发布；7 日内出回执"],
        ],
        col_widths=[2.5, 4.5, 5.0, 5.5],
    )
    blank(doc)

    # ====== 七、付款 ======
    heading(doc, "七、款项支付与发票", level=1)
    clause(doc, "7.1", "乙方应于本协议签订之日起 5 个工作日内，将全部赞助款项一次性汇入甲方指定账户：")
    add_kv_table(
        doc,
        rows=[
            ["收款单位", "________________________________"],
            ["开户银行", "________________________________"],
            ["银行账号", "________________________________"],
            ["税号", "________________________________"],
        ],
    )
    clause(doc, "7.2", "甲方在款项到账后 10 个工作日内向乙方开具等额合法有效的增值税普通发票或增值税专用发票（发票内容：会议服务费 / 赞助费）。")
    clause(doc, "7.3", "本协议项下任何因银行手续费、税费产生的费用，由各自一方承担。")
    blank(doc)

    # ====== 八、知识产权 ======
    heading(doc, "八、知识产权与品牌使用", level=1)
    clause(doc, "8.1", "乙方授权甲方在本次峰会的宣传、报道、白皮书、官网、大屏、议程手册、媒体通稿、回顾视频等场景中使用乙方「绿城中国」及「绿城·潮鸣外滩」商标与企业简介，授权范围严格限于本次峰会及其衍生宣传内容，授权期限为本协议签订之日起至本次峰会结束后 12 个月。")
    clause(doc, "8.2", "甲方拥有本次峰会名称、视觉系统、白皮书、回顾视频等衍生品的完整知识产权。乙方在使用大会名称、Logo、嘉宾形象等元素进行二次宣传前，应事先取得甲方书面同意（含微信书面同意），并不得作出有损大会形象或误导性的表述。")
    clause(doc, "8.3", "未经对方事先书面同意，任何一方不得在本协议约定的合作范围之外，单独或与第三方就本次峰会进行商业开发、销售衍生品或对外授权。")
    clause(doc, "8.4", "本协议项下双方就合作过程中形成的现场摄影 / 摄像、回顾视频、白皮书署名作品等成果物，知识产权归甲方所有，乙方享有非独家、不可转授权的免费使用权，使用范围限于乙方品牌内部宣传与销售案场展示。")
    blank(doc)

    # ====== 九、保密 ======
    heading(doc, "九、保密条款", level=1)
    clause(doc, "9.1", "双方对在合作过程中接触到的对方商业秘密、客户名录、嘉宾联系方式、未公开的合作金额、未公开的物料底稿等信息负有保密义务，保密期限自本协议签订之日起 3 年。")
    clause(doc, "9.2", "双方依照《中华人民共和国个人信息保护法》《中华人民共和国数据安全法》之规定处理本次合作中涉及的个人信息：未经嘉宾事先同意，不得将嘉宾联系方式用于本次峰会以外的任何商业推广。乙方对其通过展位/案场收集到的意向客户个人信息，应单独取得当事人书面或电子告知同意，并不得提供给本协议双方以外的第三方。")
    clause(doc, "9.3", "保密义务不因本协议解除或终止而失效。")
    blank(doc)

    # ====== 十、违约责任 ======
    heading(doc, "十、违约责任", level=1)
    clause(doc, "10.1", "乙方未按约定时间支付赞助款项的，每逾期一日按未付款项的 0.5% 向甲方支付违约金；逾期超过 10 日的，甲方有权解除本协议并不退还任何已收款项（如有），同时乙方应另行赔偿甲方因此遭受的直接经济损失。")
    clause(doc, "10.2", "甲方未按本协议第三条 / 第六条 / 附表一约定提供权益且无正当理由的，应按未履行权益对应价值的 100% 向乙方退款，或在双方协商一致后以同等价值的下一届同类活动权益补偿。")
    clause(doc, "10.3", "任何一方违反第八条（知识产权）、第九条（保密）的，应向守约方支付不低于人民币 50,000 元的违约金；造成实际损失超过违约金的，超过部分应另行赔偿。")
    clause(doc, "10.4", "任何一方解除本协议或终止合作的，应提前 5 个自然日书面通知对方，并就已发生的费用据实结算。")
    blank(doc)

    # ====== 十一、不可抗力 ======
    heading(doc, "十一、不可抗力", level=1)
    clause(doc, "11.1", "因不可抗力（包括但不限于地震、洪水、台风等自然灾害；战争、暴乱、罢工；政府管制、公共卫生事件、电力中断、网络中断等）导致大会延期或取消，使本协议无法继续履行的，双方互不承担违约责任；已支付款项可顺延至甲方下一届同类活动，或按未执行权益比例退还。")
    clause(doc, "11.2", "遭受不可抗力的一方应在事件发生后 5 个自然日内书面通知对方，并提供有效证明文件。")
    blank(doc)

    # ====== 十二、合规与廉洁 ======
    heading(doc, "十二、合规与廉洁条款", level=1)
    clause(doc, "12.1", "双方承诺在履行本协议过程中严格遵守中华人民共和国反贿赂、反洗钱、反不正当竞争等相关法律法规，不得直接或间接向对方工作人员（含其近亲属、关联方）给予任何形式的回扣、佣金、礼品、招待等不正当利益。")
    clause(doc, "12.2", "任一方发现对方有违反本条约定的行为，有权单方解除本协议，并要求对方赔偿因此遭受的全部损失。")
    blank(doc)

    # ====== 十三、争议解决 ======
    heading(doc, "十三、争议解决与法律适用", level=1)
    clause(doc, "13.1", "本协议适用中华人民共和国法律。")
    clause(doc, "13.2", "因本协议引起的或与本协议有关的任何争议，双方应首先友好协商解决；协商不成的，任何一方均有权向甲方所在地有管辖权的人民法院提起诉讼。")
    blank(doc)

    # ====== 十四、其他 ======
    heading(doc, "十四、其他", level=1)
    clause(doc, "14.1", "本附件四作为主协议的组成部分，与主协议具有同等法律效力；本附件与主协议正文条款不一致的，以本附件约定为准；本附件未尽事宜，适用主协议约定及双方另行签订的书面补充协议。")
    clause(doc, "14.2", "本附件自双方加盖公章（或合同章）并由授权代表签字之日起生效，至本次峰会回执出具及款项结算完毕之日终止；但第八条（知识产权）、第九条（保密）、第十条（违约责任）等条款的效力不因本附件终止而失效。")
    clause(doc, "14.3", "本附件一式肆份，甲乙双方各执贰份，具有同等法律效力。")
    clause(doc, "14.4", "本协议任何修改或补充均应采取书面形式，并经双方授权代表签字盖章后生效。")
    clause(doc, "14.5", "双方通过电子邮件、企业微信、钉钉等电子形式签署的扫描件 / 加盖电子签章的文件，与纸质原件具有同等法律效力。")
    blank(doc)

    # ====== 签署页 ======
    doc.add_page_break()
    heading(doc, "签署页（Signature Page）", level=0, color=GREEN, center=True)
    blank(doc)
    para(
        doc,
        "（以下无正文，本页为附件四《晚宴冠名战略合作伙伴专项合作协议》之签署页）",
        align="center", first_indent=False, size=10, color=GREY,
    )
    blank(doc, 2)

    # Signature table
    sig = doc.add_table(rows=6, cols=2)
    sig.style = "Table Grid"
    sig.autofit = False
    for row in sig.rows:
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(8.0)
    set_cell_text(sig.rows[0].cells[0], "甲方（盖章）：", bold=True, size=11)
    set_cell_text(sig.rows[0].cells[1], "乙方（盖章）：", bold=True, size=11)
    set_cell_text(sig.rows[1].cells[0], "北京大学经济学院上海校友会 /\n复旦大学住房政策研究中心", size=10.5)
    set_cell_text(sig.rows[1].cells[1], "绿城中国控股有限公司 /\n____________________________（项目主体）", size=10.5)
    set_cell_text(sig.rows[2].cells[0], "\n\n\n（公章 / 合同章位置）\n\n", size=10.5, align_center=True)
    set_cell_text(sig.rows[2].cells[1], "\n\n\n（公章 / 合同章位置）\n\n", size=10.5, align_center=True)
    set_cell_text(sig.rows[3].cells[0], "授权代表签字：______________________", size=11)
    set_cell_text(sig.rows[3].cells[1], "授权代表签字：______________________", size=11)
    set_cell_text(sig.rows[4].cells[0], "职务：__________________", size=11)
    set_cell_text(sig.rows[4].cells[1], "职务：__________________", size=11)
    set_cell_text(sig.rows[5].cells[0], "签字日期：______年______月______日", size=11)
    set_cell_text(sig.rows[5].cells[1], "签字日期：______年______月______日", size=11)

    blank(doc, 2)
    para(
        doc,
        "见证 / 在场人（如有）：______________________  　　日期：______年______月______日",
        first_indent=False, size=10.5,
    )
    blank(doc)

    # ====== 附表一 ======
    doc.add_page_break()
    heading(doc, "附表一  晚宴冠名战略合作伙伴专项权益执行清单",
            level=0, color=GREEN, center=True)
    para(
        doc,
        "Annex 4 — Schedule 1: Rights Execution Checklist",
        align="center", size=10, color=GREY, first_indent=False,
    )
    blank(doc)
    para(
        doc,
        "本附表逐项列示本协议第三条项下的权益执行明细。请甲乙双方在每项右侧「确认」列签字 / 盖章确认；如该项无需执行，请打「×」。",
    )

    rows = [
        ["1", "晚宴冠名权全程统称（KV / 口播 / 媒体口径）", "全场", "甲方", "□ 已确认"],
        ["2", "晚宴 KV 主视觉及 LED 大屏品牌锁屏", "1 套（16:9 1920×1080 / 条屏 6144×1080，最终以场地像素比为准）", "甲方", "□ 已确认"],
        ["3", "晚宴席卡 logo + 品牌主张植入", "全场（90×55mm 双面）", "甲方", "□ 已确认"],
        ["4", "晚宴桌卡 logo 植入", "全场（A5 折立 148×210mm 或亚克力 200×150mm）", "活动公司", "□ 已确认"],
        ["5", "晚宴菜单 logo 植入", "全场（210×285mm 或 285×210mm）", "活动公司", "□ 已确认"],
        ["6", "晚宴项目宣传片轮播", "≤ 60 秒循环", "甲方上屏 / 乙方供片", "□ 已确认"],
        ["7", "晚宴主持人 3 段口播", "开场 + 引荐 + 散场", "甲方", "□ 已确认"],
        ["8", "川总专场宣讲", "15 分钟 + Q&A 5 分钟", "乙方", "□ 已确认"],
        ["9", "晚宴入场券 / 主桌席位", "合计 8 张（主桌 3 + 销售/接待 5）", "甲方", "□ 已确认"],
        ["10", "主背景板钻石级 logo + 晚宴冠名副标", "1 处", "甲方", "□ 已确认"],
        ["11", "议程手册整版广告 + 折页夹页", "整版 A4，出血 3mm", "甲方设计 / 印刷单位", "□ 已确认"],
        ["12", "主会场宣传片轮播", "16:9 1920×1080 ≤ 60s", "甲方", "□ 已确认"],
        ["13", "500 份手拎袋夹页（折页 + 户型图）", "297×210mm 三折页 × 500", "印刷单位 / 活动公司", "□ 已确认"],
        ["14", "白皮书扉页联合署名", "1 处", "甲方", "□ 已确认"],
        ["15", "现场展位 · 电梯口", "1 处（销售 2 人 + 易拉宝×2 + KT 板）", "乙方 + 活动公司", "□ 已确认"],
        ["16", "现场展位 · 展场内", "1 处（销售 2 人 + 易拉宝×2 + KT 板）", "乙方 + 活动公司", "□ 已确认"],
        ["17", "论坛后主持人口播 + 引导员", "1 次主口播 + 引导员 3 名", "甲方", "□ 已确认"],
        ["18", "项目案场接驳（华为尊界 8–10 辆 + 考斯特补位）", "由乙方承担费用", "乙方", "□ 已确认"],
        ["19", "项目案场接待（沙盘 + 样板间）", "由乙方承担费用", "乙方", "□ 已确认"],
        ["20", "媒体通稿 ≥ 5 家头部财经/地产媒体", "标题级或副标题级", "甲方媒体组", "□ 已确认"],
        ["21", "朋友圈九宫格", "9 张 1080×1080 PNG", "甲方媒体组", "□ 已确认"],
        ["22", "回顾视频片头 / 片尾鸣谢", "1 条 mp4", "甲方", "□ 已确认"],
        ["23", "双校长三角校友产业联盟战略合作伙伴永久入册 + 牌匾", "永久入册", "甲方", "□ 已确认"],
        ["24", "执行回执（图片 + 媒体链接 + 监测）", "7 个自然日内", "甲方", "□ 已确认"],
    ]
    add_table(
        doc,
        headers=["序号", "权益项 / 物料", "规格 / 数量", "责任方", "双方确认"],
        rows=rows,
        col_widths=[1.2, 6.5, 4.5, 3.0, 2.5],
    )
    blank(doc)

    sig2 = doc.add_table(rows=2, cols=2)
    sig2.style = "Table Grid"
    sig2.autofit = False
    for row in sig2.rows:
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(8.0)
    set_cell_text(sig2.rows[0].cells[0], "甲方授权代表（签字 / 盖章）：\n\n\n日期：______年______月______日", bold=True, size=11)
    set_cell_text(sig2.rows[0].cells[1], "乙方授权代表（签字 / 盖章）：\n\n\n日期：______年______月______日", bold=True, size=11)
    set_cell_text(sig2.rows[1].cells[0], "（附表一与本附件协议正文具有同等法律效力）", size=10, color=GREY, align_center=True)
    set_cell_text(sig2.rows[1].cells[1], "（本附表共 1 页）", size=10, color=GREY, align_center=True)

    out = "/workspace/deliverables/附件四-绿城中国-潮鸣外滩-晚宴冠名战略合作伙伴专项合作协议.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
