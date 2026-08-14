#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成王德峰老师 10 月 31 日活动合作方案：对外版 + 内部版 Word。"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output"

NAVY = RGBColor(0x1A, 0x3A, 0x5C)
NAVY_HEX = "1A3A5C"
GOLD = RGBColor(0x8B, 0x69, 0x14)
GOLD_HEX = "8B6914"
INK = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED = RGBColor(0x8B, 0x1E, 0x1E)
RED_HEX = "8B1E1E"
CREAM_HEX = "F7F3EA"
SOFT_HEX = "EEF3F8"
AMBER_HEX = "F8F1E3"
ROSE_HEX = "F8EEEE"
GREEN_HEX = "E8F2EA"
ROW_ALT_HEX = "F5F7FA"


def set_run_font(run, *, name="微软雅黑", size=11, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge not in kwargs:
            continue
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn(f"w:{key}"), str(value))


def set_table_full_width(table, widths_cm):
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = int(sum(widths_cm) * 567)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[idx])
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths_cm[idx] * 567)))
            tc_w.set(qn("w:type"), "dxa")
            v_align = OxmlElement("w:vAlign")
            v_align.set(qn("w:val"), "center")
            tc_pr.append(v_align)


def prevent_row_split(row):
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=9, color=MUTED)


def setup_section(doc, header_text, confidential=False):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run(header_text)
    set_run_font(run, size=9, color=RED if confidential else MUTED, bold=confidential)
    hp.paragraph_format.space_after = Pt(2)

    p_pr = hp._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), RED_HEX if confidential else NAVY_HEX)
    pbdr.append(bottom)
    p_pr.append(pbdr)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = fp.add_run("第 ")
    set_run_font(left, size=9, color=MUTED)
    add_page_number(fp)
    right = fp.add_run(" 页")
    set_run_font(right, size=9, color=MUTED)


def add_para(
    doc,
    text,
    *,
    size=11,
    bold=False,
    color=INK,
    align=None,
    space_before=0,
    space_after=8,
    first_line=False,
    line_spacing=1.35,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if first_line:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_rich(doc, segments, *, size=11, space_before=0, space_after=8, first_line=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.35
    if first_line:
        pf.first_line_indent = Cm(0.74)
    for text, bold, color in segments:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color or INK)
    return p


def add_h1(doc, text):
    return add_para(
        doc,
        text,
        size=16,
        bold=True,
        color=NAVY,
        space_before=16,
        space_after=8,
    )


def add_h2(doc, text):
    return add_para(
        doc,
        text,
        size=13,
        bold=True,
        color=NAVY,
        space_before=12,
        space_after=6,
    )


def add_h3(doc, text):
    return add_para(
        doc,
        text,
        size=12,
        bold=True,
        color=GOLD,
        space_before=8,
        space_after=4,
    )


def add_quote_box(doc, title, body, fill=CREAM_HEX, title_color=GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "12", "color": GOLD_HEX, "space": "0"},
        left={"val": "single", "sz": "24", "color": GOLD_HEX, "space": "0"},
        bottom={"val": "single", "sz": "12", "color": GOLD_HEX, "space": "0"},
        right={"val": "single", "sz": "12", "color": GOLD_HEX, "space": "0"},
    )
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(8)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(title)
    set_run_font(r1, size=11, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing = 1.3
    r2 = p2.add_run(body)
    set_run_font(r2, size=11, color=INK)
    set_table_full_width(table, [16.4])
    add_para(doc, "", size=6, space_after=6)
    return table


def add_alert_box(doc, title, body, fill=ROSE_HEX, accent=RED_HEX):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "8", "color": accent, "space": "0"},
        left={"val": "single", "sz": "24", "color": accent, "space": "0"},
        bottom={"val": "single", "sz": "8", "color": accent, "space": "0"},
        right={"val": "single", "sz": "8", "color": accent, "space": "0"},
    )
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(8)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(title)
    set_run_font(r1, size=11, bold=True, color=RED)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing = 1.3
    r2 = p2.add_run(body)
    set_run_font(r2, size=11, color=INK)
    set_table_full_width(table, [16.4])
    add_para(doc, "", size=6, space_after=6)


def add_bullets(doc, items, *, size=11, color=INK):
    for item in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(0.5)
        pf.first_line_indent = Cm(-0.4)
        pf.space_before = Pt(1)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.3
        run = p.add_run("●  " + item)
        set_run_font(run, size=size, color=color)


def cell_text(cell, text, *, bold=False, size=10, color=INK, align=None, fill=None):
    cell.text = ""
    if fill:
        shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, *, header_fill=NAVY_HEX, emphasize_col=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell_text(
            table.rows[0].cells[i],
            header,
            bold=True,
            size=10,
            color=WHITE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=header_fill,
        )
    prevent_row_split(table.rows[0])
    for r_idx, row in enumerate(rows):
        fill = ROW_ALT_HEX if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            is_emph = emphasize_col is not None and c_idx == emphasize_col
            cell_text(
                table.rows[r_idx + 1].cells[c_idx],
                str(val),
                bold=is_emph or c_idx == 0,
                size=10,
                color=NAVY if is_emph else INK,
                align=WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT,
                fill=fill,
            )
        prevent_row_split(table.rows[r_idx + 1])
    set_table_full_width(table, widths)
    add_para(doc, "", size=6, space_after=8)
    return table


def add_cover_line(doc, text, *, size=12, bold=False, color=MUTED, space_after=4):
    add_para(
        doc,
        text,
        size=size,
        bold=bold,
        color=color,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=space_after,
    )


def new_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.font.color.rgb = INK
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return doc


def build_external():
    doc = new_doc()
    setup_section(doc, "王德峰老师 10 月 31 日活动 · 联合运营合作方案（对外沟通版）")

    add_para(doc, "", size=14, space_after=18)
    add_para(doc, "王德峰老师", size=14, bold=True, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(
        doc,
        "2026 年 10 月 31 日大型活动",
        size=22,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_para(
        doc,
        "联合运营合作方案",
        size=26,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )
    add_cover_line(doc, "对外沟通版", size=13, bold=True, color=GOLD, space_after=18)
    add_quote_box(
        doc,
        "合作一句话",
        "贵方提供王德峰老师及内容 / IP 资源，我方负责商业化、组织、渠道和现场交付。"
        "双方按实际贡献分钱，把 10 月 31 日做成可复制的第一期联合项目。",
    )
    add_para(doc, "", size=8, space_after=24)
    add_cover_line(doc, "角色定位：联合主办 / 独家运营   ·   非渠道代理", size=12, bold=True, color=NAVY)
    add_cover_line(doc, "版本：V1.0　　日期：2026 年 8 月", size=10.5, color=MUTED, space_after=2)
    add_cover_line(doc, "本文仅供双方沟通合作结构使用，具体条款以正式协议为准。", size=9.5, color=MUTED)

    doc.add_page_break()

    add_h1(doc, "一、合作背景与共识")
    add_para(
        doc,
        "王德峰老师 10 月 31 日的活动，具备做成一场高质量人文思想活动的基础。老师本人的学术影响力、思想深度和公众号召力，是这次活动最核心的内容资产。要把这场活动从“有老师”变成“有收入、有口碑、可复制”，还需要场地、票务、招商、传播、客户服务和现场履约一整套商业化能力。",
        first_line=True,
    )
    add_para(
        doc,
        "双方此前沟通中，贵方提出过两种合作方向：一是将当天活动授权给我方运营；二是双方共同推进，按报名来源进行分成。我方认同“一起把这件事做成”的方向，同时建议把合作身份从“帮忙卖票”提升为“联合主办、联合运营”。这样更匹配双方真实贡献，也更有利于把 10 月 31 日做成后续系列的第一期，而不是一锤子买卖。",
        first_line=True,
    )
    add_quote_box(
        doc,
        "本次沟通希望达成的共识",
        "不把合作做成一次性买断，也不把我方理解成票务渠道。把 10 月 31 日设立为独立项目：贵方负责老师、课程内容及其原有资源；我方负责活动运营、城市资源、企业客户、场地、传播和商务开发。双方按收入来源、成本承担和履约责任分别结算。",
    )

    add_h1(doc, "二、合作定位：联合运营，而不是渠道代理")
    add_para(
        doc,
        "如果只按“渠道代理”理解，分成逻辑会变成：谁有老师资源，谁拿大头；谁卖票，谁拿佣金。这会低估把一场大型活动真正落地所需要的工作量，也不利于双方把事情做成。",
        first_line=True,
    )
    add_para(doc, "因此，本方案建议双方首先确认身份：", first_line=True)
    add_quote_box(
        doc,
        "核心定位",
        "我方不是票务渠道，而是联合运营方 / 联合主办方之一。贵方是内容与老师资源方。双方共同对 10 月 31 日项目负责，但分工不同、贡献不同、分成也不同。",
        fill=SOFT_HEX,
        title_color=NAVY,
    )
    add_h2(doc, "2.1 为什么必须按来源结算，而不是笼统五五分")
    add_para(
        doc,
        "本方案明确建议：不要笼统谈“利润五五分”，也不要笼统谈“流水五五分”。一定要按照收入来源、谁承担成本、谁负责履约，分别结算。",
        first_line=True,
    )
    add_para(
        doc,
        "原因很直接。一场大型活动的收入，至少会来自个人报名、共同渠道、赞助招商、企业包场等不同路径；成本则至少包括老师相关费用、场地、制作、人员、投放、客服和退款风险。如果把所有收入混在一起再五五分，最后很容易出现一种失衡：我方找场地、组团队、做传播、卖票、找赞助、处理退款，贵方以老师资源参与，但结算时无法体现双方真实贡献。",
        first_line=True,
    )

    add_h1(doc, "三、双方分工")
    add_para(
        doc,
        "建议把 10 月 31 日活动成立为一个独立项目。双方在项目中的角色如下。",
        first_line=True,
    )
    add_table(
        doc,
        ["事项", "贵方（资源方）", "我方（联合运营方）"],
        [
            ["老师确认与内容", "确认王德峰老师出席、主题、时长、形式", "配合议程包装、对外表述与现场呈现"],
            ["IP 与宣传授权", "授权使用老师姓名、肖像、主题做活动宣传", "统一视觉、文案、海报、传播节奏"],
            ["场地与现场", "如有既有场地资源可共同评估", "负责场地洽谈、布置、流程、人员与现场交付"],
            ["票务与报名", "负责贵方自有渠道的报名转化", "负责售票系统、客服、退改、现场核验"],
            ["企业与团体", "对接贵方既有机构资源", "开发企业包场、团体客户、商务合作"],
            ["赞助招商", "对接贵方既有赞助线索", "负责赞助方案、商务谈判与履约"],
            ["客户服务", "协助老师侧沟通与内容答疑", "负责报名咨询、社群运营、到场服务"],
            ["成本与风险", "承担老师相关约定成本", "承担场地、制作、人员、运营等落地成本"],
            ["主业协同", "可持续聚焦科学课程等既有业务", "把本次活动做成可复制的商业化样板"],
        ],
        [4.2, 6.1, 6.1],
    )
    add_para(
        doc,
        "这样分工的好处是：贵方不必把精力从既有科学课程培训中抽走，去承担自己并不主攻的场地、组织、票务和现场交付；我方则可以按联合主办标准投入，而不是按“帮忙卖几张票”的力度投入。",
        first_line=True,
        color=MUTED,
        size=10.5,
    )

    add_h1(doc, "四、推荐方案：按来源分成的联合运营")
    add_h2(doc, "4.1 建议采用的主方案")
    add_para(
        doc,
        "我方建议优先采用以下结构，作为双方合作的基础方案。",
        first_line=True,
    )
    add_table(
        doc,
        ["收入来源", "我方比例", "贵方比例", "对应逻辑"],
        [
            ["我方独立开发的报名", "60%", "40%", "含获客、运营、客服、现场与风险"],
            ["贵方独立开发的报名", "25%", "75%", "尊重老师资源与贵方原有渠道"],
            ["双方共同渠道产生的报名", "50%", "50%", "共同开发、共同维护"],
            ["我方引入的赞助", "70%", "30%", "商务开发与赞助履约主要由我方完成"],
            ["贵方引入的赞助", "25%", "75%", "贵方线索，我方协助方案与执行"],
            ["我方开发的团体客户 / 企业包场", "60%", "40%", "企业开发、方案定制与交付在我方"],
        ],
        [4.6, 2.6, 2.6, 6.6],
        emphasize_col=1,
    )
    add_para(
        doc,
        "上表是书面建议比例，不是“渠道佣金”。我方 60% 对应的是：获客 + 活动运营 + 场地统筹 + 商务 + 客户服务 + 现场执行 + 风险承担。贵方 40% 对应的是老师资源、内容价值和原有渠道。如果后续双方确认进入更深度、更长期的联合，比例可以在此基础上微调，但不宜再回到“渠道代理 20%”的框架。",
        first_line=True,
    )
    add_h2(doc, "4.2 结算原则")
    add_bullets(
        doc,
        [
            "先分清收入来源，再分成。每一笔报名、赞助、包场都要标注来源，双方定期对账。",
            "先明确成本承担，再谈净额。场地、制作、人员、投放、平台佣金、税费、退款，原则上由实际发生方承担；共同决策的费用双方书面确认后分摊。",
            "先确认履约责任，再确认分成。谁承诺了场地、票务、赞助权益、现场服务，谁对相应履约负责。",
            "老师相关费用、差旅、接待等，单列并事先书面确认，不混入模糊的“项目利润”。",
        ],
    )

    add_h1(doc, "五、备选机制：阶梯分成")
    add_para(
        doc,
        "如果贵方希望在项目冷启动阶段获得更清晰的保障感，同时希望活动做大后老师 IP 的价值能够进一步提高，我方建议在主方案之外增加阶梯分成，而不是改回大额一次性买断。",
        first_line=True,
    )
    add_table(
        doc,
        ["活动票务收入区间", "我方", "贵方", "设计用意"],
        [
            ["0–10 万元部分", "70%", "30%", "冷启动阶段运营投入和风险最大"],
            ["10–30 万元部分", "60%", "40%", "项目跑通后回归联合运营均衡点"],
            ["30 万元以上部分", "50%", "50%", "卖爆后老师 IP 价值充分体现"],
        ],
        [4.4, 2.4, 2.4, 7.2],
        header_fill="3D5A40",
    )
    add_para(
        doc,
        "阶梯分成比一次性收取固定授权费更健康。项目还没卖起来时，是运营侧最辛苦、风险最大的阶段；一旦项目卖好，老师资源的价值会明显体现出来，贵方比例随之提高。双方都是在分享结果，而不是一方先把风险转给另一方。",
        first_line=True,
    )

    add_h1(doc, "六、关于一次性买断授权")
    add_para(
        doc,
        "贵方此前提出过将 10 月 31 日当天活动全部授权给我方、并收取约 10–18 万元的设想。我方理解这是希望获得确定回报。但从把活动做成的角度，我方不建议现在就采用大额一次性买断。",
        first_line=True,
    )
    add_para(
        doc,
        "原因不在于否定老师资源的价值，而在于当前关键经营变量尚未闭合：老师出席的书面确认、活动内容与时长、可售座位、票价体系、历史同类活动实际转化、独家宣传与售票权限、姓名肖像使用权，以及老师因故取消时的退还机制。在这些事项明确之前，买断本质上是运营方先支付一笔保底，再自行承担场地、宣传、人员和退款风险。这会抑制我方继续加码投入，也不利于双方把盘做大。",
        first_line=True,
    )
    add_h2(doc, "6.1 若仍需讨论买断，需先闭合的事项")
    add_table(
        doc,
        ["序号", "需事先明确的事项", "为什么必须先明确"],
        [
            ["1", "王德峰老师确定出席的书面确认", "没有书面确认，授权本身无法落地"],
            ["2", "活动内容、时长、形式", "决定票价、座位和传播卖点"],
            ["3", "可售座位数量", "决定收入上限"],
            ["4", "票价体系", "决定测算能否成立"],
            ["5", "同类公开商业活动的实际报名情况", "判断市场真实转化，而不是感觉"],
            ["6", "是否允许我方独家招商、宣传、售票", "没有独家，买断没有商业意义"],
            ["7", "姓名、肖像、活动内容的宣传使用权", "没有授权，无法做正式传播"],
            ["8", "老师临时取消时授权费是否 100% 退还", "这是买断能否谈的前提条件"],
        ],
        [1.6, 7.0, 7.8],
    )
    add_para(
        doc,
        "因此，本方案建议：不采用 10–18 万元固定授权费，改为按来源分成，必要时叠加阶梯分成。如贵方确需一定确定性，可在正式协议中另行协商小额基础成本分摊，而不是先做大额买断。",
        first_line=True,
    )

    add_h1(doc, "七、把 10 月 31 日做成第一期，而不是一次性生意")
    add_para(
        doc,
        "我方认为，这是这次合作里最值得争取的一层。10 月 31 日不应该只是联合卖票，而应该是双方第一次联合项目。",
        first_line=True,
    )
    add_table(
        doc,
        ["层次", "内容", "对双方的意义"],
        [
            ["第一期", "10 月 31 日王德峰老师活动", "验证分工、分成、运营和转化"],
            ["第二期", "复盘后共同开发后续哲学 / 人文场次", "把单场做成系列"],
            ["长期", "探索哲学、人文与科学课程的协同转化", "贵方主业获得稳定入口，我方获得可持续运营权"],
        ],
        [3.2, 6.6, 6.6],
    )
    add_quote_box(
        doc,
        "给双方的共同价值",
        "贵方可以继续把重心放在科学课程培训等主业上，同时把王德峰老师活动做成高质量流量入口；我方则负责把高端知识产品做成可售、可交付、可复盘的商业化体系。第一期成功后，再共同开发后续课程，而不是把所有压力压在 10 月 31 日这一天。",
    )

    add_h1(doc, "八、项目资产与客户资源保护")
    add_para(
        doc,
        "联合运营要能长期走下去，必须把项目资产讲清楚。以下条款建议对等写入协议，保护的是双方，不是单方。",
        first_line=True,
    )
    add_table(
        doc,
        ["事项", "建议原则"],
        [
            ["客户数据归属", "项目报名数据由双方共有；各自原有客户仍归各自。未经书面同意，不得把对方原有客户挪作他用。"],
            ["微信群归属", "本次活动官方群为项目资产，合作期内由联合运营方维护，任何一方不得擅自解散或转移。"],
            ["报名数据使用权", "仅用于本次活动履约、现场服务及双方书面同意的后续联合项目。"],
            ["后续课程二次转化", "基于本次活动产生的新客户，二次转化需双方事先约定分成，不得绕开对方单独成交。"],
            ["赞助商资源保护期", "我方引入的赞助商，保护期内对方不得绕开我方直接签约同类项目；反之亦然。"],
            ["不得绕开成交", "合作期内及保护期内，任何一方不得绕开另一方，就本次老师资源或本次客户直接成交。"],
        ],
        [4.4, 12.0],
    )
    add_para(
        doc,
        "这些条款看起来像“约束”，实际上是为了让双方敢于把真资源拿出来：我方敢于把场地、企业客户、票务体系和运营团队投入进来；贵方也敢于把老师资源和原有渠道投入进来。没有保护条款，联合运营很难做成深度合作。",
        first_line=True,
    )

    add_h1(doc, "九、成本、对账与风险安排")
    add_h2(doc, "9.1 成本透明")
    add_bullets(
        doc,
        [
            "设立项目独立账目，门票、赞助、包场、退款、成本分类入账。",
            "超过事先约定金额的支出，须双方书面确认后发生。",
            "每周或每两周对账一次，活动结束后 15 个工作日内完成结算。",
            "平台手续费、支付通道费、税费按实际发生列支，不计入任何一方“分成让利”。",
        ],
    )
    add_h2(doc, "9.2 风险安排")
    add_bullets(
        doc,
        [
            "老师因故不能出席：已售票退款路径、沉没成本分担、是否改期，须事先写明。",
            "场地变更或不可抗力：双方按共同决策原则处理，不单方面归责。",
            "退票与投诉：由负责售票和现场的运营方处理，规则对外统一。",
            "宣传口径：凡使用王德峰老师姓名、肖像、观点的内容，须符合双方确认的授权范围。",
        ],
    )

    add_h1(doc, "十、建议的合作路径")
    add_table(
        doc,
        ["步骤", "事项", "目标结果"],
        [
            ["1", "确认合作身份", "书面确认我方为联合主办 / 独家运营方之一，而非渠道代理"],
            ["2", "确认老师与内容", "出席书面确认、主题、时长、形式、宣传授权"],
            ["3", "确认分成结构", "按来源分成 + 可选阶梯分成，不采用大额买断"],
            ["4", "确认成本与对账", "项目独立账、成本清单、结算周期"],
            ["5", "确认资产保护", "客户、社群、赞助、二次转化和不得绕开条款"],
            ["6", "签署合作协议", "以本方案为讨论底稿，形成可执行合同"],
            ["7", "启动第一期筹备", "场地、票价、视觉、招商、售票同步展开"],
        ],
        [2.2, 5.6, 8.6],
    )

    add_h1(doc, "十一、给贵方的合作建议（小结）")
    add_para(
        doc,
        "如果用一页纸把本方案说完，就是下面四句话：",
        first_line=True,
    )
    add_bullets(
        doc,
        [
            "请把我方看成联合运营方，而不是卖票渠道。",
            "10 月 31 日用按来源分成，而不是 10–18 万元买断。",
            "我方独立开发的报名按 60% : 40%，我方引入的赞助按 70% : 30%；贵方自己的报名，贵方拿大头。",
            "这一场是第一期。做成了，再一起做后续哲学、人文和科学课程的商业化。",
        ],
    )
    add_quote_box(
        doc,
        "我方希望推进的结果",
        "不交大额买断费，确立联合运营身份，按贡献分钱，成本透明，客户资源受到保护。这样贵方老师资源的价值能随项目放大而提高，我方也愿意按主办标准持续投入。双方不是在争“20% 还是 70% 的渠道佣金”，而是在共同定义一种可长期复制的合作模式。",
    )
    add_para(
        doc,
        "以上方案供双方讨论。具体比例、授权范围和合同条款，可在确认合作身份后进一步细化。",
        first_line=True,
        space_before=8,
    )
    add_para(
        doc,
        "（本文件为沟通方案，不构成要约。最终合作内容以双方签署的协议为准。）",
        size=9.5,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=18,
    )

    path = OUT / "王德峰老师10月31日活动_联合运营合作方案_对外版.docx"
    doc.save(path)
    return path


def build_internal():
    doc = new_doc()
    setup_section(
        doc,
        "内部留存 · 请勿外传  |  王德峰老师 10 月 31 日活动合作谈判方案",
        confidential=True,
    )

    add_para(doc, "", size=12, space_after=10)
    add_para(
        doc,
        "【内部留存 · 请勿发给对方】",
        size=12,
        bold=True,
        color=RED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_para(doc, "王德峰老师 · 2026 年 10 月 31 日", size=14, bold=True, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(
        doc,
        "合作谈判内部方案",
        size=26,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_cover_line(doc, "给自己看的版本  ·  谈判策略 / 底线 / 话术 / 合同要点", size=12, bold=True, color=NAVY)
    add_para(doc, "", size=8, space_after=10)
    add_alert_box(
        doc,
        "使用提醒",
        "本文件只给己方内部使用。对外请使用《联合运营合作方案（对外版）》。内部版写了对方经验不足、买断不划算、70% 锚点、50% 底线、防被掏空条款等，发出去会直接破坏谈判。",
    )
    add_cover_line(doc, "版本：V1.0　　日期：2026 年 8 月　　编制：己方内部", size=10.5, color=MUTED)

    doc.add_page_break()

    add_h1(doc, "一、局势判断")
    add_para(
        doc,
        "对方想和你合作王德峰老师 10 月 31 日的大活动。他原来有两个方案：方案一，把当天活动全部授权给你，向你要 10–18 万；方案二，双方一起做，从你这边报名里抽成 20%。你的直觉是：如果真合作，你这边报名至少应抽 50%–70%。",
        first_line=True,
    )
    add_para(
        doc,
        "从现在掌握的情况看，对方完全没有组织这类活动的经验，连场地、筹划、现场交付都不在他的能力圈里。他真正想做的，是把重心放到其他科学课程培训上。也就是说：老师资源在他手里，但把“资源”变成“收入”的大量工作，很可能由你来完成。",
        first_line=True,
    )
    add_quote_box(
        doc,
        "内部结论（先记住这一句）",
        "不建议现在直接付 10–18 万买断当天授权。更合适的是把合作重新定义成：他提供王德峰老师及内容 / IP 资源，你负责商业化、组织、渠道和现场交付，按实际贡献分钱，而不是你替他承担全部经营风险。",
    )
    add_h2(doc, "1.1 这件事真正的利害")
    add_para(
        doc,
        "你最怕的不是这一次少赚几万元，而是：你辛辛苦苦给他搭完场地、微信群、客户名单、企业资源、票务体系和运营团队，10 月 31 日之后，他拿着这套体系自己继续卖其他课程。所以合同里的数据归属、社群归属、二次转化权和不得绕开条款，可能比你究竟拿 60% 还是 70% 更重要。",
        first_line=True,
    )

    add_h1(doc, "二、对方两个原方案为什么都不能直接接")
    add_h2(doc, "2.1 方案一：10–18 万买断")
    add_para(
        doc,
        "买断的本质，是你给他保底，他把风险全部转嫁给你。",
        first_line=True,
        bold=True,
    )
    add_para(
        doc,
        "举例：你付 15 万授权费，场地再投入 5 万，宣传、制作、人员再投入 5 万，活动开始前就已经承担约 25 万风险。按一线城市 3 小时讲座的合理实收客单约 520 元计，扣掉平台、退款、税费和渠道后，单个付费用户实际贡献大约只有 460–480 元。250,000 ÷ 470 ≈ 532 名付费用户才能覆盖这 25 万。若只能卖掉容量的 80%，场馆至少需要约 665 座。没有这个把握，买断不成立。更不能为了覆盖买断费，把普通票硬抬到 699–999 元。",
        first_line=True,
    )
    add_table(
        doc,
        ["项目", "金额 / 数量", "含义"],
        [
            ["授权费（取其例）", "15 万元", "你先付给对方的保底"],
            ["场地", "5 万元", "你还要再垫"],
            ["宣传 / 制作 / 人员", "5 万元", "你还要再垫"],
            ["活动前累计风险", "约 25 万元", "还没开场就已经出去了"],
            ["客单价假设（实收）", "520 元", "一线 3 小时讲座合理目标，不是 699"],
            ["扣费后单人贡献", "约 470 元", "已计平台、退款、税费、渠道"],
            ["覆盖 25 万所需付费人数", "约 532 人", "售出率 80% 时场馆约需 665 座"],
        ],
        [5.4, 4.4, 6.6],
        header_fill=RED_HEX,
    )
    add_h2(doc, "2.2 方案二：你的报名只抽 20%")
    add_para(
        doc,
        "20% 只有在一种身份下才说得通：你只是票务渠道，对方负责场地、组织、客服、现场和风险。现在对方这些都没有，还准备把精力放到别的科学课程上。再给你 20%，就是让你干主办的活、拿渠道的钱。",
        first_line=True,
    )
    add_para(
        doc,
        "所以不要跟他在“20% 还是 70%”上硬磕，那是在他设定的渠道框架里谈判。你要先改身份：我不是票务渠道，我是联合运营方。这是整个谈判最重要的一句话。",
        first_line=True,
    )

    add_h1(doc, "三、谈判总策略：先改身份，再谈比例")
    add_table(
        doc,
        ["层次", "对外怎么说", "对内真正要的"],
        [
            ["身份", "联合主办 / 独家运营方之一", "不要被定义成代理、渠道、分销"],
            ["结构", "按收入来源分别结算", "避免笼统利润五五分、流水五五分"],
            ["比例", "书面给 60% : 40%", "口头可从 70% 起，底线 50%"],
            ["买断", "不建议现在买断", "除非独家权利完整且测算过关，否则不碰"],
            ["周期", "10 月 31 日是第一期", "做成可复制体系，而不是一次性生意"],
            ["资产", "客户和数据对等保护", "防他拿着你搭的体系去卖科学课"],
        ],
        [3.0, 6.7, 6.7],
    )
    add_quote_box(
        doc,
        "对外方案已经帮你把身份改好了",
        "发给对方的 Word，写的是 60% : 40%，语气是联合运营、优势互补、可长期合作。没有写他没经验，没有写 70% 锚点和 50% 底线，也没有写“防他掏空”。口头谈判时，你可以按下面的锚点来。",
    )

    add_h1(doc, "四、数字怎么谈")
    add_h2(doc, "4.1 书面给 60%，口头从 70% 起，底线 50%")
    add_para(
        doc,
        "不要一上来死咬 70%。70% 是谈判锚点，60% 是目标成交点，50% 是心理底线。对外稿已经按 60% 写好，逻辑是：考虑老师资源价值，深度合作可以落到 60% : 40%。",
        first_line=True,
    )
    add_table(
        doc,
        ["收入来源", "口头锚点（己方）", "书面目标", "底线（己方）", "对方大概拿"],
        [
            ["己方独立报名", "70%", "60%", "50%", "30%–50%"],
            ["对方独立报名", "30%", "25%", "20%", "70%–80%"],
            ["共同渠道报名", "50%", "50%", "50%", "50%"],
            ["己方引入赞助", "80%", "70%", "70%", "20%–30%"],
            ["对方引入赞助", "30%", "25%", "20%", "70%–80%"],
            ["己方企业包场", "70%", "60%", "50%", "30%–50%"],
        ],
        [3.6, 3.2, 2.8, 3.2, 3.6],
        emphasize_col=2,
    )
    add_para(
        doc,
        "口头可以这样说：“我这边独立开发的报名和客户，我方 70%、你方 30%。考虑老师资源价值，如果最终双方深度合作，可以调整到 60% : 40%。” 这比直接说“我要 70%”更有逻辑。因为你的 70% 不是渠道佣金，而是获客 + 活动运营 + 场地 + 商务 + 客户服务 + 现场执行 + 风险承担。",
        first_line=True,
    )
    add_h2(doc, "4.2 三种方案的选择顺序")
    add_table(
        doc,
        ["排序", "方案", "什么时候用", "内部态度"],
        [
            ["第一选择", "联合运营 + 按来源分成", "现在就应该推这个", "优先成交"],
            ["第二选择", "小额保底 3–5 万 + 高比例分成", "对方一定要“有点保障”时", "可以让，但不主动先出"],
            ["第三选择", "真正买断 10–18 万", "只有独家权利完整，且你有把握做到 30 万、50 万甚至更高收入", "按当前信息，现在不接"],
        ],
        [2.4, 5.0, 5.6, 3.4],
        header_fill="3D5A40",
    )
    add_alert_box(
        doc,
        "第一方案风险最大，现在不要接",
        "这里的“第一方案”是对方说的买断，不是你的第一选择。按目前信息：他没组织经验、没场地、重心在别的课，还要你先掏 10–18 万。这是让你替他保底。",
    )

    add_h1(doc, "五、如果对方坚持买断，先要这 8 个数字")
    add_para(
        doc,
        "不要在“给不给 15 万”上立刻回答。先让他把下面这些事项拿出来。拿不出来，就说明买断还不成立。",
        first_line=True,
    )
    add_table(
        doc,
        ["序号", "必须先拿到的东西", "过关标准"],
        [
            ["1", "王德峰老师确定出席的书面确认", "有老师或合法授权方签字 / 盖章，不是口头“差不多”"],
            ["2", "活动内容、时长、形式", "主题、时长、互动形式明确，能用来做传播和定价"],
            ["3", "可售座位", "有场地方案或明确座位数，能算收入上限"],
            ["4", "票价体系", "有建议票价或允许你定价的书面授权"],
            ["5", "过去类似公开商业活动的实际报名人数", "要实际人数，不要“应该能火”"],
            ["6", "是否允许你独家招商、宣传、售票", "必须独家。不独家就没有买断的意义"],
            ["7", "姓名、肖像、活动内容宣传权", "必须允许你做正式传播，否则 15 万买了个不能说的活动"],
            ["8", "老师临时取消，10–18 万是否 100% 退还", "必须 100% 退还。这一条尤其重要"],
        ],
        [1.6, 7.4, 7.4],
        header_fill=RED_HEX,
    )
    add_para(
        doc,
        "第 8 条是买断的生死线。没有 100% 退还，等于你既买了不确定性，又买了沉没成本。场地、宣传、人员的钱已经出去，授权费还不退，这笔合作不能做。",
        first_line=True,
    )

    add_h1(doc, "六、阶梯分成：给对方“保障感”，但不先掏 15 万")
    add_para(
        doc,
        "如果他希望得到更多保障，就用阶梯分成替代固定授权费。对外稿已经写了同一张表，内部要理解它的作用：把冷启动风险留在你可控的分成里，而不是先现金出去。",
        first_line=True,
    )
    add_table(
        doc,
        ["票务收入", "己方", "对方", "你为什么接受"],
        [
            ["0–10 万元部分", "70%", "30%", "这阶段你最辛苦、风险最大"],
            ["10–30 万元部分", "60%", "40%", "项目跑通，回到目标成交点"],
            ["30 万元以上部分", "50%", "50%", "卖爆后让出一点，换他继续配合长期合作"],
        ],
        [4.4, 2.8, 2.8, 6.4],
    )
    add_para(
        doc,
        "小额保底只在对方咬住“必须先有一笔钱”时再用。上限建议 3–5 万，并且要对应老师相关真实成本，而不是空的“授权费”。保底可以抵扣后续分成，避免你既付了保底又按全额再分一次。",
        first_line=True,
    )

    add_h1(doc, "七、把 10 月 31 日升级成体系，而不是帮他卖一张票")
    add_para(
        doc,
        "对方真正想做科学课程培训，说明王德峰老师活动可能只是他的流量入口。这对你反而是机会：你不是去抢他的课，而是帮他搭一套高端知识产品的商业化体系。这时你谈 60% 甚至 70%，逻辑会顺很多。",
        first_line=True,
    )
    add_table(
        doc,
        ["如果只做成", "如果升级做成"],
        [
            ["帮他卖王德峰老师的票", "帮他搭一套高端知识产品商业化体系"],
            ["一次性分成争议", "第一期联合项目，后续可复制"],
            ["他拿老师，你干活", "他负责老师和内容，你负责城市、企业、场地、传播、商务"],
            ["活动结束各走各的", "第一期成功后共同开发后续哲学、人文、科学课程"],
        ],
        [8.2, 8.2],
    )
    add_para(
        doc,
        "对外稿已经用“第一期”把这件事写柔了。内部要清楚：这既是加分项，也是防护项。没有后续合作约束，你搭的体系就是在给他做基建。",
        first_line=True,
    )

    add_h1(doc, "八、合同里必须写死的条款")
    add_para(
        doc,
        "下面这几条的价值，可能比 60% 还是 70% 更大。对外稿用了“对等保护”的口气；内部要按防被掏空来执行。",
        first_line=True,
    )
    add_table(
        doc,
        ["条款", "必须写成什么样", "为什么"],
        [
            ["客户数据归属", "项目新客户共有；各自原有客户仍归各自。禁止把对方原有名单挪走。", "防止他直接拿走你的企业客户"],
            ["微信群归属", "官方群为项目资产，由你方维护。禁止擅自解散、转移群主、导出名单另建群。", "群往往是后续转化的真正入口"],
            ["报名数据使用权", "仅限本次履约和书面同意的联合项目。禁止用于对方单独的科学课销售。", "这是最可能被“顺手转化”的地方"],
            ["后续课程二次转化权", "基于本次活动产生的新客户，二次转化必须双方事先分成。", "否则你等于给他免费获客"],
            ["赞助商保护期", "你引入的赞助商，至少 12 个月不得被绕开签约同类项目。", "赞助关系一旦暴露，很容易被截胡"],
            ["不得绕开直接成交", "合作期 + 保护期内，不得绕开你就老师资源或本次客户成交。", "没有这一条，分成条款会被架空"],
            ["独家运营范围", "10 月 31 日该场次的场地、票务、招商、传播以你方为主执行。", "避免他一边合作一边另找渠道"],
            ["退出与交接", "合作终止后，群、数据、物料如何交接，写清楚时间和范围。", "防止不欢而散后资产被单方面控制"],
        ],
        [3.6, 6.8, 6.0],
        header_fill=RED_HEX,
    )

    add_h1(doc, "九、谈判话术（可直接用）")
    add_h3(doc, "开场：先改身份")
    add_quote_box(
        doc,
        "可直接说",
        "我不是来做票务渠道的，也不是来买一场授权的。我的理解是：你提供王德峰老师和内容资源，我来做商业化、组织、渠道和现场交付。我们是联合运营，不是我帮你卖票。",
    )
    add_h3(doc, "被问“为什么不是 20%”")
    add_quote_box(
        doc,
        "可直接说",
        "如果我只负责转发报名链接，20% 说得通。但场地、团队、海报、社群、售票、赞助、退款和现场，这些如果主要由我来做，20% 就不是合作，是让我替你承担经营。所以要按来源分：我自己带来的客户，我拿 60% 到 70%；你自己带来的客户，你拿大头。",
    )
    add_h3(doc, "被问“那你就买断吧，10–18 万”")
    add_quote_box(
        doc,
        "可直接说",
        "买断可以谈，但要先把老师书面确认、内容、座位、票价、历史转化、独家权利和取消退款这几件事闭合。否则我付的不是授权费，是替你保底。更健康的做法是不收大额买断，改成阶梯分成：前面你我共担启动，后面卖好了你的比例自然提高。",
    )
    add_h3(doc, "把长期合作抛出来")
    add_quote_box(
        doc,
        "可直接说",
        "10 月 31 日不要做成一次性生意。你真正想做的是科学课程，这场活动可以成为入口。我负责把运营、城市资源、企业客户和票务体系搭起来。第一期做成了，后面哲学、人文、科学课程可以一起做。这样你不用从主业里分心去办活动，我也不只是帮你卖一张票。",
    )
    add_h3(doc, "对方开始谈五五分时")
    add_quote_box(
        doc,
        "可直接说",
        "五五分可以出现在共同开发的客户上，也可以出现在项目卖爆之后的增量上。但不要一上来把所有流水或利润五五分。我找来的企业包场和赞助，如果也五五分，等于我的商务成果被平均掉了。按来源分，对你我都更清楚。",
    )

    add_h1(doc, "十、红线与让步清单")
    add_h2(doc, "10.1 可以让的")
    add_bullets(
        doc,
        [
            "己方独立报名从 70% 让到 60%，必要时让到 50%。低于 50% 就不再是联合运营。",
            "对方自己带来的报名，可以让他拿 75%–80%，你拿 20%–25%。这反而显得公平。",
            "阶梯分成里，30 万以上部分可以五五开，用来换他接受前面高分成。",
            "如果他非要“有一笔钱”，最多接受 3–5 万小额保底，且抵扣后续分成。",
            "品牌露出、联合主办署名、老师接待规格，这些可以大方给。",
        ],
    )
    add_h2(doc, "10.2 不能让的")
    add_bullets(
        doc,
        [
            "现在就付 10–18 万买断，且没有书面出席确认和 100% 退还。",
            "被定义成渠道代理，只拿你自己报名的 20%–30%，还要你去找场地、组班子。",
            "笼统利润五五分 / 流水五五分，成本却主要由你承担。",
            "客户名单、微信群、赞助商可以被他单方面拿走，用于科学课或其他项目。",
            "没有“不得绕开成交”和二次转化分成，却让你把企业客户带进场。",
            "老师取消不退费，或宣传权不完整，导致你无法正式做传播。",
        ],
    )
    add_h2(doc, "10.3 对方可能出现的几种反应，以及你怎么接")
    add_table(
        doc,
        ["对方说法", "真实含义", "你怎么接"],
        [
            ["老师是我的资源，所以我拿大头", "想用 IP 覆盖你的全部劳动", "IP 当然值钱，所以你自己的报名你拿大头；我自己做成的部分，按运营贡献分。"],
            ["你先付 15 万，后面都是你的", "把不确定性卖给你", "可以，请先给书面确认、独家权和取消全额退款，并一起把座位和票价算平。算不平就不要买断。"],
            ["那就五五分吧，简单", "听起来公平，其实让你承担成本后并不公平", "共同客户可以五五。我独立开发的报名和赞助不行。成本也必须先讲清楚谁承担。"],
            ["先把这场办了，合同以后再补", "事后一定会对你不利", "可以先出筹备清单，但售票、招商、场地签约前必须把分成和数据条款签掉。"],
            ["群和客户当然是我的，老师是我请来的", "这是最危险的信号", "老师资源归你，项目新客户共有。否则我等于在给你的科学课做地推。"],
        ],
        [4.4, 4.4, 7.6],
    )

    add_h1(doc, "十一、会前准备清单")
    add_table(
        doc,
        ["类别", "你要准备好的东西"],
        [
            ["身份", "对外方案 Word；口头一句：我是联合运营方，不是渠道。"],
            ["数字", "记熟：书面 60%、口头 70%、底线 50%；赞助 70%；票档 299/399/499/699/999；买断测算 25 万 / 约 532 人。"],
            ["问题", "把第 8 问打印或记在手机里，对方一提买断就按表问。"],
            ["替代方案", "阶梯分成表、3–5 万保底抵扣分成，作为第二选择，不要先抛。"],
            ["合同底线", "数据、群、二次转化、赞助保护期、不得绕开，六条缺一就缓签。"],
            ["长期叙事", "第一期 → 后续人文 / 哲学 → 和他的科学课协同，用来抬高你的位置。"],
            ["不说出口的", "不要说他没经验、不会办活动、只会借老师收钱。用“各擅所长”代替。"],
        ],
        [3.4, 13.0],
    )

    add_h1(doc, "十二、内部决策摘要")
    add_para(doc, "如果站在你的位置，三个方案排序非常明确：", first_line=True)
    add_bullets(
        doc,
        [
            "第一选择：联合运营 + 按来源分成。自己的报名 60%–70%，对方自己的报名你拿 20%–30%；共同客户五五开；你引进的赞助约 70% 归你。",
            "第二选择：小额保底 + 高比例分成。最多承担 3–5 万元基础授权 / 老师相关成本，再进行销售分成，而不是 10–18 万直接买断。",
            "第三选择：真正买断。只有对方愿意给你完整独家商业权利，而且经过测算你有高度把握做到比如 30 万、50 万甚至更高收入，才考虑 10–18 万元授权费。",
        ],
    )
    add_quote_box(
        doc,
        "把谈判目标定成这一句",
        "不交大额买断费，取得联合运营身份，己方客户至少 60%，己方赞助至少 70%，共同项目成本透明，客户资源受到保护。这会比单纯去和他争“20% 还是 70%”更有利，因为你是在重新定义整个合作模式，而不是跟他讨更高的渠道佣金。",
    )
    add_alert_box(
        doc,
        "发出文件前再看一眼",
        "对外版可以发。内部版不要发、不要转发到有对方的群、不要把本页的锚点和红线截图给任何人。两份文件已经分开，就是为了避免谈着谈着把底牌交出去。",
    )
    add_para(
        doc,
        "（内部工作稿，不构成对外承诺。）",
        size=9.5,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=16,
    )

    path = OUT / "王德峰老师10月31日活动_合作谈判方案_内部版.docx"
    doc.save(path)
    return path


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    external = build_external()
    internal = build_internal()
    print(f"已生成对外版：{external}")
    print(f"已生成内部版：{internal}")


if __name__ == "__main__":
    main()
