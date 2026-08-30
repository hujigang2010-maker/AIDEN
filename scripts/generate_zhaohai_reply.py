#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成致赵海的微信回复稿、可转发确认函，以及内部发送说明。"""

import sys
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
import generate_wangdefeng_proposal as g  # noqa: E402
from generate_external_outputs import _mini_table, _p, set_cell_margins  # noqa: E402

OUT = Path(__file__).resolve().parents[1] / "output"

WECHAT_PATH = OUT / "致赵海_微信回复稿.txt"
LETTER_PATH = OUT / "致赵海_8月29日沟通后合作路径确认.docx"
INTERNAL_PATH = OUT / "致赵海回复_内部说明.docx"

# 第一条：对人，给台阶，把「完全按你意思来」收成「两条线并行」。
WECHAT_BUBBLE_1 = """赵海好。

昨天聊完，我回去又过了一遍。书院系统课已经开起来了，这个我理解，也不会拿低价场去冲你们主课。

不过 10 月 31 日这场，我这边还是按公开体验场来做，不改成只推 3000 / 13000 的码。399、499 是进门的人，系统课是愿意深学的人，两拨客不一样；停掉体验场，两边都会薄。

路径我写清楚了，你方便转给团队看。哪几点能定，回我一声就行。"""

# 第二条：可直接转给心仪 / 书院负责人，不提内部判断。
WECHAT_BUBBLE_2 = """【10 月 31 日合作路径，请转达】

一、公开体验场继续做
10 月 31 日仍做 2.5–3 小时公开场。标准票 499 不涨。组织、场地、渠道由我这边按联合主办对接。

二、系统课作续课，不互相替代
3000 / 13000 作为体验后的续课。有意向的，我按你们规则推整课，转化部分渠道费 20%。不把公开场取消，也不改成只做渠道。

三、群
官方交流群由我企业微信建、我做群主不转让；你们做管理员，内容、老师动态你们发。群名用「活动交流群」。续课群同样我做群主；群里先只提老师和时间，不对外报价。

四、此前确认函四点仍有效
票价 499；分成要么两边都不调、要么两边一起调（不能只把我这边降到 55%、你们客户仍 70/30）；群主在我；视觉保留联合主办。

我这边最近节奏会稳一些，系统课人数不替你们冲刺。先把路径定住，再谈场地和开售。"""

# 对方如果只回一句「那你就推我们的课」，用这条钉住。
WECHAT_FALLBACK = """公开场和系统课是两件事，可以一起做，不能互相替代。系统课转化 20% 我可以配合；10 月 31 日这场仍按 499 公开体验场走，群主在我。你帮我跟团队说一声。"""


def write_wechat_txt() -> Path:
    text = "\n".join(
        [
            "致赵海｜微信回复稿（直接复制，分两条发）",
            "日期：2026年8月30日",
            "说明：先发第一条，等他回「好」或「你发来我转」，再发第二条。第二条可直接转给对方团队。",
            "",
            "======== 第一条（对人） ========",
            "",
            WECHAT_BUBBLE_1,
            "",
            "======== 第二条（可转发） ========",
            "",
            WECHAT_BUBBLE_2,
            "",
            "======== 若对方仍要你只推系统课 ========",
            "",
            WECHAT_FALLBACK,
            "",
        ]
    )
    WECHAT_PATH.write_text(text, encoding="utf-8")
    return WECHAT_PATH


def build_letter() -> Path:
    doc = g.new_doc()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.4)
    section.footer_distance = Cm(0.35)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("致赵海  ·  可转达团队  ·  2026年8月30日")
    g.set_run_font(run, size=8, color=g.MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("本函为沟通确认文件，不构成合同。最终以双方签署的协议为准。请就下列路径一并书面回复。")
    g.set_run_font(fr, size=7.5, color=g.MUTED)

    g.add_para(
        doc,
        "10 月 31 日王德峰老师活动｜沟通后合作路径确认",
        size=15,
        bold=True,
        color=g.NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0,
        space_after=3,
        line_spacing=1.0,
    )
    g.add_para(
        doc,
        "公开体验场与系统课分层并行，不互相替代。公开场标准票 499 元；官方群由我方任群主。",
        size=9,
        bold=True,
        color=g.GOLD,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
        line_spacing=1.05,
    )
    g.add_para(
        doc,
        "赵海好。感谢昨日沟通。书院系统课已经开班，价格保护的考虑我完全理解。回去对齐后确认：公开体验场继续做，系统课作为续课转化。两者并行，不互相替代。",
        size=9,
        space_before=0,
        space_after=4,
        line_spacing=1.1,
    )

    g.add_para(doc, "一、公开体验场：10 月 31 日继续做", size=11, bold=True, color=g.NAVY, space_before=2, space_after=3, line_spacing=1.0)
    g.add_para(
        doc,
        "2.5–3 小时公开讲座仍按原计划举办。399、499 是进门的人，3000、13000 是愿意深学的人。停掉体验场，两边都会薄。组织、场地、渠道由我方按联合主办对接，不是单纯票务渠道。",
        size=8.5,
        space_before=0,
        space_after=3,
        line_spacing=1.08,
    )
    _mini_table(
        doc,
        ["事项", "我方确认"],
        [
            ["日期与形态", "2026 年 10 月 31 日，2.5–3 小时公开体验场"],
            ["标准票", "499 元，不改为 599 元主力，也不取消公开场"],
            ["买断", "本次不按 10–18 万买断当天授权"],
        ],
        [3.6, 14.9],
    )

    g.add_para(doc, "二、系统课：作为续课转化，渠道费 20%", size=11, bold=True, color=g.NAVY, space_before=7, space_after=3, line_spacing=1.0)
    g.add_para(
        doc,
        "单课 3000 元、系列课 13000 元可作为体验后的续课。我方按贵方规则推介整课，转化部分渠道费 20%。该项是附带合作，不能替代公开体验场，也不能把 10 月 31 日让给系统课二期单独使用。系统课人数不替贵方冲刺凑班。",
        size=8.5,
        space_before=0,
        space_after=3,
        line_spacing=1.08,
    )
    box = doc.add_table(rows=1, cols=1)
    cell = box.cell(0, 0)
    g.shade_cell(cell, g.ROSE_HEX)
    g.set_cell_border(
        cell,
        top={"val": "single", "sz": "6", "color": g.RED_HEX, "space": "0"},
        left={"val": "single", "sz": "16", "color": g.RED_HEX, "space": "0"},
        bottom={"val": "single", "sz": "6", "color": g.RED_HEX, "space": "0"},
        right={"val": "single", "sz": "6", "color": g.RED_HEX, "space": "0"},
    )
    cell.text = ""
    set_cell_margins(cell, top=40, bottom=36, left=80, right=80)
    _p(
        cell,
        "我方不能接受：取消公开体验场，改成只推系统课二维码、只拿 20% 渠道费。系统课 20% 可以配合，但不能替代联合主办。",
        size=8.5,
        after=0,
    )
    g.set_table_full_width(box, [18.5])

    g.add_para(doc, "三、官方社群：我方任群主，贵方任管理员", size=11, bold=True, color=g.NAVY, space_before=7, space_after=3, line_spacing=1.0)
    _mini_table(
        doc,
        ["事项", "我方确认"],
        [
            ["建群 / 群名", "我方企业微信创建；「活动交流群」，不叫课程群"],
            ["权限", "我方任群主且不转让；贵方任管理员，内容与老师动态由贵方发"],
            ["续课群", "同样由我方任群主；群内先只提老师与时间，不对外报价"],
            ["边界", "不得擅自解散、转群主、导出名单另建群，或在群内单独销售其他课程"],
        ],
        [3.6, 14.9],
    )

    g.add_para(doc, "四、此前确认函四点仍有效", size=11, bold=True, color=g.NAVY, space_before=7, space_after=3, line_spacing=1.0)
    g.add_para(
        doc,
        "8 月 19 日确认函中的票价、分成、群主、视觉，不因系统课开班而单边作废。分成若下调，须对等让利。",
        size=8.5,
        space_before=0,
        space_after=3,
        line_spacing=1.08,
    )

    next_box = doc.add_table(rows=1, cols=4)
    items = [
        ("① 票价", "标准票 499 元；599 仅作优选票。"),
        ("② 分成", "选方案 A，或选方案 B（两边同步调）。"),
        ("③ 社群", "我方企微建群并任群主，贵方任管理员。"),
        ("④ 视觉", "双方书面确认后对外，保留联合主办。"),
    ]
    fills = [g.SOFT_HEX, g.CREAM_HEX, g.GREEN_HEX, g.AMBER_HEX]
    for i, ((title, body), fill) in enumerate(zip(items, fills)):
        c = next_box.cell(0, i)
        c.text = ""
        g.shade_cell(c, fill)
        set_cell_margins(c, top=46, bottom=46, left=60, right=60)
        _p(c, title, size=8.5, bold=True, color=g.NAVY, after=2)
        _p(c, body, size=7.5, after=0)
    g.set_table_full_width(next_box, [4.625, 4.625, 4.625, 4.625])

    g.add_para(
        doc,
        "请一并确认后回复。我方（联合运营方）　　胡继刚　　2026 年 8 月 30 日",
        size=9,
        bold=True,
        color=g.NAVY,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        space_before=8,
        space_after=0,
        line_spacing=1.0,
    )

    doc.save(LETTER_PATH)
    return LETTER_PATH


def build_internal() -> Path:
    doc = g.new_doc()
    g.setup_section(doc, "致赵海回复 · 内部说明（不要发给对方）", confidential=True)

    g.add_para(
        doc,
        "给赵海怎么回：内部说明",
        size=18,
        bold=True,
        color=g.NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    g.add_para(
        doc,
        "只给胡总看。不要把本文、心仪评价、养病细节发给赵海或对方团队。",
        size=10,
        bold=True,
        color=g.RED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )

    g.add_h1(doc, "一、为什么这样回")
    g.add_para(
        doc,
        "8 月 29 日当面，对方希望你停掉 399 / 499，改推 3000 / 13000，渠道费 20%。你当时偏软，说过「完全按你意思来」。这条微信的作用，是把口头让步收成「两条线并行」，同时给赵海面子，方便他转给心仪和书院负责人。",
    )
    g.add_para(
        doc,
        "赵海是对接人、相对听话。对人要暖，对事要硬。硬的部分放在第二条，让他去转，不要你去跟心仪硬碰。",
    )

    g.add_h1(doc, "二、怎么发")
    rows = [
        ["先发第一条", "对人、给台阶。等他回「好」或「你发来我转」。"],
        ["再发第二条", "可转发正文。需要书面时，再补发确认函 Word。"],
        ["不要一条发完", "长文容易被截成「胡总不同意」，赵海不好转。"],
        ["不要解释反悔", "说「回去对齐后把可落地方案写清楚」，不说「昨天说错了」。"],
    ]
    g.add_table(doc, ["步骤", "做法"], rows, [4.0, 12.5])

    g.add_h1(doc, "三、给了什么，没给什么")
    g.add_table(
        doc,
        ["给了（让他好转）", "没给（底线）"],
        [
            ["理解系统课价格保护，不拿低价场去冲主课", "取消 10 月 31 日公开体验场"],
            ["系统课转化 20%，按他们规则推整课", "把联合主办改成纯渠道"],
            ["他们做管理员，内容、广告他们发", "群主让出，或把名单交给对方另建群"],
            ["续课群只提老师和时间，不对外报价", "10 月 31 日让给系统课二期单独使用"],
            ["节奏稳一些，不替他们冲刺凑 30 人", "口头承诺帮系统课凑满班"],
        ],
        [8.25, 8.25],
    )

    g.add_h1(doc, "四、对方顶回来怎么说")
    g.add_table(
        doc,
        ["对方说", "你回"],
        [
            ["你昨天不是答应了吗？", "当面听清楚了你们的难处。回去对齐后，能落地的是两条线并行，不是取消公开场。"],
            ["低价场会冲我们 3000 的课。", "客群不同。体验场是进门，系统课是深学。停掉进门，转化更少。"],
            ["那就只帮我们推课，拿 20%。", "系统课 20% 可以配合；10 月 31 日这场仍按 499 公开场走，群主在我。"],
            ["群主给我们做吧。", "内容你们发没问题。建群和群主在我，这是联合主办的基本条件。"],
            ["你先帮我们凑 10 个人。", "我这边在调养，系统课人数不替你们冲刺。公开场组织我按联合主办对接。"],
            ["分成你降到 55%，我们客户还是 70。", "要么两边都不调，要么两边一起调。单边下调不能接受。"],
        ],
        [6.0, 10.5],
    )

    g.add_h1(doc, "五、微信里不要出现的词")
    g.add_para(
        doc,
        "心仪、没经验、指挥、掏空、买名单、面神经、五级、养病细节、10–18 万太贵、他们不会办活动。这些只留在内部。对外只说：联合主办、两条线、499、群主、20% 是附带。",
    )

    g.add_para(
        doc,
        "内部文件，勿外传。",
        size=9,
        color=g.RED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=16,
    )

    doc.save(INTERNAL_PATH)
    return INTERNAL_PATH


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    wechat = write_wechat_txt()
    letter = build_letter()
    internal = build_internal()
    print(f"已生成微信稿：{wechat}")
    print(f"已生成确认函：{letter}")
    print(f"已生成内部说明：{internal}")


if __name__ == "__main__":
    main()
