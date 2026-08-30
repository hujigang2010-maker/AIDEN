#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""按胡总口径重写：主推 10 月 31 日活动，不主推二期整课；续课群我做，观点不让。"""

import sys
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
import generate_wangdefeng_proposal as g  # noqa: E402
from generate_external_outputs import _mini_table, _p, set_cell_margins  # noqa: E402

OUT = Path(__file__).resolve().parents[1] / "output"

WECHAT_PATH = OUT / "致赵海_按活动为主回复稿.txt"
LETTER_PATH = OUT / "致赵海_活动为主合作路径确认.docx"
INTERNAL_PATH = OUT / "致赵海_活动为主_内部说明.docx"

WECHAT_BUBBLE_1 = """赵海好。

三点看过了。中欧那边恭喜。分流的担心我理解，所以我有一个建议：续课群我来建、我做群主。群里先只提老师和时间，不对外报价，也不铺主课。这样主课价格保护住，活动这边也能把人留下来。

不过有一点要说清楚：我这边主推的是 10 月 31 日这场公开活动，不主推二期整课。先推主课、一个月后再说单课，这个我有保留。活动场还是要做，日期还是 10 月 31 日。

路径写在下面，你转给团队。"""

WECHAT_BUBBLE_2 = """【回复三点，请转达】

一、我主推 10 月 31 日公开活动，不主推二期整课
公开活动保留，而且要办。2.5–3 小时公开讲座，标准票 499，日期 10 月 31 日。
近期可以先不铺 299 低价引流，避免碰到主课锚点；但筹备、建群、预告推的是这场活动，不是帮二期招生。
9 月 30 日可以作正式开售节点，不是把活动改期，更不是等主课招满再谈。
王老师可以是二期亮点；10 月 31 日对外场次仍是公开活动。

二、主课渠道费：不作为我这边的工作
阶梯比例收到了。我这边不按这个任务去招主课，也不承诺人数。
若有人自己问到系统课，再单独说。续课群不是主课招生群。

三、续课群我来做；观点仍是活动为主
同意：我企业微信建群、我做群主不转让；你做管理员，内容、老师动态你发。群名用「活动交流群」。续课群同样我做群主。
群里先只提老师和时间，不对外报价。这是保护主课，也是把活动的人留住，不是把群改成二期招生。

公开活动仍按联合主办对接场地、组织和渠道。此前确认函四点仍然有效。"""

WECHAT_FALLBACK = """续课群我做、群主在我，群里不卖主课。我主推 10 月 31 日公开活动，不主推二期整课。渠道费不作为我的招生任务。活动日期不改。"""


def write_wechat_txt() -> Path:
    text = "\n".join(
        [
            "致赵海｜按「活动为主」口径的回复稿（直接复制，分两条发）",
            "说明：先发第一条。这版不再答应去推他们的大课。",
            "",
            "======== 第一条（对人） ========",
            "",
            WECHAT_BUBBLE_1,
            "",
            "======== 第二条（可转发） ========",
            "",
            WECHAT_BUBBLE_2,
            "",
            "======== 若对方仍要你先推主课 ========",
            "",
            WECHAT_FALLBACK,
            "",
        ]
    )
    WECHAT_PATH.write_text(text, encoding="utf-8")
    return WECHAT_PATH


def _rose_box(doc, text):
    box = doc.add_table(rows=1, cols=1)
    cell = box.cell(0, 0)
    g.shade_cell(cell, g.ROSE_HEX)
    g.set_cell_border(
        cell,
        top={"val": "single", "sz": "6", "color": g.RED_HEX, "space": "0"},
        left={"val": "single", "sz": "16", "color": g.RED_HEX, "space": "0"},
        bottom={"val": "single", "sz": "6", "color": g.RED_HEX, "space": "0"},
        right={"val": "single", "sz": "6", "color": g.RED_HEX, "space": "0"},
    )
    cell.text = ""
    set_cell_margins(cell, top=40, bottom=36, left=80, right=80)
    _p(cell, text, size=8.5, after=0)
    g.set_table_full_width(box, [18.5])
    return box


def build_letter() -> Path:
    doc = g.new_doc()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.4)
    section.footer_distance = Cm(0.35)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("致赵海  ·  可转达团队  ·  2026年8月30日")
    g.set_run_font(run, size=8, color=g.MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("本函为沟通确认文件，不构成合同。最终以双方签署的协议为准。请就下列路径一并书面回复。")
    g.set_run_font(fr, size=7.5, color=g.MUTED)

    g.add_para(
        doc,
        "10 月 31 日王德峰老师活动｜合作路径确认（活动为主）",
        size=14.5,
        bold=True,
        color=g.NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0,
        space_after=3,
        line_spacing=1.0,
    )
    g.add_para(
        doc,
        "我方主推 10 月 31 日公开活动，不主推二期整课。续课群由我方建群并任群主，群内不销售主课。",
        size=9,
        bold=True,
        color=g.GOLD,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
        line_spacing=1.05,
    )
    g.add_para(
        doc,
        "赵海好。贵方三点收悉。中欧学员缴费，祝贺。分流担心，我方建议用续课群来接：我方建群并任群主，群内先只提老师与时间，不对外报价、不铺主课。公开活动仍要办，我方不主推二期整课。",
        size=9,
        space_before=0,
        space_after=4,
        line_spacing=1.1,
    )

    g.add_para(doc, "一、主推公开活动，不主推二期整课", size=11, bold=True, color=g.NAVY, space_before=2, space_after=3, line_spacing=1.0)
    _mini_table(
        doc,
        ["事项", "我方确认"],
        [
            ["主推什么", "10 月 31 日 2.5–3 小时公开活动，标准票 499 元"],
            ["不主推什么", "二期整课（3000 / 13000）；不按「先推主课」执行"],
            ["日期", "2026 年 10 月 31 日，不改期，不等主课招满再谈"],
            ["低价票", "近期可不铺 299 引流票，避免碰到主课锚点"],
            ["开售", "9 月 30 日可作正式开售节点，不是改期"],
            ["与二期同日", "对外场次是公开活动；系统课学员可内部入场"],
        ],
        [3.6, 14.9],
    )

    g.add_para(doc, "二、主课渠道费：不作为我方工作", size=11, bold=True, color=g.NAVY, space_before=7, space_after=3, line_spacing=1.0)
    g.add_para(
        doc,
        "阶梯比例已收到。我方不按此任务招主课，不承诺人数。若有人自己问到系统课，再单独沟通。续课群不是主课招生群。",
        size=8.5,
        space_before=0,
        space_after=3,
        line_spacing=1.08,
    )
    _rose_box(
        doc,
        "我方不能接受：先推主课、一个月后再说公开活动；把我方改成主课招生渠道；把续课群当成二期招生群。续课群我方愿意做，观点仍是活动为主。",
    )

    g.add_para(doc, "三、续课群我方来做，群主不转让", size=11, bold=True, color=g.NAVY, space_before=7, space_after=3, line_spacing=1.0)
    _mini_table(
        doc,
        ["事项", "双方可对齐"],
        [
            ["建群 / 群主", "我方企业微信创建，我方任群主且不转让"],
            ["管理员", "贵方任管理员，内容与老师动态由贵方发"],
            ["群名", "「活动交流群」；续课群同样由我方任群主"],
            ["群内规则", "先只提老师与时间，不对外报价，不铺主课二维码"],
            ["作用", "保护主课价格，并把公开活动的人留下来，不改成二期招生群"],
        ],
        [3.6, 14.9],
    )

    g.add_para(doc, "四、此前确认函仍有效", size=11, bold=True, color=g.NAVY, space_before=7, space_after=3, line_spacing=1.0)
    next_box = doc.add_table(rows=1, cols=4)
    items = [
        ("① 活动", "10 月 31 日公开活动，标准票 499。"),
        ("② 主课", "不主推整课，渠道费不作任务。"),
        ("③ 续课群", "我方建群任群主，群内不卖主课。"),
        ("④ 角色", "联合主办，不是主课渠道。"),
    ]
    fills = [g.SOFT_HEX, g.CREAM_HEX, g.GREEN_HEX, g.AMBER_HEX]
    for i, ((title, body), fill) in enumerate(zip(items, fills)):
        c = next_box.cell(0, i)
        c.text = ""
        g.shade_cell(c, fill)
        set_cell_margins(c, top=46, bottom=46, left=60, right=60)
        _p(c, title, size=8.5, bold=True, color=g.NAVY, after=2)
        _p(c, body, size=7.5, after=0)
    g.set_table_full_width(next_box, [4.625, 4.625, 4.625, 4.625])

    g.add_para(
        doc,
        "请一并确认后回复。我方（联合运营方）　　胡继刚　　2026 年 8 月 30 日",
        size=9,
        bold=True,
        color=g.NAVY,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        space_before=8,
        space_after=0,
        line_spacing=1.0,
    )

    doc.save(LETTER_PATH)
    return LETTER_PATH


def build_internal() -> Path:
    doc = g.new_doc()
    g.setup_section(doc, "活动为主口径 · 内部说明（不要发给对方）", confidential=True)

    g.add_para(
        doc,
        "对你这版想法的反馈",
        size=18,
        bold=True,
        color=g.NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    g.add_para(
        doc,
        "只给胡总看。上一版把渠道费写太实，这版已改回：主推活动，续课群我做，观点不让。",
        size=10,
        bold=True,
        color=g.RED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )

    g.add_h1(doc, "一、你的判断是对的")
    g.add_para(
        doc,
        "不想推他们的大课、想推 10 月 31 日这场活动，这个分歧要写进回复里，不能含糊。他们三点的目标，就是让你 9 月去填二期。你一旦答应「先推主课」或把 20% / 25% / 30% 当成自己的工作，公开活动就会被往后拖，拖到他们不需要你。",
    )
    g.add_para(
        doc,
        "上一版为了给台阶，写了「渠道费同意作附带」「499 可抵系统课」。对外像是你愿意帮他们卖 13000。按你现在的口径，这两句都拿掉。",
    )

    g.add_h1(doc, "二、续课群为什么能「保留观点」")
    g.add_para(
        doc,
        "续课群是你让出去的那一步，也是你真正要的那一步。让出去的是：群里不报价、不铺主课，照顾他们怕分流。要回来的是：群主在你，名单在你，人是为 10 月 31 日活动留下来的。后续有人自己问大课，再个案谈，你没有义务帮他们冲 30 人。",
    )
    g.add_table(
        doc,
        ["续课群要写成", "续课群不要写成"],
        [
            ["我来建、我做群主，不转让", "你们发内容就等于这是你们的招生群"],
            ["只提老师和时间，不对外报价", "群里铺 3000 / 13000 二维码"],
            ["把公开活动的人留下来", "先帮二期凑人，活动以后再说"],
            ["有人自己问主课，再单独说", "按阶梯任务去招主课"],
        ],
        [8.25, 8.25],
    )

    g.add_h1(doc, "三、这版让了什么，没让什么")
    g.add_table(
        doc,
        ["让了", "没让"],
        [
            ["续课群我做，群里不卖主课", "主推二期整课"],
            ["近期可不铺 299 引流票", "9 月改去招主课"],
            ["9 月 30 日可作活动开售节点", "活动改期或等主课招满再谈"],
            ["渠道费收到了，不接成任务", "承诺人数、按 20–30% 去招"],
            ["恭喜中欧一句", "被带成「那你也赶紧推主课」"],
        ],
        [8.25, 8.25],
    )

    g.add_h1(doc, "四、怎么发")
    g.add_table(
        doc,
        ["步骤", "做法"],
        [
            ["先发第一条", "建议是续课群；保留的是不推大课、活动日期不改。"],
            ["再发第二条", "可转发。三条对着他们的三点回，主推什么写在第一条标题里。"],
            ["不要用上一版", "不要发「渠道费同意作附带」「499 可抵系统课」那一稿。"],
        ],
        [4.0, 12.5],
    )

    g.add_h1(doc, "五、对方再顶怎么说")
    g.add_table(
        doc,
        ["对方说", "你回"],
        [
            ["那你先帮我们推主课。", "大课我不主推。我主推 10 月 31 日活动。续课群我做，群里不卖主课。"],
            ["你不做渠道，阶梯就作废了。", "可以。我本来也不是按这个任务来的。有人自己问，再单独说。"],
            ["续课群里我们要发课程。", "老师动态可以发。价格和报名码不在群里铺。要报的私聊。"],
            ["活动现在推会分流。", "所以群里不报价，也不铺 299。推的是活动预告，不是拿低价打你们主课。"],
        ],
        [6.2, 10.3],
    )

    g.add_para(
        doc,
        "内部文件，勿外传。请用本口径，不要用上一版「渠道费作附带」的回复。",
        size=9,
        color=g.RED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=16,
    )

    doc.save(INTERNAL_PATH)
    return INTERNAL_PATH


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    wechat = write_wechat_txt()
    letter = build_letter()
    internal = build_internal()
    print(f"已生成微信稿：{wechat}")
    print(f"已生成确认函：{letter}")
    print(f"已生成内部说明：{internal}")


if __name__ == "__main__":
    main()
