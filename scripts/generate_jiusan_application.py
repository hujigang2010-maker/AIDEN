#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""根据《九三学社申请人社登记表》（2020年版）结构，结合简历生成胡继刚填表稿。"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "九三学社申请人社登记表-胡继刚.docx"


def set_run_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def clear_cell(cell):
    cell.text = ""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)


def write_cell(cell, text, *, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, shade=None, center_v=True):
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.15
    run = p.add_run(text if text is not None else "")
    set_run_font(run, size=size, bold=bold)
    set_cell_borders(cell)
    if shade:
        set_cell_shading(cell, shade)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center" if center_v else "top")
    tcPr.append(vAlign)


def label_cell(cell, text, size=9):
    write_cell(cell, text, bold=True, size=size, align=WD_ALIGN_PARAGRAPH.CENTER, shade="E8EEF7")


def value_cell(cell, text, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    write_cell(cell, text, bold=False, size=size, align=align)


def merge(table, r1, c1, r2, c2):
    a = table.cell(r1, c1)
    b = table.cell(r2, c2)
    a.merge(b)
    return a


def add_title(doc, text, size=18, space_after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=RGBColor(0x1A, 0x3A, 0x5C))
    return p


def add_note(doc, text, size=9, color=RGBColor(0x66, 0x66, 0x66)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)
    return p


def add_heading_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color=RGBColor(0x1A, 0x3A, 0x5C))
    return p


def add_para(doc, text, *, size=10.5, bold=False, first_line=False, space_after=4):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)


def build_cover(doc):
    add_title(doc, "九三学社", size=22, space_after=4)
    add_title(doc, "申请人社登记表", size=20, space_after=18)

    cover = doc.add_table(rows=3, cols=2)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_col_widths(cover, [3.2, 10.5])
    label_cell(cover.cell(0, 0), "申请人", size=11)
    value_cell(cover.cell(0, 1), "胡继刚", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    label_cell(cover.cell(1, 0), "填表日期", size=11)
    value_cell(cover.cell(1, 1), "2026年7月30日", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    label_cell(cover.cell(2, 0), "所属组织（盖章）", size=10)
    value_cell(cover.cell(2, 1), "（由接收组织填写盖章）", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in cover.rows:
        row.height = Cm(1.1)

    add_note(doc, "九三学社中央委员会组织部印制（2020年版）填表稿")
    add_note(
        doc,
        "说明：本文件依据《九三学社申请人社登记表》栏目，结合申请人简历整理。"
        "人名统一为“胡继刚”（不以简历中的“汤先生”称呼填写）。"
        "基层组织意见、地方组织审批意见、省级备案时间由社组织填写，本申请稿留空。"
        "照片、介绍人第二人及个别未知家庭成员信息请申请人自行补全。",
    )
    doc.add_page_break()


def build_basic(doc):
    add_heading_line(doc, "一、基本信息")

    t = doc.add_table(rows=10, cols=8)
    t.style = "Table Grid"
    set_col_widths(t, [2.0, 2.2, 1.6, 2.2, 1.8, 2.2, 1.6, 2.4])

    # row0 姓名/性别/出生日期/照片
    label_cell(t.cell(0, 0), "姓名")
    value_cell(t.cell(0, 1), "胡继刚", align=WD_ALIGN_PARAGRAPH.CENTER)
    label_cell(t.cell(0, 2), "性别")
    value_cell(t.cell(0, 3), "男", align=WD_ALIGN_PARAGRAPH.CENTER)
    label_cell(t.cell(0, 4), "出生日期")
    value_cell(merge(t, 0, 5, 0, 6), "1987年4月25日", align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(merge(t, 0, 7, 3, 7), "照\n片\n\n（自行粘贴）", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, shade="F5F5F5")

    # row1 民族/籍贯/出生地
    label_cell(t.cell(1, 0), "民族")
    value_cell(t.cell(1, 1), "汉族", align=WD_ALIGN_PARAGRAPH.CENTER)
    label_cell(t.cell(1, 2), "籍贯")
    value_cell(t.cell(1, 3), "山东青岛", align=WD_ALIGN_PARAGRAPH.CENTER)
    label_cell(t.cell(1, 4), "出生地")
    value_cell(merge(t, 1, 5, 1, 6), "吉林磐石", align=WD_ALIGN_PARAGRAPH.CENTER)

    # row2 身份证/健康
    label_cell(t.cell(2, 0), "身份证号")
    value_cell(merge(t, 2, 1, 2, 3), "220283198704250614", align=WD_ALIGN_PARAGRAPH.CENTER)
    label_cell(t.cell(2, 4), "健康状况")
    value_cell(merge(t, 2, 5, 2, 6), "良好", align=WD_ALIGN_PARAGRAPH.CENTER)

    # row3 入社时间/基层组织
    label_cell(t.cell(3, 0), "入社时间")
    value_cell(merge(t, 3, 1, 3, 3), "（审批通过后由组织填写）", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    label_cell(t.cell(3, 4), "所属基层组织")
    value_cell(merge(t, 3, 5, 3, 6), "（由组织填写）", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

    # row4 参加工作/行政级别/任现职级/是否退休
    label_cell(t.cell(4, 0), "参加工作时间")
    value_cell(t.cell(4, 1), "2011年6月", align=WD_ALIGN_PARAGRAPH.CENTER)
    label_cell(t.cell(4, 2), "行政级别")
    value_cell(t.cell(4, 3), "无", align=WD_ALIGN_PARAGRAPH.CENTER)
    label_cell(t.cell(4, 4), "任现职级时间")
    value_cell(t.cell(4, 5), "2021年6月", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    label_cell(t.cell(4, 6), "是否退休")
    value_cell(t.cell(4, 7), "否", align=WD_ALIGN_PARAGRAPH.CENTER)

    # 修正：row4 是否退休单独占一格更清晰——已写入
    # row5 职称
    label_cell(t.cell(5, 0), "专业技术职称")
    value_cell(merge(t, 5, 1, 5, 2), "副高级 / 副教授级高级工程师", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    label_cell(t.cell(5, 3), "专业技术职称2")
    value_cell(merge(t, 5, 4, 5, 5), "中级 / 会计师", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    label_cell(t.cell(5, 6), "学术称号")
    value_cell(t.cell(5, 7), "无", align=WD_ALIGN_PARAGRAPH.CENTER)

    # row6 全日制教育
    label_cell(merge(t, 6, 0, 6, 1), "全日制教育\n学历/学位")
    value_cell(t.cell(6, 2), "本科 / 学士", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    label_cell(t.cell(6, 3), "毕业院校系及专业")
    value_cell(merge(t, 6, 4, 6, 7), "中国海洋大学 · 土木工程", size=9)

    # row7 在职教育
    label_cell(merge(t, 7, 0, 7, 1), "在职教育\n学历/学位")
    value_cell(t.cell(7, 2), "硕士研究生 / 硕士", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    label_cell(t.cell(7, 3), "毕业院校系及专业")
    value_cell(merge(t, 7, 4, 7, 7), "复旦大学 · 工商管理（财务金融方向）", size=9)

    # row8 工作单位/部门/职务
    label_cell(t.cell(8, 0), "工作单位")
    value_cell(merge(t, 8, 1, 8, 3), "复旦大学住房政策研究中心", size=9)
    label_cell(t.cell(8, 4), "工作部门")
    value_cell(t.cell(8, 5), "研究中心", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    label_cell(t.cell(8, 6), "职务")
    value_cell(t.cell(8, 7), "秘书长", align=WD_ALIGN_PARAGRAPH.CENTER)

    # row9 是否中共/党派交叉/手机/邮箱
    label_cell(t.cell(9, 0), "是否中共")
    value_cell(t.cell(9, 1), "否", align=WD_ALIGN_PARAGRAPH.CENTER)
    label_cell(t.cell(9, 2), "党派交叉")
    value_cell(t.cell(9, 3), "无", align=WD_ALIGN_PARAGRAPH.CENTER)
    label_cell(t.cell(9, 4), "手机号")
    value_cell(t.cell(9, 5), "18678408669", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    label_cell(t.cell(9, 6), "电子邮箱")
    value_cell(t.cell(9, 7), "262782809@qq.com", size=8)

    # 地址表
    addr = doc.add_table(rows=2, cols=6)
    addr.style = "Table Grid"
    set_col_widths(addr, [2.0, 6.5, 1.8, 2.5, 1.8, 1.8])
    label_cell(addr.cell(0, 0), "单位地址")
    value_cell(addr.cell(0, 1), "上海市杨浦区（复旦大学住房政策研究中心）", size=9)
    label_cell(addr.cell(0, 2), "单位电话")
    value_cell(addr.cell(0, 3), "18678408669", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    label_cell(addr.cell(0, 4), "单位邮编")
    value_cell(addr.cell(0, 5), "200433", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    label_cell(addr.cell(1, 0), "家庭地址")
    value_cell(addr.cell(1, 1), "上海市杨浦区爱国路389号", size=9)
    label_cell(addr.cell(1, 2), "家庭电话")
    value_cell(addr.cell(1, 3), "18678408669", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    label_cell(addr.cell(1, 4), "家庭邮编")
    value_cell(addr.cell(1, 5), "200090", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading_line(doc, "二、社会职务")
    social = doc.add_table(rows=5, cols=6)
    social.style = "Table Grid"
    set_col_widths(social, [3.2, 2.2, 2.0, 1.6, 1.6, 5.4])
    headers = ["社会组织类别/名称", "级别", "职务", "省份", "城市", "备注"]
    for i, h in enumerate(headers):
        label_cell(social.cell(0, i), h, size=8)
    # 社会职务以公开报道可交叉印证项优先；工商联商会/省商会等简历兼职仍可填，政协界别留空
    rows = [
        ("上海市杨浦区科技企业联合会", "区级", "执行会长", "上海", "上海", "政协界别：无"),
        ("复旦MBA不动产资产管理协会", "校级/协会", "秘书长", "上海", "上海", "政协界别：无"),
        ("上海市工商联房地产商会", "市级", "秘书长", "上海", "上海", "政协界别：无；简历兼任"),
        ("上海山东省商会", "市级", "理事", "上海", "上海", "政协界别：无；简历兼任"),
    ]
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            value_cell(social.cell(r, c), val, size=8, align=WD_ALIGN_PARAGRAPH.CENTER if c else WD_ALIGN_PARAGRAPH.LEFT)

    add_para(doc, "特邀（约）职务：", bold=True, size=10, space_after=2)
    add_para(
        doc,
        "上海市城市更新研究会城市更新评审专家；上海节能减排工程技术协会碳中和节能评审专家；"
        "上海市虹口区科技企业联合会科技企业专业评审专家；复旦大学研究生管理联考面试官。",
        size=10,
        space_after=4,
    )
    add_para(doc, "其他职务：", bold=True, size=10, space_after=2)
    add_para(
        doc,
        "上海市人才引进；项目管理工程师任职资格（上海市住房和城乡建设委员会评定）。"
        "（人才引进属资质认证，不写入获奖情况。）",
        size=10,
    )


def build_expertise_and_resume(doc):
    add_heading_line(doc, "三、业务专长")
    exp = doc.add_table(rows=2, cols=2)
    exp.style = "Table Grid"
    set_col_widths(exp, [3.5, 12.5])
    label_cell(exp.cell(0, 0), "业务分类")
    label_cell(exp.cell(0, 1), "专长详情")
    value_cell(exp.cell(1, 0), "应用经济学\n（兼土木工程）", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    value_cell(
        exp.cell(1, 1),
        "不动产投资拓展与特殊资产业务（募投管退）；住房政策与租赁市场研究；城市更新与产城融合；"
        "商业综合体/物流地产/冷链产业园落地；股权收并购、法院拍卖及不良资产处置；政企协同与资源整合。",
        size=9,
    )
    exp.rows[1].height = Cm(2.2)

    add_heading_line(doc, "四、简历（学习简历 + 工作简历）")
    resume_lines = [
        "2007.09-2011.07  中国海洋大学土木工程专业学习，获工学学士学位",
        "2011.06-2012.10  中南建设集团股份有限公司/中南控股集团总裁办公室战略企划中心 战略企划管理主管",
        "2012.10-2015.10  中南建设集团股份有限公司新项目发展中心 投资拓展经理",
        "2015.11-2017.02  新城控股集团股份有限公司战略投资中心 投资拓展资深专业经理",
        "2017.03-2021.06  万科企业集团股份有限公司上海区域江浙事业部 投资副总经理",
        "2018.09-2021.06  复旦大学工商管理（财务金融方向）硕士研究生在职学习，获硕士学位",
        "2021.06-至今      复旦大学住房政策研究中心 秘书长（副教授级高级工程师）",
    ]
    for line in resume_lines:
        add_para(doc, line, size=10.5, space_after=2)

    add_note(
        doc,
        "补充说明：工作期间主导靖江印象城、嘉兴海宁、镇江大港、常州金坛、临沂相关合作项目，"
        "以及上海龙湖滟澜山法拍收并购、万纬嘉兴平湖/上海南桥冷链产业园等拓展落地；"
        "兼任杨浦区科技企业联合会执行会长、上海市工商联房地产商会秘书长等社会职务。",
        size=9,
    )


def build_family_awards(doc):
    add_heading_line(doc, "五、家庭主要成员及重要社会关系情况")
    fam = doc.add_table(rows=4, cols=5)
    fam.style = "Table Grid"
    set_col_widths(fam, [2.2, 2.4, 2.8, 2.8, 5.8])
    for i, h in enumerate(["称谓", "姓名", "出生日期", "政治面貌", "工作单位及职务"]):
        label_cell(fam.cell(0, i), h, size=9)
    value_cell(fam.cell(1, 0), "配偶", align=WD_ALIGN_PARAGRAPH.CENTER)
    value_cell(fam.cell(1, 1), "马喜艳", align=WD_ALIGN_PARAGRAPH.CENTER)
    value_cell(fam.cell(1, 2), "", align=WD_ALIGN_PARAGRAPH.CENTER)
    value_cell(fam.cell(1, 3), "", align=WD_ALIGN_PARAGRAPH.CENTER)
    value_cell(fam.cell(1, 4), "联系电话：18653268620", size=9)
    for r in range(2, 4):
        for c in range(5):
            value_cell(fam.cell(r, c), "")

    add_heading_line(doc, "六、海外关系")
    add_para(doc, "无。", size=10.5)

    add_heading_line(doc, "七、获奖情况（省部级以上）")
    award = doc.add_table(rows=3, cols=6)
    award.style = "Table Grid"
    set_col_widths(award, [3.0, 3.2, 2.0, 1.8, 3.0, 2.8])
    for i, h in enumerate(["奖项名称", "获奖项目名称", "获奖时间", "获奖级别", "授予单位", "备注"]):
        label_cell(award.cell(0, i), h, size=8)
    # 无省部级以上奖励则留空；评职称企业内部奖、人才引进认证均不填本栏
    for r in range(1, 3):
        for c in range(6):
            value_cell(award.cell(r, c), "")
    add_note(
        doc,
        "按填写说明，本栏原则上填写省部级以上奖励。评职称所用单位内部表彰、上海市人才引进认证均不填入本栏。",
        size=9,
    )


def build_patents_and_achievements(doc):
    add_heading_line(doc, "八、专利情况")
    pat = doc.add_table(rows=3, cols=5)
    pat.style = "Table Grid"
    set_col_widths(pat, [2.4, 3.2, 6.0, 2.0, 2.4])
    for i, h in enumerate(["专利类别", "专利号", "专利名称", "时间", "备注"]):
        label_cell(pat.cell(0, i), h, size=8)
    patents = [
        ("实用新型", "2020205***", "一种建筑工程管理用功能脚架", "2020", "简历脱敏号，请补全完整专利号"),
        ("实用新型", "20202055***", "一种建筑工程用防护围栏", "2020", "简历脱敏号，请补全完整专利号"),
    ]
    for r, row in enumerate(patents, start=1):
        for c, v in enumerate(row):
            value_cell(pat.cell(r, c), v, size=8, align=WD_ALIGN_PARAGRAPH.CENTER if c != 2 else WD_ALIGN_PARAGRAPH.LEFT)

    add_heading_line(doc, "九、专业技术工作及成果")
    achievements = [
        "1. 研究论文：《万科物流地产平台业务发展战略研究》（知网收录）；"
        "《房屋建筑工种人工智能技术的应用》（《住宅与房地产》）；"
        "《铝模板在高层住宅建筑工程中的应用》（《中国新技术新产品》）；"
        "《基于BIM的装配式钢结构建筑施工新技术与管理研究》（题名以正式发表稿为准）。",
        "2. 实用新型专利2项：一种建筑工程管理用功能脚架；一种建筑工程用防护围栏。",
        "3. 主要从事住房政策研究与不动产投资拓展相关工作，工作经历涵盖商业综合体、产城融合、"
        "物流地产及冷链产业园等项目拓展，熟悉重资产勾地、轻资产品牌输出、委托代建、"
        "股权收并购、法院拍卖与不良资产处置等业务路径。",
        "4. 代表性项目（节选，据简历）：万科靖江印象城；嘉兴海宁商住项目；镇江大港万科金域蓝湾；"
        "常州金坛理想城；临沂河东鲁商新都会；上海龙湖滟澜山（法院拍卖收并购）；"
        "万纬嘉兴平湖/上海南桥冷链园；新城临沂吾悦广场、青岛吾悦广场；中南如皋世纪城、苏州中南锦苑等。",
        "5. 公开活动与职务相关成果（有媒体/官网报道可核验）："
        "2024年6月主持“他山之石——中国投资者的海外不动产战略布局”主题分享；"
        "2025年5月参加“2025全球新经济增长引擎峰会”圆桌交流；"
        "2025年5月出席上海市工商联房地产商会资产管理分会成立大会暨不动产资产管理高质量发展论坛，"
        "并参与相关战略合作仪式；"
        "2026年3月参加“北欧创新国际会客厅”揭牌交流；"
        "2026年5月主持“2026人工智能商业化落地与硬核投资破局峰会”圆桌对话"
        "“从算力引擎到新质资产（AI全产业链的商业化实战）”，并以主办方代表身份发言。",
    ]
    for text in achievements:
        add_para(doc, text, size=10.5, space_after=4)

    add_heading_line(doc, "十、入社介绍人")
    intro = doc.add_table(rows=3, cols=3)
    intro.style = "Table Grid"
    set_col_widths(intro, [3.5, 8.5, 4.0])
    for i, h in enumerate(["姓名", "工作单位及职务", "关系"]):
        label_cell(intro.cell(0, i), h, size=9)
    # 原扫描件手写信息予以保留
    value_cell(intro.cell(1, 0), "王航", align=WD_ALIGN_PARAGRAPH.CENTER)
    value_cell(intro.cell(1, 1), "上海元能智能科技有限公司 总经理", size=9)
    value_cell(intro.cell(1, 2), "校友", align=WD_ALIGN_PARAGRAPH.CENTER)
    value_cell(intro.cell(2, 0), "（第二介绍人待填写）", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    value_cell(intro.cell(2, 1), "", size=9)
    value_cell(intro.cell(2, 2), "", size=9)
    add_note(doc, "说明：第一介绍人信息按原登记表手写栏转录；按规定一般需两名社员作介绍人，第二人请补全。", size=9)

    add_heading_line(doc, "十一、组织意见（由社组织填写，本申请稿留空）")
    org = doc.add_table(rows=4, cols=2)
    org.style = "Table Grid"
    set_col_widths(org, [3.5, 12.5])
    for r, name in enumerate(["介绍人意见", "基层组织意见", "地方组织审批意见", "省级组织部门备案时间"]):
        label_cell(org.cell(r, 0), name, size=9)
        value_cell(org.cell(r, 1), "", size=9)
        org.rows[r].height = Cm(1.3)


def build_appendix_checklist(doc):
    doc.add_page_break()
    add_heading_line(doc, "附件：填表核对与待补全事项")
    add_para(doc, "已据简历填入、且人名统一为胡继刚的主要栏目：", bold=True, size=10.5)
    items_done = [
        "封面申请人、填表日期",
        "姓名、性别、出生日期、民族、籍贯、出生地、身份证号、健康状况",
        "参加工作时间、任现职级时间、是否退休、专业技术职称及职称2",
        "全日制/在职学历学位与毕业院校专业",
        "工作单位、职务、联系方式、家庭住址",
        "社会职务、特邀评审专家职务",
        "业务专长、学习与工作简历",
        "专利情况、专业技术工作及成果",
        "第一介绍人（王航）",
    ]
    for x in items_done:
        add_para(doc, f"· {x}", size=10, space_after=1)

    add_para(doc, "建议申请人自行补全：", bold=True, size=10.5, space_after=2)
    items_todo = [
        "一寸或规定规格照片",
        "完整专利号（简历中为脱敏号）",
        "第二入社介绍人姓名、单位职务及关系",
        "配偶出生日期、政治面貌、工作单位及职务",
        "其他家庭主要成员/重要社会关系",
        "单位详细门牌地址与家庭邮编（如有）",
        "省部级以上获奖证明（如有）",
        "所属基层组织由接收方盖章确认",
    ]
    for x in items_todo:
        add_para(doc, f"· {x}", size=10, space_after=1)

    add_para(doc, "联系方式备查：", bold=True, size=10.5, space_after=2)
    add_para(doc, "手机：18678408669；邮箱：262782809@qq.com；微信：hu262782809", size=10)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    build_cover(doc)
    build_basic(doc)
    build_expertise_and_resume(doc)
    build_family_awards(doc)
    build_patents_and_achievements(doc)
    build_appendix_checklist(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"saved: {OUT}")


if __name__ == "__main__":
    build()
