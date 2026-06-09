"""Generate the Greentown standard sponsorship contract (GTH-zzc-ppch-016-2022a)
with all blanks filled in for the 100,000 RMB Greentown · 潮鸣外滩 sponsorship of
the 2026 AI Commercialization Summit, where Greentown is 甲方 (sponsor) and the
summit organizing committee is 乙方 (organizer / event party).

All filled-in values are highlighted in RED bold text with YELLOW background to
clearly mark which fields were added (vs. blanks that may still need review)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX


DARK = RGBColor(0x00, 0x00, 0x00)
RED = RGBColor(0xC0, 0x00, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)


def _force_font(run, font="宋体"):
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rpr.append(rfonts)


def set_default_font(doc, name="宋体"):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def add_run(p, text, bold=False, color=DARK, size=10.5, font="宋体",
            highlight=False):
    """Add a run; if highlight=True, render red bold on yellow background."""
    run = p.add_run(text)
    if highlight:
        run.bold = True
        run.font.color.rgb = RED
        run.font.size = Pt(size)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        _force_font(run, font=font)
    else:
        run.bold = bold
        run.font.color.rgb = color
        run.font.size = Pt(size)
        _force_font(run, font=font)
    return run


def add_para(doc, segments, align=None, first_indent=False, line_spacing=1.4,
             space_after=2):
    """segments is a list of tuples (text, is_filled). is_filled=True will
    render as red bold on yellow highlight."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(21)
    for seg in segments:
        if isinstance(seg, str):
            add_run(p, seg)
        else:
            text, is_filled = seg[0], seg[1]
            bold = seg[2] if len(seg) > 2 else False
            add_run(p, text, bold=bold, highlight=is_filled)
    return p


def add_heading_line(doc, text, size=12, center=False, bold=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    _force_font(run)
    return p


def add_title(doc, text, size=20):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    _force_font(run, font="黑体")
    return p


def blank(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)


def set_page_header(section, header_text):
    """Set the page header with the standard template code."""
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(header_text)
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    _force_font(run)

    # Add page number on the right via a tab
    tab_run = p.add_run("\t第 ")
    tab_run.font.size = Pt(9)
    tab_run.font.color.rgb = GREY
    _force_font(tab_run)

    # PAGE field
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run = p.add_run()
    page_run.font.size = Pt(9)
    page_run.font.color.rgb = GREY
    _force_font(page_run)
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_sep)
    page_run._r.append(fld_end)

    suffix_run = p.add_run(" 页 / 共 11 页")
    suffix_run.font.size = Pt(9)
    suffix_run.font.color.rgb = GREY
    _force_font(suffix_run)


def build():
    doc = Document()
    set_default_font(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)
    set_page_header(section, "标准文本编码 GTH-zzc-ppch-016-2022a")

    # ============================================================
    # 第 1 页：封面页
    # ============================================================
    blank(doc, 3)
    add_para(doc,
             [("合同编号：", False),
              ("________________________________", False)],
             align="center", line_spacing=2.0)
    blank(doc, 2)
    add_title(doc, "活动赞助合同", size=24)
    blank(doc, 4)

    add_para(doc,
             [("甲方：", False),
              ("________________________________", False)],
             line_spacing=2.0)
    add_para(doc,
             [("乙方：", False),
              ("上海市杨浦区科技企业联合会", True)],
             line_spacing=2.0)
    add_para(doc,
             [("签约地点：", False),
              ("上海市", True)],
             line_spacing=2.0)

    doc.add_page_break()

    # ============================================================
    # 第 2 页：正文起始 — 抬头 + 双方信息 + 第一条
    # ============================================================
    add_para(doc,
             [("本合同由以下两方于", False),
              ("二〇二六", True),
              ("年", False),
              ("五", True),
              ("月", False),
              ("二十一", True),
              ("日在", False),
              ("上海", True),
              ("省", False),
              ("上海", True),
              ("市", False),
              ("虹口", True),
              ("区签订：", False)])
    blank(doc)

    add_heading_line(doc, "甲方：________________________________", size=11)
    add_para(doc,
             [("住所：", False),
              ("________________________________", False)])
    add_para(doc,
             [("法定代表人：", False),
              ("________________________________", False)])
    blank(doc)

    add_heading_line(doc, "乙方：上海市杨浦区科技企业联合会", size=11)
    add_para(doc,
             [("住所：", False),
              ("上海市杨浦区国泰路 11 号 复旦科技园副楼一楼 1008 室", True)])
    add_para(doc,
             [("法定代表人：", False),
              ("朱震", True)])
    blank(doc)

    add_para(doc,
             [("根据《中华人民共和国民法典》等法律、法规的规定，本着公正、公平、合理、自愿的原则，就甲方为乙方组织的", False),
              ("「2026 人工智能商业化落地与硬核投资破局峰会」（含晚宴及现场展示等专属冠名权益）", True),
              ("活动提供赞助等事宜，经过双方友好协商一致，签订本合同，以资信守。", False)])
    blank(doc)

    add_heading_line(doc, "第一条  活动概况", size=12)
    add_para(doc,
             [("1、活动名称：", False),
              ("2026 人工智能商业化落地与硬核投资破局峰会", True),
              ("（下称\u201C本次活动\u201D）", False)])
    add_para(doc,
             [("2、活动地点：", False),
              ("上海·北外滩·一滴水（具体场地位置由乙方另行书面确认）", True)])
    add_para(doc,
             [("3、活动时间：", False),
              ("2026 年 5 月 22 日 13:30 – 20:30（含主论坛 + VIP 闭门晚宴；以正式邀请函为准）", True)])
    add_para(doc,
             [("4、活动具体内容：见附件", False)])
    add_para(doc,
             [("5、主办单位：", False),
              ("上海市杨浦区科技企业联合会", True)])
    add_para(doc,
             [("6、承办单位：", False),
              ("上海市杨浦区科技企业联合会", True)])
    blank(doc)

    add_heading_line(doc, "第二条  赞助方式", size=12)
    add_para(doc,
             [("1、甲方对本次活动的赞助方式为：", False),
              ("提供赞助费（现金）", True)])

    add_para(doc,
             [("2、如本次活动的赞助方式为提供赞助费的，赞助费总额为¥：", False),
              ("100,000", True),
              ("元，大写：人民币", False),
              ("壹拾万", True),
              ("元整。", False),
              ("乙方上海市杨浦区科技企业联合会系民办非企业单位（民办非企业法人），按小规模纳税人税务管理，仅向甲方开具增值税普通发票（不开具增值税专用发票）。乙方就本合同项下收入的开票方式按以下二者之一执行（具体以乙方在首次开票时的实际税务申报口径为准）：", True),
              ("", False)])
    add_para(doc,
             [("（1）", False),
              ("免税开票：", True),
              ("若乙方就本合同项下收入符合非营利组织增值税免税条件（依据：财税〔2014〕13 号《关于非营利组织免税资格认定管理有关问题的通知》、财税〔2016〕36 号《关于全面推开营业税改征增值税试点的通知》附件 3 及其他非营利组织相关税收规定），乙方按免税口径开具增值税普通发票，发票适用税率栏标注为「免税」，", True),
              ("票面金额为合同总价款 100,000 元，无需另行加收增值税，乙方实际到账金额为 100,000 元。", True)])
    add_para(doc,
             [("（2）", False),
              ("1% 简易征收开票：", True),
              ("若乙方就本合同项下收入因税务政策调整或资质未办妥等原因不符合免税条件，按小规模纳税人简易计税方式适用 1% 增值税征收率开具增值税普通发票，合同总价款 100,000 元保持不变，其中不含税价款约 99,009.90 元，增值税税金约 990.10 元，乙方实际到账金额相应为含税价款。", True)])
    add_para(doc,
             [("本合同履行期间，如遇增值税征收率调整或国家针对非营利组织、民非组织的税收政策变更，", False),
              ("合同总价款 100,000 元保持不变", True),
              ("，开票口径、不含税价款及税金按当时有效的税法规定相应调整。", False)])

    doc.add_page_break()

    # ============================================================
    # 第 3 页
    # ============================================================
    add_para(doc,
             [("3、赞助费按以下第", False),
              ("(2)", True),
              ("种方式支付。若本合同生效日期晚于下述付款期限届满之日，则付款期限相应顺延至本合同生效后 15 个工作日内：", False)])

    add_para(doc, [("（1）分期支付：", False)])
    add_para(doc,
             [("本合同生效之日起", False),
              ("____", False),
              ("个工作日内，甲方支付给乙方本合同总金额的", False),
              ("____", False),
              ("%，即人民币", False),
              ("____", False),
              ("元；", False)])
    add_para(doc,
             [("本次活动结束并经", False),
              ("____", False),
              ("方验收合格后", False),
              ("____", False),
              ("个工作日内，甲方支付给乙方本合同总金额的", False),
              ("____", False),
              ("%，即人民币", False),
              ("____", False),
              ("元。", False)])

    add_para(doc,
             [("（2）", False),
              ("乙方向甲方提交执行报告（含现场图片、嘉宾合影、媒体链接等代表性证据材料）后", True),
              ("，本次活动结束并经甲方验收合格后", False),
              ("15", True),
              ("个工作日内一次性支付。", False),
              ("【视同验收机制】甲方应在收到乙方执行报告后 3 个工作日内书面提出具体异议，逾期未提出书面异议的，视为验收合格；甲方就个别细项的异议不影响无争议部分款项的支付。", True),
              ("（鉴于活动已于 2026 年 5 月 22 日实际完成，本次合作采用方式 (2) 后置一次性付款；甲方应支付乙方合同总价款 100,000 元）", True)])
    add_para(doc,
             [("（3）其他方式：", False),
              ("____________________________________________", False),
              ("。", False)])

    add_para(doc,
             [("4、每次付款前", False),
              ("5", True),
              ("个工作日，乙方应向甲方提供等额合法合规的增值税", False),
              ("普通", True),
              ("发票及服务报告，", False),
              ("【发票合规边界】甲方不得以非实质性的发票形式问题（如版式、抬头错别字等）拒付款项；甲方仅就发票真伪、税号一致性、票面金额一致性等实质性问题享有合理的异议权，并应于收到发票之日起 3 个工作日内书面提出，逾期未提出视为发票接收无异议。", True),
              ("如乙方向甲方开具虚假、伪造或变造的发票，则乙方应自甲方通知之日起七天内向甲方支付该发票票面金额的 20%作为违约金，并重新开具合法有效的发票；如乙方支付的违约金不足以弥补甲方因此遭受的损失，甲方有权向乙方继续追偿。如需开具红字增值税专用发票等，双方应提供方便、互相配合。", False)])
    add_para(doc,
             [("发票内容：", False),
              ("以乙方实际交易内容、税务资质及国家税法规定为准；甲方不得要求乙方开具与实际业务不符的发票。常见可适用科目包括「会议服务费」/「会务咨询费」/「赞助费」，最终以乙方税务申报口径为准。", True),
              ("。", False)])

    add_para(doc, [("甲方开票信息如下：", False)])
    add_para(doc,
             [("公司名称：", False),
              ("________________________________", False),
              ("；", False)])
    add_para(doc,
             [("税务登记证号：", False),
              ("________________________________", False),
              ("；", False)])
    add_para(doc,
             [("地址：", False),
              ("________________________________", False),
              ("；", False)])
    add_para(doc,
             [("电话：", False),
              ("________________________________", False),
              ("；", False)])
    add_para(doc,
             [("开户行账号：", False),
              ("________________________________", False),
              ("；", False)])
    add_para(doc,
             [("开户行全称：", False),
              ("________________________________", False),
              ("。", False)])
    add_para(doc, [("甲方如变更开票信息，应及时书面通知乙方。", False)])

    add_para(doc, [("5、乙方的收款账户为：", False)])
    add_para(doc,
             [("户名：", False),
              ("上海市杨浦区科技企业联合会", True),
              ("；", False)])
    add_para(doc,
             [("账号：", False),
              ("03354200040012533", True),
              ("；", False)])
    add_para(doc,
             [("开户行：", False),
              ("中国农业银行股份有限公司上海营口支行", True),
              ("。", False)])
    add_para(doc,
             [("乙方变更收款账户的，应当提前 5 个工作日书面告知甲方并确认甲方已经签收，否则甲方向上述收款账户支付的款项均视为甲方支付本合同约定的赞助费。", False)])

    doc.add_page_break()

    # ============================================================
    # 第 4 页 — 第三条
    # ============================================================
    add_heading_line(doc, "第三条  甲方的权利和义务", size=12)
    add_para(doc, [("1、甲方应该按照合同的约定履行赞助方的义务。", False)])
    add_para(doc, [("2、甲方有权对活动中涉及本次活动的名称、商标、标签、标志等设计方案以及甲方宣传内容、印有甲方品牌资料的物资设计等提出合理修改意见，乙方应予以执行。", False)])
    add_para(doc, [("3、甲方在根据合同约定进行自主宣传时，有权使用本次活动的名称、商标、标签设计、品名、图案、广告、照片、视频等相关资料，该等使用包括生产、销售或许可生产、销售带有活动相关名称、商标、标识、图案、照片、视频的宣传品、促销品。", False)])
    add_para(doc, [("4、如本合同约定的赞助方式为提供赞助费的，则甲方有权对赞助费的使用情况进行监督和检查。", False)])
    add_para(doc,
             [("5、甲方有权委派", False),
              ("6", True),
              ("名领导/员工代表参与本次活动，并发表讲话，向参会人员介绍和推广甲方及其关联公司、运作的项目、企业文化和成就等。（其中现场服务 2 人 + 销售 2 人 + 甲方指派代表 2 人；晚宴实际出席人数另行确认；具体名单于活动前 3 个工作日双方确认）", True)])
    add_para(doc, [("6、甲方有权对本次活动的有关内容及组织流程、环节设置等事项进行了解并提出合理修改意见，乙方应予以协助修改及配合。", False)])

    add_para(doc, [("7、甲方作为本次活动的赞助方，还享有以下权益：", False)])
    add_para(doc,
             [("1) 甲方拥有本次活动", False),
              ("晚宴部分", True),
              ("的冠名权，本次活动", False),
              ("晚宴专场", True),
              ("统一名称为\u201C", False),
              ("晚宴冠名战略合作伙伴", True),
              ("\u201D，对外宣传及媒体报道时", False),
              ("晚宴部分", True),
              ("均应统一以这个名称进行展示；（本款限定为晚宴冠名，主会议名称不予冠名变更）", True)])
    add_para(doc, [("2) 甲方有权利将甲方及其关联公司名称、LOGO 和运作项目广告等相关信息植入本次活动宣传、活动组织所需物料（包括但不限于邀请函、宣传册、参赛证、活动主背景板、横幅、宣传旗帜、海报、奖品、纪念品等）中。", False)])
    add_para(doc, [("3) 甲方及其关联公司有权结合自身公司需求，在本次活动前期、现场及后期就本次活动内容进行自主宣传，包括但不限于通过报纸、杂志、广播、电视、网络等方式进行宣传。", False)])
    add_para(doc,
             [("4) 其他：", False),
              ("甲方在本次活动中享有以下四大模块综合权益（全包制 · 不作分项标价）：", True)])
    add_para(doc,
             [("（a）", False),
              ("晚宴冠名和专属权益：", True),
              ("晚宴主题冠名、专场 PPT 宣讲、视频轮播、主持人口播、席卡等物料植入；", True)])
    add_para(doc,
             [("（b）", False),
              ("主会场权益：", True),
              ("1 号位展台、项目物料植入参会嘉宾手拎袋、视频轮播奖项颁发；", True)])
    add_para(doc,
             [("（c）", False),
              ("专场项目参观邀约：", True),
              ("论坛/颁奖结束后主持人鸣谢口播 + 现场动线引导，定向邀约意向嘉宾前往项目案场；", True)])
    add_para(doc,
             [("（d）", False),
              ("宣发配合：", True),
              ("现场摄影、朋友圈九宫格、回顾视频项目鸣谢、执行回执（图片 + 合影 + 执行报告）；", True)])
    add_para(doc,
             [("（e）", False),
              ("现场展位：", True),
              ("甲方共享有 3 个现场活动展位（位置由乙方与甲方现场协商确定）；", True)])
    add_para(doc,
             [("（f）", False),
              ("证书奖项：", True),
              ("乙方为甲方颁发\u201C2026 年度智慧人居新质资产领军企业 暨卓越战略合作伙伴\u201D证书 / 奖项一份。", True)])

    add_heading_line(doc, "第四条  乙方的权利和义务", size=12)
    add_para(doc,
             [("1、乙方应按本合同约定如期、合规组织本次活动，并配合甲方完成 logo、品牌简介、甲方指派代表宣讲 PPT、户型折页等权益落地所需的素材确认与上线工作；乙方就邀请的（客户）出席率尽合理之力达到", False),
              ("80", True),
              ("%以上。", False),
              ("（参会名单及嘉宾联络方式涉及个人信息保护，根据《中华人民共和国个人信息保护法》第十三条、第二十三条之规定，未经嘉宾本人单独授权同意，乙方不得向第三方提供，故本合同不就参会名单及嘉宾构成作硬性提交承诺，相关意向客户对接由乙方在现场协助促成。）", True)])

    doc.add_page_break()

    # ============================================================
    # 第 5 页
    # ============================================================
    add_para(doc, [("2、乙方有权决定或修改本次活动有关内容及组织流程，但应提前征得甲方书面同意，甲方应配合乙方或乙方委托的承办方及现场工作人员的安排。", False)])
    add_para(doc,
             [("3、", False),
              ("乙方应按本合同第二条约定的资金用途（场地租赁、舞美搭建、嘉宾接待、物料制作及现场运营等直接费用），将甲方提供的赞助费专款专用于本次活动当期支出；乙方对资金使用承担合理说明义务，可按甲方书面要求提供资金用途的合理说明文件（如分项费用汇总、主要供应商发票影像等）。", True),
              ("（本次合作属于带有曝光对价的商业赞助，非慈善捐赠或政府专项拨款；乙方按约完成本合同第三条项下各项权益交付后，赞助费所有权即归属乙方；乙方作为独立法人，其内部财务账目、商业秘密及经营自主权受法律保护，不接受穿透式审计或对账，但配合合理范围内的资金用途说明请求。）", True)])
    add_para(doc,
             [("4、", False),
              ("本活动举办过程中，因乙方组织不当、故意或重大过失造成的人身损害或者财产损失事故，由乙方承担相应责任；因甲方原因（含甲方人员行为）造成的损失由甲方自行承担；因第三方原因（含场地方、第三方供应商、不可控因素、意外事件）导致的损失，由相应责任方按法定责任划分承担。乙方仅就因自身过错造成的、直接、实际、可预见的损失向甲方承担赔偿责任。", True)])
    add_para(doc,
             [("5、", False),
              ("乙方及其官方可控渠道（含联合会官网、官方公众号、官方视频号等）不得发布、转载、链接等关于甲方及其关联方品牌及其下属项目的不实诋毁性报道；如乙方发现第三方媒体、论坛、个人评论对甲方及其关联方有不实负面报道，乙方应协助甲方采取措施（如向平台投诉要求屏蔽或删除等）并通知甲方。乙方不承担非因乙方故意或重大过失造成的第三方独立侵权及损失赔偿责任。", True)])
    add_para(doc,
             [("6、乙方应按照甲方的要求对甲方及其关联方品牌及其下属项目在本次活动的前期宣传、现场及后期宣传中进行推广。", False)])
    add_para(doc, [("7、乙方应确保本次活动的计划性、完整性、条理性、有序性、合法性和合理性，并在甲方的监督指导下，严格按照本合同约定的活动时间、活动内容、活动要求举办。", False)])
    add_para(doc, [("8、乙方作为本次活动的主办方，或经主办方授权作为本次活动的资源经营机构及平台（已在本合同签订前向甲方提供主办方的相关授权文件），全权负责本次活动赞助招商各项事宜。", False)])
    add_para(doc,
             [("9、其他约定：", False),
              ("____________________（无；如有特别约定，双方另行书面确认）", True),
              ("。", False)])
    blank(doc)

    add_heading_line(doc, "第五条  知识产权", size=12)
    add_para(doc, [("1、本合同所涉甲方商标、宣传物料、设计图纸、广告、摄影作品等相关资料的知识产权及相应的财产权利归甲方所有。未经甲方书面同意，乙方不得擅自或者委托、授权第三方使用，否则，乙方应将因此所获利益返还给甲方，并应按本合同总价款的 30%向甲方支付违约金。", False)])
    add_para(doc,
             [("2、", False),
              ("乙方就其自行提供的活动内容（不含甲方提供的物料、品牌素材、宣传文案等）不侵犯任何第三方的知识产权或在先合法权利作出声明。", True),
              ("如因乙方自行提供的内容导致第三方提起维权主张的，乙方应在自身过错范围内承担相应法律责任。甲方提供的物料、品牌素材、宣传文案等导致的第三方维权主张，由甲方自行承担。", False)])

    doc.add_page_break()

    # ============================================================
    # 第 6 页
    # ============================================================
    add_para(doc, [("乙方就本款项下违约对甲方承担的赔偿责任，不超过本合同总价款的 30%；甲方提供的物料导致的第三方维权主张所产生的损失，由甲方独立承担。", False)])
    blank(doc)

    add_heading_line(doc, "第六条  保密义务", size=12)
    add_para(doc, [("1、甲乙双方应对其通过订立和履行本合同而获悉的对方的商业秘密、未对外公开的重要信息严格保密，未经对方事先书面同意，不得以任何方式进行利用、向任何第三方披露或以其他方式予以公开。", False)])
    add_para(doc, [("2、未经对方许可，任何一方不得向第三方（有关法律、法规、政府部门、证券交易所或其他监管机构要求和双方的法律、会计、商业及其他顾问、雇员除外）泄露本合同的任何内容以及本合同的签订及履行情况，以及通过签订和履行本合同而获知的对方及对方关联公司的任何信息。", False)])
    add_para(doc, [("3、除订立与履行本合同项下义务之需要外，未经对方事先同意，任何一方不得擅自使用、复制对方的商标、标志、商业信息、广告样稿资料、技术及其他资料。", False)])
    add_para(doc, [("4、甲乙双方的上述保密义务不因本合同的无效、终止或被解除而终止。", False)])
    blank(doc)

    add_heading_line(doc, "第七条  违约责任", size=12)
    add_para(doc,
             [("1、甲方因自身原因无故逾期支付合同价款的，每逾期一日，按逾期应付未付金额的", False),
              ("千分之一", True),
              ("向乙方支付违约金；逾期超过 15 个自然日的，乙方有权解除本合同并要求甲方一次性支付全部应付未付款项及违约金。", False)])
    add_para(doc,
             [("2、", False),
              ("【双方确认】鉴于第四条第 1 项乙方仅就出席率承担「尽合理之力」的努力义务（非保证义务），且参会人数受活动行业自然波动、嘉宾临时变更等不可控因素影响，本款约定：除乙方故意或重大过失导致核心权益无法交付外，参会人数波动不构成乙方违约，亦不得作为甲方拒付、扣款或解除本合同的依据；本款关于「出席率违约金 / 不一致率违约金」相应不再适用，留为占位以备后续双方另行书面约定。", True)])
    add_para(doc,
             [("3、", False),
              ("（鉴于第四条第 1 项乙方义务已基于《中华人民共和国个人信息保护法》第十三条、第二十三条之规定，未对参会名单及嘉宾联络方式作硬性提交承诺，本款关于「未在签订时提供参会名单」的违约责任不再适用，留为占位以备后续双方另行书面约定。）", True)])

    doc.add_page_break()

    # ============================================================
    # 第 7 页
    # ============================================================
    add_para(doc,
             [("4、如乙方未按照本合同约定的用途使用甲方提供的赞助费或赞助物资等的，则甲方有权解除本合同，在本合同解除之日起的 5 个工作日内，乙方应返还甲方已提供的赞助费或赞助物资，同时向甲方支付", False),
              ("5,000", True),
              ("元违约金，并赔偿甲方因此遭受的全部损失。如逾期返还甲方已支付的赞助费、逾期支付违约金或逾期返还赞助物资的，则每逾期一日，乙方应向甲方支付", False),
              ("100", True),
              ("元违约金。", False)])
    add_para(doc,
             [("5、", False),
              ("【活动已完成 · 本款不再适用】鉴于本次活动已于 2026 年 5 月 22 日实际圆满完成，本款关于「活动无法如期举行」「逾期超过 10 日未举行」「活动取消违约金」等未来履约场景的违约责任不再适用，留为占位以备后续年度活动另行书面约定。", True),
              ("（如未来双方就同类活动续签合同，前述违约场景应另行约定违约责任，且因不可抗力（参见第八条）导致活动无法举行的，乙方不承担违约责任。）", True)])
    add_para(doc, [("6、乙方违反保密义务的，应当按照本合同总价款的 30%向甲方支付违约金，如不足以弥补甲方损失的，应予补足。", False)])
    add_para(doc, [("7、除本合同另有约定外，本合同签订后，双方应严格履行合同义务，任何一方单方面终止合同的，应按照本合同总价款的 20%向对方支付违约金。", False)])
    add_para(doc,
             [("8、", False),
              ("除前述约定外，任何一方有其他违反本合同约定的情形的，违约方应赔偿给守约方因此遭受的直接、实际、可预见的经济损失。", True),
              ("（双方明确：守约方主张的赔偿范围不包括间接损失、机会损失、商誉损失、惩罚性赔偿及《中华人民共和国民法典》第 584 条规定的违约方在订立合同时不可预见的损失。）", True)])
    add_para(doc,
             [("9、", False),
              ("【单方扣款限制】甲方拟从应付未付乙方的合同款项中扣除本合同项下乙方应付的违约金及赔偿金等，应经乙方书面确认或经生效的人民法院判决/仲裁裁决文书确认后方可抵扣；不得未经协商和书面确认径行扣款。", True)])
    add_para(doc,
             [("10、", False),
              ("双方按本合同约定的违约责任主张赔偿时，除律师费、诉讼费、保全费等与争议解决直接相关的合理费用外，不得扩大解释、不得主张不可预见或间接性损失；甲方追究违约责任或索赔的具体范围以本合同约定为限。", True)])
    add_para(doc,
             [("11、", False),
              ("【双方违约责任上限】除因故意、重大过失、保密义务（第六条）、知识产权侵权（第五条）情形外，乙方就本合同项下任何违约（含累计违约）向甲方承担的赔偿责任及违约金总额，不超过本合同总价款的 30%（即不超过人民币 30,000 元）；超过部分乙方不予承担。本款同样适用于甲方，甲方就其违约向乙方承担的赔偿责任及违约金总额亦不超过本合同总价款的 30%。", True)])

    doc.add_page_break()

    # ============================================================
    # 第 8 页
    # ============================================================
    blank(doc)

    add_heading_line(doc, "第八条  不可抗力", size=12)
    add_para(doc, [("1、不可抗力指由于不能预见、不能避免和不能克服的自然原因或社会原因，致使本合同不能履行或者不能完全履行的情形，包括战争、暴乱、空中飞行物体坠落及因此或其他非甲、乙双方责任造成的爆炸、火灾，严重的自然灾害，法律规定的其他情形等。", False)])
    add_para(doc, [("2、遭遇不可抗力的一方，应立即书面通知另一方，并应在不可抗力发生后 15 天内，向另一方提供不可抗力发生地政府部门出具的证明合同不能履行或需要延期履行、部分履行的有效证明文件。", False)])
    add_para(doc, [("3、双方按不可抗力对合同履行的影响程度协商决定是否解除合同、延期履行合同或者部分履行合同。", False)])
    add_para(doc, [("4、对因不可抗力未能履约给另一方造成的经济损失，该履约方不负赔偿责任，但本合同另有约定的除外。", False)])
    add_para(doc,
             [("5、关于新型冠状病毒引发的肺炎疫情的特别约定", False),
              ("【活动已完成 · 本款不再适用】鉴于本次活动已于 2026 年 5 月 22 日实际圆满完成，本款关于新冠疫情期间的费用承担、现场管控等约定不再适用，留为占位仅作模板格式保留。如未来双方就同类活动续签合同，相关疫情或公共卫生事件下的费用承担应由双方按公平原则另行协商分担。", True)])
    blank(doc)

    add_heading_line(doc, "第九条  通知与送达", size=12)

    doc.add_page_break()

    # ============================================================
    # 第 9 页
    # ============================================================
    add_para(doc, [("1、本合同项下任何一方向对方发出的通知、信件、数据电文等，应当发送至本合同下列约定的地址、联系人和通信终端。一方当事人变更名称、地址、联系人或通信终端的，应当在变更后 3 日内及时书面通知对方当事人，对方当事人实际收到书面变更通知前的送达仍为有效送达。任何根据本合同发出的文件，均应当采取书面形式。", False)])
    add_para(doc, [("2、联系人及联系方式：", False)])
    add_para(doc,
             [("甲方联系人：", False),
              ("________________________________", False),
              ("；", False)])
    add_para(doc,
             [("联系电话：", False),
              ("________________________________", False),
              ("；", False)])
    add_para(doc,
             [("联系地址：", False),
              ("________________________________", False),
              ("。", False)])
    add_para(doc,
             [("乙方联系人：", False),
              ("胡继刚", True),
              ("；", False)])
    add_para(doc,
             [("联系电话：", False),
              ("13262607888", True),
              ("；", False)])
    add_para(doc,
             [("联系地址：", False),
              ("上海市杨浦区国泰路 11 号 复旦科技园副楼一楼 1008 室", True),
              ("。", False)])
    add_para(doc, [("3、双方确认上述送达地址之适用范围包括双方非诉时各类通知、协议等文件以及就合同发生纠纷时相关文件和法律文书的送达，同时包括因履行本合同产生诉讼进入仲裁、民事诉讼程序后的一审、二审、再审和执行程序时需要送达的资料和文书（包括但不限于诉讼参加人提交的材料、裁决机构调取的材料、应诉通知书、举证通知书、合议庭告知书、传票、开庭通知书、判决书、裁定书、调解书、限期履行通知书等文书和资料）的送达。", False)])
    add_para(doc, [("4、一方采取邮寄方式通知的，上述通讯地址为送达地址，以该通讯地址寄出 3 日（本省）或 7 日后（外省）视为送达。对于双方在合同中明确约定的送达地址，法院进行送达时可直接邮寄送达，即使当事人未能收到法院邮寄送达的文书，由于双方的约定，也应当视为送达。", False)])
    add_para(doc, [("5、本约定作为合同中独立存在的条款，不受合同其他条款效力的影响。", False)])
    blank(doc)

    add_heading_line(doc, "第十条  其他事项", size=12)
    add_para(doc,
             [("1、因本合同引起的一切争议，双方首先应当友好协商解决；协商不成，双方确认向", False),
              ("乙方所在地（即上海市杨浦区）", True),
              ("有管辖权的人民法院提起诉讼。", False)])
    add_para(doc, [("2、本合同的附件为本合同的有效组成部分，与本合同具有同等法律效力。", False)])
    add_para(doc, [("3、本合同经双方盖章后生效。", False)])

    doc.add_page_break()

    # ============================================================
    # 第 10 页
    # ============================================================
    add_para(doc,
             [("4、本合同一式", False),
              ("肆", True),
              ("份，甲方执", False),
              ("贰", True),
              ("份、乙方执", False),
              ("贰", True),
              ("份，均具同等法律效力。", False)])
    blank(doc)

    add_para(doc, [("（以下无正文）", False)], align="center")

    doc.add_page_break()

    # ============================================================
    # 第 11 页：签署页
    # ============================================================
    add_para(doc, [("（本页为《活动赞助合同》签署页，无正文）", False)],
             align="center")
    blank(doc, 4)

    add_para(doc,
             [("甲方（盖章）：", False),
              ("________________________________", False)],
             line_spacing=2.0)
    blank(doc, 4)
    add_para(doc,
             [("法定代表人或授权代表（签名或盖章）：", False),
              ("________________________________", False)],
             line_spacing=2.0)
    blank(doc, 4)

    add_para(doc,
             [("乙方（盖章）：", False),
              ("上海市杨浦区科技企业联合会", True)],
             line_spacing=2.0)
    blank(doc, 4)
    add_para(doc,
             [("法定代表人或授权代表（签名或盖章）：", False),
              ("胡继刚", True)],
             line_spacing=2.0)

    out = "/workspace/deliverables/绿城活动赞助合同(GTH-zzc-ppch-016-2022a)-补充填写版.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
