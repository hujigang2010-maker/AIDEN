"""Generate a concise (2-5 pages) version of the Tencent Cloud Strategic
Sponsorship Agreement, used solely for showing during Greentown's
bid-evaluation process. Keeps the legally meaningful clauses, drops the
detailed schedules / material specs / receipt template."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DARK_BLUE = RGBColor(0x00, 0x3A, 0x99)
TENCENT_BLUE = RGBColor(0x00, 0x52, 0xD9)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)


def _force_font(run, font="宋体"):
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rpr.append(rfonts)


def set_default_font(doc, name="宋体"):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def heading(doc, text, level=1, color=None, center=False, font="黑体"):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(18)
    elif level == 1:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    if color is not None:
        run.font.color.rgb = color
    _force_font(run, font=font)
    return p


def para(doc, text, bold=False, size=10.5, color=DARK, align=None,
         first_indent=True, font="宋体", line_spacing=1.4):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(2)
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    _force_font(run, font=font)
    return p


def clause(doc, num, text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(2)
    run_num = p.add_run(num + "  ")
    run_num.bold = True
    run_num.font.size = Pt(size)
    _force_font(run_num)
    run = p.add_run(text)
    run.font.size = Pt(size)
    _force_font(run)
    return p


def blank(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, color=DARK, size=10.5,
              align_center=False, font="宋体"):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    _force_font(run, font=font)


def build():
    doc = Document()
    set_default_font(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # ===== 抬头 =====
    heading(doc, "2026 人工智能商业化落地与硬核投资破局峰会",
            level=2, color=GREY, center=True)
    heading(doc, "《战略协办赞助合作协议》",
            level=0, color=DARK_BLUE, center=True)
    blank(doc)
    para(doc,
         "合同编号：AIBIZ-2026-SP-013      签署地点：上海市      签署日期：2026 年 5 月 12 日",
         first_indent=False, size=10, align="center", color=GREY)
    blank(doc)

    # ===== 协议主体 =====
    para(doc,
         "甲方（主办方 / 组委会）：人工智能商业化落地峰会组委会 / 复旦大学住房政策研究中心 / 上海市杨浦区科技企业联合会（联合主办）",
         bold=True, first_indent=False, size=10.5)
    para(doc,
         "乙方（战略协办方）：腾讯云计算（北京）有限责任公司",
         bold=True, first_indent=False, size=10.5)
    blank(doc)

    # ===== 鉴于 =====
    para(doc, "鉴于：", bold=True, first_indent=False)
    para(doc,
         "1. 甲方拟于 2026 年 5 月在上海·北外滩核心地标举办「2026 人工智能商业化落地与硬核投资破局峰会」（以下简称「本次峰会」），聚焦人工智能在新质生产力、大模型商业化与产业数字化转型等领域的落地实践；")
    para(doc,
         "2. 乙方作为国内领先的云计算与人工智能基础设施服务商，希望通过本次峰会向核心政企客户与产业生态展示其在算力与大模型领域的商业化能力；")
    para(doc,
         "3. 双方依据《中华人民共和国民法典》《广告法》《反不正当竞争法》《商标法》及其他相关法律法规，本着平等自愿、互利共赢的原则，就乙方战略协办本次峰会的相关事宜，达成如下协议，以资共同遵守。")
    blank(doc)

    # ===== 第一条 =====
    heading(doc, "第一条  合作级别与赞助标的", level=1, color=DARK_BLUE)
    clause(doc, "1.1",
           "合作级别：乙方为本次峰会唯一的「首席战略合作伙伴」及「算力生态独家伙伴」，享有同一级别下不可被并列、不可被超越的最高排序权益。")
    clause(doc, "1.2",
           "赞助标的额：乙方为本次峰会提供专项赞助资金，总计金额为人民币（大写）叁拾万元整（小写：￥300,000.00 元），含税。")
    clause(doc, "1.3",
           "支付方式：乙方应于本协议签署并经双方盖章生效后 15 个工作日内，将上述款项一次性汇入甲方指定的对公账户。甲方在收款后向乙方开具相应金额的合规发票（发票内容：会议服务费 / 赞助费）。")
    clause(doc, "1.4",
           "资金用途：上述款项专款专用于本次峰会当期的场地租赁、舞美搭建、嘉宾接待、物料制作及现场运营等直接费用。")
    blank(doc)

    # ===== 第二条 =====
    heading(doc, "第二条  甲方权益交付（全包制 · 不作分项标价）", level=1, color=DARK_BLUE)
    para(doc,
         "为对应乙方上述赞助投入，甲方在本次峰会及其后续生态互动中，向乙方综合交付包括但不限于以下权益。双方确认：下列权益作为一揽子综合对价整体交付，不再就单项权益进行分项估值与单独定价，亦不作为单项退款依据。")
    clause(doc, "2.1",
           "核心现场曝光：主背景板、官方议程、邀请函、白皮书、媒体通稿、签到背板及核心传播物料中的「首席战略合作伙伴」顶级 logo 并排首位露出。")
    clause(doc, "2.2",
           "议题主导与高管发声：甲方为乙方核心高管在主会场保留 1 个主旨演讲席位或 1 个圆桌核心席位（具体形式由双方协商）。")
    clause(doc, "2.3",
           "顶尖圈层深度对接：甲方依托自身政企与智库资源，为乙方定向引荐具备真实数字化转型与算力需求的大型重资产企业，安排闭门深度对接环节。")
    clause(doc, "2.4",
           "算力生态延伸权益（自带，不计入第 1.2 条赞助标的额）：乙方有权在峰会后向甲方生态内符合资质的科创企业与「超级个体」定向发放腾讯云专属算力包及大模型 API 调用支持，甲方通过自有渠道配合宣发。")
    clause(doc, "2.5",
           "智库长效背书（自带，不计入第 1.2 条赞助标的额）：乙方自动成为「见微知海新质商业生态」首批年度战略理事单位，享有后续长三角系列闭门局、专题沙龙及行业研究发布会的优先参与权。")
    blank(doc)

    # ===== 第三条 =====
    heading(doc, "第三条  双方义务", level=1, color=DARK_BLUE)
    clause(doc, "3.1",
           "甲方义务：（1）严格按本协议第二条约定，向乙方完整、及时交付各项权益；（2）保障峰会按计划如期举办，并维护峰会的行业影响力与规格档次；（3）在峰会筹备、执行、后续传播全过程中，确保乙方作为「首席战略合作伙伴」及「算力生态独家伙伴」的排他性与唯一性，不引入与乙方在云计算及算力领域形成直接竞争关系的同级别合作方；（4）妥善管理乙方支付的赞助资金，确保专款专用。")
    clause(doc, "3.2",
           "乙方义务：（1）按本协议约定按时足额支付赞助款项；（2）配合甲方完成 logo、高管简介、演讲主题等权益落地所需的素材提供与确认工作；（3）不得利用本次峰会平台从事违反国家法律法规及公序良俗的活动。")
    blank(doc)

    # ===== 第四条 =====
    heading(doc, "第四条  知识产权、保密及独立比对声明", level=1, color=DARK_BLUE)
    clause(doc, "4.1",
           "知识产权：乙方授权甲方在本次峰会的宣传、报道、白皮书、官网、议程手册、媒体通稿、回顾视频等场景中使用乙方商标与企业简介；授权范围限于本次峰会及其衍生宣传内容，授权期限至本次峰会结束后 12 个月。")
    clause(doc, "4.2",
           "商业机密保护：未经甲乙双方书面同意，任何一方不得向任何第三方（包括但不限于其他赞助商、合作方、媒体及外部审计机构）披露本协议项下的资金打款凭证、财务流水、内部核算细则及权益分项估值等敏感信息，保密期限自本协议签订之日起 3 年。")
    clause(doc, "4.3",
           "独立比对声明：甲方有权在对外招商、生态展示及向同级别合作方进行横向拉齐时，合法引用本协议第一条项下的合作总金额（叁拾万元）及合作级别（首席战略合作伙伴 + 算力生态独家伙伴）作为优秀案例说明，以证明本次峰会的市场估值与顶尖行业含金量。该等引用不视为对本条第 4.2 款保密义务的违反。")
    clause(doc, "4.4",
           "合规承诺：双方承诺，本协议项下的资金往来与权益交付均合法合规，不存在任何商业贿赂、利益输送及违反反不正当竞争法、反商业贿赂相关规定的情形。")
    blank(doc)

    # ===== 第五条 =====
    heading(doc, "第五条  违约责任与不可抗力", level=1, color=DARK_BLUE)
    clause(doc, "5.1",
           "任何一方未按本协议约定履行义务，给对方造成损失的，应承担相应的违约责任并赔偿对方因此遭受的直接经济损失。")
    clause(doc, "5.2",
           "若因不可抗力（包括但不限于自然灾害、政府管制、重大公共卫生事件等）导致本次峰会延期或无法举办，双方应友好协商，按已发生的实际支出比例处理已支付的赞助款项；剩余款项可顺延用于甲方下一届同级别峰会，或经双方书面同意后退还乙方。")
    blank(doc)

    # ===== 第六条 =====
    heading(doc, "第六条  争议解决与法律适用", level=1, color=DARK_BLUE)
    clause(doc, "6.1",
           "本协议的签订、履行、解释及争议解决均适用中华人民共和国法律。")
    clause(doc, "6.2",
           "因本协议产生或与本协议有关的任何争议，双方应首先通过友好协商解决；协商不成的，任何一方均有权向甲方所在地有管辖权的人民法院提起诉讼。")
    blank(doc)

    # ===== 第七条 =====
    heading(doc, "第七条  签署与生效", level=1, color=DARK_BLUE)
    clause(doc, "7.1",
           "生效条件：本协议自甲乙双方法定代表人或授权代表签字并加盖公章（或合同专用章）之日起生效。")
    clause(doc, "7.2",
           "物理要求：本协议正本须加盖骑缝章，一式贰份，甲乙双方各执壹份，具有同等法律效力。")
    clause(doc, "7.3",
           "附件效力：本协议如有附件、补充协议或经双方书面确认的往来函件（含电子邮件、企业微信、钉钉等电子形式），均为本协议不可分割的组成部分，与本协议具有同等法律效力。")
    blank(doc)

    # ===== 签署页 =====
    para(doc, "（以下无正文，下为签署页）",
         align="center", first_indent=False, size=9, color=GREY)
    blank(doc)

    sig = doc.add_table(rows=5, cols=2)
    sig.style = "Table Grid"
    sig.autofit = False
    for row in sig.rows:
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(8.0)
    cell_text(sig.rows[0].cells[0], "甲方（盖章）：", bold=True, size=11)
    cell_text(sig.rows[0].cells[1], "乙方（盖章）：", bold=True, size=11)
    cell_text(sig.rows[1].cells[0],
              "人工智能商业化落地峰会组委会 /\n复旦大学住房政策研究中心 /\n上海市杨浦区科技企业联合会",
              size=10.5)
    cell_text(sig.rows[1].cells[1],
              "腾讯云计算（北京）有限责任公司",
              size=10.5)
    cell_text(sig.rows[2].cells[0],
              "\n\n\n（公章 / 合同章位置）\n\n",
              size=10.5, align_center=True)
    cell_text(sig.rows[2].cells[1],
              "\n\n\n（公章 / 合同章位置）\n\n",
              size=10.5, align_center=True)
    cell_text(sig.rows[3].cells[0],
              "授权代表（签字）：______________________",
              size=11)
    cell_text(sig.rows[3].cells[1],
              "授权代表（签字）：______________________",
              size=11)
    cell_text(sig.rows[4].cells[0],
              "日期：______年______月______日",
              size=11)
    cell_text(sig.rows[4].cells[1],
              "日期：______年______月______日",
              size=11)

    out = "/workspace/deliverables/腾讯云-首席战略协办伙伴合作协议(精简版).docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
