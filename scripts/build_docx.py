# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from content import (
    AGENT_JOIN,
    DISCLAIMER,
    EVENT,
    FORWARD_LONG,
    FORWARD_REPLAY,
    FORWARD_SHORT,
    GUESTS,
    SERIES_NOTES,
    WATCH_FLOW,
    WHY_WATCH,
)

OUT = Path(__file__).resolve().parents[1] / "exports" / "飞书DemoDay4_观摩备忘录.docx"
NAVY = RGBColor(0x14, 0x3A, 0x7A)
BLUE = RGBColor(0x1F, 0x5F, 0xBF)
DARK = RGBColor(0x1A, 0x2A, 0x3A)
GREY = RGBColor(0x5A, 0x6A, 0x7A)
WARN = RGBColor(0x6A, 0x5A, 0x3A)


def set_run(run, size=12, bold=False, color=DARK, name="微软雅黑"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def p(doc, text, size=12, bold=False, color=DARK, space=8, align="left"):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = 1.25
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return para


def h(doc, text, level=1):
    size = {1: 20, 2: 16, 3: 13}[level]
    color = NAVY if level == 1 else BLUE
    p(doc, text, size=size, bold=True, color=color, space=10)


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(item)
        set_run(run, size=12)


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    p(doc, "AIDEN 观摩材料  ·  非飞书官方议程", size=11, color=GREY)
    p(doc, EVENT["title"], size=26, bold=True, color=RGBColor(0x0B, 0x16, 0x2C), space=4)
    p(doc, "观摩备忘录  ·  转发口径  ·  Agent 入会对照清单", size=14, color=BLUE, space=12)
    p(doc, DISCLAIMER, size=10, color=WARN, space=16)

    h(doc, "1. 先看结论")
    bullets(
        doc,
        [
            "这是飞书 AI Builder Demo Day 第四场，主题「让 AI 接住你的日常小麻烦」。",
            "官方分享日历：2026 年 7 月 2 日（周四）11:00–12:30（GMT+8），组织者 Ni Dan，时长 90 分钟。",
            "打开日历页时状态为「已结束」。先点原链接，确认有没有回放或加场。",
            "用户预告：张咋啦、向阳乔木会出 Demo；新能力「Agent 入会」大概率按惯例现场展示。",
            "上一场容易约满。要看直播，第一件事是加会议，不是写观后感。",
        ],
    )

    h(doc, "2. 官方已核对信息")
    rows = [
        ("系列", EVENT["series"]),
        ("场次标题", EVENT["title_share"]),
        ("时间", "2026-07-02 11:00–12:30（北京时间）"),
        ("时长", EVENT["duration"]),
        ("形式", EVENT["format"]),
        ("组织者", EVENT["organizer"]),
        ("状态", EVENT["status_zh"]),
        ("预约 / 入会", EVENT["share_url"]),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for j, cell in enumerate(table.rows[i].cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run(run, size=11, bold=(j == 0))
    p(doc, "", space=8)

    h(doc, "3. 为什么值得看")
    bullets(doc, WHY_WATCH)

    h(doc, "4. 系列背景（只写能核到的）")
    bullets(doc, SERIES_NOTES)
    p(doc, f"第二期回放：{EVENT['replay_ep2']['title']}  {EVENT['replay_ep2']['url']}")
    p(doc, f"检索入口：{EVENT['bilibili_search']}", space=12)

    h(doc, "5. 预告嘉宾怎么看")
    p(doc, "以下不是官方嘉宾表，是转发预告 + 公开资料，用来准备耳朵，不当成台词。", size=11, color=WARN)
    for g in GUESTS:
        h(doc, f"{g['name']}（{g['aka']}）", 2)
        p(doc, g["why"])
        bullets(doc, g["watch"])

    h(doc, "6. Agent 入会：本场最大看点")
    p(doc, AGENT_JOIN["source"], size=11, color=WARN)
    p(doc, AGENT_JOIN["what"])
    h(doc, "公开能力边界", 3)
    p(doc, f"技能名：{AGENT_JOIN['official_skill']}", bold=True)
    p(doc, AGENT_JOIN["join_cmd"])
    p(doc, "能做", bold=True, space=4)
    bullets(doc, AGENT_JOIN["can"])
    p(doc, "先别指望", bold=True, space=4)
    bullets(doc, AGENT_JOIN["cannot"])
    p(doc, "现场只问", bold=True, space=4)
    bullets(doc, AGENT_JOIN["watch_questions"])

    h(doc, "7. 观摩节奏")
    for phase, items in WATCH_FLOW.items():
        h(doc, phase, 3)
        bullets(doc, items)

    h(doc, "8. 可直接转发")
    h(doc, "短口径", 3)
    p(doc, FORWARD_SHORT)
    h(doc, "完整口径", 3)
    p(doc, FORWARD_LONG)
    h(doc, "已结束 / 回放口径", 3)
    p(doc, FORWARD_REPLAY)

    h(doc, "9. 会后 24 小时最小复现")
    bullets(
        doc,
        [
            "开一场只有自己人的测试会，不要用客户会、董事会、人事面谈。",
            "若演示了 Agent 入会：记下会议号位数、是否要密码、是否要主持人放行。",
            "对照：入会是否出现在参会列表、会中能否发消息、纪要从哪一页导出。",
            "把「可抄三步」写进飞书文档，下一场会前当 brief 丢给 Agent。",
            "复现失败也要记：失败点往往比 Demo 更有用。",
        ],
    )

    h(doc, "10. 使用边界")
    bullets(
        doc,
        [
            "本备忘录供内部转发和观摩，不代表飞书官方立场。",
            "不要把预告嘉宾写成已确认赞助或联名。",
            "不要把 CLI 命令写成「本场唯一官方玩法」。现场以产品经理演示为准。",
            "涉及真实会议内容时，遵守所在组织的录屏、纪要和外发制度。",
        ],
    )
    p(doc, "生成说明：口径来自用户转发原文 + 飞书分享日历 INIT_DATA + 公开检索。", size=10, color=GREY, space=0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"wrote {OUT}")
    return OUT


if __name__ == "__main__":
    build()
