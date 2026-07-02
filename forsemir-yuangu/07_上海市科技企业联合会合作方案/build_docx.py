"""Build 《上海市科技企业联合会合作服务方案》正式 Word 文档."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name('上海市科技企业联合会合作服务方案.docx')


def set_run(run, *, size=11, bold=False, color=None, font_name='微软雅黑'):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font_name)
    rfonts.set(qn('w:ascii'), font_name)
    rfonts.set(qn('w:hAnsi'), font_name)


def add_para(doc, text, *, size=11, bold=False, color=None, align=None,
             indent_first=True, space_after=4, left_indent=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent:
        p.paragraph_format.left_indent = Pt(left_indent)
    if indent_first:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    size = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 4)
    run = p.add_run(text)
    set_run(run, size=size, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))
    return p


def make_table(doc, headers, rows, widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.autofit = False
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    hdr = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = h
        for p in hdr[j].paragraphs:
            for r in p.runs:
                set_run(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc_pr = hdr[j]._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '142C5E')
        tc_pr.append(shd)
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, v in enumerate(row):
            cells[j].text = v
            for p in cells[j].paragraphs:
                for r in p.runs:
                    set_run(r, size=10.5)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # --- Cover ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(30)
    r = title.add_run('上 海 市 科 技 企 业 联 合 会')
    set_run(r, size=22, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))
    title.add_run('\n')
    r2 = title.add_run('合 作 服 务 方 案')
    set_run(r2, size=22, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Shanghai Science & Technology Enterprise Association — Cooperation Service Plan')
    set_run(r, size=11, color=RGBColor(0x55, 0x60, 0x7A))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(28)
    for line in [
        '版本:v1.0',
        '起草方:胡教授团队',
        '签署日期:__________ 年 ____ 月 ____ 日',
    ]:
        r = meta.add_run(line)
        set_run(r, size=11)
        meta.add_run('\n')

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(30)
    r = note.add_run('本方案为商务洽谈核心条款描述,不构成最终法律约束力的合作合同,最终以双方法务确认并签署的正式协议文本为准。')
    set_run(r, size=10, color=RGBColor(0x99, 0x33, 0x33))

    doc.add_page_break()

    # --- 一、合作主体 ---
    heading(doc, '一、合作主体', level=1)
    make_table(doc, ['角色', '主体'], [
        ['甲方', '上海市科技企业联合会'],
        ['乙方', '胡教授团队(具体合作主体以双方书面确认为准)'],
    ], widths_cm=[3.5, 12.5])

    # --- 二、服务内容 ---
    heading(doc, '二、服务内容', level=1)

    heading(doc, '1. 科技企业服务中心建设', level=2)
    add_para(doc, '建设目标:协助甲方落地"科技企业服务中心"实体平台或品牌化运营场所,作为甲方产业服务能力对外输出的统一窗口。', indent_first=False)
    add_para(doc, '核心工作:', indent_first=False, bold=True)
    for line in [
        '(1) 中心定位、功能模块、空间动线与服务清单的总体规划设计;',
        '(2) 服务体系搭建 — 含财税、法律、知识产权、政府补贴申报、人才与签证、投融资、品牌公关、数字化工具、培训认证等九大类标准化服务路径;',
        '(3) 中心 VI 与对外品牌识别系统设计支持;',
        '(4) 挂牌仪式策划与对外正式发布;',
        '(5) 首批入驻企业邀约(不少于 20 家科技 / 产业企业)。',
    ]:
        add_para(doc, line, indent_first=False, left_indent=20)
    add_para(doc, '交付周期:合同生效后 90 日内完成中心挂牌运营。', indent_first=False)
    add_para(doc, '配套人员:项目经理 1 人全程驻场 + 兼职专家顾问按需调用。', indent_first=False)

    heading(doc, '2. 年度品牌活动执行', level=2)
    add_para(doc, '活动数量:全年承办 6 场科技主题品牌活动,按季均匀分布、每两个月 1 场。', indent_first=False)
    add_para(doc, '单场规模:到场目标产业客户不少于 30 家;目标媒体声量不少于 100 万次曝光。', indent_first=False)
    add_para(doc, '服务范围:主题策划、议程设计、嘉宾邀约、场地与物料、现场执行、传播投放、政府与协会资源对接、活动后传播复盘。', indent_first=False)
    add_para(doc, '建议主题矩阵(首年 6 场):', indent_first=False, bold=True)
    for i, t in enumerate([
        'AI 与硬科技商业化落地',
        '潮玩文创与新消费',
        '出海加速(北欧 / 日韩 / 东南亚)',
        '投融资路演与产业基金对接',
        '政策红利解读与补贴申报',
        'ESG / 双碳与产业可持续发展',
    ], start=1):
        add_para(doc, f'第 {i} 场:{t}', indent_first=False, left_indent=20)
    add_para(doc, '打包价:全年 6 场合计 24 万元(含场地、物料、嘉宾接待、传播投放等全部直接成本,不另行加收)。', indent_first=False)

    heading(doc, '3. 资源对接中介服务', level=2)
    add_para(doc, '服务内容:为甲方撮合外部产业资源,覆盖企业引进、投融资对接、政府合作、产业链上下游配对等。', indent_first=False)
    add_para(doc, '结算口径:按实际成功撮合落地的合作基数(元)核算中介服务费。', indent_first=False)
    add_para(doc, '收费费率:每万元合作基数收取 2 元中介服务费(即 0.02%)。', indent_first=False)
    add_para(doc, '结算节奏:按月或按项目结算,按实际撮合成果支付,不设保底亦不设上限。', indent_first=False)
    add_para(doc, '归属保护:乙方提交甲方备案的"对接资源清单"内的客户,自首次接触起 12 个月内成交的,均视为乙方撮合成果。', indent_first=False)

    # --- 三、项目报价 ---
    heading(doc, '三、项目报价', level=1)
    make_table(doc, ['项目', '明细', '金额(人民币)'], [
        ['服务中心建设', '基础规划、定位设计、挂牌落地、首批入驻企业邀约', '6 万元(一次性)'],
        ['年度活动执行', '6 场科技主题品牌活动(打包)',                  '24 万元(年度)'],
        ['中介服务费',   '按实际对接合作基数核算(2 元/万元基数)',          '详见第四部分计算逻辑'],
        ['基础合计',     '服务中心 + 年度活动',                          '30 万元 + 中介费'],
        ['示例总价',     '基础 + 2 亿元基数对应中介费',                   '34 万元(示例)'],
    ], widths_cm=[3.5, 8.5, 4.5])

    # --- 四、补充说明 ---
    heading(doc, '四、补充说明', level=1)

    heading(doc, '1. 中介服务费的计算逻辑', level=2)
    add_para(doc, '计算公式:中介服务费 = 实际对接合作基数(元) ÷ 10,000 × 2 元/万元', indent_first=False, bold=True)
    add_para(doc, '示例对照表(仅供参考):', indent_first=False)
    make_table(doc, ['实际对接基数', '中介服务费'], [
        ['5,000 万元', '1 万元'],
        ['2 亿元 (本方案示例)', '4 万元'],
        ['5 亿元', '10 万元'],
        ['10 亿元', '20 万元'],
        ['20 亿元', '40 万元'],
    ], widths_cm=[8.0, 8.0])
    add_para(doc, '以 2 亿元基数为标的:200,000,000 ÷ 10,000 × 2 = 40,000 元 = 4 万元。该费率(0.02%)远低于市场上常规招商代理费(通常 1-2%),是面向甲方作为长期战略合作伙伴的特别优惠定价。', indent_first=False)

    heading(doc, '2. 费用结算节奏', level=2)
    for line in [
        '· 服务中心建设费 6 万元:合同签订后 30 日内一次性支付。',
        '· 年度活动费 24 万元:',
        '    方式 A:按季度均分(每季度 6 万元,分 4 次支付);',
        '    方式 B:单场结算(4 万元/场,活动验收后 15 日内支付)。',
        '· 中介服务费:按实际撮合落地项目按月汇总结算,月度结算单经甲方书面确认后 15 日内支付。',
    ]:
        add_para(doc, line, indent_first=False)

    heading(doc, '3. 不重复计费原则', level=2)
    for line in [
        '(1) 服务中心建设期间产生的活动,若纳入年度 6 场活动范围内,则不另行加收活动费;',
        '(2) 中介服务费仅针对甲方原未自行对接的新增资源,避免与年度活动撮合成果重复结算;',
        '(3) 同一资源在 12 个月内重复撮合的,仅在首次成交时收取中介费。',
    ]:
        add_para(doc, line, indent_first=False)

    heading(doc, '4. 未达标退款机制', level=2)
    for line in [
        '(1) 服务中心未在 90 日内完成挂牌运营的,已收取的 6 万元按未完成天数比例退还;',
        '(2) 6 场年度活动若到场企业数未达 30 家/场,对应单场活动费扣减 50%;',
        '(3) 中介服务费完全与撮合成果挂钩,不存在未达标情形。',
    ]:
        add_para(doc, line, indent_first=False)

    heading(doc, '5. 变更与终止条款', level=2)
    for line in [
        '(1) 服务期内甲方调整活动频次或服务范围的,由双方协商签订补充协议;',
        '(2) 任何一方提前终止合作的,已发生服务费按履约比例结算;',
        '(3) 中介服务费在合作终止后 12 个月内最终成交的,仍按本方案标准支付(成果归属保护期)。',
    ]:
        add_para(doc, line, indent_first=False)

    heading(doc, '6. 保密与知识产权', level=2)
    for line in [
        '(1) 双方对在合作过程中知悉的对方商业秘密、客户名单、资源清单承担保密义务,保密期自本方案生效之日起至合作终止后 5 年;',
        '(2) 乙方为甲方设计的服务中心 VI、活动主题、传播物料的知识产权归属由双方在合同附件中另行约定。',
    ]:
        add_para(doc, line, indent_first=False)

    heading(doc, '7. 生效条件', level=2)
    add_para(doc, '本方案经双方法定代表人或授权代表签字盖章后生效;本方案的核心条款将作为正式合作协议的基础文本。', indent_first=False)

    # --- 签署页 ---
    doc.add_page_break()
    heading(doc, '签 署 页', level=1)
    sig = doc.add_table(rows=3, cols=2)
    sig.autofit = False
    for col in sig.columns:
        for cell in col.cells:
            cell.width = Cm(8)
    rows = [
        ('甲方:上海市科技企业联合会',          '乙方:__________(胡教授团队)'),
        ('法定代表人/授权代表:__________',     '法定代表人/授权代表:__________'),
        ('(盖章)　　　 日期:____ 年 ____ 月 ____ 日',
         '(盖章)　　　 日期:____ 年 ____ 月 ____ 日'),
    ]
    for i, (a, b) in enumerate(rows):
        sig.cell(i, 0).text = a
        sig.cell(i, 1).text = b
        for cell in sig.rows[i].cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for r in para.runs:
                    set_run(r, size=11)
                para.paragraph_format.space_after = Pt(8)

    doc.save(OUT)
    print(f'Wrote {OUT}')


if __name__ == '__main__':
    main()
