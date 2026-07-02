"""元谷项目媒体宣传方案 (媒体服务) — 发危建平总版.

基于媒体服务协议内容整理, 作为可选的媒体宣传服务模块:
  一、服务范围
  二、服务周期
  三、服务内容及报价
  四、传播策略说明 (中央/主流媒体引领 + 社交媒体双轮驱动)
  五、媒体资源清单
框架与外发口径一致:胡教授团队(代表复旦大学住房政策研究中心、上海市科技企业联合会), 第三方服务, 呈危建平总。
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name('元谷媒体宣传方案(媒体服务).docx')


def set_run(run, *, size=11, bold=False, color=None, font='微软雅黑'):
    run.font.name = font; run.font.size = Pt(size); run.font.bold = bold
    if color is not None: run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {}); rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font); rfonts.set(qn('w:ascii'), font); rfonts.set(qn('w:hAnsi'), font)


def add_para(doc, text, *, size=11, bold=False, color=None, align=None, indent_first=False, space_after=4, left_indent=0):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent: p.paragraph_format.left_indent = Pt(left_indent)
    if indent_first: p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text); set_run(run, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}; size = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 4)
    run = p.add_run(text); set_run(run, size=size, bold=True, color=RGBColor(0x0F, 0x24, 0x4E))
    return p


def make_table(doc, headers, rows, widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.autofit = False
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for cell in t.columns[i].cells: cell.width = Cm(w)
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
        for p in t.rows[0].cells[j].paragraphs:
            for r in p.runs: set_run(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.oxml import OxmlElement
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '0F244E'); t.rows[0].cells[j]._tc.get_or_add_tcPr().append(shd)
    for i, rrow in enumerate(rows, start=1):
        for j, v in enumerate(rrow):
            t.rows[i].cells[j].text = v
            for p in t.rows[i].cells[j].paragraphs:
                for r in p.runs: set_run(r, size=10.5)
            t.rows[i].cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return t


def main():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = '微软雅黑'; s.font.size = Pt(11)
    sec = doc.sections[0]; sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.0); sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # 封面
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(48)
    r = title.add_run('元谷项目媒体宣传方案'); set_run(r, size=28, bold=True, color=RGBColor(0x0F, 0x24, 0x4E))
    title.add_run('\n'); r2 = title.add_run('(媒体服务)'); set_run(r2, size=20, bold=True, color=RGBColor(0xF2, 0x7E, 0x2D))
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('中央级 + 全国性 + 上海主流媒体宣发  ·  社交媒体流量投放  ·  舆情管理'); set_run(r, size=13, color=RGBColor(0x66, 0x70, 0x86))
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(32)
    for line in ['出品方(乙方):胡教授团队 (代表复旦大学住房政策研究中心、上海市科技企业联合会)',
                 '呈:危建平总  ·  作为元谷项目招商运营的媒体宣传服务模块']:
        r = meta.add_run(line); set_run(r, size=11, color=RGBColor(0x55, 0x60, 0x7A)); meta.add_run('\n')
    note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(28)
    r = note.add_run('本方案为可选的媒体宣传服务模块, 与招商佣金、沙龙、挂牌等并行, 用于放大元谷项目品牌声量、赋能招商。')
    set_run(r, size=10, color=RGBColor(0x99, 0x33, 0x33))
    doc.add_page_break()

    # 引言
    heading(doc, '媒体宣传:为招商蓄势、为品牌立势', level=1)
    add_para(doc, '元谷项目定位『大零号湾文创融合核心区 + 上海市唯一科技时尚特色小镇』, IP + AI 双轨产业策源高地。招商成败很大程度取决于品牌声量与产业势能。乙方(胡教授团队, 代表复旦大学住房政策研究中心、上海市科技企业联合会)依托自身在中央级、全国性、上海主流媒体的资源, 为元谷项目提供专业媒体宣传服务, 以『主流媒体权威引领 + 社交媒体精准触达』双轮驱动, 放大项目声量、赋能招商转化。', indent_first=True)

    # 一、服务范围
    heading(doc, '一、服务范围', level=1)
    for line in [
        '1. 参与商议当月广告宣传重点, 向甲方推荐适合的媒体;',
        '2. 为甲方进行媒体策划、媒体公关、媒体联系等服务;',
        '3. 根据甲方宣传重点, 为甲方撰写相关宣传稿;',
        '4. 为甲方在主流媒体上联系、沟通并发表企业(项目)形象报道;',
        '5. 为甲方解答媒体传播中遇到的问题;',
        '6. 负责为甲方收集媒体市场信息以及行业媒体资料, 并对此提供点对点咨询;',
        '7. 受甲方委托, 策划组织媒体发布会、媒体沙龙、媒体行等;',
        '8. 合作结束后, 提供完整投放报告(含发布链接证明、曝光量证明和投放分析等)。',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    # 二、服务周期
    heading(doc, '二、服务周期', level=1)
    add_para(doc, '本方案约定服务周期自合同签订起至 2026 年 12 月 31 日。服务期满后, 双方可根据合作情况商谈续约事宜, 乙方享有优先续约权。', indent_first=True)
    add_para(doc, '与招商排期衔接:媒体宣传节奏建议紧扣两个硬节点 —— 2026/9/30(2,000㎡ 签约)与 2027/5/1(项目开业), 在关键节点前集中投放, 放大招商势能。', indent_first=True)

    # 三、服务内容及报价
    heading(doc, '三、服务内容及报价', level=1)
    add_para(doc, '根据双方商议, 乙方为甲方提供服务及价格如下(合作内容可根据甲方宣传需要调整):', indent_first=True)
    make_table(doc, ['服务项目', '内容说明', '数量', '费用'], [
        ['原创内容生产',
         '结合项目全新品牌定位与品牌内涵, 制定宣传报道计划, 提供品牌宣传稿件和市场营销活动稿件的原创采写服务',
         '3 篇', '折后总价 5 万元\n(媒体总曝光量预计不低于 150 万)'],
        ['中央级 / 全国性 / 上海区域主流媒体宣发',
         '根据项目宣传需求, 在中央级媒体、全国性财经/政经类媒体、上海区域主流媒体进行内容宣发',
         '10 篇', '打包报价(按媒体档位据实核算)'],
        ['全国主流社交媒体平台流量投放',
         '宣传平台:抖音、今日头条等;广告形式:信息流广告;投放区域:上海',
         '按投放量', '按投放量计\n(乙方负责精准流量投放;抖音短视频素材拍摄制作费用不含在内)'],
        ['营销策划与舆情管理咨询服务',
         '根据项目需求, 提供营销策划咨询服务, 并配合协调处理媒体舆情等',
         '不超 2 次', '含于媒体服务包'],
    ], widths_cm=[3.6, 6.6, 1.8, 4.0])
    add_para(doc, '说明:上表『原创内容生产 3 篇 / 5 万元(曝光≥150 万)』为明确报价;其余项目可按甲方宣传需要打包或据实核算。建议以媒体宣传服务包形式一次性确定当期投放范围与预算, 便于执行与结算。', indent_first=True, size=10.5, color=RGBColor(0x66, 0x70, 0x86))

    doc.add_page_break()

    # 四、传播策略说明
    heading(doc, '四、传播策略说明', level=1)
    heading(doc, '1. 主流媒体权威引领', level=2)
    add_para(doc, '在传播策略上, 首先应充分重视中央级媒体、区域主流媒体等的引领作用。主流媒体因其独特的信息采集能力、权威性以及专业性, 能够主动且直接地发声, 对网络信息环境产生显著影响, 进而决定舆情的发展方向。因此, 传播工作规划上, 首先将重点围绕全国性和区域性主流媒体而展开。', indent_first=True)
    add_para(doc, '中央级、全国性媒体包括但不限于:', indent_first=False, bold=True)
    add_para(doc, '央视新闻客户端、央广网、人民网、新华网、中新网、第一财经、21 世纪经济报道、澎湃新闻、界面新闻、中国日报、上海证券报、中国证券报 等。', indent_first=False, left_indent=15)

    heading(doc, '2. 社交媒体双轮驱动', level=2)
    add_para(doc, '在社交媒体投放策略上, 应充分利用今日头条的深度阅读场景与抖音的高频互动优势, 实现『广度覆盖 + 深度触达』的双轮驱动。通过信息流广告的精准算法推荐, 有效锁定目标受众, 提升品牌在核心人群中的可见度与影响力, 为后续业务(招商)转化奠定坚实的流量基础与认知铺垫, 确保营销预算的高效利用与投资回报最大化。', indent_first=True)

    heading(doc, '3. 与招商 / 活动的协同', level=2)
    for line in [
        '媒体宣传与 6 场产业沙龙、5·22 AI 商业化峰会、5 项产业牌照挂牌协同发声, 形成『活动造势 + 媒体放大』闭环;',
        '重大节点(9/30 签约节点、5/1 开业)前集中投放, 把品牌声量直接转化为招商线索;',
        '联动福布斯产业影响力奖、北欧创新国际会客厅等国际化资源, 提升项目在高端产业人群中的认知。',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    # 五、媒体资源清单
    heading(doc, '五、媒体资源清单(部分)', level=1)
    make_table(doc, ['层级', '媒体(部分)', '作用'], [
        ['中央级', '央视新闻客户端、央广网、人民网、新华网、中新网、中国日报', '权威背书 / 定调舆情'],
        ['全国财经政经', '第一财经、21 世纪经济报道、澎湃新闻、界面新闻、上海证券报、中国证券报', '产业价值传播 / 招商公信力'],
        ['上海区域主流', '解放日报、文汇报、上观新闻、闵行区人民政府官网等', '属地政府背书 / 区域影响'],
        ['社交媒体', '抖音(信息流)、今日头条(深度阅读)', '精准流量 / 目标受众触达'],
    ], widths_cm=[2.6, 8.0, 3.4])

    # 六、投放报告
    heading(doc, '六、投放报告与结算', level=1)
    for line in [
        '合作结束后(或按月/按节点), 乙方向甲方提供完整投放报告, 含:发布链接证明、曝光量证明、投放数据分析等;',
        '结算方式:媒体服务费建议按服务包一次性或按节点分期支付, 具体以双方约定为准;',
        '本媒体宣传服务为可选模块, 与招商佣金、沙龙执行费、挂牌费等相互独立, 不影响其他服务的计费。',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    add_para(doc, '', indent_first=False)
    add_para(doc, '本方案为商务沟通材料, 具体媒体投放范围、篇数与报价以双方签署的媒体服务协议为准。', indent_first=True, size=10, color=RGBColor(0x99, 0x33, 0x33))

    doc.save(OUT)
    print(f'Wrote {OUT}')


if __name__ == '__main__':
    main()
