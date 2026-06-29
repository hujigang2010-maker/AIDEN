"""元谷招商运营合作协议 — 第三方服务版 (发危建平总).

要点:
  - 乙方 = 胡教授团队 (代表复旦大学住房政策研究中心、上海市科技企业联合会), 第三方专业服务机构
  - 不涉及新设合资公司
  - 含排他性条款
  - 服务费用: 无月费 (不签对赌); 招商佣金; 沙龙费单独收取一次性; 挂牌 10 万/项可选挂牌前付清; 超额奖励适度正常付
  - 动态租金平衡 (招商期 1.5-1.8, 长期稳定 2.0-2.5, 保底 2.0 含物业)
  - 称谓统一 危建平总
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name('元谷招商运营合作协议(第三方服务版).docx')


def set_run(run, *, size=11, bold=False, color=None, font='微软雅黑'):
    run.font.name = font; run.font.size = Pt(size); run.font.bold = bold
    if color is not None: run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {}); rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font); rfonts.set(qn('w:ascii'), font); rfonts.set(qn('w:hAnsi'), font)


def add_para(doc, text, *, size=11, bold=False, color=None, align=None, indent_first=True, space_after=4, left_indent=0):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent: p.paragraph_format.left_indent = Pt(left_indent)
    if indent_first: p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text); set_run(run, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}; size = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 4)
    run = p.add_run(text); set_run(run, size=size, bold=True, color=RGBColor(0x0F, 0x24, 0x4E))
    return p


def article(doc, num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'第{num}条　{title}')
    set_run(run, size=13, bold=True, color=RGBColor(0x0F, 0x24, 0x4E))


def clause(doc, marker, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(20); p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{marker}　{text}'); set_run(run, size=11)


def make_table(doc, headers, rows, widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.autofit = False
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for cell in t.columns[i].cells: cell.width = Cm(w)
    hdr = t.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = h
        for p in hdr[j].paragraphs:
            for r in p.runs: set_run(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.oxml import OxmlElement
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '0F244E'); hdr[j]._tc.get_or_add_tcPr().append(shd)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.rows[i].cells[j].text = v
            for p in t.rows[i].cells[j].paragraphs:
                for r in p.runs: set_run(r, size=10.5)
            t.rows[i].cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return t


def main():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = '微软雅黑'; s.font.size = Pt(11)
    sec = doc.sections[0]; sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.0); sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # 封面
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(46)
    r = title.add_run('元谷项目 4#+5# 楼约 2 万方产业研发办公')
    set_run(r, size=20, bold=True, color=RGBColor(0x0F, 0x24, 0x4E))
    title.add_run('\n'); r2 = title.add_run('招 商 运 营 服 务 合 作 协 议')
    set_run(r2, size=22, bold=True, color=RGBColor(0x0F, 0x24, 0x4E))
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('(第三方专业服务版 · 草案)'); set_run(r, size=13, color=RGBColor(0x66, 0x70, 0x86))
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(36)
    for line in ['委托方(甲方):元谷项目方 (由危建平总主导)',
                 '服务方(乙方):胡教授团队 (代表复旦大学住房政策研究中心、上海市科技企业联合会)',
                 '签署日期:__________ 年 ____ 月 ____ 日']:
        r = meta.add_run(line); set_run(r, size=12); meta.add_run('\n')
    note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(40)
    r = note.add_run('本协议为商务草案, 不构成最终法律约束力, 最终以双方签署的正式文本为准。')
    set_run(r, size=10, color=RGBColor(0x99, 0x33, 0x33))
    doc.add_page_break()

    # 缔约方
    heading(doc, '缔约方')
    add_para(doc, '甲方(委托方):__________ (元谷项目方, 由危建平总主导决策, 以下简称"甲方")', indent_first=False)
    add_para(doc, '    住所地:__________ ; 统一社会信用代码:__________ ; 授权代表:危建平', indent_first=False)
    add_para(doc, '乙方(服务方):胡教授团队 (以下简称"乙方")', indent_first=False)
    add_para(doc, '    乙方为第三方专业招商运营服务机构, 代表复旦大学住房政策研究中心、上海市科技企业联合会, 提供产业招商运营服务。', indent_first=False)
    add_para(doc, '    住所地:__________ ; 统一社会信用代码:__________ ; 授权代表:胡教授', indent_first=False)

    add_para(doc, '鉴于:', bold=True, indent_first=False)
    clause(doc, '(1)', '甲方为元谷项目(上海市闵行区元江路-剑川路, 大零号湾文创融合核心区)的项目方, 拟就 4# 楼 5F 及以上 + 5# 楼 5F 及以上潮玩产业集群共约 2 万㎡产业研发办公空间(以下简称"项目范围")开展招商运营;')
    clause(doc, '(2)', '乙方作为第三方专业招商运营服务机构, 在产业招商、政府政策对接、品牌活动、产业资源导入等方面具备专业能力与稀缺资源(包括但不限于仲量联行爬楼大数据、产业基金、腾讯算力、北欧创新国际会客厅、福布斯产业影响力奖、复旦大学住房政策研究中心、上海市科技企业联合会等);')
    clause(doc, '(3)', '甲方拟委托乙方以第三方专业服务方式, 独家承担项目范围的招商运营服务, 双方本着平等互利原则达成如下协议:')

    # 第一条 服务性质
    article(doc, '一', '服务性质与合作方式')
    clause(doc, '1.1', '本协议项下合作为甲方委托乙方提供第三方专业招商运营服务的服务合作关系, 双方不因本协议设立任何合资公司、合伙企业或其他共同经营实体。')
    clause(doc, '1.2', '乙方以自身名义及专业团队向甲方提供服务, 自主承担其团队的组织、管理与成本, 服务成果归于本协议约定的招商运营目标。')
    clause(doc, '1.3', '乙方提供服务时, 可使用其代表的复旦大学住房政策研究中心、上海市科技企业联合会的相关资源与背书, 用于本项目招商与政府对接。')

    # 第二条 项目范围
    article(doc, '二', '项目范围')
    clause(doc, '2.1', '项目范围:元谷项目 4# 楼 5F 及以上潮玩产业集群(约 1 万㎡)与 5# 楼 5F 及以上潮玩产业集群(约 1 万㎡), 合计约 2 万㎡产业研发办公空间。')
    clause(doc, '2.2', '业务定位:双方一致将项目范围打造为 IP + AI 双轨产业研发园区, 其中 4# 楼 5F+ 以 AI 为主轴, 5# 楼 5F+ 以 IP 为主轴。')
    clause(doc, '2.3', '本协议不涉及元谷项目商业部分(约 5.2 万㎡)及其他楼栋, 该部分由甲方另行安排。')

    # 第三条 服务范围
    article(doc, '三', '乙方服务范围')
    clause(doc, '3.1', '产业招商策划:招商定位、客户画像、招商方案、招商物料(招商手册、话术、政策汇编等)的策划与提供。')
    clause(doc, '3.2', '招商执行:依托仲量联行爬楼大数据等资源, 形成精准目标客户清单并开展招商, 推动租赁合同签订。')
    clause(doc, '3.3', '政府与政策对接:对接闵行区科委/商务委/街道及上海市相关部门, 协助申请高新认定、专精特新、创新券、文创扶持、特色小镇专项等政策。')
    clause(doc, '3.4', '品牌与活动:策划承办产业沙龙、产业牌照挂牌、5·22 AI 商业化峰会借势等品牌活动, 提升项目声量与招商势能。')
    clause(doc, '3.5', '产业资源导入:导入产业基金、腾讯算力补贴、北欧创新国际会客厅、福布斯产业影响力奖等资源, 赋能项目招商。')

    # 第四条 租金策略 (动态平衡)
    article(doc, '四', '租金策略(动态租金平衡)')
    clause(doc, '4.1', '双方一致同意, 项目范围采取动态租金平衡策略, 以"先低价快速招满、后逐步抬升租金"为原则:')
    clause(doc, '      (a)', '招商期(满租前):项目范围平均租金控制在 1.5 - 1.8 元/㎡/天, 以快速招满、形成产业氛围为目标;')
    clause(doc, '      (b)', '长期稳定(满租后):随产业生态成熟逐步抬升至 2.0 - 2.5 元/㎡/天;')
    clause(doc, '      (c)', '保底锚点:以甲方(危建平总)已确定的 2.0 元/㎡/天(含物业费)为长期稳定的保底锚点。')
    clause(doc, '4.2', '具体租赁定价由乙方在上述区间内根据客户面积、租期、产业属性灵活确定, 报甲方备案。')
    clause(doc, '4.3', '抬租机制:满租后, 乙方协助甲方在租约续签或新签时逐步将租金推向长期稳定区间, 实现项目方租金收益的持续增长。')

    # 第五条 服务费用 (核心, 无月费)
    article(doc, '五', '服务费用(无月费 · 按成果付费)')
    clause(doc, '5.1', '双方确认:本协议不设固定月费, 乙方不要求任何形式的保底或对赌;乙方完全按招商成果及约定的专项服务收费。')
    clause(doc, '5.2', '招商佣金(核心收入):甲方就乙方促成的每一份项目范围内租赁合同(新签或续签), 按下列阶梯向乙方支付招商佣金:')
    clause(doc, '      (a)', '≤ 2,000㎡ 的租户:佣金 = 实际成交年租金 × 1.5 个月;')
    clause(doc, '      (b)', '2,001 - 5,000㎡ 的租户:佣金 = 实际成交年租金 × 1.75 个月;')
    clause(doc, '      (c)', '> 5,000㎡ 的租户:佣金 = 实际成交年租金 × 2.0 个月。')
    clause(doc, '      (d)', '招商佣金于租赁合同签订且租户起租之日起 30 日内由甲方支付;实际成交年租金以租赁合同载明金额为准(招商期低价租金据实计算)。')
    clause(doc, '5.3', '沙龙执行费:乙方承办的产业沙龙(全年 6 场, 每场到场目标产业客户不少于 30 家), 执行费由乙方单独收取, 不与任何第三方分润。')
    clause(doc, '      (a)', '收费标准:6 场打包人民币 30 万元(即 5 万元/场);')
    clause(doc, '      (b)', '★ 支付方式:建议于协议签署并启动后一次性付清, 以减少后续财务往来;')
    clause(doc, '      (c)', '备选方式:如甲方需要, 可分两次支付(启动时付 50%、第三场沙龙完成后付 50%)。')
    clause(doc, '5.4', '挂牌费(可选):乙方可协助甲方促成下列产业牌照挂牌, 挂牌费为人民币 10 万元/项, 由甲方按需选择挂牌项目:')
    clause(doc, '      (a)', '可选挂牌:① AI 潮玩产业基地(中国动漫集团);② 潮玩次元商业专委会(中国百货商业协会);③ 复旦大学住房政策研究中心·元谷分中心;④ 上海市科技企业联合会·元谷产业基地;⑤ 福布斯产业影响力奖·元谷专场;')
    clause(doc, '      (b)', '★ 支付方式:甲方选定挂牌项目后, 于该项目正式挂牌前一次性付清对应挂牌费(最多 5 项, 合计上限 50 万元)。')
    clause(doc, '5.5', '超额奖励(适度):项目范围招商出租率达到 90% 及以上时, 甲方向乙方支付适度超额奖励人民币 20 万元;')
    clause(doc, '      (a)', '支付方式:按正常节奏分期支付(可随招商佣金尾款分期结算), 不要求一次性支付。')
    clause(doc, '5.6', '除上述费用外, 乙方不再向甲方收取月费、管理费、咨询费等其他固定费用。')

    # 第六条 排他性
    article(doc, '六', '排他性与独家授权')
    clause(doc, '6.1', '★ 独家招商运营权:在本协议有效期内, 甲方授予乙方对项目范围(4# 楼 5F+ + 5# 楼 5F+ 约 2 万㎡)的独家招商运营权;甲方不得就项目范围另行委托任何第三方招商运营机构, 亦不得自行绕开乙方与乙方提报客户名单内的客户签约。')
    clause(doc, '6.2', '客户归属保护:乙方向甲方书面提报并备案的"招商客户名单"内的客户, 自首次接触起 24 个月内成交的, 均视为乙方招商成果, 甲方应按第五条向乙方支付招商佣金;甲方绕开签约的, 按对应佣金标准的 1.5 倍向乙方赔付。')
    clause(doc, '6.3', '区域竞业:在本协议有效期内, 就甲方在大零号湾区域内的其他同类产业研发招商需求, 甲方应优先以书面形式询价于乙方, 乙方在 30 日内未明确接受的, 甲方有权另行安排。')
    clause(doc, '6.4', '资源排他:乙方为本项目导入的仲量联行爬楼大数据、产业基金、腾讯算力、福布斯、北欧会客厅等资源, 在协议期内甲方不得绕开乙方直接对接用于项目范围外的其他用途。')

    # 第七条 双方权利义务
    article(doc, '七', '双方权利义务')
    clause(doc, '7.1', '甲方义务:')
    clause(doc, '      (a)', '在协议生效后 30 日内向乙方开放项目范围招商所需的楼盘资料、平面图、价格政策及客户对接通道;')
    clause(doc, '      (b)', '为乙方派驻团队提供项目范围内必要的办公与接待条件;')
    clause(doc, '      (c)', '按本协议第五条约定及时支付各项服务费用;')
    clause(doc, '      (d)', '配合乙方完成产业牌照挂牌、品牌活动、政府对接及联合宣传。')
    clause(doc, '7.2', '乙方义务:')
    clause(doc, '      (a)', '以专业能力开展招商, 在协议约定节点推动签约;')
    clause(doc, '      (b)', '将仲量联行爬楼大数据等资源用于本项目精准招商, 形成不少于 200 家目标客户清单;')
    clause(doc, '      (c)', '承办不少于 6 场产业沙龙(每场到场目标产业客户不少于 30 家);')
    clause(doc, '      (d)', '协助甲方完成选定的产业牌照挂牌;')
    clause(doc, '      (e)', '执行动态租金平衡策略, 兼顾招商速度与项目方长期租金收益。')

    # 第八条 工作目标 (非对赌)
    article(doc, '八', '工作目标(参考性, 非对赌)')
    clause(doc, '8.1', '双方以下列招商节点作为共同努力的工作目标(因乙方不收取月费、不设保底, 下列目标为参考性目标, 不构成对乙方的对赌或扣罚条款):')
    clause(doc, '      (a)', '2026 年 9 月 30 日前:项目范围累计签约不少于 2,000㎡(可含直播基地、共享设计中心等业态);')
    clause(doc, '      (b)', '2027 年 5 月 1 日前:配合项目开业, 项目范围累计签约率达到 50% 及以上;')
    clause(doc, '      (c)', '协议期内:推动项目范围出租率达到 90% 及以上。')
    clause(doc, '8.2', '上述目标的达成需甲方按约履行配合义务;因甲方原因(如楼宇交付延期、价格政策未到位等)导致目标延后的, 相应顺延。')

    # 第九条 协议期限
    article(doc, '九', '协议期限')
    clause(doc, '9.1', '本协议有效期为 24 个月, 自双方签署之日起算。')
    clause(doc, '9.2', '协议届满前 90 日, 双方可协商续期;若协议期内项目范围出租率达到 90% 以上, 乙方在同等条件下享有续约优先权。')

    # 第十条 知识产权与保密
    article(doc, '十', '知识产权与保密')
    clause(doc, '10.1', '乙方为本项目提供的招商物料、策略方案、活动方案的知识产权归属由双方另行约定;乙方导入的各项资源(数据、基金、牌照渠道、活动 IP)的所有权仍归原权利人, 甲方仅在本项目范围内使用。')
    clause(doc, '10.2', '双方对在合作过程中知悉的对方商业秘密、客户名单、价格政策、财务信息承担保密义务, 期限为协议存续期间及终止后 5 年。')

    # 第十一条 违约与解除
    article(doc, '十一', '违约与解除')
    clause(doc, '11.1', '甲方未按约支付招商佣金、沙龙费、挂牌费等费用的, 每逾期一日按应付金额的 0.05% 支付违约金;经书面催告 30 日仍未支付的, 乙方有权解除协议并要求甲方支付已发生服务对应的全部费用。')
    clause(doc, '11.2', '甲方违反第六条排他性约定的, 应按相应佣金标准的 1.5 倍向乙方赔付。')
    clause(doc, '11.3', '因不可抗力或政策重大变化导致项目范围无法开展招商的, 双方可协商解除, 已发生服务费用据实结算。')

    # 第十二条 争议解决
    article(doc, '十二', '法律适用与争议解决')
    clause(doc, '12.1', '本协议适用中华人民共和国法律。')
    clause(doc, '12.2', '凡因本协议引起的争议, 双方应友好协商;协商不成的, 提交上海国际经济贸易仲裁委员会(上海国际仲裁中心)按其届时有效的仲裁规则在上海仲裁。')

    # 第十三条 一般条款
    article(doc, '十三', '一般条款')
    clause(doc, '13.1', '本协议自双方授权代表签字并加盖公章之日起生效。')
    clause(doc, '13.2', '本协议一式四份, 双方各执两份, 具有同等法律效力。')
    clause(doc, '13.3', '本协议附件与正文具有同等法律效力;未尽事宜由双方另行书面约定。')

    # 签署页
    doc.add_page_break()
    heading(doc, '签 署 页', level=1)
    sig = doc.add_table(rows=3, cols=2); sig.autofit = False
    for col in sig.columns:
        for cell in col.cells: cell.width = Cm(8)
    rows = [
        ('甲方:__________ (元谷项目方)', '乙方:胡教授团队'),
        ('授权代表:危建平', '授权代表:胡教授'),
        ('(盖章)　　　 日期:____ 年 ____ 月 ____ 日', '(盖章)　　　 日期:____ 年 ____ 月 ____ 日'),
    ]
    for i, (a, b) in enumerate(rows):
        sig.cell(i, 0).text = a; sig.cell(i, 1).text = b
        for cell in sig.rows[i].cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for r in para.runs: set_run(r, size=11)
                para.paragraph_format.space_after = Pt(8)

    # 附件
    doc.add_page_break()
    heading(doc, '附件:服务费用速查表', level=1)
    make_table(doc, ['费用类别', '标准', '支付方式'], [
        ['月费', '★ 无 (不设固定月费, 不签对赌)', '—'],
        ['招商佣金(核心)', '实际成交年租金的 1.5 / 1.75 / 2.0 个月(按面积)', '起租后 30 日内'],
        ['沙龙执行费', '6 场打包 30 万(乙方单独收取, 不分润)', '一次性付清(或分两次)'],
        ['挂牌费(可选)', '10 万元/项 × 选定项数(最多 5 项 = 50 万)', '挂牌前一次性付清'],
        ['超额奖励(适度)', '出租率 ≥ 90% → 20 万', '正常分期支付'],
    ], widths_cm=[3.5, 8.0, 4.5])

    heading(doc, '附件:租金动态平衡速查', level=1)
    make_table(doc, ['阶段', '租金区间', '说明'], [
        ['招商期(满租前)', '1.5 - 1.8 元/㎡/天', '先低价快速招满'],
        ['长期稳定(满租后)', '2.0 - 2.5 元/㎡/天', '逐步抬升'],
        ['保底锚点', '2.0 元/㎡/天(含物业费)', '危建平总已确定'],
    ], widths_cm=[4.0, 4.5, 7.5])

    doc.save(OUT)
    print(f'Wrote {OUT}')


if __name__ == '__main__':
    main()
