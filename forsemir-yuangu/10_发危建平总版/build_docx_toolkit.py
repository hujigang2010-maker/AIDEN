"""元谷招商材料工具包 — 精选脱敏版 (发危建平总).

精选, 不大而全:
  一、项目一页通
  二、招租亮点与租金动态平衡
  三、五项产业牌照 (可选)
  四、6 场产业沙龙
  五、政策亮点 (闵行 + 上海市)
  六、精选招商问答 (30 问)
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name('元谷招商材料工具包(精选脱敏版).docx')


def set_run(run, *, size=11, bold=False, color=None, font='微软雅黑'):
    run.font.name = font; run.font.size = Pt(size); run.font.bold = bold
    if color is not None: run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {}); rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font); rfonts.set(qn('w:ascii'), font); rfonts.set(qn('w:hAnsi'), font)


def add_para(doc, text, *, size=11, bold=False, color=None, align=None, indent_first=False, space_after=4, left_indent=0):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent: p.paragraph_format.left_indent = Pt(left_indent)
    if indent_first: p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text); set_run(run, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}; size = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 4)
    run = p.add_run(text); set_run(run, size=size, bold=True, color=RGBColor(0x0F, 0x24, 0x4E))
    return p


def make_table(doc, headers, rows, widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.autofit = False
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for cell in t.columns[i].cells: cell.width = Cm(w)
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
        for p in t.rows[0].cells[j].paragraphs:
            for r in p.runs: set_run(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.oxml import OxmlElement
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '0F244E'); t.rows[0].cells[j]._tc.get_or_add_tcPr().append(shd)
    for i, rrow in enumerate(rows, start=1):
        for j, v in enumerate(rrow):
            t.rows[i].cells[j].text = v
            for p in t.rows[i].cells[j].paragraphs:
                for r in p.runs: set_run(r, size=10.5)
            t.rows[i].cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return t


def qa(doc, q, a):
    pq = doc.add_paragraph(); pq.paragraph_format.space_after = Pt(2)
    rq = pq.add_run('Q  ' + q); set_run(rq, size=11, bold=True, color=RGBColor(0xF2, 0x7E, 0x2D))
    pa = doc.add_paragraph(); pa.paragraph_format.space_after = Pt(8); pa.paragraph_format.left_indent = Pt(15)
    ra = pa.add_run('A  ' + a); set_run(ra, size=11)


def main():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = '微软雅黑'; s.font.size = Pt(11)
    sec = doc.sections[0]; sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.0); sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # 封面
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(48)
    r = title.add_run('元谷 2 万方招商材料工具包'); set_run(r, size=28, bold=True, color=RGBColor(0x0F, 0x24, 0x4E))
    title.add_run('\n'); r2 = title.add_run('精 选 脱 敏 版'); set_run(r2, size=22, bold=True, color=RGBColor(0xF2, 0x7E, 0x2D))
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('项目一页通 · 招租亮点 · 产业牌照 · 产业沙龙 · 政策亮点 · 媒体宣传 · 精选问答'); set_run(r, size=13, color=RGBColor(0x66, 0x70, 0x86))
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(32)
    for line in ['出品方:胡教授团队 (代表复旦大学住房政策研究中心、上海市科技企业联合会)',
                 '呈:危建平总  ·  本材料已脱敏, 用于商务汇报与招商沟通']:
        r = meta.add_run(line); set_run(r, size=11, color=RGBColor(0x55, 0x60, 0x7A)); meta.add_run('\n')
    doc.add_page_break()

    # 一、项目一页通
    heading(doc, '一、项目一页通', level=1)
    make_table(doc, ['项', '内容'], [
        ['项目', '森马 (上海) 国际运营中心 · 元谷项目'],
        ['位置', '上海市闵行区元江路-剑川路, 大零号湾文创融合核心区'],
        ['身份', '大零号湾核心区 + 上海市唯一科技时尚特色小镇'],
        ['总建面', '22 万㎡ (1#-6# 共 6 栋)'],
        ['本期招商', '4# 楼 5F+ + 5# 楼 5F+ 共约 2 万㎡产业研发办公'],
        ['产业定位', 'IP + AI 双轨 (4# 楼 AI 主轴 / 5# 楼 IP 主轴)'],
        ['交通', '15 号线元江路站 TOD, 单日客流 5-7 万人次'],
        ['覆盖', '15 分钟车行覆盖 24 万居住 + 12 万产业办公人口'],
        ['开业', '2027/5/1'],
        ['服务方', '胡教授团队 (第三方专业招商运营服务机构)'],
    ], widths_cm=[3.5, 12.5])

    # 二、招租亮点与租金动态平衡
    heading(doc, '二、招租亮点与租金动态平衡', level=1)
    heading(doc, '2.1 楼宇与定位', level=2)
    make_table(doc, ['楼栋', '面积', '主轴', '同栋协同'], [
        ['4# 楼 5F+', '约 1 万㎡', 'AI 主轴 (国际创意层)', '4F 直播中心 / 5F AI 共享设计中心 / DIY 中心'],
        ['5# 楼 5F+', '约 1 万㎡', 'IP 主轴 (产业总部层)', '5F 潮玩产业展厅 / 1-4F 动漫书店'],
    ], widths_cm=[2.8, 2.5, 4.0, 6.7])

    heading(doc, '2.2 租金动态平衡 (核心)', level=2)
    make_table(doc, ['阶段', '租金区间', '说明'], [
        ['招商期 (满租前)', '1.5 - 1.8 元/㎡/天', '先低价快速招满, 形成产业氛围'],
        ['长期稳定 (满租后)', '2.0 - 2.5 元/㎡/天', '满租后逐步抬升'],
        ['保底锚点', '2.0 元/㎡/天 (含物业费)', '危建平总已确定, 长期底线'],
    ], widths_cm=[4.0, 4.5, 7.5])
    add_para(doc, '策略逻辑:先以 1.5-1.8 元低价快速招满, 再借满租势能与产业生态逐步抬升至 2.0-2.5 元, 长期锚定保底 2.0 元(含物业), 实现项目方租金动态增收。', indent_first=True)

    heading(doc, '2.3 招商亮点 (我方服务能力)', level=2)
    for line in [
        '★ 仲量联行爬楼大数据:200+ 家精准目标客户清单, 转化率 +30% (我方核心优势)',
        '腾讯算力补贴:新注册 AI 公司 3 个月合同免费 / 算力 85 折',
        '产业基金:返投落地, 资本驱动招商',
        '北欧创新国际会客厅 + 福布斯产业影响力奖:国际化品牌势能',
        '复旦大学住房政策研究中心 + 上海市科技企业联合会:学术背书 + 政府对接',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    doc.add_page_break()

    # 三、五项产业牌照
    heading(doc, '三、5 项产业牌照 (可选 · 10 万/项)', level=1)
    make_table(doc, ['序号', '挂牌名称', '出牌方', '对入驻企业的价值'], [
        ['①', 'AI 潮玩产业基地', '中国动漫集团', '国家级 AI+IP 产业认证 / 政策对接'],
        ['②', '潮玩次元商业专委会', '中国百货商业协会', '潮玩零售生态 / 政府对话渠道'],
        ['③', '复旦大学住房政策研究中心·元谷分中心', '复旦大学', '学术背书 / 政策研究 / 高净值人脉'],
        ['④', '上海市科技企业联合会·元谷产业基地', '上海市科企联', '上海科技企业生态 / 政府补贴申报'],
        ['⑤', '福布斯产业影响力奖·元谷专场', '福布斯', '国际品牌势能 / 年度评选 IP'],
    ], widths_cm=[1.2, 6.0, 3.3, 5.5])
    add_para(doc, '说明:挂牌费 10 万元/项, 由项目方按需选择, 选定后挂牌前一次性付清(最多 5 项)。', indent_first=True)

    # 四、6 场产业沙龙
    heading(doc, '四、6 场产业沙龙 (每场 ≥ 30 个目标产业客户)', level=1)
    make_table(doc, ['#', '时间', '主题', '联办方'], [
        ['#1', '第 1 月', 'AI + 潮玩 (借势 5·22 峰会)', '中动漫 + 腾讯'],
        ['#2', '第 3 月', '潮玩出海', '北欧创新国际会客厅 + 福布斯'],
        ['#3', '第 5 月', '投融资路演', '产业基金 + 银行 + 券商'],
        ['#4', '第 7 月', '设计与创意', '上海交大设计学院 + 上海市科企联'],
        ['#5', '第 9 月', '内容 IP 与 Z 世代', '中百协潮玩次元专委 + 中动漫'],
        ['#6', '第 11 月', '政策补贴与小镇', '闵行科协 + 复旦住房政策研究中心'],
    ], widths_cm=[1.2, 2.8, 5.5, 6.5])
    add_para(doc, '沙龙执行费由我方单独收取(不与他方分润), 6 场打包一次性付清。每场触达 ≥30 家产业客户 + ≥100 万次媒体曝光。', indent_first=True)

    # 五、政策亮点
    heading(doc, '五、政策亮点 (闵行 + 上海市)', level=1)
    make_table(doc, ['政策', '出台单位', '扶持力度'], [
        ['大零号湾建设专项', '闵行区政府', '最高 1,000 万'],
        ['科技时尚特色小镇专项', '闵行区政府', '200-500 万 (元谷独占)'],
        ['闵行区 AI 行动方案', '闵行区科委', '100-500 万'],
        ['闵行区文创产业扶持', '闵行区文广旅', '最高 200 万'],
        ['闵行区科技创新券', '闵行区科委', '10-50 万/年'],
        ['上海市 AI 大模型专项', '上海市经信委', '最高 1,000 万'],
        ['高新技术企业认定', '上海市科技局', '所得税 15% (vs 25%)'],
        ['上海市数字内容产业专项', '上海市文旅局', '100-500 万'],
        ['腾讯算力补贴', '云厂商合作', '3 月合同免费 / 算力 85 折'],
    ], widths_cm=[5.5, 4.0, 6.5])
    add_para(doc, '我方代表复旦大学住房政策研究中心、上海市科技企业联合会, 可协助入驻企业一站式代办上述政策申报, 加速 30%。', indent_first=True)

    doc.add_page_break()

    # 六、媒体宣传 (可选服务模块)
    heading(doc, '六、媒体宣传服务(可选模块)', level=1)
    add_para(doc, '以『主流媒体权威引领 + 社交媒体精准触达』双轮驱动, 放大元谷项目品牌声量、赋能招商;详见《元谷媒体宣传方案(媒体服务)》。', indent_first=True)
    heading(doc, '6.1 服务内容与报价', level=2)
    make_table(doc, ['服务项目', '数量', '费用 / 说明'], [
        ['原创内容生产(品牌+营销稿原创采写)', '3 篇', '折后 5 万元(媒体总曝光量 ≥ 150 万)'],
        ['中央级/全国/上海主流媒体宣发', '10 篇', '打包报价(按媒体档位据实核算)'],
        ['社交媒体流量投放(抖音/今日头条信息流, 上海)', '按投放量', '按投放量计(短视频拍摄制作费不含)'],
        ['营销策划与舆情管理咨询', '不超 2 次', '含于媒体服务包'],
    ], widths_cm=[6.0, 2.0, 8.0])
    heading(doc, '6.2 传播策略与媒体资源', level=2)
    for line in [
        '中央级/全国性媒体:央视新闻、央广网、人民网、新华网、中新网、第一财经、21 世纪经济报道、澎湃、界面、中国日报、上海证券报、中国证券报 等;',
        '社交媒体:抖音(信息流)+ 今日头条(深度阅读), 精准算法锁定目标受众;',
        '服务周期:自签订起至 2026/12/31, 乙方享优先续约权;合作结束提供完整投放报告(链接+曝光+分析);',
        '与 6 场沙龙 + 5·22 峰会 + 5 项挂牌协同, 在 9/30、5/1 关键节点前集中投放, 把声量转化为招商线索。',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    doc.add_page_break()

    # 七、精选招商问答 30 问
    heading(doc, '七、精选招商问答 (30 问)', level=1)
    qas = [
        ('元谷项目是什么?', '森马集团出品的科技+时尚产业综合体, 总建面 22 万㎡, 位于上海闵行大零号湾核心区。本期聚焦 4#+5# 楼 5F+ 共约 2 万㎡产业研发办公, 定位 IP + AI 双轨产业园区。'),
        ('在什么位置?', '上海市闵行区元江路-剑川路, 15 号线元江路站 TOD 上盖。距紫竹高新区 5km、虹桥机场 17km、浦东机场 36km。'),
        ('什么时候开业?', '2027 年 5 月 1 日, 当前招商已启动。'),
        ('为什么选这个位置?', '大零号湾核心区(闵行五大中心之一) + 上海唯一科技时尚特色小镇, 政策含金量高, TOD 上盖客流大。'),
        ('4# 楼和 5# 楼有什么区别?', '4# 楼 5F+ 偏 AI(国际创意层, 同栋有直播中心 + AI 共享设计中心);5# 楼 5F+ 偏 IP(产业总部层, 同栋有潮玩展厅 + 动漫书店)。'),
        ('面积可以分割吗?', '可以, 200㎡ - 5,000㎡ 灵活分割, 头部客户可整层整栋。'),
        ('租金是多少?', '招商期平均 1.5-1.8 元/㎡/天(先低价招满), 长期稳定 2.0-2.5 元/㎡/天;保底 2.0 元含物业。'),
        ('为什么招商期租金低?', '动态租金平衡策略:先以低价快速招满、形成产业氛围, 再借满租势能逐步抬升租金, 对早期入驻客户是红利。'),
        ('物业费另收吗?', '保底锚点 2.0 元已含物业费;招商期低价为净租金或含物业, 以租赁合同为准。'),
        ('押金和预付多少?', '押金 = 月租金 × 2 个月;预付租金 3 个月。大客户可议。'),
        ('租期多久?', '最短 1 年, 推荐 3 年(可 3+2 续约), 头部客户可议 5 年。'),
        ('有免租期吗?', '有, 按面积阶梯 15 天 - 3 个月;装修期不计租金。'),
        ('有装修补贴吗?', '大客户专享:≥2,000㎡ 享 100-300 元/㎡ 装修补贴。'),
        ('入驻有什么政府补贴?', '高新认定(15% 所得税)、专精特新、创新券、闵行专项、AI 大模型专项(最高 1,000 万)、文创扶持等, 我方一站式代办。'),
        ('腾讯算力补贴是真的吗?', '是, 新注册 AI 公司 3 个月合同免费或算力 85 折, 我方一站式办理。'),
        ('有产业牌照吗?', '5 项可选:AI 潮玩产业基地、潮玩次元商业专委会、复旦住房政策研究中心元谷分中心、上海市科企联元谷产业基地、福布斯产业影响力奖。'),
        ('入驻能用产业牌照背书吗?', '可以, 符合条件的入驻企业可获牌照成员单位背书, 用于名片/官网/媒体宣传。'),
        ('每年有什么活动?', '6 场产业沙龙(每场 ≥30 家客户)+ 潮玩设计大赛 + 福布斯榜单 + 北欧外事接待 + 5·22 AI 峰会借势。'),
        ('我能在活动上发言吗?', '可以, 入驻企业可申请主题演讲席 / 圆桌嘉宾席 / 路演席。'),
        ('有配套设施吗?', '共享直播中心、AI 共享设计中心、AI 打样 DIY 中心、潮玩选品中心、潮玩艺术中心、潮玩产业展厅。'),
        ('科技企业服务中心做什么?', '提供注册落户、财税法、知识产权、政府补贴申报、人才签证、投融资、品牌公关、数字化工具、培训认证等服务。'),
        ('注册地不在上海能享受补贴吗?', '需在闵行区(元谷)设立子公司, 子公司单独享受政策, 我方可代办注册。'),
        ('员工落户能支持吗?', '可以, 提供居住证积分、留学生/海归落户、外籍签证代办;员工可申请闵行人才公寓(房租减免 30-50%)。'),
        ('AI 大模型公司适合吗?', '非常适合:腾讯算力补贴 + 上海 AI 大模型专项(最高 1,000 万)+ AI 共享设计中心 + AI 潮玩产业基地牌照。'),
        ('IP 公司怎么对接潮玩零售?', '一站式:IP 创意(展厅)→ 打样(DIY 中心)→ 制造(物流中心)→ 销售(选品中心)→ 直播(直播中心)。'),
        ('元谷 vs 周边园区?', '与零号湾/紫竹同价(2.0-2.5 元), 但 IP+AI 双轨 + 5 项挂牌 + TOD 上盖 + 国际化资源, 综合性价比更高;比张江/漕河泾顶级园区便宜 30-50%。'),
        ('元谷 vs 杨浦 AI+IP 园区?', '杨浦 1 万方已验证打法可行;元谷 2 万方为升级版, 加上闵行政策红利 + TOD + 国际化背书, 优势更强。'),
        ('签约流程是什么?', '看楼 → 商务沟通 → 条款谈判 → 签约 → 起租装修, 整体 1.5-2 个月可入驻。'),
        ('入驻后有人服务吗?', '有, 配专属客户经理, 负责政策对接 / 配套使用 / 活动邀请 / 投融资对接, 全程伴随。'),
        ('谁负责招商?', '由胡教授团队(第三方专业招商运营服务机构, 代表复旦住房政策研究中心、上海市科企联)全权负责产业招商。'),
    ]
    for i, (q, a) in enumerate(qas, start=1):
        qa(doc, f'{i:02d}. {q}', a)

    doc.save(OUT)
    print(f'Wrote {OUT}')


if __name__ == '__main__':
    main()
