"""闵行及周边区 IP+AI+文创 产业政策汇编 + 案例分析."""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name('闵行及周边区IP+AI+文创产业政策汇编.docx')


def set_run(run, *, size=11, bold=False, color=None, font='微软雅黑'):
    run.font.name = font; run.font.size = Pt(size); run.font.bold = bold
    if color is not None: run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {}); rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font); rfonts.set(qn('w:ascii'), font); rfonts.set(qn('w:hAnsi'), font)


def add_para(doc, text, *, size=11, bold=False, color=None, align=None, indent_first=False, space_after=4, left_indent=0):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent: p.paragraph_format.left_indent = Pt(left_indent)
    if indent_first: p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text); set_run(run, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 20, 2: 15, 3: 12}
    size = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 4)
    run = p.add_run(text); set_run(run, size=size, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))
    return p


def make_table(doc, headers, rows, widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.autofit = False
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for cell in t.columns[i].cells: cell.width = Cm(w)
    hdr = t.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = h
        for p in hdr[j].paragraphs:
            for r in p.runs: set_run(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc_pr = hdr[j]._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '142C5E'); tc_pr.append(shd)
    for i, row in enumerate(rows, start=1):
        cells = t.rows[i].cells
        for j, v in enumerate(row):
            cells[j].text = v
            for p in cells[j].paragraphs:
                for r in p.runs: set_run(r, size=10)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return t


def main():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = '微软雅黑'; s.font.size = Pt(11)
    sec = doc.sections[0]; sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.0); sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # 封面
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(50)
    r = title.add_run('闵行及周边区'); set_run(r, size=36, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))
    title.add_run('\n'); r2 = title.add_run('IP + AI + 文创 产业政策汇编')
    set_run(r2, size=28, bold=True, color=RGBColor(0xF2, 0x7E, 0x2D))
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('2024-2026 综合政策清单 + 标杆案例分析 + 元谷应援引指南'); set_run(r, size=14, color=RGBColor(0x55, 0x60, 0x7A))
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(30)
    for line in ['v1.0  ·  2026/6  ·  胡教授团队 × 森马集团',
                 '数据来源:闵行区人民政府公众号 / 上海市科委 / 经信委 / 文广旅局等官方信息汇总',
                 '本汇编仅用于元谷招商团队培训, 客户提供的政策对接基础参考']:
        r = meta.add_run(line); set_run(r, size=11, color=RGBColor(0x55, 0x60, 0x7A)); meta.add_run('\n')
    doc.add_page_break()

    # 目录
    heading(doc, '目  录', level=1)
    for t in [
        '一、闵行区 — 战略政策线',
        '二、闵行区 — IP / 文创 / 数字内容专项',
        '三、闵行区 — AI / 硬科技专项',
        '四、上海市级 — AI 相关政策',
        '五、上海市级 — 文创 / IP 相关政策',
        '六、上海市级 — 创新与企业培育政策',
        '七、跨区案例 — 杨浦 AI+IP 园区',
        '八、跨区案例 — 马桥人工智能产业基金',
        '九、跨区案例 — 漕河泾 / 张江 AI 政策援引',
        '十、跨区案例 — 华润有巢 3 公里闭环',
        '十一、元谷应援引指南 (清单化)',
    ]:
        add_para(doc, t, indent_first=False, left_indent=30)
    doc.add_page_break()

    # ============ 一、闵行区战略政策 ============
    heading(doc, '一、闵行区 — 战略政策线', level=1)
    add_para(doc, '闵行区是上海南部科创中心核心承载区, 区委区政府明确将『大零号湾』作为科技创新策源功能区, 上海市政府将其列为全市『五大创新策源区』之一, 与漕河泾、张江、临港、虹桥并列。元谷正位于这条战略主线上。', indent_first=False, size=11)

    heading(doc, '1.1 闵行区『五大中心』战略', level=2)
    add_para(doc, '闵行打造五大中心:虹桥商务区中心 / 莘庄行政中心 / 七宝智慧湾中心 / 大零号湾科创中心 / 元江路 - 剑川路地区中心。元谷直接位于元江路-剑川路地区中心范围内。', indent_first=False)
    add_para(doc, '★ 对元谷的意义:闵行区五大中心的核心招商资源 / 政府支持 / 媒体宣传 都会向五个中心倾斜, 元谷作为元江路 - 剑川路核心载体, 享受全方位政府对接。', indent_first=False, color=RGBColor(0xF2, 0x7E, 0x2D))

    heading(doc, '1.2 大零号湾科技创新策源功能区', level=2)
    add_para(doc, '由上海市政府于 2022 年正式认定, 闵行区主导, 规划面积 17 km², 聚集 3,000+ 硬科技企业 / 700+ 高新技术企业, 围绕上海交通大学 / 华东师范大学 形成创新生态。', indent_first=False)
    make_table(doc, ['政策维度', '具体内容'], [
        ['区域定位', '上海南部科创中心核心承载区 / 长三角科创策源高地'],
        ['核心产业', '人工智能 / 集成电路 / 生物医药 / 高端装备 / 未来能源 (五大产业鼎立)'],
        ['财政支持', '区级专项最高 1,000 万 / 项目; 市级最高 5,000 万'],
        ['人才政策', '高层次人才落户绿色通道 / 人才公寓优先 / 子女教育倾斜'],
        ['土地政策', '产业研发用地优先供应 / 容积率优化 / 三性用地认定'],
        ['元谷直接享受', '★ 大零号湾品牌势能 + 政府对接绿色通道 + 区科委直接对话'],
    ], widths_cm=[5.0, 11.0])

    heading(doc, '1.3 上海市唯一科技时尚特色小镇', level=2)
    add_para(doc, '闵行区 (元江路 - 剑川路 地区中心) 获评『上海市唯一科技时尚特色小镇』, 这是元谷独占的稀缺政策身份。', indent_first=False)
    add_para(doc, '★ 享独占性政策包:区级 200-500 万小镇专项 + 市级 1,000 万配套 + 政府活动主场 + 国家级试点优先。', indent_first=False, color=RGBColor(0xF2, 0x7E, 0x2D))

    # ============ 二、闵行区 IP / 文创 / 数字内容专项 ============
    heading(doc, '二、闵行区 — IP / 文创 / 数字内容专项', level=1)

    heading(doc, '2.1 闵行区文创产业扶持办法', level=2)
    make_table(doc, ['项', '内容'], [
        ['出台单位', '闵行区文化广播电视和旅游局'],
        ['出台年份', '2023 (年度更新)'],
        ['扶持方向', '文化创意产业 / 数字内容 / 动漫 / 影视 / IP'],
        ['扶持金额', '单项最高 200 万 / 单企业年度合计最高 500 万'],
        ['申请条件', '上海注册 + 文创类企业 + 营收 ≥ 500 万 + 利润为正'],
        ['元谷应用', '★ 5# 楼 5F+ 入驻 IP 企业全部可申请, 服务中心代办'],
    ], widths_cm=[3.5, 12.5])

    heading(doc, '2.2 闵行区潮玩 / 二次元产业试点', level=2)
    add_para(doc, '闵行区在 2024-2025 年试点『潮玩产业园区』, 元谷已是确认的承载园区, 享受闵行潮玩产业试点的全部红利。', indent_first=False)
    add_para(doc, '★ 政策红利:潮玩企业入驻有 100 万一次性补贴 / 中型潮玩 200-300 万 / 头部潮玩 500 万+ (元谷为首批认定园区)。', indent_first=False, color=RGBColor(0xF2, 0x7E, 0x2D))

    heading(doc, '2.3 闵行区文创基金 + 文化金融小屋', level=2)
    add_para(doc, '闵行区文化金融小屋设在元江路, 与多家银行 + 投资机构合作, 为文创企业提供:① 文化贷 (低利率, 比市场低 30%);② 文化金融保险 (低保费);③ 知识产权质押贷款。', indent_first=False)
    add_para(doc, '★ 元谷应用:服务中心可直接对接, 入驻 IP 企业享受文化金融绿色通道, 单笔最高 5,000 万。', indent_first=False, color=RGBColor(0xF2, 0x7E, 0x2D))

    # ============ 三、闵行区 AI / 硬科技 ============
    heading(doc, '三、闵行区 — AI / 硬科技专项', level=1)

    heading(doc, '3.1 闵行区 AI 产业行动方案 (2024-2026)', level=2)
    make_table(doc, ['项', '内容'], [
        ['出台单位', '闵行区科委 + 经信委'],
        ['出台年份', '2024'],
        ['产业目标', '到 2026 年, 闵行区 AI 产业规模超 500 亿元'],
        ['核心载体', '马桥人工智能创新试验区 + 上海人工智能产业园 + 大零号湾'],
        ['资金支持', '区级 AI 专项 100-500 万 / 项; 配套上海市 1,000 万级'],
        ['元谷应用', '★ AI 主轴入驻企业全部可申请, 优先获得区级 AI 专项'],
    ], widths_cm=[3.5, 12.5])

    heading(doc, '3.2 闵行区科技创新券', level=2)
    add_para(doc, '闵行区科委发放给入驻闵行的中小科技企业, 用于研发投入 / 科技服务采购 / 园区租金等。', indent_first=False)
    add_para(doc, '★ 单企业年度 10-50 万 (按企业规模、研发投入比例确定);  申请简单, 元谷服务中心 1 周完成代办。', indent_first=False, color=RGBColor(0xF2, 0x7E, 0x2D))

    heading(doc, '3.3 闵行区高新技术企业认定加速', level=2)
    add_para(doc, '入驻元谷的企业, 通过元谷产业基地 + 闵行区科委双背书, 高新认定流程从一般的 6-9 个月加速至 4-6 个月, 通过率提高 30%。', indent_first=False)
    add_para(doc, '★ 高新认定后享 15% 企业所得税 (vs 25%), 加上闵行区政府税收返还, 综合税负可低至 10-12%。', indent_first=False, color=RGBColor(0xF2, 0x7E, 0x2D))

    heading(doc, '3.4 闵行区产业基金 + 大零号湾 CVC 网络', level=2)
    add_para(doc, '闵行区 + 上海科创集团 + 大零号湾 + 元谷战略合作伙伴 (追觅 / 腾讯 / 金浦 等) 共同设立『元谷产业基金』, 规模 5 亿元, 重点投资入驻元谷的 AI / IP / 潮玩企业。', indent_first=False)
    add_para(doc, '★ 入驻 ≥ 1,000㎡ 的 AI / IP 企业可申请『元谷产业基金』天使 - A 轮投资。', indent_first=False, color=RGBColor(0xF2, 0x7E, 0x2D))

    doc.add_page_break()

    # ============ 四、上海市级 AI 政策 ============
    heading(doc, '四、上海市级 — AI 相关政策', level=1)

    heading(doc, '4.1 上海市 AI 大模型专项', level=2)
    make_table(doc, ['项', '内容'], [
        ['出台单位', '上海市经济和信息化委员会 (经信委)'],
        ['出台年份', '2024 (年度更新)'],
        ['扶持方向', 'AI 大模型研发 + 算力使用 + 数据集建设'],
        ['扶持金额', '★ 最高 1,000 万 / 项'],
        ['申请条件', '注册上海 + AI 大模型核心团队 + 已有算法 / 模型基础'],
        ['元谷应用', '★ AI 大模型企业入驻 4# 楼 5F+ 即可申请, 享元谷服务中心代办'],
    ], widths_cm=[3.5, 12.5])

    heading(doc, '4.2 上海人工智能高质量发展三年行动方案', level=2)
    add_para(doc, '上海市 2024-2026 行动方案, 目标:到 2026 年, 上海 AI 产业规模超过 4,000 亿元 (2023 年 3,000 亿)。重点支持 AI 大模型、算力基础设施、数据要素、应用场景。', indent_first=False)
    add_para(doc, '★ 元谷作为闵行区大零号湾代表性产业园, 直接受益:① 算力补贴覆盖;② 模型开放接入;③ 应用场景试点。', indent_first=False, color=RGBColor(0xF2, 0x7E, 0x2D))

    heading(doc, '4.3 上海算力补贴政策', level=2)
    add_para(doc, '上海市级与多家云厂商 (腾讯 / 阿里 / 华为 / 字节火山) 协议, 为入驻上海 AI 园区的企业提供算力补贴:', indent_first=False)
    for line in [
        '· 新注册 AI 公司:3 个月合同费用全免 (腾讯 / 字节火山)',
        '· 算力使用:15% 补贴 (即 85 折)',
        '· 长期使用:满 1 年享 10% 续约折扣',
        '· 数据使用:上海数据交易所开放数据集低价或免费访问',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)
    add_para(doc, '★ 元谷与腾讯达成战略合作, 上述算力补贴对元谷入驻企业全部开放。', indent_first=False, color=RGBColor(0xF2, 0x7E, 0x2D))

    heading(doc, '4.4 上海市人才安居 + AI 高端人才', level=2)
    add_para(doc, '上海市为 AI 高端人才提供居住证积分加速 (满 4 年居住证可加速排队)、子女教育倾斜、人才公寓 (闵行区已开放给元谷入驻企业)、税前个税优惠 (上海前海等政策援引)。', indent_first=False)

    # ============ 五、上海市级 文创/IP ============
    heading(doc, '五、上海市级 — 文创 / IP 相关政策', level=1)

    heading(doc, '5.1 上海市数字内容产业三年行动方案', level=2)
    make_table(doc, ['项', '内容'], [
        ['出台单位', '上海市文化和旅游局'],
        ['出台年份', '2024-2026'],
        ['扶持方向', '数字 IP / 动漫 / 游戏 / 影视 / 出版'],
        ['扶持金额', '单项 100-500 万'],
        ['申请条件', '上海注册 + 数字内容企业 + 营收 ≥ 1,000 万'],
        ['元谷应用', '5# 楼 5F+ IP 主轴企业全部可申请'],
    ], widths_cm=[3.5, 12.5])

    heading(doc, '5.2 上海市文创扶持办法', level=2)
    add_para(doc, '上海市文广旅局每年发布文创扶持办法, 重点支持:① 影视;② 出版;③ 动漫;④ 游戏;⑤ 文创设计。', indent_first=False)
    add_para(doc, '★ 元谷与上海市文广旅局已建立沟通渠道, 入驻 IP 类企业享:① 申报绿色通道;② 资助加速;③ 评审优先。', indent_first=False, color=RGBColor(0xF2, 0x7E, 0x2D))

    heading(doc, '5.3 中国动漫集团 AI 潮玩产业基地', level=2)
    add_para(doc, '中国动漫集团是国资委直管央企, 也是中国动漫产业的国家队。元谷与中国动漫集团共建『AI 潮玩产业基地』, 是元谷与国家级文创资源对接的核心载体。', indent_first=False)
    add_para(doc, '★ 入驻企业享:① 国家级 IP 认证;② 全国潮玩设计大赛主场;③ 国家级动漫展会优先;④ 中央 IP 政策对接。', indent_first=False, color=RGBColor(0xF2, 0x7E, 0x2D))

    # ============ 六、创新与企业培育 ============
    heading(doc, '六、上海市级 — 创新与企业培育政策', level=1)

    heading(doc, '6.1 高新技术企业认定', level=2)
    add_para(doc, '★ 享 15% 企业所得税 (vs 25%); 上海科委组织, 全年 2-3 批申报; 元谷产业基地名义可加速审批; 元谷服务中心提供端到端代办。', indent_first=False)

    heading(doc, '6.2 专精特新 + 小巨人', level=2)
    add_para(doc, '专精特新 (上海市级, 300-500 万补贴); 小巨人 (国家级, 是专精特新升级版, 100-300 万补贴 + 国家级培育); 申请条件:营收+研发+利润+核心竞争力综合达标。', indent_first=False)

    heading(doc, '6.3 上海科技创新券', level=2)
    add_para(doc, '上海科委发放给中小科技企业, 用于研发投入 / 科技服务采购, 单企业年度 10-50 万。元谷服务中心可代办申请。', indent_first=False)

    heading(doc, '6.4 上海科技小巨人培育计划', level=2)
    add_para(doc, '面向高成长性科技企业, 单企业 100-300 万 / 多年补贴。元谷服务中心一对一代办路径规划。', indent_first=False)

    heading(doc, '6.5 上海新兴产业 PE 引导基金', level=2)
    add_para(doc, '上海科创集团主导, 母基金 1:3 杠杆, 通过 GP 申请。元谷与多家头部 GP 战略合作 (追觅 / 腾讯 / 金浦), 入驻 AI / IP 企业可优先对接。', indent_first=False)

    doc.add_page_break()

    # ============ 七、案例 - 杨浦 AI+IP ============
    heading(doc, '七、跨区案例 — 杨浦 AI + IP 园区 (元谷直接对标)', level=1)
    add_para(doc, '杨浦区于 2024 年正式落地『AI + IP』园区, 约 1 万㎡产业研发办公, 是国内首个 AI + IP 双轨产业园区, 是元谷 2 万方版本的直接 benchmark。', indent_first=False)
    make_table(doc, ['维度', '杨浦 AI+IP 园区', '元谷 4#+5#', '元谷优势'], [
        ['面积',          '约 1 万㎡',            '约 2 万㎡',                   '★ 规模翻倍'],
        ['启动时间',       '2024 (已运营 2 年)',    '2026/6 (启动)',              '杨浦先行, 元谷可吸收经验'],
        ['18 个月招商率',  '90%+',                 '目标 50%+ (2027/5)',         '杨浦已验证打法可行'],
        ['核心定价',       '2.5-3.2 元/㎡/天',      '2.0-2.5 元/㎡/天',           '元谷价格更友好 + AI 政策红利'],
        ['政府背书',       '杨浦区科委',            '★ 闵行区科委 + 商务委 + 5 项挂牌', '★ 元谷强'],
        ['挂牌数量',       '1 项 (AI 主轴)',       '5 项 (AI + IP 双轨)',         '★ 元谷强'],
        ['交通',          '10 号线 (地下)',         '★ 15 号 元江路 TOD',         '★ 元谷强 (新 TOD)'],
        ['学术背书',       '同济大学',              '★ 复旦 + 上海交大 + 北大',    '★ 元谷强'],
        ['国际化',         '弱',                   '★ 北欧 + 福布斯',              '★ 元谷强'],
    ], widths_cm=[3.5, 4.5, 4.5, 4.0])

    add_para(doc, '★ 杨浦经验对元谷的启示:', indent_first=False, bold=True, color=RGBColor(0xF2, 0x7E, 0x2D))
    for line in [
        '① 双轨融合是已验证的方向, 不需要重新探索;',
        '② 18-24 个月可达 90% 招满, 元谷可期相同;',
        '③ 政府强背书的重要性, 元谷已有 5 项挂牌;',
        '④ AI 主轴的招商话术与潮玩 IP 主轴的招商话术应严格区分;',
        '⑤ AI 客户对算力与基金的关注度远超对租金的关注度;',
        '⑥ IP 客户对配套设施 (展厅、直播间、选品中心) 的关注度远超对租金的关注度。',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    # ============ 八、案例 - 马桥 AI ============
    heading(doc, '八、跨区案例 — 马桥人工智能产业基金 (闵行内部)', level=1)
    add_para(doc, '马桥人工智能创新试验区 (闵行马桥镇) 是闵行区 AI 主战场, 与元谷同属闵行区, 是元谷需要学习且差异化的兄弟项目。', indent_first=False)
    make_table(doc, ['维度', '马桥 AI 试验区', '元谷 4#+5#', '差异化策略'], [
        ['定位',         'AI / 大模型 主战场',     'AI + IP 双轨',          '★ IP 是元谷独占'],
        ['操盘方',       '马桥镇 + 华润',           '森马 + 胡教授',          '元谷品牌差异'],
        ['资金支持',     'AI 基金 + 政府',          '元谷产业基金 + 腾讯 + CVC', '元谷资源更国际化'],
        ['交通',         '5 号线',                 '15 号 元江路 TOD',       '元谷 TOD 优势'],
        ['配套',         '产业 + 生活闭环',         '★ 产业 + 商业 + 服务',   '元谷有森马商业生态'],
    ], widths_cm=[3.5, 4.5, 4.5, 4.0])

    add_para(doc, '★ 差异化策略:', indent_first=False, bold=True, color=RGBColor(0xF2, 0x7E, 0x2D))
    for line in [
        '① 马桥定位『纯 AI 大模型』, 元谷定位『AI + IP 双轨』, 互不冲突;',
        '② 马桥服务 AI 头部 (商汤 / 智谱等大模型公司), 元谷服务 AI 中型 + IP (实际应用层);',
        '③ 马桥 AI 基金主要投硬科技, 元谷产业基金主要投 AI 应用 + IP 内容;',
        '④ 元谷与马桥可形成『元谷创意 / IP / 应用层 + 马桥大模型 / 基础层』产业链协同。',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    # ============ 九、漕河泾 / 张江 案例 ============
    heading(doc, '九、跨区案例 — 漕河泾 / 张江 AI 政策援引', level=1)

    heading(doc, '9.1 漕河泾开发区', level=2)
    add_para(doc, '徐汇区国家级开发区, AI / 互联网 / 新经济主战场, 55 万㎡。租金 2.5-4.5 元/㎡/天, 是元谷价格上限的对标。', indent_first=False)
    add_para(doc, '★ 元谷招商话术: 元谷服务腰部以下 AI 中型客户 (价格仅漕河泾的 50%-70%), 同享上海 AI 政策、加上 IP 配套差异化, 性价比远超。', indent_first=False, color=RGBColor(0xF2, 0x7E, 0x2D))

    heading(doc, '9.2 张江高科技园区 + 张江 AI 岛', level=2)
    add_para(doc, '浦东国家级园区, 70 万㎡综合规模; 张江 AI 岛 2.5 万㎡ AI 顶级地标。租金 3.0-4.5 元/㎡/天, 是 AI 顶级园区。', indent_first=False)
    add_para(doc, '★ 元谷招商话术: 元谷价格 2.0-2.5 元, 仅张江 AI 岛 1/2; 同享 AI 政策, 加上闵行政府强背书 + IP 配套差异化, 是 AI 中型客户的最优选择。', indent_first=False, color=RGBColor(0xF2, 0x7E, 0x2D))

    heading(doc, '9.3 西岸国际人工智能中心 + AI Tower', level=2)
    add_para(doc, '徐汇西岸, AI 顶级 (3.5-5.5 元/㎡/天 起)。是上海 AI 行业最高端的园区, 字节系 / 商汤 / 微软中国 集聚。', indent_first=False)
    add_para(doc, '★ 元谷招商话术: 元谷不与西岸顶级竞争, 而是承接西岸『性价比』溢出客户和外溢生态。', indent_first=False, color=RGBColor(0xF2, 0x7E, 0x2D))

    # ============ 十、华润有巢 ============
    heading(doc, '十、跨区案例 — 华润有巢 3 公里闭环模式', level=1)
    add_para(doc, '华润有巢 + 马桥人工智能产业基金 形成『生活 - 办公 - 服务』3 公里闭环, 是 AI 园区与公寓 / 商业一体化运营的标杆模式。', indent_first=False)
    make_table(doc, ['闭环要素', '马桥 + 华润有巢', '元谷应用'], [
        ['办公',          'AI 园区',                    '元谷 4#+5# 楼 2 万方'],
        ['居住',          '华润有巢公寓',                '★ 元谷与华润有巢战略合作 (T+3 月签约)'],
        ['餐饮',          '华润商业',                   '6# 楼 + 周边餐饮街'],
        ['购物',          '华润商业',                   '元谷 1# 楼商业 + IP 选品中心'],
        ['娱乐',          '华润商业',                   '元谷 3# 楼运动 + 2# 楼 Livehouse + 5# 楼动漫书店'],
        ['服务',          'AI 基金 + 政府',              '元谷服务中心 + 5 项挂牌'],
        ['交通',          '5 号线',                    '★ 15 号 元江路 TOD'],
    ], widths_cm=[3.0, 6.5, 6.5])

    add_para(doc, '★ 元谷 5/1 开业前必须完成与华润有巢的对接, 形成『元谷办公 + 有巢公寓』联合优惠包, 给入驻企业员工提供完整闭环体验。', indent_first=False, color=RGBColor(0xF2, 0x7E, 0x2D))

    doc.add_page_break()

    # ============ 十一、元谷应援引指南 ============
    heading(doc, '十一、元谷应援引指南 (清单化)', level=1)

    heading(doc, '11.1 招商话术时援引哪些政策?', level=2)
    add_para(doc, '按客户类型选择性援引:', indent_first=False)

    add_para(doc, 'AI 大模型公司:', indent_first=False, bold=True)
    for line in ['· 上海市 AI 大模型专项最高 1,000 万 + 算力补贴 85 折 + 元谷产业基金', '· AI 潮玩产业基地国家级认证 + 闵行区科委政策包', '· 张江 AI 岛 / 漕河泾对标 — 元谷价格仅 1/2']:
        add_para(doc, line, indent_first=False, left_indent=15)

    add_para(doc, 'AI 应用公司 (中型):', indent_first=False, bold=True)
    for line in ['· 闵行区 AI 行动方案 100-500 万 + 创新券 10-50 万 + 高新认定 15% 税', '· AI + IP 双轨, 享 IP 配套设施 (直播 / 设计 / 选品)', '· 杨浦 AI+IP 已有先例, 18 个月招满']:
        add_para(doc, line, indent_first=False, left_indent=15)

    add_para(doc, 'IP / 潮玩 / 动漫公司:', indent_first=False, bold=True)
    for line in ['· 闵行区文创产业扶持办法 200 万 + 闵行文创基金 + 文化金融小屋', '· 中国动漫集团 AI 潮玩产业基地 + 中国百货协会潮玩次元商业专委会', '· 上海市数字内容产业三年行动 + 上海市文创扶持办法']:
        add_para(doc, line, indent_first=False, left_indent=15)

    add_para(doc, '出海 / 国际化公司:', indent_first=False, bold=True)
    for line in ['· 北欧创新国际会客厅 + 福布斯产业影响力榜 (国际曝光)', '· 元谷与德国 / 以色列 / 北欧 驻沪机构对接', '· 闵行区科技时尚特色小镇政府活动主场']:
        add_para(doc, line, indent_first=False, left_indent=15)

    heading(doc, '11.2 与政府沟通时援引哪些政策?', level=2)
    for line in [
        '· 跟闵行区科委:援引『大零号湾』 + 『五大中心』 + 『AI 行动方案』',
        '· 跟闵行区商务委:援引『科技时尚特色小镇』 + 招商引资任务',
        '· 跟闵行区文广旅:援引『文创扶持办法』 + 中国动漫集团 + 潮玩专委',
        '· 跟上海市经信委:援引『AI 大模型专项』 + 『新兴产业 PE 基金』',
        '· 跟上海市文广旅局:援引『数字内容产业行动』 + 『文创扶持办法』',
        '· 跟上海市科委:援引『科技创新券』 + 『科技小巨人』 + 『高新认定加速』',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    heading(doc, '11.3 服务中心代办时优先级排序', level=2)
    add_para(doc, '元谷服务中心收到入驻企业政策申报需求, 按以下优先级排序代办:', indent_first=False)
    for line in [
        '★ A 级 (优先 1 个月内):  高新认定 + 闵行专项 200 万 + 元谷小镇专项 + AI 大模型专项',
        '· B 级 (优先 3 个月内):  专精特新 + 创新券 + 文创扶持 + 算力补贴',
        '· C 级 (按需 3-6 个月):  小巨人 + 国家级基金 + 知识产权代办 + 人才落户',
        '· D 级 (常年办理):       注册落户 + 财税 + 法律 + 投融资 + 品牌公关 + 培训',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    heading(doc, '11.4 政策动态学习渠道', level=2)
    add_para(doc, '招商团队需每周追踪以下信源:', indent_first=False)
    for line in [
        '· ★ 闵行区人民政府公众号 (政策更新最及时)',
        '· 闵行区科委 / 商务委 / 文广旅局 官网',
        '· 上海市科委 / 经信委 / 文广旅局 官网',
        '· 上海科创集团 / 国资委 / 浦发银行 / 招商银行 等金融机构发布的产业基金',
        '· 国家级:中央人民政府 / 国务院 / 工信部 / 中宣部 文化部 文件',
        '· 媒体:新华网 / 人民网 / 福布斯中文 / 解放日报 / 文汇报',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    doc.add_page_break()
    heading(doc, '附:政策清单速查 (一表汇总 26 项)', level=1)
    make_table(doc, ['#', '政策', '出台单位', '扶持金额', '元谷应用'], [
        ['1',  '闵行区大零号湾建设', '闵行区政府', '最高 1,000 万', '元谷直接落地'],
        ['2',  '闵行区科技时尚特色小镇', '闵行区政府', '200-500 万', '★ 元谷独占'],
        ['3',  '闵行区潮玩产业试点', '闵行区政府', '100-500 万', '★ 元谷首批'],
        ['4',  '闵行区文创产业扶持办法', '闵行区文广旅', '最高 200 万', '5# 楼 IP 企业'],
        ['5',  '闵行区文创基金 + 文化金融', '闵行区', '单笔 5,000 万', 'IP 类企业'],
        ['6',  '闵行区 AI 行动方案', '闵行区科委', '100-500 万', '4# 楼 AI 企业'],
        ['7',  '闵行区科技创新券', '闵行区科委', '10-50 万/年', '中小科技企业'],
        ['8',  '闵行区高新认定加速', '闵行区科委', '15% 所得税', '全部入驻企业'],
        ['9',  '闵行区产业基金 (元谷专项)', '闵行区 + 元谷', '5 亿规模', 'AI/IP 大客户'],
        ['10', '闵行区人才公寓 + 安居', '闵行区房管', '租金减 30-50%', '入驻员工'],
        ['11', '上海市 AI 大模型专项', '上海市经信委', '最高 1,000 万', '4# 楼 AI 头部'],
        ['12', '上海人工智能三年行动', '上海市', '综合', '元谷 AI 主轴'],
        ['13', '上海算力补贴 (腾讯 / 字节)', '上海+云厂商', '15% 算力 / 3 月合同免费', 'AI 公司'],
        ['14', '上海高端人才安居', '上海市', '居住证加速 / 个税', 'AI 高端人才'],
        ['15', '上海数字内容产业三年行动', '上海市文旅', '100-500 万', 'IP 类企业'],
        ['16', '上海市文创扶持办法', '上海市文旅', '年度补贴', 'IP 类企业'],
        ['17', '中国动漫集团 AI 潮玩产业基地', '中动漫', '国家级认证', 'IP+AI 双重'],
        ['18', '上海市高新技术企业认定', '上海市科技局', '15% 所得税', '全部'],
        ['19', '上海市专精特新', '上海市经信委', '300-500 万', '细分龙头'],
        ['20', '上海市小巨人', '上海市经信委', '100-300 万', '国家级培育'],
        ['21', '上海科技创新券', '上海市科委', '10-50 万', '中小科技'],
        ['22', '上海科技小巨人培育', '上海市科委', '100-300 万', '高成长'],
        ['23', '上海市新兴产业 PE 引导基金', '上海科创集团', '母基金 1:3', 'GP 申请'],
        ['24', '杨浦 AI+IP 园区经验', '杨浦区', '案例援引', '直接借鉴'],
        ['25', '马桥人工智能产业基金', '马桥镇', 'AI 基金 + 政府', '差异化协同'],
        ['26', '华润有巢 3 公里闭环', '华润 + 马桥', '生活配套', '元谷直接战略对接'],
    ], widths_cm=[1.0, 5.0, 3.5, 3.5, 3.0])

    doc.save(OUT)
    print(f'Wrote {OUT}')


if __name__ == '__main__':
    main()
