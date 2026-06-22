"""元谷 2 万方招商手册 (文字稿) — 用于设计公司外包美化的内容底稿."""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name('元谷2万方招商手册(文字稿).docx')


def set_run(run, *, size=11, bold=False, color=None, font='微软雅黑'):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {}); rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font); rfonts.set(qn('w:ascii'), font); rfonts.set(qn('w:hAnsi'), font)


def add_para(doc, text, *, size=11, bold=False, color=None, align=None, indent_first=True, space_after=4, left_indent=0):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent: p.paragraph_format.left_indent = Pt(left_indent)
    if indent_first: p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text); set_run(run, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 20, 2: 15, 3: 12}
    size = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 4)
    run = p.add_run(text); set_run(run, size=size, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))
    return p


def make_table(doc, headers, rows, widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.autofit = False
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for cell in t.columns[i].cells: cell.width = Cm(w)
    hdr = t.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = h
        for p in hdr[j].paragraphs:
            for r in p.runs: set_run(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc_pr = hdr[j]._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '142C5E'); tc_pr.append(shd)
    for i, row in enumerate(rows, start=1):
        cells = t.rows[i].cells
        for j, v in enumerate(row):
            cells[j].text = v
            for p in cells[j].paragraphs:
                for r in p.runs: set_run(r, size=10.5)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return t


def main():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = '微软雅黑'; s.font.size = Pt(11)
    sec = doc.sections[0]; sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.0); sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # 封面
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(50)
    r = title.add_run('元谷'); set_run(r, size=44, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))
    title.add_run('\n'); r2 = title.add_run('4# + 5# 楼 2 万方产业研发办公')
    set_run(r2, size=22, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))
    title.add_run('\n'); r3 = title.add_run('招  商  手  册')
    set_run(r3, size=22, bold=True, color=RGBColor(0xF2, 0x7E, 0x2D))

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('IP + AI 双轨 · 长三角首个 AI+潮玩 双牌照产业策源高地'); set_run(r, size=14, color=RGBColor(0x55, 0x60, 0x7A))

    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(30)
    for line in ['Semir 森马集团 出品  ·  上海闵行 大零号湾', '2026/6 v1.0  ·  本手册仅用于招商沟通']:
        r = meta.add_run(line); set_run(r, size=12); meta.add_run('\n')
    doc.add_page_break()

    # ====== 一、项目总览 ======
    heading(doc, '一、项目总览', level=1)
    add_para(doc, '元谷 (Yuangu)，全称『森马 (上海) 国际运营中心』，定位为长三角首个 AI + IP 双轨产业策源高地，由森马集团出品、胡教授团队联合运营，坐落于上海市闵行区元江路-剑川路核心区，正处于『大零号湾』科技创新策源功能区与『上海市唯一科技时尚特色小镇』双重身份叠加之上。', indent_first=False)
    add_para(doc, '总建筑面积 22 万㎡，由 1#–6# 共 6 栋建筑构成完整产业生态。本手册聚焦 4# 楼 5F 及以上潮玩产业集群（约 1 万㎡）与 5# 楼 5F 及以上潮玩产业集群（约 1 万㎡）合计 2 万㎡产业研发办公空间的招商招租。', indent_first=False)

    make_table(doc, ['项目核心数据', '数值'], [
        ['总建筑面积',         '22 万㎡'],
        ['本手册招商范围',     '2 万㎡ (4# 楼 5F+ 1 万㎡ + 5# 楼 5F+ 1 万㎡)'],
        ['产业研发用地属性',    '已确认 (双方一致主推 AI + IP 双轨)'],
        ['交通条件',           '15 号线 元江路站 TOD 项目, 单日客流 5-7 万人次'],
        ['15 分钟车行覆盖',    '24 万居住人口 + 12 万产业办公人口'],
        ['距紫竹高新区',       '5 km / 车程约 10 分钟'],
        ['距虹桥机场',         '17 km / 机场快线约 45 分钟'],
        ['距浦东机场',         '36 km / 车程约 1.5 小时'],
        ['项目开业日',         '2027/5/1'],
        ['本手册版本',         'v1.0  ·  2026 年 6 月'],
    ], widths_cm=[6.0, 10.0])

    # ====== 二、战略定位 ======
    heading(doc, '二、战略定位', level=1)
    heading(doc, '1. 双重战略身份', level=2)
    add_para(doc, '·『大零号湾』文创融合核心区 — 闵行区五大中心之一，能级比肩漕河泾、张江', indent_first=False)
    add_para(doc, '·『上海市唯一科技时尚特色小镇』— 政策与品牌双重独占', indent_first=False)

    heading(doc, '2. 产业愿景', level=2)
    add_para(doc, '构建『时尚研发』与『文创转化』双向赋能的示范性融合标杆，打造 AI + IP 双赛道协同发展、上海首个 1 万㎡ 以上 AI+IP 综合产业园区。', indent_first=False)

    heading(doc, '3. 三大战略目标', level=2)
    for line in ['① 国际化产业枢纽 (4# 楼 5F+ — 北欧 / 日韩 / 东南亚 IP 首站)', '② 品牌运营高地 (5# 楼 5F+ — 头部企业 + 中型潮玩运营企业)', '③ IP 创制中心 (4# 楼 5F AI 共享设计 + 5# 楼 5F 潮玩产业展厅)']:
        add_para(doc, line, indent_first=False, left_indent=15)

    heading(doc, '4. 五大优势', level=2)
    for line in ['创意集聚', '源头孵化', '场景体验', '生态复合', '集约选品']:
        add_para(doc, '· ' + line, indent_first=False, left_indent=15)

    doc.add_page_break()

    # ====== 三、IP + AI 双轨产业地图 ======
    heading(doc, '三、IP + AI 双轨产业地图', level=1)
    heading(doc, '1. 4# 楼 5F+ ★ AI 主轴 (国际 IP 创意层)', level=2)
    add_para(doc, '面积:约 1 万㎡  |  建议日租金:2.2 – 2.5 元/㎡/天  |  年租金:803 – 913 元/㎡', indent_first=False, bold=True)
    add_para(doc, '同栋协同配套:', indent_first=False, bold=True)
    for line in ['· 4F 直播中心 (1,000㎡+，与绮丽少女女团联合直播)',
                 '· 5F AI 共享设计中心 (联动腾讯 + 上海交大设计学院, 提供 AI 工作站和 AI 设计算力)',
                 '· 5F AI 共享打样 & DIY 产品中心 (帮助 IP / 潮玩 / 设计师快速打样, 提高面世效率)']:
        add_para(doc, line, indent_first=False, left_indent=15)
    add_para(doc, '目标客户画像:', indent_first=False, bold=True)
    for line in ['· AI 大模型应用公司 (含 AIGC、AI Agent、垂直行业 AI)',
                 '· AI 设计 / AI 内容生成公司',
                 '· 国际 IP (北欧、日韩、东南亚) 进入中国首站',
                 '· 海外 AI 公司中国办公室',
                 '· 字节 / MiniMax / 月之暗面 / 360AI / 智谱 / 海外 AI 公司']:
        add_para(doc, line, indent_first=False, left_indent=15)

    heading(doc, '2. 5# 楼 5F+ ★ IP 主轴 (潮玩产业总部层)', level=2)
    add_para(doc, '面积:约 1 万㎡  |  建议日租金:2.0 – 2.2 元/㎡/天  |  年租金:730 – 803 元/㎡', indent_first=False, bold=True)
    add_para(doc, '同栋协同配套:', indent_first=False, bold=True)
    for line in ['· 1-4F 动漫书店 + 休闲娱乐 (天然消费场景)',
                 '· 5F 潮玩产业展厅 (聚焦品牌 IP 叙事、行业交流、渠道拓展)',
                 '· 临近 4# 楼 1-3F 潮玩艺术中心']:
        add_para(doc, line, indent_first=False, left_indent=15)
    add_para(doc, '目标客户画像:', indent_first=False, bold=True)
    for line in ['· 头部央企 (中国动漫集团旗下子公司 / 中字头数字央企)',
                 '· 中型潮玩运营企业 (盲盒 / 手办 / 数字 IP / 衍生品)',
                 '· 内容公司 (动漫 / 视频 / 直播 / 短视频 / 出版)',
                 '· 行业协会 (潮玩次元商业专委会 / 中国百货商业协会)',
                 '· B 站生态企业 + 上海美影厂周边企业']:
        add_para(doc, line, indent_first=False, left_indent=15)

    heading(doc, '3. 双轨融合协同点', level=2)
    for line in [
        '· AI + IP 联合招商 — 招商时强调 4#/5# 楼为一个完整生态, 客户可跨楼协同',
        '· 同楼层互通 — 4#5F+ AI 设计中心与 5#5F+ 潮玩展厅形成『从设计到展陈』全链条',
        '· 共享配套 — 直播 / 选品 / DIY / 服务中心 跨楼共享, 入驻企业可灵活使用',
        '· 共享算力 — 通过腾讯算力补贴政策, 入驻 AI 企业享 15% 算力折扣',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    doc.add_page_break()

    # ====== 四、租赁条款 ======
    heading(doc, '四、租赁条款与定价', level=1)
    make_table(doc, ['面积档位', '日租金 (元/㎡/天)', '年租金 (元/㎡)', '物业费 (元/㎡/月)', '装修标准', '免租期'], [
        ['≤ 500㎡ 小型',       '2.2 – 2.5', '803 – 913', '10', '精装交付',     '15 天'],
        ['501 – 2,000㎡ 中型',  '2.0 – 2.3', '730 – 840', '10', '简装/精装可选', '1 个月'],
        ['2,001 – 5,000㎡ 大型', '2.0 – 2.2', '730 – 803', '10', '客户定制',     '2 个月'],
        ['> 5,000㎡ 头部/旗舰',   '★ 1.8 – 2.0 (议价)', '657 – 730', '10', '客户定制 + 设计补贴', '3 个月'],
    ], widths_cm=[3.6, 3.0, 2.6, 2.4, 2.4, 2.0])

    heading(doc, '签约条款', level=2)
    for line in [
        '· 起租日: 合同签订后 30 日内入驻, 装修期不计租金',
        '· 押金: 月租金 × 2 个月',
        '· 预付租金: 3 个月',
        '· 租期: 建议 3 + 2 (3 年首期 + 续约 2 年) ; 头部客户可议为 5 年',
        '· 年度递增: 3% (前 3 年) , 中型以上可议',
        '· 装修补贴包 (≥ 2,000㎡ 客户专享): 100-300 元/㎡ 定制补贴, 装修期延长 30 天',
        '· 大客户专属包 (≥ 5,000㎡): 装修补贴 + 服务中心 1 年免费 + 政府对接绿色通道',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    doc.add_page_break()

    # ====== 五、五项产业牌照 ======
    heading(doc, '五、5 项产业牌照与挂牌资源', level=1)
    add_para(doc, '元谷已确认与下列 5 家国家级 / 上海市级机构共建产业牌照, 是入驻企业最重要的政府背书与品牌势能。', indent_first=False)
    make_table(doc, ['序号', '挂牌名称', '出牌方', '对入驻企业的价值'], [
        ['①', 'AI 潮玩产业基地',              '中国动漫集团',           '国家级 AI+IP 产业认证 / 政策对接 / 产业大会主场'],
        ['②', '潮玩次元商业专委会',           '中国百货商业协会',         '聚集潮玩零售生态 / 政府对话渠道'],
        ['③', '复旦大学住房政策研究中心 元谷分中心', '复旦大学住房政策研究中心', '学术背书 / 政策研究 / 高净值人脉'],
        ['④', '上海市科技企业联合会 元谷产业基地', '上海市科技企业联合会',     '上海科技企业生态导流 / 政府补贴申报'],
        ['⑤', '福布斯产业影响力奖 元谷专场',     '福布斯',                   '国际品牌势能 / 年度评选 IP / 国际媒体曝光'],
    ], widths_cm=[1.5, 5.5, 3.5, 5.5])

    heading(doc, '挂牌如何转化为客户红利', level=2)
    for line in [
        '· ★ 申请高新认定 / 专精特新 — 元谷产业基地名义可加速审批',
        '· ★ 政府补贴绿色通道 — 元谷服务中心代办, 提交时间从 6 月缩短至 2 月',
        '· ★ 福布斯年度评选 — 入驻企业可参评, 上榜后国际媒体声量翻倍',
        '· ★ 行业大会主场 — 中国动漫集团 / 中百协年度大会落地元谷',
        '· ★ 复旦学术资源 — 政策预研 / 行业报告 / 高净值人脉对接',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    doc.add_page_break()

    # ====== 六、共享配套与服务 ======
    heading(doc, '六、共享配套与服务体系', level=1)
    heading(doc, '1. 共享配套 (元谷自营业态)', level=2)
    make_table(doc, ['配套', '面积', '落位', '功能'], [
        ['IP 潮玩选品&仓储式零售中心', '约 5,000㎡', '1# 4F', '华东首个 IP 潮玩选品 + 仓储式销售空间'],
        ['共享直播中心',              '约 1,000㎡', '4# 4F', '潮玩电商直播, 联动绮丽少女女团'],
        ['AI 共享设计中心',           '约 800㎡',   '4# 5F', '联动腾讯 + 高校 AI 工作站'],
        ['AI 共享打样 / DIY 中心',     '约 600㎡',   '4# 5F', '帮助 IP 快速打样, 提高面世效率'],
        ['潮玩艺术中心',              '约 2,000㎡', '4# 1-3F', '南上海首个特色艺术文化体验空间'],
        ['潮玩产业展厅',              '约 800㎡',   '5# 5F', '聚焦品牌 IP 叙事 / 行业交流 / 渠道拓展'],
        ['代运营物流中心',            '联动森马二期', '森马二期仓库', '智能仓储 + 物流代运营'],
    ], widths_cm=[5.0, 2.5, 2.5, 6.0])

    heading(doc, '2. 元谷科技企业服务中心 (B2B 增值服务)', level=2)
    add_para(doc, '元谷设立『科技企业服务中心』提供 9 大类标准化 + 项目制服务, 入驻企业首年部分服务免费:', indent_first=False)
    for line in [
        '① 注册落户 / 政策红利申报代办 (5,000 - 30,000 元/项)',
        '② 财税 / 法律顾问 / 税筹 (1,500 - 10,000 元/月)',
        '③ 知识产权 (商标 / 专利 / 潮玩 IP 维权)',
        '④ 政府补贴申报 (高新 / 专精特新 / 创新券 / 文创基金, 8-15% 提成)',
        '⑤ 人才与签证 (居住证积分 / 落户 / 外籍签证)',
        '⑥ 投融资 (路演 / FA / 并购)',
        '⑦ 品牌与公关 (含潮玩出海)',
        '⑧ 数字化工具 (SaaS / AI 设计工作站)',
        '⑨ 培训与认证 (潮玩产业认证 / 出海实操营)',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    heading(doc, '3. 政策与资本配套', level=2)
    for line in [
        '★ 腾讯算力补贴: 新注册 AI 公司 3 个月合同免费 / 算力费用 85 折',
        '★ CVC 创业投资基金: 入驻 AI / IP 企业天使-A 轮投资优先',
        '★ 闵行区科委 / 商务委政策包: 高新认定 / 专精特新 / 创新券 / 闵行专项',
        '★ 上海市 AI 大模型专项: 最高 1,000 万/项目',
        '★ 闵行区人才公寓: 高端人才租房减免 30-50%',
        '★ 政府绿色通道: 行政审批 / 企业代办 / 各项申报',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    doc.add_page_break()

    # ====== 七、6 场产业沙龙年度日历 ======
    heading(doc, '七、年度产业活动日历', level=1)
    add_para(doc, '元谷年度承办 6 场产业沙龙 + 1 届全国潮玩设计大赛 + 1 届福布斯榜单, 每场沙龙到场目标产业客户不少于 30 家:', indent_first=False)
    make_table(doc, ['#', '时间', '主题', '联办方'], [
        ['#1', 'T+1 月  (2026/7)',  'AI + 潮玩 跨界融合 (借势 5/22 峰会)', '中动漫 + 腾讯 + AI 腾讯生态'],
        ['#2', 'T+3 月  (2026/9)',  '潮玩出海 (北欧 / 日韩 / 东南亚)',      '北欧创新国际会客厅 + 福布斯'],
        ['#3', 'T+5 月  (2026/11)', '投融资路演',                          '追觅 + 招商银行 + 长江证券 + 金浦'],
        ['#4', 'T+7 月  (2027/1)',  '设计与创意',                          '上海交大设计学院 + 上海市科企联'],
        ['#5', 'T+9 月  (2027/3)',  '内容 IP 与 Z 世代',                  '中百协潮玩次元专委 + 中动漫'],
        ['#6', 'T+11 月 (2027/5)',  '政策补贴与小镇',                      '闵行科协 + 复旦住房政策研究中心'],
    ], widths_cm=[1.5, 4.0, 5.5, 5.0])

    add_para(doc, '入驻企业可优先获得活动主办权 / 主题演讲机会 / 媒体曝光配额; 单场触达 ≥ 30 家产业客户 + ≥ 100 万次媒体曝光。', indent_first=False)

    # ====== 八、招商优惠政策 ======
    heading(doc, '八、招商优惠政策 (大客户专享)', level=1)
    make_table(doc, ['客户档位', '面积', '优惠包'], [
        ['头部 / 央企 / 旗舰', '> 5,000㎡', '★ 装修补贴 300 元/㎡ + 免租 3 个月 + 服务中心 1 年免费 + 政府绿色通道 + AI 算力补贴 + CVC 基金对接'],
        ['大型 / 中型',         '2,001 - 5,000㎡', '装修补贴 200 元/㎡ + 免租 2 个月 + 服务中心半价 + 政府对接'],
        ['中型 / 中小型',       '501 - 2,000㎡', '装修补贴 100 元/㎡ + 免租 1 个月 + 服务中心标准价'],
        ['小型 / 创新',         '≤ 500㎡',    '免租 15 天 + 服务中心标准价'],
    ], widths_cm=[4.0, 3.0, 9.0])

    # ====== 九、典型客户案例 (脱敏) ======
    heading(doc, '九、典型客户案例 (脱敏)', level=1)
    for line in [
        '案例 1  某 AI 大模型公司 (中字头数字央企)  |  2,000㎡ × 2 个单元  |  4# 楼 5F+  |  来源:中国动漫集团牌照锚定',
        '案例 2  某 AI 潮玩品牌  |  1,500㎡  |  4# 楼 5F+  |  来源:5/22 峰会 + 追觅基金返投',
        '案例 3  某国漫 IP 公司  |  1,200㎡  |  5# 楼 5F+  |  来源:中动漫推荐 + 中百协对接',
        '案例 4  某 AI 设计公司  |  800㎡  |  4# 楼 5F  |  来源:AI 腾讯 + 上海交大联合推荐',
        '案例 5  盲盒新锐品牌集群  |  200-500㎡ × 5 家  |  5# 楼 5F+  |  来源:仲量联行爬楼大数据',
        '案例 6  设计师工作室集群  |  200-500㎡ × 10 家  |  4# 楼 5F+ + 5# 楼 5F+  |  来源:6 场沙龙转化',
        '案例 7  财税 / IP / 出海咨询服务机构  |  200-500㎡ × 15 家  |  4# 楼 5F+ + 5# 楼 5F+  |  来源:服务中心生态',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    doc.add_page_break()

    # ====== 十、招商联系 ======
    heading(doc, '十、招商联系', level=1)
    make_table(doc, ['模块', '负责人', '联系方式'], [
        ['产业招商 (2 万方主战场)',     '胡教授团队',          '__________ / __________'],
        ['商业招商 (5.2 万方)',         '森马商业部 (发哥 / 周志超)', '__________ / __________'],
        ['战略与合作',                  '威总 / 森马集团',       '__________ / __________'],
        ['政府对接',                    '胡教授 + 危总团队',      '__________'],
        ['项目地址',                    '上海市闵行区元江路 / 剑川路 (15 号线 元江路站 TOD)', ''],
    ], widths_cm=[5.0, 5.5, 6.0])

    add_para(doc, '本招商手册仅用于招商沟通, 不构成法律要约; 最终条款以正式租赁合同为准。', indent_first=False)

    doc.save(OUT)
    print(f'Wrote {OUT}')


if __name__ == '__main__':
    main()
