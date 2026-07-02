"""招商策略与排期 - 文字执行版 (Word, PPT 同步长文本)."""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name('招商策略与排期(文字执行版).docx')


def set_run(run, *, size=11, bold=False, color=None, font='微软雅黑'):
    run.font.name = font; run.font.size = Pt(size); run.font.bold = bold
    if color is not None: run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {}); rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font); rfonts.set(qn('w:ascii'), font); rfonts.set(qn('w:hAnsi'), font)


def add_para(doc, text, *, size=11, bold=False, color=None, align=None, indent_first=False, space_after=4, left_indent=0):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent: p.paragraph_format.left_indent = Pt(left_indent)
    if indent_first: p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text); set_run(run, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 20, 2: 15, 3: 12}
    size = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 4)
    run = p.add_run(text); set_run(run, size=size, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))
    return p


def make_table(doc, headers, rows, widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.autofit = False
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for cell in t.columns[i].cells: cell.width = Cm(w)
    hdr = t.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = h
        for p in hdr[j].paragraphs:
            for r in p.runs: set_run(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc_pr = hdr[j]._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '142C5E'); tc_pr.append(shd)
    for i, row in enumerate(rows, start=1):
        cells = t.rows[i].cells
        for j, v in enumerate(row):
            cells[j].text = v
            for p in cells[j].paragraphs:
                for r in p.runs: set_run(r, size=10)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return t


def main():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = '微软雅黑'; s.font.size = Pt(11)
    sec = doc.sections[0]; sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.0); sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # 封面
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(50)
    r = title.add_run('元谷项目 4#+5# 楼 2 万方'); set_run(r, size=30, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))
    title.add_run('\n'); r2 = title.add_run('IP + AI 双轨招商策略与排期')
    set_run(r2, size=26, bold=True, color=RGBColor(0xF2, 0x7E, 0x2D))
    title.add_run('\n'); r3 = title.add_run('文 字 执 行 版')
    set_run(r3, size=20, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('与 PPT 同步的长文本执行手册 · 团队内部派工'); set_run(r, size=14, color=RGBColor(0x55, 0x60, 0x7A))
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(30)
    for line in ['v1.0  ·  2026/6  ·  胡教授团队 × 森马集团',
                 '本文档与 PPT 内容完全一致, 可作为团队内部派工 / 培训 / 客户深度对话使用']:
        r = meta.add_run(line); set_run(r, size=11, color=RGBColor(0x55, 0x60, 0x7A)); meta.add_run('\n')
    doc.add_page_break()

    # 目录
    heading(doc, '目  录', level=1)
    toc = [
        '0. 项目背景与战略转向',
        '一、执行摘要',
        '二、项目实测',
        '三、SWOT 分析',
        '四、Phase 1:策略与定位',
        '五、Phase 2:招商执行 (重头戏)',
        '六、Phase 3:品牌与活动',
        '七、Phase 4:商业条款',
        '八、Phase 5:落地推进',
        '九、月度签约率推进表',
        '十、12 月费用与现金流',
        '十一、风险与对冲',
        '十二、投决建议',
    ]
    for t in toc:
        add_para(doc, t, indent_first=False, left_indent=30)
    doc.add_page_break()

    # ====== 0. 项目背景与战略转向 ======
    heading(doc, '0. 项目背景与战略转向', level=1)
    add_para(doc, '本方案产生于 2026 年 6 月的项目战略复盘会议, 与会方包括威总 (森马集团战略主导)、发哥 (商业部总牵头)、周志超 (商业部)、发言人 1 (产业部执行人)、胡教授 (产业部运营顾问/科技企业服务中心负责人) 等。', indent_first=True)
    add_para(doc, '会议核心达成的共识:', indent_first=False, bold=True)
    for line in [
        '★ 一、战略主轴由 "纯二次元" 升级为 "IP + AI 双轨"。理由是: 二次元企业规模过小 (一般 200-500㎡)、支付能力弱、难获政府政策与创新基金, 与森马商业难以协同。改为 IP+AI 双轨后, AI 主轴负责拿政府政策、拿基金、拿税收贡献, IP 副轴 兼顾威总战略文化导向且杨浦区已有 1 万方 AI+IP 园区先例。',
        '★ 二、明确两个硬节点: 2026/9/30 完成 2,000㎡ 签约 (可含直播基地、共享设计中心擦边球); 2027/5/1 项目开业且 2 万方达到 50%+ 签约 + 消费氛围。',
        '★ 三、招商客户策略调整: 优先招 2,000-20,000㎡ 大客户, 行业转化率新基准 110:1 (招 1 单需见 110 家潜客)。',
        '★ 四、深度绑定闵行区政府: 必须接入闵行区科委 / 商务委 / 街道 三轨, 否则招商节奏会慢 50% 以上。',
        '★ 五、引入资源升级: 腾讯算力补贴 + CVC 创业投资基金 + 欧洲驻沪机构 + 福布斯产业影响力榜 + 复旦大学住房政策研究中心 + 上海市科技企业联合会。',
        '★ 六、确认 5 项产业牌照规划: AI 潮玩产业基地 + 潮玩次元商业专委会 + 复旦分中心 + 上海市科企联 + 福布斯, 一次性挂牌奖励 5 × 30 万 = 150 万。',
    ]:
        add_para(doc, line, indent_first=False, left_indent=10, space_after=6)

    doc.add_page_break()

    # ====== 一、执行摘要 ======
    heading(doc, '一、执行摘要', level=1)
    add_para(doc, '元谷项目 4# 楼 5F+ 和 5# 楼 5F+ 共约 2 万㎡产业研发办公, 是元谷整盘 22 万㎡中产业心脏。本方案以 IP+AI 双轨为战略主轴, 通过基础月费 + 招商佣金 + 5 项挂牌 + 6 场沙龙的组合, 在 24 个月内实现满租, 24 月双向账本:甲方付出 ≈ 714 万元, 甲方收入 ≈ 2,769 万元, 投入产出比 1:3.88, 同时为森马贡献年化租金 1,606 万元 + 资产估值增量 2.47 亿元。', indent_first=True)

    make_table(doc, ['维度', '关键数据'], [
        ['业务范围', '4# 楼 5F+ 约 1 万㎡ + 5# 楼 5F+ 约 1 万㎡, 合计约 2 万㎡产业研发办公'],
        ['战略主轴', '★ IP + AI 双轨 (AI 拿政策/基金/税收, IP 兼顾战略文化导向)'],
        ['硬节点 1',  '2026/9/30 → 2,000㎡ 签约 (T+100 天)'],
        ['硬节点 2',  '2027/5/1 → 项目开业 (T+314 天, 50%+ 签约)'],
        ['全周期目标', '2028/5/1 → 90%+ 满租 (T+679 天)'],
        ['团队配置', '2 人驻场 + CSO 顾问 (轻骑兵)'],
        ['基础月费', '12 万元/月 (区间 6-18 万, 推荐 2 人)'],
        ['招商佣金', '1.5 / 1.75 / 2.0 个月年租金 (按面积阶梯)'],
        ['挂牌奖励', '30 万元/项 × 5 项 = 150 万元'],
        ['沙龙费',   '5 万元/场 × 6 场 = 30 万元'],
        ['超额奖',   '24 月满租率 ≥ 95% → 100 万一次性'],
        ['ROI',     '1 : 3.88 (24 个月)'],
    ], widths_cm=[3.5, 12.5])

    doc.add_page_break()

    # ====== 二、项目实测 ======
    heading(doc, '二、项目实测', level=1)

    heading(doc, '2.1 楼宇结构', level=2)
    add_para(doc, '元谷整盘 22 万㎡, 由 1#-6# 共 6 栋建筑构成:', indent_first=True)
    make_table(doc, ['楼栋', '1-3F', '4F', '5F+'], [
        ['1# 楼', '零售', '零售', '森马集团总部办公'],
        ['2# 楼', '二次元主题 Livehouse / 秀场', '—', '—'],
        ['3# 楼', '1-4F 休闲运动 / 萌宠空间', '休闲', '酒店'],
        ['★ 4# 楼', '潮玩艺术中心', '直播中心', '★ 潮玩产业集群 (本方案 1 万㎡)'],
        ['★ 5# 楼', '动漫书店 / 休闲娱乐', '动漫延伸', '★ 潮玩产业集群 (本方案 1 万㎡)'],
        ['6# 楼', '1-5F 品质生活 / 特色餐饮 / 服务配套 / 商务宴请', '—', '—'],
    ], widths_cm=[2.5, 6.5, 3.5, 6.5])

    heading(doc, '2.2 区位与客流', level=2)
    add_para(doc, '元谷位于上海市闵行区元江路 - 剑川路核心区, 享受三重区位优势:', indent_first=True)
    for line in [
        '★ TOD 优势: 15 号线元江路站 TOD 项目上盖, 单日客流 5-7 万人次',
        '★ 区位优势: 距紫竹高新区 5 km (车程 10 min), 距虹桥机场 17 km (机场快线 45 min), 距浦东机场 36 km',
        '★ 人口覆盖: 15 分钟车行覆盖 24 万居住人口 + 12 万产业办公人口',
        '★ 板块定位: 闵行区五大中心之一, 元江路 - 剑川路地区中心, 比肩漕河泾、张江',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    heading(doc, '2.3 市场租金对标', level=2)
    add_para(doc, '基于 2024-2026 年大零号湾 / 紫竹高新区主流园区公开放租信息:', indent_first=True)
    make_table(doc, ['对标园区', '日租金 (元)', '类别'], [
        ['零号湾全球创新创业集聚区', '2.0-2.5', '标杆'],
        ['大零号湾科创成果转化中心', '2.0-2.5', '标杆'],
        ['华谊万创新所',          '2.2',     '主流'],
        ['上海人工智能产业园',     '2.2',     '主流'],
        ['紫竹信息数码港 (5A 甲级)', '2.1-2.5', '主流'],
        ['紫竹数字创意港',         '2.0-3.0', '★ 文创对标'],
        ['龙湖蓝海引擎 / 金领谷',   '1.5-2.8', '可对标'],
        ['云境 443 / 夏日汇 (高端)', '2.3-4.5', '上限'],
        ['★ 元谷 4# 楼 5F+ (AI 主轴)', '2.2-2.5', '中位偏上'],
        ['★ 元谷 5# 楼 5F+ (IP 主轴)', '2.0-2.2', '中位'],
    ], widths_cm=[6.0, 3.0, 3.0])
    add_para(doc, '满租推算 (甲方视角): 2 万㎡ × 365 天 × 2.2 元/㎡/天 ≈ 1,606 万元/年。', indent_first=True, bold=True)

    doc.add_page_break()

    # ====== 三、SWOT ======
    heading(doc, '三、SWOT 矩阵', level=1)
    make_table(doc, ['', '内部 (Internal)', '外部 (External)'], [
        ['有利',
         'S 优势:森马底盘 + 双牌照已确认 + 4#/5# 楼产业配套齐全 (选品/直播/AI 设计/展厅/艺术中心) + 胡教授团队市级科委资源',
         'O 机会:AI+潮玩双赛道政策红利 + 5/22 峰会高净值人脉 + 复旦/北大/上海交大学术资源可平移 + 杨浦区先例'],
        ['不利',
         'W 劣势:新建项目品牌势能尚弱 + TOD 工程未完工 + 潮玩生态需培育 + 团队部分成员无产业招商经验',
         'T 威胁:周边新园区 (云境 443 / 夏日汇) 竞争 + 政策调整 + 招商节奏不及预期 + 转化率 110:1 严峻'],
    ], widths_cm=[3.0, 6.5, 6.5])

    add_para(doc, '应对策略 (写进协议):', indent_first=False, bold=True)
    for line in [
        '· 用『基础月费保底 + 阶梯佣金』对冲 W (招商节奏风险)',
        '· 用『5/22 峰会 + 6 场沙龙 + 5 项挂牌』转化 O (政策红利) 为实际客户',
        '· 用 AI 潮玩产业基地『差异化定位』对冲 T (周边竞争)',
        '· 用『团队 2 人 + CSO 顾问 + 兼职专家池』对冲 W (经验不足风险)',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    # ====== 四、Phase 1 策略与定位 ======
    heading(doc, '四、Phase 1:策略与定位', level=1)

    heading(doc, '4.1 总定位 — AI 潮玩产业基地', level=2)
    add_para(doc, '元谷 4# + 5# 楼 5F+ 共 2 万㎡产业研发办公, 总定位为『AI 潮玩产业基地 — 长三角首个 AI+潮玩 双牌照产业策源高地』, 以中国动漫集团『AI 潮玩产业基地』+ 中国百货商业协会『潮玩次元商业专委会』为核心牌照, 与上海交大设计学院 / 闵行科协 / 森马集团 联合建设。', indent_first=True)

    heading(doc, '4.2 楼宇分工', level=2)
    make_table(doc, ['楼栋', '昵称', '主轴', '同栋协同配套', '目标客户'], [
        ['4# 楼 5F+', '国际 IP 创意层', '★ AI 主轴', '4F 直播 + 5F AI 共享设计 + 5F DIY 中心', 'AI 大模型 / AIGC / AI 设计 / 国际 IP'],
        ['5# 楼 5F+', '潮玩产业总部层', '★ IP 主轴', '5F 潮玩展厅 + 1-4F 动漫书店', '头部央企 + 中型潮玩 / 内容 / 动漫'],
    ], widths_cm=[3.0, 3.0, 2.5, 4.5, 4.0])

    heading(doc, '4.3 客户五档配比', level=2)
    add_para(doc, '严格按照森马原产业规划的五档配比执行:', indent_first=True)
    make_table(doc, ['占比', '类型', '数量 / 户型', '总面积'], [
        ['10%', '头部央企 / 行业协会 (导向)', '3 家 / 2,000㎡',   '6,000㎡ → 主要在 5# 楼'],
        ['10%', '共享配套服务体系 (吸附点)',  '3 家 / 700㎡ 均值', '2,100㎡ → 4# 楼 + 5# 楼'],
        ['20%', '中型潮玩运营企业 (基础)',    '4-6 家 / 1,000㎡', '4,000-6,000㎡ → 主要在 5# 楼'],
        ['20%', '中小型潮业服务机构 (血肉)',  '15 家 / 300㎡',    '4,500㎡ → 主要在 4# 楼'],
        ['40%', '小型潮玩运营企业 (骨架)',    '30 家 / 200-500㎡', '6,000-15,000㎡ → 4# 楼 + 5# 楼'],
    ], widths_cm=[1.5, 5.5, 3.5, 5.0])

    doc.add_page_break()

    # ====== 五、Phase 2 招商执行 ======
    heading(doc, '五、Phase 2:招商执行 (重头戏)', level=1)

    heading(doc, '5.1 四级招商漏斗', level=2)
    make_table(doc, ['层级', '招商动作', '核心资源', '目标产能', '周期'], [
        ['L1 牌照锚定',   '签头部央企 / 行业协会作为旗舰',  'AI 潮玩产业基地 + 潮玩次元专委 双牌照',   '3 家 × 2,000㎡ = 6,000㎡', 'T+3 月'],
        ['L2 资本招商',   '追觅基金返投 + AI 腾讯生态导流', '追觅 + 5/22 峰会 LP 资源 + CVC 基金',     '4-6 家 × 1,000㎡ = 5,000㎡', 'T+6 月'],
        ['L3 大数据爬楼', '仲量联行爬楼数据 + 上门拜访',    '仲量联行爬楼数据 (¥2.6 万已购入)',         '30 家 × 200-500㎡ = 6,000-8,000㎡', 'T+9 月'],
        ['L4 活动带流',   '6 场沙龙 + 福布斯 + 北欧外事',    '5/22 峰会 + 6 场沙龙',                    '15 家 × 200-500㎡ = 3,000-4,000㎡', 'T+12 月'],
    ], widths_cm=[3.0, 5.0, 5.0, 4.0, 2.0])

    heading(doc, '5.2 5/22 AI 商业化峰会借势路径', level=2)
    add_para(doc, '2026 年 5 月 22 日由北京大学经济学院上海校友会 + 复旦大学住房政策研究中心主办的『2026 人工智能商业化落地与硬核投资破局峰会』, 是元谷品牌势能的最大借势机会, 5 件事必做:', indent_first=True)
    for line in [
        '① 元谷设峰会专属展位 + 闭门 1V1 招商台 (现场锁定 30-50 位 LP / 潜在客户)',
        '② 颁奖嵌入 — 增设『AI 潮玩产业影响力榜』, 由元谷主办、福布斯背书',
        '③ 主办方资源平移 — 复旦住房政策中心、北大经济学院上海校友会、中行/长江/招行/金浦/铂帝 战略合作',
        '④ 峰会嘉宾转化沙龙 #1 主题嘉宾 — 直接转化为元谷 T+1 月首场 AI+潮玩沙龙嘉宾',
        '⑤ 峰会闭门晚宴 → 形成『元谷创始合伙人圈』 → 12 个月后归入元谷年度大会',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    heading(doc, '5.3 招商六大『出彩点』 (VS 普通招商)', level=2)
    for line in [
        '①『牌照即招商』— AI 潮玩产业基地 + 潮玩次元专委 双牌照前置, 客户『送进来』不是『拉进来』',
        '②『基金即招商』— 追觅基金 1:1.5 返投绑定, 资本驱动签约 (不只是租约, 还有股权)',
        '③『数据即招商』— 仲量联行爬楼数据 (¥2.6 万已购) → 200 家精准客户清单 + 转化率提升 30%',
        '④『峰会即招商』— 5/22 AI 峰会 200+ VIP, 1 天锁定 30-50 客户',
        '⑤『沙龙即招商』— 6 场产业沙龙, 每场 30+ 目标客户, 一场签 1-2 家 = 5-12 家直接成果',
        '⑥『学术即招商』— 复旦 + 上海交大 + 北大 三所高校背书招商',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)
    add_para(doc, '→ 普通园区只做 ③④, 元谷六维齐发 → 招商速度领先 6-12 个月, 单方租金溢价 0.2-0.4 元/㎡/天。', indent_first=True, bold=True, color=RGBColor(0xF2, 0x7E, 0x2D))

    heading(doc, '5.4 9/30 节点深拆 — 2,000㎡ 怎么来', level=2)
    make_table(doc, ['来源', '面积估算', '落地概率', '兜底方案'], [
        ['1-2 家中大型 AI/IP 客户',     '1,000-2,000㎡',   '60%',   '若未签, 直播 + 设计中心兜底'],
        ['共享直播基地落地',            '800-1,000㎡',     '★ 90%', '森马已具备意愿, 7 月启动硬件'],
        ['AI 共享设计中心落地',          '500-800㎡',      '★ 80%', '联动腾讯/上海交大, 8 月挂牌'],
        ['产业服务中心入驻',            '300-500㎡',      '100%',  '★ 乙方自营业态, 100% 计入'],
        ['首批小型 AI/IP 客户 3-5 家',   '600-1,500㎡',    '50%',   'L4 沙龙带流的快速转化'],
    ], widths_cm=[5.5, 3.0, 2.0, 5.5])

    doc.add_page_break()

    # ====== 六、Phase 3 品牌与活动 ======
    heading(doc, '六、Phase 3:品牌与活动', level=1)

    heading(doc, '6.1 6 场产业沙龙', level=2)
    add_para(doc, '每场沙龙到场目标产业客户不少于 30 家, 每场触达 300-500 人, 媒体声量 ≥ 100 万次曝光:', indent_first=True)
    make_table(doc, ['#', '时间', '主题', '联办方'], [
        ['#1', 'T+1 月 (2026/7)',  'AI + 潮玩 跨界融合 (借势 5/22 峰会)', '中动漫 + AI 腾讯'],
        ['#2', 'T+3 月 (2026/9)',  '潮玩出海 (北欧 / 日韩 / 东南亚)',      '北欧创新国际会客厅 + 福布斯'],
        ['#3', 'T+5 月 (2026/11)', '投融资路演',                          '追觅 + 招商银行 + 长江证券 + 金浦'],
        ['#4', 'T+7 月 (2027/1)',  '设计与创意',                          '上海交大设计学院 + 上海市科企联'],
        ['#5', 'T+9 月 (2027/3)',  '内容 IP 与 Z 世代',                  '中百协潮玩次元专委 + 中动漫'],
        ['#6', 'T+11 月 (2027/5)', '政策补贴与小镇',                      '闵行科协 + 复旦住房政策中心'],
    ], widths_cm=[1.5, 4.0, 5.5, 5.0])

    heading(doc, '6.2 5 项挂牌', level=2)
    make_table(doc, ['#', '挂牌名称', '出牌方', '落地节点', '单项奖励'], [
        ['①', 'AI 潮玩产业基地',                '中国动漫集团',           'T+3 月',  '30 万'],
        ['②', '潮玩次元商业专委会',             '中国百货商业协会',         'T+3 月',  '30 万'],
        ['③', '复旦大学住房政策研究中心 元谷分中心', '复旦大学住房政策研究中心',  'T+9 月',  '30 万'],
        ['④', '上海市科技企业联合会 元谷产业基地', '上海市科技企业联合会',     'T+6 月',  '30 万'],
        ['⑤', '福布斯产业影响力奖 元谷专场',     '福布斯',                  'T+12 月', '30 万'],
        ['',  '5 项合计',                       '',                       '',        '150 万'],
    ], widths_cm=[1.0, 6.0, 4.5, 2.0, 2.5])

    # ====== 七、Phase 4 商业条款 ======
    heading(doc, '七、Phase 4:商业条款', level=1)

    heading(doc, '7.1 基础月费 — 1-3 人配置阶梯', level=2)
    make_table(doc, ['配置', '人员', '月度刚性成本', '建议月费', '毛利率'], [
        ['1 人轻配',           '招商单兵 + CSO 顾问 (折半)', '45,000 元', '60,000 元/月', '25%'],
        ['★ 2 人推荐配',       '招商 + 活动策划 + CSO 顾问 + 数据接口', '95,000 元', '120,000 元/月', '21%'],
        ['3 人重配',           '招商 + 活动 + 基金投后/政府关系 + CSO', '140,000 元', '180,000 元/月', '22%'],
    ], widths_cm=[3.5, 6.0, 2.5, 2.5, 1.5])
    add_para(doc, '★ 推荐方案:2 人配置 + 12 万元/月 → 24 个月合计基础月费 288 万元。覆盖刚性成本 9.5 万 + 安全垫 2.5 万 (毛利率 21%)。', indent_first=True, bold=True, color=RGBColor(0xF2, 0x7E, 0x2D))

    heading(doc, '7.2 招商佣金阶梯', level=2)
    make_table(doc, ['面积档位', '日租金', '年租金', '佣金月数', '佣金 (元/㎡)'], [
        ['≤ 2,000㎡ 小型',        '2.0', '730 元/㎡',   '1.5 个月',  '91'],
        ['2,001-5,000㎡ 中型',    '2.2', '803 元/㎡',   '1.75 个月', '117'],
        ['> 5,000㎡ 头部 / 央企', '2.5', '913 元/㎡',   '2.0 个月',  '152'],
    ], widths_cm=[4.0, 2.5, 3.0, 2.5, 2.5])
    add_para(doc, '★ 满租 2 万㎡ × 按 1.75 个月加权平均 → 佣金累计 ≈ 234 万元 (24 个月)。返投基金客户额外加成 +0.25 个月。', indent_first=True, color=RGBColor(0xF2, 0x7E, 0x2D))

    heading(doc, '7.3 KPI 与对赌', level=2)
    for line in [
        '· T+6 月签约 ≥ 4,000㎡ / T+12 月 ≥ 10,000㎡ (4# 楼基本满租) / T+24 月 ≥ 18,000㎡ (90% 满租视为达标)',
        '· 沙龙:每场 ≥ 30 目标产业客户 (现场签到 + 闭环管理)',
        '· 挂牌:T+12 月 ≥ 3 项 / T+24 月 ≥ 5 项',
        '· 未达标的, 当期月费保留 50%, 待达成后补足',
        '· 超额奖励:24 月满租率 ≥ 95% → 一次性奖励 100 万',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    doc.add_page_break()

    # ====== 八、Phase 5 落地推进 ======
    heading(doc, '八、Phase 5:落地推进', level=1)
    make_table(doc, ['节点', '关键动作', '里程碑产出'], [
        ['30 天 (5/22 当周内)', 'MOU 签订 + 2-3 人战队就位 + 仲量联行数据接入 + 5/22 峰会借势完成', '首批 20 家意向客户清单'],
        ['60 天',               '沙龙 #1 (AI+潮玩) 完成 + AI 潮玩产业基地挂牌筹备', '签约首单 + 媒体头部曝光'],
        ['90 天',               'AI 潮玩产业基地正式挂牌 + 潮玩次元专委挂牌 + 4,000㎡ 签约', 'L1 牌照客户 3 家入驻'],
        ['180 天',              '沙龙过半 (3 场) + 上海市科企联挂牌 + 8,000㎡ 签约', 'L2 资本招商客户落地'],
        ['365 天',              '6 场沙龙完成 + 5 项挂牌全部落地 + 4# 楼 1 万方满租', '首年报告 + 续约谈判'],
        ['730 天',              '5# 楼 1 万方满租 + 沙龙年度 IP 化 + 福布斯榜单元谷专场上线', '★ 2 万方目标完成'],
    ], widths_cm=[3.5, 7.5, 5.0])

    add_para(doc, '团队配置 (2 人核心 + CSO + 兼职专家池):', indent_first=False, bold=True)
    for line in [
        '· 产业招商经理 (全职):22K 底薪 + 10K 绩效 + 招商提成',
        '· 国际合作 & 活动策划 (全职):20K 底薪 + 8K 绩效 + 沙龙分润',
        '· CSO (胡教授本人):每周 ≥ 2 个工作日 + 重大节点全勤, 月顾问费 25K',
        '· 兼职专家池 (项目制):财税 / 法律 / 知识产权 / 政府补贴顾问',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    # ====== 九、月度签约率推进 ======
    heading(doc, '九、月度签约率推进表 (按 803 元/㎡ 估算)', level=1)
    make_table(doc, ['月份', '本月新增 (㎡)', '累计签约 (㎡)', '累计签约率', '累计年租金 (元)', '里程碑'], [
        ['2026/7',  '300',   '300',     '1.5%',  '240,900',    '战队就位 + 物料'],
        ['2026/8',  '800',   '1,100',   '5.5%',  '883,300',    'L1 牌照 + 大客户接触'],
        ['2026/9',  '1,200', '2,300',   '11.5%', '1,846,900',  '★ 9/30 硬节点达成'],
        ['2026/10', '1,000', '3,300',   '16.5%', '2,649,900',  'L3 爬楼启动'],
        ['2026/11', '1,000', '4,300',   '21.5%', '3,452,900',  '沙龙 #3 投融资'],
        ['2026/12', '1,200', '5,500',   '27.5%', '4,416,500',  'Q4 收官'],
        ['2027/1',  '1,200', '6,700',   '33.5%', '5,380,100',  '新年首单'],
        ['2027/2',  '1,000', '7,700',   '38.5%', '6,183,100',  '—'],
        ['2027/3',  '1,200', '8,900',   '44.5%', '7,146,700',  '—'],
        ['2027/4',  '1,500', '10,400',  '★ 52%', '8,351,200',  '★ 5/1 节点'],
        ['2027/5',  '1,500', '11,900',  '59.5%', '9,555,700',  '开业典礼'],
        ['2027/6',  '1,300', '13,200',  '66.0%', '10,599,600', '—'],
    ], widths_cm=[2.0, 2.5, 2.5, 2.5, 3.0, 3.5])

    # ====== 十、12 月费用与现金流 ======
    heading(doc, '十、12 月费用与现金流 (向森马提案口径)', level=1)
    make_table(doc, ['费用类别', '金额 (元)', '支付节奏', '森马承担'], [
        ['团队月费 12 万 × 12',          '1,440,000', '按月预付',          '100%'],
        ['招商佣金 (累计 11,900㎡)',     '1,400,000', '起租后 30 日内',     '100%'],
        ['5 项挂牌 (一次性)',            '1,500,000', '挂牌 30 日内',      '100%'],
        ['6 场沙龙执行',                  '300,000',  '按场结算',          '100%'],
        ['5/1 开业典礼',                  '800,000',  '开业前 30 日内',     '100%'],
        ['装修补贴 + 物料 + 设计',         '500,000',  '实报实销',          '100%'],
        ['接待 + 交通 + 政府关系',          '200,000',  '月度报销',          '100%'],
        ['直播 + AI 设计中心硬件',         '700,000',  '一次性',            '100%'],
        ['12 月合计',                    '6,840,000', '',                  '★ 684 万'],
    ], widths_cm=[5.5, 3.5, 3.5, 3.5])

    add_para(doc, '同期甲方现金回收: 租金累计 ≈ 956 万 + 物业 ≈ 90 万 = ≈ 1,046 万元。单 12 个月已现金回正 (1,046 万 vs 684 万)。', indent_first=True, bold=True, color=RGBColor(0xF2, 0x7E, 0x2D))

    doc.add_page_break()

    # ====== 十一、风险与对冲 ======
    heading(doc, '十一、风险与对冲', level=1)
    make_table(doc, ['风险', '影响', '对冲机制'], [
        ['9/30 节点 2,000㎡ 未达',          '森马失信 + 项目延期',   '★ 擦边球 (直播 + 设计中心) 兜底 + KPI 对赌月费保留 50%'],
        ['5/1 开业 50% 未达',               '开业氛围不足',         '★ 阶段 4 满租期延长 + 沙龙 IP 化加速'],
        ['110:1 → 130:1 (恶化)',           '招商节奏放缓 30%',     '★ 团队扩 1 人 (改 3 人重配) + 月费上调至 18 万'],
        ['闵行区政策变化',                  '补贴/牌照不落地',       '★ 备份:杨浦区 / 漕河泾 / 张江政策可援引'],
        ['AI 行业资本退潮',                 'AI 客户付款能力下降',   '★ 切换 IP 主轴比例 + 出海企业引入'],
        ['森马商业部协同延误',              '1#/6# 商业氛围不足',   '★ 5.2 万方与 2 万方独立排期'],
    ], widths_cm=[4.0, 4.5, 7.5])

    # ====== 十二、投决建议 ======
    heading(doc, '十二、投决建议', level=1)
    add_para(doc, '★ 三问决策 (建议森马本周内确定):', indent_first=False, bold=True)
    add_para(doc, '① 是否接受 IP+AI 双轨战略调整? — 建议:是 (杨浦区已有先例)', indent_first=False, left_indent=20)
    add_para(doc, '② 是否接受 2 人驻场 + 12 万元月费 + 1.5-2 月佣金 + 5 项挂牌 + 6 场沙龙条款? — 建议:是', indent_first=False, left_indent=20)
    add_para(doc, '③ 是否同意 7 月初启动 (赶 9/30 硬节点)? — 建议:是, 错过 7 月即风险陡增', indent_first=False, left_indent=20)
    add_para(doc, '', indent_first=False)
    add_para(doc, '★ 一句话总结:', indent_first=False, bold=True)
    add_para(doc, '以 IP+AI 双轨战略, 借势 5/22 峰会余热与闵行区科委政策红利,', indent_first=True)
    add_para(doc, '100 天达成 9/30 节点 2,000㎡, 314 天达成 5/1 开业 50% 签约,', indent_first=True)
    add_para(doc, '730 天达成 2 万方满租, 为森马贡献年化租金 1,606 万元 + 资产估值 2.47 亿元。', indent_first=True)
    add_para(doc, '', indent_first=False)
    add_para(doc, '★ 决策成本对照:', indent_first=False, bold=True)
    for line in [
        '─ 12 个月投入: ≈ 684 万元 (月费 + 佣金 + 挂牌 + 沙龙 + 开业)',
        '─ 12 个月收入: ≈ 1,046 万元 (租金 + 物业)',
        '─ 12 个月净收益: ≈ +362 万元 (单 12 月已正)',
        '─ 24 月 ROI: 1:3.88',
        '─ 永续资产估值: 2.47 亿元 (8% 折现)',
    ]:
        add_para(doc, line, indent_first=False, left_indent=15)

    doc.save(OUT)
    print(f'Wrote {OUT}')


if __name__ == '__main__':
    main()
