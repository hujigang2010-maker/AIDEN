"""Build the 4#+5# 楼 2 万方专项合作协议 .docx draft.

Single-quoted Python strings throughout so Chinese text can contain
ASCII " freely.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name('4#5#楼专项合作协议(草案).docx')


def set_run(run, *, size=11, bold=False, color=None, font_name='微软雅黑'):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font_name)
    rfonts.set(qn('w:ascii'), font_name)
    rfonts.set(qn('w:hAnsi'), font_name)


def add_para(doc, text, *, size=11, bold=False, color=None, align=None,
             indent_first=True, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if indent_first:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    size = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 4)
    run = p.add_run(text)
    set_run(run, size=size, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))
    return p


def article(doc, num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'第{num}条　{title}')
    set_run(run, size=13, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))


def clause(doc, marker, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{marker}　{text}')
    set_run(run, size=11)


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # --- Cover ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(40)
    r = title.add_run('元谷项目 4#+5# 楼约 2 万方产业研发办公')
    set_run(r, size=22, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))
    title.add_run('\n')
    r2 = title.add_run('专 项 招 商 运 营 合 作 协 议(草案)')
    set_run(r2, size=22, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Buildings 4 & 5 Industry R&D Office Leasing & Operation Cooperation Agreement (Draft)')
    set_run(r, size=12, color=RGBColor(0x55, 0x60, 0x7A))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(36)
    for line in [
        '项目名称:森马(上海)国际运营中心 元谷项目',
        '业务范围:4# 楼 5F+ + 5# 楼 5F+ 共约 2 万㎡产业研发办公',
        '协议版本:v1.0  ·  起草方:胡教授团队',
        '签署日期:__________ 年 ____ 月 ____ 日',
    ]:
        r = meta.add_run(line)
        set_run(r, size=12)
        meta.add_run('\n')

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(40)
    r = note.add_run('本协议为商务谈判草案,不构成法律约束力,最终以双方法务确认并签署的正式文本为准。')
    set_run(r, size=10, color=RGBColor(0x99, 0x33, 0x33))

    doc.add_page_break()

    # --- 缔约方 ---
    heading(doc, '缔约方')
    add_para(doc, '甲方(委托方/业主方):森马集团股份有限公司(以下简称"甲方"或"森马")')
    add_para(doc, '    住所地:__________ ; 统一社会信用代码:__________ ; 法定代表人:__________')
    add_para(doc, '乙方(招商运营服务方):__________ (以下简称"乙方"或"胡教授团队")')
    add_para(doc, '    住所地:__________ ; 统一社会信用代码:__________ ; 法定代表人:胡教授')

    add_para(doc, '鉴于:', bold=True, indent_first=False)
    clause(doc, '(1)', '甲方为元谷项目(位于上海市闵行区元江路-剑川路地区中心,"大零号湾"文创融合核心区)的资产持有方与品牌方,项目总建面约 22 万㎡, 商业建面约 5.2 万㎡;')
    clause(doc, '(2)', '本协议项下"项目范围"特指元谷项目 4# 楼 5F 及以上潮玩产业集群 (约 1 万㎡) 与 5# 楼 5F 及以上潮玩产业集群 (约 1 万㎡), 合计约 2 万㎡产业研发办公空间(以下简称"项目范围");')
    clause(doc, '(3)', '乙方在产业招商、政府关系、文创活动 IP 及高端学术资源方面拥有可投入的稀缺资源,包括但不限于:北欧创新国际会客厅、福布斯系列奖项、科技开放麦、AI 腾讯生态、仲量联行爬楼大数据(已实际投入采购成本人民币 26,000 元)、追觅科技基金、复旦大学住房政策研究中心、上海市科技企业联合会等;')
    clause(doc, '(4)', '甲方拟将项目范围内的招商运营业务委托乙方独家承担, 双方就此达成如下协议:')

    # --- 第一条 业务范围 ---
    article(doc, '一', '业务范围与定位')
    clause(doc, '1.1', '业务范围(招商运营独家委托):')
    clause(doc, '      (a)', '元谷项目 4# 楼 5F 及以上潮玩产业集群约 1 万㎡;')
    clause(doc, '      (b)', '元谷项目 5# 楼 5F 及以上潮玩产业集群约 1 万㎡;')
    clause(doc, '      (c)', '合计约 2 万㎡产业研发办公空间。')
    clause(doc, '1.2', '业务定位:双方一致将项目范围打造为 "AI 潮玩产业基地 — 长三角首个 AI+潮玩 双牌照产业策源高地", 配套设施包括: 4F 直播中心、4# 楼 5F AI 共享设计中心、4# 楼 5F AI 共享打样及 DIY 中心、5# 楼 5F 潮玩产业展厅。')
    clause(doc, '1.3', '独家性:在本协议存续期内, 甲方授予乙方对项目范围的独家招商运营权 (不可撤销); 甲方不得就项目范围委托任何第三方招商运营机构。')
    clause(doc, '1.4', '客户画像 (与甲方原产业规划一致):10% 头部央企/行业协会 (3 家 × 2,000㎡); 10% 共享配套服务 (3 家 × 2,000㎡); 20% 中型潮玩运营企业 (4-6 家 × 5,000㎡); 20% 中小型潮业服务机构 (15 家 × 200-500㎡); 40% 小型潮玩运营企业 (30 家 × 200-500㎡)。')

    # --- 第二条 服务期限 ---
    article(doc, '二', '服务期限')
    clause(doc, '2.1', '本协议自双方签署之日起生效, 期限为 24 个月。')
    clause(doc, '2.2', '届满前 90 日, 双方可协商续期; 若 24 个月内项目范围招商出租率达到 90% 以上, 乙方有权要求按相同或更优条款续期 24 个月。')
    clause(doc, '2.3', '本协议另有解除条款约定外, 任何一方不得单方提前终止。')

    # --- 第三条 基础服务费 ---
    article(doc, '三', '基础服务费 (月费 / Retainer)')
    clause(doc, '3.1', '人员配置约定:乙方按下列三档之一向项目范围派驻团队 (具体档位由双方在签署时勾选):')
    clause(doc, '      (a)', '1 人轻配:产业招商经理 × 1 + CSO 顾问 (折半投入);')
    clause(doc, '      (b)', '★ 2 人推荐配 (建议):产业招商经理 + 国际合作&活动策划 + CSO 顾问 (全额投入);')
    clause(doc, '      (c)', '3 人重配:招商 + 活动 + 基金投后&政府关系 + CSO 顾问。')
    clause(doc, '3.2', '基础月费金额:')
    clause(doc, '      (a)', '1 人轻配:人民币 6 万元/月;')
    clause(doc, '      (b)', '★ 2 人推荐配:人民币 12 万元/月 (推荐选择);')
    clause(doc, '      (c)', '3 人重配:人民币 18 万元/月。')
    clause(doc, '3.3', '支付节奏:每自然月 5 日前预付当月月费; 逾期支付的, 每逾期一日按当月月费的 0.05% 计违约金。')
    clause(doc, '3.4', '调整机制:经双方一致书面同意, 每 12 个月可对月费金额进行一次评审与调整。')
    clause(doc, '3.5', '基础月费覆盖范围:乙方派驻人员薪酬及绩效、CSO 顾问费、仲量联行爬楼大数据接口运维费、必要差旅与接待物料、月度管理协调费;不包含招商佣金、挂牌奖励、沙龙执行费及其他单项收费。')

    # --- 第四条 招商佣金 ---
    article(doc, '四', '招商佣金 (1.5-2 个月年租金)')
    clause(doc, '4.1', '甲方同意:乙方每促成一份新签或续签的项目范围内租赁合同, 按下列阶梯向乙方支付招商佣金:')
    clause(doc, '      (a)', '≤ 2,000㎡ 的小型租户:佣金 = 实际成交年租金 × 1.5 个月;')
    clause(doc, '      (b)', '2,001-5,000㎡ 的中型租户:佣金 = 实际成交年租金 × 1.75 个月;')
    clause(doc, '      (c)', '> 5,000㎡ 的头部 / 央企 / 行业协会租户:佣金 = 实际成交年租金 × 2.0 个月。')
    clause(doc, '4.2', '返投基金加成:凡通过追觅科技基金或其他乙方导入的产业基金返投落地的租户, 前述基础佣金额外加成 0.25 个月。')
    clause(doc, '4.3', '结算节奏:自租赁合同签订且租户起租之日起 30 日内一次性支付; 合同中途解除的, 按已履约期限对应比例处理。')
    clause(doc, '4.4', '客户归属保护:在乙方向甲方提交的"招商客户名单"内的客户, 自该客户首次接触起 24 个月内成交的, 均视为乙方业绩; 甲方不得绕开乙方签约, 否则按本条 4.1 标准的 1.5 倍向乙方赔付。')

    # --- 第五条 挂牌奖励 ---
    article(doc, '五', '挂牌奖励 (5 项, 一次性激励)')
    clause(doc, '5.1', '甲方同意:乙方主导促成下列任一挂牌正式落地的, 甲方向乙方一次性支付奖励人民币 30 万元/项:')
    clause(doc, '      (a)', '"AI 潮玩产业基地" 牌照 (中国动漫集团);')
    clause(doc, '      (b)', '"潮玩次元商业专委会" 牌照 (中国百货商业协会);')
    clause(doc, '      (c)', '"复旦大学住房政策研究中心 · 元谷分中心" 挂牌;')
    clause(doc, '      (d)', '"上海市科技企业联合会 · 元谷产业基地" 挂牌;')
    clause(doc, '      (e)', '"福布斯产业影响力奖 · 元谷专场" 挂牌。')
    clause(doc, '5.2', '上限:5 项挂牌奖励合计上限为人民币 150 万元 (5 × 30 万); 续挂 (第 2 年起) 按 50% 收取, 即 15 万元/项。')
    clause(doc, '5.3', '结算节奏:自挂牌或获奖正式公告之日起 30 日内由甲方一次性支付。')
    clause(doc, '5.4', '附加奖项:除上述 5 项外, 若乙方促成其他省级及以上政府奖项 / 行业牌照落地, 双方可另行书面约定奖励金额。')

    # --- 第六条 沙龙执行 ---
    article(doc, '六', '产业沙龙执行 (6 场, 每场 ≥ 30 目标产业客户)')
    clause(doc, '6.1', '乙方承诺:于服务期内每 12 个月承办不少于 6 场产业沙龙 (合计不少于 12 场), 每场到场目标产业客户不少于 30 家 (以现场签到 + 闭环管理为准)。')
    clause(doc, '6.2', '六大主题 (首年):')
    clause(doc, '      (a)', '#1 AI + 潮玩 跨界融合 (T+1 月, 借势 5/22 AI 商业化峰会);')
    clause(doc, '      (b)', '#2 潮玩出海 (T+3 月, 北欧创新国际会客厅 + 福布斯);')
    clause(doc, '      (c)', '#3 投融资路演 (T+5 月, 追觅 + 招商银行 + 长江证券 + 金浦);')
    clause(doc, '      (d)', '#4 设计与创意 (T+7 月, 上海交大设计学院 + 上海市科企联);')
    clause(doc, '      (e)', '#5 内容 IP 与 Z 世代 (T+9 月, 中百协潮玩次元专委 + 中国动漫集团);')
    clause(doc, '      (f)', '#6 政策补贴与小镇 (T+11 月, 闵行科协 + 复旦住房政策研究中心)。')
    clause(doc, '6.3', '执行费:甲方按每场人民币 5 万元支付执行费 (覆盖场地、物料、嘉宾接待、传播), 6 场合计人民币 30 万元/年。')
    clause(doc, '6.4', '收入分润:沙龙赞助、票务、政府补贴等收入由乙方统一收取, 单场净利 (收入 - 执行成本) 按 甲方 30% / 乙方 70% 分润。')
    clause(doc, '6.5', '未达标处理:单场到场目标产业客户少于 30 家的, 当场执行费扣回 50%, 净利分润同步调整为 甲方 50% / 乙方 50%。')

    # --- 第七条 KPI 与对赌 ---
    article(doc, '七', 'KPI 与业绩对赌')
    clause(doc, '7.1', '招商 KPI (累计签约面积):')
    clause(doc, '      (a)', 'T+6 月 ≥ 4,000㎡;')
    clause(doc, '      (b)', 'T+12 月 ≥ 10,000㎡ (4# 楼基本满租);')
    clause(doc, '      (c)', 'T+24 月 ≥ 18,000㎡ (4#+5# 楼 90% 满租视为达标)。')
    clause(doc, '7.2', '挂牌 KPI:T+12 月 ≥ 3 项 / T+24 月 ≥ 5 项。')
    clause(doc, '7.3', '沙龙 KPI:首年 ≥ 6 场, 每场 ≥ 30 目标产业客户。')
    clause(doc, '7.4', '对赌处理:')
    clause(doc, '      (a)', 'KPI 未达标的, 自达标偏差超过 20% 的次月起, 月费保留 50% 至达成时补足;')
    clause(doc, '      (b)', '连续 6 个月未达标的, 甲方有权解除本协议, 乙方按已结算款返还 20%;')
    clause(doc, '      (c)', '若 24 个月项目范围满租率 ≥ 95%, 甲方向乙方一次性支付奖励人民币 100 万元 (超额完成奖)。')

    # --- 第八条 资源承诺 ---
    article(doc, '八', '资源承诺')
    clause(doc, '8.1', '甲方承诺:')
    clause(doc, '      (a)', '在本协议生效后 60 日内向乙方开放项目范围的全部招商主导权及客户对接通道;')
    clause(doc, '      (b)', '配合乙方完成 5 项挂牌的对接、文件、场地及森马品牌联合传播;')
    clause(doc, '      (c)', '为乙方派驻团队提供项目范围内的免费办公空间 (建议 4# 楼 5F 设乙方运营室);')
    clause(doc, '      (d)', '在森马集团及关联方资源范围内, 为乙方业务提供必要支持。')
    clause(doc, '8.2', '乙方承诺 (量化承诺, 列入考核):')
    clause(doc, '      (a)', '本协议生效后 30 日内, 完成 2-3 人驻场战队就位;')
    clause(doc, '      (b)', '90 日内将仲量联行爬楼大数据接入项目中台 (覆盖 ≥ 200 家目标客户);')
    clause(doc, '      (c)', '90 日内挂牌 "AI 潮玩产业基地" + "潮玩次元商业专委会" 两项;')
    clause(doc, '      (d)', '12 个月内完成 5 项挂牌 / 6 场沙龙 / 4# 楼基本满租 (≥ 10,000㎡);')
    clause(doc, '      (e)', '协调 2026 年 5 月 22 日 "AI 商业化落地与硬核投资破局峰会" 资源, 为项目范围导入不少于 30 位 VIP 潜在客户。')

    # --- 第九条 知识产权 ---
    article(doc, '九', '知识产权与资源排他性')
    clause(doc, '9.1', '乙方就以下资源在项目范围内的使用, 授予甲方独家、可转授的使用权 (在协议存续期内):北欧创新国际会客厅 IP、科技开放麦 IP、福布斯系列奖项渠道、AI 腾讯导入合作、仲量联行爬楼大数据接入、追觅科技基金合作、复旦大学住房政策研究中心合作、上海市科技企业联合会合作。')
    clause(doc, '9.2', '上述资源的所有权仍归原权利人;合资公司或本协议解除的, 使用权同步解除。')
    clause(doc, '9.3', '甲方对外发布及营销中应同步标注上述资源贡献方信息, 不得擅自利用上述资源开展项目范围外的业务。')

    # --- 第十条 排他与竞业 ---
    article(doc, '十', '排他与竞业禁止')
    clause(doc, '10.1', '本协议存续期内, 双方不得直接或间接从事与本协议业务存在实质性竞争的活动;不得另行委托或受托从事元谷项目范围内的招商运营。')
    clause(doc, '10.2', '甲方在大零号湾区域内的其他自有商业项目, 需先以书面形式询价于乙方, 乙方在 30 日内未明示接受的, 甲方有权另行委托。')

    # --- 第十一条 退出 ---
    article(doc, '十一', '退出与终止')
    clause(doc, '11.1', '本协议自然届满终止, 双方可协商续期。')
    clause(doc, '11.2', '解除事由:')
    clause(doc, '      (a)', '甲方严重违约或未按约支付月费、佣金、挂牌奖励, 经书面催告 30 日仍未纠正的, 乙方有权解除;')
    clause(doc, '      (b)', '乙方连续 6 个月 KPI 未达标的, 甲方有权解除;')
    clause(doc, '      (c)', '不可抗力或政策重大变更导致项目范围无法开展招商的, 双方可协商解除。')
    clause(doc, '11.3', '解除后处理:已完成 KPI 对应的款项不予返还; 未完成部分按已履约比例处理。')

    # --- 第十二条 保密 ---
    article(doc, '十二', '保密')
    clause(doc, '12.1', '双方就在本协议项下知悉的对方商业秘密、客户名单、技术信息、财务信息及未公开战略, 承担保密义务, 期限为本协议存续期间及解除/终止后 5 年。')

    # --- 第十三条 争议 ---
    article(doc, '十三', '法律适用与争议解决')
    clause(doc, '13.1', '本协议适用中华人民共和国法律。')
    clause(doc, '13.2', '凡因本协议引起的争议, 双方应首先友好协商解决; 协商不成的, 提交上海国际经济贸易仲裁委员会 (上海国际仲裁中心) 按其届时有效的仲裁规则在上海仲裁。')

    # --- 第十四条 一般 ---
    article(doc, '十四', '一般条款')
    clause(doc, '14.1', '本协议自双方法定代表人或授权代表签字并加盖公章之日起生效。')
    clause(doc, '14.2', '本协议一式四份, 双方各执两份, 具有同等法律效力。')
    clause(doc, '14.3', '本协议附件与正文具有同等法律效力。')
    clause(doc, '14.4', '本协议未尽事宜, 由双方另行书面约定。')

    # --- 签署页 ---
    doc.add_page_break()
    heading(doc, '签 署 页', level=1)
    sig = doc.add_table(rows=3, cols=2)
    sig.autofit = False
    for col in sig.columns:
        for cell in col.cells:
            cell.width = Cm(8)
    rows = [
        ('甲方:森马集团股份有限公司',          '乙方:__________(胡教授团队)'),
        ('法定代表人/授权代表:__________',     '法定代表人/授权代表:__________'),
        ('(盖章)　　　　　 日期:____ 年 ____ 月 ____ 日',
         '(盖章)　　　　　 日期:____ 年 ____ 月 ____ 日'),
    ]
    for i, (a, b) in enumerate(rows):
        sig.cell(i, 0).text = a
        sig.cell(i, 1).text = b
        for cell in sig.rows[i].cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for r in para.runs:
                    set_run(r, size=11)
                para.paragraph_format.space_after = Pt(8)

    # --- 附件 ---
    doc.add_page_break()

    heading(doc, '附件一:项目范围与楼宇分工', level=1)
    add_para(doc, '4# 楼 5F 及以上潮玩产业集群 (约 1 万㎡) — 国际 IP 创意层, 与同栋 4F 直播中心 / 5F AI 共享设计中心 / 5F AI 共享打样 DIY 中心 联动;', indent_first=False)
    add_para(doc, '5# 楼 5F 及以上潮玩产业集群 (约 1 万㎡) — 潮玩产业总部层, 与同栋 5F 潮玩产业展厅 / 1-4F 动漫书店 + 休闲娱乐 联动;', indent_first=False)
    add_para(doc, '合计约 2 万㎡产业研发办公空间, 招商 / 续约 / 租户管理 / 退租 等业务全部由乙方独家承担。', indent_first=False)

    heading(doc, '附件二:商业条款速查表', level=1)
    add_para(doc, '· 基础月费:6 / 12 / 18 万元/月 (1/2/3 人配置, 推荐 2 人 = 12 万元/月)', indent_first=False)
    add_para(doc, '· 招商佣金:1.5 / 1.75 / 2.0 个月年租金 (按面积阶梯), 返投基金客户额外 +0.25 个月', indent_first=False)
    add_para(doc, '· 挂牌奖励:30 万元/项 × 5 项 = 上限 150 万元 (首年挂); 续挂 50% (15 万元/项)', indent_first=False)
    add_para(doc, '· 沙龙执行费:5 万元/场 × 6 场 = 30 万元/年; 净利 30/70 分润 (甲/乙)', indent_first=False)
    add_para(doc, '· 服务期限:24 个月 (首期); 续期由乙方优先选择权', indent_first=False)
    add_para(doc, '· 超额奖励:24 月满租率 ≥ 95% → 一次性奖励 100 万元', indent_first=False)

    heading(doc, '附件三:5 项挂牌清单', level=1)
    for line in [
        '① AI 潮玩产业基地 (中国动漫集团)',
        '② 潮玩次元商业专委会 (中国百货商业协会)',
        '③ 复旦大学住房政策研究中心 · 元谷分中心',
        '④ 上海市科技企业联合会 · 元谷产业基地',
        '⑤ 福布斯产业影响力奖 · 元谷专场',
    ]:
        add_para(doc, line, indent_first=False)

    heading(doc, '附件四:6 场产业沙龙清单 (首年)', level=1)
    for line in [
        '#1 T+1 月  AI + 潮玩 跨界融合 (中动漫 + AI 腾讯; 借势 5/22 峰会次月)',
        '#2 T+3 月  潮玩出海 (北欧创新国际会客厅 + 福布斯)',
        '#3 T+5 月  投融资路演 (追觅 + 招商银行 + 长江证券 + 金浦)',
        '#4 T+7 月  设计与创意 (上海交大设计学院 + 上海市科企联)',
        '#5 T+9 月  内容 IP 与 Z 世代 (中百协潮玩次元专委 + 中动漫)',
        '#6 T+11 月 政策补贴与小镇 (闵行科协 + 复旦住房政策研究中心)',
    ]:
        add_para(doc, line, indent_first=False)

    heading(doc, '附件五:合作收益测算', level=1)
    add_para(doc, '详见随附 Excel 文件 《4#5#楼合作收益测算表.xlsx》, 包括: 对甲方贡献 / 基础月费构成 / 招商佣金阶梯 / 挂牌奖励明细 / 6 场沙龙测算 / 24 月双向账本 + ROI / 招商客户管道 共 8 个 Sheet。', indent_first=False)
    add_para(doc, '基础场景下: 乙方 24 个月结算金额合计约人民币 714 万元; 甲方 24 个月直接收入合计约人民币 2,769 万元; 投入产出比 ≈ 1 : 3.88。', indent_first=False)

    doc.save(OUT)
    print(f'Wrote {OUT}')


if __name__ == '__main__':
    main()
