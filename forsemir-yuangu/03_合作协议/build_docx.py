"""Build the联合运营合作协议 .docx draft.

This is a non-binding draft tailored to the Yuangu project for use in
the joint-venture negotiation between Semir, Mr Wei's team, and Prof
Hu's team. It must be reviewed by qualified legal counsel before
execution.

Implementation note: every Python string here uses SINGLE quotes so
the Chinese content can freely include ASCII double-quotes without
collision.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name('元谷项目联合运营合资公司合作协议(草案).docx')


def set_run(run, *, size=11, bold=False, color=None, font_name='微软雅黑'):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font_name)
    rfonts.set(qn('w:ascii'), font_name)
    rfonts.set(qn('w:hAnsi'), font_name)


def add_para(doc, text, *, size=11, bold=False, color=None, align=None,
             indent_first=True, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if indent_first:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    size = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 4)
    run = p.add_run(text)
    set_run(run, size=size, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))
    return p


def article(doc, num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'第{num}条　{title}')
    set_run(run, size=13, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))


def clause(doc, marker, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{marker}　{text}')
    set_run(run, size=11)


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # --- Cover ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(40)
    r = title.add_run('元谷项目联合运营合资公司')
    set_run(r, size=24, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))
    title.add_run('\n')
    r2 = title.add_run('合 作 协 议(草案)')
    set_run(r2, size=24, bold=True, color=RGBColor(0x14, 0x2C, 0x5E))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Joint Operation JV Cooperation Agreement (Draft)')
    set_run(r, size=12, color=RGBColor(0x55, 0x60, 0x7A))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(40)
    for line in [
        '项目名称:森马(上海)国际运营中心 元谷项目',
        '协议版本:v1.0  ·  起草方:胡教授团队',
        '签署日期:__________ 年 ____ 月 ____ 日',
    ]:
        r = meta.add_run(line)
        set_run(r, size=12)
        meta.add_run('\n')

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(60)
    r = note.add_run('本协议为商务谈判草案,不构成法律约束力,最终以三方法务确认并签署的正式文本为准。')
    set_run(r, size=10, color=RGBColor(0x99, 0x33, 0x33))

    doc.add_page_break()

    # --- 缔约方 ---
    heading(doc, '缔约方')
    add_para(doc, '甲方(资产方/业主方):森马集团股份有限公司(以下简称"森马")')
    add_para(doc, '    住所地:__________ ; 统一社会信用代码:__________ ; 法定代表人:__________')
    add_para(doc, '乙方(属地资源方/共同发起方):__________ (以下简称"危总团队",以危总实际控制的主体为准)')
    add_para(doc, '    住所地:__________ ; 统一社会信用代码:__________ ; 法定代表人:__________')
    add_para(doc, '丙方(运营资源方/主导运营方):__________ (以下简称"胡教授团队",由胡教授作为创始合伙人)')
    add_para(doc, '    住所地:__________ ; 统一社会信用代码:__________ ; 法定代表人:胡教授')

    add_para(doc, '鉴于:', bold=True, indent_first=False)
    clause(doc, '(1)', '甲方为元谷项目(位于上海市闵行区元江路-剑川路地区中心,"大零号湾"文创融合核心区,总建面约 22 万㎡,商业建面约 5.2 万㎡)的资产持有方与品牌方;')
    clause(doc, '(2)', '丙方在国际化运营、产业招商、政府关系、文创活动 IP(包括但不限于"北欧创新国际会客厅"、福布斯系列奖项、"科技开放麦"、AI 腾讯生态、仲量联行爬楼大数据、追觅科技基金)等方面拥有可持续投入的稀缺资源;')
    clause(doc, '(3)', '乙方在大零号湾及闵行属地拥有政府关系、央企对接与产业资源;')
    clause(doc, '(4)', '三方一致同意按"基金 + 基地 + 活动"模式共同设立合资公司,由该合资公司作为元谷项目的独家招商运营主体。')
    add_para(doc, '经友好协商,三方就合资公司之设立及合作事宜达成如下协议:', indent_first=True)

    # --- 第一条 设立 ---
    article(doc, '一', '合资公司之设立')
    clause(doc, '1.1', '公司名称(拟定):上海元谷招商运营管理有限公司(最终以登记机关核准为准)。')
    clause(doc, '1.2', '注册地:上海市闵行区(建议落地于元谷项目内)。')
    clause(doc, '1.3', '注册资本:人民币 ____ 万元,分期实缴,首期 30% 在公司设立后 30 个工作日内完成。')
    clause(doc, '1.4', '出资及股权比例(建议方案,最终以三方书面确认为准):')
    clause(doc, '      (a)', '甲方(森马):51%;以现金 + 招商运营独家授权(含元谷项目商业 5.2 万㎡ 5 年内招商主导权)作价出资。')
    clause(doc, '      (b)', '乙方(危总团队):19%;以现金 + 属地政府与央企资源作价出资。')
    clause(doc, '      (c)', '丙方(胡教授团队):30%(含 5% 期权池);以现金 + 资源出资(资源出资清单详见附件一)。')
    clause(doc, '1.5', '经营期限:自公司成立之日起 10 年;期限届满前 12 个月由三方协商续期事宜。')

    # --- 第二条 业务范围 ---
    article(doc, '二', '合资公司业务范围')
    clause(doc, '2.1', '元谷项目商业 5.2 万㎡ 的招商策划、招商执行、租户管理与续约。')
    clause(doc, '2.2', '元谷项目共享配套业态的直营运营,包括但不限于:')
    clause(doc, '      (a)', 'IP 潮玩选品 & 仓储式零售中心(约 5,000㎡);')
    clause(doc, '      (b)', '动漫潮玩谷主题街区(约 3,000㎡);')
    clause(doc, '      (c)', '潮玩艺术中心(约 2,000㎡);')
    clause(doc, '      (d)', '动漫主题书店(约 1,500㎡);')
    clause(doc, '      (e)', '森马展厅 & 二次元 Livehouse(约 700㎡)。')
    clause(doc, '2.3', '"北欧创新国际会客厅"元谷站之运营(设于 4# 楼 1-3F 潮玩艺术中心 + 4F 直播中心)。')
    clause(doc, '2.4', '园区品牌活动 IP 的策划与运营,包括"科技开放麦"系列、"全国潮玩设计大赛"、福布斯榜单元谷专场等。')
    clause(doc, '2.5', '产业牌照(包括但不限于"AI 潮玩产业基地"、"潮玩次元商业专委会")的申报、挂牌与维护。')
    clause(doc, '2.6', '招商投资基金(与追觅科技基金等)合作的"返投落地"工作; AI 腾讯生态导流; 仲量联行爬楼大数据接入。')
    clause(doc, '2.7', '三方一致同意:在合资公司经营期内,甲方就元谷项目商业部分的招商运营,授予合资公司独家、不可撤销的运营权;甲方不得将相同业务交由第三方运营机构。')

    # --- 第三条 治理 ---
    article(doc, '三', '治理结构')
    clause(doc, '3.1', '董事会:3 名董事,甲方委派 2 名(含董事长 1 名),乙方与丙方各委派 1 名(其中丙方董事由胡教授本人或其指定人担任)。')
    clause(doc, '3.2', '首席战略运营官(CSO):由丙方指定胡教授担任,全面负责合资公司日常运营; CSO 任期不少于 5 年,非因法定解除事由不得撤换。')
    clause(doc, '3.3', '管理层任命:总经理由 CSO 提名,董事会过半数通过;财务负责人由甲方提名。')
    clause(doc, '3.4', '重大事项三方一致同意(保护性条款):包括但不限于注册资本变更、新增股东、对外担保、单笔超过 ____ 万元的资本性支出、年度预算审批、章程修订、合并/分立/清算/解散。')

    # --- 第四条 月费 ---
    article(doc, '四', '运营服务费(月费 Retainer)')
    clause(doc, '4.1', '甲方同意:自合资公司成立之日起,按月向合资公司支付运营服务费("月费")。')
    clause(doc, '4.2', '月费金额:人民币 ____ 万元/月(v1.1 建议区间 20-35 万元/月,推荐 28 万元/月; 区间已基于大零号湾 / 紫竹高新区主流办公&研发租金 1.8-2.5 元/㎡/天进行成本反算,正式签署前以三方书面确认为准)。详见附件二《合作收益测算模型》Sheet 03。')
    clause(doc, '4.3', '支付节奏:每自然月 5 日前预付当月月费;逾期支付的,每逾期一日按当月月费的 0.05% 计违约金。')
    clause(doc, '4.4', '调整机制:经三方一致书面同意,每 12 个月可对月费金额进行一次评审与调整;调整范围以同期 CPI ± 8% 为参考区间。')
    clause(doc, '4.5', '胡教授团队保证:以月费支付的服务范围至少包括(详见附件一):')
    clause(doc, '      (a)', '胡教授作为 CSO 每周不少于 2 个工作日投入;')
    clause(doc, '      (b)', '1 名产业招商经理 + 1 名国际合作及活动策划 + 1 名基金投后及政府关系,全职驻场;')
    clause(doc, '      (c)', '仲量联行爬楼大数据接入与运维;')
    clause(doc, '      (d)', '北欧创新国际会客厅外事接待与 IP 引入;')
    clause(doc, '      (e)', 'AI 腾讯生态对接与共享设计中心运营。')

    # --- 第五条 招商佣金 ---
    article(doc, '五', '招商佣金')
    clause(doc, '5.1', '甲方同意:合资公司每促成一份新签或续签的元谷项目商业租赁合同(不含森马自用部分),按下列阶梯向合资公司支付招商佣金,胡教授团队按其股权比例享有相应分润:')
    clause(doc, '      (a)', '≤ 2,000㎡ 的小型租户:佣金 = 实际成交年租金 × 1.0 个月;')
    clause(doc, '      (b)', '2,001-5,000㎡ 的中型租户:佣金 = 实际成交年租金 × 1.5 个月;')
    clause(doc, '      (c)', '> 5,000㎡ 的头部/央企/行业协会租户:佣金 = 实际成交年租金 × 2.5 个月。')
    clause(doc, '5.2', '返投基金加成:凡通过追觅科技基金等合资公司导入的返投基金落地的租户,前述基础佣金额外加成 0.5 个月。')
    clause(doc, '5.3', '结算节奏:自租赁合同签订且租户起租之日起 30 日内一次性支付;合同中途解除的,按已履约期限对应比例处理。')
    clause(doc, '5.4', '归属保护:在合资公司向甲方提交的"招商客户名单"内的客户,自该客户首次接触起 24 个月内成交的,均视为合资公司业绩;甲方不得绕开合资公司签约。')

    # --- 第六条 奖项 / 挂牌 ---
    article(doc, '六', '专项奖项与挂牌激励')
    clause(doc, '6.1', '凡由合资公司或胡教授团队主导促成的下列事项,甲方应按下列标准向合资公司支付一次性激励,由合资公司按股比分配:')
    clause(doc, '      (a)', '"AI 潮玩产业基地"牌照(中国动漫集团)正式挂牌:人民币 30 万元;')
    clause(doc, '      (b)', '"潮玩次元商业专委会"牌照(中国百货商业协会)正式挂牌:人民币 30 万元;')
    clause(doc, '      (c)', '上海"科技时尚特色小镇"市级或区级奖项落地:人民币 40 万元/项;')
    clause(doc, '      (d)', '福布斯系列奖项/榜单 元谷专场发布或上榜:人民币 20 万元/项;')
    clause(doc, '      (e)', '国家级科技/文创奖项(含国家文创基金、工信部专项等):人民币 30 万元/项。')
    clause(doc, '6.2', '若同一事项在多个口径下重复计奖,以最高一档为准,不重复计酬。')
    clause(doc, '6.3', '结算节奏:自挂牌或获奖正式公告之日起 30 日内由甲方一次性支付。')

    # --- 第七条 活动 ---
    article(doc, '七', '活动运营收入')
    clause(doc, '7.1', '合资公司主导策划与执行的园区活动(包括但不限于科技开放麦、北欧外事接待、全国潮玩设计大赛、AI 共享设计中心工作坊等),其单场净收入(赞助 + 票务 + 政府补贴 - 直接成本)由合资公司单独核算。')
    clause(doc, '7.2', '活动相关政府补贴、企业赞助、票务收入由合资公司统一收取并列入合资公司营收。')
    clause(doc, '7.3', '甲方有权获得活动现场森马品牌曝光的优先权,并配合活动联合传播。')

    # --- 第八条 科技企业服务中心 ---
    article(doc, '八', '科技企业服务中心 (对外增值服务收入)')
    clause(doc, '8.1', '合资公司在元谷项目内 (建议位于 4# 楼或 5# 楼 5F+ 潮玩产业集群) 设立"元谷科技企业服务中心",为元谷项目入驻企业及大零号湾区域内潜在企业提供增值服务,该业务收入完全归属合资公司所有,不向甲方另行收取费用。')
    clause(doc, '8.2', '服务中心服务范围至少包括(详见附件五《科技企业服务中心服务清单》):')
    clause(doc, '      (a)', '注册落户、政策红利申报代办;')
    clause(doc, '      (b)', '财税、法律顾问及税筹;')
    clause(doc, '      (c)', '知识产权 (商标 / 专利 / 潮玩 IP 维权);')
    clause(doc, '      (d)', '政府补贴申报 (高新技术、专精特新、上海创新券、文创基金等);')
    clause(doc, '      (e)', '人才与签证 (居住证积分、落户、外籍工作签证);')
    clause(doc, '      (f)', '投融资 (路演对接、FA、并购);')
    clause(doc, '      (g)', '品牌与公关 (含潮玩出海推广);')
    clause(doc, '      (h)', '数字化工具 (SaaS、AI 设计工作站, 与 AI 腾讯 / 共享设计中心联动);')
    clause(doc, '      (i)', '培训与认证 (潮玩产业认证、出海实操营)。')
    clause(doc, '8.3', '收费模式:服务中心针对不同服务采用一次性、月费、项目制或提成等多种收费方式;政府补贴申报类服务可按补贴金额的 8-15% 提成,具体收费标准由 CSO 提报董事会备案后实施。')
    clause(doc, '8.4', '排他性:本协议存续期间, 甲方在元谷项目内不得另行设立或委托第三方提供与服务中心实质性竞争的服务,亦不得将本应通过服务中心提供的产业服务直接对接其他乙方机构。')
    clause(doc, '8.5', '收入归属:服务中心营收完全归合资公司,在合资公司净利润口径中按本协议第九条进行股权分红。三方一致同意:服务中心营收不计入向甲方收取的月费、招商佣金、奖项激励或活动收入科目,亦不影响上述科目的计算与结算。')
    clause(doc, '8.6', '量化承诺:丙方保证服务中心首年营收不低于人民币 150 万元,第二年不低于 300 万元,第三年不低于 600 万元;未达标的,丙方应于年度审计后 60 日内向合资公司补足差额的 30%(对应丙方持股部分)。')

    # --- 第九条 资源投入 ---
    article(doc, '九', '三方资源投入与承诺')
    clause(doc, '9.1', '甲方承诺:')
    clause(doc, '      (a)', '在合资公司成立后 90 日内向合资公司开放元谷项目商业部分的全部招商主导权;')
    clause(doc, '      (b)', '向合资公司开放元谷项目共享配套业态的直营授权;')
    clause(doc, '      (c)', '在森马集团及关联方资源范围内,为合资公司业务提供必要支持。')
    clause(doc, '9.2', '乙方承诺:')
    clause(doc, '      (a)', '协助合资公司对接闵行区、大零号湾管委会等政府关系及央企资源;')
    clause(doc, '      (b)', '协助申报"科技时尚特色小镇"等政策包及相关补贴。')
    clause(doc, '9.3', '丙方承诺(量化承诺,列入考核):')
    clause(doc, '      (a)', '12 个月内挂牌"北欧创新国际会客厅"元谷站;')
    clause(doc, '      (b)', '90 日内将仲量联行爬楼大数据接入合资公司中台;')
    clause(doc, '      (c)', '12 个月内通过追觅科技基金完成首期返投落地不少于 3 家潮玩企业;')
    clause(doc, '      (d)', '12 个月内挂牌不少于 2 项产业牌照或福布斯榜单;')
    clause(doc, '      (e)', '12 个月内举办不少于 10 场"科技开放麦";')
    clause(doc, '      (f)', '科技企业服务中心首年营收不低于 150 万元(详见第八条 8.6)。')

    # --- 第十条 利润 ---
    article(doc, '十', '利润分配与亏损分担')
    clause(doc, '10.1', '合资公司的可分配利润("分红")按三方实际持股比例分配。')
    clause(doc, '10.2', '三方一致同意:在合资公司成立后的前 24 个月内,原则上将不少于 60% 的可分配利润留存于合资公司用于业务发展。')
    clause(doc, '10.3', '亏损按持股比例分担,但任何一方的责任不超过其实缴出资额。')

    # --- 第十一条 团队 ---
    article(doc, '十一', '团队招聘与人员管理')
    clause(doc, '11.1', '合资公司核心团队由 CSO 主导招聘,包括但不限于:产业招商经理、国际合作 & 活动策划、基金投后 & 政府关系,及行政/财务岗位。岗位说明详见附件三《核心团队 JD 与薪酬》。')
    clause(doc, '11.2', '核心员工的薪酬由合资公司负担,并由 CSO 报董事会备案。')
    clause(doc, '11.3', '合资公司可设立 5% 的员工持股计划(ESOP),由丙方持股部分中预留。')

    # --- 第十二条 知识产权 ---
    article(doc, '十二', '知识产权与资源排他性')
    clause(doc, '12.1', '胡教授团队就以下资源在元谷项目场景内的使用,授予合资公司独家、可转授的使用权(在合资公司经营期内):北欧创新国际会客厅 IP、科技开放麦 IP、福布斯系列奖项渠道、AI 腾讯导入合作、仲量联行爬楼大数据接入、追觅科技基金合作。')
    clause(doc, '12.2', '上述资源的所有权仍归原权利人;合资公司经营期届满或本协议解除的,使用权同步解除。')
    clause(doc, '12.3', '合资公司对外发布及营销中应同步标注上述资源贡献方信息。')

    # --- 第十三条 排他与竞业 ---
    article(doc, '十三', '排他与竞业禁止')
    clause(doc, '13.1', '合资公司经营期内,三方不得直接或间接从事与合资公司业务存在实质性竞争的活动;不得另行委托或受托从事元谷项目商业部分的招商运营。')
    clause(doc, '13.2', '甲方在大零号湾区域内的其他自有商业项目,需先以书面形式询价于合资公司,合资公司在 30 日内未明示接受的,甲方有权另行委托。')

    # --- 第十四条 退出 ---
    article(doc, '十四', '退出机制')
    clause(doc, '14.1', '合资公司成立满 36 个月后,任何一方可发起股权转让;其他股东在同等条件下享有优先购买权。')
    clause(doc, '14.2', '若合资公司连续 12 个月未能完成关键 KPI(详见附件二),甲方有权回购丙方股权(回购价不低于丙方实缴出资额);丙方亦有权在该情形下要求乙方与甲方按比例回购。')
    clause(doc, '14.3', '本协议任一方严重违约,且经书面催告 30 日仍未纠正的,守约方有权解除本协议,并要求违约方按附件二所列损失予以赔偿。')

    # --- 第十五条 保密 ---
    article(doc, '十五', '保密')
    clause(doc, '15.1', '三方就在本协议项下知悉的对方商业秘密、客户名单、技术信息、财务信息及未公开战略,承担保密义务,期限为本协议存续期间及解除/终止后 5 年。')

    # --- 第十六条 争议 ---
    article(doc, '十六', '法律适用与争议解决')
    clause(doc, '16.1', '本协议适用中华人民共和国法律。')
    clause(doc, '16.2', '凡因本协议引起的争议,三方应首先友好协商解决;协商不成的,提交上海国际经济贸易仲裁委员会(上海国际仲裁中心)按其届时有效的仲裁规则在上海仲裁。')

    # --- 第十七条 一般 ---
    article(doc, '十七', '一般条款')
    clause(doc, '17.1', '本协议自三方法定代表人或授权代表签字并加盖公章之日起生效。')
    clause(doc, '17.2', '本协议一式六份,三方各执两份,具有同等法律效力。')
    clause(doc, '17.3', '本协议附件与正文具有同等法律效力。')
    clause(doc, '17.4', '本协议未尽事宜,由三方另行书面约定。')

    # --- 签署页 ---
    doc.add_page_break()
    heading(doc, '签 署 页', level=1)
    sig = doc.add_table(rows=4, cols=2)
    sig.autofit = False
    for col in sig.columns:
        for cell in col.cells:
            cell.width = Cm(8)
    rows = [
        ('甲方:森马集团股份有限公司',          '乙方:__________(危总团队)'),
        ('法定代表人/授权代表:__________',     '法定代表人/授权代表:__________'),
        ('(盖章)　　　　　　 日期:____ 年 ____ 月 ____ 日',
         '(盖章)　　　　　　 日期:____ 年 ____ 月 ____ 日'),
        ('丙方:__________(胡教授团队)　　法定代表人/授权代表:__________　　(盖章)　　　　　　 日期:____ 年 ____ 月 ____ 日', ''),
    ]
    for i, (a, b) in enumerate(rows):
        if i == 3:
            sig.cell(i, 0).merge(sig.cell(i, 1))
            sig.cell(i, 0).text = a
        else:
            sig.cell(i, 0).text = a
            sig.cell(i, 1).text = b
        for cell in sig.rows[i].cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for r in para.runs:
                    set_run(r, size=11)
                para.paragraph_format.space_after = Pt(8)

    # --- 附件 ---
    doc.add_page_break()
    heading(doc, '附件一:丙方资源出资清单', level=1)
    for line in [
        '1. 北欧创新国际会客厅 IP & 国际外事网络。',
        '2. 福布斯系列奖项及榜单合作渠道。',
        '3. 科技开放麦活动 IP(多季历史活动数据)。',
        '4. AI 腾讯生态导流合作。',
        '5. 仲量联行爬楼大数据(已实际投入采购成本人民币 26,000 元)。',
        '6. 追觅科技基金"返投落地"招商合作通道。',
        '7. 中国百货商业协会/中国动漫集团 两大产业牌照对接资源。',
        '8. 胡教授个人品牌、学术 & 行业网络(按附议详细 listing)。',
    ]:
        add_para(doc, line, indent_first=False)

    heading(doc, '附件二:合作收益测算模型', level=1)
    add_para(doc, '详见随附 Excel 文件 《合作收益测算模型.xlsx》(含 8 个 Sheet:封面/收入总览/市场租金对标/月费测算/招商佣金阶梯/活动+奖项/科技企业服务中心/股权分红+三年累计/敏感性分析)。基础场景下首年向胡教授团队结算金额合计约人民币 1,002 万元;另合资公司层面新增对外科技企业服务中心营收约人民币 300 万元(三年累计约 1,090 万元), 通过股权分红体现于丙方收入。', indent_first=False)

    heading(doc, '附件三:核心团队 JD 与薪酬', level=1)
    add_para(doc, '详见随附文档 《04_团队招聘/团队招聘计划.md》。包括 3 个全职岗位的岗位职责、任职要求、薪酬包以及 1 个共享行政岗位的设置建议。', indent_first=False)

    heading(doc, '附件四:关键 KPI', level=1)
    for line in [
        '1. 首年招商成交面积 ≥ 12,000㎡。',
        '2. 首年挂牌产业牌照 ≥ 2 项。',
        '3. 首年举办品牌活动 ≥ 14 场(含 12 场科技开放麦 + 1 届潮玩大赛 + 1 场福布斯发布)。',
        '4. 首年完成基金返投落地 ≥ 3 家潮玩企业。',
        '5. 北欧创新国际会客厅元谷站 12 个月内挂牌运营。',
        '6. 科技企业服务中心首年营收 ≥ 150 万元 (服务户数 ≥ 30 家)。',
    ]:
        add_para(doc, line, indent_first=False)

    heading(doc, '附件五:科技企业服务中心服务清单', level=1)
    add_para(doc, '详见随附文档 《05_附件/科技企业服务中心服务清单.md》。该清单覆盖 9 大类服务及对应收费区间, 由 CSO 提报合资公司董事会备案后实施, 每年度可根据市场调整一次。', indent_first=False)
    intro = '九大类服务:① 注册落户; ② 财税法; ③ 知识产权; ④ 政府补贴申报; ⑤ 人才与签证; ⑥ 投融资; ⑦ 品牌与公关; ⑧ 数字化工具; ⑨ 培训与认证。'
    add_para(doc, intro, indent_first=False)
    add_para(doc, '该业务对外收费, 不向甲方收取费用, 营收完全归属合资公司; 丙方按 30% 持股享有相应分红。', indent_first=False)

    doc.save(OUT)
    print(f'Wrote {OUT}')


if __name__ == '__main__':
    main()
