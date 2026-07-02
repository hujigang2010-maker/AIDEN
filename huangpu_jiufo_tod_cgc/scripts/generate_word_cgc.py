# -*- coding: utf-8 -*-
"""生成 Word《合作意向说明与授权(提资)证明申请函》(含证明书建议格式附件)。"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import content_cgc as C

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "deliverables")
FONT = "微软雅黑"
NAVY = RGBColor(0x0F, 0x2A, 0x4A)
GOLD = RGBColor(0xA8, 0x80, 0x2A)


def cn(run, size=None, bold=None, color=None, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def para(doc, text, size=12, bold=False, color=None, align=None,
         space_after=8, indent=None, line=1.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    cn(r, size=size, bold=bold, color=color)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line
    if indent:
        p.paragraph_format.first_line_indent = Pt(indent)
    return p


def heading(doc, text, size=15, space_before=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    cn(r, size=size, bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(8)
    return p


def main():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.6)
        sec.left_margin = Cm(3.0); sec.right_margin = Cm(3.0)

    # 抬头机构
    para(doc, C.OWNER_ORG, size=13, bold=True, color=GOLD,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, line=1.2)
    para(doc, "CGC · 文创与IP专委会", size=9, color=GOLD,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, line=1.0)
    # 分隔线(用底部边框)
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pbdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "12",
        qn("w:space"): "1", qn("w:color"): "A8802A"})
    pbdr.append(bottom); pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(16)

    # 标题
    para(doc, C.LETTER_TITLE, size=17, bold=True, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20, line=1.4)

    # 正文
    body = C.LETTER_BODY
    para(doc, body[0].format(cp=C.COUNTERPARTY), size=12, bold=True, space_after=10)
    for seg in body[1:]:
        para(doc, seg, size=12, indent=24)

    # 落款
    doc.add_paragraph()
    for i, line in enumerate(C.LETTER_SIGN):
        para(doc, line, size=12, bold=(i == 0),
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4, line=1.3)

    # ---------- 附件:证明书建议格式 ----------
    doc.add_page_break()
    para(doc, "附件", size=11, bold=True, color=GOLD, space_after=6)
    para(doc, C.CERT_TITLE, size=16, bold=True, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8, line=1.4)
    para(doc, "证　明", size=20, bold=True, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    for seg in C.CERT_BODY:
        para(doc, seg, size=12.5, indent=25, line=1.7)
    doc.add_paragraph(); doc.add_paragraph()
    for line in C.CERT_SIGN:
        para(doc, line, size=12.5, align=WD_ALIGN_PARAGRAPH.RIGHT,
             space_after=10, line=1.4)
    para(doc, "（注：本证明为建议格式，抬头单位、具体表述及盖章方式可由贵方按实际情况调整。）",
         size=10, color=RGBColor(0x88, 0x88, 0x88), space_after=4, line=1.3)

    os.makedirs(OUT_DIR, exist_ok=True)
    out = os.path.join(OUT_DIR, "CGC合作意向与授权提资证明申请函.docx")
    doc.save(out)
    print("saved:", out)


if __name__ == "__main__":
    main()
