# -*- coding: utf-8 -*-
"""生成《服务建议书》Word 文档(含项目概要、提资清单与报价)。"""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import content as C

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "deliverables")
CN_FONT = "微软雅黑"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)


def set_cn_font(run, size=None, bold=None, color=None, font=CN_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def para(doc, text, size=12, bold=False, color=None, align=None,
         space_after=6, first_line_indent=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=bold, color=color)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(first_line_indent)
    return p


def heading(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12.5}
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, size=sizes[level], bold=True, color=ACCENT)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    return p


def bullet(doc, text, size=12):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_cn_font(r, size=size)
    p.paragraph_format.space_after = Pt(4)
    return p


def style_table(table, header_row=True):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_cn_font(r, size=10.5, bold=(header_row and i == 0))
            if header_row and i == 0:
                shading = cell._tc.get_or_add_tcPr()
                el = shading.makeelement(qn("w:shd"), {
                    qn("w:val"): "clear", qn("w:fill"): "1F4E79"})
                shading.append(el)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # ---------------- 封面 ----------------
    for _ in range(5):
        doc.add_paragraph()
    para(doc, C.PROJECT_NAME, size=22, bold=True, color=ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    para(doc, "策划服务建议书", size=28, bold=True, color=ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    para(doc, "(含提资清单与服务报价)", size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=60)
    for _ in range(6):
        doc.add_paragraph()
    para(doc, f"提供方:{C.PROVIDER_LINE}", size=13, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    para(doc, f"提交对象:{C.CLIENT}", size=13,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    para(doc, C.DOC_DATE, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ---------------- 一、项目概要 ----------------
    heading(doc, "一、项目概要", 1)
    heading(doc, "1.1 项目背景与概念", 2)
    para(doc, f"{C.CONCEPT_TITLE}:", size=12, bold=True)
    for pt in C.CONCEPT_POINTS:
        bullet(doc, pt)

    heading(doc, "1.2 产业方向", 2)
    para(doc, C.INDUSTRY_INTRO, size=12)
    for name, detail in C.INDUSTRY_GROUPS:
        bullet(doc, f"{name}:{detail}")

    heading(doc, "1.3 下阶段重点", 2)
    para(doc, C.NEXT_STAGE_TEXT, size=12, first_line_indent=24)

    # ---------------- 二、服务范围 ----------------
    heading(doc, "二、服务范围(策划服务八大模块)", 1)
    for i, (name, detail) in enumerate(C.SERVICE_MODULES, 1):
        para(doc, f"{i}. {name}", size=12, bold=True, space_after=2)
        para(doc, detail, size=11.5, first_line_indent=24)

    # ---------------- 三、成果交付 ----------------
    heading(doc, "三、成果交付", 1)
    for d in C.DELIVERABLES:
        bullet(doc, d)

    # ---------------- 四、提资清单 ----------------
    doc.add_page_break()
    heading(doc, "四、提资清单(请委托方提供的资料)", 1)
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    for j, t in enumerate(["序号", "类别", "资料名称", "说明", "优先级"]):
        hdr[j].text = t
    for i, (cat, name, desc, prio) in enumerate(C.INFO_REQUEST_ITEMS, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = cat
        row[2].text = name
        row[3].text = desc
        row[4].text = prio
    style_table(table)
    para(doc, "", size=6)
    para(doc, "提资说明:", size=12, bold=True)
    for n in C.INFO_REQUEST_NOTES:
        bullet(doc, n, size=11.5)

    # ---------------- 五、服务报价 ----------------
    heading(doc, "五、服务报价", 1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    for j, t in enumerate(["工作阶段", "主要工作内容", "主要成果", "报价(万元)"]):
        hdr[j].text = t
    for stage, work, output, fee in C.QUOTATION_ITEMS:
        row = table.add_row().cells
        row[0].text = stage
        row[1].text = work
        row[2].text = output
        row[3].text = str(fee)
    row = table.add_row().cells
    row[0].text = "合计"
    row[3].text = f"{C.QUOTATION_TOTAL}"
    style_table(table)
    for cell in table.rows[-1].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                set_cn_font(r, size=10.5, bold=True)
    para(doc, "", size=6)
    para(doc, "报价说明:", size=12, bold=True)
    for n in C.QUOTATION_NOTES:
        bullet(doc, n, size=11.5)

    # ---------------- 六、团队与联系方式 ----------------
    heading(doc, "六、服务团队与联系方式", 1)
    para(doc,
         "本项目由复旦大学住房政策研究中心与上海市杨浦区科技企业联合会联合组建服务团队:"
         "研究中心负责政策研究、项目定位与策划方案编制,"
         "科技企业联合会负责产业资源组织、企业对接与招商建议。",
         size=12, first_line_indent=24)
    para(doc, "联系方式:另行提供(以正式合同联络人为准)。", size=12,
         first_line_indent=24)

    os.makedirs(OUT_DIR, exist_ok=True)
    out = os.path.join(OUT_DIR, "广州知识城全球自贸365街区_服务建议书.docx")
    doc.save(out)
    print("saved:", out)


if __name__ == "__main__":
    main()
