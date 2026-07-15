# -*- coding: utf-8 -*-
"""广州黄埔区九佛TOD「全球自贸365街区」项目——统一(内容级合并)策划服务建议书 Word。

将原两套材料在内容层面合并、去重并优化,统一到唯一项目名称,商务大气版式。
"""
import os

from docx import Document
from docx.shared import Cm, RGBColor

import content_unified as C
import content_cgc as CGC
from style_docx import (BLUE, GOLD, GRAY, NAVY, WD_ALIGN_PARAGRAPH,
                        banner, bullet, h1, h2, para, styled_table)

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "deliverables")
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHTGOLD = RGBColor(0xE8, 0xC9, 0x77)


def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    banner(doc, [
        (C.PROJECT_NAME, 24, WHITE, True),
        ("GLOBAL FREE TRADE 365 BLOCK · HUANGPU JIUFO TOD", 10, LIGHTGOLD, True),
    ])
    para(doc, "", space_after=2)
    banner(doc, [
        ("策划服务建议书", 26, NAVY, True),
        ("（含提资清单与服务报价 · 附市场数据）", 12, GRAY, False),
    ], fill_hex="F6EEDA")
    for _ in range(7):
        doc.add_paragraph()
    para(doc, f"提供方:{C.PROVIDER_LINE}", size=13, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(doc, f"提交对象:{C.CLIENT}", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(doc, C.DOC_DATE, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def body(doc):
    # 一、项目概要
    h1(doc, "项目概要", no="一")
    h2(doc, C.CONCEPT_TITLE)
    for pt in C.CONCEPT_POINTS:
        bullet(doc, pt)
    h2(doc, "产业方向")
    para(doc, C.INDUSTRY_INTRO, size=11.5)
    for name, detail in C.INDUSTRY_GROUPS:
        bullet(doc, f"{name}:{detail}")
    h2(doc, "下阶段重点")
    para(doc, C.NEXT_STAGE_TEXT, size=11.5, indent=24)

    # 二、服务范围
    h1(doc, "服务范围(策划服务八大模块)", no="二")
    for i, (name, detail) in enumerate(C.SERVICE_MODULES, 1):
        para(doc, f"{i}. {name}", size=12, bold=True, color=BLUE, space_after=2)
        para(doc, detail, size=11, indent=24, space_after=6)

    # 三、成果交付
    h1(doc, "成果交付", no="三")
    for d in C.DELIVERABLES:
        bullet(doc, d)

    # 四、提资清单
    h1(doc, "提资清单(请委托方提供的资料)", no="四")
    rows = [(str(i), cat, name, desc, prio) for i, (cat, name, desc, prio)
            in enumerate(C.INFO_REQUEST_ITEMS, 1)]
    styled_table(doc, ["序号", "类别", "资料名称", "说明", "优先级"], rows,
                 col_widths_cm=[1.1, 2.6, 3.8, 5.8, 1.5],
                 center_cols={0, 4}, body_size=9.5, header_size=10)
    para(doc, "", size=4)
    for n in C.INFO_REQUEST_NOTES:
        bullet(doc, n, size=10.5)

    # 五、服务报价
    h1(doc, "服务报价", no="五")
    qrows = [(s, w, o, str(f)) for (s, w, o, f) in C.QUOTATION_ITEMS]
    qrows.append(("合计", "", "", str(C.QUOTATION_TOTAL)))
    styled_table(doc, ["工作阶段", "主要工作内容", "主要成果", "报价(万元)"], qrows,
                 col_widths_cm=[3.3, 5.6, 3.6, 2.3], center_cols={3},
                 body_size=9.5, header_size=10, last_row_bold=True)
    para(doc, "", size=4)
    for n in C.QUOTATION_NOTES:
        bullet(doc, n, size=10.5)
    doc.add_page_break()


def appendix(doc):
    h1(doc, "附件 · 市场数据", no="六")
    h2(doc, "附件一  2025年广州出口TOP20品类")
    rows = [(str(r[0]), r[1], r[2], r[3], r[4]) for r in CGC.EXPORT_TOP20]
    styled_table(doc, ["排名", "品类", "代表品牌", "广州口岸交易额", "核心说明"], rows,
                 col_widths_cm=[1.1, 3.2, 3.2, 2.8, 4.5], center_cols={0, 3},
                 body_size=8.8, header_size=9.5)
    para(doc, "说明:" + CGC.EXPORT_TOP20_NOTE, size=9.5, color=GRAY, space_before=4)
    doc.add_page_break()
    h2(doc, "附件二  2025年前10月广州消费类出口20强(单位:亿元)")
    rows = [(str(r[0]), r[1], r[2], r[3], r[4], r[5]) for r in CGC.CONSUMER_TOP20]
    styled_table(doc, ["排名", "品类", "出口额", "同比", "核心出口品牌", "核心市场"], rows,
                 col_widths_cm=[1.1, 3.0, 1.8, 1.6, 4.0, 3.0],
                 center_cols={0, 2, 3}, body_size=8.8, header_size=9.5)
    para(doc, "说明:" + CGC.CONSUMER_TOP20_NOTE, size=9.5, color=GRAY, space_before=4)
    doc.add_page_break()
    h2(doc, "附件三  楼层功能布局")
    rows = [(r[0], r[1], r[2], r[3]) for r in CGC.FLOOR_LAYOUT]
    styled_table(doc, ["区位楼层", "功能定位", "核心业态 / 服务", "数据支撑与参考案例"],
                 rows, col_widths_cm=[2.6, 3.2, 5.2, 4.2], center_cols={0},
                 body_size=8.8, header_size=9.5)
    doc.add_page_break()


def team(doc):
    h1(doc, "服务团队与联系方式", no="七")
    para(doc,
         "本项目由复旦大学住房政策研究中心与上海市杨浦区科技企业联合会联合组建服务团队:"
         "研究中心负责政策研究、项目定位与策划方案编制;科技企业联合会负责产业资源组织、"
         "企业对接与招商建议,共同为项目提供一体化策划服务。", size=11.5, indent=24)
    para(doc, "联系方式:另行提供(以正式合同联络人为准)。", size=11.5, indent=24)


def main():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.4); sec.bottom_margin = Cm(2.4)
        sec.left_margin = Cm(2.6); sec.right_margin = Cm(2.6)
    cover(doc)
    body(doc)
    appendix(doc)
    team(doc)
    os.makedirs(OUT_DIR, exist_ok=True)
    out = os.path.join(OUT_DIR, "广州黄埔九佛TOD全球自贸365街区_策划服务建议书.docx")
    doc.save(out)
    print("saved:", out)


if __name__ == "__main__":
    main()
