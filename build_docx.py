"""Generate a well-formatted Word document for the Tencent Cloud sponsorship agreement.

Targets formal Chinese contract typesetting:
- Body: 宋体 (SimSun) 小四 (12pt), 1.5x line spacing, 2-char first-line indent
- Headings: 黑体 (SimHei), centered/bold, hierarchical sizes
- Emphasis (**...**) rendered as bold inline runs
- A4 page, 2.5cm margins
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DOC_PATH = "/workspace/腾讯云战略协办合作协议.docx"

BODY_FONT_EN = "Times New Roman"
BODY_FONT_CN = "宋体"
HEAD_FONT_CN = "黑体"


def set_run_fonts(run, size_pt, *, bold=False, cn_font=BODY_FONT_CN, en_font=BODY_FONT_EN, color=None):
    run.font.name = en_font
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), en_font)
    rfonts.set(qn("w:hAnsi"), en_font)
    rfonts.set(qn("w:eastAsia"), cn_font)
    rfonts.set(qn("w:cs"), en_font)


def set_paragraph_format(p, *, align=None, line_spacing=1.5, space_before=0, space_after=6,
                         first_line_indent_chars=0, left_indent_cm=None):
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if left_indent_cm is not None:
        pf.left_indent = Cm(left_indent_cm)
    if first_line_indent_chars:
        ppr = p._element.get_or_add_pPr()
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            ppr.append(ind)
        ind.set(qn("w:firstLineChars"), str(first_line_indent_chars * 100))
        ind.set(qn("w:firstLine"), "0")


def add_runs_with_bold(paragraph, text, *, size_pt=12, base_bold=False, cn_font=BODY_FONT_CN):
    """Parse **bold** segments and emit runs."""
    parts = []
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end == -1:
                parts.append((text[i:], base_bold))
                break
            parts.append((text[i + 2:end], True))
            i = end + 2
        else:
            nxt = text.find("**", i)
            if nxt == -1:
                parts.append((text[i:], base_bold))
                break
            parts.append((text[i:nxt], base_bold))
            i = nxt
    for chunk, bold in parts:
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        set_run_fonts(run, size_pt, bold=bold, cn_font=cn_font)


def add_title(doc, text, *, size_pt=22):
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                         space_before=0, space_after=6)
    run = p.add_run(text)
    set_run_fonts(run, size_pt, bold=True, cn_font=HEAD_FONT_CN)


def add_subtitle(doc, text, *, size_pt=18):
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                         space_before=0, space_after=18)
    run = p.add_run(text)
    set_run_fonts(run, size_pt, bold=True, cn_font=HEAD_FONT_CN)


def add_meta_line(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.5,
                         space_before=0, space_after=2)
    add_runs_with_bold(p, text, size_pt=11)


def add_section_heading(doc, text, *, size_pt=14):
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                         space_before=12, space_after=8)
    run = p.add_run(text)
    set_run_fonts(run, size_pt, bold=True, cn_font=HEAD_FONT_CN)


def add_body_paragraph(doc, text, *, indent_chars=2, bold=False, align=None):
    p = doc.add_paragraph()
    set_paragraph_format(p,
                         align=align or WD_ALIGN_PARAGRAPH.JUSTIFY,
                         line_spacing=1.5,
                         space_before=0, space_after=6,
                         first_line_indent_chars=indent_chars)
    add_runs_with_bold(p, text, size_pt=12, base_bold=bold)
    return p


def add_numbered_item(doc, idx, title_text, body_text=None):
    """Numbered item like:  1. **合作级别**  + body paragraph(s)."""
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
                         space_before=4, space_after=4,
                         first_line_indent_chars=2)
    num_run = p.add_run(f"{idx}. ")
    set_run_fonts(num_run, 12, bold=True)
    add_runs_with_bold(p, title_text, size_pt=12, base_bold=True)
    if body_text:
        bp = doc.add_paragraph()
        set_paragraph_format(bp, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
                             space_before=0, space_after=6,
                             first_line_indent_chars=2)
        add_runs_with_bold(bp, body_text, size_pt=12)


def add_bullet(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
                         space_before=0, space_after=3,
                         left_indent_cm=1.4)
    run = p.add_run("• ")
    set_run_fonts(run, 12)
    add_runs_with_bold(p, text, size_pt=12)


def add_hr(doc):
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.0, space_before=4, space_after=4)
    pPr = p._element.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_blank(doc, size_pt=12):
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.0, space_before=0, space_after=0)
    run = p.add_run("")
    set_run_fonts(run, size_pt)


def build():
    doc = Document()

    # --- Page setup: A4, 2.5cm margins
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.7)
        section.right_margin = Cm(2.7)

    # --- Default style
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT_EN
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_CN)

    # =========================================================
    # Title block
    # =========================================================
    add_title(doc, "2026 人工智能商业化落地峰会", size_pt=20)
    add_subtitle(doc, "《战略协办合作协议》", size_pt=22)

    add_hr(doc)

    add_meta_line(doc, "合同编号： AIBIZ-2026-SP-001")
    add_meta_line(doc, "签署地点： 上海市")
    add_meta_line(doc, "签署日期： 2026 年 ___ 月 ___ 日")

    add_hr(doc)
    add_blank(doc, 6)

    # =========================================================
    # Parties
    # =========================================================
    add_section_heading(doc, "协议主体")

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5, space_before=0, space_after=4, first_line_indent_chars=0)
    add_runs_with_bold(p, "**甲方（主办方 / 组委会）：**", size_pt=12)
    add_bullet(doc, "人工智能商业化落地峰会组委会")
    add_bullet(doc, "复旦大学住房政策研究中心")
    add_bullet(doc, "上海市杨浦区科技企业联合会")

    add_blank(doc, 4)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5, space_before=0, space_after=4, first_line_indent_chars=0)
    add_runs_with_bold(p, "**乙方（战略协办方）：**", size_pt=12)
    add_bullet(doc, "腾讯云计算（北京）有限责任公司")

    add_blank(doc, 4)

    add_body_paragraph(doc, "鉴于：", indent_chars=0, bold=True)
    add_numbered_item(doc, 1,
                      "甲方拟于 2026 年在上海北外滩核心地标举办"
                      "“2026 人工智能商业化落地峰会”（以下简称“本次峰会”），"
                      "聚焦人工智能在新质生产力、大模型商业化、产业数字化转型等领域的落地实践；",
                      None)
    add_numbered_item(doc, 2,
                      "乙方作为国内领先的云计算与人工智能基础设施服务商，"
                      "希望通过本次峰会向核心政企客户与产业生态展示其在算力与大模型领域的商业化能力；",
                      None)
    add_numbered_item(doc, 3,
                      "双方本着平等自愿、互利共赢的原则，"
                      "就乙方战略协办本次峰会的相关事宜达成如下协议，以资共同遵守。",
                      None)

    # =========================================================
    # Article 1
    # =========================================================
    add_section_heading(doc, "第一条　合作级别与赞助标的")

    add_numbered_item(doc, 1, "合作级别",
                      "乙方为本次峰会唯一的“**首席战略合作伙伴**”及"
                      "“**算力生态独家伙伴**”，享有同一级别下不可被并列、不可被超越的最高排序权益。")
    add_numbered_item(doc, 2, "赞助标的额",
                      "乙方为本次峰会提供专项赞助资金，"
                      "总计金额为 **人民币 300,000 元（大写：叁拾万元整）**。")
    add_numbered_item(doc, 3, "资金定向用途",
                      "双方明确，**上述人民币 300,000 元赞助款项仅用于且全部用于本次峰会当期的专项支出**，"
                      "具体包括但不限于：")
    add_bullet(doc, "北外滩核心地标场地租赁与场地保障费用；")
    add_bullet(doc, "主会场高规格舞美设计、视觉搭建与现场技术支持费用；")
    add_bullet(doc, "顶尖政企大咖、智库专家及核心嘉宾的接待、出行与安保运营费用；")
    add_bullet(doc, "峰会当期主视觉物料制作、官方传播与现场运营所产生的其他直接费用。")

    add_numbered_item(doc, 4, "支付方式与时点",
                      "乙方应于本协议签署并经双方盖章生效后 15 个工作日内，"
                      "将上述款项一次性汇入甲方指定的对公账户。"
                      "甲方在收款后向乙方开具相应金额的合规发票。")

    # =========================================================
    # Article 2
    # =========================================================
    add_section_heading(doc, "第二条　甲方权益交付明细（全包制，不作分项标价）")

    add_body_paragraph(doc,
        "为对应乙方上述赞助投入，甲方在本次峰会及其后续生态互动中，向乙方综合交付以下顶级权益。"
        "双方确认，下列权益作为一揽子综合对价整体交付，"
        "不再就单项权益进行分项估值与单独定价：")

    add_numbered_item(doc, 1, "核心现场曝光",
                      "乙方享有峰会主视觉背景板、官方议程、官方新闻通稿、现场签到背板及核心传播物料中的"
                      "“首席战略协办方”顶级 Logo 并排露出权益，排位优先于本次峰会其他所有合作方。")
    add_numbered_item(doc, 2, "议题主导与高管发声",
                      "甲方为乙方核心高管在峰会主会场保留 1 个主旨演讲席位或 1 个核心圆桌对话席位"
                      "（具体形式由双方协商确定），"
                      "用于全方位展示乙方在新质生产力、大模型及云计算商业化落地领域的最新成果。")
    add_numbered_item(doc, 3, "顶尖圈层深度对接",
                      "甲方依托自身政企资源与智库资源，在峰会当期为乙方定向引荐具备真实数字化转型与算力需求的大型重资产企业，"
                      "包括但不限于头部算力需求方、高端汽车流通集团、头部地产与城市运营集团、金融及产业资本机构等，"
                      "并安排闭门深度对接环节。")
    add_numbered_item(doc, 4, "算力生态专属延伸权益",
                      "作为首席算力生态独家伙伴，乙方有权在峰会后，"
                      "向甲方生态内符合资质的科创企业与“超级个体”定向发放腾讯云专属算力包及大模型 API 调用支持，"
                      "甲方将通过自有渠道予以全程配合宣发。"
                      "该项权益作为乙方首席算力生态伙伴身份的自带延伸权益，"
                      "不计入本协议第一条项下的赞助标的额。")
    add_numbered_item(doc, 5, "智库长效背书",
                      "乙方自动成为“见微知海新质商业生态”首批年度战略理事单位，"
                      "共享复旦大学住房政策研究中心宏观经济与产业研究成果，"
                      "享有后续长三角系列闭门局、专题沙龙及行业研究发布会的优先参与权。"
                      "该项长效权益同样作为乙方战略级合作身份的自带权益，"
                      "不计入本协议第一条项下的赞助标的额。")

    # =========================================================
    # Article 3
    # =========================================================
    add_section_heading(doc, "第三条　双方义务")

    add_numbered_item(doc, 1, "甲方义务", None)
    add_bullet(doc, "严格按本协议第二条约定，向乙方完整、及时交付各项权益；")
    add_bullet(doc, "保障峰会按计划如期举办，并维护峰会的行业影响力与规格档次；")
    add_bullet(doc, "在峰会筹备、执行、后续传播全过程中，确保乙方作为“首席战略协办方”及“算力生态独家伙伴”的排他性与唯一性，不引入与乙方在云计算及算力领域形成直接竞争关系的同级别合作方；")
    add_bullet(doc, "妥善管理乙方支付的赞助资金，确保专款专用于本次峰会当期支出。")

    add_numbered_item(doc, 2, "乙方义务", None)
    add_bullet(doc, "按本协议约定按时足额支付赞助款项；")
    add_bullet(doc, "配合甲方完成 Logo、高管简介、演讲主题等权益落地所需的素材提供与确认工作；")
    add_bullet(doc, "不得利用本次峰会平台从事违反国家法律法规及公序良俗的活动。")

    # =========================================================
    # Article 4
    # =========================================================
    add_section_heading(doc, "第四条　保密及合规条款")

    add_numbered_item(doc, 1, "商业机密保护",
                      "鉴于本协议体现了甲乙双方最核心的战略合作对价，"
                      "**未经甲乙双方书面同意，任何一方不得向任何第三方（包括但不限于其他赞助商、合作方、媒体及外部审计机构）"
                      "披露本协议项下的资金打款凭证、财务流水、内部核算细则及权益分项估值等敏感信息。**")
    add_numbered_item(doc, 2, "独立比对声明",
                      "甲方有权在对外招商、生态展示及向同级别合作方进行横向拉齐时，"
                      "合法引用本协议第一条项下的合作总金额及合作级别作为优秀案例说明，"
                      "以证明本次峰会的市场估值与顶尖行业含金量。"
                      "该等引用不视为对本条第 1 款保密义务的违反。")
    add_numbered_item(doc, 3, "合规承诺",
                      "双方承诺，本协议项下的资金往来与权益交付均合法合规，"
                      "不存在任何商业贿赂、利益输送及违反反不正当竞争法、反商业贿赂相关规定的情形。")

    # =========================================================
    # Article 5
    # =========================================================
    add_section_heading(doc, "第五条　违约责任")

    add_numbered_item(doc, 1,
                      "任何一方未按本协议约定履行义务，"
                      "给对方造成损失的，应承担相应的违约责任并赔偿对方因此遭受的直接经济损失。",
                      None)
    add_numbered_item(doc, 2,
                      "若因不可抗力（包括但不限于自然灾害、政府管制、重大公共卫生事件等）"
                      "导致本次峰会延期或无法举办，双方应友好协商，"
                      "按已发生的实际支出比例处理已支付的赞助款项；"
                      "剩余款项可顺延用于甲方下一届同级别峰会，或经双方书面同意后退还乙方。",
                      None)

    # =========================================================
    # Article 6
    # =========================================================
    add_section_heading(doc, "第六条　争议解决")

    add_body_paragraph(doc,
        "本协议的签订、履行、解释及争议解决均适用中华人民共和国法律。"
        "因本协议产生或与本协议有关的任何争议，双方应首先通过友好协商解决；"
        "协商不成的，任何一方均有权向甲方所在地有管辖权的人民法院提起诉讼。")

    # =========================================================
    # Article 7
    # =========================================================
    add_section_heading(doc, "第七条　签署与生效")

    add_numbered_item(doc, 1, "生效条件：",
                      "本协议自甲乙双方法定代表人或授权代表签字并加盖公章（或合同专用章）之日起生效。")
    add_numbered_item(doc, 2, "落款与盖章：",
                      "甲方加盖组委会公章或复旦大学住房政策研究中心公章；乙方加盖公章或合同专用章。")
    add_numbered_item(doc, 3, "物理要求：",
                      "本协议正本须加盖骑缝章，一式贰份，甲乙双方各执壹份，具有同等法律效力。")
    add_numbered_item(doc, 4, "附件效力：",
                      "本协议如有附件、补充协议或经双方书面确认的往来函件，"
                      "均为本协议不可分割的组成部分，与本协议具有同等法律效力。")

    # =========================================================
    # Signature page
    # =========================================================
    doc.add_page_break()
    add_section_heading(doc, "签　署　页", size_pt=16)
    add_blank(doc, 12)

    # 甲方
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.8, space_before=8, space_after=8,
                         first_line_indent_chars=0)
    add_runs_with_bold(p, "**甲方（盖章）：**", size_pt=13)

    add_body_paragraph(doc,
        "人工智能商业化落地峰会组委会 / 复旦大学住房政策研究中心 / 上海市杨浦区科技企业联合会",
        indent_chars=0)

    add_blank(doc, 14)
    add_body_paragraph(doc, "授权代表（签字）：______________________", indent_chars=0)
    add_blank(doc, 6)
    add_body_paragraph(doc, "日　　　　　期：　　　　年　　　月　　　日", indent_chars=0)

    add_blank(doc, 18)
    add_hr(doc)
    add_blank(doc, 12)

    # 乙方
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.8, space_before=8, space_after=8,
                         first_line_indent_chars=0)
    add_runs_with_bold(p, "**乙方（盖章）：**", size_pt=13)

    add_body_paragraph(doc, "腾讯云计算（北京）有限责任公司", indent_chars=0)

    add_blank(doc, 14)
    add_body_paragraph(doc, "授权代表（签字）：______________________", indent_chars=0)
    add_blank(doc, 6)
    add_body_paragraph(doc, "日　　　　　期：　　　　年　　　月　　　日", indent_chars=0)

    add_blank(doc, 24)

    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing=1.5, space_before=12, space_after=0,
                         first_line_indent_chars=0)
    run = p.add_run("（以下无正文）")
    set_run_fonts(run, 11)
    run.italic = True

    doc.save(DOC_PATH)
    print(f"OK -> {DOC_PATH}")


if __name__ == "__main__":
    build()
