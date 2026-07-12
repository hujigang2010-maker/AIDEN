#!/usr/bin/env python3
"""Convert the WAIC 北外滩双夜 markdown plan into a Word (.docx) file.

Handles the subset of Markdown used in the source document:
headings (#/##/###), tables, bullet & checkbox lists, blockquotes,
horizontal rules, inline bold, and plain paragraphs.
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "WAIC2026北外滩双夜活动整合方案.md"
OUT = "WAIC2026北外滩双夜活动整合方案.docx"


def add_inline(paragraph, text):
    """Render **bold** segments within a line."""
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "":
            continue
        run = paragraph.add_run(seg)
        if i % 2 == 1:  # captured group = bold
            run.bold = True


def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def split_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def is_separator(cells):
    return all(re.fullmatch(r":?-{2,}:?", c.replace(" ", "")) for c in cells if c)


def main():
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line
        if stripped == "":
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            if level == 1:
                h = doc.add_heading("", level=0)
                add_inline(h, text)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                h = doc.add_heading("", level=min(level - 1, 4))
                add_inline(h, text)
            i += 1
            continue

        # Tables
        if is_table_row(line) and i + 1 < n and is_table_row(lines[i + 1]) and is_separator(split_row(lines[i + 1])):
            header = split_row(line)
            rows = []
            i += 2  # skip header + separator
            while i < n and is_table_row(lines[i]):
                rows.append(split_row(lines[i]))
                i += 1
            ncols = len(header)
            table = doc.add_table(rows=1, cols=ncols)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for c in range(ncols):
                hdr[c].text = ""
                p = hdr[c].paragraphs[0]
                add_inline(p, header[c] if c < len(header) else "")
                for r in p.runs:
                    r.bold = True
            for row in rows:
                cells = table.add_row().cells
                for c in range(ncols):
                    cells[c].text = ""
                    add_inline(cells[c].paragraphs[0], row[c] if c < len(row) else "")
            doc.add_paragraph("")
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run("  ".join(quote_lines))
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # Checkbox list
        m = re.match(r"^-\s+\[[ x]\]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, "☐ " + m.group(1))
            i += 1
            continue

        # Bullet list
        m = re.match(r"^-\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, m.group(1))
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(OUT)
    print(f"Saved {OUT}")


if __name__ == "__main__":
    sys.exit(main())
